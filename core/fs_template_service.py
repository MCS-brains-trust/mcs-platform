"""
Financial Statement Template Service — docxtpl pipeline.

Replaces the old python-docx programmatic approach (core/docgen.py) with
Jinja2-templated .docx files rendered via docxtpl.

Functions:
  format_amount           — Decimal → financial string
  aggregate_tb_lines      — Group/sum TB lines by normalised account name
  build_company_context   — Full Jinja2 context for company entity
  build_trust_context     — Full Jinja2 context for trust entity
  build_sole_trader_context — Full Jinja2 context for sole trader entity
  render_template         — Load + render a .docx template via DocxTemplate
  generate_financial_statements — Orchestrate all templates for a FY
  assemble_pdf_package    — Generate clean PDFs, merge into single package
"""
import io
import logging
import os
import tempfile
from collections import OrderedDict
from decimal import Decimal, ROUND_HALF_UP

from django.conf import settings
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.libreoffice_utils import convert_docx_to_pdf

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# 1. format_amount
# ---------------------------------------------------------------------------
def format_amount(value, show_negative_brackets=True):
    """Format a Decimal to financial string — whole dollars.

    - Zero / None → "—"
    - Negative with brackets → "(1,234)"
    - Positive → "1,234"
    - Whole dollars, comma separators, no $ sign
    """
    if value is None:
        return "—"
    d = Decimal(str(value)).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
    if d == 0:
        return "—"
    if d < 0 and show_negative_brackets:
        return f"({abs(d):,.0f})"
    return f"{d:,.0f}"


# ---------------------------------------------------------------------------
# 2. aggregate_tb_lines
# ---------------------------------------------------------------------------
def aggregate_tb_lines(queryset):
    """Group TrialBalanceLine queryset by normalised account_name, sum amounts.

    Grouping key: account_name.strip().lower()
    Display name: most frequent original casing within group.

    Returns list of dicts:
        [{"account_name": str, "cy_amount": Decimal, "py_amount": Decimal}, ...]
    """
    raw_count = 0
    agg = OrderedDict()      # norm_key → {cy, py, names: {name: count}}
    for line in queryset:
        raw_count += 1
        norm = line.account_name.strip().lower()
        cy = line.closing_balance
        py = line.prior_debit - line.prior_credit
        if norm in agg:
            agg[norm]["cy"] += cy
            agg[norm]["py"] += py
            agg[norm]["names"][line.account_name] = (
                agg[norm]["names"].get(line.account_name, 0) + 1
            )
        else:
            agg[norm] = {
                "cy": cy,
                "py": py,
                "names": {line.account_name: 1},
            }

    result = []
    for norm, data in agg.items():
        display_name = max(data["names"], key=data["names"].get)
        result.append({
            "account_name": display_name,
            "cy_amount": data["cy"],
            "py_amount": data["py"],
        })

    logger.info(
        "aggregate_tb_lines: %d raw lines -> %d aggregated rows",
        raw_count, len(result),
    )
    return result


# ---------------------------------------------------------------------------
# helpers — TB section extraction (mirrors docgen._get_tb_sections logic)
# ---------------------------------------------------------------------------
def _get_tb_sections(fy):
    """Extract trial balance lines grouped into financial statement sections."""
    lines = fy.trial_balance_lines.order_by("account_code").all()
    sections = {
        "trading_income": [],
        "cogs": [],
        "income": [],
        "expenses": [],
        "current_assets": [],
        "noncurrent_assets": [],
        "current_liabilities": [],
        "noncurrent_liabilities": [],
        "equity": [],
    }

    for line in lines:
        try:
            code_num = int(line.account_code.split(".")[0])
        except (ValueError, TypeError):
            continue

        cy = line.closing_balance
        py = line.prior_debit - line.prior_credit
        entry = {
            "account_code": line.account_code,
            "account_name": line.account_name,
            "cy_amount": cy,
            "py_amount": py,
        }

        name_lower = line.account_name.lower()
        is_cogs = any(kw in name_lower for kw in [
            "cost of", "opening stock", "closing stock", "purchases", "stock on hand",
        ])

        if code_num < 1000:
            is_other_income = any(kw in name_lower for kw in [
                "interest", "other", "fbt", "contribution", "dividend", "sundry",
            ])
            is_trading = any(kw in name_lower for kw in [
                "sales", "income", "takings", "revenue", "accommodation",
                "conference", "meals", "bar", "trading",
            ])
            if is_other_income:
                sections["income"].append(entry)
            elif is_trading:
                sections["trading_income"].append(entry)
            else:
                sections["income"].append(entry)
        elif code_num < 1200:
            sections["cogs"].append(entry)
        elif code_num < 2000:
            if is_cogs:
                sections["cogs"].append(entry)
            else:
                sections["expenses"].append(entry)
        elif code_num < 2500:
            sections["current_assets"].append(entry)
        elif code_num < 3000:
            sections["noncurrent_assets"].append(entry)
        elif code_num < 3500:
            sections["current_liabilities"].append(entry)
        elif code_num < 4000:
            sections["noncurrent_liabilities"].append(entry)
        elif code_num < 5000:
            # Income tax accounts (4100-4149) ARE equity appropriation items.
            # They must be included — they appear as "Less: Income tax on
            # profit" and reduce Total Equity so it balances to Net Assets.
            # Do NOT exclude them.
            sections["equity"].append(entry)
        elif code_num < 6000:
            sections["cogs"].append(entry)
        elif 9000 <= code_num < 10000:
            # Capital accounts (9000 series) — trust unit holder / beneficiary
            # capital provisioned by capital_account_service.  These are equity.
            sections["equity"].append(entry)

    # Aggregate lines with the same account within each section.
    # Primary merge key: account_code (stable across renames in Xero/QBO).
    # Fallback for blank codes: case-insensitive, whitespace-normalised name.
    # Display name preference: names from lines with non-zero CY data are
    # weighted higher so that renamed accounts show the current-year name.
    for key in sections:
        raw = sections[key]
        if not raw:
            continue
        agg = OrderedDict()
        name_counts = {}
        for entry in raw:
            code = entry.get("account_code", "").strip()
            merge_key = code if code else entry["account_name"].strip().lower()
            weight = 10 if entry["cy_amount"] != 0 else 1
            if merge_key in agg:
                agg[merge_key]["cy_amount"] += entry["cy_amount"]
                agg[merge_key]["py_amount"] += entry["py_amount"]
                name_counts[merge_key][entry["account_name"]] = (
                    name_counts[merge_key].get(entry["account_name"], 0) + weight
                )
            else:
                agg[merge_key] = {
                    "account_name": entry["account_name"],
                    "cy_amount": entry["cy_amount"],
                    "py_amount": entry["py_amount"],
                }
                name_counts[merge_key] = {entry["account_name"]: weight}
        for mk in agg:
            best = max(name_counts[mk], key=name_counts[mk].get)
            agg[mk]["account_name"] = best
        sections[key] = list(agg.values())
        logger.debug(
            "_get_tb_sections [%s]: %d raw -> %d aggregated",
            key, len(raw), len(agg),
        )

    return sections


def _sum_section(items, field="cy_amount"):
    """Sum a list of dicts by field."""
    return sum(item[field] for item in items)


# Placeholder for ampersand to survive docxtpl XML rendering.
# docxtpl's Jinja2→XML pipeline strips bare "&" from values.
# We replace "&" with this placeholder before template rendering,
# then restore it in _post_process_fs_doc.
_AMP_PLACEHOLDER = "\u00a7AMP\u00a7"  # §AMP§ — won't appear in real data


def _safe_amp(text):
    """Replace '&' with a placeholder that survives docxtpl rendering."""
    if text and "&" in str(text):
        return str(text).replace("&", _AMP_PLACEHOLDER)
    return text


def _get_firm_context():
    """Return firm branding fields from FirmSettings for document context.

    Loads the singleton FirmSettings record and returns a dict of firm_*
    keys ready to be unpacked into any document context.  Falls back to
    MC & S defaults if the record has not yet been configured so that
    existing documents are never broken during the migration period.
    """
    try:
        from core.models import FirmSettings
        fs = FirmSettings.get()
        _name = _safe_amp(fs.firm_name or "MC & S Pty Ltd")
        _addr = " ".join(filter(None, [fs.firm_address_1, fs.firm_address_2]))
        return {
            # Legacy keys (kept for backward compat with existing templates)
            "firm_name": _name,
            "firm_address_1": fs.firm_address_1 or "PO Box 4440",
            "firm_address_2": fs.firm_address_2 or "Dandenong South VIC 3164",
            "firm_phone": fs.firm_phone or "(03) 9794 0000",
            "firm_email": fs.firm_email or "info@mcands.com.au",
            "firm_website": fs.firm_website or "",
            "firm_logo_path": fs.logo_path,
            "firm_logo_url": fs.logo_url or "",
            "firm_abn": fs.firm_abn or "",
            # New spec keys (practice_* namespace)
            "practice_name": _name,
            "practice_legal_name": _safe_amp(fs.firm_legal_name or fs.firm_name or "MC & S Pty Ltd"),
            "practice_abn": fs.firm_abn or "",
            "practice_registered_address": _addr,
            "practice_phone": fs.firm_phone or "",
            "practice_email": fs.firm_email or "",
            "practice_website": fs.firm_website or "",
            "practice_logo_url": fs.logo_url or "",
            "practice_logo_path": fs.logo_path,
            "practice_tax_agent_number": fs.tax_agent_number or "",
            "practice_bas_agent_number": fs.bas_agent_number or "",
            "practice_asic_agent_number": fs.asic_agent_number or "",
            "practice_signatory_name": fs.signatory_name or "",
            "practice_signatory_designation": fs.signatory_designation or "",
            "practice_professional_body": fs.professional_body or "CPA Australia",
            "practice_membership_number": fs.membership_number or "",
            "practice_independence_maintained": fs.practice_independence_maintained,
            "practice_compilation_report_name": fs.compilation_report_name or fs.firm_name or "",
            "practice_legal_disclaimer": fs.document_disclaimer or "",
        }
    except Exception:
        logger.warning("FirmSettings unavailable — using MC & S defaults", exc_info=True)
        return {
            "firm_name": _safe_amp("MC & S Pty Ltd"),
            "firm_address_1": "PO Box 4440",
            "firm_address_2": "Dandenong South VIC 3164",
            "firm_phone": "(03) 9794 0000",
            "firm_email": "info@mcands.com.au",
            "firm_website": "",
            "firm_logo_path": None,
            "firm_logo_url": "",
            "firm_abn": "",
            "practice_name": _safe_amp("MC & S Pty Ltd"),
            "practice_legal_name": _safe_amp("MC & S Pty Ltd"),
            "practice_abn": "",
            "practice_registered_address": "",
            "practice_phone": "",
            "practice_email": "",
            "practice_website": "",
            "practice_logo_url": "",
            "practice_logo_path": None,
            "practice_tax_agent_number": "",
            "practice_bas_agent_number": "",
            "practice_asic_agent_number": "",
            "practice_signatory_name": "",
            "practice_signatory_designation": "",
            "practice_professional_body": "CPA Australia",
            "practice_membership_number": "",
            "practice_independence_maintained": True,
            "practice_compilation_report_name": "",
            "practice_legal_disclaimer": "",
        }


def _format_lines(items, credit_normal=False):
    """Add formatted amount strings to each item dict.

    For credit-normal accounts (liabilities, equity, income/revenue),
    the raw TB amount is negative (debit - credit).  Setting
    ``credit_normal=True`` negates the value before formatting so these
    accounts display as positive numbers in the financial statements.
    """
    for item in items:
        cy = item["cy_amount"]
        py = item["py_amount"]
        if credit_normal:
            cy = -cy if cy else cy
            py = -py if py else py
        item["cy_formatted"] = format_amount(cy)
        item["py_formatted"] = format_amount(py)
        # Protect ampersand in account names from XML stripping
        item["account_name"] = _safe_amp(item.get("account_name", ""))
    return items


def _classify_current_asset(name_lower):
    """Classify a current asset into a sub-group by keyword matching."""
    if any(kw in name_lower for kw in ["cash", "bank", "petty cash", "on hand"]):
        return "Cash Assets"
    if any(kw in name_lower for kw in ["debtor", "receivable", "trade debtor"]):
        return "Receivables"
    return "Other Current Assets"


def _classify_current_liability(name_lower):
    """Classify a current liability into a sub-group by keyword matching.

    Tax Liabilities are checked FIRST because accounts like "GST payable"
    contain "payable" which would otherwise match the Payables group.
    """
    # Tax / ATO statutory obligations — check first (higher priority)
    if any(kw in name_lower for kw in [
        "gst", "payg", "tax", "taxation", "withholding", "bas", "ato", "clearing",
    ]):
        return "Tax Liabilities"
    # Trade payables / creditors
    if any(kw in name_lower for kw in [
        "creditor", "payable", "accrual", "accounts payable", "trade creditor",
        "sundry creditor",
    ]):
        return "Payables"
    return "Other Current Liabilities"


def _build_subgrouped_items(items, classify_fn, credit_normal=False):
    """Group items into sub-categories and add formatted amounts.

    Returns a list of dicts suitable for the template. Each entry is either:
      - A sub-heading row: {"is_heading": True, "account_name": "Cash Assets"}
      - A line item row:   {"account_name": ..., "cy_formatted": ..., "py_formatted": ...}
      - A subtotal row:    {"is_subtotal": True, "cy_formatted": ..., "py_formatted": ...}
    """
    from collections import OrderedDict

    groups = OrderedDict()
    for item in items:
        group = classify_fn(item["account_name"].lower())
        groups.setdefault(group, []).append(item)

    # If there's only one group, return a flat list (no sub-headings needed)
    if len(groups) <= 1:
        return _format_lines(list(items), credit_normal=credit_normal)

    result = []
    for group_name, group_items in groups.items():
        # Sub-heading row
        result.append({
            "is_heading": True,
            "account_name": group_name,
            "cy_formatted": "",
            "py_formatted": "",
            "note_ref": "",
        })
        # Line items
        formatted = _format_lines(list(group_items), credit_normal=credit_normal)
        result.extend(formatted)
        # Subtotal row
        sub_cy = sum(item["cy_amount"] for item in group_items)
        sub_py = sum(item["py_amount"] for item in group_items)
        if credit_normal:
            sub_cy = -sub_cy if sub_cy else sub_cy
            sub_py = -sub_py if sub_py else sub_py
        result.append({
            "is_subtotal": True,
            "account_name": "",
            "cy_formatted": format_amount(sub_cy),
            "py_formatted": format_amount(sub_py),
            "note_ref": "",
        })

    return result


def _has_prior_year(fy):
    """Check if there is prior year data."""
    if not fy.prior_year:
        return False
    return fy.prior_year.trial_balance_lines.exists()


def _format_acn_abn(acn, abn):
    """Build a combined 'ACN: xxx / ABN: xxx' display string."""
    parts = []
    if acn:
        d = "".join(c for c in str(acn) if c.isdigit())
        if len(d) == 9:
            parts.append(f"ACN: {d[:3]} {d[3:6]} {d[6:9]}")
        elif d:
            parts.append(f"ACN: {d}")
    if abn:
        d = "".join(c for c in str(abn) if c.isdigit())
        if len(d) == 11:
            parts.append(f"ABN: {d[:2]} {d[2:5]} {d[5:8]} {d[8:11]}")
        elif d:
            parts.append(f"ABN: {d}")
    return " / ".join(parts) if parts else ""


# ---------------------------------------------------------------------------
# 2b. Note map computation — shared between context builder and notes generator
# ---------------------------------------------------------------------------
def _compute_note_map(sections, entity_type, has_income_tax):
    """Compute the sequential note numbering map from TB section data.

    Returns (note_map, note_lookup) where:
      note_map   = [(1, 'policies'), (2, 'receivables'), ...]
      note_lookup = {'receivables': 2, 'ppe': 3, ...}
    """
    # Trade Receivables trigger
    has_trade_debtors = any(
        ("trade" in i["account_name"].lower() and "debtor" in i["account_name"].lower())
        and (i["cy_amount"] != 0 or i["py_amount"] != 0)
        for i in sections["current_assets"]
    )

    # PPE trigger — any non-current asset that is PPE at cost with non-zero balance
    has_ppe = False
    for item in sections["noncurrent_assets"]:
        nl = item["account_name"].lower()
        is_depr = any(kw in nl for kw in ["accumulated", "amortisation", "depreciation"]) or nl.startswith("less:")
        is_deposit = "deposit" in nl
        is_cost = any(kw in nl for kw in [
            "equipment", "vehicle", "furniture", "building", "fixture",
            "plant", "motor", "computer", "office", "at cost",
        ]) and not is_depr and not is_deposit
        if is_cost and (item["cy_amount"] != 0 or item["py_amount"] != 0):
            has_ppe = True
            break

    # Related party triggers
    has_mgmt_fees = any(
        "management" in i["account_name"].lower() and "fee" in i["account_name"].lower()
        and (i["cy_amount"] != 0 or i["py_amount"] != 0)
        for i in sections["expenses"]
    )
    has_director_loan = (
        entity_type != "sole_trader" and
        any(
            "loan" in i["account_name"].lower() and "director" in i["account_name"].lower()
            and (i["cy_amount"] != 0 or i["py_amount"] != 0)
            for i in sections["noncurrent_liabilities"]
        )
    )
    has_related_loans = any(
        "loan" in i["account_name"].lower()
        and any(kw in i["account_name"].lower() for kw in ["majoti", "ets", "related"])
        and (i["cy_amount"] != 0 or i["py_amount"] != 0)
        for i in sections["noncurrent_liabilities"]
    )
    has_related_party = has_mgmt_fees or has_director_loan or has_related_loans

    is_company = entity_type == "company"

    note_map = []
    n = 1
    note_map.append((n, "policies")); n += 1
    if has_trade_debtors:
        note_map.append((n, "receivables")); n += 1
    if has_ppe:
        note_map.append((n, "ppe")); n += 1
    if has_related_party:
        note_map.append((n, "related_party")); n += 1
    if has_income_tax:
        note_map.append((n, "income_tax")); n += 1
    if is_company:
        note_map.append((n, "events")); n += 1

    note_lookup = {note_type: note_num for note_num, note_type in note_map}
    return note_map, note_lookup


def _assign_note_refs(items, note_lookup, classify_fn):
    """Add a 'note_ref' key to each item dict based on a classification function.

    classify_fn(account_name_lower) -> note_type string or None.
    """
    for item in items:
        nl = item["account_name"].lower()
        note_type = classify_fn(nl)
        if note_type and note_type in note_lookup:
            item["note_ref"] = str(note_lookup[note_type])
        else:
            item["note_ref"] = ""
    return items


# ---------------------------------------------------------------------------
# 3. build_company_context
# ---------------------------------------------------------------------------
def build_company_context(financial_year, include_watermark=True):
    """Build full Jinja2 context dict for a company entity."""
    from core.models import EntityOfficer

    fy = financial_year
    entity = fy.entity
    sections = _get_tb_sections(fy)
    has_prior = _has_prior_year(fy)
    has_trading = len(sections["trading_income"]) > 0 or len(sections["cogs"]) > 0

    # P&L calculations
    # Trading income, other income are credit-normal (raw TB value is negative)
    trading_income = _format_lines(sections["trading_income"], credit_normal=True)
    cogs = _format_lines(sections["cogs"])
    income = _format_lines(sections["income"], credit_normal=True)
    expenses = _format_lines(sections["expenses"])

    total_trading_income_cy = -_sum_section(sections["trading_income"])
    total_trading_income_py = -_sum_section(sections["trading_income"], "py_amount")
    total_cogs_cy = _sum_section(sections["cogs"])
    total_cogs_py = _sum_section(sections["cogs"], "py_amount")
    gross_profit_cy = total_trading_income_cy - total_cogs_cy
    gross_profit_py = total_trading_income_py - total_cogs_py

    total_income_cy = -_sum_section(sections["income"])
    total_income_py = -_sum_section(sections["income"], "py_amount")
    total_expenses_cy = _sum_section(sections["expenses"])
    total_expenses_py = _sum_section(sections["expenses"], "py_amount")

    if has_trading:
        net_profit_pretax_cy = gross_profit_cy + total_income_cy - total_expenses_cy
        net_profit_pretax_py = gross_profit_py + total_income_py - total_expenses_py
    else:
        net_profit_pretax_cy = total_income_cy - total_expenses_cy
        net_profit_pretax_py = total_income_py - total_expenses_py

    # Extract income tax from equity section — companies only.
    # Trusts, partnerships, and sole traders do not pay income tax at entity
    # level, so the reclassification must be skipped for those entity types.
    income_tax_cy = Decimal("0")
    income_tax_py = Decimal("0")
    has_income_tax = False

    entity_type = entity.entity_type
    if entity_type == "company":
        equity_without_tax = []
        for item in sections["equity"]:
            code_str = item.get("account_code", "")
            name_lower = item.get("account_name", "").lower()
            try:
                code_num = int(code_str.split(".")[0]) if code_str else 0
            except (ValueError, TypeError):
                code_num = 0
            is_tax = (4100 <= code_num <= 4149) or any(
                kw in name_lower for kw in ["income tax", "tax on profit", "tax expense"]
            )
            if is_tax:
                income_tax_cy += item["cy_amount"] if item["cy_amount"] else Decimal("0")
                income_tax_py += item["py_amount"] if item["py_amount"] else Decimal("0")
            else:
                equity_without_tax.append(item)
        sections["equity"] = equity_without_tax
        has_income_tax = income_tax_cy != 0 or income_tax_py != 0

    # After-tax profit (for non-companies, income_tax is zero so this is same as pretax)
    net_profit_cy = net_profit_pretax_cy - income_tax_cy
    net_profit_py = net_profit_pretax_py - income_tax_py

    # Retained profit roll-forward
    # ─────────────────────────────────────────────────────────────────────────
    # StatementHub uses an UNCLOSED trial balance convention: current-year
    # profit/loss remains in the P&L accounts and has NOT been transferred
    # into the retained earnings account.  This means:
    #
    #   cy_amount of retained/accumulated accounts
    #       = the OPENING retained balance for the current year
    #       = the CLOSING retained balance of the prior year
    #
    #   py_amount of retained/accumulated accounts
    #       = the OPENING retained balance for the prior year
    #       = the CLOSING retained balance of the year before that
    #
    # Roll-forward identity (unclosed TB):
    #   retained_profit_opening_cy  = cy_amount (negated, credit-normal)
    #   retained_profit_closing_cy  = opening + after-tax profit − dividends
    #
    #   retained_profit_opening_py  = py_amount (negated, credit-normal)
    #   retained_profit_closing_py  = py_opening + net_profit_py − dividends_py
    #
    # Credit-normal convention: equity accounts carry negative raw values for
    # credit (positive) balances, so we negate to get display amounts.
    # ─────────────────────────────────────────────────────────────────────────
    _retained_opening_raw_cy = Decimal("0")   # cy_amount = opening for current year
    _retained_opening_raw_py = Decimal("0")   # py_amount = opening for prior year
    _dividends_cy = Decimal("0")
    _dividends_py = Decimal("0")
    for _item in sections["equity"]:
        _name_l = _item.get("account_name", "").lower()
        if any(kw in _name_l for kw in ["retained", "accumulated"]):
            _retained_opening_raw_cy += _item.get("cy_amount", Decimal("0")) or Decimal("0")
            _retained_opening_raw_py += _item.get("py_amount", Decimal("0")) or Decimal("0")
        elif any(kw in _name_l for kw in ["dividend", "distribution"]):
            _dividends_cy += abs(_item.get("cy_amount", Decimal("0")) or Decimal("0"))
            _dividends_py += abs(_item.get("py_amount", Decimal("0")) or Decimal("0"))

    # Convert credit-normal raw to positive display amounts
    # For unclosed TB: cy_amount IS the opening balance (profit not yet transferred)
    retained_profit_opening_cy = -_retained_opening_raw_cy
    retained_profit_opening_py = -_retained_opening_raw_py

    # Total available = opening + after-tax profit
    total_available_cy = retained_profit_opening_cy + net_profit_cy
    total_available_py = retained_profit_opening_py + net_profit_py

    # Closing retained = total available minus dividends
    retained_profit_closing_cy = total_available_cy - _dividends_cy
    retained_profit_closing_py = total_available_py - _dividends_py

    # Pre-closing TB: add current year profit (after tax) to equity if BS won't balance
    _test_equity = -_sum_section(sections["equity"])
    _test_liab = -(_sum_section(sections["current_liabilities"])
                    + _sum_section(sections["noncurrent_liabilities"]))
    _test_assets = (_sum_section(sections["current_assets"])
                    + _sum_section(sections["noncurrent_assets"]))
    _test_net_assets = _test_assets - _test_liab
    if abs(_test_net_assets - _test_equity) > 1:
        # TB not yet closed — inject current year profit / (loss) (after tax)
        sections["equity"].append({
            "account_name": "Current year profit / (loss)",
            "cy_amount": -net_profit_cy,   # credit-normal convention
            "py_amount": -net_profit_py,
        })

    # Post-injection balance sheet integrity check.
    # Tolerance of $1 matches the pre-injection threshold above — sub-dollar
    # differences are expected rounding artefacts from the net profit
    # calculation (income and expense cents that don't perfectly cancel).
    _final_equity = -_sum_section(sections["equity"])
    _final_net_assets = _test_net_assets  # unchanged by equity injection
    if abs(_final_net_assets - _final_equity) > Decimal("1"):
        logger.warning(
            "Balance Sheet integrity failure for %s FY %s: "
            "Net Assets=%s, Total Equity=%s, diff=%s",
            entity.entity_name, fy.pk,
            _final_net_assets, _final_equity,
            _final_net_assets - _final_equity,
        )
        # Surface as a CRITICAL Eva finding
        try:
            from core.models import EvaReview, EvaFinding
            _active_review = EvaReview.objects.filter(
                financial_year=fy,
            ).order_by("-created_at").first()
            if _active_review:
                import hashlib, json as _json
                _fp = hashlib.sha256(_json.dumps({
                    "entity_id": str(entity.pk),
                    "financial_year_id": str(fy.pk),
                    "rule_category": "balance_sheet_integrity",
                }, sort_keys=True).encode()).hexdigest()
                EvaFinding.objects.update_or_create(
                    eva_review=_active_review,
                    check_name="balance_sheet_integrity",
                    defaults={
                        "severity": "critical",
                        "title": "Balance Sheet does not reconcile",
                        "plain_english_explanation": (
                            f"Net Assets ({format_amount(_final_net_assets)}) does not equal "
                            f"Total Equity ({format_amount(_final_equity)}). "
                            f"Difference: {format_amount(_final_net_assets - _final_equity)}."
                        ),
                        "recommendation": (
                            "Investigate equity section for missing components such as "
                            "unitholders' capital, retained earnings, or current year profit. "
                            "The financial statements cannot be finalised until this is resolved."
                        ),
                        "source": "risk_engine",
                        "fingerprint": _fp,
                    },
                )
        except Exception as _e:
            logger.error("Could not create balance_sheet_integrity finding: %s", _e)

    # Compute note_map and assign note_ref to items for the Note column
    note_map, note_lookup = _compute_note_map(sections, entity_type, has_income_tax)

    def _classify_note(nl):
        """Return note_type for an account name, or None."""
        if "trade" in nl and "debtor" in nl:
            return "receivables"
        if any(kw in nl for kw in [
            "equipment", "vehicle", "furniture", "building", "fixture",
            "plant", "motor", "computer", "office", "at cost",
        ]):
            # Exclude depreciation/amortisation lines — note ref goes on cost only
            if not any(kw in nl for kw in ["accumulated", "amortisation", "depreciation"]) and not nl.startswith("less:"):
                return "ppe"
        if "deposit" in nl:
            return "ppe"
        if "loan" in nl and ("director" in nl or "majoti" in nl or "ets" in nl or "related" in nl):
            return "related_party"
        if "management" in nl and "fee" in nl:
            return "related_party"
        return None

    _assign_note_refs(sections["current_assets"], note_lookup, _classify_note)
    _assign_note_refs(sections["noncurrent_assets"], note_lookup, _classify_note)
    _assign_note_refs(sections["noncurrent_liabilities"], note_lookup, _classify_note)
    _assign_note_refs(sections["expenses"], note_lookup, _classify_note)
    # Sections that don't have note refs — add empty note_ref
    for sec_key in ["trading_income", "cogs", "income", "current_liabilities", "equity"]:
        for item in sections[sec_key]:
            if "note_ref" not in item:
                item["note_ref"] = ""

    # For P&L rendering: merge trading income & COGS into income & expenses
    # since the P&L template has a single Income and Expenses section
    if has_trading:
        rendered_income = trading_income + income
        rendered_total_income_cy = total_trading_income_cy + total_income_cy
        rendered_total_income_py = total_trading_income_py + total_income_py
        rendered_expenses = cogs + expenses
        rendered_total_expenses_cy = total_cogs_cy + total_expenses_cy
        rendered_total_expenses_py = total_cogs_py + total_expenses_py
    else:
        rendered_income = income
        rendered_total_income_cy = total_income_cy
        rendered_total_income_py = total_income_py
        rendered_expenses = expenses
        rendered_total_expenses_cy = total_expenses_cy
        rendered_total_expenses_py = total_expenses_py

    # Suppress zero line items — if both CY and PY are zero, exclude the row.
    # Keep sub-headings and subtotals (they have is_heading/is_subtotal flags).
    def _is_zero_line(item):
        if item.get("is_heading") or item.get("is_subtotal"):
            return False
        cy = item.get("cy_amount", 0) or 0
        py = item.get("py_amount", 0) or 0
        return cy == 0 and py == 0
    rendered_income = [i for i in rendered_income if not _is_zero_line(i)]
    rendered_expenses = [i for i in rendered_expenses if not _is_zero_line(i)]

    # Balance Sheet — sub-grouped current assets and current liabilities
    current_assets = _build_subgrouped_items(
        sections["current_assets"], _classify_current_asset)
    noncurrent_assets = _format_lines(sections["noncurrent_assets"])
    # Liabilities and equity are credit-normal (raw TB value is negative)
    current_liabilities = _build_subgrouped_items(
        sections["current_liabilities"], _classify_current_liability, credit_normal=True)
    noncurrent_liabilities = _format_lines(sections["noncurrent_liabilities"], credit_normal=True)
    equity = _format_lines(sections["equity"], credit_normal=True)

    total_current_assets_cy = _sum_section(sections["current_assets"])
    total_current_assets_py = _sum_section(sections["current_assets"], "py_amount")
    total_noncurrent_assets_cy = _sum_section(sections["noncurrent_assets"])
    total_noncurrent_assets_py = _sum_section(sections["noncurrent_assets"], "py_amount")
    total_assets_cy = total_current_assets_cy + total_noncurrent_assets_cy
    total_assets_py = total_current_assets_py + total_noncurrent_assets_py

    total_current_liab_cy = -_sum_section(sections["current_liabilities"])
    total_current_liab_py = -_sum_section(sections["current_liabilities"], "py_amount")
    total_noncurrent_liab_cy = -_sum_section(sections["noncurrent_liabilities"])
    total_noncurrent_liab_py = -_sum_section(sections["noncurrent_liabilities"], "py_amount")
    total_liab_cy = total_current_liab_cy + total_noncurrent_liab_cy
    total_liab_py = total_current_liab_py + total_noncurrent_liab_py

    # Zero section suppression flags for Balance Sheet
    has_noncurrent_assets = total_noncurrent_assets_cy != 0 or total_noncurrent_assets_py != 0
    has_noncurrent_liabilities = total_noncurrent_liab_cy != 0 or total_noncurrent_liab_py != 0

    net_assets_cy = total_assets_cy - total_liab_cy
    net_assets_py = total_assets_py - total_liab_py

    total_equity_cy = -_sum_section(sections["equity"])
    total_equity_py = -_sum_section(sections["equity"], "py_amount")

    # Officers
    directors = EntityOfficer.objects.filter(
        entity=entity,
        role__in=["director"],
        date_ceased__isnull=True,
    ).order_by("display_order", "full_name")

    year_end = fy.end_date
    year_str = str(year_end.year) if year_end else ""
    prior_year_str = str(year_end.year - 1) if year_end else ""
    date_text = f"For the Year Ended {year_end.strftime('%d %B %Y')}" if year_end else ""

    # Signing / declaration date — use finalised_at if available, else today
    from datetime import date as _date
    if fy.finalised_at:
        signing_date = fy.finalised_at.date().strftime("%-d %B %Y")
    else:
        signing_date = _date.today().strftime("%-d %B %Y")

    # Format ACN/ABN with proper spacing for display
    abn_raw = entity.abn or ""
    acn_raw = entity.acn or ""
    abn_digits = "".join(c for c in str(abn_raw) if c.isdigit())
    acn_digits = "".join(c for c in str(acn_raw) if c.isdigit())
    abn_formatted = (
        f"{abn_digits[:2]} {abn_digits[2:5]} {abn_digits[5:8]} {abn_digits[8:]}"
        if len(abn_digits) == 11 else abn_raw
    )
    acn_formatted = (
        f"{acn_digits[:3]} {acn_digits[3:6]} {acn_digits[6:]}"
        if len(acn_digits) == 9 else acn_raw
    )

    context = {
        "entity_name": _safe_amp(entity.entity_name),
        "trading_as": _safe_amp(entity.trading_as or ""),
        "abn": abn_formatted,
        "acn": acn_formatted,
        "acn_abn": _safe_amp(_format_acn_abn(acn_raw, abn_raw)),
        "entity_type": entity.entity_type,
        "year": year_str,
        "financial_year": year_str,
        "prior_year": prior_year_str,
        "date_text": date_text,
        "financial_year_end": year_end.strftime("%d %B %Y") if year_end else "",
        "year_end_date": year_end.strftime("%d %B %Y") if year_end else "",
        "has_prior": has_prior,
        "has_trading": has_trading,
        "watermark": "DRAFT" if include_watermark else "",
        # P&L
        "trading_income": trading_income,
        "cogs": cogs,
        "income": rendered_income,
        "expenses": rendered_expenses,
        "total_trading_income_cy": format_amount(total_trading_income_cy),
        "total_trading_income_py": format_amount(total_trading_income_py),
        "total_cogs_cy": format_amount(total_cogs_cy),
        "total_cogs_py": format_amount(total_cogs_py),
        "gross_profit_cy": format_amount(gross_profit_cy),
        "gross_profit_py": format_amount(gross_profit_py),
        "total_income_cy": format_amount(rendered_total_income_cy),
        "total_income_py": format_amount(rendered_total_income_py),
        "total_expenses_cy": format_amount(rendered_total_expenses_cy),
        "total_expenses_py": format_amount(rendered_total_expenses_py),
        "net_profit_pretax_cy": format_amount(net_profit_pretax_cy),
        "net_profit_pretax_py": format_amount(net_profit_pretax_py),
        "has_income_tax": has_income_tax,
        "income_tax_cy": format_amount(-income_tax_cy) if income_tax_cy else "-",
        "income_tax_py": format_amount(-income_tax_py) if income_tax_py else "-",
        "net_profit_cy": format_amount(net_profit_cy),
        "net_profit_py": format_amount(net_profit_py),
        # Summary P&L — retained profit appropriation
        "retained_profit_opening_cy": format_amount(retained_profit_opening_cy),
        "retained_profit_opening_py": format_amount(retained_profit_opening_py),
        "total_available_cy": format_amount(total_available_cy),
        "total_available_py": format_amount(total_available_py),
        "retained_profit_closing_cy": format_amount(retained_profit_closing_cy),
        "retained_profit_closing_py": format_amount(retained_profit_closing_py),
        # Balance Sheet
        "current_assets": current_assets,
        "noncurrent_assets": noncurrent_assets,
        "has_noncurrent_assets": has_noncurrent_assets,
        "current_liabilities": current_liabilities,
        "noncurrent_liabilities": noncurrent_liabilities,
        "has_noncurrent_liabilities": has_noncurrent_liabilities,
        "equity": equity,
        "total_current_assets_cy": format_amount(total_current_assets_cy),
        "total_current_assets_py": format_amount(total_current_assets_py),
        "total_noncurrent_assets_cy": format_amount(total_noncurrent_assets_cy),
        "total_noncurrent_assets_py": format_amount(total_noncurrent_assets_py),
        "total_assets_cy": format_amount(total_assets_cy),
        "total_assets_py": format_amount(total_assets_py),
        "total_current_liab_cy": format_amount(total_current_liab_cy),
        "total_current_liab_py": format_amount(total_current_liab_py),
        "total_noncurrent_liab_cy": format_amount(total_noncurrent_liab_cy),
        "total_noncurrent_liab_py": format_amount(total_noncurrent_liab_py),
        "total_liabilities_cy": format_amount(total_liab_cy),
        "total_liabilities_py": format_amount(total_liab_py),
        "net_assets_cy": format_amount(net_assets_cy),
        "net_assets_py": format_amount(net_assets_py),
        "total_equity_cy": format_amount(total_equity_cy),
        "total_equity_py": format_amount(total_equity_py),
        # Declaration
        "declaration_title": "Directors' Declaration",
        "compilation_responsible_party": "directors",
        "directors": [
            {"name": d.full_name, "title": d.title or "Director"}
            for d in directors
        ],
        # Signing / declaration date (for Compilation Report "Dated:" line)
        "signing_date": signing_date,
        # Firm details — loaded from FirmSettings singleton (white-label support)
        **_get_firm_context(),
    }

    # Add format_amount as a Jinja2 filter
    context["format_amount"] = format_amount

    # Note map for notes generator and statement tables
    context["_note_map"] = note_map
    context["_note_lookup"] = note_lookup

    # Raw data for programmatic notes generation (not used by Jinja2 templates)
    context["_sections"] = sections
    context["_entity"] = entity
    context["_fy"] = fy
    context["_income_tax_cy"] = income_tax_cy
    context["_income_tax_py"] = income_tax_py
    context["_has_income_tax"] = has_income_tax
    context["_total_revenue_cy"] = (
        total_trading_income_cy + total_income_cy if has_trading
        else total_income_cy
    )
    context["_total_revenue_py"] = (
        total_trading_income_py + total_income_py if has_trading
        else total_income_py
    )

    return context


# ---------------------------------------------------------------------------
# 4. build_trust_context
# ---------------------------------------------------------------------------
def build_trust_context(financial_year, include_watermark=True):
    """Build full Jinja2 context dict for a trust entity."""
    from core.models import EntityOfficer

    # Start with company context as base (shared structure)
    context = build_company_context(financial_year, include_watermark)

    entity = financial_year.entity

    # Override trust-specific fields
    context["declaration_title"] = "Trustee's Declaration"
    context["compilation_responsible_party"] = "director of the trustee company"

    # Relabel unit holder capital/loan accounts in the equity section.
    # Accounts containing "loan" or "fund" keywords that also contain a
    # beneficiary/unit holder name are relabelled to
    # "Funds loaned to trust — [Unit Holder Name]".
    beneficiary_names_for_match = []
    _all_officers = EntityOfficer.objects.filter(
        entity=entity,
        role__in=["beneficiary", "unit_holder"],
        date_ceased__isnull=True,
    )
    for _off in _all_officers:
        if _off.full_name:
            beneficiary_names_for_match.append(_off.full_name)

    _loan_keywords = ["loan", "fund", "capital introduced", "funds introduced"]
    equity_items = context.get("equity", [])
    for item in equity_items:
        acct_lower = item.get("account_name", "").lower()
        if any(kw in acct_lower for kw in _loan_keywords):
            # Try to match a beneficiary name in the account name
            matched_name = None
            for bn in beneficiary_names_for_match:
                if bn.lower() in acct_lower:
                    matched_name = bn
                    break
            if matched_name:
                item["account_name"] = f"Funds loaned to trust — {matched_name}"
            else:
                # Generic relabel if no name match found
                item["account_name"] = "Funds loaned to trust"

    # Get trustees and beneficiaries — check both legacy `role` and `roles` JSONField
    from django.db import models as django_models
    trustees = EntityOfficer.objects.filter(
        entity=entity, date_ceased__isnull=True,
    ).filter(
        django_models.Q(role="trustee") | django_models.Q(roles__contains="trustee")
    ).order_by("display_order", "full_name")

    beneficiaries = EntityOfficer.objects.filter(
        entity=entity, date_ceased__isnull=True,
    ).filter(
        django_models.Q(role__in=["beneficiary", "unit_holder"])
        | django_models.Q(roles__contains="beneficiary")
        | django_models.Q(roles__contains="unit_holder")
    ).order_by("display_order", "full_name")

    context["directors"] = [
        {"name": t.full_name, "title": t.title or "Trustee"}
        for t in trustees
    ]

    # Build declaration_signatories — individual directors of the corporate trustee
    from django.db import models as _m
    _trustee_officer = EntityOfficer.objects.filter(
        entity=entity, date_ceased__isnull=True,
    ).filter(
        _m.Q(role="trustee") | _m.Q(roles__contains="trustee")
    ).first()
    _trustee_company = _trustee_officer.full_name if _trustee_officer else (
        entity.trustee_name or ""
    )
    _signatories = EntityOfficer.objects.filter(
        entity=entity, is_signatory=True, date_ceased__isnull=True,
    ).order_by("display_order", "full_name")
    context["declaration_signatories"] = [
        {
            "name": o.full_name,
            "trustee_company": _trustee_company,
            "trust_name": entity.entity_name,
            "is_chairperson": o.is_chairperson,
        }
        for o in _signatories
    ]

    # Distribution data — read from selected TaxPlanningScenario (Stage 2)
    from core.models import TrustWorkspace
    net_profit_raw = Decimal("0")
    sections = _get_tb_sections(financial_year)
    total_income = -_sum_section(sections["trading_income"]) + -_sum_section(sections["income"])
    total_expenses = _sum_section(sections["expenses"]) + _sum_section(sections["cogs"])
    net_profit_raw = total_income - total_expenses

    distributions = []
    _missing_names = []
    workspace = getattr(financial_year, "trust_workspace", None)
    scenario = workspace.selected_tax_scenario if workspace else None

    if scenario and scenario.distributions:
        # Use selected Tax Planning scenario
        scenario_total = sum(
            Decimal(str(e.get("proposed_distribution", 0)))
            for e in scenario.distributions
            if Decimal(str(e.get("proposed_distribution", 0))) > 0
        )
        for entry in scenario.distributions:
            amount = Decimal(str(entry.get("proposed_distribution", 0)))
            if amount <= 0:
                continue
            officer = EntityOfficer.objects.filter(pk=entry.get("beneficiary_id")).first()
            if officer:
                name = (officer.full_name or "").strip()
                if not name:
                    _missing_names.append(officer.pk)
            else:
                name = ""
            pct = (amount / scenario_total * 100) if scenario_total else Decimal("0")
            pct_display = f"{pct:.2f}"
            amount_display = f"{amount:,.0f}"
            distributions.append({
                "beneficiary_name": name or "— Name missing —",
                "percentage": pct_display,
                "amount": amount_display,
                "amount_raw": amount,
            })
        # Override total_distribution to reflect actual scenario total
        total_for_summary = scenario_total
    else:
        # Fallback — no scenario selected; use static distribution_percentage
        for ben in beneficiaries:
            name = (ben.full_name or "").strip()
            if not name:
                _missing_names.append(ben.pk)
            pct = ben.distribution_percentage or Decimal("0")
            amount = (net_profit_raw * pct / 100).quantize(
                Decimal("0.01"), rounding=ROUND_HALF_UP
            )
            pct_display = f"{pct:.2f}" if pct else "0.00"
            if amount == 0:
                amount_display = "-"
            elif amount < 0:
                amount_display = f"({abs(amount):,.0f})"
            else:
                amount_display = f"{amount:,.0f}"
            distributions.append({
                "beneficiary_name": name or "— Name missing —",
                "percentage": pct_display,
                "amount": amount_display,
                "amount_raw": amount,
            })
        total_for_summary = net_profit_raw

    # Raise CRITICAL Eva finding if any beneficiary name is missing
    if _missing_names:
        logger.warning(
            "Missing beneficiary name(s) for entity %s FY %s: officer PKs %s",
            entity.entity_name, financial_year.pk, _missing_names,
        )
        try:
            from core.models import EvaReview, EvaFinding
            import hashlib, json as _json
            _active_review = EvaReview.objects.filter(
                financial_year=financial_year,
            ).order_by("-created_at").first()
            if _active_review:
                _fp = hashlib.sha256(_json.dumps({
                    "entity_id": str(entity.pk),
                    "financial_year_id": str(financial_year.pk),
                    "rule_category": "beneficiary_name_missing",
                }, sort_keys=True).encode()).hexdigest()
                EvaFinding.objects.update_or_create(
                    eva_review=_active_review,
                    check_name="beneficiary_name_missing",
                    defaults={
                        "severity": "critical",
                        "title": "Beneficiary name missing",
                        "plain_english_explanation": (
                            f"{len(_missing_names)} beneficiary record(s) have no name. "
                            "The Distribution Summary cannot be completed accurately."
                        ),
                        "recommendation": (
                            "Update the beneficiary records with full names before "
                            "generating the financial statements."
                        ),
                        "source": "risk_engine",
                        "fingerprint": _fp,
                    },
                )
        except Exception as _e:
            logger.error("Could not create beneficiary_name_missing finding: %s", _e)

    # Debug log for distribution context verification
    logger.info(
        "Distribution context for %s FY %s: %d beneficiaries, net_profit_raw=%s, "
        "distributions=%s",
        entity.entity_name, financial_year.pk, len(distributions),
        net_profit_raw,
        [{k: v for k, v in d.items() if k != "amount_raw"} for d in distributions],
    )

    context["beneficiaries"] = distributions
    # Format total distribution as whole dollars for the Distribution Summary
    if total_for_summary == 0:
        context["total_distribution"] = "-"
    elif total_for_summary < 0:
        context["total_distribution"] = f"({abs(total_for_summary):,.0f})"
    else:
        context["total_distribution"] = f"{total_for_summary:,.0f}"

    return context


# ---------------------------------------------------------------------------
# 5. build_sole_trader_context
# ---------------------------------------------------------------------------
def build_sole_trader_context(financial_year, include_watermark=True):
    """Build full Jinja2 context dict for a sole trader entity."""
    from core.models import EntityOfficer

    context = build_company_context(financial_year, include_watermark)

    entity = financial_year.entity

    context["declaration_title"] = "Proprietor Declaration"
    context["compilation_responsible_party"] = "owner"

    proprietor = EntityOfficer.objects.filter(
        entity=entity,
        role="sole_trader",
        date_ceased__isnull=True,
    ).first()

    if proprietor:
        context["directors"] = [
            {"name": proprietor.full_name, "title": "Proprietor"}
        ]
    else:
        context["directors"] = [
            {"name": entity.entity_name, "title": "Proprietor"}
        ]

    return context


# ---------------------------------------------------------------------------
# Post-processing — borders, page numbers, page-break prevention, ampersand
# ---------------------------------------------------------------------------
# Handiledger row type classification for post-processor border application.
# Section totals: single above + double below on amount cols
_SECTION_TOTAL_LABELS = [
    "total income", "total expenses", "total revenue",
    "total current assets", "total non-current assets",
    "total current liabilities", "total non-current liabilities",
    "total equity",
    "gross profit",
]

# Major totals: double below only on amount cols
_MAJOR_TOTAL_LABELS = [
    "total assets", "total liabilities", "net assets",
]

# Grand totals: double below only on amount cols
_GRAND_TOTAL_LABELS = [
    "net profit", "net loss", "net profit / (loss)", "net profit/(loss)",
    "operating profit after income tax",
    "operating profit before income tax",
]

# All summary labels (union of above — used for detecting any total row)
_SUMMARY_LABELS = _SECTION_TOTAL_LABELS + _MAJOR_TOTAL_LABELS + _GRAND_TOTAL_LABELS

_SUB_HEADING_LABELS = [
    "cash assets", "receivables", "other current assets",
    "payables", "tax liabilities", "other current liabilities",
]


def _post_process_fs_doc(buffer, doc_type, has_prior=True):
    """Post-process a rendered financial statement .docx.

    Applied to ALL document types to fix:
      - Firm-name ampersand stripped by XML rendering (all doc types).
      - Page number in footer (all doc types).
      - Ruling lines / borders on summary rows (P&L, BS, Summary).
      - Duplicate "Net Profit / (Loss)" standalone paragraph (P&L).
      - cantSplit + keepNext on summary rows (P&L, BS).
      - keepNext on section heading paragraphs before tables.
      - Inline "(PY: xxx)" values in totals rows (BS).
    """
    import copy
    import re
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(buffer)

    # ------------------------------------------------------------------
    # Restore ampersands: both the §AMP§ placeholder (from _safe_amp)
    # and the firm-name "MC S" pattern (from direct XML stripping).
    # ------------------------------------------------------------------
    def _restore_amps(text):
        if not text:
            return text
        text = text.replace(_AMP_PLACEHOLDER, "&")
        text = text.replace("MC  S Pty Ltd", "MC & S Pty Ltd")
        text = text.replace("MC S Pty Ltd", "MC & S Pty Ltd")
        return text

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.text:
                run.text = _restore_amps(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text:
                            run.text = _restore_amps(run.text)
    # Also restore in headers
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            for paragraph in hdr.paragraphs:
                for run in paragraph.runs:
                    if run.text:
                        run.text = _restore_amps(run.text)

    # ------------------------------------------------------------------
    # Remove any existing PAGE field footers — page numbers are stamped
    # on the final merged PDF so they run continuously.
    # ------------------------------------------------------------------
    for section in doc.sections:
        footer = section.footer
        paras_to_remove = []
        for p in footer.paragraphs:
            has_page_field = any(
                el.tag == qn('w:fldChar') or el.tag == qn('w:instrText')
                for el in p._p.iter()
            )
            if has_page_field:
                paras_to_remove.append(p._p)
        for p_el in paras_to_remove:
            p_el.getparent().remove(p_el)

    # ------------------------------------------------------------------
    # P&L / BS / Summary: borders, page-break fixes, duplicate removal
    # ------------------------------------------------------------------
    if doc_type in ("DETAILED_PL", "BALANCE_SHEET", "SUMMARY_PL"):

        # Remove duplicate standalone "Net Profit / (Loss)" paragraph
        # that precedes the table containing the same label + values.
        if doc_type == "DETAILED_PL":
            body = doc.element.body
            paragraphs_to_remove = []
            all_elements = list(body)
            for i, el in enumerate(all_elements):
                if el.tag != qn('w:p'):
                    continue
                para_text = ''.join(
                    t.text or '' for t in el.iter(qn('w:t'))
                ).strip()
                if para_text == "Net Profit / (Loss)":
                    # Only remove if followed by a table that also has this label
                    for j in range(i + 1, min(i + 3, len(all_elements))):
                        if all_elements[j].tag == qn('w:tbl'):
                            tbl_text = ''.join(
                                t.text or '' for t in all_elements[j].iter(qn('w:t'))
                            )
                            if "Net Profit" in tbl_text:
                                paragraphs_to_remove.append(el)
                            break
            for el in paragraphs_to_remove:
                body.remove(el)

        # keepNext on section heading paragraphs before tables
        body = doc.element.body
        all_elements = list(body)
        for i, el in enumerate(all_elements):
            if el.tag == qn('w:p') and i + 1 < len(all_elements):
                next_el = all_elements[i + 1]
                if next_el.tag == qn('w:tbl'):
                    para_text = ''.join(
                        t.text or '' for t in el.iter(qn('w:t'))
                    ).strip()
                    if para_text and len(para_text) < 50:
                        pPr = el.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            el.insert(0, pPr)
                        if pPr.find(qn('w:keepNext')) is None:
                            keepNext = OxmlElement('w:keepNext')
                            keepNext.set(qn('w:val'), '1')
                            pPr.append(keepNext)

        # Fix 8: Section integrity — set keepWithNext on every row in each
        # table EXCEPT the last row.  This chains all rows together so the
        # entire section moves to the next page if it does not fit.
        for table in doc.tables:
            rows = list(table.rows)
            for idx, row in enumerate(rows):
                is_last = (idx == len(rows) - 1)
                for cell in row.cells:
                    for para in cell.paragraphs:
                        ppPr = para._p.get_or_add_pPr()
                        # Remove existing keepNext first
                        for existing_kn in ppPr.findall(qn('w:keepNext')):
                            ppPr.remove(existing_kn)
                        if not is_last:
                            kn = OxmlElement('w:keepNext')
                            kn.set(qn('w:val'), '1')
                            ppPr.append(kn)

        # Process table rows: borders + cantSplit + inline PY fix
        for table in doc.tables:
            for row in table.rows:
                if not row.cells:
                    continue
                first_cell_text = row.cells[0].text.strip().lower()

                is_section_total = any(lbl in first_cell_text for lbl in _SECTION_TOTAL_LABELS)
                is_major_total = any(lbl in first_cell_text for lbl in _MAJOR_TOTAL_LABELS)
                is_grand_total = any(lbl in first_cell_text for lbl in _GRAND_TOTAL_LABELS)
                is_summary = is_section_total or is_major_total or is_grand_total

                if is_summary:
                    # cantSplit — prevent the row itself from splitting
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    for existing in trPr.findall(qn('w:cantSplit')):
                        trPr.remove(existing)
                    cantSplit = OxmlElement('w:cantSplit')
                    cantSplit.set(qn('w:val'), '1')
                    trPr.append(cantSplit)

                    # Determine border style per Handiledger standard
                    if is_section_total:
                        # Section subtotal: single above + single below
                        amount_top = {'val': 'single', 'sz': '6'}
                        amount_bot = {'val': 'single', 'sz': '6'}
                    else:
                        # Major total / grand total: single above + double below
                        amount_top = {'val': 'single', 'sz': '6'}
                        amount_bot = {'val': 'double', 'sz': '8'}

                    num_cells = len(row.cells)
                    for cell_idx, cell in enumerate(row.cells):
                        is_amount_col = cell_idx >= num_cells - 2
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        existing_borders = tcPr.find(qn('w:tcBorders'))
                        if existing_borders is not None:
                            tcPr.remove(existing_borders)
                        tcBorders = OxmlElement('w:tcBorders')
                        tcPr.append(tcBorders)
                        if not is_amount_col:
                            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                                el = OxmlElement(f'w:{side}')
                                el.set(qn('w:val'), 'none')
                                el.set(qn('w:sz'), '0')
                                el.set(qn('w:color'), 'auto')
                                tcBorders.append(el)
                            continue
                        # Amount columns: apply per row type
                        top_el = OxmlElement('w:top')
                        top_el.set(qn('w:val'), amount_top['val'])
                        top_el.set(qn('w:sz'), amount_top['sz'])
                        top_el.set(qn('w:space'), '0')
                        top_el.set(qn('w:color'), '000000')
                        tcBorders.append(top_el)
                        bot_el = OxmlElement('w:bottom')
                        bot_el.set(qn('w:val'), amount_bot['val'])
                        bot_el.set(qn('w:sz'), amount_bot['sz'])
                        bot_el.set(qn('w:space'), '0')
                        bot_el.set(qn('w:color'), '000000')
                        tcBorders.append(bot_el)
                        for side in ('left', 'right', 'insideH', 'insideV'):
                            el = OxmlElement(f'w:{side}')
                            el.set(qn('w:val'), 'none')
                            el.set(qn('w:sz'), '0')
                            el.set(qn('w:color'), 'auto')
                            tcBorders.append(el)
                    # Bold all text in summary rows
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True

                # Sub-group headings (Cash Assets, Receivables, etc.) — bold only
                is_sub_heading = any(
                    lbl == first_cell_text for lbl in _SUB_HEADING_LABELS
                )
                if is_sub_heading:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True

                # Sub-group subtotal rows — empty label, amounts present,
                # single top border on AMOUNT columns only.
                # Exclude column-header rows (year number or "$" in cell 2).
                _cell2 = row.cells[2].text.strip() if len(row.cells) >= 3 else ""
                _is_col_header = _cell2.startswith("$") or (_cell2.isdigit() and len(_cell2) == 4)
                if (not first_cell_text
                        and len(row.cells) >= 3
                        and _cell2
                        and not _is_col_header):
                    num_cells = len(row.cells)
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.bold = True
                    # Border on last 2 cells (amount columns) only
                    for cell_idx in range(max(0, num_cells - 2), num_cells):
                        tc = row.cells[cell_idx]._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcBorders = tcPr.find(qn('w:tcBorders'))
                        if tcBorders is None:
                            tcBorders = OxmlElement('w:tcBorders')
                            tcPr.append(tcBorders)
                        top_el = tcBorders.find(qn('w:top'))
                        if top_el is None:
                            top_el = OxmlElement('w:top')
                            tcBorders.append(top_el)
                        top_el.set(qn('w:val'), 'single')
                        top_el.set(qn('w:sz'), '4')
                        top_el.set(qn('w:space'), '0')
                        top_el.set(qn('w:color'), '000000')

                # Indent "Less:" rows (accumulated depreciation/amortisation)
                if first_cell_text.startswith("less:"):
                    label_para = row.cells[0].paragraphs[0]
                    from docx.shared import Cm as _Cm
                    label_para.paragraph_format.left_indent = _Cm(0.5)

                # Fix inline PY values in Balance Sheet totals
                if doc_type == "BALANCE_SHEET":
                    _fix_inline_py_in_row(row, qn, OxmlElement, copy, re)

    # ------------------------------------------------------------------
    # Fix 8: Section integrity for short documents — keep all paragraphs
    # together so they never split across pages.
    # ------------------------------------------------------------------
    if doc_type in ("DECLARATION", "COMPILATION", "NOTES"):
        body_paras = doc.paragraphs
        for idx, para in enumerate(body_paras):
            is_last = (idx == len(body_paras) - 1)
            ppPr = para._p.get_or_add_pPr()
            # Remove existing keepNext
            for existing_kn in ppPr.findall(qn('w:keepNext')):
                ppPr.remove(existing_kn)
            if not is_last:
                kn = OxmlElement('w:keepNext')
                kn.set(qn('w:val'), '1')
                ppPr.append(kn)

    # ------------------------------------------------------------------
    # Comparative column suppression
    # When has_prior=False, strip the last column from every financial
    # table in the document.  This handles templates that were designed
    # with a comparative column but are being rendered for a first-year
    # entity that has no prior-year trial balance data.
    # ------------------------------------------------------------------
    if not has_prior:
        _strip_comparative_column(doc, qn)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def _strip_comparative_column(doc, qn):
    """Remove the last column from every table in the document.

    Used when has_prior=False to suppress comparative year columns that
    may have been rendered by the Jinja2 template.
    Tables with only 1 or 2 columns are left untouched (they are likely
    notes or label-only tables, not financial statement tables).
    Also removes trailing tab characters from header/footer paragraphs
    that would leave a dangling tab stop for the missing column.
    """
    from docx.oxml.ns import qn as _qn
    # Strip last column from tables with 3+ columns
    for table in doc.tables:
        if len(table.columns) < 3:
            continue
        for row in table.rows:
            tr = row._tr
            cells = list(tr.findall(qn('w:tc')))
            if not cells:
                continue
            last_tc = cells[-1]
            # Skip if this cell has a gridSpan > 1 (merged header cell)
            tcPr = last_tc.find(qn('w:tcPr'))
            if tcPr is not None:
                gs = tcPr.find(qn('w:gridSpan'))
                if gs is not None:
                    try:
                        span_val = int(gs.get(qn('w:val'), '1'))
                        if span_val > 1:
                            # Reduce gridSpan by 1 instead of removing
                            gs.set(qn('w:val'), str(span_val - 1))
                            continue
                    except (ValueError, TypeError):
                        pass
            tr.remove(last_tc)
        # Update tblGrid to remove last gridCol
        tbl = table._tbl
        tblGrid = tbl.find(qn('w:tblGrid'))
        if tblGrid is not None:
            gridCols = list(tblGrid.findall(qn('w:gridCol')))
            if gridCols:
                tblGrid.remove(gridCols[-1])

    # Strip trailing tab from header/footer paragraphs that had a
    # right-aligned prior-year column tab stop.
    def _strip_trailing_tab(para):
        for run in para.runs:
            if run.text and run.text.endswith('\t'):
                run.text = run.text.rstrip('\t')

    for section in doc.sections:
        for container in [section.header, section.footer]:
            for para in container.paragraphs:
                _strip_trailing_tab(para)

    # Also strip trailing tabs from body paragraphs (column headers)
    for para in doc.paragraphs:
        _strip_trailing_tab(para)


def _fix_inline_py_in_row(row, qn, OxmlElement, copy, re):
    """Split a merged cell containing 'CY (PY: xxx)' into separate CY and PY cells."""
    tr = row._tr
    tcs = list(tr.findall(qn('w:tc')))

    for tc in tcs:
        # Collect all text from the cell
        text = ''.join(t.text or '' for t in tc.iter(qn('w:t')))
        match = re.search(r'\(PY:\s*(.+?)\)\s*$', text)
        if not match:
            continue

        py_text = match.group(1).strip()
        cy_text = text[:match.start()].strip()

        # Check for gridSpan (merged cell)
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is not None:
                span = int(gridSpan.get(qn('w:val'), '1'))
                if span > 1:
                    if span - 1 > 1:
                        gridSpan.set(qn('w:val'), str(span - 1))
                    else:
                        tcPr.remove(gridSpan)

        # Set this cell's text to CY value only
        _set_cell_text(tc, cy_text, qn)

        # Create new cell for PY value (clone formatting from CY cell)
        new_tc = copy.deepcopy(tc)
        # Remove gridSpan from new cell
        new_tcPr = new_tc.find(qn('w:tcPr'))
        if new_tcPr is not None:
            new_gs = new_tcPr.find(qn('w:gridSpan'))
            if new_gs is not None:
                new_tcPr.remove(new_gs)

        _set_cell_text(new_tc, py_text, qn)
        tc.addnext(new_tc)
        break  # Only fix one cell per row


def _set_cell_text(tc, text, qn):
    """Set the text of all runs in a table cell, preserving formatting."""
    first_set = False
    for p in tc.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                if not first_set:
                    t.text = text
                    first_set = True
                else:
                    t.text = ''
    if not first_set:
        # No runs found — create one
        p_els = tc.findall(qn('w:p'))
        if p_els:
            r_el = OxmlElement('w:r')
            t_el = OxmlElement('w:t')
            t_el.text = text
            r_el.append(t_el)
            p_els[0].append(r_el)


# ---------------------------------------------------------------------------
# 6. render_template
# ---------------------------------------------------------------------------
def render_template(template_db_record, context):
    """Load .docx via DocxTemplate, render with Jinja2 context, return BytesIO."""
    from docxtpl import DocxTemplate

    if template_db_record is None:
        raise ValueError("render_template called with None template record")

    if not template_db_record.template_file:
        raise ValueError(
            f"Template file not found for {template_db_record.document_type} / "
            f"{template_db_record.entity_type}: FileField is empty "
            f"(record pk={template_db_record.pk})"
        )

    try:
        template_path = template_db_record.template_file.path
    except ValueError:
        raise ValueError(
            f"Template file not found for {template_db_record.document_type} / "
            f"{template_db_record.entity_type}: no file associated with FileField "
            f"(record pk={template_db_record.pk})"
        )

    if not os.path.exists(template_path):
        raise ValueError(
            f"Template file not found for {template_db_record.document_type} / "
            f"{template_db_record.entity_type}: {template_path} does not exist on disk "
            f"(record pk={template_db_record.pk})"
        )

    tpl = DocxTemplate(template_path)

    # Inject InlineImage logo if context carries a builder reference
    # (new-style callers via DocumentContextBuilder)
    builder = context.pop("__builder__", None)
    if builder is not None:
        builder.tpl = tpl
        # Re-resolve logo now that tpl is available
        from core.models import FirmSettings
        firm = FirmSettings.get()
        context["practice_logo"] = builder._resolve_logo_for_docx(firm)

    # Use the StatementHub Jinja2 environment with all custom filters
    from core.document_context_builder import get_jinja_env
    jinja_env = get_jinja_env()

    tpl.render(context, jinja_env=jinja_env)

    buffer = io.BytesIO()
    tpl.save(buffer)
    buffer.seek(0)

    # Cleanup temp logo files if builder was used
    if builder is not None:
        builder.cleanup()

    return buffer


# ---------------------------------------------------------------------------
# 6b. Programmatic Notes Document Generator
# ---------------------------------------------------------------------------

# Reuse constants from generate_fs_templates for visual consistency
_NOTES_FONT = "Times New Roman"
_NOTES_FONT_SIZE = Pt(10)
_NOTES_MARGIN_TOP = Cm(1.6)     # 16mm
_NOTES_MARGIN_BOTTOM = Cm(1.7)  # 17mm
_NOTES_MARGIN_LEFT = Cm(2.0)    # 20mm
_NOTES_MARGIN_RIGHT = Cm(2.4)   # 24mm
_NOTES_COL_WIDTHS_3 = [Cm(10), Cm(3), Cm(3)]  # label, CY, PY


def _notes_spacer(doc, size_pt=8):
    """Insert a blank paragraph as a visual spacer.

    More reliable than space_before/space_after when LibreOffice is in
    the PDF conversion chain — LO sometimes strips paragraph spacing
    from python-docx documents.
    """
    p = doc.add_paragraph()
    run = p.add_run("")
    run.font.name = _NOTES_FONT
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _notes_keep_with_next(para):
    """Set keepNext on a paragraph so it stays with the following element."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    if pPr.find(qn('w:keepNext')) is None:
        kn = OxmlElement('w:keepNext')
        kn.set(qn('w:val'), '1')
        pPr.append(kn)


def _notes_keep_together(para):
    """Set keepLines on a paragraph so it does not split across pages."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    if pPr.find(qn('w:keepLines')) is None:
        kl = OxmlElement('w:keepLines')
        kl.set(qn('w:val'), '1')
        pPr.append(kl)


def _notes_table_keep_with_next(table):
    """Set keepNext on all paragraphs in the last row of a table."""
    rows = list(table.rows)
    if not rows:
        return
    for cell in rows[-1].cells:
        for para in cell.paragraphs:
            _notes_keep_with_next(para)


def _notes_add_para(doc, text, bold=False, italic=False, size=None,
                    alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=None,
                    space_before=None, left_indent=None):
    """Add a styled paragraph to the notes document."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = _NOTES_FONT
    run.font.size = size or _NOTES_FONT_SIZE
    run.bold = bold
    run.font.italic = italic
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    return p


def _notes_add_table_row(table, label, cy_str, py_str, bold=False,
                         indent=None):
    """Add a data row to a 3-column notes table."""
    row = table.add_row()
    for i, text in enumerate([label, cy_str, py_str]):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.text = ""
        run = p.add_run(text)
        run.font.name = _NOTES_FONT
        run.font.size = _NOTES_FONT_SIZE
        run.bold = bold
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        if i >= 1:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if indent:
        row.cells[0].paragraphs[0].paragraph_format.left_indent = indent
    return row


def _notes_apply_subtotal_border(row):
    """Single top border on amount columns (cols 1, 2)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for i in [1, 2]:
        tc = row.cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)


def _notes_apply_grand_total_border(row):
    """Single top + double bottom on amount columns (cols 1, 2)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for i in [1, 2]:
        tc = row.cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), '000000')
        tcBorders.append(top)
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'double')
        bot.set(qn('w:sz'), '12')
        bot.set(qn('w:space'), '0')
        bot.set(qn('w:color'), '000000')
        tcBorders.append(bot)


def _notes_create_table(doc, has_prior=True):
    """Create a 3-column borderless table for notes (label, CY, PY).

    Column widths are set explicitly so the label column is wide enough
    to fit 'Less: Accumulated depreciation' on a single line.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    cols = 3 if has_prior else 2
    table = doc.add_table(rows=0, cols=cols)
    table.autofit = False
    tbl = table._tbl
    tblPr = tbl.tblPr
    # Full width
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9356')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    # Fixed layout so Word/LO respects our column widths
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    existing_layout = tblPr.find(qn('w:tblLayout'))
    if existing_layout is not None:
        tblPr.remove(existing_layout)
    tblPr.append(tblLayout)
    # Set column widths via tblGrid
    tblGrid = OxmlElement('w:tblGrid')
    widths = _NOTES_COL_WIDTHS_3 if has_prior else _NOTES_COL_WIDTHS_3[:2]
    for w in widths:
        gridCol = OxmlElement('w:gridCol')
        # Convert Cm to twips (1 cm = 567 twips)
        twips = int(w.cm * 567)
        gridCol.set(qn('w:w'), str(twips))
        tblGrid.append(gridCol)
    existing_grid = tbl.find(qn('w:tblGrid'))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    # tblGrid must come after tblPr
    tbl.insert(list(tbl).index(tblPr) + 1, tblGrid)
    # Remove all borders
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    return table


def _fmt_note_amount(value):
    """Format Decimal for notes tables — same convention as format_amount."""
    if value is None:
        return "-"
    d = Decimal(str(value)).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
    if d == 0:
        return "-"
    if d < 0:
        return f"({abs(d):,.0f})"
    return f"{d:,.0f}"


def _fmt_dollar(value):
    """Format Decimal with $ prefix for prose. Returns 'nil' for zero."""
    if value is None or value == 0:
        return "nil"
    d = Decimal(str(value)).quantize(Decimal("1"), rounding=ROUND_HALF_UP)
    if d == 0:
        return "nil"
    if d < 0:
        return f"$({abs(d):,.0f})"
    return f"${d:,.0f}"


def _generate_notes_document(context):
    """Build the Notes to Financial Statements as a python-docx Document.

    Called at render time with full context including raw TB sections.
    Returns a BytesIO buffer containing the .docx.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    entity = context["_entity"]
    fy = context["_fy"]
    sections = context["_sections"]
    entity_type = entity.entity_type
    entity_name = entity.entity_name
    income_tax_cy = context["_income_tax_cy"]
    income_tax_py = context["_income_tax_py"]
    has_income_tax = context["_has_income_tax"]
    total_revenue_cy = context["_total_revenue_cy"]
    total_revenue_py = context["_total_revenue_py"]

    year_end = fy.end_date
    year_str = str(year_end.year) if year_end else ""
    prior_year_str = str(year_end.year - 1) if year_end else ""
    date_text = f"For the Year Ended {year_end.strftime('%d %B %Y')}" if year_end else ""
    has_prior = context.get("has_prior", False)

    # --- Compute note triggers from TB sections ---
    # Trade Receivables
    trade_debtors = []
    provision_doubtful = None
    for item in sections["current_assets"]:
        nl = item["account_name"].lower()
        if "trade" in nl and "debtor" in nl:
            trade_debtors.append(item)
        elif "provision" in nl and "doubtful" in nl:
            if item["cy_amount"] != 0 or item["py_amount"] != 0:
                provision_doubtful = item
    has_trade_debtors = any(
        i["cy_amount"] != 0 or i["py_amount"] != 0 for i in trade_debtors
    )

    # PPE — group into asset classes by TB proximity.
    # In the TB, a cost account is immediately followed by its depreciation
    # account. Walk the NCA list in order and build (cost, depr) pairs.
    ppe_classes = []       # list of {"cost": item, "depr": item_or_None}
    ppe_deposit = []
    nca_items = sections["noncurrent_assets"]
    i_nca = 0
    while i_nca < len(nca_items):
        item = nca_items[i_nca]
        nl = item["account_name"].lower()
        is_depr = any(kw in nl for kw in [
            "accumulated", "amortisation", "depreciation",
        ]) or nl.startswith("less:")
        is_deposit = "deposit" in nl
        is_ppe = any(kw in nl for kw in [
            "equipment", "vehicle", "furniture", "building", "fixture",
            "plant", "motor", "computer", "office", "at cost",
        ]) or is_depr or is_deposit

        if not is_ppe:
            i_nca += 1
            continue

        if is_deposit:
            ppe_deposit.append(item)
            i_nca += 1
            continue

        if is_depr:
            # Orphan depreciation with no preceding cost — add as depr-only
            ppe_classes.append({"cost": None, "depr": item})
            i_nca += 1
            continue

        # Cost account — check if next item is its depreciation pair
        depr_item = None
        if i_nca + 1 < len(nca_items):
            next_item = nca_items[i_nca + 1]
            next_nl = next_item["account_name"].lower()
            next_is_depr = any(kw in next_nl for kw in [
                "accumulated", "amortisation", "depreciation",
            ]) or next_nl.startswith("less:")
            if next_is_depr:
                depr_item = next_item
                i_nca += 1  # skip the depr item in the outer loop

        ppe_classes.append({"cost": item, "depr": depr_item})
        i_nca += 1

    has_ppe = any(
        pair["cost"] is not None and
        (pair["cost"]["cy_amount"] != 0 or pair["cost"]["py_amount"] != 0)
        for pair in ppe_classes
    )

    # Related Party — management fees
    mgmt_fee_items = []
    for item in sections["expenses"]:
        nl = item["account_name"].lower()
        if "management" in nl and "fee" in nl and (
            "majoti" in nl or "related" in nl
        ):
            mgmt_fee_items.append(item)
    has_mgmt_fees = any(
        i["cy_amount"] != 0 or i["py_amount"] != 0 for i in mgmt_fee_items
    )

    # Related Party — director loans
    director_loan_items = []
    related_loan_items = []
    for item in sections["noncurrent_liabilities"]:
        nl = item["account_name"].lower()
        if "loan" in nl and "director" in nl:
            director_loan_items.append(item)
        elif "loan" in nl and any(kw in nl for kw in [
            "majoti", "ets", "related",
        ]):
            related_loan_items.append(item)
    has_director_loan = (
        entity_type != "sole_trader" and
        any(i["cy_amount"] != 0 or i["py_amount"] != 0 for i in director_loan_items)
    )
    has_related_loans = any(
        i["cy_amount"] != 0 or i["py_amount"] != 0 for i in related_loan_items
    )
    has_related_party = has_mgmt_fees or has_director_loan or has_related_loans

    is_company = entity_type == "company"

    # --- Note map — use precomputed from context, or compute fresh ---
    note_map = context.get("_note_map")
    if not note_map:
        note_map, _ = _compute_note_map(sections, entity_type, has_income_tax)

    def _note_num_for(note_type):
        for n, t in note_map:
            if t == note_type:
                return n
        return None

    # --- Build document ---
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = _NOTES_FONT
    style.font.size = _NOTES_FONT_SIZE
    for section in doc.sections:
        section.top_margin = _NOTES_MARGIN_TOP
        section.bottom_margin = _NOTES_MARGIN_BOTTOM
        section.left_margin = _NOTES_MARGIN_LEFT
        section.right_margin = _NOTES_MARGIN_RIGHT

    # --- Header (match other statements exactly) ---
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    p1 = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p1.text = ""
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(entity_name)
    r1.font.name = _NOTES_FONT
    r1.font.size = Pt(11)
    r1.bold = True
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(0)

    abn_raw = entity.abn or ""
    abn_digits = "".join(c for c in str(abn_raw) if c.isdigit())
    abn_formatted = (
        f"{abn_digits[:2]} {abn_digits[2:5]} {abn_digits[5:8]} {abn_digits[8:]}"
        if len(abn_digits) == 11 else abn_raw
    )
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"ABN {abn_formatted}")
    r2.font.name = _NOTES_FONT
    r2.font.size = Pt(9)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)

    p3 = header.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Notes to the Financial Statements")
    r3.font.name = _NOTES_FONT
    r3.font.size = Pt(9)
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)

    p4 = header.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(date_text)
    r4.font.name = _NOTES_FONT
    r4.font.size = Pt(9)
    # space_after = ~1cm (284,200 EMUs) — gap between header text and rule
    p4.paragraph_format.space_after = Emu(284200)
    p4.paragraph_format.space_before = Pt(0)
    # Horizontal rule
    pPr = p4._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:space'), '1')
    bottom_border.set(qn('w:color'), '000000')
    pBdr.append(bottom_border)
    pPr.append(pBdr)

    # Watermark placeholder
    pw = header.add_paragraph()
    pw.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rw = pw.add_run(context.get("watermark", ""))
    rw.font.name = _NOTES_FONT
    rw.font.size = Pt(14)
    rw.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    rw.bold = True
    pw.paragraph_format.space_after = Pt(0)
    pw.paragraph_format.space_before = Pt(0)

    # --- Footer ---
    # Centre-aligned Times New Roman italic 9pt with 0.5pt top border
    # (horizontal rule above), matching Handiledger standard.
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "These financial statements are unaudited. They must be read in "
        "conjunction with the attached Accountant\u2019s Compilation Report "
        "and Notes which form part of these financial statements."
    )
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(9)
    fr.font.italic = True
    _fpPr = fp._p.get_or_add_pPr()
    for _existing in _fpPr.findall(qn('w:pBdr')):
        _fpPr.remove(_existing)
    _fpBdr = OxmlElement('w:pBdr')
    _fpTop = OxmlElement('w:top')
    _fpTop.set(qn('w:val'), 'single')
    _fpTop.set(qn('w:sz'), '4')
    _fpTop.set(qn('w:space'), '4')
    _fpTop.set(qn('w:color'), '000000')
    _fpBdr.append(_fpTop)
    _fpPr.append(_fpBdr)

    # 0.5cm gap between header rule and first body content
    _p0 = doc.add_paragraph()
    _p0.paragraph_format.space_before = Emu(142100)
    _p0.paragraph_format.space_after = Pt(0)

    # ==================================================================
    # NOTE 1: Statement of Significant Accounting Policies
    # ==================================================================
    n1 = _note_num_for("policies")
    _notes_add_para(doc, f"Note {n1}: Statement of Significant Accounting Policies",
                    bold=True, space_before=0, space_after=6)

    # Opening paragraph — entity type specific
    if entity_type == "company":
        _notes_add_para(
            doc,
            "The financial statements are special purpose financial statements "
            "prepared in order to satisfy the financial reporting requirements of the "
            "Corporations Act 2001. The directors have determined that the entity is "
            "not a reporting entity.",
            space_before=6, space_after=6)
    elif entity_type == "trust":
        _notes_add_para(
            doc,
            "The financial statements are special purpose financial statements "
            "prepared in order to satisfy the financial reporting requirements of the "
            "trust deed. The trustees have determined that the entity is not a "
            "reporting entity.",
            space_before=6, space_after=6)
    elif entity_type == "partnership":
        _notes_add_para(
            doc,
            "The financial statements are special purpose financial statements "
            "prepared in order to satisfy the financial reporting requirements of the "
            "partnership agreement. The partners have determined that the entity is "
            "not a reporting entity.",
            space_before=6, space_after=6)
    else:  # sole_trader
        _notes_add_para(
            doc,
            "The financial statements are special purpose financial statements "
            "prepared in order to satisfy the information needs of the proprietor.",
            space_before=6, space_after=6)

    _notes_add_para(
        doc,
        "The financial statements have been prepared on an accruals basis and are "
        "based on historical costs.",
        space_after=6)

    _notes_add_para(
        doc,
        "The following significant accounting policies have been adopted in the "
        "preparation and presentation of the financial statements:",
        space_after=8)

    # Sub-clauses — lettered sequentially
    policy_letter = ord("a")

    # (a) Revenue Recognition — always
    _notes_spacer(doc, 4)
    _notes_add_para(doc, f"{chr(policy_letter)}) Revenue Recognition",
                    bold=True, space_before=10, space_after=2)
    _notes_add_para(
        doc,
        "Revenue is recognised when the entity satisfies a performance obligation "
        "by transferring a promised good or service to a customer.",
        space_before=2, space_after=4)
    policy_letter += 1

    # Income Tax — companies only
    if is_company:
        _notes_spacer(doc, 4)
        _notes_add_para(doc, f"{chr(policy_letter)}) Income Tax",
                        bold=True, space_before=10, space_after=2)
        _notes_add_para(
            doc,
            "The income tax expense for the year comprises current income tax expense. "
            "Current income tax expense reflects the current year tax payable based on "
            "taxable income for the year.",
            space_before=2, space_after=4)
        policy_letter += 1

    # GST — always
    _notes_spacer(doc, 4)
    _notes_add_para(doc, f"{chr(policy_letter)}) Goods and Services Tax (GST)",
                    bold=True, space_before=10, space_after=2)
    _notes_add_para(
        doc,
        "Revenues, expenses and assets are recognised net of the amount of GST. "
        "Receivables and payables are stated with the amount of GST included.",
        space_before=2, space_after=4)
    policy_letter += 1

    # PPE — only if has_ppe
    if has_ppe:
        _notes_spacer(doc, 4)
        _notes_add_para(doc, f"{chr(policy_letter)}) Property, Plant and Equipment",
                        bold=True, space_before=10, space_after=2)
        _notes_add_para(
            doc,
            "Property, plant and equipment are carried at cost less any subsequent "
            "accumulated depreciation and impairment losses. Depreciation is calculated "
            "on a diminishing value basis over the estimated useful life of the asset.",
            space_before=2, space_after=4)
        policy_letter += 1

    # ==================================================================
    # NOTE 2: Trade Receivables
    # ==================================================================
    if has_trade_debtors:
        _notes_spacer(doc, 8)
        n = _note_num_for("receivables")
        heading_p = _notes_add_para(doc, f"Note {n}: Trade Receivables",
                                    bold=True, space_before=18, space_after=6)

        tbl = _notes_create_table(doc, has_prior)
        total_cy = Decimal("0")
        total_py = Decimal("0")
        for item in trade_debtors:
            cy = abs(item["cy_amount"]) if item["cy_amount"] else Decimal("0")
            py = abs(item["py_amount"]) if item["py_amount"] else Decimal("0")
            total_cy += cy
            total_py += py
            _notes_add_table_row(tbl, "Trade debtors",
                                 _fmt_note_amount(cy), _fmt_note_amount(py))

        if provision_doubtful:
            pcy = -abs(provision_doubtful["cy_amount"]) if provision_doubtful["cy_amount"] else Decimal("0")
            ppy = -abs(provision_doubtful["py_amount"]) if provision_doubtful["py_amount"] else Decimal("0")
            total_cy += pcy
            total_py += ppy
            _notes_add_table_row(tbl, "Less: Provision for doubtful debts",
                                 _fmt_note_amount(pcy), _fmt_note_amount(ppy))

            # Subtotal rule only when there are multiple lines above Total
            sub_row = _notes_add_table_row(tbl, "", _fmt_note_amount(total_cy),
                                           _fmt_note_amount(total_py))
            _notes_apply_subtotal_border(sub_row)

        total_row = _notes_add_table_row(tbl, "Total", _fmt_note_amount(total_cy),
                                         _fmt_note_amount(total_py), bold=True)
        _notes_apply_grand_total_border(total_row)

        _notes_add_para(
            doc,
            "Trade receivables are non-interest bearing and are generally on 30 to "
            "90 day terms. An allowance for doubtful debts is made when there is "
            "objective evidence that a trade receivable is impaired.",
            space_before=4, space_after=6)

    # ==================================================================
    # NOTE 3: Property, Plant and Equipment
    # ==================================================================
    if has_ppe:
        _notes_spacer(doc, 8)
        n = _note_num_for("ppe")
        _notes_add_para(doc, f"Note {n}: Property, Plant and Equipment",
                        bold=True, space_before=18, space_after=6)

        tbl = _notes_create_table(doc, has_prior)

        # Render each cost/depreciation pair from TB proximity matching
        for pair in ppe_classes:
            cost_item = pair["cost"]
            depr_item = pair["depr"]

            if cost_item is None:
                continue  # orphan depreciation — skip

            cost_cy = abs(cost_item["cy_amount"]) if cost_item["cy_amount"] else Decimal("0")
            cost_py = abs(cost_item["py_amount"]) if cost_item["py_amount"] else Decimal("0")
            _notes_add_table_row(tbl, cost_item["account_name"],
                                 _fmt_note_amount(cost_cy), _fmt_note_amount(cost_py))

            depr_cy = Decimal("0")
            depr_py = Decimal("0")
            if depr_item:
                depr_cy = abs(depr_item["cy_amount"]) if depr_item["cy_amount"] else Decimal("0")
                depr_py = abs(depr_item["py_amount"]) if depr_item["py_amount"] else Decimal("0")
                # Use the actual account name from the TB
                _notes_add_table_row(tbl, depr_item["account_name"],
                                     _fmt_note_amount(-depr_cy),
                                     _fmt_note_amount(-depr_py),
                                     indent=Cm(0.5))

            net_cy = cost_cy - depr_cy
            net_py = cost_py - depr_py
            nbv_row = _notes_add_table_row(tbl, "Net book value",
                                           _fmt_note_amount(net_cy),
                                           _fmt_note_amount(net_py), bold=True)
            _notes_apply_subtotal_border(nbv_row)
            _notes_apply_grand_total_border(nbv_row)

            # Blank spacer row
            _notes_add_table_row(tbl, "", "", "")

        # Deposits (non-depreciable)
        for item in ppe_deposit:
            val_cy = abs(item["cy_amount"]) if item["cy_amount"] else Decimal("0")
            val_py = abs(item["py_amount"]) if item["py_amount"] else Decimal("0")
            _notes_add_table_row(tbl, item["account_name"],
                                 _fmt_note_amount(val_cy), _fmt_note_amount(val_py))

        _notes_add_para(
            doc,
            "All plant and equipment is stated at historical cost less depreciation. "
            "Depreciation is calculated on a diminishing value basis at rates determined "
            "by the Australian Taxation Office.",
            space_before=4, space_after=6)

    # ==================================================================
    # NOTE 4: Related Party Transactions
    # ==================================================================
    if has_related_party:
        _notes_spacer(doc, 8)
        n = _note_num_for("related_party")
        _notes_add_para(doc, f"Note {n}: Related Party Transactions",
                        bold=True, space_before=18, space_after=6)

        # Entity-type language
        if entity_type == "company":
            entity_leader = "director of the company"
        elif entity_type == "trust":
            entity_leader = "trustee of the trust"
        elif entity_type == "partnership":
            entity_leader = "partner in the partnership"
        else:
            entity_leader = "proprietor"

        sub_letter = ord("a")

        # Management Fees
        if has_mgmt_fees:
            for item in mgmt_fee_items:
                if item["cy_amount"] == 0 and item["py_amount"] == 0:
                    continue
                # Derive counterparty from account name
                raw_name = item["account_name"]
                counterparty = raw_name
                matched = False
                for prefix in ["Management fees - ", "Mgmt fee - ", "Management fee - ",
                                "Management fees- ", "Mgmt fees - "]:
                    if raw_name.lower().startswith(prefix.lower()):
                        counterparty = raw_name[len(prefix):].strip()
                        matched = True
                        break
                if not matched and " - " in raw_name:
                    counterparty = raw_name.split(" - ", 1)[1].strip()

                _notes_add_para(
                    doc,
                    f"({chr(sub_letter)}) Management Fees \u2014 {counterparty}",
                    bold=True, space_before=10, space_after=2)

                fee_cy = abs(item["cy_amount"]) if item["cy_amount"] else Decimal("0")
                fee_py = abs(item["py_amount"]) if item["py_amount"] else Decimal("0")
                _notes_add_para(
                    doc,
                    f"During the year {entity_name} was charged management fees by "
                    f"{counterparty}, a related party. Management fees charged during "
                    f"the year were {_fmt_dollar(fee_cy)} "
                    f"({prior_year_str}: {_fmt_dollar(fee_py)}).",
                    space_before=2, space_after=6)
                sub_letter += 1

        # Director Loans (skip for sole traders)
        if has_director_loan:
            for item in director_loan_items:
                if item["cy_amount"] == 0 and item["py_amount"] == 0:
                    continue
                # Em-dash in heading to match other sub-sections
                heading_name = item["account_name"].replace(" - ", " \u2014 ")
                _notes_add_para(
                    doc,
                    f"({chr(sub_letter)}) {heading_name}",
                    bold=True, space_before=10, space_after=2)

                # Balance sign: in NCL section, cy_amount = debit - credit
                # Negative = credit balance = entity owes director (liability)
                # Positive = debit balance = director owes entity (asset)
                bal_cy = abs(item["cy_amount"]) if item["cy_amount"] else Decimal("0")
                bal_py = abs(item["py_amount"]) if item["py_amount"] else Decimal("0")

                if item["cy_amount"] and item["cy_amount"] > 0:
                    # Debit balance — director owes entity
                    _notes_add_para(
                        doc,
                        f"The {entity_leader} has borrowed funds from {entity_name}. "
                        f"The amount outstanding at year end was {_fmt_dollar(bal_cy)} "
                        f"({prior_year_str}: {_fmt_dollar(bal_py)}).",
                        space_before=2, space_after=6)
                else:
                    _notes_add_para(
                        doc,
                        f"{entity_name} has a loan with a {entity_leader}. "
                        f"The balance outstanding at year end was "
                        f"{_fmt_dollar(bal_cy)} ({prior_year_str}: {_fmt_dollar(bal_py)}). "
                        f"The loan is unsecured, interest free and repayable on demand.",
                        space_before=2, space_after=6)
                sub_letter += 1

        # Related Entity Loans
        if has_related_loans:
            for item in related_loan_items:
                if item["cy_amount"] == 0 and item["py_amount"] == 0:
                    continue
                # Derive counterparty
                raw_name = item["account_name"]
                counterparty = raw_name
                if " - " in raw_name:
                    counterparty = raw_name.split(" - ", 1)[1].strip()

                _notes_add_para(
                    doc,
                    f"({chr(sub_letter)}) Loan \u2014 {counterparty}",
                    bold=True, space_before=10, space_after=2)

                bal_cy = abs(item["cy_amount"]) if item["cy_amount"] else Decimal("0")
                bal_py = abs(item["py_amount"]) if item["py_amount"] else Decimal("0")

                if item["cy_amount"] and item["cy_amount"] > 0:
                    # Debit = entity advanced funds (asset)
                    _notes_add_para(
                        doc,
                        f"{entity_name} has advanced funds to {counterparty}, a related "
                        f"party. The amount receivable at year end was "
                        f"{_fmt_dollar(bal_cy)} ({prior_year_str}: {_fmt_dollar(bal_py)}). "
                        f"The amount is unsecured, interest free and repayable on demand.",
                        space_before=2, space_after=6)
                else:
                    # Credit = entity owes (liability)
                    _notes_add_para(
                        doc,
                        f"{entity_name} has a loan with {counterparty}, a related party. "
                        f"The balance outstanding at year end was "
                        f"{_fmt_dollar(bal_cy)} ({prior_year_str}: {_fmt_dollar(bal_py)}). "
                        f"The loan is unsecured, interest free and repayable on demand.",
                        space_before=2, space_after=6)
                sub_letter += 1

    # ==================================================================
    # NOTE 5: Income Tax (companies only)
    # ==================================================================
    if has_income_tax:
        _notes_spacer(doc, 8)
        n = _note_num_for("income_tax")
        _notes_add_para(doc, f"Note {n}: Income Tax",
                        bold=True, space_before=18, space_after=6)

        _notes_add_para(doc, "The income tax expense for the year comprises:",
                        space_after=6)

        tbl = _notes_create_table(doc, has_prior)
        cte_row = _notes_add_table_row(tbl, "Current tax expense",
                                       _fmt_note_amount(income_tax_cy),
                                       _fmt_note_amount(income_tax_py))
        _notes_apply_subtotal_border(cte_row)
        total_row = _notes_add_table_row(tbl, "Income tax expense",
                                         _fmt_note_amount(income_tax_cy),
                                         _fmt_note_amount(income_tax_py), bold=True)
        _notes_apply_grand_total_border(total_row)

        # Tax rate
        rate_cy = 25 if abs(total_revenue_cy) < 50_000_000 else 30
        rate_py = 25 if abs(total_revenue_py) < 50_000_000 else 30

        _notes_add_para(
            doc,
            f"The income tax provision has been calculated at the applicable corporate "
            f"tax rate of {rate_cy}% on the estimated taxable profit for the year.",
            space_before=4, space_after=6)

        _notes_add_para(
            doc,
            f"The applicable tax rate is {rate_cy}% ({prior_year_str}: {rate_py}%) "
            f"being the corporate tax rate for base rate entities.",
            space_after=6)

    # ==================================================================
    # NOTE 6: Events After the Reporting Date (companies only)
    # ==================================================================
    if is_company:
        _notes_spacer(doc, 8)
        n = _note_num_for("events")
        _notes_add_para(doc, f"Note {n}: Events After the Reporting Date",
                        bold=True, space_before=18, space_after=6)

        _notes_add_para(
            doc,
            "The directors are not aware of any matter or circumstance that has "
            "arisen since the end of the financial year that has significantly "
            "affected or may significantly affect the operations of the entity, "
            "the results of those operations, or the state of affairs of the "
            "entity in future years.",
            space_before=4, space_after=6)

    # ==================================================================
    # POST-PROCESSING: Keep each note together on a single page.
    # Scan the document body, identify note boundaries, estimate heights,
    # and insert explicit page breaks where a note would overflow.
    # Also apply keepNext + keepLines on every element within each note.
    # ==================================================================
    import re as _re

    body = doc.element.body
    all_elements = list(body)

    # Identify note boundary indices — each "Note N:" heading starts a note
    note_heading_pattern = _re.compile(r'^Note \d+:')
    note_start_indices = []
    for idx, el in enumerate(all_elements):
        if el.tag == qn('w:p'):
            text = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
            if note_heading_pattern.match(text):
                note_start_indices.append(idx)

    # Build list of (start_idx, end_idx) for each note
    note_ranges = []
    for i, start in enumerate(note_start_indices):
        # Include the spacer paragraph immediately before the heading (if any)
        actual_start = start
        if start > 0:
            prev_el = all_elements[start - 1]
            if prev_el.tag == qn('w:p'):
                prev_text = ''.join(t.text or '' for t in prev_el.iter(qn('w:t'))).strip()
                if prev_text == '':
                    actual_start = start - 1

        if i + 1 < len(note_start_indices):
            # End at the element before the next note's spacer/heading
            next_start = note_start_indices[i + 1]
            # Check if the element before the next heading is a spacer
            if next_start > 0:
                prev_next = all_elements[next_start - 1]
                if prev_next.tag == qn('w:p'):
                    prev_text = ''.join(t.text or '' for t in prev_next.iter(qn('w:t'))).strip()
                    if prev_text == '':
                        end = next_start - 1
                    else:
                        end = next_start
                else:
                    end = next_start
            else:
                end = next_start
        else:
            end = len(all_elements)
        note_ranges.append((actual_start, end))

    def _estimate_element_height(el):
        """Estimate height of a body element in cm.

        Deliberately conservative (overestimates) so that page breaks
        are inserted early rather than late — a note that starts on a
        fresh page with room to spare is better than one that splits.
        """
        if el.tag == qn('w:tbl'):
            rows = el.findall(qn('w:tr'))
            # 0.55cm per row accounts for cell padding and borders
            return len(rows) * 0.55
        elif el.tag == qn('w:p'):
            text = ''.join(t.text or '' for t in el.iter(qn('w:t'))).strip()
            if not text:
                return 0.4  # spacer paragraph
            # Page text width ~16cm, but Calibri 10pt renders ~60 chars
            # per line at that width.  Add extra for paragraph spacing.
            lines = max(1, len(text) / 60)
            return lines * 0.45 + 0.3  # +0.3cm for space_before/after
        return 0.6

    def _set_keep_on_paragraph_element(p_el):
        """Set keepNext and keepLines on a w:p element."""
        pPr = p_el.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p_el.insert(0, pPr)
        # keepLines
        if pPr.find(qn('w:keepLines')) is None:
            kl = OxmlElement('w:keepLines')
            kl.set(qn('w:val'), '1')
            pPr.append(kl)
        # keepNext
        if pPr.find(qn('w:keepNext')) is None:
            kn = OxmlElement('w:keepNext')
            kn.set(qn('w:val'), '1')
            pPr.append(kn)

    def _remove_keep_next(p_el):
        """Remove keepNext from the last element so it can break after."""
        pPr = p_el.find(qn('w:pPr'))
        if pPr is not None:
            for kn in pPr.findall(qn('w:keepNext')):
                pPr.remove(kn)

    # Apply keep-together to each note block
    for start, end in note_ranges:
        note_elements = all_elements[start:end]
        for i_el, el in enumerate(note_elements):
            is_last = (i_el == len(note_elements) - 1)
            if el.tag == qn('w:p'):
                _set_keep_on_paragraph_element(el)
                if is_last:
                    _remove_keep_next(el)
            elif el.tag == qn('w:tbl'):
                # Apply to every cell paragraph in every row
                for tr in el.findall(qn('w:tr')):
                    for tc in tr.findall(qn('w:tc')):
                        for p in tc.findall(qn('w:p')):
                            _set_keep_on_paragraph_element(p)
                # Remove keepNext from last row's paragraphs if this is the last element
                if is_last:
                    last_tr = el.findall(qn('w:tr'))
                    if last_tr:
                        for tc in last_tr[-1].findall(qn('w:tc')):
                            for p in tc.findall(qn('w:p')):
                                _remove_keep_next(p)

    # Height-based page break insertion.
    # A4 = 29.7cm, margins top+bottom = 4cm, header ~2.5cm, footer ~1.5cm
    # → usable body height ≈ 21.7cm.  Use 21cm as a conservative limit.
    USABLE_HEIGHT_CM = 21.0
    page_pos = 0.0

    for i, (start, end) in enumerate(note_ranges):
        note_height = sum(
            _estimate_element_height(all_elements[j])
            for j in range(start, end)
        )

        if page_pos > 0 and (page_pos + note_height) > USABLE_HEIGHT_CM:
            # Insert page break before this note's first element
            target_el = all_elements[start]
            br_para = OxmlElement('w:p')
            br_pPr = OxmlElement('w:pPr')
            # Minimal spacing on the break paragraph
            sp_before = OxmlElement('w:spacing')
            sp_before.set(qn('w:before'), '0')
            sp_before.set(qn('w:after'), '0')
            br_pPr.append(sp_before)
            br_para.append(br_pPr)
            br_r = OxmlElement('w:r')
            br_el = OxmlElement('w:br')
            br_el.set(qn('w:type'), 'page')
            br_r.append(br_el)
            br_para.append(br_r)
            body.insert(list(body).index(target_el), br_para)
            page_pos = note_height
        else:
            page_pos += note_height

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# 6b. _generate_depreciation_report
# ---------------------------------------------------------------------------
def _generate_depreciation_report(context):
    """Build the Depreciation Report as a python-docx Document (landscape).

    Assets are grouped by category with subtotals per category.
    Columns: Asset | Total Cost | Priv% | OWDV | Disp Date | Disp Consid |
             Add Date | Add Cost | Dep Value | T | Rate | Deprec | Priv | CWDV

    Returns a BytesIO buffer containing the .docx, or None if no assets exist.
    """
    from docx import Document
    from docx.shared import Cm, Pt, Emu, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from collections import OrderedDict
    from decimal import Decimal
    from core.models import DepreciationAsset

    fy = context["_fy"]
    entity = context["_entity"]

    assets_qs = DepreciationAsset.objects.filter(
        financial_year=fy
    ).order_by("category", "display_order", "asset_name", "id")

    if not assets_qs.exists():
        return None

    # Deduplicate by (category, asset_name) — first occurrence wins.
    # Safeguards against duplicate DepreciationAsset records created by
    # repeated imports/backfills (bulk_create without existence check).
    seen = set()
    assets = []
    for a in assets_qs:
        key = (a.category, a.asset_name)
        if key in seen:
            continue
        seen.add(key)
        assets.append(a)

    # Group by category
    categories = OrderedDict()
    for asset in assets:
        categories.setdefault(asset.category, []).append(asset)

    def _fmt(val):
        if val is None or val == 0:
            return ""
        return f"{val:,.0f}"

    def _fmt_rate(val):
        if val is None or val == 0:
            return "0.00"
        return f"{val:.2f}"

    def _fmt_date(d):
        if d is None:
            return ""
        return d.strftime("%d/%m/%y")

    FONT = "Times New Roman"
    FONT_SZ = Pt(7)
    FONT_SZ_HDR = Pt(7)
    FONT_SZ_TITLE = Pt(9)
    HEADER_BG = "D9E2F3"

    # Column widths (cm) — 14 columns to fit A4 landscape
    COL_W = [5.0, 1.7, 0.9, 1.7, 1.4, 1.4, 1.4, 1.4, 1.4, 0.5, 1.1, 1.7, 1.1, 1.7]
    COL_HEADERS = [
        "Asset", "Total\nCost", "Priv\n%", "OWDV",
        "Disp\nDate", "Disp\nConsid",
        "Add\nDate", "Add\nCost",
        "Dep\nValue", "T", "Rate", "Deprec", "Priv", "CWDV",
    ]

    def _shade_cell(cell, hex_color):
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    def _apply_top_border(row):
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '6')
            top.set(qn('w:space'), '0')
            top.set(qn('w:color'), '000000')
            tcBorders.append(top)

    doc = Document()

    # Landscape A4 setup
    from docx.shared import Cm as _Cm
    for section in doc.sections:
        section.page_width = _Cm(29.7)
        section.page_height = _Cm(21.0)
        section.top_margin = _Cm(1.6)
        section.bottom_margin = _Cm(1.7)
        section.left_margin = _Cm(2.0)
        section.right_margin = _Cm(2.4)
        section.orientation = 1  # WD_ORIENT.LANDSCAPE

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = FONT_SZ

    # Repeating header
    header = doc.sections[0].header
    header.is_linked_to_previous = False
    for para in list(header.paragraphs):
        para.clear()
    _entity_name = context.get("entity_name", entity.entity_name)
    _abn = context.get("abn", "")
    _date_text = context.get("date_text", "")
    _watermark = context.get("watermark", "")

    p1 = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(_entity_name)
    r1.font.name = FONT; r1.font.size = Pt(10); r1.bold = True
    p1.paragraph_format.space_after = Pt(0)
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"ABN {_abn}" if _abn else "")
    r2.font.name = FONT; r2.font.size = Pt(8)
    p2.paragraph_format.space_after = Pt(0)
    p3 = header.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Depreciation Report")
    r3.font.name = FONT; r3.font.size = Pt(8)
    p3.paragraph_format.space_after = Pt(0)
    p4 = header.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(_date_text)
    r4.font.name = FONT; r4.font.size = Pt(8)
    # space_after = ~1cm (284,200 EMUs) — gap between header text and rule
    p4.paragraph_format.space_after = Emu(284200)
    pPr = p4._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bb = OxmlElement('w:bottom')
    bb.set(qn('w:val'), 'single'); bb.set(qn('w:sz'), '6')
    bb.set(qn('w:space'), '1'); bb.set(qn('w:color'), '000000')
    pBdr.append(bb); pPr.append(pBdr)
    if _watermark:
        pw = header.add_paragraph()
        pw.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rw = pw.add_run(_watermark)
        rw.font.name = FONT; rw.font.size = Pt(12)
        rw.font.color.rgb = RGBColor(0xFF, 0x00, 0x00); rw.bold = True
        pw.paragraph_format.space_after = Pt(0)

    # Footer — centred Times New Roman italic 9pt with 0.5pt top border
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        "These financial statements are unaudited. They must be read in conjunction with the "
        "attached Accountant\u2019s Compilation Report and Notes which form part of these financial statements."
    )
    fr.font.name = "Times New Roman"
    fr.font.size = Pt(9)
    fr.font.italic = True
    _fpPr = fp._p.get_or_add_pPr()
    for _existing in _fpPr.findall(qn('w:pBdr')):
        _fpPr.remove(_existing)
    _fpBdr = OxmlElement('w:pBdr')
    _fpTop = OxmlElement('w:top')
    _fpTop.set(qn('w:val'), 'single')
    _fpTop.set(qn('w:sz'), '4')
    _fpTop.set(qn('w:space'), '4')
    _fpTop.set(qn('w:color'), '000000')
    _fpBdr.append(_fpTop)
    _fpPr.append(_fpBdr)

    # Grand totals accumulators
    grand_cost = Decimal("0")
    grand_owdv = Decimal("0")
    grand_deprec = Decimal("0")
    grand_priv = Decimal("0")
    grand_cwdv = Decimal("0")
    grand_add = Decimal("0")
    grand_disp = Decimal("0")

    first_category = True
    for cat_name, cat_assets in categories.items():
        if not first_category:
            doc.add_page_break()
        else:
            # 0.5cm gap between header rule and first body content (first page only)
            _p0 = doc.add_paragraph()
            _p0.paragraph_format.space_before = Emu(142100)
            _p0.paragraph_format.space_after = Pt(0)
        first_category = False

        # Category heading
        ph = doc.add_paragraph()
        rh = ph.add_run(cat_name)
        rh.font.name = FONT; rh.font.size = Pt(9); rh.bold = True; rh.underline = True
        ph.paragraph_format.space_after = Pt(2)

        # Table
        table = doc.add_table(rows=1, cols=len(COL_HEADERS))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Fix column widths
        for row in table.rows:
            for i, w in enumerate(COL_W):
                row.cells[i].width = Cm(w)

        # Header row
        hdr = table.rows[0].cells
        for i, hdr_text in enumerate(COL_HEADERS):
            cell = hdr[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hdr_text)
            run.font.size = FONT_SZ_HDR; run.font.name = FONT; run.bold = True
            _shade_cell(cell, HEADER_BG)

        # Category totals
        cat_cost = Decimal("0")
        cat_owdv = Decimal("0")
        cat_deprec = Decimal("0")
        cat_priv = Decimal("0")
        cat_cwdv = Decimal("0")
        cat_add = Decimal("0")
        cat_disp = Decimal("0")

        for asset in cat_assets:
            row = table.add_row()
            for i, w in enumerate(COL_W):
                row.cells[i].width = Cm(w)
            vals = [
                asset.asset_name,
                _fmt(asset.total_cost),
                f"{asset.private_use_pct:.2f}" if asset.private_use_pct else "",
                _fmt(asset.opening_wdv),
                _fmt_date(asset.disposal_date),
                _fmt(asset.disposal_consideration),
                _fmt_date(asset.addition_date),
                _fmt(asset.addition_cost),
                _fmt(asset.depreciable_value),
                asset.get_method_display()[0] if asset.method else "",
                _fmt_rate(asset.rate),
                _fmt(asset.depreciation_amount),
                _fmt(asset.private_depreciation),
                _fmt(asset.closing_wdv),
            ]
            for i, val in enumerate(vals):
                cell = row.cells[i]
                cell.width = Cm(COL_W[i])
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(str(val))
                run.font.size = FONT_SZ; run.font.name = FONT

            cat_cost += asset.total_cost or Decimal("0")
            cat_owdv += asset.opening_wdv or Decimal("0")
            cat_deprec += asset.depreciation_amount or Decimal("0")
            cat_priv += asset.private_depreciation or Decimal("0")
            cat_cwdv += asset.closing_wdv or Decimal("0")
            cat_add += asset.addition_cost or Decimal("0")
            cat_disp += asset.disposal_consideration or Decimal("0")

        # Subtotals row
        sub_row = table.add_row()
        for i, w in enumerate(COL_W):
            sub_row.cells[i].width = Cm(w)
        sub_vals = [
            "Subtotals",
            _fmt(cat_cost), "", _fmt(cat_owdv),
            "", _fmt(cat_disp),
            "", _fmt(cat_add),
            "", "", "",
            _fmt(cat_deprec), _fmt(cat_priv), _fmt(cat_cwdv),
        ]
        for i, val in enumerate(sub_vals):
            cell = sub_row.cells[i]
            cell.width = Cm(COL_W[i])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = FONT_SZ; run.font.name = FONT; run.bold = True
        _apply_top_border(sub_row)

        # Net depreciation line
        net_dep = cat_deprec - cat_priv
        pn = doc.add_paragraph()
        rn = pn.add_run(f"Deduct Private Portion: {_fmt(cat_priv)}     Net Depreciation: {_fmt(net_dep)}")
        rn.font.size = Pt(8); rn.font.name = FONT; rn.bold = True; rn.underline = True
        pn.paragraph_format.space_after = Pt(4)

        # Accumulate grand totals
        grand_cost += cat_cost; grand_owdv += cat_owdv
        grand_deprec += cat_deprec; grand_priv += cat_priv
        grand_cwdv += cat_cwdv; grand_add += cat_add; grand_disp += cat_disp

    # Grand totals section
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    pg = doc.add_paragraph()
    rg = pg.add_run("Grand Totals")
    rg.font.name = FONT; rg.font.size = Pt(9); rg.bold = True
    pg.paragraph_format.space_after = Pt(2)

    gt_table = doc.add_table(rows=1, cols=len(COL_HEADERS))
    gt_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gt_table.autofit = False
    for row in gt_table.rows:
        for i, w in enumerate(COL_W):
            row.cells[i].width = Cm(w)
    gt_vals = [
        "Total All Categories",
        _fmt(grand_cost), "", _fmt(grand_owdv),
        "", _fmt(grand_disp),
        "", _fmt(grand_add),
        "", "", "",
        _fmt(grand_deprec), _fmt(grand_priv), _fmt(grand_cwdv),
    ]
    gt_row = gt_table.rows[0]
    for i, val in enumerate(gt_vals):
        cell = gt_row.cells[i]
        cell.width = Cm(COL_W[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(val))
        run.font.size = FONT_SZ; run.font.name = FONT; run.bold = True
    _apply_top_border(gt_row)

    # Net depreciation grand total
    grand_net = grand_deprec - grand_priv
    pgn = doc.add_paragraph()
    rgn = pgn.add_run(
        f"Total Private Portion: {_fmt(grand_priv)}     Total Net Depreciation: {_fmt(grand_net)}"
    )
    rgn.font.size = Pt(9); rgn.font.name = FONT; rgn.bold = True; rgn.underline = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# 7. generate_financial_statements
# ---------------------------------------------------------------------------
DOCUMENT_TYPE_ORDER = [
    "COVER",
    "DETAILED_PL",
    "BALANCE_SHEET",
    "SUMMARY_PL",
    "DEPRECIATION_REPORT",
    "NOTES",
    "DECLARATION",
    "DISTRIBUTION",
    "COMPILATION",
]


def generate_financial_statements(financial_year_id, include_watermark=True):
    """Orchestrate all templates for a financial year.

    Returns dict of document_type → BytesIO.
    """
    from core.models import FinancialStatementTemplate, FinancialYear

    fy = FinancialYear.objects.select_related(
        "entity", "entity__client", "prior_year",
    ).get(pk=financial_year_id)
    entity = fy.entity
    entity_type = entity.entity_type

    # Build context based on entity type
    context_builders = {
        "company": build_company_context,
        "trust": build_trust_context,
        "sole_trader": build_sole_trader_context,
    }
    builder = context_builders.get(entity_type, build_company_context)
    context = builder(fy, include_watermark=include_watermark)

    # ── DocumentContextBuilder enrichment ────────────────────────────────────
    # Merge the new practice_* namespace and Jinja2 filter keys on top of the
    # legacy context without replacing any existing keys.  This allows existing
    # templates to keep working while new templates can use the richer spec.
    try:
        from core.document_context_builder import DocumentContextBuilder
        dcb = DocumentContextBuilder(entity, financial_year=fy)
        enriched = dcb.build("financial_statements")
        # Only inject keys that are NOT already present (legacy keys take priority)
        for k, v in enriched.items():
            if k not in context:
                context[k] = v
        # Always override practice_* keys (new namespace — no conflict risk)
        for k, v in enriched.items():
            if k.startswith("practice_") or k in (
                "signatory_name", "signatory_designation",
                "tax_agent_number", "bas_agent_number", "asic_agent_number",
                "professional_body", "membership_number",
                "practice_independence_maintained",
                "basis_of_preparation", "basis_of_preparation_note",
                "is_going_concern", "has_events_after_balance_date",
                "compilation_report_name",
            ):
                context[k] = v
        logger.debug("DocumentContextBuilder enrichment applied for FY %s", fy.pk)
    except Exception as _dcb_exc:
        logger.warning(
            "DocumentContextBuilder enrichment failed for FY %s — "
            "falling back to legacy context: %s", fy.pk, _dcb_exc
        )

    # Get active templates for this entity type
    templates = FinancialStatementTemplate.objects.filter(
        entity_type=entity_type,
        is_active=True,
    )

    # Skip SUMMARY_PL for non-company entities
    # Skip DISTRIBUTION for non-trust entities
    skip_types = set()
    if entity_type != "company":
        skip_types.add("SUMMARY_PL")
    if entity_type != "trust":
        skip_types.add("DISTRIBUTION")

    has_prior = context.get("has_prior", False)

    results = {}
    for doc_type in DOCUMENT_TYPE_ORDER:
        if doc_type in skip_types:
            continue

        # DEPRECIATION_REPORT: programmatic generation from DepreciationAsset records
        if doc_type == "DEPRECIATION_REPORT":
            try:
                buffer = _generate_depreciation_report(context)
                if buffer is None:
                    logger.info(
                        "No depreciation assets for FY %s — skipping DEPRECIATION_REPORT",
                        fy.pk,
                    )
                else:
                    buffer = _post_process_fs_doc(buffer, doc_type, has_prior=has_prior)
                    results[doc_type] = buffer
                    logger.info("Generated programmatic DEPRECIATION_REPORT for FY %s", fy.pk)
            except Exception as e:
                logger.error(
                    "Failed to generate DEPRECIATION_REPORT for FY %s: %s", fy.pk, e
                )
            continue

        # NOTES: use programmatic generation instead of static template
        if doc_type == "NOTES":
            try:
                buffer = _generate_notes_document(context)
                buffer = _post_process_fs_doc(buffer, doc_type, has_prior=has_prior)
                results[doc_type] = buffer
                logger.info("Generated programmatic NOTES for FY %s", fy.pk)
            except Exception as e:
                logger.error(
                    "Failed to generate NOTES for FY %s: %s", fy.pk, e
                )
            continue

        # DISTRIBUTION: programmatic generation — docxtpl strips column widths
        if doc_type == "DISTRIBUTION":
            try:
                buffer = _build_distribution_docx(fy, context)
                # DISTRIBUTION is already PDF (reportlab) — skip docx post-processing
                results[doc_type] = buffer
                logger.info("Generated programmatic DISTRIBUTION for FY %s", fy.pk)
            except Exception as e:
                logger.error(
                    "Failed to generate DISTRIBUTION for FY %s: %s", fy.pk, e
                )
            continue

        tmpl = templates.filter(document_type=doc_type).first()
        if not tmpl:
            logger.warning(
                "No active template for %s/%s — skipping",
                doc_type, entity_type,
            )
            continue

        try:
            buffer = render_template(tmpl, context)
            buffer = _post_process_fs_doc(buffer, doc_type, has_prior=has_prior)
            results[doc_type] = buffer
            logger.info("Rendered %s for FY %s", doc_type, fy.pk)
        except Exception as e:
            logger.error(
                "Failed to render %s for FY %s: %s", doc_type, fy.pk, e
            )

    return results


# ---------------------------------------------------------------------------
# Page-number stamping — absolute numbering on the final merged PDF
# ---------------------------------------------------------------------------
def _stamp_page_numbers(pdf_bytes):
    """No-op passthrough — page numbers are intentionally suppressed.

    Previously stamped centred page numbers at the bottom of every page
    via a reportlab overlay. Removed per client requirement: no page
    numbers anywhere in generated FS / package outputs. The call sites
    (generate_combined_pdf, build_package_bundle) are retained so the
    function can be re-enabled by restoring the original body.
    """
    return pdf_bytes


def _build_distribution_docx(financial_year, context):
    """Generate Distribution Summary as PDF directly via reportlab.

    Returns BytesIO containing PDF bytes (NOT docx).
    python-docx cannot reliably produce 3-column tables that survive
    LibreOffice conversion on this server.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    import io as _io

    entity = financial_year.entity
    end_date = financial_year.end_date
    date_text = f"For the Year Ended {end_date.strftime('%d %B %Y')}" if end_date else ""
    beneficiaries = context.get("beneficiaries", [])
    total_distribution = context.get("total_distribution", "—")

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.0 * cm, rightMargin=2.4 * cm,
        topMargin=1.6 * cm, bottomMargin=1.7 * cm,
    )

    styles = getSampleStyleSheet()
    # Helvetica = Palatino substitute (reportlab built-in)
    # Times-Roman = Times New Roman substitute
    centre = ParagraphStyle(
        "centre", parent=styles["Normal"],
        alignment=TA_CENTER, fontName="Helvetica", fontSize=11)
    centre_sm = ParagraphStyle(
        "centre_sm", parent=styles["Normal"],
        alignment=TA_CENTER, fontName="Helvetica", fontSize=11)
    centre_bold = ParagraphStyle(
        "centre_bold", parent=styles["Normal"],
        alignment=TA_CENTER, fontName="Helvetica-Bold", fontSize=13)
    normal = ParagraphStyle(
        "norm", parent=styles["Normal"],
        fontName="Times-Roman", fontSize=10)
    bold_style = ParagraphStyle(
        "bold", parent=styles["Normal"],
        fontName="Times-Bold", fontSize=10)
    small_italic = ParagraphStyle(
        "si", parent=styles["Normal"],
        fontName="Helvetica-Oblique", fontSize=9)

    abn_raw = entity.abn or ""
    abn_digits = "".join(c for c in str(abn_raw) if c.isdigit())
    abn_fmt = (
        f"{abn_digits[:2]} {abn_digits[2:5]} {abn_digits[5:8]} {abn_digits[8:]}"
        if len(abn_digits) == 11 else abn_raw
    )

    story = []
    story.append(Paragraph(_safe_amp(entity.entity_name), centre_bold))
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph(f"ABN {abn_fmt}", centre_sm))
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph("Beneficiaries Distribution Summary", centre_sm))
    story.append(Spacer(1, 2 * mm))
    story.append(Paragraph(date_text, centre_sm))
    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph(
        f"Net Income Available for Distribution: {total_distribution}",
        bold_style))
    story.append(Spacer(1, 6 * mm))

    # Table data
    page_width = A4[0] - 2.5 * cm - 2 * cm
    col_widths = [page_width * 0.60, page_width * 0.20, page_width * 0.20]

    table_data = [["Beneficiary", "Percentage", "Amount\n$"]]
    for b in beneficiaries:
        table_data.append([
            b.get("beneficiary_name", ""),
            f"{b.get('percentage', '')}%",
            b.get("amount", ""),
        ])
    table_data.append(["Total", "100%", total_distribution])

    tbl = Table(table_data, colWidths=col_widths)
    tbl.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTNAME", (0, 1), (-1, -2), "Times-Roman"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("FONTNAME", (0, -1), (-1, -1), "Times-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("ALIGN", (0, 0), (0, -1), "LEFT"),
        ("LINEBELOW", (0, 0), (-1, 0), 0.5, colors.black),
        ("LINEABOVE", (0, -1), (-1, -1), 0.5, colors.black),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(tbl)

    # Per-page footer — centred Times Italic 9pt with 0.5pt horizontal rule
    # above, matching the Handiledger-aligned docx footer.
    FOOTER_TEXT = (
        "These financial statements are unaudited. They must be read "
        "in conjunction with the attached Accountant\u2019s Compilation "
        "Report and Notes which form part of these financial statements."
    )

    def _draw_footer(canvas, _doc):
        canvas.saveState()
        page_w = A4[0]
        left_x = 2.0 * cm
        right_x = page_w - 2.4 * cm
        # Horizontal rule just above the footer text
        canvas.setLineWidth(0.5)
        canvas.setStrokeColorRGB(0, 0, 0)
        canvas.line(left_x, 1.3 * cm, right_x, 1.3 * cm)
        # Footer text (centred, Times Italic 9pt)
        canvas.setFont("Times-Italic", 9)
        canvas.setFillColorRGB(0, 0, 0)
        # Wrap to two lines if wider than usable area
        usable_w = right_x - left_x
        text_w = canvas.stringWidth(FOOTER_TEXT, "Times-Italic", 9)
        if text_w <= usable_w:
            canvas.drawCentredString(page_w / 2, 0.9 * cm, FOOTER_TEXT)
        else:
            # Split at the roughly-midway "the" word for a balanced break
            split = FOOTER_TEXT.rfind("conjunction with")
            line1 = FOOTER_TEXT[:split].strip()
            line2 = FOOTER_TEXT[split:].strip()
            canvas.drawCentredString(page_w / 2, 1.0 * cm, line1)
            canvas.drawCentredString(page_w / 2, 0.65 * cm, line2)
        canvas.restoreState()

    doc.build(story, onFirstPage=_draw_footer, onLaterPages=_draw_footer)
    buf.seek(0)
    logger.info(
        "Built programmatic DISTRIBUTION (PDF) for %s: %d beneficiaries",
        entity.entity_name, len(beneficiaries),
    )
    return buf


# _reapply_distribution_widths — kept for reference but no longer called.
# Programmatic generation via _build_distribution_docx bypasses docxtpl entirely.
def _reapply_distribution_widths(docx_path):
    """Re-apply explicit column widths to Distribution Summary table.

    docxtpl strips w:tcW attributes during Jinja2 XML processing, causing
    LibreOffice to collapse the amount column during PDF conversion.
    Must be called on the rendered docx before LibreOffice conversion.
    """
    from docx import Document as _DocxDoc
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlEl

    col_widths = [5580, 1860, 1920]  # dxa (twips)
    doc = _DocxDoc(docx_path)
    if not doc.tables:
        return
    table = doc.tables[0]
    for row in table.rows:
        for k, cell in enumerate(row.cells):
            if k >= len(col_widths):
                continue
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for existing in tcPr.findall(_qn('w:tcW')):
                tcPr.remove(existing)
            tcW = _OxmlEl('w:tcW')
            tcW.set(_qn('w:w'), str(col_widths[k]))
            tcW.set(_qn('w:type'), 'dxa')
            tcPr.append(tcW)
    doc.save(docx_path)
    logger.info("Re-applied column widths to DISTRIBUTION docx: %s", docx_path)


# ---------------------------------------------------------------------------
# generate_combined_pdf — render each template to PDF individually, merge
# ---------------------------------------------------------------------------
def generate_combined_pdf(financial_year_id, include_watermark=True, exclude_types=None):
    """Generate all templates, convert each to PDF, merge into single PDF BytesIO.

    Args:
        exclude_types: optional set of document type keys to skip (e.g. {"DECLARATION"})

    Returns a BytesIO containing the merged PDF bytes.
    """
    from pypdf import PdfWriter, PdfReader

    logger.info("generate_combined_pdf called for FY %s (watermark=%s)",
                financial_year_id, include_watermark)

    docs = generate_financial_statements(financial_year_id, include_watermark)

    logger.info("generate_financial_statements returned %d documents: %s",
                len(docs), list(docs.keys()))

    if not docs:
        raise RuntimeError("No templates rendered — check template registration")

    excluded = exclude_types or set()
    ordered_keys = [dt for dt in DOCUMENT_TYPE_ORDER if dt in docs and dt not in excluded]
    logger.info("Ordered keys for PDF merge: %s", ordered_keys)

    if not ordered_keys:
        raise RuntimeError("No templates rendered")

    writer = PdfWriter()
    tmpdir = tempfile.mkdtemp(prefix="shub_combined_pdf_")
    pdfs_merged = 0

    try:
        for doc_type in ordered_keys:
            buffer = docs[doc_type]

            # DISTRIBUTION is already PDF (generated by reportlab) — merge directly
            if doc_type == "DISTRIBUTION":
                buffer.seek(0)
                dist_reader = PdfReader(buffer)
                for page in dist_reader.pages:
                    writer.add_page(page)
                pdfs_merged += 1
                logger.info("Merged DISTRIBUTION PDF directly (reportlab) (%d pages)",
                            len(dist_reader.pages))
                continue

            docx_path = os.path.join(tmpdir, f"{doc_type}.docx")
            with open(docx_path, "wb") as f:
                f.write(buffer.read())

            try:
                convert_docx_to_pdf(docx_path, tmpdir, timeout=60)
            except RuntimeError:
                logger.error("LibreOffice not available — skipping %s", doc_type)
                continue

            pdf_path = os.path.join(tmpdir, f"{doc_type}.pdf")
            if os.path.exists(pdf_path):
                reader = PdfReader(pdf_path)
                for page in reader.pages:
                    writer.add_page(page)
                pdfs_merged += 1
                logger.info("Merged %s into combined PDF (%d pages)", doc_type, len(reader.pages))
            else:
                logger.warning("PDF conversion produced no output for %s", doc_type)

        if pdfs_merged == 0:
            raise RuntimeError("No templates could be converted to PDF")

        output = io.BytesIO()
        writer.write(output)
        raw_bytes = output.getvalue()

        # Stamp continuous page numbers on the merged PDF
        stamped = _stamp_page_numbers(raw_bytes)

        result = io.BytesIO(stamped)
        result.seek(0)
        logger.info("generate_combined_pdf complete: %d documents, %d bytes",
                    pdfs_merged, len(stamped))
        return result
    finally:
        import shutil
        shutil.rmtree(tmpdir, ignore_errors=True)


# ---------------------------------------------------------------------------
# 8. assemble_pdf_package
# ---------------------------------------------------------------------------
def assemble_pdf_package(financial_year_id):
    """Generate all docs with include_watermark=False, convert to PDF, merge.

    Returns bytes of the merged PDF.
    """
    docs = generate_financial_statements(
        financial_year_id, include_watermark=False,
    )

    if not docs:
        logger.warning("No documents generated for FY %s", financial_year_id)
        return None

    try:
        from PyPDF2 import PdfMerger
    except ImportError:
        logger.error("PyPDF2 not available — cannot merge PDFs")
        return None

    merger = PdfMerger()
    tmpdir = tempfile.mkdtemp(prefix="shub_fs_pkg_")
    pdfs_added = 0

    for doc_type in DOCUMENT_TYPE_ORDER:
        if doc_type not in docs:
            continue

        buffer = docs[doc_type]
        docx_path = os.path.join(tmpdir, f"{doc_type}.docx")
        with open(docx_path, "wb") as f:
            f.write(buffer.read())

        # Convert to PDF via LibreOffice
        pdf_path = os.path.join(tmpdir, f"{doc_type}.pdf")
        try:
            convert_docx_to_pdf(docx_path, tmpdir, timeout=60)
        except RuntimeError:
            logger.error("LibreOffice not available — skipping PDF conversion for %s", doc_type)
            continue

        if os.path.exists(pdf_path):
            try:
                merger.append(pdf_path)
                pdfs_added += 1
            except Exception as e:
                logger.error("Failed to append PDF %s: %s", doc_type, e)
        else:
            logger.warning("PDF conversion produced no output for %s", doc_type)

    if pdfs_added == 0:
        merger.close()
        return None

    output = io.BytesIO()
    merger.write(output)
    merger.close()
    output.seek(0)

    logger.info(
        "Assembled PDF package for FY %s: %d documents",
        financial_year_id, pdfs_added,
    )
    return output.read()
