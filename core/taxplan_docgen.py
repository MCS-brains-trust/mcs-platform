"""
Trust Tax Planning — Document Generation

Generates:
1. Trust Election (s97/streaming) — Word document
2. Tax Planning Summary (client-facing) — Word document

Uses python-docx to build documents programmatically with merge fields.

Merge fields supported:
  {{trust_name}}                  — Entity name
  {{trustee_name}}                — Trustee(s) names
  {{chairperson_name}}            — Chairperson name
  {{resolution_date}}             — 30 June YYYY
  {{financial_year_end}}          — 30 June YYYY
  {{distributable_income}}        — Formatted dollar amount
  {{beneficiary_table}}           — Block-level table of distributions
  {{total_tax_payable}}           — Formatted dollar amount
  {{scenario_name}}               — Active scenario name (if any)
  {{accountant_recommendation}}   — From Section 5 notes
  {{streaming_election_details}}  — Block-level streaming details
"""
import io
from decimal import Decimal
from datetime import date

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# =============================================================================
# Constants
# =============================================================================
FONT_NAME = "Times New Roman"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_HEADING = Pt(14)
FONT_SIZE_SUBHEADING = Pt(12)
FONT_SIZE_SMALL = Pt(9)


def _set_run_font(run, size=FONT_SIZE_BODY, bold=False, italic=False, color=None):
    """Apply consistent font styling to a run."""
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_heading(doc, text, level=1):
    """Add a styled heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        _set_run_font(run, FONT_SIZE_HEADING, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        _set_run_font(run, FONT_SIZE_SUBHEADING, bold=True)
    else:
        _set_run_font(run, FONT_SIZE_BODY, bold=True)
    return p


def _add_para(doc, text, bold=False, italic=False, alignment=None):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, bold=bold, italic=italic)
    if alignment:
        p.alignment = alignment
    return p


def _fmt_money(amount):
    """Format a Decimal as $X,XXX.XX."""
    if amount is None:
        return "$0.00"
    return f"${amount:,.2f}"


def _officer_has_role(officer, role_value):
    """Check if an officer has a specific role."""
    if officer.roles:
        return role_value in officer.roles
    return officer.role == role_value


def _get_trust_context(financial_year_id):
    """
    Gather all merge field data for a trust financial year.
    Returns a dict of merge field values.
    """
    from core.models import (
        FinancialYear, EntityOfficer,
        TaxPlanningWorksheet, TaxPlanningBeneficiaryRow,
    )

    fy = FinancialYear.objects.select_related("entity").get(pk=financial_year_id)
    entity = fy.entity

    # Officers
    all_officers = EntityOfficer.objects.filter(
        entity=entity, date_ceased__isnull=True
    ).order_by("display_order", "full_name")

    trustees = [o for o in all_officers if _officer_has_role(o, "trustee")]
    chairperson = None
    for o in all_officers:
        if _officer_has_role(o, "chairperson"):
            chairperson = o
            break
    if not chairperson:
        for o in all_officers:
            if getattr(o, "is_chairperson", False):
                chairperson = o
                break

    # Build trustee name string
    trustee_names = [t.full_name for t in trustees]
    if len(trustee_names) == 1:
        trustee_str = trustee_names[0]
    elif len(trustee_names) == 2:
        trustee_str = f"{trustee_names[0]} and {trustee_names[1]}"
    elif len(trustee_names) > 2:
        trustee_str = ", ".join(trustee_names[:-1]) + f", and {trustee_names[-1]}"
    else:
        trustee_str = "(No trustees found)"

    chairperson_name = chairperson.full_name if chairperson else "(No chairperson found)"

    # FY year
    year_digits = "".join(c for c in fy.year_label if c.isdigit())
    fy_year = year_digits if year_digits else str(fy.end_date.year)

    # Worksheet data
    try:
        worksheet = TaxPlanningWorksheet.objects.get(financial_year=fy)
    except TaxPlanningWorksheet.DoesNotExist:
        worksheet = None

    rows = []
    total_tax = Decimal("0")
    total_distributed = Decimal("0")
    if worksheet:
        for row in worksheet.beneficiary_rows.select_related("beneficiary").order_by(
            "beneficiary__full_name"
        ):
            rows.append(row)
            total_distributed += row.proposed_distribution
            if row.beneficiary_type != "trust":
                total_tax += row.net_tax_payable

    distributable = worksheet.distributable_income if worksheet else Decimal("0")
    recommendation = worksheet.recommendation_notes if worksheet else ""

    return {
        "fy": fy,
        "entity": entity,
        "trust_name": entity.entity_name,
        "trustee_name": trustee_str,
        "chairperson_name": chairperson_name,
        "resolution_date": f"30 June {fy_year}",
        "financial_year_end": f"30 June {fy_year}",
        "fy_year": fy_year,
        "distributable_income": distributable,
        "total_tax_payable": total_tax,
        "total_distributed": total_distributed,
        "beneficiary_rows": rows,
        "accountant_recommendation": recommendation,
        "worksheet": worksheet,
    }


# =============================================================================
# Trust Election (s97) Generator
# =============================================================================

def generate_trust_election(financial_year_id):
    """
    Generate a Trust Election (s97/streaming) document.

    Returns:
        io.BytesIO buffer containing the .docx file.
    """
    ctx = _get_trust_context(financial_year_id)
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY

    # Title
    _add_heading(doc, "TRUST ELECTION", level=1)
    _add_heading(doc, "Pursuant to Section 97 of the Income Tax Assessment Act 1936", level=2)
    doc.add_paragraph()

    # Trust details
    _add_para(doc, f"Trust Name: {ctx['trust_name']}", bold=True)
    _add_para(doc, f"Trustee: {ctx['trustee_name']}")
    _add_para(doc, f"Financial Year Ending: {ctx['financial_year_end']}")
    doc.add_paragraph()

    # Resolution
    _add_para(doc, "RESOLUTION", bold=True)
    doc.add_paragraph()

    _add_para(
        doc,
        f"The Trustee of {ctx['trust_name']} hereby resolves, in accordance with "
        f"the Trust Deed and Section 97 of the Income Tax Assessment Act 1936, "
        f"that the net income of the trust estate for the year ending "
        f"{ctx['financial_year_end']} be distributed as follows:"
    )
    doc.add_paragraph()

    # Beneficiary distribution table
    _add_para(doc, "Distribution of Net Income", bold=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ["Beneficiary", "Type", "Amount", "% Share"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                _set_run_font(run, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Dark background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="333333"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    distributable = ctx["distributable_income"]
    for row in ctx["beneficiary_rows"]:
        data_row = table.add_row()
        data_row.cells[0].text = row.beneficiary.full_name
        data_row.cells[1].text = row.get_beneficiary_type_display()
        data_row.cells[2].text = _fmt_money(row.proposed_distribution)
        data_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        pct = (
            (row.proposed_distribution / distributable * 100).quantize(Decimal("0.01"))
            if distributable > 0
            else Decimal("0")
        )
        data_row.cells[3].text = f"{pct}%"
        data_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = "TOTAL"
    for p in total_row.cells[0].paragraphs:
        for run in p.runs:
            _set_run_font(run, bold=True)
    total_row.cells[2].text = _fmt_money(ctx["total_distributed"])
    total_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for p in total_row.cells[2].paragraphs:
        for run in p.runs:
            _set_run_font(run, bold=True)
    total_row.cells[3].text = "100.00%"
    total_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Streaming election details
    ws = ctx["worksheet"]
    if ws and (ws.capital_gains > 0 or ws.franked_dividends > 0):
        _add_para(doc, "STREAMING ELECTION", bold=True)
        doc.add_paragraph()

        _add_para(
            doc,
            f"The Trustee further resolves that the following income components "
            f"be specifically streamed to the nominated beneficiaries in accordance "
            f"with Division 6E of the Income Tax Assessment Act 1936:"
        )
        doc.add_paragraph()

        if ws.capital_gains > 0:
            _add_para(doc, f"Capital Gains: {_fmt_money(ws.capital_gains)}")

        if ws.franked_dividends > 0:
            _add_para(doc, f"Franked Dividends: {_fmt_money(ws.franked_dividends)}")
            _add_para(doc, f"Franking Credits: {_fmt_money(ws.franking_credits)}")

        doc.add_paragraph()

    # Fallback clause
    _add_para(
        doc,
        "In the event of an increase or decrease in the income of the Trust Fund "
        "for any reason, it was also resolved that such increase or decrease be "
        "distributed in the same proportions as set out above."
    )
    doc.add_paragraph()

    _add_para(
        doc,
        "In the event any apportionment of income above is found or held to be "
        "invalid or ineffective, it was resolved that such amount be distributed "
        "equally between all beneficiaries."
    )
    doc.add_paragraph()
    doc.add_paragraph()

    # Signature block
    _add_para(doc, "Signed as a true and correct record.")
    doc.add_paragraph()
    doc.add_paragraph()
    _add_para(doc, "______________________________")
    _add_para(doc, ctx["chairperson_name"])
    _add_para(doc, "Chairperson / Trustee")
    doc.add_paragraph()
    _add_para(doc, f"Date: {ctx['resolution_date']}")

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# Tax Planning Summary (Client-Facing) Generator
# =============================================================================

def generate_tax_planning_summary(financial_year_id):
    """
    Generate a Tax Planning Summary document (client-facing).

    Returns:
        io.BytesIO buffer containing the .docx file.
    """
    ctx = _get_trust_context(financial_year_id)
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY

    # Title
    _add_heading(doc, "TAX PLANNING SUMMARY", level=1)
    _add_heading(doc, ctx["trust_name"], level=2)
    _add_para(
        doc,
        f"Financial Year Ending {ctx['financial_year_end']}",
        italic=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()

    # Prepared by
    _add_para(doc, "Prepared by: M C & S Pty Ltd", italic=True)
    _add_para(doc, f"Date: {date.today().strftime('%d %B %Y')}", italic=True)
    doc.add_paragraph()

    # Section 1: Distributable Income
    _add_para(doc, "1. DISTRIBUTABLE INCOME", bold=True)
    doc.add_paragraph()

    ws = ctx["worksheet"]
    if ws:
        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"

        items = [
            ("Net Profit Before Distributions", ws.net_profit_before_distributions),
            ("Add: Non-Deductible Expenses", ws.non_deductible_expenses),
            ("Less: Non-Assessable Income", ws.non_assessable_income),
        ]
        for i, (label, val) in enumerate(items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = _fmt_money(val)
            table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Distributable income (bold)
        table.rows[3].cells[0].text = "Trust Distributable Income"
        table.rows[3].cells[1].text = _fmt_money(ws.distributable_income)
        table.rows[3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for cell in table.rows[3].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    _set_run_font(run, bold=True)

        # Streamable components
        components = []
        if ws.capital_gains > 0:
            components.append(f"Capital Gains: {_fmt_money(ws.capital_gains)}")
        if ws.franked_dividends > 0:
            components.append(f"Franked Dividends: {_fmt_money(ws.franked_dividends)}")
        if ws.franking_credits > 0:
            components.append(f"Franking Credits: {_fmt_money(ws.franking_credits)}")

        table.rows[4].cells[0].text = "Of which (streamable):"
        table.rows[4].cells[1].text = "; ".join(components) if components else "Nil"
        for p in table.rows[4].cells[0].paragraphs:
            for run in p.runs:
                _set_run_font(run, italic=True, size=FONT_SIZE_SMALL)

    doc.add_paragraph()

    # Section 2: Proposed Distribution & Tax Impact
    _add_para(doc, "2. PROPOSED DISTRIBUTION & TAX IMPACT", bold=True)
    doc.add_paragraph()

    # Beneficiary table
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Beneficiary", "Type", "Distribution", "Est. Tax", "Eff. Rate", "Notes"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                _set_run_font(run, bold=True, size=FONT_SIZE_SMALL)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="333333"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)

    for row in ctx["beneficiary_rows"]:
        data_row = table.add_row()
        data_row.cells[0].text = row.beneficiary.full_name
        data_row.cells[1].text = row.get_beneficiary_type_display()

        data_row.cells[2].text = _fmt_money(row.proposed_distribution)
        data_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        if row.beneficiary_type == "trust":
            data_row.cells[3].text = "Refer to sub-trust"
            data_row.cells[4].text = "—"
            data_row.cells[5].text = "Separate tax plan required"
        else:
            data_row.cells[3].text = _fmt_money(row.net_tax_payable)
            data_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            eff_rate = (
                f"{(row.effective_tax_rate * 100).quantize(Decimal('0.01'))}%"
                if row.effective_tax_rate
                else "0.00%"
            )
            data_row.cells[4].text = eff_rate
            data_row.cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            notes = []
            if row.beneficiary_type == "company":
                if row.company_tax_rate_override:
                    notes.append(f"Non-base rate ({row.company_tax_rate_override * 100:.0f}%)")
                else:
                    notes.append("Base rate entity (25%)")
            data_row.cells[5].text = "; ".join(notes) if notes else ""

    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = "TOTAL"
    for p in total_row.cells[0].paragraphs:
        for run in p.runs:
            _set_run_font(run, bold=True)
    total_row.cells[2].text = _fmt_money(ctx["total_distributed"])
    total_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_row.cells[3].text = _fmt_money(ctx["total_tax_payable"])
    total_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in [total_row.cells[2], total_row.cells[3]]:
        for p in cell.paragraphs:
            for run in p.runs:
                _set_run_font(run, bold=True)

    doc.add_paragraph()

    # Section 3: Summary
    _add_para(doc, "3. SUMMARY", bold=True)
    doc.add_paragraph()

    distributable = ctx["distributable_income"]
    undistributed = distributable - ctx["total_distributed"]

    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = "Table Grid"

    summary_items = [
        ("Total Distributable Income", _fmt_money(distributable)),
        ("Total Proposed Distributions", _fmt_money(ctx["total_distributed"])),
        ("Undistributed Balance", _fmt_money(undistributed)),
        ("Total Estimated Tax", _fmt_money(ctx["total_tax_payable"])),
        (
            "Weighted Effective Tax Rate",
            f"{(ctx['total_tax_payable'] / ctx['total_distributed'] * 100).quantize(Decimal('0.01'))}%"
            if ctx["total_distributed"] > 0
            else "0.00%",
        ),
    ]
    for i, (label, val) in enumerate(summary_items):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = val
        summary_table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Section 4: Recommendation
    _add_para(doc, "4. ACCOUNTANT'S RECOMMENDATION", bold=True)
    doc.add_paragraph()

    recommendation = ctx["accountant_recommendation"]
    if recommendation:
        # Strip HTML tags for plain text output
        import re
        clean_text = re.sub(r"<[^>]+>", "", recommendation)
        clean_text = clean_text.replace("&nbsp;", " ").replace("&amp;", "&")
        clean_text = clean_text.replace("&lt;", "<").replace("&gt;", ">")
        _add_para(doc, clean_text)
    else:
        _add_para(doc, "(No recommendation recorded.)", italic=True)

    doc.add_paragraph()

    # Disclaimer
    _add_para(doc, "DISCLAIMER", bold=True)
    _add_para(
        doc,
        "This tax planning summary is based on the information provided and current "
        "tax legislation as at the date of preparation. It is intended as a guide only "
        "and does not constitute formal tax advice. Actual tax outcomes may differ based "
        "on individual circumstances, changes in legislation, or ATO rulings. "
        "Please contact our office to discuss any questions.",
        italic=True,
    )

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
