"""
StatementHub — Management Accounts Generator

Generates period-scoped, watermarked management accounts (Cover + B/S + P&L)
for any date range within a financial year.  Supports three TB sources:

  1. Xero / QBO connected  — transient API pull (never saved to DB)
  2. Bank statement upload  — derived from coded transactions ≤ period_end
  3. Manual TB upload       — date picker disabled; full-year only

The output is a stripped version of the annual financial statements:
  ✓ Cover Page  ✓ Balance Sheet  ✓ Profit & Loss
  ✗ Notes  ✗ Director's Declaration  ✗ Compilation Report

Every page carries a red header watermark:
  "DRAFT | Generated DD MMM YYYY HH:MM by Name"
"""
import io
import logging
from datetime import date, datetime
from decimal import Decimal, ROUND_HALF_UP
from collections import OrderedDict

from django.utils import timezone

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from .models import (
    Entity, FinancialYear, TrialBalanceLine,
)
from .docgen import (
    FONT_NAME, FONT_SIZE_BODY, FONT_SIZE_HEADING, FONT_SIZE_SUBHEADING,
    FONT_SIZE_SMALL, FONT_SIZE_FOOTER,
    FIRM_NAME, FIRM_ADDRESS_1, FIRM_ADDRESS_2, FIRM_PHONE, FIRM_EMAIL, FIRM_WEBSITE,
    _set_run_font, _add_paragraph, _add_centered_heading, _add_horizontal_line,
    _get_logo_path, _start_report_section, _has_prior_year, _has_cogs,
    _add_detailed_pnl, _add_detailed_balance_sheet, _add_trading_account,
    _build_note_registry,
)
from .table_helpers import FinancialTable

logger = logging.getLogger(__name__)


# =============================================================================
# TB Source Detection
# =============================================================================

def detect_tb_source(entity):
    """
    Determine the trial balance data source for an entity.

    Returns one of: 'XERO', 'QBO', 'MYOB', 'BANK_DERIVED', 'MANUAL'
    """
    # Check for active cloud accounting connection
    try:
        from integrations.models import AccountingConnection
        conn = AccountingConnection.objects.filter(
            entity=entity,
            status=AccountingConnection.Status.ACTIVE,
        ).first()
        if conn:
            return conn.provider.upper()  # 'XERO', 'QUICKBOOKS', 'MYOB'
    except Exception:
        pass

    # Check for bank statement data
    try:
        from review.models import PendingTransaction
        has_bank_data = PendingTransaction.objects.filter(
            job__entity=entity,
            is_confirmed=True,
        ).exists()
        if has_bank_data:
            return 'BANK_DERIVED'
    except Exception:
        pass

    return 'MANUAL'


# =============================================================================
# Transient TB — Xero / QBO API Pull
# =============================================================================

def fetch_transient_tb_from_cloud(entity, period_start, period_end):
    """
    Pull a trial balance from the connected cloud accounting provider
    as at period_end.  Returns a list of dicts with keys:
        account_code, account_name, debit, credit

    CRITICAL: This data is NEVER written to the database.
    """
    from integrations.models import AccountingConnection
    from integrations.providers import get_provider

    conn = AccountingConnection.objects.filter(
        entity=entity,
        status=AccountingConnection.Status.ACTIVE,
    ).first()
    if not conn:
        raise ValueError("No active cloud accounting connection found.")

    provider = get_provider(conn.provider)

    # Refresh token if needed
    if conn.token_expires_at and conn.token_expires_at <= timezone.now():
        try:
            new_tokens = provider.refresh_access_token(conn.refresh_token)
            conn.access_token = new_tokens['access_token']
            conn.refresh_token = new_tokens.get('refresh_token', conn.refresh_token)
            conn.token_expires_at = new_tokens.get('expires_at')
            conn.save(update_fields=['access_token', 'refresh_token', 'token_expires_at'])
        except Exception as e:
            raise ValueError(f"Token refresh failed: {e}")

    raw_lines = provider.fetch_trial_balance(
        conn.access_token, conn.tenant_id, period_end
    )
    return raw_lines


def _cloud_tb_to_sections(raw_lines):
    """
    Convert raw cloud TB lines (list of dicts) into the sections dict
    expected by docgen rendering functions.

    Uses the same account code ranges as _get_tb_sections in docgen.py.
    """
    sections = {
        "trading_income": [], "cogs": [], "income": [], "expenses": [],
        "current_assets": [], "noncurrent_assets": [],
        "current_liabilities": [], "noncurrent_liabilities": [],
        "equity": [],
    }

    for line in raw_lines:
        code = line.get("account_code", "")
        name = line.get("account_name", "")
        debit = Decimal(str(line.get("debit", 0)))
        credit = Decimal(str(line.get("credit", 0)))
        current_amount = debit - credit
        prior_amount = Decimal("0")  # No prior for transient pull
        entry = (code, name, current_amount, prior_amount)

        try:
            code_num = int(code.split('.')[0])
        except (ValueError, TypeError):
            continue

        name_lower = name.lower()
        is_cogs = any(kw in name_lower for kw in
                      ("cost of", "opening stock", "closing stock", "purchases",
                       "stock on hand"))

        if code_num < 1000:
            is_other = any(kw in name_lower for kw in
                           ("interest", "other", "fbt", "contribution",
                            "dividend", "sundry"))
            is_trading = any(kw in name_lower for kw in
                             ("sales", "income", "takings", "revenue",
                              "accommodation", "conference", "meals", "bar",
                              "trading"))
            if is_other:
                sections["income"].append(entry)
            elif is_trading:
                sections["trading_income"].append(entry)
            else:
                sections["income"].append(entry)
        elif code_num < 1200:
            sections["cogs"].append(entry)
        elif code_num < 2000:
            sections["cogs"].append(entry) if is_cogs else sections["expenses"].append(entry)
        elif code_num < 2500:
            sections["current_assets"].append(entry)
        elif code_num < 3000:
            sections["noncurrent_assets"].append(entry)
        elif code_num < 3500:
            sections["current_liabilities"].append(entry)
        elif code_num < 4000:
            sections["noncurrent_liabilities"].append(entry)
        elif code_num < 5000:
            sections["equity"].append(entry)
        elif code_num < 6000:
            sections["cogs"].append(entry)

    return sections


# =============================================================================
# Bank-Derived TB
# =============================================================================

def build_bank_derived_tb(entity, fy, period_end):
    """
    Aggregate coded bank transactions with date <= period_end into a
    provisional trial balance.  Includes opening balances from prior year
    rollforward for balance sheet accounts.

    Returns the same sections dict as _get_tb_sections.
    """
    from review.models import PendingTransaction
    from collections import defaultdict

    # 1. Gather all confirmed transactions for this entity
    #    NOTE: PendingTransaction.date is a CharField stored as "dd/mm/yyyy",
    #    so we cannot use date__lte / date__gte at the DB level. Instead, we
    #    fetch all confirmed txns and filter by parsed date in Python.
    from datetime import datetime as _dt

    all_txns = PendingTransaction.objects.filter(
        job__entity=entity,
        is_confirmed=True,
    ).select_related('job')

    def _parse_txn_date(date_str):
        """Parse dd/mm/yyyy or yyyy-mm-dd date strings."""
        for fmt in ("%d/%m/%Y", "%Y-%m-%d"):
            try:
                return _dt.strptime(date_str.strip(), fmt).date()
            except (ValueError, AttributeError):
                continue
        return None

    # Aggregate by account code
    account_totals = defaultdict(lambda: {"debit": Decimal("0"), "credit": Decimal("0"), "name": ""})

    for txn in all_txns:
        # Filter by date in Python
        txn_date = _parse_txn_date(txn.date)
        if txn_date is None:
            continue
        if txn_date < fy.start_date or txn_date > period_end:
            continue

        code = txn.confirmed_code or txn.ai_suggested_code or ""
        name = txn.confirmed_name or txn.ai_suggested_name or ""
        if not code:
            continue

        amount = txn.amount or Decimal("0")
        if amount >= 0:
            # Positive = deposit/income → credit
            account_totals[code]["credit"] += abs(amount)
        else:
            # Negative = expense → debit
            account_totals[code]["debit"] += abs(amount)

        if not account_totals[code]["name"]:
            account_totals[code]["name"] = name

    # 2. Add opening balances from prior year TB (balance sheet accounts only)
    if fy.prior_year:
        prior_lines = TrialBalanceLine.objects.filter(
            financial_year=fy.prior_year
        )
        for line in prior_lines:
            try:
                code_num = int(line.account_code.split('.')[0])
            except (ValueError, TypeError):
                continue
            # Only balance sheet accounts (2000+)
            if code_num >= 2000:
                account_totals[line.account_code]["debit"] += line.debit
                account_totals[line.account_code]["credit"] += line.credit
                if not account_totals[line.account_code]["name"]:
                    account_totals[line.account_code]["name"] = line.account_name

    # 3. Convert to sections dict
    sections = {
        "trading_income": [], "cogs": [], "income": [], "expenses": [],
        "current_assets": [], "noncurrent_assets": [],
        "current_liabilities": [], "noncurrent_liabilities": [],
        "equity": [],
    }

    for code, data in sorted(account_totals.items()):
        current_amount = data["debit"] - data["credit"]
        entry = (code, data["name"], current_amount, Decimal("0"))

        try:
            code_num = int(code.split('.')[0])
        except (ValueError, TypeError):
            continue

        name_lower = data["name"].lower()
        is_cogs = any(kw in name_lower for kw in
                      ("cost of", "opening stock", "closing stock", "purchases",
                       "stock on hand"))

        if code_num < 1000:
            is_other = any(kw in name_lower for kw in
                           ("interest", "other", "fbt", "contribution",
                            "dividend", "sundry"))
            is_trading = any(kw in name_lower for kw in
                             ("sales", "income", "takings", "revenue",
                              "accommodation", "conference", "meals", "bar",
                              "trading"))
            if is_other:
                sections["income"].append(entry)
            elif is_trading:
                sections["trading_income"].append(entry)
            else:
                sections["income"].append(entry)
        elif code_num < 1200:
            sections["cogs"].append(entry)
        elif code_num < 2000:
            sections["cogs"].append(entry) if is_cogs else sections["expenses"].append(entry)
        elif code_num < 2500:
            sections["current_assets"].append(entry)
        elif code_num < 3000:
            sections["noncurrent_assets"].append(entry)
        elif code_num < 3500:
            sections["current_liabilities"].append(entry)
        elif code_num < 4000:
            sections["noncurrent_liabilities"].append(entry)
        elif code_num < 5000:
            sections["equity"].append(entry)
        elif code_num < 6000:
            sections["cogs"].append(entry)

    return sections


# =============================================================================
# Manual TB — Period-Scoped from Existing TB Lines
# =============================================================================

def build_manual_tb_sections(fy):
    """
    For manual TB uploads, use the existing full-year TB data.
    This is the same as _get_tb_sections in docgen.py but imported here
    for consistency.
    """
    from .docgen import _get_tb_sections
    return _get_tb_sections(fy)


# =============================================================================
# Cover Page — Management Accounts
# =============================================================================

def _add_mgmt_cover_page(doc, entity, period_start, period_end, generated_by_name):
    """
    Add the management accounts cover page.

    Layout:
        [MC&S Logo]
        [ENTITY NAME]
        ABN [ABN]
        DRAFT
        MANAGEMENT ACCOUNTS
        For the period [start] to [end]
        Prepared: [datetime] by [name]
        These accounts have not been finalised or reviewed.
    """
    now = timezone.localtime(timezone.now())

    # Add MC&S logo — centered
    logo_path = _get_logo_path()
    if logo_path:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(6))

    # Spacing after logo
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(28)

    # Entity name
    _add_centered_heading(doc, entity.entity_name, size=Pt(16), bold=True, space_after=4)

    # Trading As
    if entity.trading_as:
        _add_centered_heading(doc, f"Trading As {entity.trading_as}",
                              size=Pt(14), bold=False, space_after=4)

    # ABN
    if entity.abn:
        _add_centered_heading(doc, f"ABN {entity.abn}",
                              size=Pt(11), bold=False, space_after=8)

    # Spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)

    # DRAFT warning — red, bold, above Management Accounts
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("DRAFT")
    _set_run_font(run, size=Pt(12), bold=True)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # "Management Accounts" title
    _add_centered_heading(doc, "Management Accounts",
                          size=Pt(14), bold=True, space_after=4)

    # Period text — formatted dates
    start_str = period_start.strftime('%-d %B %Y')
    end_str = period_end.strftime('%-d %B %Y')
    _add_centered_heading(doc, f"For the period {start_str} to {end_str}",
                          size=Pt(11), bold=False, space_after=12)

    # Prepared line
    prepared_str = now.strftime('%-d %B %Y %H:%M')
    _add_centered_heading(doc, f"Prepared: {prepared_str} by {generated_by_name}",
                          size=Pt(10), bold=False, space_after=2)

    # Disclaimer
    _add_centered_heading(doc, "These accounts have not been finalised or reviewed.",
                          size=Pt(10), bold=False, space_after=0)

    # Push firm details to bottom
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)

    # Firm details — consolidated into fewer lines
    _add_centered_heading(doc, FIRM_NAME, size=Pt(10), bold=False, space_after=0)
    _add_centered_heading(doc, f"{FIRM_ADDRESS_1}, {FIRM_ADDRESS_2}",
                          size=Pt(9), bold=False, space_after=2)
    _add_centered_heading(doc, f"{FIRM_PHONE}  |  {FIRM_EMAIL}  |  {FIRM_WEBSITE}",
                          size=Pt(9), bold=False, space_after=0)

    doc.add_page_break()


# =============================================================================
# Header Watermark — Red, Bold, 9pt, Right-Aligned
# =============================================================================

def _inject_mgmt_watermark(doc, generated_by_name):
    """
    Inject a red, bold, 9pt, right-aligned watermark into the page header
    of every section.  This survives LibreOffice PDF conversion.

    Text: "DRAFT | Generated DD MMM YYYY HH:MM by Name"
    """
    now = timezone.localtime(timezone.now())
    timestamp = now.strftime('%-d %b %Y %H:%M')
    watermark_text = (
        f"DRAFT | "
        f"Generated {timestamp} by {generated_by_name}"
    )

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Insert watermark as the FIRST paragraph in the header
        # We need to prepend, not append, so the watermark appears above
        # any existing header content
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # If the paragraph already has content, insert a new one before it
        if p.text.strip():
            new_p = parse_xml(
                f'<w:p {nsdecls("w")}>'
                f'  <w:pPr><w:jc w:val="right"/></w:pPr>'
                f'</w:p>'
            )
            p._element.addprevious(new_p)
            # Now add the run to the new paragraph
            from docx.text.paragraph import Paragraph
            p = Paragraph(new_p, header)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

        run = p.add_run(watermark_text)
        _set_run_font(run, size=Pt(9), bold=True)
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)


# (Bank statement disclaimer removed per firm request)


# =============================================================================
# Main Generation Function
# =============================================================================

class _FYProxy:
    """
    A lightweight proxy that mimics FinancialYear for the docgen rendering
    functions, but with overridden dates for the selected period.
    """
    def __init__(self, fy, period_start, period_end):
        self._fy = fy
        self.start_date = period_start
        self.end_date = period_end
        self.entity = fy.entity
        self.prior_year = None  # No comparatives for management accounts
        self.pk = fy.pk
        self.id = fy.id
        self.year_label = fy.year_label
        self.status = fy.status
        # Determine period_type based on date range
        delta = (period_end - period_start).days + 1
        if delta <= 45:
            self.period_type = 'monthly'
        elif delta <= 105:
            self.period_type = 'quarterly'
        elif delta <= 200:
            self.period_type = 'half_year'
        elif delta <= 380:
            self.period_type = 'annual'
        else:
            self.period_type = 'interim'

    def __getattr__(self, name):
        return getattr(self._fy, name)


def generate_management_accounts(
    financial_year_id,
    period_start,
    period_end,
    user,
    output_type='bs_pnl',  # 'bs_pnl' or 'pnl_only'
):
    """
    Generate period-scoped management accounts.

    Args:
        financial_year_id: UUID of the FinancialYear
        period_start: date object — start of the reporting period
        period_end: date object — end of the reporting period
        user: Django User who triggered the generation
        output_type: 'bs_pnl' for Balance Sheet + P&L, 'pnl_only' for P&L only

    Returns:
        (buffer, tb_source) — BytesIO with the Word document, and the TB source string

    Raises:
        ValueError: if validation fails or manual TB with custom period
    """
    fy = FinancialYear.objects.select_related(
        "entity", "entity__client", "prior_year"
    ).get(pk=financial_year_id)

    entity = fy.entity
    show_cents = entity.show_cents
    generated_by_name = user.get_full_name() or user.username

    # Validate dates
    if period_end < period_start:
        raise ValueError("Period end date must be after period start date.")
    if period_end < fy.start_date or period_end > fy.end_date:
        raise ValueError("Period end date must be within the financial year.")
    if period_start < fy.start_date or period_start > fy.end_date:
        raise ValueError("Period start date must be within the financial year.")

    # Detect TB source
    tb_source = detect_tb_source(entity)

    # Build sections based on TB source
    if tb_source in ('XERO', 'QUICKBOOKS', 'MYOB'):
        # Transient API pull — NEVER saved to DB
        try:
            raw_lines = fetch_transient_tb_from_cloud(entity, period_start, period_end)
            sections = _cloud_tb_to_sections(raw_lines)
        except Exception as e:
            logger.warning(f"Cloud TB pull failed for {entity}: {e}. Falling back to manual TB.")
            # Fall back to existing TB data
            tb_source = 'MANUAL'
            sections = build_manual_tb_sections(fy)
    elif tb_source == 'BANK_DERIVED':
        sections = build_bank_derived_tb(entity, fy, period_end)
    else:
        # Manual TB — use existing full-year data
        sections = build_manual_tb_sections(fy)

    # Create FY proxy with period dates
    fy_proxy = _FYProxy(fy, period_start, period_end)

    # Build the document
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_BODY

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    has_trading = _has_cogs(sections)

    # ---- Cover Page ----
    _add_mgmt_cover_page(doc, entity, period_start, period_end, generated_by_name)

    # ---- Trading Account (if COGS exist) ----
    gross_profit = None
    gross_profit_prior = None
    if has_trading:
        gross_profit, gross_profit_prior = _add_trading_account(
            doc, entity, fy_proxy, sections, show_cents=show_cents)

    # ---- Detailed P&L ----
    net_profit, net_profit_prior = _add_detailed_pnl(
        doc, entity, fy_proxy, sections, show_cents=show_cents,
        gross_profit=gross_profit, gross_profit_prior=gross_profit_prior,
        note_registry=None)

    # ---- Balance Sheet (unless P&L only) ----
    if output_type == 'bs_pnl':
        _add_detailed_balance_sheet(
            doc, entity, fy_proxy, sections, show_cents=show_cents,
            net_profit=net_profit, net_profit_prior=net_profit_prior,
            note_registry=None)

    # ---- Inject Watermark ----
    _inject_mgmt_watermark(doc, generated_by_name)

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer, tb_source
