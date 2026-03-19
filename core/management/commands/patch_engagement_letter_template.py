"""
Management command: patch_engagement_letter_template

Patches the active engagement letter template (.docx) stored in the database
to use expanded scope of service descriptions, removes blank paragraphs in the
fees section, and applies keep-together page formatting to section headings.

This command modifies the template file in-place and updates the database record.

Usage:
    python manage.py patch_engagement_letter_template
    python manage.py patch_engagement_letter_template --dry-run
"""
import io
import os

from django.core.files.base import ContentFile
from django.core.management.base import BaseCommand, CommandError

from core.models import LegalDocumentTemplate


# ---------------------------------------------------------------------------
# Expanded scope descriptions (Jinja2-compatible — preserve {{ }} tags)
# ---------------------------------------------------------------------------
SCOPE_DESCRIPTIONS = {
    "Financial Statements": (
        "Preparation of annual financial statements in accordance with applicable Australian Accounting Standards, "
        "including the Balance Sheet, Profit & Loss Statement, Statement of Changes in Equity, and Notes to the "
        "Financial Statements. We will review the trial balance, apply any necessary year-end adjustments, and "
        "prepare the financial statements in a form suitable for director approval and, where required, lodgement "
        "with ASIC. The financial statements will reflect the entity's financial position as at the end of the "
        "relevant financial year and will be accompanied by the Directors' Declaration confirming the statements "
        "give a true and fair view. Preparation of tax effect accounting entries and deferred tax disclosures are "
        "included where applicable."
    ),
    "Income Tax Return": (
        "Preparation and electronic lodgement of the income tax return for the entity for the relevant income year "
        "ended 30 June, in accordance with the Income Tax Assessment Act 1997 and applicable ATO requirements. "
        "Our service includes the calculation of taxable income, application of available deductions and offsets, "
        "preparation of the tax reconciliation from accounting profit to taxable income, and review of carry-forward "
        "losses and franking account balances. We will advise you of the estimated tax liability or refund position "
        "prior to lodgement and notify you of the due date for payment. This service does not include the lodgement "
        "of individual tax returns for directors, shareholders, or beneficiaries unless separately engaged."
    ),
    "BAS Preparation & Lodgement": (
        "Preparation and electronic lodgement of Business Activity Statements (BAS) for each applicable reporting "
        "period (monthly, quarterly, or annually as required by the ATO). Our service includes reconciliation of "
        "GST collected and GST paid to the accounting records, review of input tax credit entitlements, calculation "
        "of PAYG withholding obligations, and preparation of the instalment activity statement where applicable. "
        "We will advise you of any net GST liability or refund prior to lodgement and notify you of payment due "
        "dates. This service assumes that the accounting records are maintained in a current and accurate state "
        "prior to each BAS period. Amendments to prior period BAS lodgements are outside the scope of this service "
        "unless separately agreed."
    ),
    "Bookkeeping": (
        "Ongoing bookkeeping services to maintain accurate and up-to-date accounting records throughout the "
        "financial year. Our service includes data entry of transactions, bank and credit card reconciliations, "
        "accounts payable and receivable processing, maintenance of the general ledger, and coding of transactions "
        "in accordance with the agreed chart of accounts. We will liaise with you to resolve any unidentified or "
        "unusual transactions and will flag any items requiring your attention. This service is performed using "
        "cloud-based accounting software as agreed. The frequency and volume of bookkeeping services will be as "
        "agreed between the parties and may be adjusted by mutual agreement during the engagement."
    ),
    "Tax Planning & Advisory": (
        "Proactive tax planning and advisory services to assist you in managing your tax position throughout the "
        "financial year. Our service includes a review of your current and projected tax position, identification "
        "of lawful tax minimisation strategies, year-end planning recommendations (including timing of income and "
        "deductions, superannuation contributions, and asset purchases), and advice on the tax implications of "
        "proposed transactions or structural changes. We will provide written recommendations where appropriate "
        "and will be available to discuss planning matters as they arise. This service does not constitute formal "
        "tax advice for the purposes of the Tax Agent Services Act 2009 unless confirmed in a separate advice "
        "letter. Implementation of any recommended strategies is subject to your approval and instruction."
    ),
    "Payroll Services": (
        "Processing of employee payroll on the agreed frequency (weekly, fortnightly, or monthly), including "
        "calculation of gross wages, superannuation guarantee contributions, PAYG withholding, and any applicable "
        "salary sacrifice or salary packaging arrangements. Our service includes preparation of payslips, "
        "Single Touch Payroll (STP) reporting to the ATO, annual payment summary reconciliation, and preparation "
        "of the payroll finalisation declaration at year-end. We will advise you of any changes to superannuation "
        "guarantee rates, minimum wage adjustments, or STP reporting requirements as they arise. This service "
        "assumes that employee details, leave balances, and entitlements are maintained and provided to us in a "
        "timely manner. Payroll tax obligations to state revenue authorities are not included unless separately "
        "agreed."
    ),
    "ASIC Annual Review & Compliance": (
        "Review of the ASIC annual statement issued to the company and payment of the annual review fee on your "
        "behalf (reimbursed by you). Our service includes confirmation that the company's registered details "
        "(officeholders, registered address, and share structure) are accurate and up to date, preparation and "
        "lodgement of any required ASIC forms to correct or update company details, and preparation of the "
        "solvency declaration by the directors as required under section 347A of the Corporations Act 2001. "
        "We will notify you of the annual review date and advise of any changes required. This service does not "
        "include advice on corporate governance, restructuring, or the preparation of special purpose financial "
        "statements for ASIC lodgement unless separately agreed."
    ),
    "Dividend Management": (
        "Preparation of the documentation required to declare and pay dividends to shareholders, including "
        "dividend resolutions (board minutes), dividend statements for each shareholder, and maintenance of the "
        "company's franking account. Our service includes calculation of the available franking credits, "
        "determination of the appropriate franking percentage, and preparation of the dividend statement in the "
        "form required by the Income Tax Assessment Act 1997. We will advise you on the tax implications of "
        "proposed dividends, including the impact on shareholders' individual tax positions where relevant. "
        "This service assumes that the company's retained earnings and franking account balance are confirmed "
        "prior to the declaration of any dividend. Dividends paid to non-resident shareholders and any "
        "applicable dividend withholding tax obligations are outside the scope of this service unless "
        "separately agreed."
    ),
    "Director's Report": (
        "Preparation of the Directors' Report as required under Part 2M.3 of the Corporations Act 2001 for "
        "the relevant financial year. Our service includes preparation of the principal activities statement, "
        "review of the operating results and financial position, identification of any significant changes in "
        "the state of affairs, and preparation of the dividends and events after balance date disclosures. "
        "Where required, we will prepare the auditor's independence declaration and the lead auditor's "
        "independence declaration for inclusion in the Directors' Report. This service is applicable to "
        "companies that are required to prepare a Directors' Report under the Corporations Act 2001. "
        "The Directors' Report will be presented to the directors for approval prior to the signing of "
        "the financial statements."
    ),
    "Compilation Report": (
        "Preparation of a compilation report in accordance with APES 315 Compilation of Financial Information. "
        "Our service includes the compilation of the financial information provided by management into financial "
        "statements, the application of appropriate accounting policies, and the preparation of the compilation "
        "report for inclusion with the financial statements. A compilation engagement does not involve the "
        "performance of audit or review procedures, and accordingly we do not express an audit opinion or "
        "review conclusion on the financial statements. The financial statements are based on the information "
        "provided by you and we are not responsible for errors or omissions resulting from incomplete or "
        "inaccurate information. The compilation report will be signed by a registered company auditor or "
        "a qualified accountant as required."
    ),
    "Trust Distribution Planning": (
        "Preparation of the trust distribution resolution and supporting documentation for the relevant "
        "income year, including identification of the distributable income of the trust, allocation of "
        "income to beneficiaries, and preparation of the trustee resolution to distribute income. "
        "Our service includes review of the trust deed to confirm the trustee's distribution powers, "
        "assessment of the tax implications of proposed distributions (including the application of "
        "Division 7A, section 100A, and the trust tax provisions), and preparation of the distribution "
        "minutes and beneficiary statements. We will advise you of the required resolution date (no later "
        "than 30 June) and the tax consequences of the proposed distribution. This service does not include "
        "the preparation of beneficiary tax returns unless separately engaged."
    ),
    "Trust Deed Review": (
        "Review of the trust deed and any amending deeds to confirm that the trust structure remains "
        "appropriate for the entity's current circumstances and tax planning objectives. Our service "
        "includes identification of any provisions that may restrict the trustee's distribution powers, "
        "assessment of the trust's compliance with the trust tax provisions (including the definition of "
        "net income and the streaming provisions), and recommendations for any amendments required. "
        "This service does not constitute legal advice and we recommend that you obtain independent "
        "legal advice before making any amendments to the trust deed."
    ),
    "Division 7A Monitoring": (
        "Ongoing monitoring and management of Division 7A loan accounts to ensure compliance with the "
        "minimum yearly repayment requirements under section 109N of the Income Tax Assessment Act 1936. "
        "Our service includes calculation of the minimum yearly repayment for each loan account, "
        "preparation of the Division 7A loan agreement (where required), review of the distributable "
        "surplus position, and identification of any deemed dividend exposure. We will advise you of "
        "the required repayment amount prior to 30 June each year and will flag any accounts that are "
        "at risk of triggering a deemed dividend. This service assumes that all loans and unpaid present "
        "entitlements are disclosed to us in a timely manner."
    ),
    "FBT": (
        "Preparation and lodgement of the Fringe Benefits Tax (FBT) return for the FBT year ended "
        "31 March, including identification and valuation of all fringe benefits provided to employees "
        "and their associates during the year. Our service includes review of motor vehicle usage, "
        "expense payments, loans, and other benefits, calculation of the FBT liability using the "
        "applicable valuation methods, and preparation of the employee payment summaries where required. "
        "We will advise you of any FBT planning opportunities and will notify you of the lodgement and "
        "payment due dates. This service assumes that all relevant information regarding benefits "
        "provided to employees is disclosed to us in full."
    ),
    "TPAR": (
        "Preparation and lodgement of the Taxable Payments Annual Report (TPAR) for the relevant "
        "income year, reporting payments made to contractors in the building and construction, "
        "cleaning, courier, road freight, information technology, and security industries as required "
        "by the ATO. Our service includes review of contractor payment records, confirmation of "
        "contractor ABNs, and preparation of the TPAR in the format required by the ATO. We will "
        "advise you of the lodgement due date (28 August each year) and will notify you of any "
        "contractors for whom an ABN has not been quoted. This service assumes that contractor "
        "payment records are maintained and provided to us in a complete and accurate form."
    ),
}


SECTION_HEADINGS = {
    "1. SCOPE OF SERVICES",
    "2. OUR FEES",
    "3. YOUR RESPONSIBILITIES",
    "4. OUR RESPONSIBILITIES",
    "5. INDEPENDENCE",
    "6. CONFIDENTIALITY",
    "7. CONFLICTS OF INTEREST",
    "8. ELECTRONIC COMMUNICATION",
    "9. PRIVACY",
    "10. LIMITATION OF LIABILITY",
    "11. DISPUTE RESOLUTION",
    "12. PROFESSIONAL STANDARDS",
    "13. COMMENCEMENT",
    "ACCEPTANCE OF TERMS",
}


def _set_keep_with_next(paragraph):
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    kwn = OxmlElement("w:keepNext")
    pPr.append(kwn)
    kt = OxmlElement("w:keepLines")
    pPr.append(kt)


def _set_keep_together(paragraph):
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    kt = OxmlElement("w:keepLines")
    pPr.append(kt)


def _set_table_keep_together(table):
    from docx.oxml import OxmlElement
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        trPr.append(cant)


def patch_document(doc):
    """Apply all improvements to a docx Document object."""
    from docx.shared import Pt

    # 1. Expand scope descriptions in the scope table
    # Find the scope table (has 'Service' and 'Description' in header row)
    scope_table = None
    for table in doc.tables:
        for row in table.rows:
            cells_text = [c.text.strip() for c in row.cells]
            if "Service" in cells_text and "Description" in cells_text:
                scope_table = table
                break
        if scope_table:
            break

    if scope_table:
        for row in scope_table.rows:
            service_cell = row.cells[1]
            desc_cell = row.cells[2]
            service_name = service_cell.text.strip()
            if service_name in SCOPE_DESCRIPTIONS:
                for para in desc_cell.paragraphs:
                    para.clear()
                para = desc_cell.paragraphs[0]
                run = para.add_run(SCOPE_DESCRIPTIONS[service_name])
                run.font.size = Pt(9)
                _set_keep_together(para)
            for para in service_cell.paragraphs:
                _set_keep_together(para)
        _set_table_keep_together(scope_table)

    # 2. Remove extra blank paragraphs in fees section
    in_fees = False
    blank_count = 0
    paras_to_remove = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if "2. OUR FEES" in text or text == "OUR FEES":
            in_fees = True
            blank_count = 0
            continue
        if in_fees:
            if text.startswith("3.") or "YOUR RESPONSIBILITIES" in text:
                in_fees = False
                break
            if text == "":
                blank_count += 1
                if blank_count > 1:
                    paras_to_remove.append(para._p)
            else:
                blank_count = 0
    for p_elem in paras_to_remove:
        p_elem.getparent().remove(p_elem)

    # 3. Apply keep_with_next to section headings
    for para in doc.paragraphs:
        text = para.text.strip()
        is_heading = any(text == h or text.startswith(h) for h in SECTION_HEADINGS)
        if not is_heading and len(text) > 2 and text[0].isdigit() and "." in text[:3]:
            is_heading = True
        if is_heading:
            _set_keep_with_next(para)

    # 4. Apply keep_together to paragraphs immediately after headings
    paragraphs = list(doc.paragraphs)
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        if i > 0:
            prev_text = paragraphs[i - 1].text.strip()
            is_after_heading = (
                any(prev_text == h or prev_text.startswith(h) for h in SECTION_HEADINGS)
                or (len(prev_text) > 2 and prev_text[0].isdigit() and "." in prev_text[:3])
            )
            if is_after_heading and text:
                _set_keep_together(para)

    # 5. Apply cantSplit to all tables
    for table in doc.tables:
        _set_table_keep_together(table)

    return doc


class Command(BaseCommand):
    help = (
        "Patch the active engagement letter template with expanded scope descriptions, "
        "fixed fees section layout, and keep-together page formatting."
    )

    def add_arguments(self, parser):
        parser.add_argument(
            "--dry-run",
            action="store_true",
            help="Show what would be done without making changes.",
        )

    def handle(self, *args, **options):
        from docx import Document

        dry_run = options["dry_run"]

        template = LegalDocumentTemplate.objects.filter(
            document_type="engagement_letter",
            is_active=True,
        ).first()

        if not template:
            raise CommandError(
                "No active engagement letter template found. "
                "Run seed_engagement_letter_template first."
            )

        if not template.template_file:
            raise CommandError(f"Template record {template.pk} has no file attached.")

        self.stdout.write(
            f"Found template: '{template.name}' v{template.version} "
            f"(id={template.pk}, file={template.template_file.name})"
        )

        if dry_run:
            self.stdout.write(self.style.WARNING("Dry run — no changes made."))
            return

        # Load, patch, save back
        template.template_file.seek(0)
        doc = Document(template.template_file)
        doc = patch_document(doc)

        # Save to BytesIO
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        # Overwrite the file in storage
        original_name = os.path.basename(template.template_file.name)
        template.template_file.delete(save=False)
        template.template_file.save(
            original_name,
            ContentFile(buf.read()),
            save=True,
        )

        self.stdout.write(
            self.style.SUCCESS(
                f"Template patched successfully: '{template.name}' v{template.version}"
            )
        )
        self.stdout.write(
            "All future engagement letters will use the expanded scope descriptions."
        )
