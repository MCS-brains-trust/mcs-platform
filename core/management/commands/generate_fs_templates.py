"""
Management command: generate_fs_templates

Builds default .docx template files using docxtpl-compatible Jinja2 merge
fields and registers them in the FinancialStatementTemplate model.

Usage:
    python3 manage.py generate_fs_templates
    python3 manage.py generate_fs_templates --force   # overwrite existing
"""
import os

from django.conf import settings
from django.core.management.base import BaseCommand

from docx import Document
from docx.shared import Inches, Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.models import FinancialStatementTemplate


# Template layout constants — Handiledger reference standard
FONT_HEADING = "Arial"  # Entity name, ABN, title, section headers, footer (Handiledger-aligned)
FONT_BODY = "Times New Roman"       # Line items, column headers, totals, notes
FONT_NAME = FONT_BODY               # Legacy alias used by existing helpers
FONT_SIZE = Pt(10)
FONT_SIZE_HEADING = Pt(11)
PAGE_MARGIN_TOP = Cm(1.6)       # 16mm
PAGE_MARGIN_BOTTOM = Cm(1.7)    # 17mm
PAGE_MARGIN_LEFT = Cm(2.0)      # 20mm
PAGE_MARGIN_RIGHT = Cm(2.4)     # 24mm

# Column widths for 4-column table: name(8.5cm) note(1.5cm) cy(3cm) py(3cm)
COL_WIDTHS = [Cm(8.5), Cm(1.5), Cm(3), Cm(3)]

# Entity types that get templates
ENTITY_TYPES = ["company", "trust", "sole_trader", "partnership"]

# Document types — Compilation Report last per APES 315
DOC_TYPES = [
    ("COVER", "Cover Page"),
    ("DETAILED_PL", "Detailed Profit and Loss Statement"),
    ("BALANCE_SHEET", "Detailed Balance Sheet"),
    ("SUMMARY_PL", "Summary P&L"),
    ("DEPRECIATION_REPORT", "Depreciation Report"),
    ("NOTES", "Notes to Financial Statements"),
    ("DECLARATION", "Declaration"),
    ("DISTRIBUTION", "Distribution Summary"),
    ("COMPILATION", "Compilation Report"),
]

# Which doc types apply to which entity types
ENTITY_DOC_TYPES = {
    "company": ["COVER", "DETAILED_PL", "BALANCE_SHEET", "SUMMARY_PL",
                 "DEPRECIATION_REPORT", "NOTES", "DECLARATION", "COMPILATION"],
    "trust": ["COVER", "DETAILED_PL", "BALANCE_SHEET",
              "DEPRECIATION_REPORT", "NOTES",
              "DECLARATION", "COMPILATION", "DISTRIBUTION"],
    "sole_trader": ["COVER", "DETAILED_PL", "BALANCE_SHEET",
                     "DEPRECIATION_REPORT", "NOTES",
                     "DECLARATION", "COMPILATION"],
    "partnership": ["COVER", "DETAILED_PL", "BALANCE_SHEET",
                     "DEPRECIATION_REPORT", "NOTES",
                     "DECLARATION", "COMPILATION"],
}


def _set_page_setup(doc):
    """Set A4 portrait page size and margins per DOCGEN.md spec."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = PAGE_MARGIN_TOP
        section.bottom_margin = PAGE_MARGIN_BOTTOM
        section.left_margin = PAGE_MARGIN_LEFT
        section.right_margin = PAGE_MARGIN_RIGHT


def _set_default_font(doc):
    """Set Calibri 10pt as default font."""
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE


def _add_para(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT,
              size=None, color=None, keep_with_next=False):
    """Add a styled paragraph.

    Default space_after = Pt(3) — gives ~15pt effective row height on
    10pt body text, matching the Handiledger reference. Callers that need
    different spacing can override `p.paragraph_format.space_after` on
    the returned paragraph.
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size or FONT_SIZE
    run.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if keep_with_next:
        pPr = p._p.get_or_add_pPr()
        keepNext = OxmlElement('w:keepNext')
        keepNext.set(qn('w:val'), '1')
        pPr.append(keepNext)
    return p


def _set_table_full_width(table):
    """Set a document-body table width to full page text width (9356 twips = 16cm)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9356')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def _clear_table_borders(table):
    """
    Explicitly set all table-level borders to nil AND zero cell margins.

    LibreOffice inherits w:tblBorders from the table style (e.g. 'Table Grid')
    and renders them even when individual cells declare w:nil.  Clearing the
    table-level borders AND zeroing tblCellMar eliminates both the unwanted
    border lines and the excessive vertical gaps between sections.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove and replace tblBorders
    for tag in ('w:tblBorders', 'w:tblCellMar'):
        existing = tblPr.find(qn(tag))
        if existing is not None:
            tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)
    # Zero top/bottom cell margins so table sits flush (no gap between sections)
    tblCellMar = OxmlElement('w:tblCellMar')
    for side in ('top', 'bottom'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), '0')
        el.set(qn('w:type'), 'dxa')
        tblCellMar.append(el)
    tblPr.append(tblCellMar)


# ---------------------------------------------------------------------------
# Border helpers — Handiledger professional standard
# ---------------------------------------------------------------------------

# Row type constants for border application
ROW_TYPE_DATA = "data"
ROW_TYPE_HEADER = "header"
ROW_TYPE_SUBCATEGORY_SUBTOTAL = "subcategory_subtotal"
ROW_TYPE_SECTION_TOTAL = "section_total"
ROW_TYPE_MAJOR_TOTAL = "major_total"
ROW_TYPE_GRAND_TOTAL = "grand_total"

_NIL = {"val": "none", "sz": "0", "color": "auto"}
_SINGLE = {"val": "single", "sz": "6", "color": "000000"}
_DOUBLE = {"val": "double", "sz": "8", "color": "000000"}


def _apply_cell_border(cell, **kwargs):
    """Apply borders to a cell.

    kwargs: named sides (top, bottom, left, right, insideH, insideV) with
    {val, sz, color} dicts for sides that should be visible.

    ALL sides not explicitly passed are set to nil so LibreOffice cannot
    inherit borders from the table style on any unspecified edge.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if side in kwargs:
            attrs = kwargs[side]
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), attrs.get('val', 'single'))
            el.set(qn('w:sz'), str(attrs.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), attrs.get('color', '000000'))
        else:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            el.set(qn('w:sz'), '0')
            el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _apply_row_borders(row, row_type=ROW_TYPE_DATA):
    """Apply Handiledger-standard borders to a row based on row_type.

    Border rules (cols 2 & 3 only — cols 0 & 1 always nil):
      data / heading       → all nil
      header               → bottom single on cols 2,3
      subcategory_subtotal → top single on cols 2,3
      section_total        → top single + bottom double on cols 2,3
      major_total          → bottom double on cols 2,3
      grand_total          → bottom double on cols 2,3
    """
    num_cells = len(row.cells)

    # Determine border kwargs for amount columns based on row_type
    if row_type == ROW_TYPE_HEADER:
        amount_kwargs = {}
    elif row_type == ROW_TYPE_SUBCATEGORY_SUBTOTAL:
        amount_kwargs = {"top": _SINGLE}
    elif row_type == ROW_TYPE_SECTION_TOTAL:
        # Handiledger: section subtotal = single above + single below
        amount_kwargs = {"top": _SINGLE, "bottom": _SINGLE}
    elif row_type in (ROW_TYPE_MAJOR_TOTAL, ROW_TYPE_GRAND_TOTAL):
        # Handiledger: major/grand total = single above + double below
        amount_kwargs = {"top": _SINGLE, "bottom": _DOUBLE}
    else:
        amount_kwargs = {}  # data / heading — all nil

    for idx, cell in enumerate(row.cells):
        if idx >= num_cells - 2 and amount_kwargs:
            _apply_cell_border(cell, **amount_kwargs)
        else:
            _apply_cell_border(cell)  # nil all borders

    # Bold all text for total rows
    if row_type not in (ROW_TYPE_DATA, ROW_TYPE_HEADER):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True


# ---------------------------------------------------------------------------
# Page number footer helper
# ---------------------------------------------------------------------------
def _add_page_number_footer(doc):
    """Placeholder — page numbers are stamped on the final merged PDF.

    This sets up a minimal footer so the section knows it has one,
    but does NOT embed a PAGE field (those restart per-document).
    """
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    # Empty footer — page number will be stamped by _stamp_page_numbers()
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""


def _add_total_row(doc, label, cy_tag, py_tag, size=None,
                   row_type=ROW_TYPE_SECTION_TOTAL, grand_total=False):
    """Add a single-row 4-column table for a total/summary line.

    row_type controls the Handiledger border style applied.
    grand_total is a legacy parameter — if True, overrides row_type to GRAND_TOTAL.
    """
    if grand_total:
        row_type = ROW_TYPE_GRAND_TOTAL
    font_size = size or FONT_SIZE
    table = doc.add_table(rows=1, cols=4, style='Normal Table')
    _set_table_full_width(table)
    _clear_table_borders(table)
    table.autofit = False
    for i, width in enumerate(COL_WIDTHS):
        table.columns[i].width = width
    row = table.rows[0]
    # Prevent this summary row from splitting across pages
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    cantSplit = OxmlElement('w:cantSplit')
    cantSplit.set(qn('w:val'), '1')
    trPr.append(cantSplit)
    row.cells[0].text = label
    row.cells[1].text = ""
    row.cells[2].text = cy_tag
    row.cells[3].text = py_tag
    for i in range(4):
        for p in row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = font_size
                run.bold = True
    _apply_row_borders(row, row_type)


def _add_repeating_header(doc, document_title, date_field="{{ date_text }}"):
    """Add a repeating page header: entity name, ABN, doc title, date, horizontal rule.

    Uses Jinja2 variables rendered by docxtpl. Repeats on every page.
    Also includes the DRAFT watermark (hidden when context is empty).
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    header.is_linked_to_previous = False

    # Clear any existing header content
    for para in list(header.paragraphs):
        para.clear()

    # Entity name — Arial 13pt Bold, centred
    p1 = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p1.text = ""
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run("{{ entity_name }}")
    run.font.name = FONT_HEADING
    run.font.size = Pt(13)
    run.bold = True
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(0)

    # ABN — Arial 11pt, centred
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("ABN {{ abn }}")
    run2.font.name = FONT_HEADING
    run2.font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(0)

    # Document title — Arial 11pt, centred
    p3 = header.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(document_title)
    run3.font.name = FONT_HEADING
    run3.font.size = Pt(11)
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.space_before = Pt(0)

    # Date / period — Arial 11pt, centred. Pt(6) gap before the rule.
    p4 = header.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(date_field)
    run4.font.name = FONT_HEADING
    run4.font.size = Pt(11)
    p4.paragraph_format.space_after = Pt(6)
    p4.paragraph_format.space_before = Pt(0)

    # Separate rule paragraph — empty, with bottom border (0.5pt black)
    p_rule = header.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after = Pt(0)
    pPr_rule = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '4')    # 4/8 = 0.5pt
    bottom_border.set(qn('w:space'), '0')
    bottom_border.set(qn('w:color'), '000000')
    pBdr.append(bottom_border)
    pPr_rule.append(pBdr)

    # DRAFT watermark — right-aligned, red, only visible when non-empty
    pw = header.add_paragraph()
    pw.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runw = pw.add_run("{{ watermark }}")
    runw.font.name = FONT_NAME
    runw.font.size = Pt(14)
    runw.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    runw.bold = True
    pw.paragraph_format.space_after = Pt(0)
    pw.paragraph_format.space_before = Pt(0)


def _add_footer(doc, text="These financial statements are unaudited. They must be read in conjunction with the attached Accountant\u2019s Compilation Report and Notes which form part of these financial statements."):
    """Add standard footer with full unaudited disclaimer text.

    - Centre-aligned, Times New Roman italic 9pt (Handiledger standard).
    - Top border (0.5pt black) creates a horizontal rule above the text,
      repeating on every page because this is a section footer.
    """
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = FONT_BODY  # Times New Roman
    run.font.size = Pt(9)
    run.font.italic = True

    # Top border on the footer paragraph — 0.5pt solid black, full width
    pPr = p._p.get_or_add_pPr()
    # Remove any existing pBdr first to avoid duplicates
    for existing in pPr.findall(qn('w:pBdr')):
        pPr.remove(existing)
    pBdr = OxmlElement('w:pBdr')
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '4')       # 4/8 = 0.5pt
    top_border.set(qn('w:space'), '4')
    top_border.set(qn('w:color'), '000000')
    pBdr.append(top_border)
    pPr.append(pBdr)


def _add_spacer(doc, pts=4):
    """Add a minimal spacer paragraph with controlled height to separate sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(pts * 2)))  # half-points
    rPr = OxmlElement('w:rPr')
    rPr.append(sz)
    pPr.append(rPr)
    return p


def _add_financial_table(doc, section_title, items_tag, total_label, total_cy_tag, total_py_tag):
    """Add a 4-column financial table with Jinja2 for-loop.

    The section title (e.g. 'Income', 'Expenses') is placed as the FIRST ROW
    of the table rather than as a standalone paragraph before the table.
    This eliminates the LibreOffice paragraph-before-table bottom-border
    artefact that produced an unwanted horizontal line above the year headers.
    """
    # Build the table — combined heading+years row + dollar row + data rows + total
    table = doc.add_table(rows=1, cols=4, style='Normal Table')
    _set_table_full_width(table)
    _clear_table_borders(table)
    table.autofit = False
    for i, width in enumerate(COL_WIDTHS):
        table.columns[i].width = width

    # Row 0 — Combined section heading + column headers (Handiledger layout)
    # cells[0]: section title (Arial Bold 11pt, left-aligned)
    # cells[1-3]: Note / year / prior_year (Arial Bold 9pt, right-aligned)
    heading_row = table.rows[0]
    heading_row.cells[0].text = section_title
    heading_row.cells[1].text = "Note"
    heading_row.cells[2].text = "{{ year }}"
    heading_row.cells[3].text = "{{ prior_year }}"
    # Style cell 0 (section title) — Arial Bold 11pt
    for p in heading_row.cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.name = FONT_HEADING
            run.font.size = FONT_SIZE_HEADING
            run.bold = True
    # Style cells 1-3 (column headers) — Arial Bold 9pt
    for i in range(1, 4):
        for p in heading_row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = FONT_HEADING
                run.font.size = Pt(9)
                run.bold = True
    for i in range(4):
        _apply_cell_border(heading_row.cells[i])
    # Keep heading row with the dollar row below it
    tr_heading = heading_row._tr
    trPr_heading = tr_heading.get_or_add_trPr()
    cantSplit_heading = OxmlElement('w:cantSplit')
    cantSplit_heading.set(qn('w:val'), '1')
    trPr_heading.append(cantSplit_heading)

    # Row 1 — Dollar sign row: "$" below each year column
    dollar_row = table.add_row()
    dollar_row.cells[0].text = ""
    dollar_row.cells[1].text = ""
    dollar_row.cells[2].text = "$"
    dollar_row.cells[3].text = "$"
    for i in range(4):
        for p in dollar_row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = FONT_HEADING  # Arial
                run.font.size = Pt(9)
                run.bold = True
    # No borders on dollar row — matches Handiledger
    _apply_row_borders(dollar_row, ROW_TYPE_DATA)
    # Keep dollar row with first data row
    tr_dollar = dollar_row._tr
    trPr_dollar = tr_dollar.get_or_add_trPr()
    kn_dollar = OxmlElement('w:keepNext')
    kn_dollar.set(qn('w:val'), '1')
    trPr_dollar.append(kn_dollar)

    # Row 2 — {%tr for %} tag in its own row (docxtpl requirement)
    for_row = table.add_row()
    for_row.cells[0].text = "{%tr for item in " + items_tag + " %}"
    # Set minimum row height so docxtpl loop row doesn't create a visible gap
    tr_for = for_row._tr
    trPr_for = tr_for.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '1')
    trHeight.set(qn('w:hRule'), 'exact')
    trPr_for.append(trHeight)
    for cell in for_row.cells:
        _apply_cell_border(cell)

    # Row 3 — data row with item fields
    data_row = table.add_row()
    data_row.cells[0].text = "{{ item.account_name }}"
    data_row.cells[1].text = "{{ item.note_ref }}"
    data_row.cells[2].text = "{{ item.cy_formatted }}"
    data_row.cells[3].text = "{{ item.py_formatted }}"
    for i in range(4):
        for p in data_row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)  # ~15pt effective row height (Handiledger)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
        _apply_cell_border(data_row.cells[i])

    # Row 3 — {%tr endfor %} tag in its own row
    endfor_row = table.add_row()
    endfor_row.cells[0].text = "{%tr endfor %}"
    # Set minimum row height so endfor row doesn't create a visible gap
    tr_end = endfor_row._tr
    trPr_end = tr_end.get_or_add_trPr()
    trHeight_end = OxmlElement('w:trHeight')
    trHeight_end.set(qn('w:val'), '1')
    trHeight_end.set(qn('w:hRule'), 'exact')
    trPr_end.append(trHeight_end)
    for cell in endfor_row.cells:
        _apply_cell_border(cell)

    # Total row — subtotal borders (single top + single bottom on amount cells)
    total_row = table.add_row()
    total_row.cells[0].text = total_label
    total_row.cells[2].text = total_cy_tag
    total_row.cells[3].text = total_py_tag
    for i in range(4):
        for p in total_row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                run.bold = True
    # Section subtotal: single above + single below on amount cells (Handiledger)
    _apply_row_borders(total_row, ROW_TYPE_SECTION_TOTAL)

    return table


# ---------------------------------------------------------------------------
# Template builders
# ---------------------------------------------------------------------------
def _build_cover(entity_type):
    """Build cover page template — first page of the assembled package."""
    doc = Document()
    _set_default_font(doc)
    _set_page_setup(doc)

    # No repeating header on cover page — leave header empty
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    # --- Logo at top: prefer FirmSettings upload, fall back to static files ---
    import os as _os
    logo_path = None
    try:
        from core.models import FirmSettings
        logo_path = FirmSettings.get().logo_path
    except Exception:
        pass
    if not logo_path:
        logo_candidates = [
            _os.path.join(settings.BASE_DIR, "static", "img", "mcs_logo.png"),
            _os.path.join(settings.BASE_DIR, "static", "MCSlogo.png"),
        ]
        for candidate in logo_candidates:
            if _os.path.isfile(candidate):
                logo_path = candidate
                break

    _add_para(doc, "", size=Pt(12))  # small top spacer
    if logo_path:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(6))

    # --- Entity details (NO ACN on cover — only ABN) ---
    _add_para(doc, "", size=Pt(16))  # spacer after logo

    p_name = _add_para(doc, "{{ entity_name }}", bold=True, size=Pt(18),
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p_name.paragraph_format.space_after = Pt(4)

    p_abn = _add_para(doc, "ABN {{ abn }}", size=Pt(11),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p_abn.paragraph_format.space_after = Pt(16)

    p_fs = _add_para(doc, "Financial Statements", size=Pt(11),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p_fs.paragraph_format.space_after = Pt(4)

    p_date = _add_para(doc, "{{ date_text }}", size=Pt(11),
                       alignment=WD_ALIGN_PARAGRAPH.CENTER)
    p_date.paragraph_format.space_after = Pt(0)

    # --- Firm contact details — pushed to lower portion of page ---
    p_spacer = _add_para(doc, "", size=Pt(8))
    p_spacer.paragraph_format.space_before = Pt(160)  # push to ~60% down

    # Single continuous block — no gap between address and contact details
    for line in [
        "{{ practice_name or firm_name }}",
        "{{ practice_registered_address or firm_address_1 }}",
        "",
        "Phone: {{ practice_phone or firm_phone }}",
        "Email: {{ practice_email or firm_email }}",
        "{% if practice_website %}Website: {{ practice_website }}{% endif %}",
    ]:
        p = _add_para(doc, line, size=Pt(9), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    # --- Page break: Contents starts on page 2 ---
    doc.add_page_break()

    # --- Contents page (page 2) ---
    # Heading — Palatino Linotype (fallback Georgia) 16pt Bold, with a
    # thin horizontal rule beneath it.
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading.add_run("Contents")
    heading_run.font.name = "Palatino Linotype"
    heading_run.font.size = Pt(16)
    heading_run.bold = True
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(48)
    # Thin horizontal rule (0.5pt, black) below the heading text
    hPr = heading._p.get_or_add_pPr()
    hBdr = OxmlElement('w:pBdr')
    hBot = OxmlElement('w:bottom')
    hBot.set(qn('w:val'), 'single')
    hBot.set(qn('w:sz'), '4')      # 4/8 = 0.5pt
    hBot.set(qn('w:space'), '4')
    hBot.set(qn('w:color'), '000000')
    hBdr.append(hBot)
    hPr.append(hBdr)

    # Build the contents list — conditional per entity_type (unchanged)
    contents = [
        "Detailed Profit and Loss Statement",
        "Detailed Balance Sheet",
    ]
    if entity_type == "company":
        contents.append("Summary Profit and Loss Statement")
    contents.append("Notes to the Financial Statements")
    if entity_type == "company":
        contents.append("Directors' Declaration")
        contents.append("Solvency Resolution")
    elif entity_type == "trust":
        contents.append("Trustee's Declaration")
        contents.append("Beneficiaries Distribution Summary")
    elif entity_type == "sole_trader":
        contents.append("Proprietor Declaration")
    elif entity_type == "partnership":
        contents.append("Partners' Declaration")
    contents.append("Compilation Report")
    contents.append("Management Representation Letter")

    # Two-column borderless table
    #   Left: document name (bold Times New Roman 11pt, left)
    #   Right: empty (reserved for future page-number wiring)
    contents_table = doc.add_table(rows=len(contents), cols=2, style='Normal Table')
    _set_table_full_width(contents_table)
    _clear_table_borders(contents_table)
    contents_table.autofit = False
    # Column widths — approx 75% / 25% of text area (16cm usable width)
    contents_table.columns[0].width = Cm(12.0)
    contents_table.columns[1].width = Cm(4.0)

    for row_idx, item in enumerate(contents):
        row = contents_table.rows[row_idx]
        # Enforce minimum row height 18pt
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '360')       # 18pt = 360 twips
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

        # Left cell — document name
        left = row.cells[0]
        left.width = Cm(12.0)
        left_p = left.paragraphs[0]
        left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_p.paragraph_format.space_before = Pt(0)
        left_p.paragraph_format.space_after = Pt(6)
        left_run = left_p.add_run(item)
        left_run.font.name = "Times New Roman"
        left_run.font.size = Pt(11)
        left_run.bold = True
        _apply_cell_border(left)

        # Right cell — intentionally empty (column preserved for future
        # page-number wiring; page numbers currently suppressed globally).
        right = row.cells[1]
        right.width = Cm(4.0)
        right_p = right.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right_p.paragraph_format.space_before = Pt(0)
        right_p.paragraph_format.space_after = Pt(6)
        _apply_cell_border(right)

    return doc


def _build_detailed_pl(entity_type):
    """Build Detailed P&L template."""
    doc = Document()
    _set_default_font(doc)
    _set_page_setup(doc)
    _add_repeating_header(doc, "Detailed Profit and Loss Statement", "{{ date_text }}")
    _add_footer(doc)

    # 0.5cm gap between header rule and first body content
    _p0 = doc.add_paragraph()
    _p0.paragraph_format.space_before = Emu(71050)
    _p0.paragraph_format.space_after = Pt(0)

    # Income section
    _add_financial_table(doc, "Income", "income", "Total Income",
                         "{{ total_income_cy }}", "{{ total_income_py }}")

    # Pt(18) gap before Expenses heading (larger than standard _add_spacer)
    _exp_spacer = doc.add_paragraph()
    _exp_spacer.paragraph_format.space_before = Pt(18)
    _exp_spacer.paragraph_format.space_after = Pt(0)

    # Expenses section
    _add_financial_table(doc, "Expenses", "expenses", "Total Expenses",
                         "{{ total_expenses_cy }}", "{{ total_expenses_py }}")
    _add_spacer(doc)

    # Net Profit section — shows tax breakdown when income tax exists
    # Pre-tax profit — grand_total=True because when has_income_tax is false
    # this IS the final "Net Profit / (Loss)" line with double bottom.
    # When has_income_tax is true, the post-processor reclassifies by label.
    _add_total_row(doc, "{% if has_income_tax %}Operating profit before income tax{% else %}Net Profit / (Loss){% endif %}",
                   "{{ net_profit_pretax_cy }}", "{{ net_profit_pretax_py }}",
                   grand_total=True)

    # Income tax line (only when tax exists) — use conditional Jinja2 block
    _add_para(doc, "{% if has_income_tax %}", size=Pt(1))

    # Income tax row
    tax_table = doc.add_table(rows=1, cols=4, style='Normal Table')
    _set_table_full_width(tax_table)
    _clear_table_borders(tax_table)
    tax_table.autofit = False
    for i, width in enumerate(COL_WIDTHS):
        tax_table.columns[i].width = width
    tax_row = tax_table.rows[0]
    tax_row.cells[0].text = "Income tax attributable to operating profit (loss)"
    tax_row.cells[2].text = "{{ income_tax_cy }}"
    tax_row.cells[3].text = "{{ income_tax_py }}"
    for i in range(4):
        for p in tax_row.cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
        _apply_cell_border(tax_row.cells[i])

    # After-tax profit (grand total — double bottom on amount columns)
    _add_total_row(doc, "Operating profit after income tax",
                   "{{ net_profit_cy }}", "{{ net_profit_py }}",
                   grand_total=True)

    _add_para(doc, "{% endif %}", size=Pt(1))

    return doc


def _build_balance_sheet(entity_type):
    """Build Balance Sheet template."""
    doc = Document()
    _set_default_font(doc)
    _set_page_setup(doc)
    _add_repeating_header(doc, "Detailed Balance Sheet", "As at {{ year_end_date }}")
    _add_footer(doc)

    # 0.5cm gap between header rule and first body content
    _p0 = doc.add_paragraph()
    _p0.paragraph_format.space_before = Emu(71050)
    _p0.paragraph_format.space_after = Pt(0)

    # Current Assets
    _add_financial_table(doc, "Current Assets", "current_assets", "Total Current Assets",
                         "{{ total_current_assets_cy }}", "{{ total_current_assets_py }}")
    _add_spacer(doc)

    # Non-Current Assets — suppressed when zero
    _add_para(doc, "{% if has_noncurrent_assets %}", size=Pt(1))
    _add_financial_table(doc, "Non-Current Assets", "noncurrent_assets",
                         "Total Non-Current Assets",
                         "{{ total_noncurrent_assets_cy }}", "{{ total_noncurrent_assets_py }}")
    _add_spacer(doc)
    _add_para(doc, "{% endif %}", size=Pt(1))

    # Total Assets — major total (double underline below)
    _add_total_row(doc, "Total Assets", "{{ total_assets_cy }}", "{{ total_assets_py }}",
                   row_type=ROW_TYPE_MAJOR_TOTAL)
    _add_spacer(doc)

    # Current Liabilities
    _add_financial_table(doc, "Current Liabilities", "current_liabilities",
                         "Total Current Liabilities",
                         "{{ total_current_liab_cy }}", "{{ total_current_liab_py }}")
    _add_spacer(doc)

    # Non-Current Liabilities — suppressed when zero
    _add_para(doc, "{% if has_noncurrent_liabilities %}", size=Pt(1))
    _add_financial_table(doc, "Non-Current Liabilities", "noncurrent_liabilities",
                         "Total Non-Current Liabilities",
                         "{{ total_noncurrent_liab_cy }}", "{{ total_noncurrent_liab_py }}")
    _add_spacer(doc)
    _add_para(doc, "{% endif %}", size=Pt(1))

    # Total Liabilities — major total (double underline below)
    _add_total_row(doc, "Total Liabilities", "{{ total_liabilities_cy }}", "{{ total_liabilities_py }}",
                   row_type=ROW_TYPE_MAJOR_TOTAL)
    _add_spacer(doc)

    # Net Assets — major total (double underline below)
    _add_total_row(doc, "Net Assets", "{{ net_assets_cy }}", "{{ net_assets_py }}",
                   size=Pt(12), row_type=ROW_TYPE_MAJOR_TOTAL)

    # Equity — flows directly after Net Assets without forced page break
    _add_financial_table(doc, "Equity", "equity", "Total Equity",
                         "{{ total_equity_cy }}", "{{ total_equity_py }}")

    return doc


def _build_summary_pl(entity_type):
    """Build Summary P&L template (companies only)."""
    doc = Document()
    _set_default_font(doc)
    _set_page_setup(doc)
    _add_repeating_header(doc, "Summary Profit and Loss Statement", "{{ date_text }}")
    _add_footer(doc)

    # 0.5cm gap between header rule and first body content
    _p0 = doc.add_paragraph()
    _p0.paragraph_format.space_before = Emu(71050)
    _p0.paragraph_format.space_after = Pt(0)

    # Summary table
    # Rows:
    #  0 — header (year / $)
    #  1 — Total Income
    #  2 — Total Expenses
    #  3 — Operating profit before income tax          [bold, subtotal border]
    #  4 — Income tax attributable to operating profit
    #  5 — Operating profit after income tax           [bold, subtotal border]
    #  6 — (spacer)
    #  7 — Retained profit at the beginning of the financial year
    #  8 — Total available for appropriation           [bold, subtotal border]
    #  9 — Retained profits at the end of the financial year [bold, grand-total border]
    rows_data = [
        ("", "{{ year }}\n$", "{{ prior_year }}\n$", False, None),
        ("Total Income", "{{ total_income_cy }}", "{{ total_income_py }}", False, None),
        ("Total Expenses", "{{ total_expenses_cy }}", "{{ total_expenses_py }}", False, None),
        ("Operating profit before income tax", "{{ net_profit_pretax_cy }}", "{{ net_profit_pretax_py }}", True, "subtotal"),
        ("Income tax attributable to operating profit (loss)", "{{ income_tax_cy }}", "{{ income_tax_py }}", False, None),
        ("Operating profit after income tax", "{{ net_profit_cy }}", "{{ net_profit_py }}", True, "subtotal"),
        ("", "", "", False, None),
        ("Retained profit at the beginning of the financial year", "{{ retained_profit_opening_cy }}", "{{ retained_profit_opening_py }}", False, None),
        ("Total available for appropriation", "{{ total_available_cy }}", "{{ total_available_py }}", True, "subtotal"),
        ("Retained profits at the end of the financial year", "{{ retained_profit_closing_cy }}", "{{ retained_profit_closing_py }}", True, "grand_total"),
    ]

    table = doc.add_table(rows=len(rows_data), cols=3, style='Normal Table')
    _set_table_full_width(table)
    _clear_table_borders(table)
    table.autofit = False
    table.columns[0].width = Cm(10)
    table.columns[1].width = Cm(3)
    table.columns[2].width = Cm(3)

    for r, (label, cy, py, is_bold, border_type) in enumerate(rows_data):
        table.rows[r].cells[0].text = label
        table.rows[r].cells[1].text = cy
        table.rows[r].cells[2].text = py
        for i in range(3):
            for p in table.rows[r].cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 1 else WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE
                    run.bold = is_bold or r == 0
        # Nil all borders on every row — _apply_cell_border with no kwargs clears all sides
        for cell in table.rows[r].cells:
            _apply_cell_border(cell)
        if border_type == "subtotal":
            _apply_row_borders(table.rows[r], row_type=ROW_TYPE_SECTION_TOTAL)
        elif border_type == "grand_total":
            _apply_row_borders(table.rows[r], row_type=ROW_TYPE_GRAND_TOTAL)

    return doc


def _build_notes(entity_type):
    """Build Notes template."""
    doc = Document()
    _set_default_font(doc)
    _set_page_setup(doc)
    _add_repeating_header(doc, "Notes to the Financial Statements", "{{ date_text }}")
    _add_footer(doc)

    # 0.5cm gap between header rule and first body content
    _p0 = doc.add_paragraph()
    _p0.paragraph_format.space_before = Emu(71050)
    _p0.paragraph_format.space_after = Pt(0)

    _add_para(doc, "Note 1: Statement of Significant Accounting Policies",
              bold=True, keep_with_next=True)
    _add_para(doc, "The financial statements are special purpose financial statements "
              "prepared in order to satisfy the financial reporting requirements of the "
              "Corporations Act 2001 or relevant trust deed. The directors/trustees have "
              "determined that the entity is not a reporting entity.")
    doc.add_paragraph("")
    _add_para(doc, "The financial statements have been prepared on an accruals basis and "
              "are based on historical costs.")
    doc.add_paragraph("")
    _add_para(doc, "The following significant accounting policies have been adopted in "
              "the preparation and presentation of the financial statements:")
    doc.add_paragraph("")

    _add_para(doc, "a) Revenue Recognition", bold=True, keep_with_next=True)
    _add_para(doc, "Revenue is recognised when the entity satisfies a performance obligation "
              "by transferring a promised good or service to a customer.")
    doc.add_paragraph("")

    _add_para(doc, "b) Income Tax", bold=True, keep_with_next=True)
    _add_para(doc, "The income tax expense for the year comprises current income tax expense. "
              "Current income tax expense reflects the current year tax payable based on "
              "taxable income for the year.")
    doc.add_paragraph("")

    _add_para(doc, "c) Goods and Services Tax (GST)", bold=True, keep_with_next=True)
    _add_para(doc, "Revenues, expenses and assets are recognised net of the amount of GST. "
              "Receivables and payables are stated with the amount of GST included.")

    return doc


def _build_declaration(entity_type):
    """Build Declaration template."""
    doc = Document()
    _set_default_font(doc)
    _set_page_setup(doc)
    _add_repeating_header(doc, "{{ declaration_title }}", "{{ date_text }}")
    _add_page_number_footer(doc)

    # 0.5cm gap between header rule and first body content
    _p0 = doc.add_paragraph()
    _p0.paragraph_format.space_before = Emu(71050)
    _p0.paragraph_format.space_after = Pt(0)

    if entity_type == "company":
        _add_para(doc,
                  "The directors of the company declare that the financial statements and "
                  "notes, as set out within this report, present fairly the company's financial "
                  "position as at {{ year_end_date }} and its performance for the year ended on "
                  "that date in accordance with the accounting policies described in Note 1 to "
                  "the financial statements.")
    elif entity_type == "trust":
        _add_para(doc,
                  "The trustee of the trust declares that the financial statements and notes, "
                  "as set out within this report, present fairly the trust's financial position "
                  "as at {{ year_end_date }} and its performance for the year ended on that date "
                  "in accordance with the accounting policies described in Note 1 to the "
                  "financial statements.")
    elif entity_type == "sole_trader":
        _add_para(doc,
                  "The proprietor declares that the financial statements and notes, as set out "
                  "within this report, present fairly the financial position as at "
                  "{{ year_end_date }} and the performance for the year ended on that date.")
    elif entity_type == "partnership":
        _add_para(doc,
                  "The partners declare that the financial statements and notes, as set out "
                  "within this report, present fairly the partnership's financial position as "
                  "at {{ year_end_date }} and its performance for the year ended on that date.")

    doc.add_paragraph("")
    _add_para(doc,
              "In the opinion of the {{ compilation_responsible_party }}, there are reasonable "
              "grounds to believe that the entity will be able to pay its debts as and when "
              "they become due and payable.")
    doc.add_paragraph("")

    _add_para(doc, "This declaration is made in accordance with a resolution of the "
              "{{ compilation_responsible_party }}.")
    doc.add_paragraph("")
    doc.add_paragraph("")

    # Signature block — trust uses corporate-trustee structure with individual signatories;
    # other entity types use the legacy directors loop.
    if entity_type == "trust":
        _add_para(doc, "{%p for sig in declaration_signatories %}")
        _add_para(doc, "____________________________")
        _add_para(doc, "{{ sig.name }}")
        _add_para(doc, "Director of {{ sig.trustee_company }}")
        _add_para(doc, "As Trustee of {{ sig.trust_name }}")
        _add_para(doc, "")
        _add_para(doc, "Dated: ___________________")
        _add_para(doc, "")
        _add_para(doc, "{%p endfor %}")
    else:
        _add_para(doc, "{%p for d in directors %}")
        _add_para(doc, "____________________________")
        _add_para(doc, "{{ d.name }}")
        _add_para(doc, "{{ d.title }}")
        _add_para(doc, "")
        _add_para(doc, "{%p endfor %}")
        _add_para(doc, "Dated: {{ signing_date }}")

    return doc


def _build_compilation(entity_type):
    """Build Compilation Report (APES 315) template — single page."""
    doc = Document()
    _set_default_font(doc)
    _set_page_setup(doc)
    _add_repeating_header(doc, "Compilation Report", "{{ date_text }}")
    _add_page_number_footer(doc)

    # 0.5cm gap between header rule and first body content
    _p0 = doc.add_paragraph()
    _p0.paragraph_format.space_before = Emu(71050)
    _p0.paragraph_format.space_after = Pt(0)

    # Tight spacing to fit on one page
    SP_BODY = Pt(4)    # space after body paragraphs
    SP_HEAD = Pt(6)    # space before section headings

    p = _add_para(doc, "To the {{ compilation_responsible_party }} of {{ entity_name }}")
    p.paragraph_format.space_after = SP_BODY

    p = _add_para(doc, "Scope", bold=True, keep_with_next=True)
    p.paragraph_format.space_before = SP_HEAD
    p.paragraph_format.space_after = Pt(2)
    p = _add_para(doc,
              "We have compiled the accompanying special purpose financial statements of "
              "{{ entity_name }}, which comprise the balance sheet as at {{ year_end_date }}, "
              "the profit and loss statement for the year then ended, and notes to the "
              "financial statements including a summary of significant accounting policies.")
    p.paragraph_format.space_after = SP_BODY

    p = _add_para(doc, "The Responsibility of the {{ compilation_responsible_party | capitalize }}", bold=True, keep_with_next=True)
    p.paragraph_format.space_before = SP_HEAD
    p.paragraph_format.space_after = Pt(2)
    p = _add_para(doc,
              "The {{ compilation_responsible_party }} of {{ entity_name }} are solely "
              "responsible for the information contained in the special purpose financial "
              "statements, and have determined that the accounting policies used are "
              "consistent and are appropriate to satisfy the requirements of the "
              "{{ compilation_responsible_party }}.")
    p.paragraph_format.space_after = SP_BODY

    p = _add_para(doc, "Our Responsibility", bold=True, keep_with_next=True)
    p.paragraph_format.space_before = SP_HEAD
    p.paragraph_format.space_after = Pt(2)
    p = _add_para(doc,
              "On the basis of information provided by the {{ compilation_responsible_party }}, "
              "we have compiled the accompanying special purpose financial statements in "
              "accordance with the applicable financial reporting framework and APES 315 "
              "Compilation of Financial Information.")
    p.paragraph_format.space_after = SP_BODY

    p = _add_para(doc,
              "We have applied our expertise in accounting and financial reporting to compile "
              "these financial statements in accordance with the applicable financial reporting "
              "framework. We have complied with the relevant ethical requirements of APES 110 "
              "Code of Ethics for Professional Accountants (including Independence Standards).")
    p.paragraph_format.space_after = SP_BODY

    p = _add_para(doc, "Assurance Disclaimer", bold=True, keep_with_next=True)
    p.paragraph_format.space_before = SP_HEAD
    p.paragraph_format.space_after = Pt(2)
    p = _add_para(doc,
              "Since a compilation engagement is not an assurance engagement, we are not "
              "required to verify the reliability, accuracy or completeness of the information "
              "provided to us by management and the {{ compilation_responsible_party }} to "
              "compile these financial statements. Accordingly, we do not express an audit "
              "opinion or a review conclusion on these financial statements.")
    p.paragraph_format.space_after = SP_BODY

    p = _add_para(doc, "The special purpose financial statements were compiled exclusively for "
              "the benefit of the {{ compilation_responsible_party }} who are responsible for "
              "the reliability, accuracy and completeness of the information compiled. We do "
              "not accept responsibility for the contents of the special purpose financial "
              "statements.", keep_with_next=True)
    p.paragraph_format.space_after = SP_BODY

    # Signature gap — leave ~45pt of whitespace for wet signature
    sig_gap = _add_para(doc, "", keep_with_next=True)
    sig_gap.paragraph_format.space_before = Pt(45)
    sig_gap.paragraph_format.space_after = Pt(0)

    # Signature line (28 underscores — matches Trustee Declaration + Mgmt Rep Letter)
    sig_line = _add_para(doc, "____________________________", keep_with_next=True)
    sig_line.paragraph_format.space_before = Pt(0)
    sig_line.paragraph_format.space_after = Pt(0)

    # Firm details block — all keep_with_next to prevent orphaning
    p = _add_para(doc, "{{ practice_name or firm_name }}", bold=True, keep_with_next=True)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p = _add_para(doc, "{{ practice_registered_address or (firm_address_1 + ', ' + firm_address_2) }}", keep_with_next=True)
    p.paragraph_format.space_after = Pt(0)
    p = _add_para(doc, "Registered Tax Agent No: {{ practice_tax_agent_number }}", keep_with_next=True)
    p.paragraph_format.space_after = Pt(0)
    p = _add_para(doc, "{{ practice_signatory_designation }}", keep_with_next=True)
    p.paragraph_format.space_after = SP_BODY
    p = _add_para(doc, "Dated: {{ signing_date }}")
    p.paragraph_format.space_after = Pt(0)

    return doc


def _build_distribution(entity_type):
    """Build Distribution Summary template (trusts only)."""
    doc = Document()
    _set_default_font(doc)
    _set_page_setup(doc)
    _add_repeating_header(doc, "Beneficiaries Distribution Summary", "{{ date_text }}")
    _add_footer(doc)

    # 0.5cm gap between header rule and first body content
    _p0 = doc.add_paragraph()
    _p0.paragraph_format.space_before = Emu(71050)
    _p0.paragraph_format.space_after = Pt(0)

    _add_para(doc, "Net Income Available for Distribution: {{ total_distribution }}", bold=True)
    doc.add_paragraph("")

    # Distribution table — explicit column widths so LibreOffice does not
    # collapse the amount column during PDF conversion.
    DIST_COL_WIDTHS = [5580, 1860, 1920]  # twips (dxa), total = 9360

    def _set_cell_width(cell, width_twips):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        # Remove any existing tcW
        existing = tcPr.find(qn('w:tcW'))
        if existing is not None:
            tcPr.remove(existing)
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_twips))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def _set_row_widths(row):
        for idx, cell in enumerate(row.cells):
            _set_cell_width(cell, DIST_COL_WIDTHS[idx])

    table = doc.add_table(rows=1, cols=3, style='Normal Table')
    # Set table preferred width
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9360')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    _clear_table_borders(table)
    table.autofit = False
    # Set column widths via python-docx API as well (belt and braces)
    table.columns[0].width = 5580
    table.columns[1].width = 1860
    table.columns[2].width = 1920

    # Header row
    hdr = table.rows[0]
    hdr.cells[0].text = "Beneficiary"
    hdr.cells[1].text = "Percentage"
    hdr.cells[2].text = "Amount\n$"
    _set_row_widths(hdr)
    for i in range(3):
        for p in hdr.cells[i].paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                run.bold = True
            if i >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Jinja2 row-level loop — {%tr for %} creates a new table row per beneficiary
    for_row = table.add_row()
    for_row.cells[0].text = "{%tr for b in beneficiaries %}"
    _set_row_widths(for_row)
    tr_for = for_row._tr
    trPr_for = tr_for.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '1')
    trHeight.set(qn('w:hRule'), 'exact')
    trPr_for.append(trHeight)

    # Data row with beneficiary fields
    data_row = table.add_row()
    data_row.cells[0].text = "{{ b.beneficiary_name }}"
    data_row.cells[1].text = "{{ b.percentage }}%"
    data_row.cells[2].text = "{{ b.amount }}"
    _set_row_widths(data_row)
    for i in range(3):
        for p in data_row.cells[i].paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
            if i >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # End loop row
    endfor_row = table.add_row()
    endfor_row.cells[0].text = "{%tr endfor %}"
    _set_row_widths(endfor_row)
    tr_end = endfor_row._tr
    trPr_end = tr_end.get_or_add_trPr()
    trHeight_end = OxmlElement('w:trHeight')
    trHeight_end.set(qn('w:val'), '1')
    trHeight_end.set(qn('w:hRule'), 'exact')
    trPr_end.append(trHeight_end)

    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = "Total"
    total_row.cells[1].text = "100%"
    total_row.cells[2].text = "{{ total_distribution }}"
    _set_row_widths(total_row)
    for i in range(3):
        for p in total_row.cells[i].paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE
                run.bold = True

    return doc


# Builder dispatch
BUILDERS = {
    "COVER": _build_cover,
    "DETAILED_PL": _build_detailed_pl,
    "BALANCE_SHEET": _build_balance_sheet,
    "SUMMARY_PL": _build_summary_pl,
    "NOTES": _build_notes,
    "DECLARATION": _build_declaration,
    "COMPILATION": _build_compilation,
    "DISTRIBUTION": _build_distribution,
}


class Command(BaseCommand):
    help = "Generate default .docx financial statement templates and register them in the database"

    def add_arguments(self, parser):
        parser.add_argument(
            "--force",
            action="store_true",
            help="Overwrite existing templates",
        )

    def handle(self, *args, **options):
        force = options["force"]
        output_dir = os.path.join(settings.MEDIA_ROOT, "fs_templates", "defaults")
        os.makedirs(output_dir, exist_ok=True)

        created = 0
        skipped = 0

        for entity_type in ENTITY_TYPES:
            applicable = ENTITY_DOC_TYPES.get(entity_type, [])
            for doc_type, doc_label in DOC_TYPES:
                if doc_type not in applicable:
                    continue

                # Check if active template already exists
                existing = FinancialStatementTemplate.objects.filter(
                    document_type=doc_type,
                    entity_type=entity_type,
                    is_active=True,
                ).first()

                if existing and not force:
                    self.stdout.write(
                        f"  SKIP {doc_type}/{entity_type} — already exists"
                    )
                    skipped += 1
                    continue

                # Build the template document
                builder = BUILDERS.get(doc_type)
                if not builder:
                    self.stdout.write(
                        self.style.WARNING(f"  No builder for {doc_type}")
                    )
                    continue

                doc = builder(entity_type)

                # Save to file
                filename = f"{doc_type}_{entity_type}.docx"
                filepath = os.path.join(output_dir, filename)
                doc.save(filepath)

                # Relative path from MEDIA_ROOT to the file on disk
                relative_path = f"fs_templates/defaults/{filename}"

                if existing and force:
                    # Update existing
                    existing.template_file = relative_path
                    existing.name = f"{doc_label} — {entity_type.replace('_', ' ').title()}"
                    existing.version = "1.0"
                    existing.save()
                    action = "UPDATED"
                else:
                    # Create new
                    tmpl = FinancialStatementTemplate(
                        name=f"{doc_label} — {entity_type.replace('_', ' ').title()}",
                        document_type=doc_type,
                        entity_type=entity_type,
                        version="1.0",
                        is_active=True,
                        template_file=relative_path,
                    )
                    tmpl.save()
                    action = "CREATED"

                self.stdout.write(
                    self.style.SUCCESS(f"  {action} {doc_type}/{entity_type}: {filename}")
                )
                created += 1

        self.stdout.write(
            self.style.SUCCESS(
                f"\nDone: {created} templates created/updated, {skipped} skipped."
            )
        )
