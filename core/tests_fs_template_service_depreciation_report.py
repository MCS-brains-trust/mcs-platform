"""
Regression test for core/fs_template_service.py:_generate_depreciation_report
=============================================================================
Ensures disposed assets (disposal_date populated) have Total Cost blanked in
the body row and excluded from category and grand total subtotals. Non-disposed
assets (including fully-depreciated-but-owned, i.e. CWDV=0 with no
disposal_date) retain Total Cost and contribute to totals.

Defect E. See audit_hazaway_depreciation_defect_e.md and
phase2_fix_depreciation_defect_e.md.

NOTE: _generate_depreciation_report reads `context["_fy"]` and
`context["_entity"]` (underscore prefix — see fs_template_service.py:3255-3256),
not the unprefixed names used by docxtpl templates.
"""
from datetime import date
from decimal import Decimal

from django.test import TestCase, override_settings
from docx import Document

from core.fs_template_service import _generate_depreciation_report
from core.models import Client, DepreciationAsset, Entity, FinancialYear


STORAGES_OVERRIDE = {
    "default": {"BACKEND": "django.core.files.storage.FileSystemStorage"},
    "staticfiles": {"BACKEND": "django.contrib.staticfiles.storage.StaticFilesStorage"},
}


@override_settings(STORAGES=STORAGES_OVERRIDE)
class DepreciationReportDisposalTests(TestCase):
    """Three asset cases: disposed in CY, fully-depreciated-owned, normally-depreciating."""

    @classmethod
    def setUpTestData(cls):
        cls.client_obj = Client.objects.create(name="Depreciation Test Client")
        cls.entity = Entity.objects.create(
            entity_name="Depreciation Test Co Pty Ltd",
            entity_type="company",
            client=cls.client_obj,
        )
        cls.fy = FinancialYear.objects.create(
            entity=cls.entity,
            year_label="FY2025",
            start_date=date(2024, 7, 1),
            end_date=date(2025, 6, 30),
        )
        # Asset 1: currently depreciating (matches Hazaway "Fuso Fighter" pattern)
        cls.normal = DepreciationAsset.objects.create(
            financial_year=cls.fy,
            category="General Pool",
            asset_name="Active Asset",
            total_cost=Decimal("100000.00"),
            opening_wdv=Decimal("60000.00"),
            depreciation_amount=Decimal("18000.00"),
            closing_wdv=Decimal("42000.00"),
            rate=Decimal("30.00"),
            disposal_date=None,
        )
        # Asset 2: fully depreciated but still owned (matches Hazaway "Ford Ranger" pattern)
        cls.fully_dep = DepreciationAsset.objects.create(
            financial_year=cls.fy,
            category="General Pool",
            asset_name="Fully Depreciated",
            total_cost=Decimal("50000.00"),
            opening_wdv=Decimal("0"),
            depreciation_amount=Decimal("0"),
            closing_wdv=Decimal("0"),
            rate=Decimal("30.00"),
            disposal_date=None,
        )
        # Asset 3: disposed during CY (matches Hazaway "Jayco / Porsche" pattern)
        cls.disposed = DepreciationAsset.objects.create(
            financial_year=cls.fy,
            category="General Pool",
            asset_name="Disposed Asset",
            total_cost=Decimal("80000.00"),
            opening_wdv=Decimal("40000.00"),
            depreciation_amount=Decimal("0"),
            closing_wdv=Decimal("0"),
            rate=Decimal("30.00"),
            disposal_date=date(2024, 9, 1),
            disposal_consideration=Decimal("45000.00"),
        )

    def _render_and_parse(self):
        """Render the report and return a parsed Document."""
        context = {"_fy": self.fy, "_entity": self.entity}
        buf = _generate_depreciation_report(context)
        self.assertIsNotNone(buf, "_generate_depreciation_report returned None")
        buf.seek(0)
        return Document(buf)

    def _find_row_cells(self, doc, asset_name):
        """Find the row whose first cell matches asset_name; return cell text list."""
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if cells and cells[0] == asset_name:
                    return cells
        self.fail(f"Row for {asset_name!r} not found in any table")

    def _find_subtotal_row(self, doc):
        """Locate the per-category Subtotals row (first cell starts with 'Subtotals')."""
        for table in doc.tables:
            for row in table.rows:
                first = row.cells[0].text.strip().lower()
                if first.startswith("subtotal"):
                    return [c.text.strip() for c in row.cells]
        self.fail("Subtotals row not found")

    def test_active_asset_keeps_total_cost(self):
        doc = self._render_and_parse()
        cells = self._find_row_cells(doc, "Active Asset")
        # Total Cost cell is index 1 per the body row construction
        self.assertIn(
            "100,000", cells[1].replace("$", ""),
            f"Active Asset Total Cost should be 100,000; got {cells[1]!r}",
        )

    def test_fully_depreciated_asset_keeps_total_cost(self):
        doc = self._render_and_parse()
        cells = self._find_row_cells(doc, "Fully Depreciated")
        self.assertIn(
            "50,000", cells[1].replace("$", ""),
            f"Fully Depreciated Total Cost should be 50,000; got {cells[1]!r}",
        )

    def test_disposed_asset_blanks_total_cost(self):
        doc = self._render_and_parse()
        cells = self._find_row_cells(doc, "Disposed Asset")
        self.assertEqual(
            cells[1].strip(), "",
            f"Disposed asset Total Cost must be blank; got {cells[1]!r}",
        )

    def test_subtotal_excludes_disposed_total_cost(self):
        doc = self._render_and_parse()
        cells = self._find_subtotal_row(doc)
        # Expected subtotal = 100,000 + 50,000 = 150,000 (80,000 disposed excluded)
        self.assertIn(
            "150,000", cells[1].replace("$", ""),
            f"Subtotal Total Cost should be 150,000; got {cells[1]!r}",
        )
