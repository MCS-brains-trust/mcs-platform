"""
Work Papers Tab — template download with pre-fill.

When an accountant clicks "Download" on a work paper template the system:
  1. Opens the uploaded Excel (.xlsx) or Word (.docx) template file.
  2. Substitutes the merge fields (entity_name, abn, financial_year, etc.)
     directly in the file in memory.
  3. Streams the modified file back as a download.

Nothing is saved to the database — the template is purely a launcher.
"""

import io
import logging
import os
import re

from django.contrib.auth.decorators import login_required
from django.http import FileResponse, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.http import require_GET

from core.models import WorkPaperTemplate

logger = logging.getLogger(__name__)


def _get_fy(request, pk):
    """Get FinancialYear with permission check."""
    from core.views import get_financial_year_for_user
    return get_financial_year_for_user(request, pk)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _build_prefill_context(entity, fy):
    """Return a dict of merge-field values for the given entity + financial year."""
    # Format ABN as XX XXX XXX XXX for readability
    raw_abn = (entity.abn or "").strip()
    if len(raw_abn) == 11:
        abn_display = f"{raw_abn[:2]} {raw_abn[2:5]} {raw_abn[5:8]} {raw_abn[8:11]}"
    else:
        abn_display = raw_abn

    return {
        "entity_name": entity.entity_name,
        "abn": abn_display,
        "financial_year": fy.year_label,
        "fy_start_date": fy.start_date.strftime("%-d %B %Y"),
        "fy_end_date": fy.end_date.strftime("%-d %B %Y"),
        "fy_year": str(fy.end_date.year),
        "prepared_by": "",  # left blank for the accountant to fill in
        "date_prepared": "",
    }


def _prefill_xlsx(template_file_path, context):
    """
    Open an Excel template, replace merge-field placeholders, and return the
    modified workbook as a BytesIO buffer.

    Merge fields can appear in any cell as plain text in the form:
      {{entity_name}}, {{abn}}, {{financial_year}}, etc.

    The function also writes to the first sheet's named ranges if they exist
    (e.g. a named range called "entity_name" will be populated directly).
    """
    import openpyxl

    wb = openpyxl.load_workbook(template_file_path)

    pattern = re.compile(r"\{\{(\w+)\}\}")

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.data_type == "s" and cell.value:
                    original = cell.value
                    def _replace(m):
                        key = m.group(1)
                        return str(context.get(key, m.group(0)))
                    new_value = pattern.sub(_replace, original)
                    if new_value != original:
                        cell.value = new_value

    # Also try to write to named ranges (if the template uses them)
    for range_name, value in context.items():
        if range_name in wb.defined_names:
            try:
                defn = wb.defined_names[range_name]
                for sheet_title, coord in defn.destinations:
                    if sheet_title in wb.sheetnames:
                        ws = wb[sheet_title]
                        ws[coord] = value
            except Exception:
                pass  # Named range may be complex; skip silently

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def _prefill_docx(template_file_path, context):
    """
    Open a Word template, replace {{merge_field}} placeholders using docxtpl,
    and return the modified document as a BytesIO buffer.
    """
    try:
        from docxtpl import DocxTemplate
        tpl = DocxTemplate(template_file_path)
        tpl.render(context)
        buf = io.BytesIO()
        tpl.save(buf)
        buf.seek(0)
        return buf
    except ImportError:
        # Fallback: use python-docx with simple text replacement
        from docx import Document
        doc = Document(template_file_path)
        pattern = re.compile(r"\{\{(\w+)\}\}")

        def _replace_runs(para):
            for run in para.runs:
                if run.text:
                    def _sub(m):
                        return str(context.get(m.group(1), m.group(0)))
                    run.text = pattern.sub(_sub, run.text)

        for para in doc.paragraphs:
            _replace_runs(para)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_runs(para)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf


# ---------------------------------------------------------------------------
# Views
# ---------------------------------------------------------------------------

@login_required
@require_GET
def workpaper_download(request, fy_pk, template_pk):
    """
    Download a work paper template pre-filled with entity/FY data.

    GET /years/<fy_pk>/workpapers/<template_pk>/download/
    """
    fy = _get_fy(request, fy_pk)
    entity = fy.entity

    template = get_object_or_404(WorkPaperTemplate, pk=template_pk, is_active=True)

    # Check entity-type restriction
    if template.entity_types and entity.entity_type not in template.entity_types:
        return HttpResponse("This template is not available for this entity type.", status=403)

    context = _build_prefill_context(entity, fy)

    try:
        file_path = template.template_file.path
    except (ValueError, FileNotFoundError):
        logger.error("WorkPaperTemplate %s has no file or file missing.", template_pk)
        return HttpResponse("Template file not found. Please contact your administrator.", status=404)

    fmt = template.file_format.lower()

    try:
        if fmt == "xlsx":
            buf = _prefill_xlsx(file_path, context)
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            ext = "xlsx"
        elif fmt == "docx":
            buf = _prefill_docx(file_path, context)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ext = "docx"
        else:
            # Unknown format — serve the raw file
            with open(file_path, "rb") as f:
                buf = io.BytesIO(f.read())
            content_type = "application/octet-stream"
            ext = os.path.splitext(file_path)[1].lstrip(".")

    except Exception as exc:
        logger.exception("Error pre-filling work paper template %s: %s", template_pk, exc)
        return HttpResponse(
            "An error occurred while preparing the template. Please try again or contact support.",
            status=500,
        )

    # Build a clean filename: "EntityName — BAS Reconciliation FY2025.xlsx"
    safe_entity = re.sub(r'[^\w\s\-]', '', entity.entity_name).strip()
    safe_name = re.sub(r'[^\w\s\-]', '', template.name).strip()
    filename = f"{safe_entity} — {safe_name} {fy.year_label}.{ext}"
    # RFC 5987 encoding for non-ASCII characters in Content-Disposition
    filename_ascii = filename.encode("ascii", "replace").decode("ascii").replace("?", "_")

    response = HttpResponse(buf.read(), content_type=content_type)
    response["Content-Disposition"] = (
        f'attachment; filename="{filename_ascii}"; '
        f"filename*=UTF-8''{filename.encode('utf-8').hex()}"
    )
    return response


@login_required
@require_GET
def workpaper_list_api(request, fy_pk):
    """
    Return JSON list of available work paper templates for a financial year.
    Used by the tab to dynamically render the template library.

    GET /api/years/<fy_pk>/workpapers/
    """
    fy = _get_fy(request, fy_pk)
    entity = fy.entity

    templates = WorkPaperTemplate.objects.filter(is_active=True)

    # Filter by entity type if restrictions are set
    result = []
    for tpl in templates:
        if tpl.entity_types and entity.entity_type not in tpl.entity_types:
            continue
        result.append({
            "id": str(tpl.pk),
            "name": tpl.name,
            "category": tpl.category,
            "category_display": tpl.get_category_display(),
            "description": tpl.description,
            "file_format": tpl.file_format,
            "download_url": f"/years/{fy_pk}/workpapers/{tpl.pk}/download/",
        })

    return JsonResponse({"templates": result})
