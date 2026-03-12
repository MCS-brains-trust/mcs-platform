"""
Eva BAS Period Commentary — Context Builder, Prompt Engineering, and Generation Logic

This module handles the generation of AI-powered period commentary for BAS periods,
transforming compliance data into client-ready advisory insights.

Five-section commentary structure:
  1. Period Snapshot (2-3 sentences, revenue headline)
  2. Revenue Analysis (2-3 significant movements)
  3. Cost & Margin Analysis (gross margin, opex trends)
  4. Items to Watch (max 4 items, plain English)
  5. Recommended Actions (1-3 specific actions)
"""
import json
import logging
import time
import traceback
from datetime import date
from decimal import Decimal

from django.conf import settings
from django.utils import timezone

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# System Prompt for BAS Period Commentary
# ---------------------------------------------------------------------------
BAS_COMMENTARY_SYSTEM_PROMPT = """You are Eva, the AI Practice Intelligence assistant for MC & S Accountants.
You are generating a BAS Period Commentary — a professional advisory document that
transforms raw BAS/GST compliance data into plain-English insights for the client.

═══════════════════════════════════════════════════════════════
DOCUMENT PURPOSE
═══════════════════════════════════════════════════════════════

This commentary accompanies the BAS lodgement and serves as a proactive advisory
touchpoint. It demonstrates the value of the accounting relationship beyond
compliance. The client is a business owner, NOT an accountant — write accordingly.

═══════════════════════════════════════════════════════════════
FIVE-SECTION STRUCTURE (MANDATORY)
═══════════════════════════════════════════════════════════════

You MUST produce exactly five sections in this order. Return your response as a
JSON object with these keys:

{
  "section_snapshot": "...",
  "section_revenue": "...",
  "section_costs": "...",
  "section_watch_items": "...",
  "section_actions": "..."
}

SECTION 1 — Period Snapshot
- 2-3 sentences maximum.
- Lead with the revenue headline figure and direction (up/down/flat vs prior period).
- Mention the net GST position (refund or payable) and the amount.
- Set the tone for the rest of the commentary.

SECTION 2 — Revenue Analysis
- Identify 2-3 significant revenue movements.
- Compare to prior period AND prior-year-same-period where data is available.
- Use dollar amounts and percentages.
- Explain what the movement MEANS for the business, not just that it happened.

SECTION 3 — Cost & Margin Analysis
- Comment on gross margin if cost of sales data is available.
- Identify the 2-3 largest operating expense movements.
- Flag any expense categories that have grown disproportionately to revenue.
- If expenses are flat but revenue grew, note the positive operating leverage.

SECTION 4 — Items to Watch
- Maximum 4 items.
- These are NOT compliance findings — they are business observations.
- Use plain English. No jargon. No legislation references.
- Examples: seasonal cash flow patterns, large one-off expenses, unusual
  creditor balances, GST refund positions that may trigger ATO review.
- This is the MOST VALUABLE section for the client. Make it count.

SECTION 5 — Recommended Actions
- 1-3 specific, actionable recommendations.
- Each action should have a clear owner (client, accountant, or both).
- Tie each recommendation back to something observed in the data.
- Include timing where relevant (e.g. "before the end of Q3", "at your next BAS review").

═══════════════════════════════════════════════════════════════
TONE AND STYLE
═══════════════════════════════════════════════════════════════

- {tone_instruction}
- Use Australian English spelling and conventions.
- Format monetary values as $X,XXX (no cents unless material).
- Use percentages with one decimal place (e.g. 12.3%).
- Do NOT use accounting jargon — translate everything into business language.
- Do NOT reference internal account codes or BAS label numbers (G1, G2, etc.).
- Do NOT mention Eva, StatementHub, or any internal systems.
- Write as if you are the accountant writing to their client.
- Address the client's business by name where natural.

═══════════════════════════════════════════════════════════════
DATA INTEGRITY
═══════════════════════════════════════════════════════════════

- ONLY reference figures that appear in the provided data context.
- If prior period or prior-year data is unavailable, acknowledge this briefly
  and focus on the current period.
- If the data is sparse (e.g. only a few transactions), keep the commentary
  proportionally brief. Do NOT pad with generic advice.
- Round figures to the nearest dollar for readability.
"""

TONE_INSTRUCTIONS = {
    "professional": (
        "Write in a professional, advisory tone. Use third person where possible. "
        "This reads like a letter from the firm to the client."
    ),
    "conversational": (
        "Write in a warm, conversational tone. Use 'you' and 'your business' freely. "
        "This reads like a trusted adviser speaking directly to the business owner."
    ),
    "technical": (
        "Write in a precise, technical tone suitable for financially literate clients. "
        "You may include more detail on GST positions and cash flow implications."
    ),
}


# ---------------------------------------------------------------------------
# Period-Scoped Context Builder
# ---------------------------------------------------------------------------
def build_period_context(financial_year, period_start, period_end, period_label=""):
    """
    Build the full data context for a BAS period commentary.

    Gathers:
    - Entity information
    - Current period GST data (from bas_utils.calculate_gst_for_period)
    - Prior period GST data (previous quarter/month)
    - Prior-year-same-period GST data (same quarter/month last year)
    - ATO benchmark data if available
    - Any existing Eva findings relevant to GST/BAS

    Returns:
        dict with structured context data
        str with formatted context for the LLM prompt
    """
    from core.models import (
        EvaReview, TrialBalanceLine, FinancialYear, BASPeriod,
    )
    from core.bas_utils import calculate_gst_for_period

    fy = financial_year
    entity = fy.entity

    context_data = {
        "entity_name": entity.entity_name,
        "entity_type": entity.get_entity_type_display(),
        "abn": entity.abn or "Not recorded",
        "financial_year": fy.year_label,
        "period_label": period_label,
        "period_start": str(period_start),
        "period_end": str(period_end),
        "is_gst_registered": entity.is_gst_registered,
    }

    sections = []

    # ── Entity Information ────────────────────────────────────────────
    sections.append(f"""=== ENTITY INFORMATION ===
Entity: {entity.entity_name}
Type: {entity.get_entity_type_display()}
ABN: {entity.abn or 'Not recorded'}
GST Registered: {'Yes' if entity.is_gst_registered else 'No'}
Financial Year: {fy.year_label}
Commentary Period: {period_label} ({period_start} to {period_end})
""")

    # ── Current Period GST Data ───────────────────────────────────────
    try:
        current_gst = calculate_gst_for_period(fy, period_start, period_end)
        bas_data = current_gst.get("bas_data", {})
        sales_lines = current_gst.get("sales_lines", [])
        purchase_lines = current_gst.get("purchase_lines", [])

        context_data["current_period"] = {
            "total_sales": str(bas_data.get("G1", 0)),
            "gst_on_sales": str(bas_data.get("1A", 0)),
            "total_purchases": str(bas_data.get("G11", 0)),
            "gst_on_purchases": str(bas_data.get("1B", 0)),
            "net_gst": str(bas_data.get("gst_payable", 0)),
            "export_sales": str(bas_data.get("G2", 0)),
            "gst_free_sales": str(bas_data.get("G3", 0)),
            "capital_purchases": str(bas_data.get("G10", 0)),
        }

        gst_section = [f"=== CURRENT PERIOD: {period_label} ==="]
        gst_section.append(f"Total Sales (G1): ${bas_data.get('G1', 0):,.0f}")
        gst_section.append(f"GST on Sales (1A): ${bas_data.get('1A', 0):,.0f}")
        gst_section.append(f"Total Purchases (G11): ${bas_data.get('G11', 0):,.0f}")
        gst_section.append(f"GST on Purchases (1B): ${bas_data.get('1B', 0):,.0f}")
        net_gst = bas_data.get("gst_payable", Decimal("0"))
        gst_section.append(f"Net GST: ${net_gst:,.0f} ({'Payable' if net_gst > 0 else 'Refund'})")

        if sales_lines:
            gst_section.append("\nRevenue Breakdown:")
            for line in sorted(sales_lines, key=lambda x: abs(x.get("amount", 0)), reverse=True)[:10]:
                gst_section.append(f"  {line.get('name', line.get('code', ''))}: ${line.get('amount', 0):,.0f}")

        if purchase_lines:
            gst_section.append("\nExpense Breakdown:")
            for line in sorted(purchase_lines, key=lambda x: abs(x.get("amount", 0)), reverse=True)[:10]:
                gst_section.append(f"  {line.get('name', line.get('code', ''))}: ${line.get('amount', 0):,.0f}")

        sections.append("\n".join(gst_section))
    except Exception as e:
        logger.error(f"BAS Commentary: current period GST error: {e}")
        sections.append(f"=== CURRENT PERIOD: {period_label} ===\nData unavailable: {e}")

    # ── Prior Period GST Data ─────────────────────────────────────────
    try:
        prior_period = _get_prior_period_dates(fy, period_start, period_end)
        if prior_period:
            pp_start, pp_end, pp_label = prior_period
            prior_gst = calculate_gst_for_period(fy, pp_start, pp_end)
            pp_bas = prior_gst.get("bas_data", {})
            pp_sales = prior_gst.get("sales_lines", [])
            pp_purchases = prior_gst.get("purchase_lines", [])

            context_data["prior_period"] = {
                "label": pp_label,
                "total_sales": str(pp_bas.get("G1", 0)),
                "gst_on_sales": str(pp_bas.get("1A", 0)),
                "total_purchases": str(pp_bas.get("G11", 0)),
                "gst_on_purchases": str(pp_bas.get("1B", 0)),
                "net_gst": str(pp_bas.get("gst_payable", 0)),
            }

            pp_section = [f"=== PRIOR PERIOD: {pp_label} ==="]
            pp_section.append(f"Total Sales (G1): ${pp_bas.get('G1', 0):,.0f}")
            pp_section.append(f"Total Purchases (G11): ${pp_bas.get('G11', 0):,.0f}")
            pp_section.append(f"Net GST: ${pp_bas.get('gst_payable', 0):,.0f}")

            if pp_sales:
                pp_section.append("\nRevenue Breakdown:")
                for line in sorted(pp_sales, key=lambda x: abs(x.get("amount", 0)), reverse=True)[:10]:
                    pp_section.append(f"  {line.get('name', line.get('code', ''))}: ${line.get('amount', 0):,.0f}")

            sections.append("\n".join(pp_section))
    except Exception as e:
        logger.warning(f"BAS Commentary: prior period data unavailable: {e}")

    # ── Prior Year Same Period ────────────────────────────────────────
    try:
        pysp = _get_prior_year_same_period(fy, period_start, period_end)
        if pysp:
            pysp_fy, pysp_start, pysp_end, pysp_label = pysp
            pysp_gst = calculate_gst_for_period(pysp_fy, pysp_start, pysp_end)
            pysp_bas = pysp_gst.get("bas_data", {})

            context_data["prior_year_same_period"] = {
                "label": pysp_label,
                "total_sales": str(pysp_bas.get("G1", 0)),
                "total_purchases": str(pysp_bas.get("G11", 0)),
                "net_gst": str(pysp_bas.get("gst_payable", 0)),
            }

            pysp_section = [f"=== PRIOR YEAR SAME PERIOD: {pysp_label} ==="]
            pysp_section.append(f"Total Sales (G1): ${pysp_bas.get('G1', 0):,.0f}")
            pysp_section.append(f"Total Purchases (G11): ${pysp_bas.get('G11', 0):,.0f}")
            pysp_section.append(f"Net GST: ${pysp_bas.get('gst_payable', 0):,.0f}")
            sections.append("\n".join(pysp_section))
    except Exception as e:
        logger.warning(f"BAS Commentary: prior year same period unavailable: {e}")

    # ── Eva Findings (BAS/GST related) ────────────────────────────────
    try:
        latest_review = EvaReview.objects.filter(
            financial_year=fy
        ).order_by("-triggered_at").first()
        if latest_review:
            gst_findings = latest_review.findings.filter(
                check_name__in=["gst_compliance", "ato_benchmarks", "bas_reconciliation"]
            )
            if gst_findings.exists():
                find_section = ["=== RELEVANT EVA FINDINGS ==="]
                for f in gst_findings:
                    find_section.append(
                        f"[{f.get_severity_display()}] {f.title}: {f.plain_english_explanation[:200]}"
                    )
                sections.append("\n".join(find_section))
    except Exception as e:
        logger.warning(f"BAS Commentary: Eva findings lookup error: {e}")

    formatted_context = "\n\n".join(sections)
    return context_data, formatted_context


def _get_prior_period_dates(fy, period_start, period_end):
    """
    Calculate the prior period dates (previous quarter or month).
    Returns (start_date, end_date, label) or None.
    """
    from dateutil.relativedelta import relativedelta

    # Determine period length
    period_days = (period_end - period_start).days
    if period_days > 80:  # Quarterly
        pp_end = period_start - relativedelta(days=1)
        pp_start = pp_end - relativedelta(months=3) + relativedelta(days=1)
        # Check if prior period is within the same FY
        if pp_start >= fy.start_date:
            quarter_num = (period_start.month - 7) // 3 + 1 if period_start.month >= 7 else (period_start.month + 5) // 3 + 1
            prior_q = quarter_num - 1
            if prior_q < 1:
                return None
            label = f"Q{prior_q} (Prior Quarter)"
            return pp_start, pp_end, label
    else:  # Monthly
        pp_end = period_start - relativedelta(days=1)
        pp_start = date(pp_end.year, pp_end.month, 1)
        if pp_start >= fy.start_date:
            import calendar
            label = f"{calendar.month_abbr[pp_start.month]} (Prior Month)"
            return pp_start, pp_end, label

    return None


def _get_prior_year_same_period(fy, period_start, period_end):
    """
    Find the prior year's financial year and calculate the same period dates.
    Returns (prior_fy, start_date, end_date, label) or None.
    """
    from core.models import FinancialYear
    from dateutil.relativedelta import relativedelta

    try:
        # Look for a FY that ended approximately one year before this one started
        prior_fy = FinancialYear.objects.filter(
            entity=fy.entity,
            end_date__lt=fy.start_date,
        ).order_by("-end_date").first()

        if not prior_fy:
            return None

        # Calculate the same period dates in the prior year
        pysp_start = period_start - relativedelta(years=1)
        pysp_end = period_end - relativedelta(years=1)

        # Verify the dates fall within the prior FY
        if pysp_start >= prior_fy.start_date and pysp_end <= prior_fy.end_date:
            label = f"Same Period {prior_fy.year_label}"
            return prior_fy, pysp_start, pysp_end, label
    except Exception as e:
        logger.warning(f"Prior year same period lookup failed: {e}")

    return None


# ---------------------------------------------------------------------------
# Commentary Generation (Background Task)
# ---------------------------------------------------------------------------

def _update_generation_step(commentary_pk, step):
    """Update the generation step in the database for progress polling."""
    from core.models import BASPeriodCommentary
    BASPeriodCommentary.objects.filter(pk=commentary_pk).update(generation_step=step)


def generate_bas_commentary(commentary_pk, user_pk):
    """
    Generate the BAS period commentary using the LLM.

    This is designed to be called from a background thread or Celery task.
    Progress is tracked via database fields on BASPeriodCommentary
    (generation_step, generation_started_at, generation_completed_at)
    so status survives server restarts and works across multiple workers.
    """
    import django
    django.setup()

    from core.models import BASPeriodCommentary, ActivityLog
    from core.ai_service import _call_llm
    from django.contrib.auth import get_user_model

    User = get_user_model()
    start_time = time.time()

    try:
        commentary = BASPeriodCommentary.objects.select_related(
            "financial_year", "financial_year__entity", "bas_period"
        ).get(pk=commentary_pk)
        user = User.objects.get(pk=user_pk)

        fy = commentary.financial_year
        entity = fy.entity

        # Mark generation as started
        commentary.generation_started_at = timezone.now()
        commentary.generation_step = "Building period context..."
        commentary.save(update_fields=["generation_started_at", "generation_step"])

        # Step 2: Build period-scoped context
        context_data, formatted_context = build_period_context(
            fy,
            commentary.period_start,
            commentary.period_end,
            commentary.period_label,
        )

        _update_generation_step(commentary_pk, "Querying Knowledge Brain...")

        # Step 3: Knowledge Brain retrieval
        kb_context = ""
        try:
            from core.eva_knowledge import retrieve_relevant_chunks, format_rag_context
            query = f"BAS GST period commentary {entity.get_entity_type_display()} {entity.entity_name}"
            chunks = retrieve_relevant_chunks(query, top_k=5)
            if chunks:
                kb_context = format_rag_context(chunks)
        except Exception as kb_err:
            logger.warning(f"BAS Commentary KB retrieval error: {kb_err}")

        _update_generation_step(commentary_pk, "Generating commentary...")

        # Step 4: Construct system prompt with tone
        tone_key = commentary.tone or "professional"
        tone_instruction = TONE_INSTRUCTIONS.get(tone_key, TONE_INSTRUCTIONS["professional"])
        system_prompt = BAS_COMMENTARY_SYSTEM_PROMPT.replace("{tone_instruction}", tone_instruction)

        # Step 5: Construct user prompt
        # Check for prior commentary for comparison context
        prior_commentary_text = ""
        try:
            prior = BASPeriodCommentary.objects.filter(
                financial_year=fy,
                status__in=["draft", "reviewed", "sent"],
                generated_at__lt=commentary.generated_at,
            ).order_by("-generated_at").first()
            if prior and prior.full_content:
                prior_commentary_text = (
                    f"\n=== PRIOR COMMENTARY ({prior.period_label}) ===\n"
                    f"{prior.full_content[:1500]}\n"
                    f"(Use this for continuity of narrative and to avoid repeating the same observations.)\n"
                )
        except Exception:
            pass

        user_prompt = f"""{formatted_context}

{kb_context}

{prior_commentary_text}

Generate the BAS Period Commentary for {entity.entity_name} covering {commentary.period_label}.

Respond with a JSON object containing exactly these five keys:
- section_snapshot
- section_revenue
- section_costs
- section_watch_items
- section_actions

Each value should be a string containing the section content in plain text (no markdown).
"""

        # Step 6: Call the LLM
        tier = "sonnet"
        response_text = _call_llm(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            tier=tier,
            temperature=0.3,
            max_tokens=3000,
        )

        # Step 7: Parse the JSON response
        result = _parse_commentary_json(response_text)

        # Step 8: Populate the five section fields
        commentary.section_snapshot = result.get("section_snapshot", "")
        commentary.section_revenue = result.get("section_revenue", "")
        commentary.section_costs = result.get("section_costs", "")
        commentary.section_watch_items = result.get("section_watch_items", "")
        commentary.section_actions = result.get("section_actions", "")

        # Step 9: Build the full_content field
        full_sections = []
        if commentary.section_snapshot:
            full_sections.append(f"**Period Snapshot**\n\n{commentary.section_snapshot}")
        if commentary.section_revenue:
            full_sections.append(f"**Revenue Analysis**\n\n{commentary.section_revenue}")
        if commentary.section_costs:
            full_sections.append(f"**Cost & Margin Analysis**\n\n{commentary.section_costs}")
        if commentary.section_watch_items:
            full_sections.append(f"**Items to Watch**\n\n{commentary.section_watch_items}")
        if commentary.section_actions:
            full_sections.append(f"**Recommended Actions**\n\n{commentary.section_actions}")

        commentary.full_content = "\n\n---\n\n".join(full_sections)

        # Step 10: Save context snapshot
        commentary.context_snapshot = context_data
        commentary.model_used = tier

        # Step 11: Update status and mark generation complete
        commentary.status = "draft"
        commentary.generation_completed_at = timezone.now()
        commentary.generation_step = ""
        commentary.save()

        duration = time.time() - start_time
        logger.info(f"BAS Commentary generated in {duration:.1f}s for {entity.entity_name} {commentary.period_label}")

        # Step 12: Log activity
        try:
            ActivityLog.objects.create(
                user=user,
                event_type="bas_commentary_generated",
                title=f"BAS Commentary Generated — {entity.entity_name}",
                description=(
                    f"Period commentary for {commentary.period_label} generated in {duration:.1f}s. "
                    f"Model: {tier}. Sections: {commentary.section_count}/5."
                ),
                entity=entity,
                financial_year=fy,
                url=f"/entities/years/{fy.pk}/",
            )
        except Exception:
            pass

        return True

    except Exception as e:
        logger.error(f"BAS Commentary generation error: {e}", exc_info=True)
        traceback.print_exc()

        try:
            BASPeriodCommentary.objects.filter(pk=commentary_pk).update(
                status="error",
                error_message=str(e)[:1000],
                generation_completed_at=timezone.now(),
                generation_step="",
            )
        except Exception:
            pass

        return False


def _parse_commentary_json(response_text):
    """Parse the LLM response as JSON, with fallback repair."""
    import re

    # Try direct parse
    try:
        # Find JSON object in the response
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            return json.loads(json_match.group())
    except json.JSONDecodeError:
        pass

    # Fallback: try to extract sections by pattern
    result = {}
    section_patterns = {
        "section_snapshot": r'"section_snapshot"\s*:\s*"((?:[^"\\]|\\.)*)"',
        "section_revenue": r'"section_revenue"\s*:\s*"((?:[^"\\]|\\.)*)"',
        "section_costs": r'"section_costs"\s*:\s*"((?:[^"\\]|\\.)*)"',
        "section_watch_items": r'"section_watch_items"\s*:\s*"((?:[^"\\]|\\.)*)"',
        "section_actions": r'"section_actions"\s*:\s*"((?:[^"\\]|\\.)*)"',
    }

    for key, pattern in section_patterns.items():
        match = re.search(pattern, response_text, re.DOTALL)
        if match:
            result[key] = match.group(1).replace('\\"', '"').replace('\\n', '\n')

    if not result:
        # Last resort: treat the entire response as the snapshot
        result["section_snapshot"] = response_text[:500]
        logger.warning("BAS Commentary: could not parse JSON, using raw text as snapshot")

    return result


# ---------------------------------------------------------------------------
# Word Document Generation
# ---------------------------------------------------------------------------
def generate_commentary_docx(commentary):
    """
    Generate a Word document (.docx) for the BAS period commentary.

    Returns the file path of the generated document.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import os
    import tempfile

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    entity = commentary.financial_year.entity

    # Title
    title = doc.add_heading(f'BAS Period Commentary', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x6B, 0x74)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(
        f'{entity.entity_name} — {commentary.period_label}'
    )
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Metadata table
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Shading'
    cells = [
        ("Entity", entity.entity_name),
        ("ABN", entity.abn or "Not recorded"),
        ("Period", f"{commentary.period_start} to {commentary.period_end}"),
        ("Prepared", commentary.generated_at.strftime("%d %B %Y") if commentary.generated_at else ""),
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        # Bold the label
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # Section helper
    def add_section(heading_text, content):
        if not content or not content.strip():
            return
        heading = doc.add_heading(heading_text, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0x0F, 0x6B, 0x74)

        for para_text in content.strip().split('\n\n'):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                p.paragraph_format.space_after = Pt(6)

    # Add sections
    add_section("Period Snapshot", commentary.section_snapshot)
    add_section("Revenue Analysis", commentary.section_revenue)
    add_section("Cost & Margin Analysis", commentary.section_costs)
    add_section("Items to Watch", commentary.section_watch_items)
    add_section("Recommended Actions", commentary.section_actions)

    # Footer disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "This commentary is prepared by MC & S Accountants based on the financial data "
        "available at the time of preparation. It is intended as general guidance and does "
        "not constitute financial advice. Please contact your accountant to discuss any items "
        "raised in this commentary."
    )
    disclaimer_run.font.size = Pt(8)
    disclaimer_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    disclaimer_run.italic = True

    # Save to temp file
    temp_dir = os.path.join(settings.MEDIA_ROOT, "bas_commentaries")
    os.makedirs(temp_dir, exist_ok=True)

    filename = (
        f"bas_commentary_{entity.entity_name.replace(' ', '_')}_"
        f"{commentary.period_label.replace(' ', '_').replace('(', '').replace(')', '')}_"
        f"v{commentary.version}.docx"
    )
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)

    return filepath
