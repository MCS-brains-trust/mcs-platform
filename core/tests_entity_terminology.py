"""A unit trust has unit holders; a discretionary trust has beneficiaries.

The resolver (core/entity_terminology.py) is unit-tested directly against
bare Entity instances. The view-level tests below guard the PRIMARY RISK
called out in Task 10: a discretionary trust must keep seeing the literal
word "Beneficiary" everywhere, unchanged, after this task's edits — only a
unit trust (`entity_type="trust_unit"`) should ever see "Unit Holder".
"""
from datetime import date
from decimal import Decimal

from django.contrib.auth import get_user_model
from django.test import TestCase
from django.urls import reverse

from core.entity_terminology import beneficiary_noun
from core.models import (
    Entity,
    EntityChartOfAccount,
    EntityOfficer,
    FinancialYear,
    TrialBalanceLine,
)


class BeneficiaryNounTests(TestCase):
    def test_unit_trust(self):
        e = Entity(entity_name="Minli", entity_type="trust_unit")
        self.assertEqual(beneficiary_noun(e), "Unit Holder")
        self.assertEqual(beneficiary_noun(e, plural=True), "Unit Holders")

    def test_discretionary_trust(self):
        e = Entity(entity_name="Vincent", entity_type="trust")
        self.assertEqual(beneficiary_noun(e), "Beneficiary")
        self.assertEqual(beneficiary_noun(e, plural=True), "Beneficiaries")

    def test_non_trust_entity_falls_back_to_beneficiary(self):
        e = Entity(entity_name="Acme Pty Ltd", entity_type="company")
        self.assertEqual(beneficiary_noun(e), "Beneficiary")


def _give_distributable_income(fy, amount):
    """Tag one revenue account and post a TB line so tax_planning_tab's
    TB-driven recalculation produces a known, non-zero distributable_income.
    """
    EntityChartOfAccount.objects.create(
        entity=fy.entity, account_code="400", account_name="Sales",
        section=EntityChartOfAccount.StatementSection.REVENUE,
    )
    TrialBalanceLine.objects.create(
        financial_year=fy, account_code="400", account_name="Sales",
        debit=Decimal("0.00"), credit=amount, closing_balance=amount,
    )


class TerminologyOnScreenTests(TestCase):
    """Officer tab and tax planning tab say the right noun for each type,
    and a discretionary trust's wording is byte-for-byte unchanged.
    """

    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="term_test_user", email="term@example.com",
            password="secret123",
            totp_secret="dummy-secret-for-test", totp_confirmed=True,
        )
        self.client.force_login(self.user)
        session = self.client.session
        session["2fa_verified"] = True
        session.save()

        self.unit_trust = Entity.objects.create(
            entity_name="Minli Enterprise Unit Trust", entity_type="trust_unit",
            assigned_accountant=self.user,
        )
        EntityOfficer.objects.create(
            entity=self.unit_trust, full_name="A Pty Ltd",
            role=EntityOfficer.OfficerRole.UNIT_HOLDER,
            roles=["unit_holder"], units_held=100,
        )

        self.disc_trust = Entity.objects.create(
            entity_name="Vincent Family Trust", entity_type="trust",
            assigned_accountant=self.user,
        )
        EntityOfficer.objects.create(
            entity=self.disc_trust, full_name="John Vincent",
            role=EntityOfficer.OfficerRole.BENEFICIARY,
            roles=["beneficiary"], distribution_percentage=Decimal("100.00"),
        )

    def test_unit_trust_officer_tab_says_unit_holder(self):
        response = self.client.get(
            reverse("core:entity_officers", kwargs={"pk": self.unit_trust.pk}),
            secure=True,
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn("Unit Holder", content)
        self.assertNotIn("Trustees & Beneficiaries", content)

    def test_discretionary_trust_officer_tab_still_says_beneficiary(self):
        response = self.client.get(
            reverse("core:entity_officers", kwargs={"pk": self.disc_trust.pk}),
            secure=True,
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn("Trustees & Beneficiaries", content)
        self.assertIn("Directors/Trustees/Beneficiaries", content)
        self.assertNotIn("Unit Holder", content)

    def test_discretionary_trust_officer_label_unchanged(self):
        response = self.client.get(
            reverse("core:entity_officers", kwargs={"pk": self.disc_trust.pk}),
            secure=True,
        )
        self.assertContains(response, "Trustee / Beneficiary")

    def test_unit_trust_officer_label_says_unit_holder(self):
        response = self.client.get(
            reverse("core:entity_officers", kwargs={"pk": self.unit_trust.pk}),
            secure=True,
        )
        self.assertContains(response, "Trustee / Unit Holder")

    def test_unit_trust_tax_planning_tab_says_unit_holder(self):
        fy = FinancialYear.objects.create(
            entity=self.unit_trust, year_label="FY2026",
            start_date=date(2025, 7, 1), end_date=date(2026, 6, 30),
        )
        _give_distributable_income(fy, Decimal("100000.00"))
        response = self.client.get(
            reverse("core:tax_planning_tab", kwargs={"pk": fy.pk}),
            secure=True,
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn("Unit Holder Tax Position", content)

    def test_discretionary_trust_tax_planning_tab_still_says_beneficiary(self):
        fy = FinancialYear.objects.create(
            entity=self.disc_trust, year_label="FY2026",
            start_date=date(2025, 7, 1), end_date=date(2026, 6, 30),
        )
        _give_distributable_income(fy, Decimal("100000.00"))
        response = self.client.get(
            reverse("core:tax_planning_tab", kwargs={"pk": fy.pk}),
            secure=True,
        )
        self.assertEqual(response.status_code, 200)
        content = response.content.decode()
        self.assertIn("Beneficiary Tax Position", content)
        self.assertNotIn("Unit Holder", content)


class EntityOfficersViewLabelMapTests(TestCase):
    """core/views.py:entity_officers no longer falls back to the generic
    "Officer" label for a unit trust.
    """

    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="label_map_test_user", email="labelmap@example.com",
            password="secret123",
            totp_secret="dummy-secret-for-test", totp_confirmed=True,
        )
        self.client.force_login(self.user)
        session = self.client.session
        session["2fa_verified"] = True
        session.save()

    def test_unit_trust_add_officer_button_says_unit_holder(self):
        entity = Entity.objects.create(
            entity_name="Standalone Unit Trust", entity_type="trust_unit",
            assigned_accountant=self.user,
        )
        response = self.client.get(
            reverse("core:entity_officers", kwargs={"pk": entity.pk}),
            secure=True,
        )
        self.assertContains(response, "Add Trustee / Unit Holder")


class TaxplanDocgenHeaderTests(TestCase):
    """The Word-document table headers in core/taxplan_docgen.py say the
    right noun too (Task 10 sites 5-6). The discretionary-specific prose
    elsewhere in this module is explicitly out of scope and untouched.
    """

    def _fy_with_officer(self, entity_type, role, roles, full_name="A Holder", **extra):
        entity = Entity.objects.create(
            entity_name="Docgen Test Entity", entity_type=entity_type,
        )
        EntityOfficer.objects.create(
            entity=entity, full_name="Trudy Trustee", role="trustee", roles=["trustee"],
        )
        EntityOfficer.objects.create(
            entity=entity, full_name=full_name, role=role, roles=roles, **extra,
        )
        fy = FinancialYear.objects.create(
            entity=entity, year_label="FY2026",
            start_date=date(2025, 7, 1), end_date=date(2026, 6, 30),
        )
        return fy

    def _header_row_text(self, buf):
        from docx import Document
        doc = Document(buf)
        return [cell.text for cell in doc.tables[0].rows[0].cells]

    def _six_col_header_row_text(self, buf):
        """The Section 2 table is always 6 columns wide; its position among
        doc.tables shifts depending on whether Section 1's table exists.
        """
        from docx import Document
        doc = Document(buf)
        table = next(t for t in doc.tables if len(t.columns) == 6)
        return [cell.text for cell in table.rows[0].cells]

    def test_trust_election_header_says_unit_holder_for_a_unit_trust(self):
        from core.taxplan_docgen import generate_trust_election
        fy = self._fy_with_officer(
            "trust_unit", EntityOfficer.OfficerRole.UNIT_HOLDER, ["unit_holder"],
            units_held=100,
        )
        buf = generate_trust_election(fy.pk)
        self.assertEqual(self._header_row_text(buf)[0], "Unit Holder")

    def test_trust_election_header_still_says_beneficiary_for_a_discretionary_trust(self):
        from core.taxplan_docgen import generate_trust_election
        fy = self._fy_with_officer(
            "trust", EntityOfficer.OfficerRole.BENEFICIARY, ["beneficiary"],
            distribution_percentage=Decimal("100.00"),
        )
        buf = generate_trust_election(fy.pk)
        self.assertEqual(self._header_row_text(buf)[0], "Beneficiary")

    def test_tax_planning_summary_header_says_unit_holder_for_a_unit_trust(self):
        from core.taxplan_docgen import generate_tax_planning_summary
        fy = self._fy_with_officer(
            "trust_unit", EntityOfficer.OfficerRole.UNIT_HOLDER, ["unit_holder"],
            units_held=100,
        )
        buf = generate_tax_planning_summary(fy.pk)
        self.assertEqual(self._six_col_header_row_text(buf)[0], "Unit Holder")

    def test_tax_planning_summary_header_still_says_beneficiary_for_a_discretionary_trust(self):
        from core.taxplan_docgen import generate_tax_planning_summary
        fy = self._fy_with_officer(
            "trust", EntityOfficer.OfficerRole.BENEFICIARY, ["beneficiary"],
            distribution_percentage=Decimal("100.00"),
        )
        buf = generate_tax_planning_summary(fy.pk)
        self.assertEqual(self._six_col_header_row_text(buf)[0], "Beneficiary")
