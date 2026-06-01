"""
Trust Distribution Tab Views
=============================

Handles the 5-stage trust distribution workflow:
  Stage 1: Income Calculation
  Stage 2: Distribution Modelling (select Tax Planning scenario + post to TB)
  Stage 3: Section 100A Assessment
  Stage 4: Trust Elections
  Stage 5: Documents
"""

import json
import logging
from decimal import Decimal, InvalidOperation

from django.contrib.auth.decorators import login_required
from django.db import models as django_models, transaction
from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
from django.views.decorators.http import require_POST

from core.models import (
    FinancialYear, TrustWorkspace, BeneficiaryProfile,
    DistributionScenario, Section100AAssessment, TrustElectionRecord,
    EntityOfficer, ActivityLog, TaxPlanningScenario,
    AdjustingJournal, JournalLine,
)

logger = logging.getLogger(__name__)
ZERO = Decimal("0")


# ---------------------------------------------------------------------------
# Workspace Management
# ---------------------------------------------------------------------------
@login_required
def trust_workspace_api(request, pk):
    """
    GET  /api/years/<pk>/trust-workspace/ — Get or create workspace
    POST /api/years/<pk>/trust-workspace/ — Update workspace fields
    """
    fy = get_object_or_404(FinancialYear, pk=pk)

    if request.method == "GET":
        workspace, created = TrustWorkspace.objects.get_or_create(
            financial_year=fy,
        )
        if created:
            # Auto-populate income streams from TB
            _auto_populate_income(workspace)
            # Auto-create beneficiary profiles from officers
            _auto_create_beneficiary_profiles(workspace)

        # Auto-skip Stage 3 (Section 100A) for unit trusts / trusts with unitholders
        if _entity_has_unitholders(fy.entity):
            if workspace.stage_3_status != TrustWorkspace.StageStatus.COMPLETED:
                workspace.stage_3_status = TrustWorkspace.StageStatus.COMPLETED
                workspace.save(update_fields=["stage_3_status", "updated_at"])

        return JsonResponse(_serialize_workspace(workspace))

    elif request.method == "POST":
        workspace = get_object_or_404(TrustWorkspace, financial_year=fy)
        data = json.loads(request.body)

        # Update income streams if provided
        if "income_streams" in data:
            workspace.income_streams = data["income_streams"]
        if "net_distributable_income" in data:
            try:
                workspace.net_distributable_income = Decimal(str(data["net_distributable_income"]))
            except (InvalidOperation, ValueError):
                pass

        workspace.save()
        return JsonResponse(_serialize_workspace(workspace))


@login_required
def trust_stage_update(request, pk, stage_num):
    """
    POST /api/years/<pk>/trust-workspace/stage/<stage_num>/
    Update a specific stage's status.
    """
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace = get_object_or_404(TrustWorkspace, financial_year=fy)
    data = json.loads(request.body)
    new_status = data.get("status", "")

    if new_status not in dict(TrustWorkspace.StageStatus.choices):
        return JsonResponse({"error": "Invalid status"}, status=400)

    if stage_num < 1 or stage_num > 6:
        return JsonResponse({"error": "Invalid stage number"}, status=400)

    field_name = f"stage_{stage_num}_status"
    setattr(workspace, field_name, new_status)
    workspace.save(update_fields=[field_name, "updated_at"])

    stage_names = {
        1: "Income Calculation", 2: "Distribution Modelling",
        3: "Section 100A Assessment", 4: "Trust Elections",
        5: "Documents",
    }

    ActivityLog.objects.create(
        user=request.user,
        event_type="trust_stage_update",
        title=f"Trust Stage {stage_num}: {stage_names.get(stage_num, '')} → {new_status}",
        description=f"Updated stage {stage_num} ({stage_names.get(stage_num, '')}) to {new_status}",
        entity=fy.entity,
        financial_year=fy,
        url=f"/entities/years/{fy.pk}/?tab=trust",
    )

    return JsonResponse({
        "status": "ok",
        "stage": stage_num,
        "new_status": new_status,
        "all_completed": workspace.all_stages_completed(),
    })


# ---------------------------------------------------------------------------
# Stage 2: Beneficiary Profiles
# ---------------------------------------------------------------------------
@login_required
def beneficiary_profiles_api(request, pk):
    """
    GET  — List all beneficiary profiles for this workspace
    POST — Update a beneficiary profile
    """
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace = get_object_or_404(TrustWorkspace, financial_year=fy)

    if request.method == "GET":
        profiles = workspace.beneficiary_profiles.select_related("beneficiary").all()
        return JsonResponse({
            "profiles": [_serialize_beneficiary_profile(p) for p in profiles],
        })

    elif request.method == "POST":
        data = json.loads(request.body)
        profile_id = data.get("id")
        if not profile_id:
            return JsonResponse({"error": "Profile ID required"}, status=400)

        profile = get_object_or_404(BeneficiaryProfile, pk=profile_id, trust_workspace=workspace)

        for field in ["beneficiary_type", "other_income", "marginal_rate",
                       "bracket_remaining", "franking_surplus", "include_in_distribution",
                       "exclusion_reason", "tax_residency"]:
            if field in data:
                val = data[field]
                if field in ("other_income", "marginal_rate", "bracket_remaining", "franking_surplus"):
                    try:
                        val = Decimal(str(val)) if val not in (None, "", "null") else None
                    except (InvalidOperation, ValueError):
                        val = None
                setattr(profile, field, val)

        profile.save()
        return JsonResponse(_serialize_beneficiary_profile(profile))


# ---------------------------------------------------------------------------
# Stage 2: Tax Planning Scenario Selection
# ---------------------------------------------------------------------------
@login_required
def tax_planning_scenarios_api(request, pk):
    """
    GET  — List TaxPlanningScenarios for this FY for Stage 2 display.
    POST — Select a scenario for distribution posting.
    """
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace, _ = TrustWorkspace.objects.get_or_create(financial_year=fy)

    if request.method == "POST":
        data = json.loads(request.body)
        scenario_id = data.get("scenario_id")
        if not scenario_id:
            return JsonResponse({"error": "scenario_id required"}, status=400)
        scenario = get_object_or_404(TaxPlanningScenario, pk=scenario_id, financial_year=fy)
        workspace.selected_tax_scenario = scenario
        workspace.save(update_fields=["selected_tax_scenario", "updated_at"])
        return JsonResponse({"status": "ok", "selected_tax_scenario_id": str(scenario.pk)})

    # GET — return all scenarios with beneficiary name resolution
    scenarios = TaxPlanningScenario.objects.filter(financial_year=fy).order_by("created_at")
    officer_map = {
        str(o.pk): o.full_name
        for o in EntityOfficer.objects.filter(entity=fy.entity)
    }

    result = []
    for sc in scenarios:
        distributions = []
        for entry in (sc.distributions or []):
            ben_id = str(entry.get("beneficiary_id", ""))
            distributions.append({
                "beneficiary_id": ben_id,
                "beneficiary_name": officer_map.get(ben_id, f"Unknown ({ben_id[:8]})"),
                "proposed_distribution": str(entry.get("proposed_distribution", 0)),
            })
        result.append({
            "id": str(sc.pk),
            "scenario_name": sc.scenario_name,
            "total_distributed": str(sc.total_distributed),
            "total_tax": str(sc.total_tax),
            "distributions": distributions,
        })

    # Gate: a live (posted, non-voided) distribution journal hides the post
    # button. Keyed off the structural flag — a voided journal restores it.
    existing_journal = AdjustingJournal.live_trust_distribution(fy)

    return JsonResponse({
        "scenarios": result,
        "selected_tax_scenario_id": str(workspace.selected_tax_scenario_id) if workspace.selected_tax_scenario_id else None,
        "distribution_journal_exists": existing_journal is not None,
        "distribution_journal_ref": existing_journal.reference_number if existing_journal else None,
    })


# ---------------------------------------------------------------------------
# Stage 3 (old): Distribution Scenarios (workspace-local, kept for compat)
# ---------------------------------------------------------------------------
@login_required
def distribution_scenarios_api(request, pk):
    """
    GET  — List all scenarios
    POST — Create or update a scenario
    """
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace = get_object_or_404(TrustWorkspace, financial_year=fy)

    if request.method == "GET":
        scenarios = workspace.scenarios.all()
        return JsonResponse({
            "scenarios": [_serialize_scenario(s) for s in scenarios],
        })

    elif request.method == "POST":
        data = json.loads(request.body)
        scenario_id = data.get("id")

        with transaction.atomic():
            if scenario_id:
                scenario = get_object_or_404(DistributionScenario, pk=scenario_id, trust_workspace=workspace)
            else:
                # Lock workspace row to prevent race condition on count check
                TrustWorkspace.objects.select_for_update().get(pk=workspace.pk)
                if workspace.scenarios.count() >= 3:
                    return JsonResponse({"error": "Maximum 3 scenarios allowed"}, status=400)
                scenario = DistributionScenario(trust_workspace=workspace)

            if "name" in data:
                scenario.name = data["name"]
            if "allocations" in data:
                scenario.allocations = data["allocations"]
            if "total_tax" in data:
                try:
                    scenario.total_tax = Decimal(str(data["total_tax"]))
                except (InvalidOperation, ValueError):
                    pass
            if "tax_saved_vs_equal" in data:
                try:
                    scenario.tax_saved_vs_equal = Decimal(str(data["tax_saved_vs_equal"]))
                except (InvalidOperation, ValueError):
                    pass

            scenario.save()
        return JsonResponse(_serialize_scenario(scenario))


@login_required
@require_POST
def confirm_scenario(request, pk, scenario_pk):
    """POST — Confirm a scenario as the final distribution."""
    fy = get_object_or_404(FinancialYear, pk=pk)

    with transaction.atomic():
        workspace = TrustWorkspace.objects.select_for_update().get(financial_year=fy)
        scenario = get_object_or_404(DistributionScenario, pk=scenario_pk, trust_workspace=workspace)

        # Unconfirm all others
        workspace.scenarios.select_for_update().update(is_confirmed=False)
        scenario.is_confirmed = True
        scenario.save(update_fields=["is_confirmed"])

        workspace.confirmed_scenario = scenario
        workspace.save(update_fields=["confirmed_scenario"])

    ActivityLog.objects.create(
        user=request.user,
        event_type="trust_scenario_confirmed",
        title=f"Distribution scenario '{scenario.name}' confirmed",
        description=f"Confirmed '{scenario.name}' as the final distribution for {fy.entity.entity_name}",
        entity=fy.entity,
        financial_year=fy,
        url=f"/entities/years/{fy.pk}/?tab=trust",
    )

    return JsonResponse({"status": "ok", "scenario_id": str(scenario.pk)})


@login_required
@require_POST
def delete_scenario(request, pk, scenario_pk):
    """POST — Delete a distribution scenario."""
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace = get_object_or_404(TrustWorkspace, financial_year=fy)
    scenario = get_object_or_404(DistributionScenario, pk=scenario_pk, trust_workspace=workspace)

    if scenario.is_confirmed:
        workspace.confirmed_scenario = None
        workspace.save(update_fields=["confirmed_scenario"])

    scenario.delete()
    return JsonResponse({"status": "ok"})


# ---------------------------------------------------------------------------
# Stage 4: Section 100A Assessments
# ---------------------------------------------------------------------------
@login_required
def section_100a_api(request, pk):
    """
    GET  — List all Section 100A assessments
    POST — Update an assessment
    """
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace = get_object_or_404(TrustWorkspace, financial_year=fy)

    if request.method == "GET":
        assessments = workspace.section_100a_assessments.select_related("beneficiary").all()

        # Filter to only beneficiaries with distribution > 0 in selected scenario
        if workspace.selected_tax_scenario:
            active_ben_ids = set()
            for entry in (workspace.selected_tax_scenario.distributions or []):
                amt = Decimal(str(entry.get("proposed_distribution", 0)))
                if amt > 0:
                    active_ben_ids.add(str(entry.get("beneficiary_id", "")))
            if active_ben_ids:
                assessments = [a for a in assessments if str(a.beneficiary_id) in active_ben_ids]

        return JsonResponse({
            "assessments": [_serialize_100a(a) for a in assessments],
        })

    elif request.method == "POST":
        data = json.loads(request.body)
        assessment_id = data.get("id")

        with transaction.atomic():
            if assessment_id:
                assessment = get_object_or_404(
                    Section100AAssessment, pk=assessment_id, trust_workspace=workspace
                )
            else:
                beneficiary_id = data.get("beneficiary_id")
                beneficiary = get_object_or_404(EntityOfficer, pk=beneficiary_id)
                assessment, _ = Section100AAssessment.objects.get_or_create(
                    trust_workspace=workspace, beneficiary=beneficiary,
                )

            for q in ["q1", "q2", "q3", "q4", "q5", "q6", "q7", "q8"]:
                if q in data:
                    setattr(assessment, q, data[q])

            if "resolution_strategy" in data:
                assessment.resolution_strategy = data["resolution_strategy"]

            assessment.save()  # risk_rating calculated in save()

            # Update overall workspace risk (inside transaction for consistency)
            _update_overall_100a_risk(workspace)

        return JsonResponse(_serialize_100a(assessment))


# ---------------------------------------------------------------------------
# Stage 5: Trust Elections
# ---------------------------------------------------------------------------
@login_required
def trust_elections_api(request, pk):
    """
    GET  — List all election records
    POST — Update an election record
    """
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace = get_object_or_404(TrustWorkspace, financial_year=fy)

    if request.method == "GET":
        elections = workspace.election_records.select_related(
            "test_individual", "related_entity"
        ).all()
        return JsonResponse({
            "elections": [_serialize_election(e) for e in elections],
        })

    elif request.method == "POST":
        data = json.loads(request.body)
        election_id = data.get("id")

        if election_id:
            election = get_object_or_404(
                TrustElectionRecord, pk=election_id, trust_workspace=workspace
            )
        else:
            election = TrustElectionRecord(trust_workspace=workspace)

        if "election_type" in data:
            election.election_type = data["election_type"]
        if "status" in data:
            election.status = data["status"]
        if "effective_date" in data and data["effective_date"]:
            election.effective_date = data["effective_date"]
        if "test_individual_id" in data and data["test_individual_id"]:
            election.test_individual_id = data["test_individual_id"]
        if "related_entity_id" in data and data["related_entity_id"]:
            election.related_entity_id = data["related_entity_id"]

        election.save()
        return JsonResponse(_serialize_election(election))


@login_required
@require_POST
def confirm_election(request, pk, election_pk):
    """POST — Confirm an election record."""
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace = get_object_or_404(TrustWorkspace, financial_year=fy)
    election = get_object_or_404(
        TrustElectionRecord, pk=election_pk, trust_workspace=workspace
    )

    election.confirmed_by = request.user
    election.confirmed_at = timezone.now()
    election.save(update_fields=["confirmed_by", "confirmed_at"])

    return JsonResponse({"status": "ok"})


# ---------------------------------------------------------------------------
# Eva Context for Trust Tab
# ---------------------------------------------------------------------------
@login_required
def trust_eva_context(request, pk):
    """
    GET /api/years/<pk>/trust-workspace/eva-context/
    Provides trust-specific context for Eva's Finalisation Gate.
    """
    fy = get_object_or_404(FinancialYear, pk=pk)

    try:
        workspace = TrustWorkspace.objects.get(financial_year=fy)
    except TrustWorkspace.DoesNotExist:
        return JsonResponse({"error": "No trust workspace found"}, status=404)

    context = {
        "workspace_status": {
            "stage_1": workspace.stage_1_status,
            "stage_2": workspace.stage_2_status,
            "stage_3": workspace.stage_3_status,
            "stage_4": workspace.stage_4_status,
            "stage_5": workspace.stage_5_status,
            "stage_6": workspace.stage_6_status,
            "all_completed": workspace.all_stages_completed(),
        },
        "income": {
            "ndi": str(workspace.net_distributable_income or 0),
            "streams": workspace.income_streams,
        },
        "confirmed_scenario": None,
        "section_100a_risk": workspace.section_100a_overall_risk,
        "beneficiary_count": workspace.beneficiary_profiles.count(),
        "elections": [],
    }

    if workspace.confirmed_scenario:
        s = workspace.confirmed_scenario
        context["confirmed_scenario"] = {
            "name": s.name,
            "allocations": s.allocations,
            "total_tax": str(s.total_tax) if s.total_tax else None,
        }

    for e in workspace.election_records.all():
        context["elections"].append({
            "type": e.get_election_type_display(),
            "status": e.get_status_display(),
            "confirmed": e.confirmed_at is not None,
        })

    return JsonResponse(context)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
@login_required
@require_POST
def trust_recalculate_income(request, pk):
    fy = get_object_or_404(FinancialYear, pk=pk)
    workspace, _ = TrustWorkspace.objects.get_or_create(financial_year=fy)
    if workspace.stage_1_status == TrustWorkspace.StageStatus.COMPLETED:
        return JsonResponse({"error": "Stage 1 is already completed. Reopen it before recalculating."}, status=400)
    _auto_populate_income(workspace)
    return JsonResponse(_serialize_workspace(workspace))


def _auto_populate_income(workspace):
    """Auto-populate income streams from trial balance."""
    from core.eva_trust_planning import _calculate_income_streams
    income_data = _calculate_income_streams(workspace.financial_year)
    workspace.income_streams = income_data["income_streams"]
    workspace.net_distributable_income = Decimal(income_data["net_distributable_income"])
    workspace.save(update_fields=["income_streams", "net_distributable_income"])


def _auto_create_beneficiary_profiles(workspace):
    """Create beneficiary profiles from entity officers."""
    entity = workspace.financial_year.entity
    officers = EntityOfficer.objects.filter(entity=entity)

    for officer in officers:
        BeneficiaryProfile.objects.get_or_create(
            trust_workspace=workspace,
            beneficiary=officer,
            defaults={
                "beneficiary_type": _map_officer_to_beneficiary_type(officer),
                "tax_residency": getattr(officer, "tax_residency", "AU") or "AU",
            },
        )


def _map_officer_to_beneficiary_type(officer):
    """Map an EntityOfficer to a BeneficiaryProfile type."""
    bt = getattr(officer, "beneficiary_type", "")
    mapping = {
        "adult": BeneficiaryProfile.BeneficiaryType.ADULT,
        "minor": BeneficiaryProfile.BeneficiaryType.MINOR,
        "company": BeneficiaryProfile.BeneficiaryType.COMPANY,
        "trust": BeneficiaryProfile.BeneficiaryType.TRUST,
        "smsf": BeneficiaryProfile.BeneficiaryType.SMSF,
    }
    return mapping.get(bt, BeneficiaryProfile.BeneficiaryType.ADULT)


def _update_overall_100a_risk(workspace):
    """Update the workspace's overall Section 100A risk rating."""
    assessments = workspace.section_100a_assessments.all()
    if not assessments:
        workspace.section_100a_overall_risk = ""
        workspace.save(update_fields=["section_100a_overall_risk"])
        return

    ratings = [a.risk_rating for a in assessments]
    if "red" in ratings:
        overall = "red"
    elif "amber" in ratings:
        overall = "amber"
    else:
        overall = "green"

    workspace.section_100a_overall_risk = overall
    workspace.save(update_fields=["section_100a_overall_risk"])


def _entity_has_unitholders(entity):
    """Return True if the entity is a unit trust or has any unit holder officers."""
    if entity.entity_type == 'trust_unit':
        return True
    return entity.officers.filter(
        django_models.Q(role='unit_holder') | django_models.Q(roles__contains='unit_holder')
    ).exists()


def _serialize_workspace(workspace):
    """Serialize a TrustWorkspace to JSON."""
    entity = workspace.financial_year.entity
    has_unitholders = _entity_has_unitholders(entity)
    return {
        "id": str(workspace.pk),
        "financial_year_id": str(workspace.financial_year_id),
        "entity_type": entity.entity_type,
        "has_unitholders": has_unitholders,
        "stages": {
            "1": {"status": workspace.stage_1_status, "name": "Income Calculation"},
            "2": {"status": workspace.stage_2_status, "name": "Distribution Modelling"},
            "3": {"status": workspace.stage_3_status, "name": "Section 100A Assessment"},
            "4": {"status": workspace.stage_4_status, "name": "Trust Elections"},
            "5": {"status": workspace.stage_5_status, "name": "Documents"},
        },
        "all_completed": workspace.all_stages_completed(),
        "net_distributable_income": str(workspace.net_distributable_income or 0),
        "income_streams": workspace.income_streams,
        "section_100a_overall_risk": workspace.section_100a_overall_risk,
        "confirmed_scenario_id": str(workspace.confirmed_scenario_id) if workspace.confirmed_scenario_id else None,
        "selected_tax_scenario_id": str(workspace.selected_tax_scenario_id) if workspace.selected_tax_scenario_id else None,
    }


def _serialize_beneficiary_profile(profile):
    """Serialize a BeneficiaryProfile to JSON."""
    return {
        "id": str(profile.pk),
        "beneficiary_id": str(profile.beneficiary_id),
        "beneficiary_name": profile.beneficiary.full_name if hasattr(profile.beneficiary, 'full_name') else str(profile.beneficiary),
        "beneficiary_type": profile.beneficiary_type,
        "other_income": str(profile.other_income) if profile.other_income else None,
        "marginal_rate": str(profile.marginal_rate) if profile.marginal_rate else None,
        "bracket_remaining": str(profile.bracket_remaining) if profile.bracket_remaining else None,
        "franking_surplus": str(profile.franking_surplus) if profile.franking_surplus else None,
        "include_in_distribution": profile.include_in_distribution,
        "exclusion_reason": profile.exclusion_reason,
        "tax_residency": profile.tax_residency,
    }


def _serialize_scenario(scenario):
    """Serialize a DistributionScenario to JSON."""
    return {
        "id": str(scenario.pk),
        "name": scenario.name,
        "allocations": scenario.allocations,
        "total_tax": str(scenario.total_tax) if scenario.total_tax else None,
        "tax_saved_vs_equal": str(scenario.tax_saved_vs_equal) if scenario.tax_saved_vs_equal else None,
        "is_confirmed": scenario.is_confirmed,
        "created_at": scenario.created_at.isoformat(),
    }


def _serialize_100a(assessment):
    """Serialize a Section100AAssessment to JSON."""
    return {
        "id": str(assessment.pk),
        "beneficiary_id": str(assessment.beneficiary_id),
        "beneficiary_name": assessment.beneficiary.full_name if hasattr(assessment.beneficiary, 'full_name') else str(assessment.beneficiary),
        "q1": assessment.q1, "q2": assessment.q2, "q3": assessment.q3, "q4": assessment.q4,
        "q5": assessment.q5, "q6": assessment.q6, "q7": assessment.q7, "q8": assessment.q8,
        "risk_rating": assessment.risk_rating,
        "resolution_strategy": assessment.resolution_strategy,
        "reviewed_by": str(assessment.reviewed_by) if assessment.reviewed_by else None,
        "reviewed_at": assessment.reviewed_at.isoformat() if assessment.reviewed_at else None,
    }


def _serialize_election(election):
    """Serialize a TrustElectionRecord to JSON."""
    return {
        "id": str(election.pk),
        "election_type": election.election_type,
        "election_type_display": election.get_election_type_display(),
        "status": election.status,
        "status_display": election.get_status_display(),
        "effective_date": str(election.effective_date) if election.effective_date else None,
        "test_individual_id": str(election.test_individual_id) if election.test_individual_id else None,
        "test_individual_name": (
            election.test_individual.full_name
            if election.test_individual and hasattr(election.test_individual, 'full_name')
            else None
        ),
        "related_entity_id": str(election.related_entity_id) if election.related_entity_id else None,
        "confirmed_by": str(election.confirmed_by) if election.confirmed_by else None,
        "confirmed_at": election.confirmed_at.isoformat() if election.confirmed_at else None,
    }


# =============================================================================
# Stage 6 — Document Generation Views
# =============================================================================

def _get_confirmed_scenario_data(workspace):
    """
    Return a list of dicts describing each beneficiary's confirmed allocation.
    Each dict: {name, type, total, streams, percentage}
    Returns (rows, total_distributed, ndi)
    """
    confirmed = workspace.confirmed_scenario
    if not confirmed or not confirmed.allocations:
        return [], Decimal("0"), workspace.net_distributable_income or Decimal("0")

    profiles = {
        str(p.beneficiary_id): p
        for p in workspace.beneficiary_profiles.select_related("beneficiary").all()
    }

    rows = []
    total_distributed = Decimal("0")
    for ben_id, streams in confirmed.allocations.items():
        total_for_ben = sum(Decimal(str(v or 0)) for v in streams.values())
        if total_for_ben <= 0:
            continue
        profile = profiles.get(str(ben_id))
        name = profile.beneficiary.full_name if profile else f"Beneficiary {str(ben_id)[:8]}"
        ben_type = profile.get_beneficiary_type_display() if profile else ""
        rows.append({
            "name": name,
            "type": ben_type,
            "total": total_for_ben,
            "streams": {k: Decimal(str(v or 0)) for k, v in streams.items()},
            "percentage": Decimal("0"),
        })
        total_distributed += total_for_ben

    rows.sort(key=lambda r: r["name"])
    if total_distributed > 0:
        for r in rows:
            r["percentage"] = (r["total"] / total_distributed * 100).quantize(Decimal("0.01"))

    return rows, total_distributed, workspace.net_distributable_income or Decimal("0")


@login_required
def trust_generate_beneficiary_statements(request, pk):
    """
    GET /api/years/<pk>/trust-workspace/generate/beneficiary-statements/
    Generate a single DOCX containing all beneficiary statements from the
    confirmed distribution scenario.
    """
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.http import HttpResponse

    fy = get_object_or_404(FinancialYear, pk=pk)
    try:
        workspace = TrustWorkspace.objects.get(financial_year=fy)
    except TrustWorkspace.DoesNotExist:
        from django.http import JsonResponse as JR
        return JR({"error": "No trust workspace found."}, status=404)

    entity = fy.entity
    fy_year = "".join(c for c in fy.year_label if c.isdigit()) or str(fy.end_date.year)
    fy_end = f"30 June {fy_year}"

    # Build rows from selected TaxPlanningScenario
    scenario = workspace.selected_tax_scenario
    rows = []
    total_distributed = Decimal("0")
    if scenario and scenario.distributions:
        officer_map = {
            str(o.pk): o.full_name
            for o in EntityOfficer.objects.filter(entity=entity)
        }
        for entry in scenario.distributions:
            amount = Decimal(str(entry.get("proposed_distribution", 0)))
            if amount > 0:
                ben_id = str(entry.get("beneficiary_id", ""))
                rows.append({
                    "name": officer_map.get(ben_id, f"Beneficiary {ben_id[:8]}"),
                    "type": entry.get("beneficiary_type", "Individual").title(),
                    "total": amount,
                })
                total_distributed += amount
        rows.sort(key=lambda r: r["name"])
        if total_distributed > 0:
            for r in rows:
                r["percentage"] = (r["total"] / total_distributed * 100).quantize(Decimal("0.01"))

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for i, row in enumerate(rows):
        if i > 0:
            doc.add_page_break()

        h = doc.add_heading("Beneficiary Distribution Statement", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Trust: {entity.entity_name}")
        doc.add_paragraph(f"Financial Year: {fy.year_label}")
        doc.add_paragraph(f"Year Ended: {fy_end}")
        doc.add_paragraph(f"Beneficiary: {row['name']}")
        if row["type"]:
            doc.add_paragraph(f"Beneficiary Type: {row['type']}")
        doc.add_paragraph(f"Share of Distribution: {row.get('percentage', '')}%")
        doc.add_paragraph("")

        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Description"
        hdr[1].text = "Amount"
        for run in hdr[0].paragraphs[0].runs:
            run.bold = True
        for run in hdr[1].paragraphs[0].runs:
            run.bold = True

        total_row = table.add_row().cells
        total_row[0].text = "Total Distribution"
        total_row[1].text = f"${row['total']:,.2f}"
        for cell in total_row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        doc.add_paragraph("")
        doc.add_paragraph(
            "This statement is prepared for income tax purposes and shows your "
            "entitlement to the net income of the trust for the year ended "
            f"{fy_end}. Please retain this statement for your tax records."
        )

    if not rows:
        doc.add_paragraph(
            "No Tax Planning scenario selected. Please complete Stage 2 "
            "(Distribution Modelling) and select a scenario before generating statements."
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    entity_name = entity.entity_name.replace(" ", "_")
    filename = f"{entity_name}_Beneficiary_Statements_{fy.year_label}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@login_required
def trust_generate_distribution_summary(request, pk):
    """
    GET /api/years/<pk>/trust-workspace/generate/distribution-summary/
    Generate a DOCX distribution summary from the confirmed scenario.
    """
    import io
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.http import HttpResponse

    fy = get_object_or_404(FinancialYear, pk=pk)
    try:
        workspace = TrustWorkspace.objects.get(financial_year=fy)
    except TrustWorkspace.DoesNotExist:
        from django.http import JsonResponse as JR
        return JR({"error": "No trust workspace found."}, status=404)

    rows, total_distributed, ndi = _get_confirmed_scenario_data(workspace)
    entity = fy.entity
    fy_year = "".join(c for c in fy.year_label if c.isdigit()) or str(fy.end_date.year)
    fy_end = f"30 June {fy_year}"
    confirmed = workspace.confirmed_scenario

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("Trust Distribution Summary", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Trust: {entity.entity_name}")
    doc.add_paragraph(f"Financial Year: {fy.year_label}  |  Year Ended: {fy_end}")
    doc.add_paragraph(f"Scenario: {confirmed.name if confirmed else 'N/A'}")
    doc.add_paragraph(f"Net Distributable Income: ${ndi:,.2f}")
    doc.add_paragraph("")

    if rows:
        # Build column headers: Stream | Ben1 | Ben2 | ... | Total
        STREAM_LABELS = {
            "ordinary": "Ordinary",
            "cgt_discount": "CGT Discount",
            "cgt_non_discount": "CGT Non-Discount",
            "franked_dividends": "Franked Dividends",
            "franking_credits": "Franking Credits",
            "tax_free": "Tax-Free",
        }
        all_streams = list(STREAM_LABELS.keys())
        col_count = 1 + len(rows) + 1  # Stream + beneficiaries + Total

        table = doc.add_table(rows=1, cols=col_count)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Income Stream"
        for i, row in enumerate(rows):
            hdr[i + 1].text = row["name"]
        hdr[-1].text = "Total"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        stream_totals = {s: Decimal("0") for s in all_streams}
        for stream_key in all_streams:
            label = STREAM_LABELS.get(stream_key, stream_key)
            r = table.add_row().cells
            r[0].text = label
            row_total = Decimal("0")
            for i, ben_row in enumerate(rows):
                amt = ben_row["streams"].get(stream_key, Decimal("0"))
                r[i + 1].text = f"${amt:,.2f}" if amt else "-"
                row_total += amt
                stream_totals[stream_key] += amt
            r[-1].text = f"${row_total:,.2f}"

        # Total row
        total_row = table.add_row().cells
        total_row[0].text = "TOTAL"
        for i, ben_row in enumerate(rows):
            total_row[i + 1].text = f"${ben_row['total']:,.2f}"
        total_row[-1].text = f"${total_distributed:,.2f}"
        for cell in total_row:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        doc.add_paragraph("")

        # Percentage summary
        doc.add_heading("Distribution Percentages", level=2)
        pct_table = doc.add_table(rows=1, cols=3)
        pct_table.style = "Table Grid"
        ph = pct_table.rows[0].cells
        ph[0].text = "Beneficiary"
        ph[1].text = "Amount"
        ph[2].text = "Percentage"
        for cell in ph:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for ben_row in rows:
            pr = pct_table.add_row().cells
            pr[0].text = ben_row["name"]
            pr[1].text = f"${ben_row['total']:,.2f}"
            pr[2].text = f"{ben_row['percentage']}%"
    else:
        doc.add_paragraph(
            "No confirmed distribution scenario found. Please complete Stage 3 "
            "(Distribution Modelling) and confirm a scenario before generating this summary."
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    entity_name = entity.entity_name.replace(" ", "_")
    filename = f"{entity_name}_Distribution_Summary_{fy.year_label}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@login_required
def trust_generate_100a_summary(request, pk):
    """
    GET /api/years/<pk>/trust-workspace/generate/100a-summary/
    Generate a DOCX Section 100A risk assessment summary.
    """
    import io
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.http import HttpResponse

    fy = get_object_or_404(FinancialYear, pk=pk)
    try:
        workspace = TrustWorkspace.objects.get(financial_year=fy)
    except TrustWorkspace.DoesNotExist:
        from django.http import JsonResponse as JR
        return JR({"error": "No trust workspace found."}, status=404)

    entity = fy.entity
    fy_year = "".join(c for c in fy.year_label if c.isdigit()) or str(fy.end_date.year)
    fy_end = f"30 June {fy_year}"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading("Section 100A Risk Assessment Summary", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Trust: {entity.entity_name}")
    doc.add_paragraph(f"Financial Year: {fy.year_label}  |  Year Ended: {fy_end}")
    overall = workspace.section_100a_overall_risk or "Not assessed"
    doc.add_paragraph(f"Overall Risk Rating: {overall.upper()}")
    doc.add_paragraph("")

    QUESTIONS = [
        ("Q1", "Was the distribution made under a reimbursement agreement?"),
        ("Q2", "Did the beneficiary receive the economic benefit of the distribution?"),
        ("Q3", "Was the distribution part of a pre-arranged plan?"),
        ("Q4", "Were funds redirected to another party?"),
        ("Q5", "Is the beneficiary a related party of the trustee?"),
        ("Q6", "Was there a tax benefit from the arrangement?"),
        ("Q7", "Is the arrangement consistent with an ordinary family dealing? (protective)"),
        ("Q8", "Does the arrangement fall within a safe harbour? (protective)"),
    ]

    assessments = workspace.section_100a_assessments.select_related("beneficiary").all()
    if assessments:
        for assessment in assessments:
            doc.add_heading(assessment.beneficiary.full_name, level=2)
            doc.add_paragraph(f"Risk Rating: {assessment.risk_rating.upper() if assessment.risk_rating else 'Not assessed'}")

            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "#"
            hdr[1].text = "Question"
            hdr[2].text = "Answer"
            for cell in hdr:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            answers = [
                assessment.q1, assessment.q2, assessment.q3, assessment.q4,
                assessment.q5, assessment.q6, assessment.q7, assessment.q8,
            ]
            for (qnum, question), answer in zip(QUESTIONS, answers):
                r = table.add_row().cells
                r[0].text = qnum
                r[1].text = question
                r[2].text = (answer or "Not answered").title()

            if assessment.resolution_strategy:
                doc.add_paragraph(f"Resolution Strategy: {assessment.resolution_strategy}")
            doc.add_paragraph("")
    else:
        doc.add_paragraph(
            "No Section 100A assessments found. This stage may have been skipped "
            "(e.g. unit trust) or not yet completed."
        )

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    entity_name = entity.entity_name.replace(" ", "_")
    filename = f"{entity_name}_Section_100A_Summary_{fy.year_label}.docx"
    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# =============================================================================
# Post Distribution — Create Journal Entries
# =============================================================================

@login_required
@require_POST
def trust_post_distribution(request, pk):
    """
    POST /api/years/<pk>/trust-workspace/post-distribution/
    Creates and posts journal entries from the selected TaxPlanningScenario.

    DR  Profit Distribution — Appropriation (4199)
    CR  Beneficiary 4004.NN loan account per beneficiary (officer-linked
        EntityChartOfAccount, "Funds loaned to trust").

    Refuses to post if any beneficiary with a positive distribution has no
    4004.NN loan account — there is no silent fallback to 3100.

    Journal is created as POSTED and written to TB via _post_journal_to_tb.
    """
    from core.views import _post_journal_to_tb

    fy = get_object_or_404(FinancialYear, pk=pk)
    try:
        workspace = TrustWorkspace.objects.get(financial_year=fy)
    except TrustWorkspace.DoesNotExist:
        return JsonResponse({"error": "No trust workspace found."}, status=404)

    scenario = workspace.selected_tax_scenario
    if not scenario:
        return JsonResponse(
            {"error": "No Tax Planning scenario selected. Select a scenario in Stage 2 first."},
            status=400,
        )

    # Idempotency guard: never stack a second distribution. Blocks while a
    # live (posted, non-voided) distribution journal exists — un-post it first.
    existing = AdjustingJournal.live_trust_distribution(fy)
    if existing:
        return JsonResponse(
            {"error": "A distribution is already posted for this year — "
                      "un-post it first."},
            status=400,
        )

    # Build distribution rows from TaxPlanningScenario.distributions
    from core.models import EntityChartOfAccount

    officer_map = {
        str(o.pk): o.full_name
        for o in EntityOfficer.objects.filter(entity=fy.entity)
    }

    # Build a lookup: {officer_pk_str: (account_code, account_name)} for the
    # beneficiary's 4004.NN loan account ("Funds loaned to trust"), resolved
    # from the officer-linked EntityChartOfAccount materialised by the
    # 4000-family provisioning (core/beneficiary_account_service.py). The
    # distribution credits this account.
    #
    # There is NO 3100 / 9003 / lowest-code fallback: a missing 4004.NN
    # account is a hard error (the gate below refuses to post) rather than a
    # silent mis-post. The lowest 4004.NN per officer wins if duplicates exist.
    ben_loan_accounts = {}
    for eca in (
        EntityChartOfAccount.objects.filter(
            entity=fy.entity,
            beneficiary_officer__isnull=False,
            account_code__startswith="4004",
            is_active=True,
        )
        .select_related("beneficiary_officer")
        .order_by("account_code")
    ):
        officer_pk = str(eca.beneficiary_officer_id)
        ben_loan_accounts.setdefault(
            officer_pk, (eca.account_code, eca.account_name)
        )

    rows = []
    missing = []  # beneficiaries with a positive distribution but no 4004.NN
    total_distributed = Decimal("0")
    for entry in (scenario.distributions or []):
        amount = Decimal(str(entry.get("proposed_distribution", 0)))
        if amount <= 0:
            continue
        ben_id = str(entry.get("beneficiary_id", ""))
        ben_name = officer_map.get(ben_id, f"Beneficiary {ben_id[:8]}")
        # Resolve the beneficiary's 4004.NN loan account. No fallback — a
        # missing account is collected and reported as a hard error below.
        loan_acct = ben_loan_accounts.get(ben_id)
        if loan_acct is None:
            missing.append(ben_name)
            continue
        rows.append({
            "name": ben_name,
            "amount": amount,
            "cr_code": loan_acct[0],
            "cr_name": loan_acct[1],
        })
        total_distributed += amount

    if missing:
        names = ", ".join(missing)
        return JsonResponse(
            {"error": (
                f"Cannot post distribution: no 4004 (Funds loaned to trust) "
                f"loan account exists for: {names}. Create each beneficiary's "
                f"4004.NN account before posting."
            )},
            status=400,
        )

    if not rows:
        return JsonResponse(
            {"error": "Selected scenario has no allocations with positive amounts."},
            status=400,
        )

    fy_year = "".join(c for c in fy.year_label if c.isdigit()) or str(fy.end_date.year)

    try:
        with transaction.atomic():
            journal = AdjustingJournal.objects.create(
                financial_year=fy,
                journal_type=AdjustingJournal.JournalType.YEAR_END,
                status=AdjustingJournal.JournalStatus.POSTED,
                is_trust_distribution=True,
                journal_date=fy.end_date,
                description=f"Trust distribution — {scenario.scenario_name}",
                narration=(
                    f"Distribution from Tax Planning scenario '{scenario.scenario_name}'. "
                    f"Total: ${total_distributed:,.2f}."
                ),
                created_by=request.user,
                posted_by=request.user,
                posted_at=timezone.now(),
                total_debit=total_distributed,
                total_credit=total_distributed,
            )

            line_num = 1
            # DR side: Retained Profits / P&L appropriation
            JournalLine.objects.create(
                journal=journal,
                line_number=line_num,
                account_code="4199",
                account_name="Profit Distribution — Appropriation",
                description=f"Trust distribution for year ended 30 June {fy_year}",
                debit=total_distributed,
                credit=Decimal("0"),
            )
            line_num += 1

            # CR side: one line per beneficiary — posts to their 4004.NN
            # loan account (guaranteed present; the gate above refused to
            # post if any beneficiary lacked one).
            for row in rows:
                JournalLine.objects.create(
                    journal=journal,
                    line_number=line_num,
                    account_code=row["cr_code"],
                    account_name=row["cr_name"],
                    description=f"{row['name']}: ${row['amount']:,.2f}",
                    debit=Decimal("0"),
                    credit=row["amount"],
                )
                line_num += 1

            # Post to trial balance
            _post_journal_to_tb(journal, fy)

        return JsonResponse({
            "success": True,
            "journal_id": str(journal.pk),
            "journal_reference": journal.reference_number,
            "total_distributed": str(total_distributed),
            "beneficiary_count": len(rows),
            "message": (
                f"Distribution journal {journal.reference_number} posted to trial balance. "
                f"Total: ${total_distributed:,.2f} across {len(rows)} beneficiaries."
            ),
        })

    except Exception as e:
        logger.exception("trust_post_distribution failed: %s", e)
        return JsonResponse({"error": str(e)}, status=500)


@login_required
@require_POST
def trust_unpost_distribution(request, pk):
    """
    POST /api/years/<pk>/trust-workspace/unpost-distribution/

    Un-post the live trust distribution for the year. Branches on FY
    editability (mirrors the journal_delete teardown but soft-voids):

      * Editable year (status draft/reopened AND locked_at/finalised_at None)
        -> VOID: strip the journal's TB effect via the source_journal FK and
        mark status='voided', keeping the journal record. The post gate
        reopens (live distribution no longer exists).
      * Finalised / locked year -> REVERSE: post a system-generated reversing
        entry (debit/credit swapped) and leave both journals on the audit
        trail. The original stays posted, so the gate remains closed (you
        cannot re-post into a locked year).

    Idempotent: 400 if there is no live distribution to un-post.
    """
    from core.views import (
        _post_journal_to_tb, _reverse_journal_tb_lines,
        _delete_orphaned_tb_lines_for_journal, _verify_tb_balance,
        _log_action,
    )

    fy = get_object_or_404(FinancialYear, pk=pk)

    if not getattr(request.user, "can_do_accounting", False):
        return JsonResponse(
            {"error": "You do not have permission to un-post journals."},
            status=403,
        )

    journal = AdjustingJournal.live_trust_distribution(fy)
    if journal is None:
        return JsonResponse(
            {"error": "No live distribution to un-post."}, status=400
        )

    # Decision 2: void only while the year is editable; otherwise reverse.
    editable = (
        fy.status in (FinancialYear.Status.DRAFT, FinancialYear.Status.REOPENED)
        and fy.locked_at is None
        and fy.finalised_at is None
    )
    orig_ref = journal.reference_number

    try:
        with transaction.atomic():
            if editable:
                # VOID — remove the journal's TB lines, keep the record.
                _reverse_journal_tb_lines(journal)
                _delete_orphaned_tb_lines_for_journal(journal)
                journal.status = AdjustingJournal.JournalStatus.VOIDED
                journal.save(update_fields=["status"])
                _verify_tb_balance(fy)
                action, ref = "voided", orig_ref
                message = (
                    f"Distribution journal {orig_ref} voided. "
                    f"The post button is available again."
                )
            else:
                # REVERSE — system-generated reversing entry; both kept.
                reversal = AdjustingJournal.objects.create(
                    financial_year=fy,
                    journal_type=AdjustingJournal.JournalType.YEAR_END,
                    status=AdjustingJournal.JournalStatus.POSTED,
                    is_trust_distribution=False,  # the reversal is not the distribution
                    journal_date=fy.end_date,
                    description=f"Reversal of {orig_ref} — Trust distribution",
                    narration=(
                        f"System-generated reversal of distribution journal "
                        f"{orig_ref} in a finalised/locked year."
                    ),
                    created_by=request.user,
                    posted_by=request.user,
                    posted_at=timezone.now(),
                    total_debit=journal.total_credit,
                    total_credit=journal.total_debit,
                )
                line_num = 1
                for src in journal.lines.order_by("line_number", "id"):
                    JournalLine.objects.create(
                        journal=reversal,
                        line_number=line_num,
                        account_code=src.account_code,
                        account_name=src.account_name,
                        description=f"Reversal: {src.description}",
                        debit=src.credit,
                        credit=src.debit,
                    )
                    line_num += 1
                _post_journal_to_tb(reversal, fy)
                _verify_tb_balance(fy)
                action, ref = "reversed", reversal.reference_number
                message = (
                    f"Distribution journal {orig_ref} reversed by system entry "
                    f"{ref} (year is finalised/locked)."
                )

        _log_action(request, "adjustment", f"Un-posted ({action}) distribution {orig_ref}")
        try:
            from core.signals import trigger_risk_recalc
            trigger_risk_recalc(fy, "distribution_unposted")
        except Exception:
            logger.exception("risk recalc after un-post failed for FY %s", fy.pk)

        return JsonResponse({
            "success": True,
            "action": action,
            "reference": ref,
            "message": message,
        })

    except Exception as e:
        logger.exception("trust_unpost_distribution failed: %s", e)
        return JsonResponse({"error": str(e)}, status=500)
