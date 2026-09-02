"""A company's financial statements, asserted figure by figure.

The 2026-08-13 exploratory test generated statements for a company with no tax
posted, no prior year, no cost of sales and no trading stock -- four branches of
build_company_context that never ran. This module covers them.

Every expected figure here is hand-computed from the trial balance and written out
in the docstring of the test that asserts it. That is deliberate: the two defects
this area has produced (7e11395, ac69078) both kept the balance sheet balancing
while being wrong, so a golden baseline would have blessed them and an
invariant-only test would have passed them.

Amounts are debit-positive; credits are negative, matching what _get_tb_sections
reads. See docs/superpowers/specs/2026-08-13-company-fs-generation-test-design.md.
"""
from datetime import date
from decimal import Decimal

from django.test import TestCase, override_settings

from core.fs_template_service import build_company_context
from core.models import (
    AccountMapping,
    Client,
    Entity,
    EntityChartOfAccount,
    FinancialYear,
    TrialBalanceLine,
)

STORAGES_OVERRIDE = {
    "default": {"BACKEND": "django.core.files.storage.FileSystemStorage"},
    "staticfiles": {"BACKEND": "django.contrib.staticfiles.storage.StaticFilesStorage"},
}

# Standard line items the fixtures map to. Created per-test rather than relied on
# from a migration, following core/tests_docgen_section_classification.py.
MAPPINGS = [
    ("IS-REV-001", "Revenue", "income_statement", "Revenue"),
    ("IS-REV-002", "Other income", "income_statement", "Revenue"),
    ("IS-COS-002", "Opening stock", "income_statement", "Cost of Sales"),
    ("IS-COS-003", "Purchases", "income_statement", "Cost of Sales"),
    ("IS-COS-004", "Closing stock", "income_statement", "Cost of Sales"),
    ("IS-EXP-001", "Accounting and professional fees", "income_statement", "Expenses"),
    ("IS-EXP-007", "Depreciation and amortisation", "income_statement", "Expenses"),
    ("IS-EXP-008", "Employee benefits expense", "income_statement", "Expenses"),
    ("BS-CA-001", "Cash and cash equivalents", "balance_sheet", "Current Assets"),
    ("BS-CA-002", "Trade and other receivables", "balance_sheet", "Current Assets"),
    ("BS-CA-003", "Inventories", "balance_sheet", "Current Assets"),
    ("BS-NCA-001", "Property, plant and equipment", "balance_sheet", "Non-Current Assets"),
    ("BS-CL-001", "Trade and other payables", "balance_sheet", "Current Liabilities"),
    ("BS-EQ-001", "Issued capital", "balance_sheet", "Equity"),
    ("BS-EQ-002", "Retained earnings", "balance_sheet", "Equity"),
    ("BS-EQ-011", "Income tax provision", "balance_sheet", "Equity"),
]

# The maximal fixture: one company, two years, each internally consistent.
# (code, name, standard_code, cy_amount, py_amount) -- debit-positive.
#
# "Internally consistent" means each column balances on its own (assets = liabilities
# + equity, income - expenses = profit, within that column) -- NOT that the two columns
# chain across years. They deliberately do not: retained_profit_closing_py (340,000,
# see ComparativeColumnTests) is not equal to retained_profit_opening_cy (205,000,
# see IncomeTaxTests), and net assets move from 440,000 (PY) to 485,000 (CY) while
# after-tax profit is only 180,000 with no dividend account to absorb the rest. That
# gap is deliberate: it lets each column's arithmetic be verified independently of the
# other, without also having to model a roll-forward between years. Do not add a
# cross-year invariant such as `retained_profit_opening_cy == retained_profit_closing_py`
# against this fixture -- it would fail, correctly, because the fixture was never built
# to chain. Reworking the fixture to chain is a separate task.
#
# Every code was verified against the production copy. Codes 0570, 1100, 1115 and
# 1130 are entity accounts rather than company-template accounts: the company
# template's cost of sales block holds only 1000 and 1126, so real companies add
# these themselves, as Berwick Mechanical Services did with 1115 Purchases.
MAXIMAL_ROWS = [
    ("0510", "Sales", "IS-REV-001", Decimal("-900000"), Decimal("-800000")),
    ("0570", "Insurance recoveries", "IS-REV-002", Decimal("-20000"), Decimal("-10000")),
    ("1100", "Opening stock", "IS-COS-002", Decimal("60000"), Decimal("50000")),
    ("1115", "Purchases", "IS-COS-003", Decimal("400000"), Decimal("360000")),
    ("1130", "Closing stock", "IS-COS-004", Decimal("-70000"), Decimal("-60000")),
    ("1510", "Accountancy", "IS-EXP-001", Decimal("10000"), Decimal("9000")),
    ("1615", "Depreciation - Plant", "IS-EXP-007", Decimal("30000"), Decimal("28000")),
    ("1965", "Wages", "IS-EXP-008", Decimal("250000"), Decimal("230000")),
    ("4110", "Income tax expense/income", "BS-EQ-011", Decimal("60000"), Decimal("50000")),
    ("2000", "Cash at bank", "BS-CA-001", Decimal("210000"), Decimal("150000")),
    ("2101", "Trade debtors", "BS-CA-002", Decimal("120000"), Decimal("100000")),
    ("2363", "Finished goods - At real value", "BS-CA-003", Decimal("70000"), Decimal("60000")),
    ("2860", "Plant & equipment (cost)", "BS-NCA-001", Decimal("300000"), Decimal("300000")),
    ("2869", "Less: Accumulated depreciation", "BS-NCA-001", Decimal("-120000"), Decimal("-90000")),
    ("3048", "Trade creditors", "BS-CL-001", Decimal("-95000"), Decimal("-80000")),
    ("4200", "Issued & paid up capital", "BS-EQ-001", Decimal("-100000"), Decimal("-100000")),
    ("4199", "Retained profits", "BS-EQ-002", Decimal("-205000"), Decimal("-197000")),
]


def mapping_for(standard_code):
    return AccountMapping.objects.get(standard_code=standard_code)


def docx_text(buffer):
    """All text in a rendered DOCX, paragraphs then table cells.

    A table row's cells are joined by newlines, so a labelled figure appears as
    "Trade debtors\\n120,000" -- which is what lets a test assert the label and
    its amount are adjacent rather than merely both present somewhere.
    """
    import docx

    buffer.seek(0)
    document = docx.Document(buffer)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def build_company_fy(rows, *, with_prior=True, entity_name="FS Gen Test Co Pty Ltd"):
    """Create a company, its two financial years, and Model A trial balance rows.

    Model A row shapes (commit cb00bf1) are what make comparatives possible without
    performing a roll-forward:

      rollover row   closing_balance = the opening balance (prior-year closing for a
                     balance-sheet account, 0 for P&L), and prior_debit/prior_credit
                     carry the prior-year figures that become the comparative column
      tb_import row  closing_balance = the year's movement

    _get_tb_sections sums closing_balance across both, so CY = opening + movement,
    and reads PY from prior_debit - prior_credit on the rollover row.

    with_prior=False omits the rollover row entirely, modelling a first-year entity
    with no comparatives.
    """
    for code, label, statement, section in MAPPINGS:
        AccountMapping.objects.get_or_create(
            standard_code=code,
            defaults={
                "line_item_label": label,
                "financial_statement": statement,
                "statement_section": section,
            },
        )

    client = Client.objects.create(name=f"{entity_name} Client")
    entity = Entity.objects.create(
        entity_name=entity_name,
        entity_type="company",
        client=client,
        abn="11000000560",
    )
    prior_fy = FinancialYear.objects.create(
        entity=entity,
        year_label="2025",
        start_date=date(2024, 7, 1),
        end_date=date(2025, 6, 30),
        status=FinancialYear.Status.FINALISED,
    )
    fy = FinancialYear.objects.create(
        entity=entity,
        year_label="2026",
        start_date=date(2025, 7, 1),
        end_date=date(2026, 6, 30),
        status=FinancialYear.Status.DRAFT,
        prior_year=prior_fy,
    )

    for code, name, standard_code, cy, py in rows:
        mapping = mapping_for(standard_code)
        EntityChartOfAccount.objects.get_or_create(
            entity=entity,
            account_code=code,
            defaults={"account_name": name, "is_active": True},
        )
        is_balance_sheet = standard_code.startswith("BS-")
        opening = py if is_balance_sheet else Decimal("0")

        if with_prior:
            # The prior financial year needs its own trial balance too:
            # _has_prior_year(fy) checks fy.prior_year.trial_balance_lines.exists(),
            # and _post_process_fs_doc strips the comparative column entirely when
            # that is False. Without this, no prior-year figure reaches any
            # rendered document no matter what prior_debit/prior_credit carry on
            # the current year's rows.
            TrialBalanceLine.objects.create(
                financial_year=prior_fy,
                account_code=code,
                account_name=name,
                mapped_line_item=mapping,
                closing_balance=py,
                debit=py if py > 0 else Decimal("0"),
                credit=-py if py < 0 else Decimal("0"),
                source="tb_import",
                is_adjustment=False,
            )
            TrialBalanceLine.objects.create(
                financial_year=fy,
                account_code=code,
                account_name=name,
                mapped_line_item=mapping,
                closing_balance=opening,
                debit=opening if opening > 0 else Decimal("0"),
                credit=-opening if opening < 0 else Decimal("0"),
                prior_debit=py if py > 0 else Decimal("0"),
                prior_credit=-py if py < 0 else Decimal("0"),
                source="rollover",
                is_adjustment=False,
            )
            movement = cy - opening
        else:
            movement = cy

        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code=code,
            account_name=name,
            mapped_line_item=mapping,
            closing_balance=movement,
            debit=movement if movement > 0 else Decimal("0"),
            credit=-movement if movement < 0 else Decimal("0"),
            prior_debit=Decimal("0"),
            prior_credit=Decimal("0"),
            source="tb_import",
            is_adjustment=False,
        )

    return fy


@override_settings(STORAGES=STORAGES_OVERRIDE)
class BuilderTests(TestCase):
    """The builder must reproduce the input figures exactly, or every scenario
    below is testing the wrong numbers."""

    def test_the_maximal_trial_balance_sums_to_zero_in_both_years(self):
        cy = sum(r[3] for r in MAXIMAL_ROWS)
        py = sum(r[4] for r in MAXIMAL_ROWS)
        self.assertEqual(cy, Decimal("0"), "current-year trial balance does not balance")
        self.assertEqual(py, Decimal("0"), "prior-year trial balance does not balance")

    def test_the_builder_reproduces_each_accounts_cy_and_py(self):
        fy = build_company_fy(MAXIMAL_ROWS)
        context = build_company_context(fy)

        by_code = {}
        for items in context["_sections"].values():
            for item in items:
                by_code[item["account_code"]] = item

        for code, _name, _std, cy, py in MAXIMAL_ROWS:
            # 4110 is deliberately absent from _sections. build_company_context
            # extracts the income tax appropriation OUT of the equity section
            # before returning, so it reaches no section at all; its figures are
            # asserted through _income_tax_cy just below.
            if code == "4110":
                continue
            with self.subTest(code=code):
                self.assertIn(code, by_code, f"{code} did not reach any section")
                self.assertEqual(by_code[code]["cy_amount"], cy)
                self.assertEqual(by_code[code]["py_amount"], py)

        self.assertEqual(context["_income_tax_cy"], Decimal("60000"))
        self.assertEqual(context["_income_tax_py"], Decimal("50000"))

    def test_the_equity_section_carries_an_injected_current_year_profit_row(self):
        """Under the unclosed-TB convention the equity accounts hold only opening
        balances, so build_company_context injects a synthetic
        'Current year profit / (loss)' row (account_code 'NET_PROFIT') to make
        equity reconcile to net assets. Anything summing the equity section must
        not then add the profit a second time."""
        context = build_company_context(build_company_fy(MAXIMAL_ROWS))
        injected = [i for i in context["_sections"]["equity"]
                    if i["account_code"] == "NET_PROFIT"]
        self.assertEqual(len(injected), 1)
        self.assertEqual(injected[0]["cy_amount"], Decimal("-180000"))

    def test_without_prior_the_comparative_column_is_zero(self):
        fy = build_company_fy(MAXIMAL_ROWS, with_prior=False)
        context = build_company_context(fy)

        for items in context["_sections"].values():
            for item in items:
                with self.subTest(code=item["account_code"]):
                    self.assertEqual(item["py_amount"], Decimal("0"))


@override_settings(STORAGES=STORAGES_OVERRIDE)
class IncomeTaxTests(TestCase):
    """4110 mapped to BS-EQ-011 is the tax journal.

    Hand-computed, current year:
        income          900,000 + 20,000            = 920,000
        cost of sales    60,000 + 400,000 - 70,000  = 390,000
        other expenses   10,000 + 30,000 + 250,000  = 290,000
        profit pretax   920,000 - 390,000 - 290,000 = 240,000
        income tax                                     60,000
        profit after tax        240,000 - 60,000    = 180,000

    The ac69078 defect made profit after tax EXCEED profit before tax, because a
    non-tax equity account was swept into the tax figure with a credit balance.
    """

    def setUp(self):
        self.context = build_company_context(build_company_fy(MAXIMAL_ROWS))

    def test_the_tax_journal_is_recognised(self):
        self.assertTrue(self.context["has_income_tax"])
        self.assertEqual(self.context["_income_tax_cy"], Decimal("60000"))
        self.assertEqual(self.context["_income_tax_py"], Decimal("50000"))
        assert_statements_are_internally_consistent(self, self.context)

    def test_tax_is_deducted_from_profit_not_added(self):
        self.assertEqual(self.context["net_profit_pretax_cy"], "240,000")
        self.assertEqual(self.context["net_profit_cy"], "180,000")

    def test_the_tax_line_displays_as_a_deduction(self):
        """format_amount(-60000) -- brackets mean deducted. A positive figure here
        is the ac69078 symptom."""
        self.assertEqual(self.context["income_tax_cy"], "(60,000)")
        self.assertEqual(self.context["income_tax_py"], "(50,000)")

    def test_the_tax_account_is_removed_from_equity(self):
        equity_codes = {i["account_code"] for i in self.context["_sections"]["equity"]}
        self.assertNotIn("4110", equity_codes)
        self.assertIn("4200", equity_codes)
        self.assertIn("4199", equity_codes)

    def test_the_retained_profit_appropriation_reconciles(self):
        """opening 205,000 + after-tax profit 180,000 = closing 385,000"""
        self.assertEqual(self.context["retained_profit_opening_cy"], "205,000")
        self.assertEqual(self.context["retained_profit_closing_cy"], "385,000")


@override_settings(STORAGES=STORAGES_OVERRIDE)
class ComparativeColumnTests(TestCase):
    """The prior column, which a first-year entity never reaches.

    Hand-computed, prior year:
        income          800,000 + 10,000            = 810,000
        cost of sales    50,000 + 360,000 - 60,000  = 350,000
        other expenses    9,000 + 28,000 + 230,000  = 267,000
        profit pretax   810,000 - 350,000 - 267,000 = 193,000
        income tax                                     50,000
        profit after tax        193,000 - 50,000    = 143,000
        net assets  150,000+100,000+60,000+300,000-90,000-80,000 = 440,000
        equity      100,000 + 197,000 + 143,000                  = 440,000
    """

    def setUp(self):
        self.context = build_company_context(build_company_fy(MAXIMAL_ROWS))

    def test_prior_year_profit_and_loss(self):
        self.assertEqual(self.context["total_income_py"], "810,000")
        self.assertEqual(self.context["total_cogs_py"], "350,000")
        self.assertEqual(self.context["total_expenses_py"], "617,000")
        self.assertEqual(self.context["net_profit_pretax_py"], "193,000")
        self.assertEqual(self.context["net_profit_py"], "143,000")

    def test_prior_year_balance_sheet(self):
        self.assertEqual(self.context["total_current_assets_py"], "310,000")
        self.assertEqual(self.context["total_noncurrent_assets_py"], "210,000")
        self.assertEqual(self.context["total_assets_py"], "520,000")
        self.assertEqual(self.context["total_liabilities_py"], "80,000")
        self.assertEqual(self.context["net_assets_py"], "440,000")
        self.assertEqual(self.context["total_equity_py"], "440,000")

    def test_prior_retained_profit_reconciles(self):
        self.assertEqual(self.context["retained_profit_opening_py"], "197,000")
        self.assertEqual(self.context["retained_profit_closing_py"], "340,000")

    def test_a_first_year_entity_reports_no_comparatives(self):
        context = build_company_context(build_company_fy(MAXIMAL_ROWS, with_prior=False))
        self.assertEqual(context["total_income_py"], "—")
        self.assertEqual(context["net_assets_py"], "—")


@override_settings(STORAGES=STORAGES_OVERRIDE)
class TradingAndCostOfSalesTests(TestCase):
    """0510 maps to IS-REV-001 (trading income); the stock and purchase lines map
    to IS-COS-*, which sets has_trading and exercises the merge branch.

    The merge is deliberate and is pinned here: a trading company's detailed P&L
    shows a single Income and a single Expenses section, with cost of sales folded
    into Expenses and NO gross profit line. Confirmed on Berwick Mechanical Services
    FY2017, whose Purchases 210,029 renders under Expenses. Decision of 2026-08-13.
    If the presentation is ever meant to change, this test changes with it -- it
    records current intent, not an accounting standard.

        cost of sales   60,000 + 400,000 - 70,000 = 390,000
        other expenses  10,000 + 30,000 + 250,000 = 290,000
        rendered total expenses                    = 680,000
    """

    def setUp(self):
        self.context = build_company_context(build_company_fy(MAXIMAL_ROWS))

    def test_trading_is_detected(self):
        self.assertTrue(self.context["has_trading"])
        assert_statements_are_internally_consistent(self, self.context)

    def test_trading_income_and_other_income_are_separated_in_the_sections(self):
        trading = {i["account_code"] for i in self.context["_sections"]["trading_income"]}
        other = {i["account_code"] for i in self.context["_sections"]["income"]}
        self.assertEqual(trading, {"0510"})
        self.assertEqual(other, {"0570"})

    def test_cost_of_sales_totals_correctly(self):
        self.assertEqual(self.context["total_cogs_cy"], "390,000")

    def test_cost_of_sales_is_merged_into_rendered_expenses(self):
        """680,000 = cost of sales 390,000 + other expenses 290,000."""
        self.assertEqual(self.context["total_expenses_cy"], "680,000")
        rendered_codes = {row["account_code"] for row in self.context["expenses"]}
        for code in ("1100", "1115", "1130"):
            self.assertIn(code, rendered_codes, "cost of sales did not merge into Expenses")

    def test_total_income_merges_trading_and_other_income(self):
        self.assertEqual(self.context["total_income_cy"], "920,000")
        rendered_codes = {row["account_code"] for row in self.context["income"]}
        self.assertEqual(rendered_codes, {"0510", "0570"})

    def test_gross_profit_is_computed_but_not_presented(self):
        """900,000 trading income - 390,000 cost of sales = 510,000. The value is
        exposed to templates; no .docx renders it, and none should."""
        self.assertEqual(self.context["gross_profit_cy"], "510,000")


@override_settings(STORAGES=STORAGES_OVERRIDE)
class TradingStockTests(TestCase):
    """The stock cycle, and its tie to the balance sheet.

        opening stock              60,000   (prior year's closing stock)
        purchases                 400,000
        less closing stock        (70,000)
        cost of sales             390,000

    Closing stock 70,000 equals finished goods carried on the balance sheet, and
    prior closing stock 60,000 equals both the prior carrying amount and the
    current year's opening stock. Those ties make the fixture self-checking.
    """

    def setUp(self):
        self.context = build_company_context(build_company_fy(MAXIMAL_ROWS))
        self.by_code = {}
        for items in self.context["_sections"].values():
            for item in items:
                self.by_code[item["account_code"]] = item

    def test_the_stock_cycle_produces_cost_of_sales(self):
        opening = self.by_code["1100"]["cy_amount"]
        purchases = self.by_code["1115"]["cy_amount"]
        closing = self.by_code["1130"]["cy_amount"]
        self.assertEqual(opening + purchases + closing, Decimal("390000"))
        self.assertEqual(self.context["total_cogs_cy"], "390,000")
        assert_statements_are_internally_consistent(self, self.context)

    def test_finished_goods_is_a_current_asset(self):
        codes = {i["account_code"] for i in self.context["_sections"]["current_assets"]}
        self.assertIn("2363", codes)
        self.assertEqual(self.context["total_current_assets_cy"], "400,000")

    def test_closing_stock_ties_to_the_balance_sheet_carrying_amount(self):
        closing_stock = -self.by_code["1130"]["cy_amount"]
        on_hand = self.by_code["2363"]["cy_amount"]
        self.assertEqual(closing_stock, on_hand)

    def test_prior_closing_stock_ties_to_current_opening_stock(self):
        prior_closing = -self.by_code["1130"]["py_amount"]
        current_opening = self.by_code["1100"]["cy_amount"]
        self.assertEqual(prior_closing, current_opening)
        self.assertEqual(prior_closing, self.by_code["2363"]["py_amount"])


def assert_statements_are_internally_consistent(t, context):
    """Three relationships that must hold for any company, whatever the figures.

    A fourth -- each note ties to its face figure -- is asserted separately in
    DocumentRenderTests.test_the_notes_tie_to_the_face_of_the_statements, since it
    needs a rendered document rather than build_company_context's raw figures.

    The third is the one the ac69078 defect broke while the balance sheet still
    balanced, which is why an invariant-only suite would have passed it and why
    these run alongside hand-computed figures rather than instead of them.
    """
    sections = context["_sections"]

    def total(key):
        return sum((i["cy_amount"] or Decimal("0")) for i in sections[key])

    assets = total("current_assets") + total("noncurrent_assets")
    liabilities = -(total("current_liabilities") + total("noncurrent_liabilities"))
    # The equity section already contains the injected 'Current year profit /
    # (loss)' row (account_code 'NET_PROFIT'), which build_company_context adds
    # because the trial balance is unclosed. Adding the profit again here would
    # double-count it.
    equity = -total("equity")

    t.assertEqual(assets - liabilities, equity,
                  "net assets do not equal total equity")
    t.assertEqual(context["net_assets_cy"], context["total_equity_cy"],
                  "the rendered net assets and total equity disagree")

    pretax = _money(context["net_profit_pretax_cy"])
    after = _money(context["net_profit_cy"])
    t.assertEqual(after, pretax - context["_income_tax_cy"],
                  "profit after tax is not profit before tax less income tax")


def _money(formatted):
    """Turn a format_amount string back into a Decimal."""
    if formatted in ("—", "-"):
        return Decimal("0")
    negative = formatted.startswith("(")
    digits = formatted.strip("()").replace(",", "")
    return Decimal(digits) * (-1 if negative else 1)


@override_settings(STORAGES=STORAGES_OVERRIDE)
class FullSetTests(TestCase):
    """Everything at once. The interactions are where both prior defects lived.

    Current year, hand-computed:
        income      900,000 + 20,000                        =   920,000
        cost of sales  60,000 + 400,000 - 70,000            =   390,000
        expenses    10,000 + 30,000 + 250,000               =   290,000
        pretax      920,000 - 680,000                       =   240,000
        tax                                                 =    60,000
        after tax   240,000 - 60,000                        =   180,000
        current assets  210,000 + 120,000 + 70,000          =   400,000
        non-current     300,000 - 120,000                   =   180,000
        total assets                                        =   580,000
        liabilities                                         =    95,000
        net assets  580,000 - 95,000                        =   485,000
        equity      100,000 + 205,000 + 180,000             =   485,000
    """

    def setUp(self):
        self.context = build_company_context(build_company_fy(MAXIMAL_ROWS))

    def test_the_profit_and_loss(self):
        self.assertEqual(self.context["total_income_cy"], "920,000")
        self.assertEqual(self.context["total_cogs_cy"], "390,000")
        self.assertEqual(self.context["total_expenses_cy"], "680,000")
        self.assertEqual(self.context["net_profit_pretax_cy"], "240,000")
        self.assertEqual(self.context["income_tax_cy"], "(60,000)")
        self.assertEqual(self.context["net_profit_cy"], "180,000")

    def test_the_balance_sheet(self):
        self.assertEqual(self.context["total_current_assets_cy"], "400,000")
        self.assertEqual(self.context["total_noncurrent_assets_cy"], "180,000")
        self.assertEqual(self.context["total_assets_cy"], "580,000")
        self.assertEqual(self.context["total_liabilities_cy"], "95,000")
        self.assertEqual(self.context["net_assets_cy"], "485,000")
        self.assertEqual(self.context["total_equity_cy"], "485,000")

    def test_depreciation_expense_matches_the_accumulated_movement(self):
        """30,000 charged for the year; accumulated moves 90,000 -> 120,000."""
        by_code = {}
        for items in self.context["_sections"].values():
            for item in items:
                by_code[item["account_code"]] = item
        charge = by_code["1615"]["cy_amount"]
        movement = -(by_code["2869"]["cy_amount"] - by_code["2869"]["py_amount"])
        self.assertEqual(charge, movement)

    def test_the_invariants_hold(self):
        assert_statements_are_internally_consistent(self, self.context)


@override_settings(STORAGES=STORAGES_OVERRIDE)
class DocumentRenderTests(TestCase):
    """A correct figure that never reaches the page is still a broken statement.

    generate_financial_statements returns rendered DOCX buffers keyed by document
    type; PDF conversion is a separate LibreOffice step downstream and is not
    covered here -- it needs an external binary and adds no figure coverage.
    """

    _text_of = staticmethod(docx_text)

    def setUp(self):
        from core.fs_template_service import generate_financial_statements

        fy = build_company_fy(MAXIMAL_ROWS)
        self.documents = generate_financial_statements(fy.pk)
        self.assertTrue(self.documents, "no documents were rendered at all")

    def test_the_profit_and_loss_totals_reach_the_document(self):
        """Document types are BALANCE_SHEET, COMPILATION, COVER, DECLARATION,
        DETAILED_PL, NOTES, SUMMARY_PL."""
        rendered = self._text_of(self.documents["DETAILED_PL"])
        for figure in ("920,000", "680,000", "240,000", "(60,000)", "180,000",
                       "810,000", "143,000"):
            with self.subTest(figure=figure):
                self.assertIn(figure, rendered,
                              f"{figure} was computed but never reached the P&L")

    def test_the_balance_sheet_totals_reach_the_document(self):
        rendered = self._text_of(self.documents["BALANCE_SHEET"])
        for figure in ("400,000", "580,000", "95,000", "485,000", "440,000"):
            with self.subTest(figure=figure):
                self.assertIn(figure, rendered,
                              f"{figure} was computed but never reached the balance sheet")

    def test_the_notes_tie_to_the_face_of_the_statements(self):
        """The spec's fourth invariant, asserted where the notes actually exist.

        Receivables 120,000 and plant 300,000 each appear in the notes AND on the
        balance sheet -- which is what "ties" means here.

        Asserted against the labelled line, not the bare figure: the bare string
        "120,000" is also a substring of the PPE note's "Less: Accumulated
        depreciation (120,000)", so a bare-figure assertion stays green even if the
        entire receivables note is deleted from _compute_note_map. The docx text
        extractor joins a table row's cells with newlines, so the label and its
        figure appear as "Trade debtors\\n120,000" and
        "Plant & equipment (cost)\\n300,000" verbatim -- asserting that adjacency
        means the test fails if either note (or its figure) goes missing.

        Inventories now ties here too. It did not when this module was written:
        _compute_note_map emitted six note types and none was inventories, so a
        company carrying trading stock got a balance sheet line with no note
        behind it. Added 2026-08-14 -- see InventoriesNoteTests for the note's
        own coverage, which is where its numbering and policy text are asserted.

        The show_note_inventory flag at document_context_builder.py:1355 is still
        consumed by nothing. It is part of a parallel set of show_note_* flags
        that no rendered document reads -- the notes are built programmatically
        by _generate_notes_document, not from a .docx template -- and cleaning
        that up is a separate task.
        """
        notes = self._text_of(self.documents["NOTES"])
        balance_sheet = self._text_of(self.documents["BALANCE_SHEET"])
        self.assertIn("Trade debtors\n120,000", notes,
                      "the receivables note's labelled line is missing")
        self.assertIn("Plant & equipment (cost)\n300,000", notes,
                      "the PPE note's labelled line is missing")
        self.assertIn("Finished goods - At real value\n70,000", notes,
                      "the inventories note's labelled line is missing")
        for figure in ("120,000", "300,000", "70,000"):
            with self.subTest(figure=figure):
                self.assertIn(figure, balance_sheet,
                              f"{figure} is in the notes but not on the balance sheet")


class CurrentAssetInventoryClassificationTests(TestCase):
    """Inventory must not be presented as cash.

    _classify_current_asset tests the cash name keywords -- which include
    "on hand" -- before anything inventory-related, so "Stock on Hand", a
    default Xero account name, groups under Cash and Cash Equivalents. The
    structured standard_code branch does not save it: that branch only
    recognises cash codes, so an account correctly mapped to BS-CA-003
    Inventories still falls through to the name keywords.
    """

    @staticmethod
    def _classify(name, standard_code=None):
        from core.fs_template_service import _classify_current_asset

        item = {"account_name": name}
        if standard_code:
            item["standard_code"] = standard_code
        return _classify_current_asset(item)

    def test_stock_on_hand_is_inventory_not_cash(self):
        self.assertEqual(self._classify("Stock on Hand"), "Inventories")

    def test_the_mapping_classifies_a_name_no_keyword_would_catch(self):
        """BS-CA-003 is the structured Inventories code, already used by
        access_ledger_import.py and integrations/xero_gl_summary.py. Once an
        accountant has mapped an account, the mapping is the source of truth and
        the name should not need to look like stock at all."""
        self.assertEqual(
            self._classify("Goods for resale", standard_code="BS-CA-003"),
            "Inventories",
        )

    def test_cash_on_hand_is_still_cash(self):
        """A guard, not a red test -- it passed before the keyword reorder and
        must keep passing after it. "Cash on hand" is the only "on hand" account
        name in production (10 trial balance lines, checked 2026-08-14), so this
        is the case the reorder could realistically have broken."""
        self.assertEqual(self._classify("Cash on hand"), "Cash Assets")
        self.assertEqual(
            self._classify("Cash on hand", standard_code="BS-CA-001"),
            "Cash Assets",
        )

    def test_the_keywords_do_not_reach_beyond_stock(self):
        """Also a guard. Work in progress on a services chart is unbilled labour
        and belongs in receivables-or-other, not inventories; "Stock options" is
        an equity instrument. Both would be caught by a looser keyword list."""
        self.assertEqual(self._classify("Work in Progress"), "Other")
        self.assertEqual(self._classify("Stock options reserve"), "Other")


@override_settings(STORAGES=STORAGES_OVERRIDE)
class InventoriesNoteTests(TestCase):
    """A company carrying trading stock gets a supporting note for it.

    The maximal fixture holds 2363 Finished goods - At real value, mapped to
    BS-CA-003, at 70,000 (CY) and 60,000 (PY). Before this note existed the
    balance sheet showed that carrying amount with nothing behind it, while
    receivables and PPE each had a note.
    """

    def setUp(self):
        self.fy = build_company_fy(MAXIMAL_ROWS)
        self.context = build_company_context(self.fy)

    def test_the_note_map_places_inventories_after_receivables(self):
        """Balance-sheet order: receivables, then inventories, then PPE. The map
        is sequential, so inserting at position 3 renumbers everything below it.

        Expected for this fixture: policies 1, receivables 2 (trade debtors
        120,000), inventories 3, ppe 4 (plant at cost), income tax 5 (4110
        posted), events 6 (company). No related party note -- the fixture has no
        management fees and no director loans.
        """
        self.assertEqual(
            self.context["_note_map"],
            [(1, "policies"), (2, "receivables"), (3, "inventories"),
             (4, "ppe"), (5, "income_tax"), (6, "events")],
        )

    def test_the_note_renders_with_its_figures_and_policy(self):
        """Asserted against the labelled line, for the reason recorded in
        DocumentRenderTests.test_the_notes_tie_to_the_face_of_the_statements: a
        bare "70,000" is a substring of nothing else here today, but the PPE
        note's accumulated depreciation once made exactly that assertion
        worthless in the receivables note.
        """
        from core.fs_template_service import generate_financial_statements

        notes = docx_text(generate_financial_statements(self.fy.pk)["NOTES"])
        self.assertIn("Note 3: Inventories", notes)
        self.assertIn("Finished goods - At real value\n70,000", notes)
        self.assertIn("60,000", notes)
        self.assertIn("lower of cost and net realisable value", notes)

    def test_the_balance_sheet_line_points_at_the_note(self):
        """The trap the related-party comment records is a "(Note N)" reference
        with no matching line in the note body. This is the mirror of it: a note
        that exists while the balance sheet line it explains never refers to it.
        """
        inventory_lines = [
            i for i in self.context["current_assets"]
            if i.get("account_name") == "Finished goods - At real value"
        ]
        self.assertEqual(len(inventory_lines), 1,
                         "the inventory line is missing from current assets")
        self.assertEqual(inventory_lines[0]["note_ref"], "3")

    def test_the_note_total_ties_to_the_balance_sheet(self):
        """70,000 CY and 60,000 PY appear in the note AND as the carrying amount
        on the face of the balance sheet -- the spec's fourth invariant."""
        from core.fs_template_service import generate_financial_statements

        documents = generate_financial_statements(self.fy.pk)
        notes = docx_text(documents["NOTES"])
        balance_sheet = docx_text(documents["BALANCE_SHEET"])
        self.assertIn("Finished goods - At real value\n70,000", notes)
        # The balance sheet carries a Note column between label and figure, so the
        # row reads label, note ref, CY -- asserting all three at once.
        self.assertIn("Finished goods - At real value\n3\n70,000", balance_sheet)
        for figure in ("70,000", "60,000"):
            with self.subTest(figure=figure):
                self.assertIn(figure, notes)
                self.assertIn(figure, balance_sheet)

    def test_a_company_without_stock_gets_no_note_and_the_old_numbering(self):
        """A guard against an over-broad trigger, not a red test.

        The inventory row is swapped for a prepayment of the same amount, so the
        balance sheet still balances and the only thing that changes is that no
        account is stock. Every note below inventories must renumber back up.
        """
        from core.fs_template_service import generate_financial_statements

        rows = [r for r in MAXIMAL_ROWS if r[0] != "2363"]
        rows.append(("2363", "Prepayments", "BS-CA-002",
                     Decimal("70000"), Decimal("60000")))
        fy = build_company_fy(rows, entity_name="No Stock Pty Ltd")

        context = build_company_context(fy)
        self.assertEqual(
            context["_note_map"],
            [(1, "policies"), (2, "receivables"), (3, "ppe"),
             (4, "income_tax"), (5, "events")],
        )
        notes = docx_text(generate_financial_statements(fy.pk)["NOTES"])
        self.assertNotIn("Inventories", notes)
        self.assertIn("Note 3: Property, Plant and Equipment", notes)
