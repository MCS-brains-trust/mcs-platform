"""
StatementHub — JSON-Driven Document Template Renderer

Reads a DocumentTemplate.structure (JSON) from PostgreSQL and produces a
Word document (.docx) via python-docx.

Supported section types:
  heading           — text + level (1-3) + alignment
  paragraph         — text with {{merge_fields}}, bold/italic/alignment
  paragraph_list    — array of paragraph items (for multi-paragraph blocks)
  table             — column defs + header row + data_source merge field
  spacer            — vertical whitespace (configurable lines)
  page_break        — insert page break
  horizontal_rule   — thin horizontal line
  conditional       — if merge_field is truthy, render children[]
  signature_block   — name_field + title_field + date_field
  disclaimer        — italic text block
  numbered_list     — ordered list items with {{merge_fields}}
  bullet_list       — unordered list items with {{merge_fields}}

Merge field resolution:
  Each document category has a dedicated resolver function that queries the
  database and returns a flat dict of {field_name: value}.  Values can be:
    - str/int/Decimal  → inserted as text
    - list[dict]       → used as table data_source rows
    - bool             → used for conditional blocks
"""
import io
import re
import logging
from decimal import Decimal

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from core.template_resolvers import _fmt_money

logger = logging.getLogger(__name__)

# =============================================================================
# Default Styles
# =============================================================================
DEFAULT_STYLES = {
    "font_name": "Arial",
    "font_size_body": 11,
    "font_size_heading": 14,
    "font_size_subheading": 12,
    "font_size_small": 9,
    "font_size_footer": 8,
    "line_spacing": 1.15,
    "heading_color": "000000",
    "body_color": "000000",
    "table_header_bg": "333333",
    "table_header_fg": "FFFFFF",
    "table_alt_row_bg": "F5F5F5",
}

# =============================================================================
# Alignment mapping
# =============================================================================
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "centre": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


# =============================================================================
# Core Renderer
# =============================================================================
class TemplateRenderer:
    """
    Renders a DocumentTemplate.structure JSON into a python-docx Document.

    Usage:
        from core.models import DocumentTemplate
        from core.template_renderer import TemplateRenderer
        from core.template_resolvers import resolve_distribution_minutes

        tpl = DocumentTemplate.get_active("distribution_minutes", "trust")
        context = resolve_distribution_minutes(financial_year_id)
        renderer = TemplateRenderer(tpl.structure, context)
        buffer = renderer.render()
    """

    def __init__(self, structure: dict, context: dict):
        """
        Args:
            structure: The JSON template definition from DocumentTemplate.structure
            context: Resolved merge field dict from a resolver function
        """
        self.structure = structure
        self.context = context
        self.styles = {**DEFAULT_STYLES, **structure.get("styles", {})}
        self.metadata = structure.get("metadata", {})
        self.doc = Document()

    def render(self) -> io.BytesIO:
        """Render the template and return a BytesIO buffer containing the .docx."""
        self._apply_page_setup()
        self._apply_default_font()

        sections = self.structure.get("sections", [])
        self._render_sections(sections)

        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    # -------------------------------------------------------------------------
    # Page Setup
    # -------------------------------------------------------------------------
    def _apply_page_setup(self):
        """Apply page setup from metadata."""
        page = self.metadata.get("page_setup", {})
        section = self.doc.sections[0]

        orientation = page.get("orientation", "portrait")
        if orientation == "landscape":
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width

        # Margins in cm
        for attr, key in [
            ("top_margin", "margin_top"),
            ("bottom_margin", "margin_bottom"),
            ("left_margin", "margin_left"),
            ("right_margin", "margin_right"),
        ]:
            val = page.get(key)
            if val is not None:
                setattr(section, attr, Cm(float(val)))

    def _apply_default_font(self):
        """Set the default font for the document."""
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = self.styles["font_name"]
        font.size = Pt(self.styles["font_size_body"])
        # Set paragraph spacing
        pf = style.paragraph_format
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

    # -------------------------------------------------------------------------
    # Section Dispatcher
    # -------------------------------------------------------------------------
    def _render_sections(self, sections: list):
        """Iterate through sections and dispatch to the appropriate renderer."""
        for section in sections:
            section_type = section.get("type", "")
            handler = getattr(self, f"_render_{section_type}", None)
            if handler:
                handler(section)
            else:
                logger.warning(f"Unknown section type: {section_type}")

    # -------------------------------------------------------------------------
    # Section Renderers
    # -------------------------------------------------------------------------
    def _render_heading(self, section: dict):
        """Render a heading section."""
        text = self._resolve_text(section.get("text", ""))
        level = section.get("level", 1)
        alignment = section.get("alignment", "left")

        # Map level to font size
        size_map = {1: "font_size_heading", 2: "font_size_subheading", 3: "font_size_body"}
        font_size = Pt(self.styles.get(size_map.get(level, "font_size_body"), 11))

        para = self.doc.add_paragraph()
        para.alignment = ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_before = Pt(section.get("space_before", 6))
        para.paragraph_format.space_after = Pt(section.get("space_after", 4))

        run = para.add_run(text)
        run.font.name = self.styles["font_name"]
        run.font.size = font_size
        run.font.bold = section.get("bold", True)
        run.font.italic = section.get("italic", False)

        color_hex = section.get("color", self.styles.get("heading_color", "000000"))
        run.font.color.rgb = RGBColor.from_string(color_hex)

        # Underline
        if section.get("underline", False):
            run.font.underline = True

    def _render_paragraph(self, section: dict):
        """Render a paragraph with merge field substitution."""
        text = self._resolve_text(section.get("text", ""))
        alignment = section.get("alignment", "left")
        bold = section.get("bold", False)
        italic = section.get("italic", False)

        para = self.doc.add_paragraph()
        para.alignment = ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        para.paragraph_format.space_before = Pt(section.get("space_before", 0))
        para.paragraph_format.space_after = Pt(section.get("space_after", 4))

        # Support mixed formatting via "runs" array
        runs_def = section.get("runs")
        if runs_def:
            for run_def in runs_def:
                run_text = self._resolve_text(run_def.get("text", ""))
                run = para.add_run(run_text)
                run.font.name = self.styles["font_name"]
                run.font.size = Pt(run_def.get("font_size", self.styles["font_size_body"]))
                run.font.bold = run_def.get("bold", bold)
                run.font.italic = run_def.get("italic", italic)
                if run_def.get("color"):
                    run.font.color.rgb = RGBColor.from_string(run_def["color"])
        else:
            run = para.add_run(text)
            run.font.name = self.styles["font_name"]
            run.font.size = Pt(section.get("font_size", self.styles["font_size_body"]))
            run.font.bold = bold
            run.font.italic = italic

    def _render_paragraph_list(self, section: dict):
        """Render multiple paragraphs from an items array."""
        items = section.get("items", [])
        for item in items:
            self._render_paragraph(item)

    def _render_table(self, section: dict):
        """Render a table from column definitions and a data_source merge field."""
        columns = section.get("columns", [])
        data_source = section.get("data_source", "")
        rows_data = self.context.get(data_source, [])

        logger.info(
            "renderer table: data_source=%s rows=%d columns=%s",
            data_source, len(rows_data) if rows_data else 0,
            [c.get("field") for c in columns],
        )
        if rows_data:
            logger.info("renderer table: first_row=%r", rows_data[0])

        if not columns:
            return

        # Create table
        num_cols = len(columns)
        table = self.doc.add_table(rows=1, cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Column widths
        for i, col_def in enumerate(columns):
            width = col_def.get("width_cm")
            if width:
                table.columns[i].width = Cm(float(width))

        # Header row
        header_bg = self.styles.get("table_header_bg", "333333")
        header_fg = self.styles.get("table_header_fg", "FFFFFF")
        for i, col_def in enumerate(columns):
            cell = table.rows[0].cells[i]
            cell.text = col_def.get("header", "")
            for p in cell.paragraphs:
                p.alignment = ALIGNMENT_MAP.get(
                    col_def.get("header_alignment", "center"),
                    WD_ALIGN_PARAGRAPH.CENTER,
                )
                for run in p.runs:
                    run.font.name = self.styles["font_name"]
                    run.font.size = Pt(self.styles.get("font_size_small", 9))
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(header_fg)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{header_bg}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        # Data rows
        for row_idx, row_data in enumerate(rows_data):
            data_row = table.add_row()
            # Alternate row shading
            alt_bg = self.styles.get("table_alt_row_bg")
            if alt_bg and row_idx % 2 == 1:
                for cell in data_row.cells:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{alt_bg}"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

            for i, col_def in enumerate(columns):
                field = col_def.get("field", "")
                cell = data_row.cells[i]
                value = row_data.get(field, "")

                # Format the value
                fmt = col_def.get("format")
                if fmt == "money" and value is not None and value != "":
                    try:
                        val = Decimal(str(value))
                        if val < 0:
                            cell_text = f"({abs(val):,.2f})"
                        elif val == 0:
                            cell_text = "-"
                        else:
                            cell_text = f"{val:,.2f}"
                    except Exception:
                        cell_text = str(value)
                elif fmt == "percentage" and value is not None and value != "":
                    try:
                        val = Decimal(str(value)) * 100
                        cell_text = f"{val:.2f}%"
                    except Exception:
                        cell_text = str(value)
                else:
                    cell_text = str(value) if value is not None else ""

                cell.text = cell_text
                alignment = col_def.get("alignment", "left")
                for p in cell.paragraphs:
                    p.alignment = ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
                    for run in p.runs:
                        run.font.name = self.styles["font_name"]
                        run.font.size = Pt(self.styles.get("font_size_small", 9))

                # Bold for specific columns
                if col_def.get("bold"):
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.font.bold = True

        # Total row (if defined)
        totals = section.get("totals")
        if totals:
            total_row = table.add_row()
            for i, col_def in enumerate(columns):
                cell = total_row.cells[i]
                field = col_def.get("field", "")
                total_val = totals.get(field)
                if total_val is not None:
                    # Resolve merge field in total value
                    if isinstance(total_val, str) and "{{" in total_val:
                        total_val = self._resolve_text(total_val)
                    cell.text = str(total_val)
                elif i == 0:
                    cell.text = totals.get("label", "TOTAL")
                for p in cell.paragraphs:
                    p.alignment = ALIGNMENT_MAP.get(
                        col_def.get("alignment", "left"),
                        WD_ALIGN_PARAGRAPH.LEFT,
                    )
                    for run in p.runs:
                        run.font.name = self.styles["font_name"]
                        run.font.size = Pt(self.styles.get("font_size_small", 9))
                        run.font.bold = True

        # Spacing after table
        self.doc.add_paragraph()

    def _render_spacer(self, section: dict):
        """Render vertical whitespace."""
        lines = section.get("lines", 1)
        for _ in range(lines):
            self.doc.add_paragraph()

    def _render_page_break(self, section: dict):
        """Insert a page break."""
        self.doc.add_page_break()

    def _render_horizontal_rule(self, section: dict):
        """Render a thin horizontal line."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        pPr = para._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="auto"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    def _render_conditional(self, section: dict):
        """Conditionally render children if a merge field is truthy."""
        field = section.get("field", "")
        negate = section.get("negate", False)
        value = self.context.get(field)

        condition = bool(value)
        if negate:
            condition = not condition

        logger.info(
            "renderer conditional: field=%s value=%r condition=%s negate=%s children=%d else=%d",
            field, value, condition, negate,
            len(section.get("children", [])),
            len(section.get("else_children", [])),
        )

        if condition:
            children = section.get("children", [])
            self._render_sections(children)
        else:
            # Render else_children if provided
            else_children = section.get("else_children", [])
            if else_children:
                self._render_sections(else_children)

    def _render_signature_block(self, section: dict):
        """Render a signature block with name, title, and date."""
        name_field = section.get("name_field", "")
        title_field = section.get("title_field", "")
        date_field = section.get("date_field", "")
        label = section.get("label", "Signed")

        name = self._resolve_field(name_field)
        title = self._resolve_field(title_field)
        date_val = self._resolve_field(date_field)

        # Signature line
        self.doc.add_paragraph()
        para = self.doc.add_paragraph()
        run = para.add_run("_" * 40)
        run.font.name = self.styles["font_name"]
        run.font.size = Pt(self.styles["font_size_body"])

        # Name
        para = self.doc.add_paragraph()
        run = para.add_run(str(name))
        run.font.name = self.styles["font_name"]
        run.font.size = Pt(self.styles["font_size_body"])
        run.font.bold = True

        # Title
        if title:
            para = self.doc.add_paragraph()
            run = para.add_run(str(title))
            run.font.name = self.styles["font_name"]
            run.font.size = Pt(self.styles["font_size_body"])

        # Date
        if date_val:
            para = self.doc.add_paragraph()
            run = para.add_run(f"Date: {date_val}")
            run.font.name = self.styles["font_name"]
            run.font.size = Pt(self.styles["font_size_body"])

    def _render_disclaimer(self, section: dict):
        """Render an italic disclaimer block."""
        text = self._resolve_text(section.get("text", ""))
        heading = section.get("heading", "DISCLAIMER")

        if heading:
            h_para = self.doc.add_paragraph()
            h_run = h_para.add_run(heading)
            h_run.font.name = self.styles["font_name"]
            h_run.font.size = Pt(self.styles["font_size_body"])
            h_run.font.bold = True

        para = self.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        run.font.name = self.styles["font_name"]
        run.font.size = Pt(self.styles.get("font_size_small", 9))
        run.font.italic = True

    def _render_numbered_list(self, section: dict):
        """Render a numbered list."""
        items = section.get("items", [])
        for i, item in enumerate(items, 1):
            text = self._resolve_text(item.get("text", "") if isinstance(item, dict) else str(item))
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(f"{i}. {text}")
            run.font.name = self.styles["font_name"]
            run.font.size = Pt(self.styles["font_size_body"])
            if isinstance(item, dict):
                run.font.bold = item.get("bold", False)
                run.font.italic = item.get("italic", False)

    def _render_bullet_list(self, section: dict):
        """Render a bullet list."""
        items = section.get("items", [])
        for item in items:
            text = self._resolve_text(item.get("text", "") if isinstance(item, dict) else str(item))
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            run = para.add_run(f"\u2022  {text}")
            run.font.name = self.styles["font_name"]
            run.font.size = Pt(self.styles["font_size_body"])
            if isinstance(item, dict):
                run.font.bold = item.get("bold", False)
                run.font.italic = item.get("italic", False)

    def _render_firm_header(self, section: dict):
        """Render the firm header block (logo area + firm details)."""
        firm_name = section.get("firm_name", "M C & S Pty Ltd")
        lines = section.get("lines", [])

        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(firm_name)
        run.font.name = self.styles["font_name"]
        run.font.size = Pt(self.styles.get("font_size_heading", 14))
        run.font.bold = True

        for line in lines:
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(self._resolve_text(line))
            run.font.name = self.styles["font_name"]
            run.font.size = Pt(self.styles.get("font_size_small", 9))

    def _render_key_value_table(self, section: dict):
        """Render a simple 2-column key-value table (e.g. summary metrics)."""
        items = section.get("items", [])
        if not items:
            return

        table = self.doc.add_table(rows=len(items), cols=2)
        table.style = "Table Grid"

        for i, item in enumerate(items):
            label = self._resolve_text(item.get("label", ""))
            value_field = item.get("value", "")
            if "{{" in str(value_field):
                value = self._resolve_text(value_field)
            elif value_field in self.context:
                # Bare field name — leave empty so the context-lookup/format
                # block below resolves it (previously str(value_field) shadowed
                # the field name and made the fallback unreachable).
                value = ""
            else:
                value = str(value_field)

            # If value is a merge field name (no braces), resolve it
            if not value and value_field in self.context:
                raw = self.context[value_field]
                fmt = item.get("format")
                if fmt == "money":
                    try:
                        val = Decimal(str(raw))
                        value = f"${val:,.2f}" if val >= 0 else f"(${abs(val):,.2f})"
                    except Exception:
                        value = str(raw)
                elif fmt == "percentage":
                    try:
                        val = Decimal(str(raw)) * 100
                        value = f"{val:.2f}%"
                    except Exception:
                        value = str(raw)
                else:
                    value = str(raw)

            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            table.rows[i].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            for cell in [table.rows[i].cells[0], table.rows[i].cells[1]]:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = self.styles["font_name"]
                        run.font.size = Pt(self.styles["font_size_body"])
                        if item.get("bold"):
                            run.font.bold = True

        self.doc.add_paragraph()

    # -------------------------------------------------------------------------
    # Merge Field Resolution
    # -------------------------------------------------------------------------
    def _resolve_text(self, text: str) -> str:
        """Replace all {{field_name}} placeholders in text with context values."""
        if not text or "{{" not in text:
            return text or ""

        def replacer(match):
            field_name = match.group(1)
            value = self.context.get(field_name, f"{{{{{field_name}}}}}")
            if isinstance(value, Decimal):
                # Route through the shared money formatter so negatives render as
                # ($100.00) rather than $-100.00.
                return _fmt_money(value)
            if isinstance(value, (list, dict)):
                return str(value)
            return str(value)

        return re.sub(r"\{\{(\w+)\}\}", replacer, text)

    def _resolve_field(self, field_name: str):
        """Resolve a single field name from context."""
        if not field_name:
            return ""
        return self.context.get(field_name, "")


# =============================================================================
# Public API
# =============================================================================
def render_document_template(template_id=None, document_category=None,
                              entity_type="", context=None,
                              financial_year_id=None) -> io.BytesIO:
    """
    High-level API to render a document from a DocumentTemplate.

    Args:
        template_id: UUID of a specific DocumentTemplate (optional)
        document_category: DocumentCategory value to look up active template
        entity_type: Entity type to filter templates
        context: Pre-resolved merge field dict (if None, will use resolver)
        financial_year_id: UUID of FinancialYear (used by resolver if context is None)

    Returns:
        io.BytesIO buffer containing the rendered .docx file
    """
    from core.models import DocumentTemplate as DTModel

    # Get the template
    if template_id:
        tpl = DTModel.objects.get(pk=template_id)
    elif document_category:
        tpl = DTModel.get_active(document_category, entity_type)
        if not tpl:
            raise ValueError(
                f"No active template found for category='{document_category}', "
                f"entity_type='{entity_type}'"
            )
    else:
        raise ValueError("Either template_id or document_category must be provided.")

    # Resolve context if not provided
    if context is None:
        if financial_year_id is None:
            raise ValueError("financial_year_id is required when context is not provided.")
        from core.template_resolvers import resolve_context
        context = resolve_context(tpl.document_category, financial_year_id)

    # Render
    renderer = TemplateRenderer(tpl.structure, context)
    return renderer.render()
