"""The balance sheet's three heading levels are visually distinct.

HandiLedger runs four sizes down the page -- DJLH Properties Pty Ltd FY2024
sets the section in Arial Bold 19, the sub-group in Arial Bold 17, the nested
tier in its body face bold at 15, and the account lines in that face at 15::

    Current Liabilities        <- section
    Payables                   <- sub-group
    Unsecured:                 <- nested tier
    Trade creditors                     1,291    12,353
                                        1,291    12,353

StatementHub styled the sub-group and the tier identically -- both simply
bolded at the body size -- so "Payables" and "Unsecured:" read as one level
and the nesting was invisible. The template is single-face Arial, so the
levels are separated by size: the section keeps FONT_SIZE_HEADING, the
sub-group sits between at ``_SUB_HEADING_SIZE``, and the tier stays at the
body size, bold, marked by its trailing colon.
"""
import io

from django.test import SimpleTestCase
from docx import Document
from docx.shared import Pt

from core.fs_template_service import _post_process_fs_doc


def _doc_with_rows(labels):
    """A 4-column balance-sheet-shaped table, one row per label."""
    document = Document()
    table = document.add_table(rows=0, cols=4)
    for label, cy, py in labels:
        cells = table.add_row().cells
        cells[0].text = label
        cells[2].text = cy
        cells[3].text = py
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _rows_by_label(document):
    out = {}
    for table in document.tables:
        for row in table.rows:
            runs = [r for p in row.cells[0].paragraphs for r in p.runs]
            out[row.cells[0].text.strip()] = runs
    return out


class HeadingLevelsAreDistinctTests(SimpleTestCase):
    def setUp(self):
        buffer = _doc_with_rows([
            ("Payables", "", ""),
            ("Unsecured:", "", ""),
            ("Trade creditors", "1,291", "12,353"),
            ("Issued Capital", "", ""),
            ("Financial Liabilities", "", ""),
            ("Secured:", "", ""),
        ])
        processed = _post_process_fs_doc(buffer, "BALANCE_SHEET")
        self.rows = _rows_by_label(Document(processed))

    def test_a_sub_group_heading_is_bold(self):
        self.assertTrue(all(r.bold for r in self.rows["Payables"]))

    def test_a_sub_group_heading_sits_above_the_body_size(self):
        sizes = {r.font.size for r in self.rows["Financial Liabilities"]}
        self.assertEqual(sizes, {Pt(10.5)})

    def test_the_equity_heading_gets_the_same_treatment(self):
        sizes = {r.font.size for r in self.rows["Issued Capital"]}
        self.assertEqual(sizes, {Pt(10.5)})

    def test_a_nested_tier_is_bold(self):
        self.assertTrue(all(r.bold for r in self.rows["Unsecured:"]))

    def test_a_nested_tier_stays_at_the_body_size(self):
        """Sizing it up would flatten the two levels back into one."""
        sizes = {r.font.size for r in self.rows["Secured:"]}
        self.assertNotIn(Pt(10.5), sizes)

    def test_an_account_line_is_left_alone(self):
        runs = self.rows["Trade creditors"]
        self.assertFalse(any(r.bold for r in runs))
        self.assertNotIn(Pt(10.5), {r.font.size for r in runs})


class OnlyTheBalanceSheetHasSubGroupsTests(SimpleTestCase):
    """Sub-group headings exist on the balance sheet and nowhere else.

    The heading labels are the balance sheet's own group names, and one of
    them is the bare word "Other" -- a perfectly ordinary expense account
    name. The styling loop runs over the Detailed P&L and Summary P&L too, so
    without a document-type gate an expense called "Other" would be bolded and
    sized up as though it headed a group that does not exist there.
    """

    def _pl_rows(self):
        buffer = _doc_with_rows([
            ("Other", "1,234", "2,345"),
            ("Unsecured:", "", ""),
        ])
        return _rows_by_label(Document(_post_process_fs_doc(buffer, "DETAILED_PL")))

    def test_a_pl_account_named_other_is_left_alone(self):
        runs = self._pl_rows()["Other"]
        self.assertFalse(any(r.bold for r in runs))
        self.assertNotIn(Pt(10.5), {r.font.size for r in runs})

    def test_a_pl_row_is_not_treated_as_a_nested_tier(self):
        runs = self._pl_rows()["Unsecured:"]
        self.assertFalse(any(r.bold for r in runs))

    def test_the_balance_sheet_still_styles_the_same_label(self):
        buffer = _doc_with_rows([("Other", "1,234", "2,345")])
        rows = _rows_by_label(
            Document(_post_process_fs_doc(buffer, "BALANCE_SHEET")))
        self.assertEqual({r.font.size for r in rows["Other"]}, {Pt(10.5)})
