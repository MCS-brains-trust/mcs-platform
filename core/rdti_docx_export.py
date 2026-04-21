"""
R&D Tax Incentive (RDTI) Drafter — .docx Export
Generates a formatted Word document matching the AusIndustry portal print layout.
"""
import io
import logging
from datetime import date

logger = logging.getLogger(__name__)

DISCLAIMER = (
    "AI-GENERATED DRAFT — CONSULTANT REVIEW REQUIRED\n\n"
    "This document is an AI-generated draft. The consultant must review for factual accuracy, "
    "completeness, and compliance with the Industry Research and Development Act 1986 and the "
    "AusIndustry Guide to Interpretation. This tool is a drafting aid; responsibility for the "
    "lodged registration remains with the registered tax agent or the R&D entity."
)


def generate_rdti_docx(application) -> bytes:
    """
    Generate a .docx export of the RDTI application.
    Returns bytes of the Word document.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("python-docx is required for RDTI export. Install with: pip install python-docx")

    doc = Document()

    # --- Page margins ---
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # --- Styles ---
    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        if level == 1:
            p.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        elif level == 2:
            p.runs[0].font.color.rgb = RGBColor(0x6c, 0x8c, 0xff)
        return p

    def add_field_block(label, content, char_count=None):
        """Add a labelled narrative field block."""
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        if char_count is not None:
            run2 = p.add_run(f"  ({char_count}/4,000 characters)")
            run2.font.size = Pt(8)
            run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            if char_count > 4000:
                run2.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        content_para = doc.add_paragraph(content or "[Not drafted]")
        content_para.style.font.size = Pt(10)

        # Add thin border below
        doc.add_paragraph()

    def add_structured_field(label, value):
        """Add a structured (non-narrative) field."""
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(str(value) if value else "—").font.size = Pt(10)

    # =========================================================
    # COVER PAGE
    # =========================================================
    doc.add_heading("R&D Tax Incentive Registration", 0)

    entity = application.financial_year.entity
    fy = application.financial_year

    add_structured_field("Entity", entity.entity_name)
    add_structured_field("ABN", application.abn or entity.abn or "—")
    add_structured_field("ACN", application.acn or entity.acn or "—")
    add_structured_field("Financial Year", fy.year_label)
    add_structured_field("Period", f"{fy.start_date} to {fy.end_date}")
    add_structured_field("Application Status", application.get_status_display())
    add_structured_field("Contact", application.contact_name or "—")
    add_structured_field("Contact Email", application.contact_email or "—")
    add_structured_field("ANZSIC Division", application.anzsic_division or "—")
    add_structured_field("ANZSIC Code", application.anzsic_code or "—")

    if application.aggregated_turnover:
        offset_type = "Refundable (turnover < $20M)" if application.is_refundable else "Non-refundable (turnover ≥ $20M)"
        add_structured_field("Aggregated Turnover", f"${application.aggregated_turnover:,.0f} — {offset_type}")

    doc.add_page_break()

    # =========================================================
    # DISCLAIMER
    # =========================================================
    p = doc.add_paragraph(DISCLAIMER)
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    doc.add_page_break()

    # =========================================================
    # PROJECTS
    # =========================================================
    for project in application.projects.prefetch_related('core_activities__supporting_activities').all():
        add_heading(f"Project: {project.project_title}", level=1)

        add_structured_field("Project Start Date", project.project_start_date or "—")
        add_structured_field("Project End Date", project.project_end_date or "—")
        add_structured_field("ANZSRC Division", project.anzsrc_division or "—")
        add_structured_field("ANZSRC Code", project.anzsrc_code or "—")
        doc.add_paragraph()

        # Project narrative fields
        add_heading("Project Narrative Fields", level=2)
        add_field_block("Objectives", project.objectives, len(project.objectives) if project.objectives else 0)
        add_field_block("Documents Kept", project.documents_kept, len(project.documents_kept) if project.documents_kept else 0)
        add_field_block("Plant and Facilities", project.plant_and_facilities, len(project.plant_and_facilities) if project.plant_and_facilities else 0)
        add_field_block("Beneficiary Description", project.beneficiary_description, len(project.beneficiary_description) if project.beneficiary_description else 0)

        # =========================================================
        # CORE ACTIVITIES
        # =========================================================
        for i, activity in enumerate(project.core_activities.all(), 1):
            doc.add_page_break()
            add_heading(f"Core Activity {i}: {activity.activity_title}", level=2)

            add_structured_field("Start Date", activity.activity_start_date or "—")
            add_structured_field("End Date", activity.activity_end_date or "—")
            add_structured_field("Performed By", activity.get_performed_by_display())
            add_structured_field("Sources Investigated", ", ".join(activity.sources_investigated) if activity.sources_investigated else "—")
            add_structured_field("Evidence Kept", ", ".join(activity.evidence_kept) if activity.evidence_kept else "—")
            doc.add_paragraph()

            # The 8 narrative fields
            narrative_fields = [
                ("Description of Core R&D Activity", activity.description),
                ("How Outcome Could Not Be Known in Advance", activity.outcome_not_known_in_advance),
                ("Why a Competent Professional Could Not Have Known", activity.competent_professional),
                ("Hypothesis", activity.hypothesis),
                ("Experiment", activity.experiment),
                ("Evaluation Method", activity.evaluation_method),
                ("Conclusions", activity.conclusions),
                ("New Knowledge Produced", activity.new_knowledge),
            ]

            for label, content in narrative_fields:
                add_field_block(label, content, len(content) if content else 0)

            # Expenditure breakdown
            expenditures = activity.expenditure_years.all()
            if expenditures:
                add_heading("Expenditure Breakdown", level=3)
                table = doc.add_table(rows=1, cols=6)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Financial Year"
                hdr[1].text = "Labour"
                hdr[2].text = "Contractors"
                hdr[3].text = "Overheads"
                hdr[4].text = "Other"
                hdr[5].text = "Total"
                for exp in expenditures:
                    row = table.add_row().cells
                    row[0].text = exp.financial_year_label
                    row[1].text = f"${exp.labour_expenditure:,.0f}" if exp.labour_expenditure else "—"
                    row[2].text = f"${exp.contractor_expenditure:,.0f}" if exp.contractor_expenditure else "—"
                    row[3].text = f"${exp.overhead_expenditure:,.0f}" if exp.overhead_expenditure else "—"
                    row[4].text = f"${exp.other_expenditure:,.0f}" if exp.other_expenditure else "—"
                    row[5].text = f"${exp.total:,.0f}"
                doc.add_paragraph()

            # Supporting Activities
            supporting = activity.supporting_activities.all()
            if supporting:
                add_heading("Supporting Activities", level=3)
                for j, sa in enumerate(supporting, 1):
                    add_heading(f"Supporting Activity {j}: {sa.activity_title}", level=4)
                    add_field_block("Description", sa.description, len(sa.description) if sa.description else 0)
                    add_field_block("Direct Relation to Core Activity", sa.direct_relation, len(sa.direct_relation) if sa.direct_relation else 0)

    # =========================================================
    # FLAGS SUMMARY
    # =========================================================
    doc.add_page_break()
    add_heading("Compliance Flags Summary", level=1)

    flags = application.flags.filter(is_resolved=False).order_by("severity", "field_name")
    if not flags.exists():
        doc.add_paragraph("No unresolved compliance flags.")
    else:
        for flag in flags:
            p = doc.add_paragraph()
            severity_colors = {
                "red": RGBColor(0xDC, 0x26, 0x26),
                "amber": RGBColor(0xD9, 0x77, 0x06),
                "green": RGBColor(0x16, 0xA3, 0x4A),
            }
            run = p.add_run(f"[{flag.severity.upper()}] {flag.field_name}: ")
            run.bold = True
            run.font.color.rgb = severity_colors.get(flag.severity, RGBColor(0, 0, 0))
            p.add_run(flag.message)
            if flag.suggestion:
                p2 = doc.add_paragraph(f"  → {flag.suggestion}")
                p2.runs[0].font.size = Pt(9)
                p2.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
