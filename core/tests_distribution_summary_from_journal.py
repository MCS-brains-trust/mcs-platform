"""The Distribution Summary reports the ledger, not a scenario nobody sets.

The summary's data source (now ``_posted_distribution_data``, formerly
``_get_confirmed_scenario_data``) read ``workspace.confirmed_scenario`` -- a
``DistributionScenario``. That field is None on every workspace on the
platform: the live flow selects a ``TaxPlanningScenario`` and posts a journal,
and nothing ever confirms a DistributionScenario. So the helper always returned
no rows, and the generated DOCX printed a heading, a Net Distributable Income
figure and an empty table for the five trusts that had genuinely distributed
(Vincent FY2024 17,834.03, Vincent FY2025 61,351.87, Dr Services FY2026
61,848.01, E & J Chiaravalle FY2025 50,613.00, Chiaravalle Family FY2025
49,242.00).

The NDI it did print came from ``workspace.net_distributable_income``, the
stored snapshot -- Minli FY2026 held 876,322.95 from superseded calculations
against a true nil, which is why the post gate already refuses to trust it.

This is the fault PR #102 fixed in ``_figures_for_year`` for the FS pack: a
document reporting the plan rather than what was posted. The journal records
per-beneficiary totals only and carries no stream character, so the summary
reports beneficiaries and amounts -- not a stream grid it would have to invent.
"""
from datetime import date
from decimal import Decimal

from django.contrib.auth import get_user_model
from django.test import TestCase

from core.models import (
    AccountMapping, AdjustingJournal, Entity, EntityOfficer, FinancialYear,
    JournalLine, TrialBalanceLine, TrustWorkspace,
)
from core.views_trust import _posted_distribution_data


class DistributionSummaryFromJournalTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="dsum", email="dsum@example.com", password="secret123",
            role="senior_accountant",
            totp_secret="dummy-secret-for-test", totp_confirmed=True,
        )
        self.entity = Entity.objects.create(
            entity_name="Summary Trust", entity_type="trust_unit",
            assigned_accountant=self.user,
        )
        self.fy = FinancialYear.objects.create(
            entity=self.entity, year_label="FY2024",
            start_date=date(2023, 7, 1), end_date=date(2024, 6, 30),
        )
        self.workspace = TrustWorkspace.objects.create(financial_year=self.fy)

    @staticmethod
    def _mapping(section):
        m, _ = AccountMapping.objects.get_or_create(
            standard_code=f"DSUM-{section}",
            defaults={"line_item_label": section.title(),
                      "financial_statement": "income_statement",
                      "statement_section": section},
        )
        return m

    def _earn(self, amount):
        TrialBalanceLine.objects.create(
            financial_year=self.fy, account_code="0630", account_name="Sales",
            closing_balance=-Decimal(amount), debit=Decimal("0"),
            credit=Decimal(amount), source="tb_import",
            mapped_line_item=self._mapping("revenue"),
        )

    def _post_distribution(self, allocations, reference="JE-002"):
        """Post a distribution the way trust_post_distribution does."""
        total = sum(Decimal(a) for _, a in allocations)
        journal = AdjustingJournal.objects.create(
            financial_year=self.fy, reference_number=reference,
            journal_type=AdjustingJournal.JournalType.GENERAL,
            status=AdjustingJournal.JournalStatus.POSTED,
            journal_date=self.fy.end_date,
            description="Trust distribution", is_trust_distribution=True,
            total_debit=total, total_credit=total,
        )
        JournalLine.objects.create(
            journal=journal, line_number=1, account_code="4199",
            account_name="Undistributed income",
            description="Trust distribution for year",
            debit=total, credit=Decimal("0"),
        )
        TrialBalanceLine.objects.create(
            financial_year=self.fy, account_code="4199",
            account_name="Undistributed income", closing_balance=total,
            debit=total, credit=Decimal("0"), source="manual_journal",
            source_journal=journal,
        )
        for i, (name, amount) in enumerate(allocations, start=2):
            officer = EntityOfficer.objects.create(
                entity=self.entity, full_name=name,
                role=EntityOfficer.OfficerRole.BENEFICIARY,
                roles=["beneficiary"],
            )
            code = f"4004.{i:02d}"
            JournalLine.objects.create(
                journal=journal, line_number=i, account_code=code,
                account_name=f"Distribution — {name}",
                description=f"{name}: ${Decimal(amount):,.2f}",
                debit=Decimal("0"), credit=Decimal(amount),
            )
            TrialBalanceLine.objects.create(
                financial_year=self.fy, account_code=code,
                account_name=f"Distribution — {name}",
                closing_balance=-Decimal(amount), debit=Decimal("0"),
                credit=Decimal(amount), source="manual_journal",
                source_journal=journal,
            )
        return journal

    def test_a_posted_distribution_produces_rows(self):
        """The five live trusts that distributed all rendered an empty table."""
        self._earn("17834.03")
        self._post_distribution([("Michael Vincent", "17834.03")])
        rows, total, ndi = _posted_distribution_data(self.workspace)
        self.assertEqual(
            len(rows), 1,
            "a posted distribution produced no rows, so the summary is blank",
        )
        self.assertEqual(rows[0]["name"], "Michael Vincent")
        self.assertEqual(rows[0]["total"], Decimal("17834.03"))
        self.assertEqual(total, Decimal("17834.03"))

    def test_every_beneficiary_of_the_journal_appears(self):
        self._earn("100000.00")
        self._post_distribution([("Ana", "60000.00"), ("Bruno", "40000.00")])
        rows, total, _ = _posted_distribution_data(self.workspace)
        self.assertEqual([r["name"] for r in rows], ["Ana", "Bruno"])
        self.assertEqual(total, Decimal("100000.00"))
        self.assertEqual(
            [r["percentage"] for r in rows],
            [Decimal("60.00"), Decimal("40.00")],
        )

    def test_the_ndi_is_recomputed_not_read_from_the_stale_snapshot(self):
        """Minli FY2026 held 876,322.95 against a true nil."""
        self._earn("216101.66")
        TrialBalanceLine.objects.create(
            financial_year=self.fy, account_code="4199",
            account_name="Undistributed income",
            closing_balance=Decimal("1628428.89"),
            debit=Decimal("1628428.89"), credit=Decimal("0"), source="rollover",
        )
        self.workspace.net_distributable_income = Decimal("876322.95")
        self.workspace.save(update_fields=["net_distributable_income"])
        _, _, ndi = _posted_distribution_data(self.workspace)
        self.assertEqual(
            ndi, Decimal("0"),
            "the summary printed a stale figure the post gate would refuse",
        )

    def test_a_reversed_distribution_reports_nothing(self):
        """What was posted, and then unposted, was not distributed."""
        self._earn("100000.00")
        journal = self._post_distribution([("Ana", "100000.00")])
        AdjustingJournal.objects.create(
            financial_year=self.fy, reference_number="JE-003",
            journal_type=AdjustingJournal.JournalType.YEAR_END,
            status=AdjustingJournal.JournalStatus.POSTED,
            journal_date=self.fy.end_date, reverses=journal,
            description="Reversal of JE-002",
        )
        rows, total, _ = _posted_distribution_data(self.workspace)
        self.assertEqual(rows, [])
        self.assertEqual(total, Decimal("0"))

    def test_nothing_posted_reports_nothing(self):
        self._earn("100000.00")
        rows, total, ndi = _posted_distribution_data(self.workspace)
        self.assertEqual(rows, [])
        self.assertEqual(total, Decimal("0"))
        self.assertEqual(ndi, Decimal("100000.00"))

    def test_the_appropriation_line_is_not_a_beneficiary(self):
        """4199 is the debit side of the journal, not a recipient."""
        self._earn("50000.00")
        self._post_distribution([("Ana", "50000.00")])
        rows, _, _ = _posted_distribution_data(self.workspace)
        self.assertNotIn(
            "4199", [r.get("account_code", "") for r in rows])
        self.assertEqual(len(rows), 1)


class DistributionSummaryDocumentTests(DistributionSummaryFromJournalTests):
    """End to end: the generated DOCX, not just the helper feeding it."""

    def setUp(self):
        super().setUp()
        self.client.force_login(self.user)
        session = self.client.session
        session["2fa_verified"] = True
        session.save()

    def _docx_text(self):
        import io
        from docx import Document
        from django.urls import reverse

        response = self.client.get(
            reverse("core:trust_generate_distribution_summary",
                    kwargs={"pk": self.fy.pk}),
            secure=True,
        )
        self.assertEqual(response.status_code, 200)
        doc = Document(io.BytesIO(response.content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(parts)

    def test_the_document_names_the_beneficiary_and_amount(self):
        self._earn("17834.03")
        self._post_distribution([("Michael Vincent", "17834.03")])
        text = self._docx_text()
        self.assertIn("Michael Vincent", text)
        self.assertIn("$17,834.03", text)
        self.assertIn("JE-002", text)

    def test_the_document_does_not_print_a_stale_figure(self):
        self._earn("216101.66")
        TrialBalanceLine.objects.create(
            financial_year=self.fy, account_code="4199",
            account_name="Undistributed income",
            closing_balance=Decimal("1628428.89"),
            debit=Decimal("1628428.89"), credit=Decimal("0"), source="rollover",
        )
        self.workspace.net_distributable_income = Decimal("876322.95")
        self.workspace.save(update_fields=["net_distributable_income"])
        text = self._docx_text()
        self.assertNotIn("876,322.95", text)
        self.assertIn("Net Distributable Income: $0.00", text)
        self.assertIn("No income was distributable", text)

    def test_a_year_with_income_but_no_posting_says_so(self):
        self._earn("100000.00")
        text = self._docx_text()
        self.assertIn("No distribution has been posted", text)
