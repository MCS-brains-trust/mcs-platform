"""MCS Platform - Core Views"""
import logging
import openpyxl
from decimal import Decimal, InvalidOperation
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db import transaction as db_transaction
from django.db.models import Q, Count, Sum
from django.http import HttpResponse, HttpResponseNotAllowed, JsonResponse
from django.views.decorators.http import require_POST
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.utils import timezone

logger = logging.getLogger("core.views")

from .models import (
    Client, Entity, FinancialYear, TrialBalanceLine,
    AccountMapping, ChartOfAccount, ClientAccountMapping, AdjustingJournal,
    JournalLine, GeneratedDocument, AuditLog, EntityOfficer,
    ClientAssociate, AccountingSoftware, MeetingNote,
    DepreciationAsset, RiskFlag, StockItem, ActivityLog, EntityChartOfAccount,
    BulkJournalUpload, BASPeriod, BankAccountMapping, BASPeriodCommentary,
    EvaReview,
)
from .forms import (
    ClientForm, EntityForm, FinancialYearForm,
    TrialBalanceUploadForm, AccountMappingForm,
    AdjustingJournalForm, JournalLineFormSet, JournalLineForm,
    EntityOfficerForm, ClientAssociateForm, AccountingSoftwareForm,
    MeetingNoteForm,
)
from django import forms
from config.authorization import get_entity_for_user, get_financial_year_for_user
import re as _re_mod


def _get_subsequent_finalised_years(fy):
    """Walk the next_year chain and return a list of subsequent finalised FYs."""
    result = []
    current = fy
    while True:
        nxt = current.next_year.first()
        if nxt and nxt.is_locked:
            result.append(nxt)
            current = nxt
        else:
            break
    return result


def _compute_amber_indicators(fy):
    """Compute amber indicators for the trial balance view.
    Returns a dict keyed by account_code with a list of trigger dicts."""
    try:
        from core.views_eva import compute_amber_indicators_for_context
        return compute_amber_indicators_for_context(fy)
    except Exception:
        return {}


def _resolve_account_name(entity, account_code, raw_name):
    """Resolve a human-readable account name, falling back through multiple sources.

    When a trial balance is imported from certain accounting packages, the
    account-name column is sometimes blank or contains a copy of the account
    code.  This helper detects that situation and attempts to find the real
    name from:
      1. EntityChartOfAccount (entity-specific)
      2. ChartOfAccount (master template for the entity type)
    If no match is found the original raw_name is returned unchanged.
    """
    # If the raw name is non-empty and clearly different from the code, keep it
    if raw_name and raw_name.strip() != account_code.strip():
        return raw_name

    # Attempt 1: entity-level chart of accounts
    ecoa = EntityChartOfAccount.objects.filter(
        entity=entity, account_code=account_code, is_active=True
    ).first()
    if ecoa and ecoa.account_name:
        return ecoa.account_name

    # Attempt 2: master chart of accounts template
    coa = ChartOfAccount.objects.filter(
        entity_type=entity.entity_type, account_code=account_code, is_active=True
    ).first()
    if coa and coa.account_name:
        return coa.account_name

    # No match found — return whatever we have (code as last resort)
    return raw_name or account_code


# ---------------------------------------------------------------------------
# HandiLedger account code range -> canonical display section.
# The numeric account code range is the AUTHORITATIVE source of truth for
# section classification on all HandiLedger / Access Ledger imports.  It
# overrides whatever statement_section the mapped_line_item carries, so that
# an account like 2475 (Retention Receivable) can never be pushed into
# Non-Current Assets simply because it was once mapped to a BS-NCA code.
# ---------------------------------------------------------------------------
_HL_RANGE_SECTION = [
    (0,    999,  'Income'),
    (1000, 1499, 'Cost of Sales'),
    (1500, 1999, 'Expenses'),
    (2000, 2499, 'Current Assets'),
    (2500, 2999, 'Non Current Assets'),
    (3000, 3499, 'Current Liabilities'),
    (3500, 3999, 'Non Current Liabilities'),
    (4000, 4999, 'Equity'),
]


def _hl_section_for_code(account_code):
    """
    Return the canonical display section for a HandiLedger account code
    based purely on the numeric range.  Returns None if the code cannot
    be resolved (e.g. non-numeric or out of the known ranges).
    """
    try:
        code_int = int(str(account_code).split('.')[0])
    except (ValueError, TypeError):
        return None
    for lo, hi, section in _HL_RANGE_SECTION:
        if lo <= code_int <= hi:
            return section
    return None


# Mapping from EntityChartOfAccount / ChartOfAccount section values to
# the display section names used in the trial balance UI.
_COA_SECTION_TO_DISPLAY = {
    'revenue': 'Income',
    'cost_of_sales': 'Cost of Sales',
    'expenses': 'Expenses',
    'current_assets': 'Current Assets',
    'non_current_assets': 'Non Current Assets',
    'assets': 'Current Assets',
    'current_liabilities': 'Current Liabilities',
    'non_current_liabilities': 'Non Current Liabilities',
    'liabilities': 'Current Liabilities',
    'equity': 'Equity',
    'capital_accounts': 'Equity',
    'pl_appropriation': 'Equity',
    'suspense': 'Unmapped',
}

# Which CoA sections are P&L (used for Net Profit on unmapped lines)
_PL_COA_SECTIONS = {'revenue', 'cost_of_sales', 'expenses'}

# Sections that are balance sheet (used for roll-forward filtering)
_BS_SECTIONS = {"assets", "liabilities", "equity", "capital_accounts",
                "current_assets", "non_current_assets",
                "current_liabilities", "non_current_liabilities"}
_BS_STATEMENTS = {"balance_sheet", "equity"}


def _is_balance_sheet_account(account_code, mapped_line_item, coa_sections):
    """Classify an account as balance sheet or P&L.

    Income statement accounts (revenue/expense) are excluded from roll
    forward — they reset to zero at year end.  Only balance sheet accounts
    (assets, liabilities, equity) carry forward.

    Uses the same three-tier classification as roll_forward:
      1. HandiLedger numeric code range (authoritative)
      2. mapped_line_item.financial_statement
      3. ChartOfAccount section lookup
      4. Numeric fallback (code >= 2000 = BS)
    """
    hl_sec = _hl_section_for_code(account_code)
    if hl_sec is not None:
        return hl_sec not in ('Income', 'Cost of Sales', 'Expenses')
    if mapped_line_item:
        return getattr(mapped_line_item, 'financial_statement', '') in _BS_STATEMENTS
    if account_code in coa_sections:
        return coa_sections[account_code] in _BS_SECTIONS
    code_prefix = account_code.split(".")[0] if account_code else ""
    return code_prefix.isdigit() and int(code_prefix) >= 2000


def _build_coa_section_lookup(entity):
    """Build a dict mapping account_code -> display_section for an entity.

    Uses the entity's chart of accounts first, then falls back to the master
    template.  Returns a dict that can be used to look up the display section
    for any account code without hitting the database per-line.
    """
    lookup = {}

    # Priority 2 first (lower priority — will be overwritten by Priority 1)
    for code, section in ChartOfAccount.objects.filter(
        entity_type=entity.entity_type, is_active=True,
    ).values_list('account_code', 'section'):
        ds = _COA_SECTION_TO_DISPLAY.get(section, 'Unmapped')
        if ds != 'Unmapped':
            lookup[code] = ds

    # Priority 1: entity-level chart of accounts (overwrites template)
    for code, section in EntityChartOfAccount.objects.filter(
        entity=entity, is_active=True,
    ).values_list('account_code', 'section'):
        ds = _COA_SECTION_TO_DISPLAY.get(section, 'Unmapped')
        if ds != 'Unmapped':
            lookup[code] = ds

    return lookup


def _infer_section_for_unmapped(entity, account_code):
    """Infer the display section for an unmapped TB line by looking up the
    account code in the entity's chart of accounts (or the master template).

    Returns a display section string (e.g. 'Income', 'Expenses', 'Current Assets')
    or 'Unmapped' if no match is found.

    NOTE: For batch processing, prefer _build_coa_section_lookup() to avoid
    N+1 queries.
    """
    # Priority 1: entity-level chart of accounts
    ecoa = EntityChartOfAccount.objects.filter(
        entity=entity, account_code=account_code, is_active=True,
    ).first()
    if ecoa and ecoa.section:
        return _COA_SECTION_TO_DISPLAY.get(ecoa.section, 'Unmapped')

    # Priority 2: master chart of accounts template
    coa = ChartOfAccount.objects.filter(
        entity_type=entity.entity_type, account_code=account_code, is_active=True,
    ).first()
    if coa and coa.section:
        return _COA_SECTION_TO_DISPLAY.get(coa.section, 'Unmapped')

    return 'Unmapped'


def _apply_journal_line_to_tb(fy, account_code, account_name, jnl_debit, jnl_credit, source='manual_journal', bulk_upload=None, description='', journal=None):
    """
    Apply a journal line to the trial balance by creating a separate
    adjustment row.  The display aggregation logic nets these rows
    against the original import row when rendering.

    The mapped_line_item is resolved using a three-tier lookup:
      1. Existing TrialBalanceLine for the same account_code in this FY
         (ensures the adjustment lands in the same section as the original)
      2. ClientAccountMapping (entity-level learned mapping)
      3. None (unmapped fallback)

    Pass journal=<AdjustingJournal instance> to link the TB line back to its
    source journal via the source_journal FK — this enables reliable reversal
    on edit without fragile debit/credit matching.
    """
    mapped_item = None

    # Priority 1: inherit from existing TB line for this account_code
    existing_tb = TrialBalanceLine.objects.filter(
        financial_year=fy,
        account_code=account_code,
        mapped_line_item__isnull=False,
    ).select_related('mapped_line_item').first()
    if existing_tb:
        mapped_item = existing_tb.mapped_line_item

    # Priority 2: fall back to ClientAccountMapping
    if not mapped_item:
        mapping = ClientAccountMapping.objects.filter(
            entity=fy.entity, client_account_code=account_code
        ).first()
        mapped_item = mapping.mapped_line_item if mapping else None

    # Resolve the account name at the source - if the caller passed a raw
    # code or blank name, look it up from EntityChartOfAccount / ChartOfAccount
    account_name = _resolve_account_name(fy.entity, account_code, account_name)

    tb_line = TrialBalanceLine.objects.create(
        financial_year=fy,
        account_code=account_code,
        account_name=account_name,
        opening_balance=Decimal('0'),
        debit=jnl_debit,
        credit=jnl_credit,
        closing_balance=jnl_debit - jnl_credit,
        mapped_line_item=mapped_item,
        is_adjustment=True,
        source=source,
        description=description,
        bulk_journal_upload=bulk_upload,
        source_journal=journal,
    )
    logger.info(
        "TB line posted: %s DR=%s CR=%s journal=%s fy=%s pk=%s source=%s",
        account_code, jnl_debit, jnl_credit,
        journal.pk if journal else None, fy.pk, tb_line.pk, source,
    )


def _post_journal_to_tb(journal, fy):
    """Post a journal's lines to the trial balance, aggregating by account code.

    Multiple journal lines targeting the same account code are summed into a
    single TB adjustment line (total debits and total credits per code).  This
    is valid accounting practice — e.g. splitting a single payment across the
    same expense account with different narrations — and must not be silently
    dropped or deduplicated.

    Wrapped in transaction.atomic() so that all TB lines are created or none
    are — a partial failure cannot leave orphaned TB lines.
    """
    from collections import OrderedDict

    agg = OrderedDict()
    for line in journal.lines.order_by("line_number", "id"):
        key = line.account_code
        if key not in agg:
            agg[key] = {
                "name": line.account_name,
                "dr": Decimal("0"),
                "cr": Decimal("0"),
            }
        agg[key]["dr"] += line.debit
        agg[key]["cr"] += line.credit

    if not agg:
        logger.warning(
            "No journal lines found for journal %s (%s) — nothing to post",
            journal.reference_number, journal.pk,
        )
        return

    logger.info(
        "Posting journal %s (%s) to TB: %d unique account codes from %d lines",
        journal.reference_number, journal.pk, len(agg),
        journal.lines.count(),
    )

    with db_transaction.atomic():
        for account_code, vals in agg.items():
            _apply_journal_line_to_tb(
                fy,
                account_code,
                vals["name"],
                vals["dr"],
                vals["cr"],
                source="manual_journal",
                description=journal.description,
                journal=journal,
            )


def _reverse_journal_tb_lines(journal):
    """Delete all TB adjustment lines created by a posted journal.

    Uses a multi-tier strategy to handle journals from different eras of the
    codebase:

    .. note:: Tier 2 and 3 MUST exclude lines linked to a BulkJournalUpload
       (bulk_journal_upload__isnull=True) — those belong to bulk uploads, not
       manual journals.  Without this guard, editing/deleting a pre-FK journal
       can accidentally delete TB lines from unrelated bulk uploads that share
       the same account codes.

      Tier 1 – FK-based:  TB lines linked via source_journal FK (reliable,
               works for journals posted after migration 0068).
      Tier 2 – Value match:  Match each JournalLine's account_code + exact
               debit/credit against unlinked (source_journal IS NULL)
               adjustment TB lines.  For pre-FK journals.
      Tier 3 – Broad match:  Match each JournalLine's account_code against
               unlinked adjustment TB lines with source='manual_journal',
               ignoring exact debit/credit.  Catches journals whose lines
               were edited after posting (values diverged).

    Returns the total number of TB lines deleted.
    """
    fy = journal.financial_year
    deleted_count = 0

    # ── Tier 1: FK-based deletion ─────────────────────────────────────
    fk_qs = TrialBalanceLine.objects.filter(
        financial_year=fy,
        is_adjustment=True,
        source_journal=journal,
    )
    tier1 = fk_qs.count()
    if tier1:
        fk_qs.delete()
        deleted_count += tier1
        logger.info(
            "Tier 1 (FK): deleted %d TB lines for journal %s (%s)",
            tier1, journal.reference_number, journal.pk,
        )
        return deleted_count  # FK path is authoritative — done

    # ── Tier 2: Exact value match (legacy, pre-FK journals) ───────────
    # IMPORTANT: exclude lines linked to a BulkJournalUpload — those
    # belong to bulk uploads, not to this manual journal.
    for jnl_line in journal.lines.all():
        adj = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=jnl_line.account_code,
            debit=jnl_line.debit,
            credit=jnl_line.credit,
            is_adjustment=True,
            source="manual_journal",
            source_journal__isnull=True,
            bulk_journal_upload__isnull=True,
        ).first()
        if adj:
            adj.delete()
            deleted_count += 1

    if deleted_count:
        logger.info(
            "Tier 2 (exact match): deleted %d TB lines for journal %s (%s)",
            deleted_count, journal.reference_number, journal.pk,
        )
        return deleted_count

    # ── Tier 3: Broad account-code match ──────────────────────────────
    # The journal's JournalLine values may have diverged from the original
    # TB lines (e.g. after a partially-failed edit, or a code path that
    # modified TB lines without updating JournalLine records).  Match by
    # account_code + source type only, deleting one TB line per journal
    # line.
    # IMPORTANT: only match source="manual_journal" and exclude lines
    # linked to bulk uploads — those belong to BulkJournalUpload, not
    # to this manual journal.
    journal_codes = list(
        journal.lines.order_by("account_code")
        .values_list("account_code", flat=True)
    )
    for code in journal_codes:
        adj = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=code,
            is_adjustment=True,
            source="manual_journal",
            source_journal__isnull=True,
            bulk_journal_upload__isnull=True,
        ).first()
        if adj:
            adj.delete()
            deleted_count += 1

    if deleted_count:
        logger.info(
            "Tier 3 (broad match): deleted %d TB lines for journal %s (%s)",
            deleted_count, journal.reference_number, journal.pk,
        )
    else:
        logger.warning(
            "No TB lines found to reverse for journal %s (%s) — all 3 tiers empty",
            journal.reference_number, journal.pk,
        )

    return deleted_count


def _delete_orphaned_tb_lines_for_journal(journal):
    """Delete orphaned TB lines that match a journal's aggregated line values.

    Targets adjustment lines with source_journal=NULL that match the journal's
    account codes and aggregated debit/credit amounts.  These are remnants of
    pre-FK posting or failed edit operations that ``_reverse_journal_tb_lines``
    may not have caught (e.g. when Tier 1 FK-based deletion succeeded but
    orphaned duplicates remain).
    """
    from collections import OrderedDict

    fy = journal.financial_year
    jnl_lines = list(journal.lines.all())
    if not jnl_lines:
        return 0

    # Aggregate by account code to match the posting logic
    agg = OrderedDict()
    for jl in jnl_lines:
        key = jl.account_code
        if key not in agg:
            agg[key] = {"dr": Decimal("0"), "cr": Decimal("0")}
        agg[key]["dr"] += jl.debit
        agg[key]["cr"] += jl.credit

    deleted = 0
    for code, vals in agg.items():
        result = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=code,
            debit=vals["dr"],
            credit=vals["cr"],
            is_adjustment=True,
            source="manual_journal",
            source_journal__isnull=True,
            bulk_journal_upload__isnull=True,
        ).delete()
        deleted += result[0]

    if deleted:
        logger.info(
            "Orphan cleanup: deleted %d orphaned TB line(s) for journal %s (%s)",
            deleted, journal.reference_number, journal.pk,
        )
    return deleted


def _verify_tb_balance(fy):
    """Check that total debits equal total credits for a financial year.

    Logs a warning if the TB is out of balance after a journal operation.
    This is a diagnostic safeguard — it does not attempt auto-repair.
    """
    totals = TrialBalanceLine.objects.filter(
        financial_year=fy,
    ).aggregate(total_dr=Sum("debit"), total_cr=Sum("credit"))

    total_dr = totals["total_dr"] or Decimal("0")
    total_cr = totals["total_cr"] or Decimal("0")

    if total_dr != total_cr:
        logger.warning(
            "TB OUT OF BALANCE after journal operation — FY %s (%s): "
            "total DR=%s, total CR=%s, difference=%s",
            fy.year_label, fy.pk, total_dr, total_cr, total_dr - total_cr,
        )


def _reverse_journal_line_from_tb(fy, account_code, jnl_debit, jnl_credit, journal=None):
    """
    Reverse a previously applied journal line by deleting its adjustment row.

    When journal is provided, the deletion is scoped to TB lines linked to
    that specific journal via source_journal FK — this is reliable even when
    multiple lines share the same account code and amounts.

    Falls back to the legacy debit/credit match when journal is None (e.g.
    for journals created before the source_journal FK was added).
    """
    if journal is not None:
        # Preferred path: delete the row linked to this exact journal
        TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=account_code,
            is_adjustment=True,
            source_journal=journal,
        ).delete()
    else:
        # Legacy fallback: match by debit/credit (may be ambiguous).
        # Exclude bulk upload lines — they belong to BulkJournalUpload.
        adj = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=account_code,
            debit=jnl_debit,
            credit=jnl_credit,
            is_adjustment=True,
            source="manual_journal",
            bulk_journal_upload__isnull=True,
        ).first()
        if adj:
            adj.delete()


def _get_or_create_tb_line(financial_year=None, account_code=None, defaults=None, fy=None):
    """
    Safely get or create a TrialBalanceLine for bank-statement pushes.

    Unlike Django's get_or_create, this handles the case where multiple
    rows already exist for the same (financial_year, account_code) — which
    is normal because journal adjustments create separate rows.  We pick
    the *first non-adjustment* row, or the first row overall, to accumulate
    bank-statement amounts into.

    When creating a new line, automatically applies any existing
    ClientAccountMapping so the line is pre-mapped.
    """
    fy_resolved = financial_year or fy
    qs = TrialBalanceLine.objects.filter(
        financial_year=fy_resolved, account_code=account_code,
    )
    # Prefer the non-adjustment (original / bank_statement) row
    tb_line = qs.filter(is_adjustment=False).first() or qs.first()
    if tb_line:
        return tb_line, False
    # No row exists — create one.
    # Apply existing ClientAccountMapping if available.
    defaults = defaults or {}
    if 'mapped_line_item' not in defaults or defaults.get('mapped_line_item') is None:
        cam = ClientAccountMapping.objects.filter(
            entity=fy_resolved.entity,
            client_account_code=account_code,
            mapped_line_item__isnull=False,
        ).select_related('mapped_line_item').first()
        if cam:
            defaults['mapped_line_item'] = cam.mapped_line_item
    tb_line = TrialBalanceLine.objects.create(
        financial_year=fy_resolved,
        account_code=account_code,
        **(defaults),
    )
    return tb_line, True


def _get_bank_mapping_for_txn(txn):
    """
    Resolve the BankAccountMapping for a given PendingTransaction.
    Lookup order:
      1. BankAccountMapping by exact (entity, bsb, account_number) match
      2. BankAccountMapping marked as default for entity
      3. BankAccountMapping with empty bsb/account_number (catch-all for CSV/Excel uploads)
      4. If only one BankAccountMapping exists for the entity, use it
      5. BankAccount model by (entity, bsb, account_number) if it has tb_account_code
      6. Any BankAccountMapping for the entity (last resort — ensures contra
         entries are never silently dropped when a mapping exists)
    Returns a duck-typed object with .tb_account_code and .tb_account_name, or None.
    """
    if not txn.job or not txn.job.entity:
        return None
    entity = txn.job.entity
    job = txn.job
    mapping = None

    # 1. Exact match by BSB + account number (when available from PDF parsing)
    if job.bsb or job.account_number:
        mapping = BankAccountMapping.objects.filter(
            entity=entity, bsb=job.bsb or '', account_number=job.account_number or '',
        ).first()

    # 2. Default mapping for the entity
    if not mapping:
        mapping = BankAccountMapping.objects.filter(
            entity=entity, is_default=True,
        ).first()

    # 3. Catch-all: mapping with empty bsb/account_number (created from CSV/Excel uploads)
    if not mapping:
        mapping = BankAccountMapping.objects.filter(
            entity=entity, bsb='', account_number='',
        ).first()

    # 4. If only one mapping exists for the entity, use it regardless
    if not mapping:
        entity_mappings = BankAccountMapping.objects.filter(entity=entity)
        if entity_mappings.count() == 1:
            mapping = entity_mappings.first()

    # 5. Fallback: check the BankAccount model (has tb_account_code field)
    if not mapping and (job.bsb or job.account_number):
        from core.models import BankAccount
        ba = BankAccount.objects.filter(
            entity=entity, bsb=job.bsb or '', account_number=job.account_number or '',
        ).first()
        if ba and ba.tb_account_code:
            mapping = ba  # BankAccount has .tb_account_code and .tb_account_name

    # 6. Last resort: pick any mapping for the entity so the contra entry is
    #    never silently dropped.  Prefers the most recently updated mapping.
    if not mapping:
        mapping = BankAccountMapping.objects.filter(
            entity=entity,
        ).order_by('-updated_at').first()

    if not mapping:
        import logging
        logger = logging.getLogger('core.views')
        logger.warning(
            f"No bank mapping found for entity {entity.pk} "
            f"(job BSB={job.bsb!r}, acc={job.account_number!r}). "
            f"Bank contra entry will be skipped for txn {txn.pk}."
        )
    return mapping


def _post_bank_contra_entry(txn, fy, bank_mapping, has_gst):
    """
    Post the bank-side (contra) entry for a transaction.

    For a payment (amount < 0): the bank account is CREDITED (money leaving).
    For a receipt (amount > 0): the bank account is DEBITED (money arriving).

    The gross amount (inclusive of GST) hits the bank account, because the
    bank statement shows the full amount that moved.
    """
    if not bank_mapping:
        import logging
        logger = logging.getLogger('core.views')
        logger.error(
            "_post_bank_contra_entry: No bank_mapping for txn %s "
            "(amount=%s). Contra entry cannot be posted — TB will be out of balance.",
            txn.pk, txn.amount,
        )
        raise ValueError(
            f"No bank account mapping found for transaction {txn.pk}. "
            f"Please configure bank account mappings before approving transactions."
        )
    gross_amount = abs(txn.amount)
    bank_code = bank_mapping.tb_account_code
    bank_name = bank_mapping.tb_account_name

    tb_line, created = _get_or_create_tb_line(
        financial_year=fy,
        account_code=bank_code,
        defaults={
            "account_name": bank_name,
            "debit": gross_amount if txn.amount > 0 else Decimal("0"),
            "credit": gross_amount if txn.amount < 0 else Decimal("0"),
            "closing_balance": gross_amount if txn.amount > 0 else -gross_amount,
            "tax_type": "",
            "source": "bank_statement",
        },
    )
    if not created:
        if txn.amount > 0:
            tb_line.debit += gross_amount
            tb_line.closing_balance += gross_amount
        else:
            tb_line.credit += gross_amount
            tb_line.closing_balance -= gross_amount
        if not tb_line.source:
            tb_line.source = 'bank_statement'
        tb_line.save()


def _reverse_bank_contra_entry(txn, fy):
    """
    Reverse the bank-side (contra) entry for a transaction by decrementing
    the accumulated values on the original TB line.

    This mirrors _post_bank_contra_entry() which accumulates into the
    original line — the reversal must undo that accumulation rather than
    creating a separate adjustment line, to prevent closing_balance drift
    after unconfirm→re-approve cycles.
    """
    if not txn.job or not txn.job.entity:
        return
    bank_mapping = _get_bank_mapping_for_txn(txn)
    if not bank_mapping:
        return
    gross_amount = abs(txn.amount)
    bank_code = bank_mapping.tb_account_code

    tb_line = TrialBalanceLine.objects.filter(
        financial_year=fy,
        account_code=bank_code,
        is_adjustment=False,
    ).first()
    if not tb_line:
        tb_line = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=bank_code,
        ).first()
    if not tb_line:
        return

    if txn.amount > 0:
        # Original posting debited bank — reverse by decrementing debit
        tb_line.debit = max(Decimal("0"), tb_line.debit - gross_amount)
        tb_line.closing_balance -= gross_amount
    else:
        # Original posting credited bank — reverse by decrementing credit
        tb_line.credit = max(Decimal("0"), tb_line.credit - gross_amount)
        tb_line.closing_balance += gross_amount
    tb_line.save(update_fields=["debit", "credit", "closing_balance"])


def _post_txn_to_tb(txn, fy, has_gst):
    """
    Centralised helper: post a single confirmed bank-statement transaction
    to the trial balance.  This ALWAYS performs the full triple-entry:

      1. Net amount  → expense / income account (Dr or Cr)
      2. GST amount  → 3380 GST payable control account (if applicable)
      3. Gross amount → bank account contra-entry (ALWAYS)

    Every approval code-path MUST call this function instead of inlining
    the posting logic, so the bank contra entry can never be skipped.

    The entire operation is wrapped in a database transaction to prevent
    partial postings (e.g. expense posted but contra skipped).

    Returns True if the TB was updated, False otherwise.
    """
    from django.db import transaction as db_transaction

    # Guard: skip if already posted to prevent double-posting
    if getattr(txn, 'posted_to_tb', False):
        return False

    amount = txn.amount
    code = txn.confirmed_code
    name = txn.confirmed_name
    tax_type = txn.confirmed_tax_type or ''

    if not code or amount == 0:
        return False

    with db_transaction.atomic():
        # --- 1. Post net amount to expense/income account ---
        net_for_tb = txn.net_amount if has_gst else abs(amount)
        tb_line, created = _get_or_create_tb_line(
            financial_year=fy,
            account_code=code,
            defaults={
                "account_name": name,
                "debit": net_for_tb if amount < 0 else Decimal("0"),
                "credit": net_for_tb if amount > 0 else Decimal("0"),
                "closing_balance": net_for_tb if amount < 0 else -net_for_tb,
                "tax_type": tax_type,
                "source": "bank_statement",
            },
        )
        if not created:
            if amount < 0:
                tb_line.debit += net_for_tb
                tb_line.closing_balance += net_for_tb
            else:
                tb_line.credit += net_for_tb
                tb_line.closing_balance -= net_for_tb
            if not tb_line.tax_type:
                tb_line.tax_type = tax_type
            if not tb_line.source:
                tb_line.source = 'bank_statement'
            tb_line.save()

        # --- 2. Post GST to 3380 GST payable control account ---
        gst_amt = txn.confirmed_gst_amount if txn.confirmed_gst_amount else Decimal('0')
        if has_gst and gst_amt > 0:
            gst_code = '3380'
            gst_name = 'GST payable control account'
            if amount > 0:
                # Income: GST Collected — credit 3380 (increases liability)
                gst_line, gst_created = _get_or_create_tb_line(
                    financial_year=fy,
                    account_code=gst_code,
                    defaults={
                        "account_name": gst_name,
                        "debit": Decimal("0"),
                        "credit": gst_amt,
                        "closing_balance": -gst_amt,
                        "tax_type": "GST on Income",
                        "source": "bank_statement",
                    },
                )
                if not gst_created:
                    gst_line.credit += gst_amt
                    gst_line.closing_balance -= gst_amt
                    gst_line.save()
            else:
                # Expense: GST Paid — debit 3380 (reduces liability / creates asset)
                gst_line, gst_created = _get_or_create_tb_line(
                    financial_year=fy,
                    account_code=gst_code,
                    defaults={
                        "account_name": gst_name,
                        "debit": gst_amt,
                        "credit": Decimal("0"),
                        "closing_balance": gst_amt,
                        "tax_type": "GST on Expenses",
                        "source": "bank_statement",
                    },
                )
                if not gst_created:
                    gst_line.debit += gst_amt
                    gst_line.closing_balance += gst_amt
                    gst_line.save()

        # --- 3. ALWAYS post the bank contra-entry (gross amount) ---
        bank_mapping = _get_bank_mapping_for_txn(txn)
        _post_bank_contra_entry(txn, fy, bank_mapping, has_gst)

        # Mark as posted to prevent double-posting
        txn.posted_to_tb = True
        txn.save(update_fields=['posted_to_tb'])

    return True


def _recalculate_bank_tb_lines(fy):
    """Recompute the debit/credit/closing_balance on bank-accumulated TB lines
    by aggregating from all currently-posted PendingTransactions.

    This is a safety-net function that ensures the TB reflects the true state
    of all posted bank transactions.  It rebuilds each non-adjustment TB line
    that has source='bank_statement' from scratch.

    Journal-created adjustment lines (source='manual_journal') are never
    touched — they are managed by _apply_journal_line_to_tb / _reverse_journal_line_from_tb.
    """
    from review.models import PendingTransaction
    from django.db.models import Sum, Q, F

    posted_txns = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=True,
        posted_to_tb=True,
    ).select_related('job')

    # Build a mapping of account_code → (total_debit, total_credit) from posted txns
    account_totals = {}  # code -> {'debit': D, 'credit': D}
    gst_debit = Decimal("0")
    gst_credit = Decimal("0")
    bank_totals = {}  # bank_code -> {'debit': D, 'credit': D}

    for txn in posted_txns:
        has_gst = txn.confirmed_gst_amount and txn.confirmed_gst_amount > 0
        net_for_tb = txn.net_amount if has_gst else abs(txn.amount)

        code = txn.confirmed_code
        if not code:
            continue

        if code not in account_totals:
            account_totals[code] = {'debit': Decimal("0"), 'credit': Decimal("0")}

        if txn.amount < 0:
            account_totals[code]['debit'] += net_for_tb
        else:
            account_totals[code]['credit'] += net_for_tb

        # GST
        if has_gst:
            gst_amt = txn.confirmed_gst_amount
            if txn.amount > 0:
                gst_credit += gst_amt
            else:
                gst_debit += gst_amt

        # Bank contra
        bank_mapping = _get_bank_mapping_for_txn(txn)
        if bank_mapping:
            bcode = bank_mapping.tb_account_code
            if bcode not in bank_totals:
                bank_totals[bcode] = {'debit': Decimal("0"), 'credit': Decimal("0")}
            gross = abs(txn.amount)
            if txn.amount > 0:
                bank_totals[bcode]['debit'] += gross
            else:
                bank_totals[bcode]['credit'] += gross

    # Now update the non-adjustment TB lines for each account code
    # Only touch lines with source='bank_statement' to avoid interfering with imports
    for code, totals in account_totals.items():
        tb_line = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=code,
            is_adjustment=False,
            source='bank_statement',
        ).first()
        if tb_line:
            ob = tb_line.opening_balance or Decimal("0")
            tb_line.debit = totals['debit']
            tb_line.credit = totals['credit']
            tb_line.closing_balance = ob + totals['debit'] - totals['credit']
            tb_line.save(update_fields=["debit", "credit", "closing_balance"])

    # GST control account
    if gst_debit or gst_credit:
        gst_line = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code="3380",
            is_adjustment=False,
            source='bank_statement',
        ).first()
        if gst_line:
            ob = gst_line.opening_balance or Decimal("0")
            gst_line.debit = gst_debit
            gst_line.credit = gst_credit
            gst_line.closing_balance = ob + gst_debit - gst_credit
            gst_line.save(update_fields=["debit", "credit", "closing_balance"])

    # Bank contra accounts
    for bcode, totals in bank_totals.items():
        bank_line = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=bcode,
            is_adjustment=False,
            source='bank_statement',
        ).first()
        if bank_line:
            ob = bank_line.opening_balance or Decimal("0")
            bank_line.debit = totals['debit']
            bank_line.credit = totals['credit']
            bank_line.closing_balance = ob + totals['debit'] - totals['credit']
            bank_line.save(update_fields=["debit", "credit", "closing_balance"])

    # Clean up orphaned reversal adjustment lines from the old pattern
    TrialBalanceLine.objects.filter(
        financial_year=fy,
        is_adjustment=True,
        source='bank_statement',
        description__startswith='Reversal of ',
    ).delete()


def _log_action(request, action, description, obj=None):
    """Create an audit log entry."""
    AuditLog.objects.create(
        user=request.user,
        action=action,
        description=description,
        affected_object_type=type(obj).__name__ if obj else "",
        affected_object_id=str(obj.pk) if obj else "",
        ip_address=request.META.get("REMOTE_ADDR"),
    )


def _aggregate_tb_lines(ordered_sections, entity=None):
    """
    Aggregate multiple TrialBalanceLine records per account_code within each
    section.  When journal entries create adjustment rows (is_adjustment=True),
    the raw queryset contains multiple rows for the same account.  This helper
    nets them into a single row per account_code, matching the behaviour of
    the on-screen Trial Balance tab.

    When *entity* is provided, account names that look like bare codes
    (e.g. "1706" stored as the name for account 1706) are resolved via
    _resolve_account_name() which checks EntityChartOfAccount and the
    master ChartOfAccount template.

    Returns a new OrderedDict of {section_name: [aggregated lines]}.
    Each aggregated line has: account_code, account_name, display_dr,
    display_cr, prior_debit, prior_credit.
    """
    from collections import OrderedDict

    aggregated = OrderedDict()
    for section_name, lines_list in ordered_sections.items():
        code_groups = OrderedDict()
        for line in lines_list:
            code_groups.setdefault(line.account_code, []).append(line)

        agg_lines = []
        for code, group in code_groups.items():
            raw_dr = Decimal('0')
            raw_cr = Decimal('0')
            agg_prior_dr = Decimal('0')
            agg_prior_cr = Decimal('0')
            first = group[0]
            for l in group:
                cy = getattr(l, '_cy', None)
                if cy is None:
                    cy = l.closing_balance or Decimal('0')
                if cy > 0:
                    raw_dr += cy
                elif cy < 0:
                    raw_cr += abs(cy)
                py = getattr(l, '_py', None)
                if py is None:
                    py = (l.prior_debit or Decimal('0')) - (l.prior_credit or Decimal('0'))
                if py > 0:
                    agg_prior_dr += py
                elif py < 0:
                    agg_prior_cr += abs(py)

            if len(group) == 1:
                # Resolve name for single-line accounts too
                if entity:
                    first.account_name = _resolve_account_name(entity, first.account_code, first.account_name)
                first._agg_dr = raw_dr
                first._agg_cr = raw_cr
                first._agg_prior_dr = agg_prior_dr
                first._agg_prior_cr = agg_prior_cr
                agg_lines.append(first)
            else:
                net = raw_dr - raw_cr
                if net >= 0:
                    agg_dr = net
                    agg_cr = Decimal('0')
                else:
                    agg_dr = Decimal('0')
                    agg_cr = abs(net)

                class _AggLine:
                    pass
                agg = _AggLine()
                agg.account_code = code
                unique_names = list(dict.fromkeys(
                    l.account_name for l in group if l.account_name
                ))
                resolved_name = unique_names[0] if unique_names else code
                if entity:
                    resolved_name = _resolve_account_name(entity, code, resolved_name)
                agg.account_name = resolved_name
                agg._agg_dr = agg_dr
                agg._agg_cr = agg_cr
                agg._agg_prior_dr = agg_prior_dr
                agg._agg_prior_cr = agg_prior_cr
                agg.mapped_line_item = first.mapped_line_item
                agg_lines.append(agg)

        aggregated[section_name] = agg_lines
    return aggregated


# ---------------------------------------------------------------------------
# Dashboard
# ---------------------------------------------------------------------------
@login_required
def dashboard(request):
    user = request.user
    if user.can_view_all_entities:
        clients = Client.objects.filter(is_active=True)
    else:
        clients = Client.objects.filter(
            Q(assigned_accountant=user) & Q(is_active=True)
        )

    # Time-based greeting
    import datetime
    now = timezone.localtime(timezone.now())
    hour = now.hour
    if hour < 12:
        greeting = "Good morning"
    elif hour < 17:
        greeting = "Good afternoon"
    else:
        greeting = "Good evening"

    # Unfinalised financial years (Draft, In Review, Finished) — grouped by client
    unfinalised_years = (
        FinancialYear.objects.filter(
            entity__client__in=clients,
            status__in=["draft", "in_review", "reopened"],
        )
        .select_related("entity", "entity__client")
        .order_by("entity__client__name", "entity__entity_name", "-end_date")
    )

    # Open audit risk flags across all clients I'm working on
    open_risk_flags = (
        RiskFlag.objects.filter(
            financial_year__entity__client__in=clients,
            status__in=["open", "reviewed"],
        )
        .select_related("financial_year", "financial_year__entity", "financial_year__entity__client")
        .order_by("-severity", "-created_at")[:20]
    )

    # Group risk flags by client/entity for display
    risk_summary = {}
    for flag in open_risk_flags:
        key = flag.financial_year.entity.entity_name
        if key not in risk_summary:
            risk_summary[key] = {
                "entity_name": flag.financial_year.entity.entity_name,
                "fy_pk": flag.financial_year.pk,
                "flags": [],
            }
        risk_summary[key]["flags"].append(flag)

    # Recent activity log
    recent_activities = (
        ActivityLog.objects.all()
        .order_by("-created_at")[:30]
    )

    # Unread notification count for the bell
    unread_count = ActivityLog.objects.filter(is_read=False).count()

    context = {
        "greeting": greeting,
        "now": now,
        "unfinalised_years": unfinalised_years,
        "risk_summary": risk_summary,
        "open_risk_count": open_risk_flags.count(),
        "recent_activities": recent_activities,
        "unread_count": unread_count,
    }
    return render(request, "core/dashboard.html", context)


# ---------------------------------------------------------------------------
# Entities (top-level — replaces Clients)
# ---------------------------------------------------------------------------
@login_required
def entity_list(request):
    """Main list page showing all entities (replaces client_list)."""
    query = request.GET.get("q", "")
    entity_type = request.GET.get("entity_type", "")
    status_filter = request.GET.get("status", "")
    show_archived = request.GET.get("show_archived", "") == "1"

    if show_archived:
        if request.user.can_view_all_entities:
            entities = Entity.objects.filter(is_archived=True)
        else:
            entities = Entity.objects.filter(
                assigned_accountant=request.user, is_archived=True
            )
    else:
        if request.user.can_view_all_entities:
            entities = Entity.objects.filter(is_archived=False)
        else:
            entities = Entity.objects.filter(
                assigned_accountant=request.user, is_archived=False
            )

    if query:
        entities = entities.filter(
            Q(entity_name__icontains=query)
            | Q(abn__icontains=query)
            | Q(trading_as__icontains=query)
        ).distinct()

    if entity_type:
        entities = entities.filter(entity_type=entity_type)

    entities = entities.annotate(num_fys=Count("financial_years"))

    context = {
        "entities": entities,
        "query": query,
        "entity_type": entity_type,
        "status_filter": status_filter,
        "show_archived": show_archived,
        "entity_types": Entity.EntityType.choices,
    }

    if request.htmx:
        return render(request, "partials/entity_list_rows.html", context)
    return render(request, "core/entity_list.html", context)


# Keep client_list as alias for backward compatibility
client_list = entity_list


@login_required
def client_create(request):
    """Redirect to entity_create — clients are no longer created directly."""
    return redirect("core:entity_create")


@login_required
def client_detail(request, pk):
    """Legacy redirect — try to find entity or client."""
    # Try as entity first
    entity = Entity.objects.filter(pk=pk).first()
    if entity:
        return redirect("core:entity_detail", pk=pk)
    # Try as client — redirect to first entity
    client = Client.objects.filter(pk=pk).first()
    if client:
        first_entity = client.entities.first()
        if first_entity:
            return redirect("core:entity_detail", pk=first_entity.pk)
    return redirect("core:entity_list")


@login_required
def client_edit(request, pk):
    """Legacy redirect."""
    return redirect("core:entity_list")


# ---------------------------------------------------------------------------
# Entities
# ---------------------------------------------------------------------------
@login_required
def entity_create(request, client_pk=None):
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to create entities.")
        return redirect("core:entity_list")

    if request.method == "POST":
        form = EntityForm(request.POST, user=request.user)
        if form.is_valid():
            entity = form.save()
            _log_action(request, "import", f"Created entity: {entity.entity_name}", entity)
            messages.success(request, f"Entity '{entity.entity_name}' created.")
            return redirect("core:entity_detail", pk=entity.pk)
    else:
        form = EntityForm(user=request.user)
    return render(request, "core/entity_form.html", {
        "form": form, "title": "Create Entity"
    })


@login_required
def entity_detail(request, pk):
    entity = get_entity_for_user(request, pk)
    # Audit log: data access
    _log_action(request, "view", f"Viewed entity: {entity.entity_name}", entity)
    financial_years = entity.financial_years.all()
    officers = entity.officers.all().order_by('date_ceased', 'full_name')
    unfinalised_count = financial_years.exclude(status="finalised").count()
    associates = entity.associates.filter(is_active=True)
    family_associates = [a for a in associates if a.is_family]
    business_associates = [a for a in associates if not a.is_family]
    # Entity-to-entity relationships (both directions)
    from core.models import EntityRelationship
    outgoing_rels = EntityRelationship.objects.filter(from_entity=entity).select_related("to_entity")
    incoming_rels = EntityRelationship.objects.filter(to_entity=entity).select_related("from_entity")
    entity_relationships = list(outgoing_rels) + list(incoming_rels)
    software_configs = entity.software_configs.all()
    from integrations.models import XeroTenant, QBTenant
    linked_xero_tenant = XeroTenant.objects.filter(entity=entity).select_related("connection").first()
    linked_qb_tenant = QBTenant.objects.filter(entity=entity).select_related("connection").first()
    meeting_notes = entity.meeting_notes.all()[:20]
    pending_followups = entity.meeting_notes.filter(
        follow_up_completed=False, follow_up_date__isnull=False
    ).order_by("follow_up_date")
    has_financial_years = financial_years.exists()

    # ── Legal document prompt (Master Spec 4.6.3) ──────────────────────
    # Surface a one-time prompt after entity creation for companies/trusts
    # so the user can consciously initiate the establishment package.
    legal_doc_prompt = None
    if not entity.legal_doc_prompt_dismissed and entity.entity_type in ("company", "trust"):
        from core.models import LegalDocument
        doc_type_map = {
            "company": ("company_establishment", "Company Establishment Package"),
            "trust": ("discretionary_trust_deed", "Discretionary Trust Deed"),
        }
        doc_type_key, doc_type_label = doc_type_map[entity.entity_type]
        # Only show if no document of this type has been generated yet
        already_generated = LegalDocument.objects.filter(
            entity=entity, document_type=doc_type_key,
        ).exists()
        if not already_generated:
            # Build the wizard URL — requires a financial year; fall back to None
            latest_fy = financial_years.order_by("-end_date").first()
            legal_doc_prompt = {
                "doc_type": doc_type_key,
                "doc_type_label": doc_type_label,
                "has_fy": latest_fy is not None,
                "wizard_url": (
                    reverse("core:legal_doc_wizard", kwargs={"pk": latest_fy.pk, "doc_type": doc_type_key})
                    if latest_fy else None
                ),
            }

    context = {
        "entity": entity,
        "financial_years": financial_years,
        "officers": officers,
        "unfinalised_count": unfinalised_count,
        "has_financial_years": has_financial_years,
        "family_associates": family_associates,
        "business_associates": business_associates,
        "entity_relationships": entity_relationships,
        "outgoing_rels": outgoing_rels,
        "incoming_rels": incoming_rels,
        "software_configs": software_configs,
        "linked_xero_tenant": linked_xero_tenant,
        "linked_qb_tenant": linked_qb_tenant,
        "meeting_notes": meeting_notes,
        "pending_followups": pending_followups,
        "legal_doc_prompt": legal_doc_prompt,
        # Governing Documents tab
        "primary_governing_doc": entity.governing_documents.filter(
            is_primary=True, status="active"
        ).first(),
        "amendment_docs": entity.governing_documents.filter(
            is_primary=False, status="active"
        ).order_by("-document_date", "-uploaded_at"),
        # Engagement Letters tab
        "engagement_letters": entity.engagement_letters.select_related(
            "financial_year", "uploaded_by"
        ).order_by("-financial_year__end_date", "-uploaded_at"),
        "generated_engagement_letters": entity.legal_documents.select_related(
            "financial_year", "generated_by"
        ).filter(
            document_type="engagement_letter"
        ).order_by("-generated_at"),
        "engagement_letters_count": entity.engagement_letters.filter(is_current=True).count(),
    }
    return render(request, "core/entity_detail.html", context)


@login_required
def entity_edit(request, pk):
    entity = get_entity_for_user(request, pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to edit entities.")
        return redirect("core:entity_detail", pk=pk)

    if request.method == "POST":
        form = EntityForm(request.POST, instance=entity, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, f"Entity '{entity.entity_name}' updated.")
            return redirect("core:entity_detail", pk=pk)
    else:
        form = EntityForm(instance=entity, user=request.user)
    return render(request, "core/entity_form.html", {
        "form": form, "title": f"Edit: {entity.entity_name}"
    })


@login_required
@require_POST
def dismiss_legal_doc_prompt(request, pk):
    """Dismiss the post-creation legal document prompt (Master Spec 4.6.3)."""
    entity = get_entity_for_user(request, pk)
    entity.legal_doc_prompt_dismissed = True
    entity.save(update_fields=["legal_doc_prompt_dismissed"])
    return JsonResponse({"ok": True})


# ---------------------------------------------------------------------------
# Financial Years
# ---------------------------------------------------------------------------
@login_required
def financial_year_create(request, entity_pk):
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=entity_pk)

    if request.method == "POST":
        form = FinancialYearForm(request.POST)
        if form.is_valid():
            fy = form.save(commit=False)
            fy.entity = entity
            # Link to prior year if exists
            prior = entity.financial_years.order_by("-end_date").first()
            if prior:
                fy.prior_year = prior
            fy.save()
            # Seed entity chart of accounts from template if not already done
            seeded = EntityChartOfAccount.seed_from_template(entity)
            if seeded:
                _log_action(request, "import", f"Seeded {seeded} chart of accounts from template for {entity.entity_name}", fy)
            _log_action(request, "import", f"Created financial year: {fy.year_label}", fy)
            messages.success(request, f"Financial year '{fy.year_label}' created.")
            return redirect("core:financial_year_detail", pk=fy.pk)
    else:
        form = FinancialYearForm()
    return render(request, "core/financial_year_form.html", {
        "form": form, "entity": entity, "title": "Create Financial Year"
    })


@login_required
def financial_year_detail(request, pk):
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    # Audit log: data access
    _log_action(request, "view", f"Viewed financial year: {fy.year_label} for {entity.entity_name}", fy)
    tb_lines = fy.trial_balance_lines.select_related("mapped_line_item").all()
    adjustments = fy.adjusting_journals.all().order_by('-posted_at', '-created_at')
    # Count unmapped accounts, excluding system/clearing accounts that don't
    # need mapping (GST clearing, bank contra accounts, etc.)
    _system_codes = {'3380', '9100', '9110'}  # GST payable / clearing accounts
    unmapped_codes = (
        tb_lines.filter(mapped_line_item__isnull=True)
        .exclude(account_code__in=_system_codes)
        .values_list('account_code', flat=True)
        .distinct()
    )
    # Also exclude codes that already have a ClientAccountMapping with a
    # mapped_line_item (the TB line just hasn't been updated yet)
    _already_mapped_codes = set(
        ClientAccountMapping.objects.filter(
            entity=fy.entity,
            client_account_code__in=unmapped_codes,
            mapped_line_item__isnull=False,
        ).values_list('client_account_code', flat=True)
    )
    # Auto-fix: apply the mapping to any TB lines that are missing it
    if _already_mapped_codes:
        for _cam in ClientAccountMapping.objects.filter(
            entity=fy.entity,
            client_account_code__in=_already_mapped_codes,
            mapped_line_item__isnull=False,
        ).select_related('mapped_line_item'):
            tb_lines.filter(
                account_code=_cam.client_account_code,
                mapped_line_item__isnull=True,
            ).update(mapped_line_item=_cam.mapped_line_item)
    unmapped_count = len(set(unmapped_codes) - _already_mapped_codes)
    # Pre-load CoA section lookup for inferring sections of unmapped lines
    _coa_lookup = _build_coa_section_lookup(fy.entity)
    documents = fy.generated_documents.all().order_by('-version', '-generated_at')

    # Group lines by section for comparative display
    from collections import OrderedDict
    SECTION_ORDER = [
        'Revenue', 'Income', 'Cost of Sales', 'Expenses',
        'Current Assets', 'Non-Current Assets',
        'Current Liabilities', 'Non-Current Liabilities',
        'Equity', 'Income Tax',
    ]
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }
    tb_lines = list(tb_lines)
    for line in tb_lines:
        if line.source == 'rollover':
            line._cy = Decimal('0')
            line._py = (line.prior_debit or Decimal('0')) - (line.prior_credit or Decimal('0'))
        else:
            line._cy = line.closing_balance or Decimal('0')
            line._py = Decimal('0')

    for line in tb_lines:
        cy = line._cy
        if cy > 0:
            line.display_dr = cy
            line.display_cr = Decimal('0')
        elif cy < 0:
            line.display_dr = Decimal('0')
            line.display_cr = abs(cy)
        else:
            line.display_dr = line.debit if line.debit else Decimal('0')
            line.display_cr = line.credit if line.credit else Decimal('0')

    sections = OrderedDict()
    for line in tb_lines:
        # HARD RULE: the HandiLedger numeric code range is authoritative.
        # This prevents accounts like 2475 (Retention Receivable, range 2000-2499)
        # from being misclassified as Non-Current Assets due to a stale mapping.
        hl_section = _hl_section_for_code(line.account_code)
        if hl_section:
            display_section = hl_section
        elif line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = _coa_lookup.get(line.account_code, 'Unmapped')
        if display_section not in sections:
            sections[display_section] = []
        sections[display_section].append(line)
    ordered_sections = OrderedDict()
    section_keys_ordered = []
    for s in SECTION_ORDER:
        ds = SECTION_DISPLAY.get(s, s)
        if ds not in section_keys_ordered:
            section_keys_ordered.append(ds)
    for key in section_keys_ordered:
        if key in sections:
            ordered_sections[key] = sections[key]
    for key in sections:
        if key not in ordered_sections:
            ordered_sections[key] = sections[key]

    # ---- Aggregate multiple sub-entries per account_code ----
    # Within each section, group lines by account_code.  If a code has
    # multiple lines (e.g. from Excel sub-accounts, bank statement coding,
    # and journal entries), create an aggregated summary row with the
    # individual lines attached as `sub_entries` for drill-down display.
    aggregated_sections = OrderedDict()
    for section_name, lines_list in ordered_sections.items():
        code_groups = OrderedDict()  # account_code -> [lines]
        for line in lines_list:
            code_groups.setdefault(line.account_code, []).append(line)

        agg_lines = []
        for code, group in code_groups.items():
            if len(group) == 1:
                # Single entry — no aggregation needed
                line = group[0]
                line.account_name = _resolve_account_name(fy.entity, line.account_code, line.account_name)
                line.sub_entries = []
                line.is_aggregated = False
                # variance_amount and variance_percentage are computed
                # by the TrialBalanceLine model @property — no need to set.
                agg_lines.append(line)
            else:
                # Multiple entries — build an aggregated summary row
                # Use the first line as the base (for mapping, risk flags, etc.)
                # but use a generic account name derived from the code
                first = group[0]
                raw_dr = sum(l.display_dr or Decimal('0') for l in group)
                raw_cr = sum(l.display_cr or Decimal('0') for l in group)
                agg_prior_dr = sum(l.prior_debit or Decimal('0') for l in group)
                agg_prior_cr = sum(l.prior_credit or Decimal('0') for l in group)

                # Net the debits and credits so adjustments reduce the
                # original balance instead of showing a separate column
                net = raw_dr - raw_cr
                if net >= 0:
                    agg_dr = net
                    agg_cr = Decimal('0')
                else:
                    agg_dr = Decimal('0')
                    agg_cr = abs(net)

                # Create a lightweight wrapper object for the aggregated row
                class AggregatedLine:
                    pass
                agg = AggregatedLine()
                agg.account_code = code
                # Use the first line's name if all names match, else use
                # a generic label showing the count
                unique_names = list(dict.fromkeys(l.account_name for l in group if l.account_name))
                if len(unique_names) == 1:
                    agg.account_name = _resolve_account_name(fy.entity, code, unique_names[0])
                else:
                    resolved_name = unique_names[0] if unique_names else code
                    agg.account_name = _resolve_account_name(fy.entity, code, resolved_name)
                agg.display_dr = agg_dr
                agg.display_cr = agg_cr
                agg.prior_debit = agg_prior_dr
                agg.prior_credit = agg_prior_cr
                agg.mapped_line_item = first.mapped_line_item
                agg.is_adjustment = any(l.is_adjustment for l in group)
                agg.prior_balance_override = any(getattr(l, 'prior_balance_override', False) for l in group)
                agg.reclassified = any(getattr(l, 'reclassified', False) for l in group)
                agg.risk_flags_list = first.risk_flags_list if hasattr(first, 'risk_flags_list') else []
                # Merge eva_flags across all sub-entries (deduplicated)
                merged_eva = []
                for l in group:
                    for f in (getattr(l, 'eva_flags', None) or []):
                        if f not in merged_eva:
                            merged_eva.append(f)
                agg.eva_flags = merged_eva
                agg.sub_entries = group
                agg.is_aggregated = True
                agg.sub_count = len(group)
                # Variance for aggregated row
                current_net = agg_dr - agg_cr
                prior_net = agg_prior_dr - agg_prior_cr
                agg.variance_amount = current_net - prior_net
                if prior_net != 0:
                    agg.variance_percentage = ((current_net - prior_net) / abs(prior_net) * 100).quantize(Decimal('0.1'))
                else:
                    agg.variance_percentage = None
                agg_lines.append(agg)

        aggregated_sections[section_name] = agg_lines

    # Recompute grand totals and net profit from AGGREGATED values so they
    # match the visible rows.  Pre-aggregation totals double-count opposite-
    # direction adjustments (e.g. a Cr adjustment on a normally-Dr account
    # would inflate both total_debit and total_credit).
    pl_sections = {'Income', 'Cost of Sales', 'Expenses'}
    total_debit = Decimal('0')
    total_credit = Decimal('0')
    total_prior_debit = Decimal('0')
    total_prior_credit = Decimal('0')
    pl_dr = Decimal('0')
    pl_cr = Decimal('0')
    pl_prior_dr = Decimal('0')
    pl_prior_cr = Decimal('0')
    for section_name, agg_lines in aggregated_sections.items():
        is_pl = section_name in pl_sections
        for line in agg_lines:
            dr = line.display_dr or Decimal('0')
            cr = line.display_cr or Decimal('0')
            py = getattr(line, '_py', None)
            if py is None:
                pdr = line.prior_debit or Decimal('0')
                pcr = line.prior_credit or Decimal('0')
            else:
                if py > 0:
                    pdr = py
                    pcr = Decimal('0')
                elif py < 0:
                    pdr = Decimal('0')
                    pcr = abs(py)
                else:
                    pdr = Decimal('0')
                    pcr = Decimal('0')
            total_debit += dr
            total_credit += cr
            total_prior_debit += pdr
            total_prior_credit += pcr
            if is_pl:
                pl_dr += dr
                pl_cr += cr
                pl_prior_dr += pdr
                pl_prior_cr += pcr
    net_profit = pl_cr - pl_dr
    prior_net_profit = pl_prior_cr - pl_prior_dr

    # Year labels
    current_year = str(fy.year_label)
    # Extract year number from label like "FY2026" or "2026"
    year_digits = ''.join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        prior_year = f"FY{int(year_digits) - 1}" if fy.year_label.startswith('FY') else str(int(year_digits) - 1)
    elif fy.prior_year:
        prior_year = str(fy.prior_year.year_label)
    else:
        prior_year = 'Prior'

    # Determine whether prior-year comparative data exists in the trial balance
    # and whether the entity wants comparatives shown.
    has_prior_tb = bool(
        entity.include_comparative_figures
        and fy.prior_year
        and fy.prior_year.trial_balance_lines.exists()
    )

    # Audit Risk flags
    risk_flags = RiskFlag.objects.filter(financial_year=fy).order_by('-severity', '-created_at')
    open_risk_count = risk_flags.filter(status='open').count()

    # Build a lookup: account_code -> list of open risk flag titles/descriptions
    # affected_accounts is a JSON list of dicts: [{"account_code": "2992", "account_name": "Clearing", ...}]
    flagged_accounts = {}  # {account_code: [{"title": ..., "severity": ..., "description": ...}]}
    for flag in risk_flags.filter(status__in=['open', 'reviewed']):
        for acc in (flag.affected_accounts or []):
            code = acc.get('account_code', '') if isinstance(acc, dict) else str(acc)
            if code:
                if code not in flagged_accounts:
                    flagged_accounts[code] = []
                flagged_accounts[code].append({
                    'title': flag.title,
                    'severity': flag.severity,
                    'description': flag.description[:120],
                })

    # Compute flagged account count (distinct account codes with open/reviewed flags)
    flagged_account_count = len(flagged_accounts)

    # Build grouped flags for the Audit Risk tab: account_code -> {name, flags[]}
    # Only include OPEN flags — resolved flags are removed from the list
    grouped_risk_flags = {}  # {account_code: {"name": ..., "code": ..., "flags": [...], "open_count": int, "max_severity": str}}
    entity_level_flags = []  # open flags with no affected_accounts
    severity_order = {"CRITICAL": 0, "HIGH": 1, "MEDIUM": 2, "LOW": 3}
    for flag in risk_flags.filter(status__in=['open', 'reviewed']):
        accounts = flag.affected_accounts or []
        if not accounts:
            entity_level_flags.append(flag)
        else:
            for acc in accounts:
                code = acc.get('account_code', '') if isinstance(acc, dict) else str(acc)
                name = acc.get('account_name', '') if isinstance(acc, dict) else ''
                if code:
                    if code not in grouped_risk_flags:
                        grouped_risk_flags[code] = {
                            "code": code,
                            "name": name,
                            "flags": [],
                            "open_count": 0,
                            "max_severity": "LOW",
                        }
                    grouped_risk_flags[code]["flags"].append(flag)
                    grouped_risk_flags[code]["open_count"] += 1
                    if severity_order.get(flag.severity, 3) < severity_order.get(grouped_risk_flags[code]["max_severity"], 3):
                        grouped_risk_flags[code]["max_severity"] = flag.severity
    # Sort grouped flags by severity (most severe first), then by account code
    sorted_grouped_flags = sorted(
        grouped_risk_flags.values(),
        key=lambda g: (severity_order.get(g["max_severity"], 3), g["code"])
    )
    entity_level_flag_count = len(entity_level_flags)

    # Annotate TB lines with risk flag info
    for line in tb_lines:
        line.risk_flags_list = flagged_accounts.get(line.account_code, [])

    # Annotate TB lines with Eva amber indicators
    from core.eva_amber import annotate_tb_lines_with_amber
    for section_name, section_lines in aggregated_sections.items():
        annotate_tb_lines_with_amber(section_lines)

    # Build Eva finding metadata lookup for TB flag rendering
    # (check_name -> {title, severity}) from the latest review's open findings
    eva_finding_meta = {}
    latest_review = (
        EvaReview.objects.filter(financial_year=fy)
        .order_by("-completed_at")
        .first()
    )
    if latest_review:
        for ef in latest_review.findings.filter(  # Sprint 1b: scope to FS domain
            domain='financial_statements',
            status__in=["open", "reopened"],
        ).values("check_name", "title", "severity"):
            eva_finding_meta[ef["check_name"]] = {
                "title": ef["title"],
                "severity": ef["severity"],
            }
    # Annotate each TB line with resolved eva_findings metadata
    for section_name, section_lines in aggregated_sections.items():
        for line in section_lines:
            raw_flags = getattr(line, "eva_flags", None) or []
            enriched = []
            for cn in raw_flags:
                meta = eva_finding_meta.get(cn)
                if meta:
                    enriched.append({
                        "check_name": cn,
                        "title": meta["title"],
                        "severity": meta["severity"],
                    })
            line.eva_finding_flags = enriched

    # Depreciation assets
    depreciation_assets = DepreciationAsset.objects.filter(financial_year=fy).select_related('source_transaction')
    dep_categories = OrderedDict()
    dep_total_opening = Decimal('0')
    dep_total_depreciation = Decimal('0')
    dep_total_closing = Decimal('0')
    for asset in depreciation_assets:
        if asset.category not in dep_categories:
            dep_categories[asset.category] = []
        dep_categories[asset.category].append(asset)
        dep_total_opening += asset.opening_wdv
        dep_total_depreciation += asset.depreciation_amount
        dep_total_closing += asset.closing_wdv

    # Stock items
    stock_items = StockItem.objects.filter(financial_year=fy)
    stock_total_opening = stock_items.aggregate(total=Sum('opening_value'))['total'] or Decimal('0')
    stock_total_closing = stock_items.aggregate(total=Sum('closing_value'))['total'] or Decimal('0')

    # Review items (pending transactions from bank statement uploads for this entity)
    from review.models import PendingTransaction, ReviewJob, ClassificationRule, EntityGSTSetting
    review_jobs = ReviewJob.objects.filter(entity=fy.entity)
    pending_review = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=False,
    ).select_related('job', 'matched_rule').order_by('date')
    confirmed_review = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=True,
    ).select_related('job').order_by('date')

    # Entity GST context for enhanced review workflow
    entity_is_gst_registered = getattr(fy.entity, 'is_gst_registered', True)
    entity_gst_registration_date = getattr(fy.entity, 'gst_registration_date', None)
    entity_gst_settings = EntityGSTSetting.objects.filter(
        entity=fy.entity
    ).filter(
        Q(financial_year=fy) | Q(financial_year__isnull=True)
    ).order_by('-created_at')
    entity_classification_rules = ClassificationRule.objects.filter(
        entity=fy.entity, is_active=True
    ).order_by('-created_at')

    # Activity / Audit trail — AuditLog + ActivityLog entries for this FY
    audit_logs_fy = AuditLog.objects.filter(
        affected_object_id=str(fy.pk),
    ).select_related('user').order_by('-timestamp')
    journal_pks = list(fy.adjusting_journals.values_list('pk', flat=True))
    journal_logs = AuditLog.objects.filter(
        affected_object_id__in=[str(pk) for pk in journal_pks],
    ).select_related('user').order_by('-timestamp') if journal_pks else AuditLog.objects.none()
    all_log_pks = set(audit_logs_fy.values_list('pk', flat=True)) | set(journal_logs.values_list('pk', flat=True))
    merged_audit = list(AuditLog.objects.filter(pk__in=all_log_pks).select_related('user').order_by('-timestamp'))
    for entry in merged_audit:
        entry.source = "audit"
        entry.sort_dt = entry.timestamp

    # Include ActivityLog entries (Eva findings, status changes, etc.)
    merged_activity = list(
        ActivityLog.objects.filter(financial_year=fy)
        .select_related('user')
        .order_by('-created_at')
    )
    for entry in merged_activity:
        entry.source = "activity"
        entry.sort_dt = entry.created_at

    activity_logs = sorted(
        merged_audit + merged_activity,
        key=lambda x: x.sort_dt,
        reverse=True,
    )

    # Check if this entity has bank statement uploads
    has_bank_statements = (
        TrialBalanceLine.objects.filter(financial_year=fy, source='bank_statement').exists()
        or review_jobs.exists()
    )

    # Bank account mapping for double-entry posting
    bank_account_mappings = BankAccountMapping.objects.filter(entity=fy.entity)
    # Try to find the mapping for the current review jobs
    active_bank_mapping = None
    if review_jobs.exists():
        latest_job = review_jobs.order_by('-received_at').first()
        if latest_job:
            active_bank_mapping = bank_account_mappings.filter(
                bsb=latest_job.bsb or '',
                account_number=latest_job.account_number or '',
            ).first()
            if not active_bank_mapping:
                active_bank_mapping = bank_account_mappings.filter(is_default=True).first()
    # Get bank/cash accounts from entity CoA for the mapping dropdown
    bank_coa_accounts = EntityChartOfAccount.objects.filter(
        entity=fy.entity, is_active=True,
        section__in=['assets', 'current_assets', 'non_current_assets'],
    ).filter(
        Q(account_name__icontains='bank') | Q(account_name__icontains='cash') |
        Q(classification__icontains='bank') | Q(classification__icontains='cash')
    ).order_by('account_code')

    # Opening balance detection: check if entity has no balance sheet balances
    needs_opening_balance = False
    opening_balance_bank_code = ''
    opening_balance_bank_name = ''
    if has_bank_statements and active_bank_mapping:
        bank_code = active_bank_mapping.tb_account_code
        opening_balance_bank_code = bank_code
        opening_balance_bank_name = active_bank_mapping.tb_account_name
        # Check if the bank account has any TB line that is NOT from bank_statement source
        # (i.e. an imported or manually entered opening balance)
        existing_bank_tb = TrialBalanceLine.objects.filter(
            financial_year=fy, account_code=bank_code,
        )
        has_non_bs_bank_line = existing_bank_tb.exclude(source='bank_statement').exists()
        # Also check if there's a prior year
        has_prior_year = FinancialYear.objects.filter(
            entity=fy.entity, end_date__lt=fy.start_date,
        ).exists()
        # Check if balance sheet has any lines at all (excluding bank_statement source)
        bs_sections = ['assets', 'current_assets', 'non_current_assets',
                       'liabilities', 'current_liabilities', 'non_current_liabilities',
                       'equity', 'capital_accounts']
        has_balance_sheet = TrialBalanceLine.objects.filter(
            financial_year=fy,
        ).exclude(source='bank_statement').exists()
        if not has_non_bs_bank_line and not has_prior_year and not has_balance_sheet:
            needs_opening_balance = True

    context = {
        "fy": fy,
        "entity": fy.entity,
        "has_bank_statements": has_bank_statements,
        "tb_lines": tb_lines,
        "tb_sections": aggregated_sections,
        "adjustments": adjustments,
        "bulk_journal_uploads": fy.bulk_journal_uploads.all().order_by('-created_at'),
        "unmapped_count": unmapped_count,
        "documents": documents,
        "total_debit": total_debit,
        "total_credit": total_credit,
        "total_prior_debit": total_prior_debit,
        "total_prior_credit": total_prior_credit,
        "net_profit": net_profit,
        "net_profit_abs": abs(net_profit),
        "prior_net_profit": prior_net_profit,
        "prior_net_profit_abs": abs(prior_net_profit),
        "current_year": current_year,
        "prior_year": prior_year,
        "has_prior_tb": has_prior_tb,
        # Audit Risk
        "risk_flags": risk_flags,
        "open_risk_count": open_risk_count,
        "flagged_account_count": flagged_account_count,
        "entity_level_flag_count": entity_level_flag_count,
        "entity_level_flags": entity_level_flags,
        "sorted_grouped_flags": sorted_grouped_flags,
        "flagged_accounts": flagged_accounts,
        # Depreciation
        "depreciation_assets": depreciation_assets,
        "dep_categories": dep_categories,
        "dep_total_opening": dep_total_opening,
        "dep_total_depreciation": dep_total_depreciation,
        "dep_total_closing": dep_total_closing,
        "dep_asset_accounts": list(
            EntityChartOfAccount.objects.filter(
                entity=fy.entity, is_active=True, section="assets",
            ).order_by("account_code").values_list("account_code", "account_name")
        ),
        "dep_expense_accounts": list(
            EntityChartOfAccount.objects.filter(
                entity=fy.entity, is_active=True, section="expenses",
            ).order_by("account_code").values_list("account_code", "account_name")
        ),
        # Stock
        "stock_items": stock_items,
        "stock_total_opening": stock_total_opening,
        "stock_total_closing": stock_total_closing,
        # Review
        "pending_review": pending_review,
        "confirmed_review": confirmed_review,
        "review_jobs": review_jobs,
        "entity_is_gst_registered": entity_is_gst_registered,
        "entity_gst_registration_date": entity_gst_registration_date,
        "entity_gst_settings": entity_gst_settings,
        "entity_classification_rules": entity_classification_rules,
        # Bank Account Mapping
        "bank_account_mappings": bank_account_mappings,
        "active_bank_mapping": active_bank_mapping,
        "bank_coa_accounts": bank_coa_accounts,
        # Opening Balance Detection
        "needs_opening_balance": needs_opening_balance,
        "opening_balance_bank_code": opening_balance_bank_code,
        "opening_balance_bank_name": opening_balance_bank_name,
        # Activity
        "activity_logs": activity_logs,
        # Chart of Accounts (entity-level)
        "entity_accounts": EntityChartOfAccount.objects.filter(
            entity=fy.entity, is_active=True
        ).select_related('maps_to').order_by('section', 'account_code'),
        "entity_accounts_count": EntityChartOfAccount.objects.filter(
            entity=fy.entity, is_active=True
        ).count(),
        "account_mappings": AccountMapping.objects.filter(
            applicable_entities__contains=fy.entity.entity_type
        ).order_by('financial_statement', 'line_item_label'),
        "active_tab": request.GET.get("tab", ""),
        # Finalise button state — show for draft, reopened, and in_review
        "can_finalise": fy.status in ("draft", "reopened", "in_review"),
        # Eva review running state (check for running/pending review records)
        "eva_review_running": fy.eva_reviews.filter(status__in=["pending", "running"]).exists(),
        # Reopen feature: list of subsequent finalised years for cascade info
        "subsequent_finalised_years": _get_subsequent_finalised_years(fy),
        # Eva Amber Indicators
        "amber_indicators": _compute_amber_indicators(fy),
        # BAS Period Commentaries for Documents tab (Client Communications)
        "bas_commentaries": BASPeriodCommentary.objects.filter(
            financial_year=fy,
            status__in=["draft", "reviewed", "sent"],
        ).select_related("generated_by", "bas_period").order_by("-generated_at"),
    }

    # --- Tax Journal context (companies only) ---
    if fy.entity.entity_type == "company" and fy.status == FinancialYear.Status.FINALISED:
        has_tax_journal = AdjustingJournal.objects.filter(
            financial_year=fy, description__icontains="Income tax",
        ).exists()
        context["show_tax_journal_btn"] = not has_tax_journal
        context["has_tax_journal"] = has_tax_journal
        if has_tax_journal:
            tax_jnl = AdjustingJournal.objects.filter(
                financial_year=fy, description__icontains="Income tax",
            ).first()
            context["tax_journal"] = tax_jnl

    # --- Management Accounts context ---
    from .mgmt_accounts import detect_tb_source
    context["tb_source"] = detect_tb_source(fy.entity)
    # Default period_end for the modal: end of most recent completed BAS period, or today
    last_bas = fy.bas_periods.filter(
        status__in=['lodged', 'ready', 'partial']
    ).order_by('-period_end').first()
    if last_bas:
        context["default_period_end"] = last_bas.period_end.isoformat()
    else:
        from datetime import date as _date
        context["default_period_end"] = min(_date.today(), fy.end_date).isoformat()
    context["fy_start_iso"] = fy.start_date.isoformat()
    context["fy_end_iso"] = fy.end_date.isoformat()

    # RDTI Drafter — pass application to template for tab badge
    context["rdti_application"] = getattr(fy, "rdti_application", None)
    return render(request, "core/financial_year_detail.html", context)


@login_required
def financial_year_finalise_full(request, pk):
    """One-click finalise: draft/reopened → in_review → finalised atomically.

    If the FY is already in_review, skips straight to finalisation.
    Runs all in_review side effects (Tier 1+2 risk recalc, Tier 3 AI)
    then finalises (locks TB, records finalised_at).
    """
    if request.method != "POST":
        return redirect("core:financial_year_detail", pk=pk)

    fy = get_financial_year_for_user(request, pk)

    if not request.user.can_finalise:
        messages.error(request, "Only senior accountants or admins can finalise.")
        return redirect("core:financial_year_detail", pk=pk)

    if fy.status not in ("draft", "reopened", "in_review"):
        messages.error(request, "This financial year cannot be finalised from its current status.")
        return redirect("core:financial_year_detail", pk=pk)

    # ── Pre-check: Trust distribution must be complete ───────────────
    if fy.entity.entity_type == "trust":
        trust_ws = getattr(fy, "trust_workspace", None)
        if not trust_ws or not trust_ws.all_stages_completed():
            messages.warning(
                request,
                "Trust Distribution tab has not been completed for this "
                "financial year. Please complete the trust distribution "
                "before finalising."
            )
            return redirect("core:financial_year_detail", pk=pk)

    # ── Step 1: Transition to in_review (if not already) ─────────────
    if fy.status in ("draft", "reopened"):
        fy.status = "in_review"
        fy.save(update_fields=["status"])

        # Replicate in_review side effects from financial_year_status
        from core.signals import trigger_risk_recalc
        trigger_risk_recalc(fy, "status_in_review", force=True)
        try:
            from .ai_service import batch_analyse_flags
            ai_result = batch_analyse_flags(fy, force=False)
            _log_action(
                request, "audit_run",
                f"Auto Tier 3 AI analysis on status→In Review: "
                f"{ai_result.get('analysed', 0)} analysed, "
                f"{ai_result.get('skipped', 0)} cached",
                fy
            )
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Auto Tier 3 on In Review failed: {e}")

    # ── Step 2: Block finalisation if trust balance sheet doesn't reconcile
    if fy.entity.entity_type == "trust":
        try:
            from core.fs_template_service import _get_tb_sections, _sum_section
            from decimal import Decimal
            _sections = _get_tb_sections(fy)
            _eq = -(
                _sum_section(_sections.get("equity", []))
                + _sum_section(_sections.get("pl_appropriation", []))
                + _sum_section(_sections.get("capital_accounts", []))
            )
            # Include net profit (revenue - expenses) in equity
            _income = (
                _sum_section(_sections.get("income", []))
                + _sum_section(_sections.get("trading_income", []))
            )
            _expenses = (
                _sum_section(_sections.get("expenses", []))
                + _sum_section(_sections.get("cogs", []))
            )
            _net_profit = -(_income + _expenses)
            _eq = _eq + _net_profit
            _liab = -(
                _sum_section(_sections.get("current_liabilities", []))
                + _sum_section(_sections.get("noncurrent_liabilities", []))
                + _sum_section(_sections.get("liabilities", []))
            )
            _assets = (
                _sum_section(_sections.get("current_assets", []))
                + _sum_section(_sections.get("noncurrent_assets", []))
            )
            _na = _assets - _liab
            if abs(_na - _eq) > Decimal("0.01"):
                messages.error(
                    request,
                    f"Cannot finalise: Net Assets ({_na:,.0f}) ≠ Total Equity ({_eq:,.0f}). "
                    f"Investigate the equity section before finalising."
                )
                return redirect("core:financial_year_detail", pk=pk)
        except Exception as _e:
            import logging
            logging.getLogger(__name__).warning("Balance sheet reconciliation check failed: %s", _e)

    # ── Step 3: Finalise (in_review → finalised) ─────────────────────
    fy.status = FinancialYear.Status.FINALISED
    fy.finalised_at = timezone.now()
    fy.save(update_fields=["status", "finalised_at"])
    fy.trial_balance_lines.update(comparatives_locked=True)

    _log_action(
        request, "status_change",
        f"{request.user.get_full_name() or request.user.email} finalised this financial year. "
        f"Trial balance locked. Eva compliance review now available.",
        fy,
    )

    ActivityLog.objects.create(
        user=request.user,
        financial_year=fy,
        entity=fy.entity,
        event_type="year_finalised",
        title=f"Finalised by {request.user.get_full_name() or request.user.username}",
        description=(
            f"Financial year finalised at {timezone.now().strftime('%d %b %Y %H:%M')}. "
            f"Eva compliance review is now available."
        ),
        url=f"/entities/years/{fy.pk}/",
    )

    messages.success(request, "Financial year finalised. Eva's compliance review is now available.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
def financial_year_status(request, pk):
    """Change the status of a financial year.

    The only valid manual transitions via the status dropdown are:
        draft → in_review
        in_review → draft
    Finalisation and reopening are handled by dedicated endpoints.
    """
    fy = get_financial_year_for_user(request, pk)
    new_status = request.POST.get("status")

    if not new_status or new_status not in dict(FinancialYear.Status.choices):
        messages.error(request, "Invalid status.")
        return redirect("core:financial_year_detail", pk=pk)

    # Allowed manual transitions via the status dropdown.
    # Finalise and reopen have dedicated endpoints but are included
    # here for admin correction purposes.
    ALLOWED_TRANSITIONS = {
        "draft": ["in_review"],
        "in_review": ["draft", "finalised"],
        "finalised": ["reopened"],
        "reopened": ["in_review"],
    }
    allowed = ALLOWED_TRANSITIONS.get(fy.status, [])
    if new_status not in allowed:
        messages.error(
            request,
            f"Cannot change status from '{fy.get_status_display()}' to '{new_status}'. "
            f"Allowed transitions: {', '.join(allowed) if allowed else 'none'}."
        )
        return redirect("core:financial_year_detail", pk=pk)

    # ── Tier 3 auto-trigger on milestone status changes ────────────
    if new_status == "in_review":
        # Auto-run Tier 1+2 to ensure flags are current
        from core.signals import trigger_risk_recalc
        trigger_risk_recalc(fy, "status_in_review", force=True)
        # Auto-run Tier 3 AI analysis on all open flags
        try:
            from .ai_service import batch_analyse_flags
            ai_result = batch_analyse_flags(fy, force=False)
            _log_action(
                request, "audit_run",
                f"Auto Tier 3 AI analysis on status→In Review: "
                f"{ai_result.get('analysed', 0)} analysed, "
                f"{ai_result.get('skipped', 0)} cached",
                fy
            )
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"Auto Tier 3 on In Review failed: {e}")

    # ── Draft revert: clear all Eva findings and risk flags ─────────
    if new_status == "draft":
        from django.db import transaction as db_transaction
        with db_transaction.atomic():
            deleted_reviews = fy.eva_reviews.all().delete()[0]  # cascades to EvaFinding
            deleted_flags = fy.risk_flags.all().delete()[0]
            # Also clear suppressions since we're resetting everything
            deleted_suppressions = fy.suppressed_findings.all().delete()[0]

            old_status = fy.status
            fy.status = new_status
            fy.save()

            _log_action(
                request, "status_change",
                f"Status reverted to Draft — all Eva findings cleared. Ready for re-review. "
                f"(deleted {deleted_reviews} Eva review(s), {deleted_flags} risk flag(s))",
                fy,
            )

            ActivityLog.objects.create(
                user=request.user,
                financial_year=fy,
                entity=fy.entity,
                event_type="general",
                title="Status reverted to Draft",
                description=(
                    f"Status reverted to Draft — all Eva findings cleared. Ready for re-review. "
                    f"Deleted {deleted_reviews} Eva review(s) and {deleted_flags} risk flag(s)."
                ),
            )

        messages.success(request, f"Status changed to {fy.get_status_display()}. All Eva findings and risk flags have been cleared.")
        return redirect("core:financial_year_detail", pk=pk)

    # Block finalisation if trust balance sheet does not reconcile
    if new_status == "finalised" and fy.entity.entity_type == "trust":
        try:
            from core.fs_template_service import _get_tb_sections, _sum_section
            from decimal import Decimal
            _sections = _get_tb_sections(fy)
            _eq = -(
                _sum_section(_sections.get("equity", []))
                + _sum_section(_sections.get("pl_appropriation", []))
                + _sum_section(_sections.get("capital_accounts", []))
            )
            # Include net profit (revenue - expenses) in equity
            _income = (
                _sum_section(_sections.get("income", []))
                + _sum_section(_sections.get("trading_income", []))
            )
            _expenses = (
                _sum_section(_sections.get("expenses", []))
                + _sum_section(_sections.get("cogs", []))
            )
            _net_profit = -(_income + _expenses)
            _eq = _eq + _net_profit
            _liab = -(
                _sum_section(_sections.get("current_liabilities", []))
                + _sum_section(_sections.get("noncurrent_liabilities", []))
                + _sum_section(_sections.get("liabilities", []))
            )
            _assets = (
                _sum_section(_sections.get("current_assets", []))
                + _sum_section(_sections.get("noncurrent_assets", []))
            )
            _na = _assets - _liab
            if abs(_na - _eq) > Decimal("0.01"):
                messages.error(
                    request,
                    f"Cannot finalise: Net Assets ({_na:,.0f}) ≠ Total Equity ({_eq:,.0f}). "
                    f"Investigate the equity section before finalising."
                )
                return redirect("core:financial_year_detail", pk=pk)
        except Exception as _e:
            logger.warning("Balance sheet reconciliation check failed: %s", _e)

    old_status = fy.status
    fy.status = new_status
    if new_status == "finalised":
        fy.finalised_at = timezone.now()
        fy.trial_balance_lines.update(comparatives_locked=True)
    fy.save()

    _log_action(
        request, "status_change",
        f"Changed {fy} from {old_status} to {new_status}", fy
    )

    ActivityLog.objects.create(
        user=request.user,
        financial_year=fy,
        entity=fy.entity,
        event_type="fy_status_changed",
        title=f"Status changed to {fy.get_status_display()}",
        description=f"Changed from {old_status} to {new_status}.",
        url=f"/entities/years/{fy.pk}/",
    )

    messages.success(request, f"Status changed to {fy.get_status_display()}.")
    return redirect("core:financial_year_detail", pk=pk)


# ============================================================
# Reopen Finalised Financial Year
# ============================================================

@login_required
def reopen_financial_year(request, pk):
    """Reopen a finalised financial year so amendments can be made.

    Optionally cascades the reopen to all subsequent financial years in
    the chain (linked via the ``prior_year`` / ``next_year`` relation).
    After reopening, the user can amend the year, re-finalise it, and
    then re-roll-forward to propagate changes to later years.
    """
    fy = get_financial_year_for_user(request, pk)

    if request.method != "POST":
        return redirect("core:financial_year_detail", pk=pk)

    # ── Permission check ─────────────────────────────────────────────
    if not request.user.can_finalise:
        messages.error(request, "Only senior accountants or admins can reopen a finalised year.")
        return redirect("core:financial_year_detail", pk=pk)
    # ── Status check ─────────────────────────────────────────────────
    if not fy.is_locked:
        messages.error(request, "This financial year is not finalised.")
        return redirect("core:financial_year_detail", pk=pk)
    # ── Collect inputs ───────────────────────────────────────────────
    reason = request.POST.get("reopen_reason", "").strip()
    if not reason:
        messages.error(request, "You must provide a reason for reopening.")
        return redirect("core:financial_year_detail", pk=pk)

    cascade = request.POST.get("cascade") == "1"

    # ── Build the list of years to reopen ────────────────────────────
    years_to_reopen = [fy]
    if cascade:
        current = fy
        while True:
            # next_year is the related_name on prior_year FK (reverse)
            nxt = current.next_year.first()
            if nxt and nxt.is_locked:
                years_to_reopen.append(nxt)
                current = nxt
            else:
                break

    # ── Perform the reopen (in reverse chronological order for safety) ──
    reopened_labels = []
    for year in reversed(years_to_reopen):
        old_status = year.status
        year.status = FinancialYear.Status.REOPENED
        year.reopened_at = timezone.now()
        year.reopened_by = request.user
        year.reopen_reason = reason
        year.finalised_at = None
        year.save()

        # Unlock trial balance comparatives
        year.trial_balance_lines.update(comparatives_locked=False)

        # Also unlock comparatives in the NEXT year (since this year's
        # closing balances feed that year's comparatives, and they are
        # now stale).
        next_fy = year.next_year.first()
        if next_fy:
            next_fy.trial_balance_lines.update(comparatives_locked=False)

        # Unlock generated documents (set back to draft)
        year.generated_documents.update(is_locked=False, status="draft")

        # Audit log
        _log_action(
            request, "reopen",
            f"Reopened {year.year_label} (was {old_status}). "
            f"Reason: {reason}. Cascade: {cascade}.",
            year,
        )
        reopened_labels.append(year.year_label)

    label_str = ", ".join(reopened_labels)
    messages.success(
        request,
        f"Successfully reopened {len(reopened_labels)} financial year(s): {label_str}. "
        f"You can now make amendments and re-finalise."
    )
    return redirect("core:financial_year_detail", pk=pk)


# ============================================================
# Re-Roll Forward — Diff & Apply (modal-based workflow)
# ============================================================

@login_required
def reroll_forward_diff(request, pk):
    """GET /api/years/<pk>/reroll-forward-diff/

    Compare this year's closing TB balances against the next year's
    opening balances and return a JSON diff.

    Income statement accounts (revenue/expense) are excluded — they reset
    to zero at year end.  Only balance sheet accounts carry forward.
    """
    current_fy = get_financial_year_for_user(request, pk)
    entity = current_fy.entity

    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied."}, status=403)

    next_fy = current_fy.next_year.first()
    if not next_fy:
        return JsonResponse({"error": "No subsequent financial year found."}, status=404)

    # Build CoA section lookup for BS/PL classification
    coa_sections = dict(
        ChartOfAccount.objects.filter(
            entity_type=entity.entity_type, is_active=True
        ).values_list("account_code", "section")
    )

    # Build aggregated closing balances for the current year (BS accounts only)
    from django.db.models import Sum as _Sum
    current_lines = (
        current_fy.trial_balance_lines
        .select_related("mapped_line_item")
        .filter(is_adjustment=False)
    )
    current_map = {}
    for line in current_lines:
        if not _is_balance_sheet_account(line.account_code, line.mapped_line_item, coa_sections):
            continue
        cb = line.closing_balance or Decimal("0")
        if line.account_code in current_map:
            current_map[line.account_code]["closing"] += cb
        else:
            current_map[line.account_code] = {
                "account_name": line.account_name,
                "closing": cb,
            }

    # Build opening balances in the next year (rollover lines only, BS only)
    next_rollover = (
        next_fy.trial_balance_lines
        .select_related("mapped_line_item")
        .filter(source="rollover", is_adjustment=False)
    )
    next_map = {}
    for line in next_rollover:
        if not _is_balance_sheet_account(line.account_code, line.mapped_line_item, coa_sections):
            continue
        next_map[line.account_code] = {
            "account_name": line.account_name,
            "opening": line.opening_balance,
            "line_id": str(line.pk),
        }

    changes = []
    for code, cur in current_map.items():
        closing = cur["closing"] or Decimal("0")
        if code in next_map:
            opening = next_map[code]["opening"] or Decimal("0")
            if closing != opening:
                changes.append({
                    "account_code": code,
                    "account_name": cur["account_name"],
                    "original_opening": str(opening),
                    "new_closing": str(closing),
                    "difference": str(closing - opening),
                    "line_id": next_map[code]["line_id"],
                })
        # New accounts in current year not in next year rollover are ignored
        # here — those require a full re-roll forward via the existing view.

    return JsonResponse({
        "status": "ok",
        "current_year": current_fy.year_label,
        "next_year": next_fy.year_label,
        "next_year_pk": str(next_fy.pk),
        "changes": changes,
        "change_count": len(changes),
    })


@login_required
@require_POST
def reroll_forward_apply(request, pk):
    """POST /api/years/<pk>/reroll-forward-apply/

    Accept the diff and update only the changed opening balance lines
    in the next financial year.

    Income statement accounts (revenue/expense) are excluded — they reset
    to zero at year end.  Only balance sheet accounts carry forward.
    """
    import json as _json
    current_fy = get_financial_year_for_user(request, pk)
    entity = current_fy.entity

    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied."}, status=403)

    next_fy = current_fy.next_year.first()
    if not next_fy:
        return JsonResponse({"error": "No subsequent financial year found."}, status=404)

    # Build CoA section lookup for BS/PL classification
    coa_sections = dict(
        ChartOfAccount.objects.filter(
            entity_type=entity.entity_type, is_active=True
        ).values_list("account_code", "section")
    )

    # Re-compute the diff to ensure consistency (don't trust client-side data).
    # Only include balance sheet accounts — P&L accounts reset to zero.
    from django.db.models import Sum as _Sum
    current_lines = (
        current_fy.trial_balance_lines
        .select_related("mapped_line_item")
        .filter(is_adjustment=False)
    )
    current_map = {}
    for line in current_lines:
        if not _is_balance_sheet_account(line.account_code, line.mapped_line_item, coa_sections):
            continue
        cb = line.closing_balance or Decimal("0")
        if line.account_code in current_map:
            current_map[line.account_code] += cb
        else:
            current_map[line.account_code] = cb

    next_rollover = (
        next_fy.trial_balance_lines
        .select_related("mapped_line_item")
        .filter(source="rollover", is_adjustment=False)
    )

    updated = []
    for line in next_rollover:
        # Skip P&L rollover lines — they should keep opening_balance=0
        if not _is_balance_sheet_account(line.account_code, line.mapped_line_item, coa_sections):
            continue
        new_closing = current_map.get(line.account_code)
        if new_closing is not None and new_closing != (line.opening_balance or Decimal("0")):
            old_opening = line.opening_balance
            line.opening_balance = new_closing
            line.closing_balance = new_closing + line.debit - line.credit
            line.save(update_fields=["opening_balance", "closing_balance"])
            updated.append({
                "account_code": line.account_code,
                "account_name": line.account_name,
                "old_opening": str(old_opening),
                "new_opening": str(new_closing),
            })

    # Log activity for both years
    if updated:
        summary_lines = "; ".join(
            f"{u['account_code']} {u['old_opening']} → {u['new_opening']}"
            for u in updated[:10]
        )
        if len(updated) > 10:
            summary_lines += f" ... and {len(updated) - 10} more"

        note = (
            f"Opening balances updated via Re-Roll Forward from {current_fy.year_label}. "
            f"{len(updated)} line(s) changed: {summary_lines}"
        )

        ActivityLog.objects.create(
            user=request.user,
            event_type="general",
            title=f"Re-Roll Forward Applied — {current_fy.year_label} → {next_fy.year_label}",
            description=note,
            entity=current_fy.entity,
            financial_year=current_fy,
            url=f"/entities/years/{current_fy.pk}/",
        )
        ActivityLog.objects.create(
            user=request.user,
            event_type="general",
            title=f"Opening Balances Updated via Re-Roll Forward from {current_fy.year_label}",
            description=note,
            entity=next_fy.entity,
            financial_year=next_fy,
            url=f"/entities/years/{next_fy.pk}/",
        )

        _log_action(
            request, "reroll_forward",
            f"Re-Roll Forward applied: {current_fy.year_label} → {next_fy.year_label}. "
            f"{len(updated)} opening balance(s) updated.",
            current_fy,
        )

    return JsonResponse({
        "status": "ok",
        "updated_count": len(updated),
        "updated": updated,
        "next_year_label": next_fy.year_label,
    })


# ============================================================
# Re-Roll Forward (full wipe & recreate — legacy view)
# ============================================================

@login_required
def reroll_forward(request, pk):
    """Re-roll forward from a finalised year into an existing subsequent year.

    Unlike the initial ``roll_forward`` which creates a brand-new FY, this
    view updates an *existing* next year by:
      1. Deleting all ``source='rollover'`` TrialBalanceLine rows in the
         next year (preserving any user-entered journals / imports).
      2. Deleting rolled-forward StockItems (notes contain 'Rolled forward').
      3. Re-running the same roll-forward logic to recreate those rows
         with the updated closing balances from the amended prior year.

    The user's own journals and imported data in the next year are preserved.
    """
    current_fy = get_financial_year_for_user(request, pk)
    entity = current_fy.entity

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    if not current_fy.is_locked:
        messages.error(request, "This financial year must be finalised before re-rolling forward.")
        return redirect("core:financial_year_detail", pk=pk)

    # Find the existing next year
    next_fy = current_fy.next_year.first()
    if not next_fy:
        messages.error(request, "No subsequent financial year found. Use Roll Forward instead.")
        return redirect("core:financial_year_detail", pk=pk)

    if request.method == "POST":
        from decimal import Decimal

        # ── Step 1: Full wipe of system-generated lines in the next year ────
        #
        # We delete ALL non-adjustment, non-bank-statement lines from the next
        # year regardless of their source tag.  This is the only reliable way
        # to handle the case where:
        #   (a) the next year was also imported from Handiledger (source='tb_import'),
        #   (b) the prior year has NEW accounts that don't exist in the next year yet,
        #   (c) the prior year has accounts whose balances changed.
        #
        # User-entered journals (is_adjustment=True) and bank statement lines
        # are explicitly preserved.
        #
        # The next year's own TB data (debit/credit/closing_balance columns)
        # is also preserved — we only wipe the opening_balance and prior-year
        # comparative columns, which are the ones we control via roll-forward.
        # The cleanest approach is to delete and recreate the non-adjustment
        # base lines so the next year's own activity (journals) sits on top.
        deleted_tb = next_fy.trial_balance_lines.filter(
            is_adjustment=False,
        ).exclude(
            source='bank_statement',
        ).delete()[0]

        deleted_stock = next_fy.stock_items.filter(
            notes__icontains="Rolled forward"
        ).delete()[0]
        # Delete previously rolled-forward depreciation assets so they are
        # recreated fresh from the (now-amended) current year's closing WDVs.
        deleted_dep = next_fy.depreciation_assets.filter(
            notes__icontains="Rolled forward"
        ).delete()[0]

        # ── Step 2: Re-run the roll-forward logic ────────────────────
        # (Replicates the core logic from roll_forward but targets the
        #  existing next_fy instead of creating a new one.)

        coa_sections = dict(
            ChartOfAccount.objects.filter(
                entity_type=entity.entity_type, is_active=True
            ).values_list("account_code", "section")
        )

        # -----------------------------------------------------------------
        # Pass 1: Classify lines and calculate net P&L
        #
        # Income statement accounts (revenue/expense) are excluded — they
        # reset to zero at year end.  Only balance sheet accounts (assets,
        # liabilities, equity) carry forward.
        # -----------------------------------------------------------------
        net_pl_result = Decimal("0")
        retained_profits_line = None
        income_tax_line = None
        bs_lines = []
        pl_lines = []

        for line in current_fy.trial_balance_lines.select_related("mapped_line_item").filter(is_adjustment=False):
            is_bs = _is_balance_sheet_account(line.account_code, line.mapped_line_item, coa_sections)

            if is_bs:
                # HARD RULE: only account 4199 is the retained profits target.
                code_prefix = line.account_code.split(".")[0] if line.account_code else ""
                if code_prefix == "4199":
                    retained_profits_line = line

                is_income_tax = False
                if line.account_name and "income tax" in line.account_name.lower():
                    is_income_tax = True
                if code_prefix == "4110":
                    is_income_tax = True
                if line.mapped_line_item and (line.mapped_line_item.standard_code or "") == "BS-EQ-011":
                    is_income_tax = True
                if is_income_tax:
                    income_tax_line = line

                bs_lines.append(line)
            else:
                pl_lines.append(line)
                net_pl_result += line.debit - line.credit

        # Pass 2: Create BS lines
        carried_bs = 0
        for line in bs_lines:
            if line.closing_balance == 0 and line != retained_profits_line:
                continue

            opening = line.closing_balance

            if line == retained_profits_line:
                tax_amount = income_tax_line.closing_balance if income_tax_line else Decimal("0")
                opening = line.closing_balance + net_pl_result + tax_amount
                if opening == 0:
                    continue

            if line == income_tax_line:
                TrialBalanceLine.objects.create(
                    financial_year=next_fy,
                    account_code=line.account_code,
                    account_name=line.account_name,
                    opening_balance=Decimal("0"),
                    debit=Decimal("0"),
                    credit=Decimal("0"),
                    closing_balance=Decimal("0"),
                    prior_debit=line.debit,
                    prior_credit=line.credit,
                    mapped_line_item=line.mapped_line_item,
                    is_adjustment=False,
                    source='rollover',
                )
                carried_bs += 1
                continue

            TrialBalanceLine.objects.create(
                financial_year=next_fy,
                account_code=line.account_code,
                account_name=line.account_name,
                opening_balance=opening,
                debit=Decimal("0"),
                credit=Decimal("0"),
                closing_balance=opening,
                prior_debit=line.debit,
                prior_credit=line.credit,
                mapped_line_item=line.mapped_line_item,
                is_adjustment=False,
                source='rollover',
            )
            carried_bs += 1

        tax_amount = income_tax_line.closing_balance if income_tax_line else Decimal("0")
        if retained_profits_line is None and (net_pl_result != 0 or tax_amount != 0):
            etype = entity.entity_type
            if etype == "trust":
                rp_name, rp_code = "Undistributed income", "4199"
            elif etype == "partnership":
                rp_name, rp_code = "Partners' current accounts", "4199"
            elif etype == "sole_trader":
                rp_name, rp_code = "Proprietor's funds", "4199"
            else:
                rp_name, rp_code = "Retained profits", "4199"

            rp_opening = net_pl_result + tax_amount
            TrialBalanceLine.objects.create(
                financial_year=next_fy,
                account_code=rp_code,
                account_name=rp_name,
                opening_balance=rp_opening,
                debit=Decimal("0"),
                credit=Decimal("0"),
                closing_balance=rp_opening,
                prior_debit=Decimal("0"),
                prior_credit=Decimal("0"),
                mapped_line_item=None,
                is_adjustment=False,
                source='rollover',
            )
            carried_bs += 1

            if income_tax_line:
                TrialBalanceLine.objects.create(
                    financial_year=next_fy,
                    account_code=income_tax_line.account_code,
                    account_name=income_tax_line.account_name,
                    opening_balance=Decimal("0"),
                    debit=Decimal("0"),
                    credit=Decimal("0"),
                    closing_balance=Decimal("0"),
                    prior_debit=income_tax_line.debit,
                    prior_credit=income_tax_line.credit,
                    mapped_line_item=income_tax_line.mapped_line_item,
                    is_adjustment=False,
                    source='rollover',
                )
                carried_bs += 1

        # Pass 3: P&L comparatives and stock conversion
        carried_pl = 0
        stock_converted = 0

        opening_stock_mapping = None
        try:
            opening_stock_mapping = AccountMapping.objects.get(standard_code="IS-COS-002")
        except AccountMapping.DoesNotExist:
            pass

        CLOSING_STOCK_KEYWORDS = [
            "closing stock", "closing work in progress", "closing wip",
            "closing raw material", "closing finished goods",
            "closing inventory",
        ]

        for line in pl_lines:
            if line.debit == 0 and line.credit == 0:
                continue

            is_closing_stock = False
            name_lower = (line.account_name or "").lower()
            if line.mapped_line_item and (line.mapped_line_item.standard_code or "") == "IS-COS-004":
                is_closing_stock = True
            if not is_closing_stock:
                for kw in CLOSING_STOCK_KEYWORDS:
                    if kw in name_lower:
                        is_closing_stock = True
                        break

            TrialBalanceLine.objects.create(
                financial_year=next_fy,
                account_code=line.account_code,
                account_name=line.account_name,
                opening_balance=Decimal("0"),
                debit=Decimal("0"),
                credit=Decimal("0"),
                closing_balance=Decimal("0"),
                prior_debit=line.debit,
                prior_credit=line.credit,
                mapped_line_item=line.mapped_line_item,
                is_adjustment=False,
                source='rollover',
            )
            carried_pl += 1

            if is_closing_stock:
                closing_amount = line.credit - line.debit
                if closing_amount != 0:
                    opening_name = line.account_name
                    for old, new in [("Closing", "Opening"), ("closing", "opening"), ("CLOSING", "OPENING")]:
                        opening_name = opening_name.replace(old, new)
                    if opening_name == line.account_name:
                        opening_name = "Opening " + line.account_name

                    opening_code = line.account_code
                    for pl_line in pl_lines:
                        pl_name_lower = (pl_line.account_name or "").lower()
                        if ("opening" in pl_name_lower and
                            ("stock" in pl_name_lower or "work in progress" in pl_name_lower
                             or "wip" in pl_name_lower or "inventory" in pl_name_lower
                             or "raw material" in pl_name_lower or "finished good" in pl_name_lower)):
                            opening_code = pl_line.account_code
                            opening_name = pl_line.account_name
                            break
                        if pl_line.mapped_line_item and (pl_line.mapped_line_item.standard_code or "") == "IS-COS-002":
                            opening_code = pl_line.account_code
                            opening_name = pl_line.account_name
                            break

                    TrialBalanceLine.objects.create(
                        financial_year=next_fy,
                        account_code=opening_code,
                        account_name=opening_name,
                        opening_balance=Decimal("0"),
                        debit=closing_amount if closing_amount > 0 else Decimal("0"),
                        credit=abs(closing_amount) if closing_amount < 0 else Decimal("0"),
                        closing_balance=Decimal("0"),
                        prior_debit=Decimal("0"),
                        prior_credit=Decimal("0"),
                        mapped_line_item=opening_stock_mapping,
                        is_adjustment=False,
                        source='rollover',
                    )
                    stock_converted += 1

        # Pass 4: Roll forward stock items
        stock_rolled = 0
        # Delete previously rolled stock items first
        next_fy.stock_items.filter(notes__icontains="Rolled forward").delete()
        for stock_item in current_fy.stock_items.all():
            if stock_item.closing_value == 0 and stock_item.closing_quantity == 0:
                continue
            StockItem.objects.create(
                financial_year=next_fy,
                item_name=stock_item.item_name,
                opening_quantity=stock_item.closing_quantity,
                opening_value=stock_item.closing_value,
                closing_quantity=Decimal("0"),
                closing_value=Decimal("0"),
                notes=f"Rolled forward from FY{current_fy.year_label}",
                display_order=stock_item.display_order,
            )
            stock_rolled += 1

        # Pass 5: Roll forward depreciation assets
        dep_rolled = 0
        for pa in current_fy.depreciation_assets.all():
            if pa.closing_wdv <= 0 and not pa.disposal_date:
                continue
            new_asset = DepreciationAsset(
                financial_year=next_fy,
                category=pa.category,
                asset_name=pa.asset_name,
                purchase_date=pa.purchase_date,
                total_cost=pa.total_cost,
                private_use_pct=pa.private_use_pct,
                opening_wdv=pa.closing_wdv,  # Prior closing becomes new opening
                method=pa.method,
                rate=pa.rate,
                display_order=pa.display_order,
                asset_account_code=pa.asset_account_code,
                asset_account_name=pa.asset_account_name,
                accum_dep_code=pa.accum_dep_code,
                accum_dep_name=pa.accum_dep_name,
                dep_expense_code=pa.dep_expense_code,
                dep_expense_name=pa.dep_expense_name,
                notes=f"Rolled forward from FY{current_fy.year_label}",
            )
            _calc_depreciation(new_asset)
            new_asset.save()
            dep_rolled += 1

        pl_direction = "profit" if net_pl_result < 0 else "loss"
        tax_msg = f" Income tax of ${abs(tax_amount):,.2f} absorbed." if tax_amount else ""
        stock_msg = f" {stock_converted} closing stock entries converted to opening stock." if stock_converted else ""
        stock_items_msg = f" {stock_rolled} stock items rolled forward." if stock_rolled else ""
        dep_msg = f" {dep_rolled} depreciation assets rolled forward." if dep_rolled else ""

        _log_action(
            request, "import",
            f"Re-rolled forward to {next_fy.year_label}: removed {deleted_tb} old rollover TB lines, "
            f"recreated {carried_bs} BS items + {carried_pl} P&L comparatives. "
            f"Net {pl_direction} of ${abs(net_pl_result):,.2f} closed to retained earnings."
            f"{tax_msg}{stock_msg}{stock_items_msg}{dep_msg}",
            next_fy,
        )
        messages.success(
            request,
            f"Re-rolled forward to {next_fy.year_label}. "
            f"Removed {deleted_tb} old rollover lines, recreated {carried_bs} BS items + "
            f"{carried_pl} P&L comparatives. Net {pl_direction} of ${abs(net_pl_result):,.2f} "
            f"less tax ${abs(tax_amount):,.2f} closed to retained earnings."
            f"{stock_msg}{stock_items_msg}{dep_msg}"
        )
        return redirect("core:financial_year_detail", pk=next_fy.pk)

    # GET: compute diff and show preview page
    from decimal import Decimal

    BS_STATEMENTS = {"balance_sheet", "equity"}
    BS_SECTIONS   = {"assets", "liabilities", "equity", "capital_accounts"}

    coa_sections = dict(
        ChartOfAccount.objects.filter(
            entity_type=entity.entity_type, is_active=True
        ).values_list("account_code", "section")
    )

    def _is_bs(line):
        hl_sec = _hl_section_for_code(line.account_code)
        if hl_sec is not None:
            return hl_sec not in ('Income', 'Cost of Sales', 'Expenses')
        if line.mapped_line_item:
            return line.mapped_line_item.financial_statement in BS_STATEMENTS
        if line.account_code in coa_sections:
            return coa_sections[line.account_code] in BS_SECTIONS
        code_prefix = line.account_code.split(".")[0] if line.account_code else ""
        return code_prefix.isdigit() and int(code_prefix) >= 2000

    # Build what the new opening balance WOULD be for each BS account in current_fy
    # (mirrors the Pass 1 / Pass 2 logic in the POST handler)
    net_pl = Decimal("0")
    retained_line = None
    income_tax_line = None
    bs_lines_cur = []
    pl_lines_cur = []
    for line in current_fy.trial_balance_lines.filter(is_adjustment=False):
        if _is_bs(line):
            code_prefix = line.account_code.split(".")[0] if line.account_code else ""
            if code_prefix == "4199":
                retained_line = line
            is_it = False
            if line.account_name and "income tax" in line.account_name.lower():
                is_it = True
            if code_prefix == "4110":
                is_it = True
            if line.mapped_line_item and (line.mapped_line_item.standard_code or "") == "BS-EQ-011":
                is_it = True
            if is_it:
                income_tax_line = line
            bs_lines_cur.append(line)
        else:
            pl_lines_cur.append(line)
            net_pl += line.debit - line.credit

    tax_amount = income_tax_line.closing_balance if income_tax_line else Decimal("0")

    # Compute proposed opening balances keyed by account_code
    proposed = {}  # account_code -> {name, opening, type: 'bs'|'pl'}
    for line in bs_lines_cur:
        if line.closing_balance == 0 and line != retained_line:
            continue
        if line == income_tax_line:
            proposed[line.account_code] = {
                "name": line.account_name,
                "proposed_opening": Decimal("0"),
                "type": "bs",
            }
            continue
        opening = line.closing_balance
        if line == retained_line:
            opening = line.closing_balance + net_pl + tax_amount
            if opening == 0:
                continue
        proposed[line.account_code] = {
            "name": line.account_name,
            "proposed_opening": opening,
            "type": "bs",
        }
    # P&L lines get zero opening (they only carry prior-year comparatives)
    for line in pl_lines_cur:
        if line.debit == 0 and line.credit == 0:
            continue
        proposed[line.account_code] = {
            "name": line.account_name,
            "proposed_opening": Decimal("0"),
            "type": "pl",
        }

    # Build current state of next_fy base lines (non-adjustment)
    existing = {}  # account_code -> opening_balance
    for line in next_fy.trial_balance_lines.filter(is_adjustment=False).exclude(source='bank_statement'):
        existing[line.account_code] = {
            "name": line.account_name,
            "current_opening": line.opening_balance or Decimal("0"),
        }

    # Compute diff
    diff_added   = []  # in proposed but not in existing
    diff_updated = []  # in both but opening_balance differs
    diff_removed = []  # in existing but not in proposed
    diff_unchanged = []  # in both and identical

    for code, pdata in proposed.items():
        if code not in existing:
            diff_added.append({
                "code": code,
                "name": pdata["name"],
                "current_opening": None,
                "proposed_opening": pdata["proposed_opening"],
                "type": pdata["type"],
            })
        else:
            cur_open = existing[code]["current_opening"]
            prop_open = pdata["proposed_opening"]
            if cur_open != prop_open:
                diff_updated.append({
                    "code": code,
                    "name": pdata["name"],
                    "current_opening": cur_open,
                    "proposed_opening": prop_open,
                    "type": pdata["type"],
                })
            else:
                diff_unchanged.append({
                    "code": code,
                    "name": pdata["name"],
                    "current_opening": cur_open,
                    "proposed_opening": prop_open,
                    "type": pdata["type"],
                })

    for code, edata in existing.items():
        if code not in proposed:
            diff_removed.append({
                "code": code,
                "name": edata["name"],
                "current_opening": edata["current_opening"],
                "proposed_opening": None,
                "type": "bs",
            })

    # Sort each list by account code for readability
    for lst in (diff_added, diff_updated, diff_removed, diff_unchanged):
        lst.sort(key=lambda x: x["code"])

    has_changes = bool(diff_added or diff_updated or diff_removed)

    return render(request, "core/reroll_forward_confirm.html", {
        "fy": current_fy,
        "next_fy": next_fy,
        "diff_added": diff_added,
        "diff_updated": diff_updated,
        "diff_removed": diff_removed,
        "diff_unchanged": diff_unchanged,
        "has_changes": has_changes,
    })


def _populate_rolled_forward_fy(current_fy, new_fy):
    """Populate a freshly-created financial year with rolled-forward data from
    ``current_fy``.  This is the single authoritative implementation of the
    five-pass roll-forward algorithm and is called by both the interactive
    ``roll_forward`` view and the engagement-letter generation path so that
    both routes produce identical results.

    Passes:
      1. Classify all base TB lines as BS or P&L; accumulate net P&L.
      2. Create BS opening-balance lines in new_fy (P&L closed to retained
         profits).
      3. Create P&L comparative-only lines (zero current, prior values set)
         and convert closing stock → opening stock.
      4. Roll forward StockItem records.
      5. Roll forward DepreciationAsset records.

    Returns a dict with summary counts for logging.
    """
    from decimal import Decimal as _D
    entity = current_fy.entity

    # Seed entity chart of accounts from template if not already done
    EntityChartOfAccount.seed_from_template(entity)

    # Build lookup of account_code -> section from ChartOfAccount
    coa_sections = dict(
        ChartOfAccount.objects.filter(
            entity_type=entity.entity_type, is_active=True
        ).values_list("account_code", "section")
    )

    # -----------------------------------------------------------------
    # Pass 1: Classify all lines and calculate net P&L result
    # -----------------------------------------------------------------
    net_pl_result = _D("0")
    retained_profits_line = None
    income_tax_line = None
    bs_lines = []
    pl_lines = []

    for line in current_fy.trial_balance_lines.select_related("mapped_line_item").filter(is_adjustment=False):
        is_bs = _is_balance_sheet_account(line.account_code, line.mapped_line_item, coa_sections)
        if is_bs:
            code_prefix = line.account_code.split(".")[0] if line.account_code else ""
            if code_prefix == "4199":
                retained_profits_line = line
            is_income_tax = False
            if line.account_name and "income tax" in line.account_name.lower():
                is_income_tax = True
            if code_prefix == "4110":
                is_income_tax = True
            if line.mapped_line_item and (line.mapped_line_item.standard_code or "") == "BS-EQ-011":
                is_income_tax = True
            if is_income_tax:
                income_tax_line = line
            bs_lines.append(line)
        else:
            pl_lines.append(line)
            net_pl_result += line.debit - line.credit

    # -----------------------------------------------------------------
    # Pass 2: Create balance sheet lines, adjusting retained profits
    # -----------------------------------------------------------------
    carried_bs = 0
    for line in bs_lines:
        if line.closing_balance == _D("0") and line != retained_profits_line:
            continue
        opening = line.closing_balance
        if line == retained_profits_line:
            tax_amount = income_tax_line.closing_balance if income_tax_line else _D("0")
            opening = line.closing_balance + net_pl_result + tax_amount
            if opening == _D("0"):
                continue
        if line == income_tax_line:
            TrialBalanceLine.objects.create(
                financial_year=new_fy,
                account_code=line.account_code,
                account_name=line.account_name,
                opening_balance=_D("0"),
                debit=_D("0"),
                credit=_D("0"),
                closing_balance=_D("0"),
                prior_debit=line.debit,
                prior_credit=line.credit,
                mapped_line_item=line.mapped_line_item,
                is_adjustment=False,
                source='rollover',
            )
            carried_bs += 1
            continue
        TrialBalanceLine.objects.create(
            financial_year=new_fy,
            account_code=line.account_code,
            account_name=line.account_name,
            opening_balance=opening,
            debit=_D("0"),
            credit=_D("0"),
            closing_balance=opening,
            prior_debit=line.debit,
            prior_credit=line.credit,
            mapped_line_item=line.mapped_line_item,
            is_adjustment=False,
            source='rollover',
        )
        carried_bs += 1

    # If no retained profits line existed, create one to hold the net P&L
    tax_amount = income_tax_line.closing_balance if income_tax_line else _D("0")
    if retained_profits_line is None and (net_pl_result != _D("0") or tax_amount != _D("0")):
        etype = entity.entity_type
        if etype == "trust":
            rp_name, rp_code = "Undistributed income", "4199"
        elif etype == "partnership":
            rp_name, rp_code = "Partners' current accounts", "4199"
        elif etype == "sole_trader":
            rp_name, rp_code = "Proprietor's funds", "4199"
        else:
            rp_name, rp_code = "Retained profits", "4199"
        rp_opening = net_pl_result + tax_amount
        TrialBalanceLine.objects.create(
            financial_year=new_fy,
            account_code=rp_code,
            account_name=rp_name,
            opening_balance=rp_opening,
            debit=_D("0"),
            credit=_D("0"),
            closing_balance=rp_opening,
            prior_debit=_D("0"),
            prior_credit=_D("0"),
            mapped_line_item=None,
            is_adjustment=False,
            source='rollover',
        )
        carried_bs += 1
        if income_tax_line:
            TrialBalanceLine.objects.create(
                financial_year=new_fy,
                account_code=income_tax_line.account_code,
                account_name=income_tax_line.account_name,
                opening_balance=_D("0"),
                debit=_D("0"),
                credit=_D("0"),
                closing_balance=_D("0"),
                prior_debit=income_tax_line.debit,
                prior_credit=income_tax_line.credit,
                mapped_line_item=income_tax_line.mapped_line_item,
                is_adjustment=False,
                source='rollover',
            )
            carried_bs += 1

    # -----------------------------------------------------------------
    # Pass 3: Create P&L comparative lines and convert closing stock
    # -----------------------------------------------------------------
    carried_pl = 0
    stock_converted = 0
    opening_stock_mapping = None
    closing_stock_mapping = None
    try:
        opening_stock_mapping = AccountMapping.objects.get(standard_code="IS-COS-002")
    except AccountMapping.DoesNotExist:
        pass
    try:
        closing_stock_mapping = AccountMapping.objects.get(standard_code="IS-COS-004")
    except AccountMapping.DoesNotExist:
        pass
    CLOSING_STOCK_KEYWORDS = [
        "closing stock", "closing work in progress", "closing wip",
        "closing raw material", "closing finished goods", "closing inventory",
    ]
    for line in pl_lines:
        if line.debit == _D("0") and line.credit == _D("0"):
            continue
        is_closing_stock = False
        name_lower = (line.account_name or "").lower()
        if line.mapped_line_item and (line.mapped_line_item.standard_code or "") == "IS-COS-004":
            is_closing_stock = True
        if not is_closing_stock:
            for kw in CLOSING_STOCK_KEYWORDS:
                if kw in name_lower:
                    is_closing_stock = True
                    break
        TrialBalanceLine.objects.create(
            financial_year=new_fy,
            account_code=line.account_code,
            account_name=line.account_name,
            opening_balance=_D("0"),
            debit=_D("0"),
            credit=_D("0"),
            closing_balance=_D("0"),
            prior_debit=line.debit,
            prior_credit=line.credit,
            mapped_line_item=line.mapped_line_item,
            is_adjustment=False,
            source='rollover',
        )
        carried_pl += 1
        if is_closing_stock:
            closing_amount = line.credit - line.debit
            if closing_amount != _D("0"):
                opening_name = line.account_name
                for old, new_word in [("Closing", "Opening"), ("closing", "opening"), ("CLOSING", "OPENING")]:
                    opening_name = opening_name.replace(old, new_word)
                if opening_name == line.account_name:
                    opening_name = "Opening " + line.account_name
                opening_code = line.account_code
                for pl_line in pl_lines:
                    pl_name_lower = (pl_line.account_name or "").lower()
                    if ("opening" in pl_name_lower and
                        ("stock" in pl_name_lower or "work in progress" in pl_name_lower
                         or "wip" in pl_name_lower or "inventory" in pl_name_lower
                         or "raw material" in pl_name_lower or "finished good" in pl_name_lower)):
                        opening_code = pl_line.account_code
                        opening_name = pl_line.account_name
                        break
                    if pl_line.mapped_line_item and (pl_line.mapped_line_item.standard_code or "") == "IS-COS-002":
                        opening_code = pl_line.account_code
                        opening_name = pl_line.account_name
                        break
                TrialBalanceLine.objects.create(
                    financial_year=new_fy,
                    account_code=opening_code,
                    account_name=opening_name,
                    opening_balance=_D("0"),
                    debit=closing_amount if closing_amount > _D("0") else _D("0"),
                    credit=abs(closing_amount) if closing_amount < _D("0") else _D("0"),
                    closing_balance=_D("0"),
                    prior_debit=_D("0"),
                    prior_credit=_D("0"),
                    mapped_line_item=opening_stock_mapping,
                    is_adjustment=False,
                    source='rollover',
                )
                stock_converted += 1

    # -----------------------------------------------------------------
    # Pass 4: Roll forward stock items (closing → opening)
    # -----------------------------------------------------------------
    stock_rolled = 0
    for stock_item in current_fy.stock_items.all():
        if stock_item.closing_value == _D("0") and stock_item.closing_quantity == _D("0"):
            continue
        StockItem.objects.create(
            financial_year=new_fy,
            item_name=stock_item.item_name,
            opening_quantity=stock_item.closing_quantity,
            opening_value=stock_item.closing_value,
            closing_quantity=_D("0"),
            closing_value=_D("0"),
            notes=f"Rolled forward from FY{current_fy.year_label}",
            display_order=stock_item.display_order,
        )
        stock_rolled += 1

    # -----------------------------------------------------------------
    # Pass 5: Roll forward depreciation assets (closing WDV → opening WDV)
    # -----------------------------------------------------------------
    dep_rolled = 0
    for pa in current_fy.depreciation_assets.all():
        if pa.closing_wdv <= 0 and not pa.disposal_date:
            continue
        new_asset = DepreciationAsset(
            financial_year=new_fy,
            category=pa.category,
            asset_name=pa.asset_name,
            purchase_date=pa.purchase_date,
            total_cost=pa.total_cost,
            private_use_pct=pa.private_use_pct,
            opening_wdv=pa.closing_wdv,
            method=pa.method,
            rate=pa.rate,
            display_order=pa.display_order,
            asset_account_code=pa.asset_account_code,
            asset_account_name=pa.asset_account_name,
            accum_dep_code=pa.accum_dep_code,
            accum_dep_name=pa.accum_dep_name,
            dep_expense_code=pa.dep_expense_code,
            dep_expense_name=pa.dep_expense_name,
            notes=f"Rolled forward from FY{current_fy.year_label}",
        )
        _calc_depreciation(new_asset)
        new_asset.save()
        dep_rolled += 1

    # -----------------------------------------------------------------
    # Pass 6: For trusts — allocate confirmed distribution to beneficiary
    # payable accounts and reduce undistributed income accordingly.
    # -----------------------------------------------------------------
    dist_rolled = 0
    TRUST_TYPES = {"trust", "trust_unit", "trust_discretionary", "trust_hybrid"}
    if entity.entity_type in TRUST_TYPES:
        try:
            from core.models import TrustWorkspace, DistributionScenario, BeneficiaryProfile
            workspace = TrustWorkspace.objects.filter(financial_year=current_fy).first()
            if workspace:
                confirmed = workspace.scenarios.filter(is_confirmed=True).first()
                if confirmed and confirmed.allocations:
                    # Build beneficiary name lookup: str(uuid) -> full_name
                    bene_ids = list(confirmed.allocations.keys())
                    officer_map = {}
                    for bp in BeneficiaryProfile.objects.filter(
                        trust_workspace=workspace
                    ).select_related("beneficiary"):
                        officer_map[str(bp.beneficiary_id)] = bp.beneficiary.full_name

                    total_distributed = _D("0")
                    suffix = 1
                    for bene_id, streams in confirmed.allocations.items():
                        bene_total = sum(
                            _D(str(v)) for v in streams.values() if v
                        )
                        if bene_total <= _D("0"):
                            continue
                        bene_name = officer_map.get(bene_id, f"Beneficiary {suffix}")
                        account_code = f"3100.{suffix:02d}"
                        account_name = f"Distribution Payable — {bene_name}"
                        TrialBalanceLine.objects.create(
                            financial_year=new_fy,
                            account_code=account_code,
                            account_name=account_name,
                            opening_balance=-bene_total,  # credit = liability
                            debit=_D("0"),
                            credit=_D("0"),
                            closing_balance=-bene_total,
                            prior_debit=_D("0"),
                            prior_credit=_D("0"),
                            mapped_line_item=None,
                            is_adjustment=False,
                            source="rollover",
                        )
                        total_distributed += bene_total
                        dist_rolled += 1
                        suffix += 1

                    # Reduce the undistributed income / retained profits line
                    # in new_fy by the total distributed amount so the BS balances.
                    if total_distributed > _D("0"):
                        rp_line = new_fy.trial_balance_lines.filter(
                            account_code="4199", source="rollover"
                        ).first()
                        if rp_line:
                            rp_line.opening_balance -= total_distributed
                            rp_line.closing_balance -= total_distributed
                            rp_line.save(update_fields=["opening_balance", "closing_balance"])
        except Exception:
            pass  # Distribution allocation is best-effort; never block roll-forward

    return {
        "carried_bs": carried_bs,
        "carried_pl": carried_pl,
        "stock_converted": stock_converted,
        "stock_rolled": stock_rolled,
        "dep_rolled": dep_rolled,
        "net_pl_result": net_pl_result,
        "tax_amount": tax_amount,
        "dist_rolled": dist_rolled,
    }


@login_required
def roll_forward(request, pk):
    """Create a new financial year from the current one, carrying closing balances."""
    current_fy = get_financial_year_for_user(request, pk)
    entity = current_fy.entity

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    # Cannot roll forward unless the current FY is finalised (locked or finalised status)
    if not current_fy.is_locked:
        messages.error(
            request,
            "Cannot roll forward: this financial year must be finalised before rolling forward."
        )
        return redirect("core:financial_year_detail", pk=pk)

    # Calculate the target year label for display
    from dateutil.relativedelta import relativedelta as _rdelta
    _new_end = current_fy.end_date + _rdelta(years=1)
    _new_label = str(_new_end.year)

    if request.method == "POST":
        # Calculate new dates (add 1 year)
        from dateutil.relativedelta import relativedelta
        from datetime import timedelta
        new_start = current_fy.end_date + relativedelta(days=1)
        new_end = new_start + relativedelta(years=1) - timedelta(days=1)
        new_label = str(new_end.year)

        # Check if already exists
        if entity.financial_years.filter(year_label=new_label).exists():
            messages.warning(request, f"{new_label} already exists for this entity.")
            return redirect("core:financial_year_detail", pk=pk)

        # Create new FY
        new_fy = FinancialYear.objects.create(
            entity=entity,
            year_label=new_label,
            start_date=new_start,
            end_date=new_end,
            prior_year=current_fy,
            status=FinancialYear.Status.DRAFT,
        )

        # Run the full five-pass roll-forward algorithm (same logic used by
        # the engagement-letter-triggered path via _populate_rolled_forward_fy)
        rf = _populate_rolled_forward_fy(current_fy, new_fy)
        carried_bs = rf["carried_bs"]
        carried_pl = rf["carried_pl"]
        stock_converted = rf["stock_converted"]
        stock_rolled = rf["stock_rolled"]
        dep_rolled = rf["dep_rolled"]
        net_pl_result = rf["net_pl_result"]
        tax_amount = rf["tax_amount"]
        dist_rolled = rf.get("dist_rolled", 0)

        total_carried = carried_bs + carried_pl
        pl_direction = "profit" if net_pl_result < 0 else "loss"
        tax_msg = f" Income tax of ${abs(tax_amount):,.2f} absorbed." if tax_amount else ""
        stock_msg = f" {stock_converted} closing stock entries converted to opening stock." if stock_converted else ""
        stock_items_msg = f" {stock_rolled} stock items rolled forward." if stock_rolled else ""
        dep_msg = f" {dep_rolled} depreciation assets rolled forward." if dep_rolled else ""
        dist_msg = f" {dist_rolled} distribution payable accounts created from confirmed trust distribution." if dist_rolled else ""
        _log_action(request, "import", f"Rolled forward to {new_label} with {carried_bs} BS items, {carried_pl} P&L comparatives. Net {pl_direction} of ${abs(net_pl_result):,.2f} closed to retained earnings.{tax_msg}{stock_msg}{stock_items_msg}{dep_msg}{dist_msg}", new_fy)
        messages.success(request, f"Rolled forward to {new_label}. {carried_bs} balance sheet items carried, {carried_pl} P&L comparatives. Net {pl_direction} of ${abs(net_pl_result):,.2f} less tax ${abs(tax_amount):,.2f} closed to retained earnings.{stock_msg}{stock_items_msg}{dep_msg}{dist_msg}")
        return redirect("core:financial_year_detail", pk=new_fy.pk)

    return render(request, "core/roll_forward_confirm.html", {
        "fy": current_fy,
        "target_year_label": _new_label,
    })


# ---------------------------------------------------------------------------
# Trial Balance Import
# ---------------------------------------------------------------------------
@login_required
def trial_balance_import(request, pk):
    fy = get_financial_year_for_user(request, pk)

    if fy.is_locked:
        messages.error(request, "Cannot import into a finalised financial year.")
        return redirect("core:financial_year_detail", pk=pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    if request.method == "POST":
        form = TrialBalanceUploadForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES["file"]

            # Validate file type and size
            if uploaded_file.size > 20 * 1024 * 1024:
                messages.error(request, "File too large. Maximum size is 20MB.")
                return redirect("core:financial_year_detail", pk=pk)

            import os as _os
            file_ext = _os.path.splitext(uploaded_file.name)[1].lower()
            if file_ext not in {".xlsx", ".xls"}:
                messages.error(request, f"Unsupported file type: {file_ext}. Only Excel files (.xlsx) are supported.")
                return redirect("core:financial_year_detail", pk=pk)

            file = uploaded_file
            try:
                # Parse the Excel and stage the data for review wizard
                raw_lines = _parse_tb_excel(fy, file)
                if not raw_lines:
                    messages.error(request, "No data rows found in the uploaded file.")
                    return redirect("core:trial_balance_import", pk=pk)

                # Merge duplicate account codes before mapping
                from core.tb_dedup import merge_duplicate_accounts
                raw_lines, merge_warnings = merge_duplicate_accounts(raw_lines)
                for w in merge_warnings:
                    messages.warning(request, w)

                # Apply learned mappings and code matching
                staged_lines = _apply_tb_learned_mappings(fy.entity, raw_lines)

                # Store in session for the review wizard
                request.session["staged_tb_import"] = {
                    "fy_pk": str(fy.pk),
                    "lines": staged_lines,
                    "filename": uploaded_file.name,
                    "merge_warnings": merge_warnings,
                }
                # Force session save to DB before redirect so the next
                # request (possibly handled by a different Gunicorn worker)
                # can read the staged data from the database backend.
                request.session.modified = True
                request.session.save()

                return redirect("core:review_tb_import", pk=pk)
            except Exception as e:
                messages.error(request, f"Import failed: {e}. Please check the file format and try again.")
    else:
        form = TrialBalanceUploadForm()

    return render(request, "core/trial_balance_import.html", {
        "form": form, "fy": fy
    })


def _process_trial_balance_upload(fy, file):
    """Process an Excel trial balance upload.

    Preserves prior-year comparative values (prior_debit, prior_credit,
    prior_closing_balance, prior_mapped_line_item) that were set during
    rollover.  Only current-year balances are replaced by the upload.

    Supports two formats:
    1. Simple TB: [code, name, (opening), debit, credit]
    2. HandiLedger Comparative TB: header rows, then
       [code, name, CY_dr, CY_cr, PY_dr, PY_cr]
    """
    import re as _re

    wb = openpyxl.load_workbook(file, read_only=True, data_only=True)
    ws = wb.active

    entity = fy.entity

    # ------------------------------------------------------------------
    # Snapshot existing comparative data BEFORE deleting lines.
    # Key = account_code.  We keep the first occurrence per code.
    # ------------------------------------------------------------------
    prior_data = {}  # account_code -> dict of comparative fields + original mapping
    for line in fy.trial_balance_lines.filter(is_adjustment=False).order_by("account_code"):
        if line.account_code not in prior_data:
            prior_data[line.account_code] = {
                "prior_debit": line.prior_debit,
                "prior_credit": line.prior_credit,
                "prior_closing_balance": line.prior_closing_balance,
                "prior_balance_override": line.prior_balance_override,
                "prior_mapped_line_item": line.prior_mapped_line_item,
                "reclassified": line.reclassified,
                "comparatives_locked": line.comparatives_locked,
                # Also preserve the current mapping and account name so
                # comparative-only lines don't become unmapped after upload
                "mapped_line_item": line.mapped_line_item,
                "account_name": line.account_name,
            }

    # Clear existing non-adjustment lines (current-year data will be
    # re-created from the Excel file; comparatives restored from snapshot)
    fy.trial_balance_lines.filter(is_adjustment=False).delete()

    imported = 0
    unmapped = 0
    errors = []
    uploaded_codes = set()  # Track which account codes came from the Excel
    comp_applied = set()   # Track codes whose comparatives have already been applied

    all_rows = list(ws.iter_rows(values_only=True))

    # --- Detect HandiLedger comparative format ---
    is_comparative = False
    data_start_row = 1  # 0-indexed; default skip first row
    section_names = {
        'income', 'expenses', 'cost of sales', 'current assets',
        'non-current assets', 'non current assets', 'fixed assets',
        'current liabilities', 'non-current liabilities',
        'non current liabilities', 'equity', 'capital',
    }

    for idx, row in enumerate(all_rows[:8]):
        row_strs = [str(c).strip().lower() if c else '' for c in row]
        year_pattern = sum(1 for s in row_strs if _re.match(r'^20\d{2}$', s))
        if year_pattern >= 2:
            is_comparative = True
            continue
        if '$ dr' in row_strs and '$ cr' in row_strs:
            is_comparative = True
            data_start_row = idx + 1
            continue
        if row_strs[0] and 'comparative' in row_strs[0]:
            is_comparative = True
            continue

    for i in range(data_start_row if is_comparative else 1, len(all_rows)):
        row = all_rows[i]
        if not row:
            continue

        col0 = str(row[0]).strip() if row[0] is not None else ''
        col1 = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ''

        if is_comparative:
            # Skip blank, section header, and totals rows
            if not col0 and not col1:
                continue
            if not col0 and col1.lower() in section_names:
                continue
            if col1.lower().startswith('total') or col1.lower().startswith('net '):
                continue
            if not col0 or not _re.match(r'^[\d]', col0):
                continue
        else:
            if not col0:
                continue

        try:
            account_code = col0
            # Strip leading zero from revenue accounts (HandiLedger 0575 -> 575)
            if account_code.startswith('0') and account_code.replace('.', '').isdigit() and len(account_code.split('.')[0]) > 1:
                parts = account_code.split('.', 1)
                parts[0] = parts[0].lstrip('0') or '0'
                account_code = '.'.join(parts)
            account_name = _resolve_account_name(entity, account_code, col1)
            if is_comparative:
                # Columns: [code, name, CY_dr, CY_cr, PY_dr, PY_cr]
                opening_balance = Decimal("0")
                debit = Decimal(str(row[2] or 0)) if len(row) > 2 and row[2] is not None else Decimal("0")
                credit = Decimal(str(row[3] or 0)) if len(row) > 3 and row[3] is not None else Decimal("0")
                file_py_dr = Decimal(str(row[4] or 0)) if len(row) > 4 and row[4] is not None else Decimal("0")
                file_py_cr = Decimal(str(row[5] or 0)) if len(row) > 5 and row[5] is not None else Decimal("0")
            elif len(row) >= 5 and row[4] is not None:
                # Legacy 5-column format with Opening Balance
                opening_balance = Decimal(str(row[2] or 0))
                debit = Decimal(str(row[3] or 0))
                credit = Decimal(str(row[4] or 0))
                file_py_dr = Decimal("0")
                file_py_cr = Decimal("0")
            else:
                # New 4-column format (no Opening Balance)
                opening_balance = Decimal("0")
                debit = Decimal(str(row[2] or 0))
                credit = Decimal(str(row[3] or 0))
                file_py_dr = Decimal("0")
                file_py_cr = Decimal("0")

            closing_balance = opening_balance + debit - credit
        except (InvalidOperation, ValueError, IndexError) as e:
            errors.append(f"Row {i + 1}: {str(e)}")
            continue

        # Try to find existing mapping for this entity (learning from prior imports)
        mapping = ClientAccountMapping.objects.filter(
            entity=entity, client_account_code=account_code
        ).first()

        mapped_item = mapping.mapped_line_item if mapping else None

        # If no existing mapping, try to match against the standard chart of accounts
        # for this entity type. This provides automatic mapping for known account codes.
        coa_match = None
        if not mapped_item:
            coa_match = ChartOfAccount.objects.filter(
                entity_type=entity.entity_type,
                account_code=account_code,
                is_active=True,
            ).first()
            if coa_match and coa_match.maps_to:
                mapped_item = coa_match.maps_to

        # Restore comparative values from snapshot — but only for the FIRST
        # row per account_code.
        if account_code not in comp_applied:
            comp = prior_data.get(account_code, {})
            comp_applied.add(account_code)
        else:
            comp = {}

        # Prior year: prefer file data (comparative TB), then DB snapshot
        if file_py_dr or file_py_cr:
            py_debit = file_py_dr
            py_credit = file_py_cr
            py_closing = file_py_dr - file_py_cr
        else:
            py_debit = comp.get("prior_debit", Decimal("0"))
            py_credit = comp.get("prior_credit", Decimal("0"))
            py_closing = comp.get("prior_closing_balance", Decimal("0"))

        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code=account_code,
            account_name=account_name,
            opening_balance=opening_balance,
            debit=debit,
            credit=credit,
            closing_balance=closing_balance,
            mapped_line_item=mapped_item,
            is_adjustment=False,
            source='tb_import',
            prior_debit=py_debit,
            prior_credit=py_credit,
            prior_closing_balance=py_closing,
            prior_balance_override=comp.get("prior_balance_override", False),
            prior_mapped_line_item=comp.get("prior_mapped_line_item"),
            reclassified=comp.get("reclassified", False),
            comparatives_locked=comp.get("comparatives_locked", False),
        )

        # Create or update client account mapping record
        ClientAccountMapping.objects.update_or_create(
            entity=entity,
            client_account_code=account_code,
            defaults={"client_account_name": account_name, "mapped_line_item": mapped_item},
        )

        uploaded_codes.add(account_code)
        imported += 1
        if not mapped_item:
            unmapped += 1

    # ------------------------------------------------------------------
    # Re-create comparative-only lines for accounts that existed in the
    # prior snapshot but were NOT in the uploaded Excel.  These are
    # typically P&L accounts from the prior year that have no current-year
    # activity yet but must appear in the comparative column.
    # ------------------------------------------------------------------
    for code, comp in prior_data.items():
        if code in uploaded_codes:
            continue  # Already handled above
        if comp["prior_debit"] == 0 and comp["prior_credit"] == 0:
            continue  # No comparative data to preserve

        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code=code,
            account_name=comp.get("account_name", ""),
            opening_balance=Decimal("0"),
            debit=Decimal("0"),
            credit=Decimal("0"),
            closing_balance=Decimal("0"),
            mapped_line_item=comp.get("mapped_line_item") or comp.get("prior_mapped_line_item"),
            is_adjustment=False,
            source='rollover',
            prior_debit=comp["prior_debit"],
            prior_credit=comp["prior_credit"],
            prior_closing_balance=comp.get("prior_closing_balance", Decimal("0")),
            prior_balance_override=comp.get("prior_balance_override", False),
            prior_mapped_line_item=comp.get("prior_mapped_line_item"),
            reclassified=comp.get("reclassified", False),
            comparatives_locked=comp.get("comparatives_locked", False),
        )

    wb.close()
    return {"imported": imported, "unmapped": unmapped, "errors": errors}


def _parse_tb_excel(fy, file):
    """Parse an Excel TB file into a list of raw line dicts (without committing).

    Supports two formats:
    1. Simple TB: [code, name, (opening), debit, credit]
    2. HandiLedger Comparative TB: header rows, then
       [code, name, CY_dr, CY_cr, PY_dr, PY_cr]
       with section header rows (Income, Expenses, etc.) mixed in.
    """
    import re as _re

    wb = openpyxl.load_workbook(file, read_only=True, data_only=True)
    ws = wb.active
    all_rows = list(ws.iter_rows(values_only=True))

    # --- Detect HandiLedger comparative format ---
    # Look for a row in the first 6 rows that contains year headers like
    # '2025', '2025', '2024', '2024' or '$ Dr', '$ Cr', '$ Dr', '$ Cr'
    is_comparative = False
    data_start_row = 1  # 0-indexed; default skip first row
    section_names = {
        'income', 'expenses', 'cost of sales', 'current assets',
        'non-current assets', 'non current assets', 'fixed assets',
        'current liabilities', 'non-current liabilities',
        'non current liabilities', 'equity', 'capital',
    }

    for idx, row in enumerate(all_rows[:8]):
        row_strs = [str(c).strip().lower() if c else '' for c in row]
        # Check for year header row like ['', '', '2025', '2025', '2024', '2024']
        year_pattern = sum(1 for s in row_strs if _re.match(r'^20\d{2}$', s))
        if year_pattern >= 2:
            is_comparative = True
            continue
        # Check for column header row like ['', '', '$ Dr', '$ Cr', '$ Dr', '$ Cr']
        if '$ dr' in row_strs and '$ cr' in row_strs:
            is_comparative = True
            data_start_row = idx + 1
            continue
        # Check for 'Comparative Trial Balance' in first cell
        if row_strs[0] and 'comparative' in row_strs[0]:
            is_comparative = True
            continue

    raw_lines = []

    if is_comparative:
        # --- HandiLedger Comparative format ---
        # Columns: [code, name, CY_dr, CY_cr, PY_dr, PY_cr]
        for idx in range(data_start_row, len(all_rows)):
            row = all_rows[idx]
            if not row:
                continue

            col0 = str(row[0]).strip() if row[0] is not None else ''
            col1 = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ''

            # Skip blank rows
            if not col0 and not col1:
                continue

            # Skip section header rows (code is empty, name is a section label)
            if not col0 and col1.lower() in section_names:
                continue

            # Skip totals/summary rows
            if col1.lower().startswith('total') or col1.lower().startswith('net '):
                continue

            # Skip non-account rows (e.g. ABN, title rows that slipped through)
            if not col0:
                continue

            # Validate that col0 looks like an account code (digits, possibly with dots)
            if not _re.match(r'^[\d]', col0):
                continue

            try:
                account_code = col0
                # Strip leading zero from revenue accounts (HandiLedger 0575 -> 575)
                if account_code.startswith('0') and account_code.replace('.', '').isdigit() and len(account_code.split('.')[0]) > 1:
                    parts = account_code.split('.', 1)
                    parts[0] = parts[0].lstrip('0') or '0'
                    account_code = '.'.join(parts)
                account_name = _resolve_account_name(fy.entity, account_code, col1)
                cy_dr = float(row[2]) if len(row) > 2 and row[2] is not None else 0.0
                cy_cr = float(row[3]) if len(row) > 3 and row[3] is not None else 0.0
                py_dr = float(row[4]) if len(row) > 4 and row[4] is not None else 0.0
                py_cr = float(row[5]) if len(row) > 5 and row[5] is not None else 0.0
                raw_lines.append({
                    "account_code": account_code,
                    "account_name": account_name,
                    "opening_balance": "0",
                    "debit": str(cy_dr),
                    "credit": str(cy_cr),
                    "prior_debit": str(py_dr),
                    "prior_credit": str(py_cr),
                })
            except (ValueError, IndexError, TypeError):
                continue
    else:
        # --- Simple TB format (original logic) ---
        for idx in range(1, len(all_rows)):  # skip header row
            row = all_rows[idx]
            if not row or not row[0]:
                continue
            try:
                account_code = str(row[0]).strip()
                # Strip leading zero from revenue accounts (HandiLedger 0575 -> 575)
                if account_code.startswith('0') and account_code.replace('.', '').isdigit() and len(account_code.split('.')[0]) > 1:
                    parts = account_code.split('.', 1)
                    parts[0] = parts[0].lstrip('0') or '0'
                    account_code = '.'.join(parts)
                raw_name = str(row[1]).strip() if len(row) > 1 and row[1] else ""
                account_name = _resolve_account_name(fy.entity, account_code, raw_name)
                if len(row) >= 5 and row[4] is not None:
                    opening_balance = str(row[2] or 0)
                    debit = str(row[3] or 0)
                    credit = str(row[4] or 0)
                else:
                    opening_balance = "0"
                    debit = str(row[2] or 0)
                    credit = str(row[3] or 0)
                raw_lines.append({
                    "account_code": account_code,
                    "account_name": account_name,
                    "opening_balance": opening_balance,
                    "debit": debit,
                    "credit": credit,
                })
            except (ValueError, IndexError):
                continue

    wb.close()
    return raw_lines


def _apply_tb_learned_mappings(entity, raw_lines):
    """
    Apply learned mappings and entity COA code matching to raw TB lines.
    Returns a list of staged line dicts ready for the review wizard.
    """
    existing_mappings = {
        cam.client_account_code: cam
        for cam in ClientAccountMapping.objects.filter(entity=entity)
        .select_related("mapped_line_item")
    }
    entity_coa = {
        ea.account_code.lower(): ea
        for ea in EntityChartOfAccount.objects.filter(entity=entity)
        .select_related("maps_to")
    }
    staged = []
    for line in raw_lines:
        code = line["account_code"]
        cam = existing_mappings.get(code)
        staged_line = {
            "account_code": code,
            "account_name": line["account_name"],
            "opening_balance": line["opening_balance"],
            "debit": line["debit"],
            "credit": line["credit"],
            "prior_debit": line.get("prior_debit", "0"),
            "prior_credit": line.get("prior_credit", "0"),
            "mapped_id": "",
            "mapped_label": "",
            "confidence": "new",
            "entity_acct_code": "",
            "entity_acct_name": "",
        }
        # Learned mappings and COA auto-suggestions are intentionally not applied here.
        # The accountant must map every account manually in the wizard.

        # Still populate entity_acct_code/name for display purposes (account identification only)
        ea = entity_coa.get(code.lower())
        if ea:
            staged_line["entity_acct_code"] = ea.account_code
            staged_line["entity_acct_name"] = ea.account_name
            # Update account_name if it's still just the raw code
            if staged_line["account_name"] == code or not staged_line["account_name"]:
                staged_line["account_name"] = ea.account_name
            if staged_line["confidence"] == "new":
                staged_line["confidence"] = "matched"
        staged.append(staged_line)
    return staged


@login_required
def review_tb_import(request, pk):
    """Review wizard for TB Excel imports — same UX as cloud import wizard."""
    import json as _json
    fy = get_financial_year_for_user(request, pk)
    staged = request.session.get("staged_tb_import")

    if not staged or staged.get("fy_pk") != str(pk):
        messages.error(request, "No staged TB import data found. Please upload again.")
        return redirect("core:trial_balance_import", pk=pk)

    lines = staged["lines"]
    entity = fy.entity

    # Standard accounts for statement line mapping dropdown
    standard_accounts = list(
        AccountMapping.objects.values("id", "standard_code", "line_item_label", "statement_section")
        .order_by("financial_statement", "display_order")
    )
    for sa in standard_accounts:
        sa["id"] = str(sa["id"])

    # Entity accounts for the searchable COA dropdown
    entity_accts = []
    for ea in EntityChartOfAccount.objects.filter(entity=entity).select_related("maps_to").order_by("account_code"):
        entity_accts.append({
            "code": ea.account_code,
            "name": ea.account_name,
            "section": ea.get_section_display(),
            "section_key": ea.section,
            "maps_to_id": str(ea.maps_to.pk) if ea.maps_to else "",
        })

    total = len(lines)
    auto_mapped = sum(1 for l in lines if l.get("mapped_id") or l.get("entity_acct_code"))
    unmapped = total - auto_mapped

    # Balance check — compute totals from staged data
    total_dr = sum(Decimal(str(l.get("debit", "0"))) for l in lines)
    total_cr = sum(Decimal(str(l.get("credit", "0"))) for l in lines)
    balance_diff = abs(total_dr - total_cr)
    TOLERANCE = Decimal("0.02")
    balance_blocked = balance_diff > TOLERANCE
    balance_warning = Decimal("0") < balance_diff <= TOLERANCE

    # Beneficiary officers for the Beneficiary column (trust entities only)
    from django.db import models as _m
    beneficiary_officers = []
    if entity.entity_type == "trust":
        beneficiary_officers = list(
            EntityOfficer.objects.filter(
                entity=entity,
                date_ceased__isnull=True,
            ).filter(
                _m.Q(role__in=["beneficiary", "unit_holder"])
                | _m.Q(roles__contains="beneficiary")
                | _m.Q(roles__contains="unit_holder")
            ).order_by("display_order", "full_name")
        )
        # Attach current beneficiary_officer_id to each line from ClientAccountMapping
        ben_map = dict(
            ClientAccountMapping.objects.filter(
                entity=entity,
                beneficiary_officer__isnull=False,
            ).values_list("client_account_code", "beneficiary_officer_id")
        )
        for line in lines:
            line["current_beneficiary_id"] = str(ben_map.get(line["account_code"], ""))

    context = {
        "fy": fy,
        "lines": lines,
        "standard_accounts_json": _json.dumps(standard_accounts),
        "entity_accounts_json": _json.dumps(entity_accts),
        "total": total,
        "auto_mapped": auto_mapped,
        "unmapped": unmapped,
        "source_name": staged.get("filename", "Excel TB"),
        "balance_total_dr": total_dr,
        "balance_total_cr": total_cr,
        "balance_diff": balance_diff,
        "balance_blocked": balance_blocked,
        "balance_warning": balance_warning,
        "beneficiary_officers": beneficiary_officers,
        "is_trust": entity.entity_type == "trust",
    }
    return render(request, "core/review_tb_import.html", context)


@login_required
@require_POST
def commit_tb_import(request, pk):
    """
    Commit the reviewed TB import. Creates TrialBalanceLine records,
    preserves comparatives, updates ClientAccountMapping (learning system),
    and triggers risk engine.
    """
    fy = get_financial_year_for_user(request, pk)
    staged = request.session.get("staged_tb_import")

    if not staged or staged.get("fy_pk") != str(pk):
        messages.error(request, "No staged TB import data found.")
        return redirect("core:financial_year_detail", pk=pk)

    if fy.is_locked:
        messages.error(request, "Cannot import into a finalised financial year.")
        return redirect("core:financial_year_detail", pk=pk)

    entity = fy.entity
    staged_lines = staged["lines"]

    # Server-side balance validation — block if DR/CR differ by > $0.02
    total_dr = sum(Decimal(str(l.get("debit", "0"))) for l in staged_lines)
    total_cr = sum(Decimal(str(l.get("credit", "0"))) for l in staged_lines)
    balance_diff = abs(total_dr - total_cr)
    TOLERANCE = Decimal("0.02")

    if balance_diff > TOLERANCE:
        messages.error(
            request,
            f"Import blocked \u2014 Trial Balance is out of balance. "
            f"Total debits ${total_dr:,.2f} vs total credits ${total_cr:,.2f} "
            f"\u2014 a difference of ${balance_diff:,.2f}. "
            f"Please correct the source data and re-import.",
        )
        return redirect("core:review_tb_import", pk=pk)

    if balance_diff > 0 and not request.POST.get("rounding_acknowledged"):
        messages.error(
            request,
            f"This TB has a minor rounding difference of ${balance_diff:,.2f}. "
            f"Please tick the rounding acknowledgement checkbox to proceed.",
        )
        return redirect("core:review_tb_import", pk=pk)

    imported = 0
    unmapped = 0
    errors = []

    # Snapshot existing comparative data BEFORE deleting lines
    prior_data = {}
    for line in fy.trial_balance_lines.filter(is_adjustment=False).order_by("account_code"):
        if line.account_code not in prior_data:
            prior_data[line.account_code] = {
                "prior_debit": line.prior_debit,
                "prior_credit": line.prior_credit,
                "prior_closing_balance": line.prior_closing_balance,
                "prior_balance_override": line.prior_balance_override,
                "prior_mapped_line_item": line.prior_mapped_line_item,
                "reclassified": line.reclassified,
                "comparatives_locked": line.comparatives_locked,
                "mapped_line_item": line.mapped_line_item,
                "account_name": line.account_name,
            }

    # Clear existing non-adjustment lines
    fy.trial_balance_lines.filter(is_adjustment=False).delete()

    uploaded_codes = set()
    comp_applied = set()

    for i, line in enumerate(staged_lines):
        mapping_id = request.POST.get(f"mapping_{i}", "").strip()
        entity_acct_code = request.POST.get(f"entity_acct_{i}", "").strip()
        mapped_item = None

        if mapping_id:
            try:
                mapped_item = AccountMapping.objects.get(pk=mapping_id)
            except AccountMapping.DoesNotExist:
                pass

        # If an entity account was assigned, look it up for its maps_to
        if entity_acct_code and not mapped_item:
            try:
                ea = EntityChartOfAccount.objects.select_related("maps_to").get(
                    entity=entity, account_code=entity_acct_code
                )
                if ea.maps_to:
                    mapped_item = ea.maps_to
            except EntityChartOfAccount.DoesNotExist:
                pass

        try:
            opening = Decimal(str(line.get("opening_balance", "0")))
            debit = Decimal(str(line.get("debit", "0")))
            credit = Decimal(str(line.get("credit", "0")))
            closing = opening + debit - credit
            account_code = line["account_code"]

            # Restore comparatives (first occurrence per code only)
            if account_code not in comp_applied:
                comp = prior_data.get(account_code, {})
                comp_applied.add(account_code)
            else:
                comp = {}

            # Prior year data: prefer file-level data (from comparative TB),
            # then fall back to DB snapshot, then zero.
            file_py_dr = Decimal(str(line.get("prior_debit", "0") or "0"))
            file_py_cr = Decimal(str(line.get("prior_credit", "0") or "0"))
            if file_py_dr or file_py_cr:
                # File has prior year data — use it
                py_debit = file_py_dr
                py_credit = file_py_cr
                py_closing = file_py_dr - file_py_cr
            else:
                # Fall back to DB snapshot
                py_debit = comp.get("prior_debit", Decimal("0"))
                py_credit = comp.get("prior_credit", Decimal("0"))
                py_closing = comp.get("prior_closing_balance", Decimal("0"))

            TrialBalanceLine.objects.create(
                financial_year=fy,
                account_code=account_code,
                account_name=_resolve_account_name(entity, line["account_code"], line["account_name"]),
                opening_balance=opening,
                debit=debit,
                credit=credit,
                closing_balance=closing,
                mapped_line_item=mapped_item,
                is_adjustment=False,
                source='tb_import',
                prior_debit=py_debit,
                prior_credit=py_credit,
                prior_closing_balance=py_closing,
                prior_balance_override=comp.get("prior_balance_override", False),
                prior_mapped_line_item=comp.get("prior_mapped_line_item"),
                reclassified=comp.get("reclassified", False),
                comparatives_locked=comp.get("comparatives_locked", False),
            )

            # Update the learning system
            resolved_name = _resolve_account_name(entity, line["account_code"], line["account_name"])
            # Beneficiary officer (trust entities — from the Beneficiary dropdown)
            _ben_officer_id = request.POST.get(f"beneficiary_{i}", "").strip() or None
            _ben_officer = None
            if _ben_officer_id:
                try:
                    _ben_officer = EntityOfficer.objects.get(pk=_ben_officer_id)
                except EntityOfficer.DoesNotExist:
                    _ben_officer = None
            ClientAccountMapping.objects.update_or_create(
                entity=entity,
                client_account_code=account_code,
                defaults={
                    "client_account_name": resolved_name,
                    "mapped_line_item": mapped_item,
                    "beneficiary_officer": _ben_officer,
                },
            )
            # ── Sync EntityChartOfAccount ──────────────────────────────────────────
            # Keep the entity CoA in sync with what's in the TB.
            # If the account already exists in the CoA, update its name only
            # (preserve section, maps_to, tax_code, and other metadata).
            # If it doesn't exist yet, create a minimal entry so it appears
            # in the CoA tab and the journal account picker.
            _section_for_code = _hl_section_for_code(account_code)
            _section_map = {
                'Income': EntityChartOfAccount.StatementSection.REVENUE,
                'Cost of Sales': EntityChartOfAccount.StatementSection.COST_OF_SALES,
                'Expenses': EntityChartOfAccount.StatementSection.EXPENSES,
                'Current Assets': EntityChartOfAccount.StatementSection.ASSETS,
                'Non-Current Assets': EntityChartOfAccount.StatementSection.ASSETS,
                'Current Liabilities': EntityChartOfAccount.StatementSection.LIABILITIES,
                'Non-Current Liabilities': EntityChartOfAccount.StatementSection.LIABILITIES,
                'Equity': EntityChartOfAccount.StatementSection.EQUITY,
            }
            default_section = _section_map.get(_section_for_code, EntityChartOfAccount.StatementSection.SUSPENSE)
            existing_coa = EntityChartOfAccount.objects.filter(
                entity=entity, account_code=account_code
            ).first()
            if existing_coa:
                # Only update the name (don't overwrite section/maps_to/tax_code)
                if existing_coa.account_name != resolved_name:
                    existing_coa.account_name = resolved_name
                    existing_coa.save(update_fields=['account_name', 'updated_at'])
            else:
                # Create a new CoA entry seeded from the TB data
                EntityChartOfAccount.objects.create(
                    entity=entity,
                    account_code=account_code,
                    account_name=resolved_name,
                    section=default_section,
                    maps_to=mapped_item,
                    is_active=True,
                    is_custom=True,
                )
            # ────────────────────────────────────────────────────────────────────
            uploaded_codes.add(account_code)
            imported += 1
            if not mapped_item:
                unmapped += 1

        except Exception as e:
            errors.append(f"Line {i + 1} ({line.get('account_code', '?')}): {str(e)}")

    # Re-create comparative-only lines for accounts not in the upload
    for code, comp in prior_data.items():
        if code in uploaded_codes:
            continue
        if comp["prior_debit"] == 0 and comp["prior_credit"] == 0:
            continue
        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code=code,
            account_name=comp.get("account_name", ""),
            opening_balance=Decimal("0"),
            debit=Decimal("0"),
            credit=Decimal("0"),
            closing_balance=Decimal("0"),
            mapped_line_item=comp.get("mapped_line_item") or comp.get("prior_mapped_line_item"),
            is_adjustment=False,
            source='rollover',
            prior_debit=comp["prior_debit"],
            prior_credit=comp["prior_credit"],
            prior_closing_balance=comp.get("prior_closing_balance", Decimal("0")),
            prior_balance_override=comp.get("prior_balance_override", False),
            prior_mapped_line_item=comp.get("prior_mapped_line_item"),
            reclassified=comp.get("reclassified", False),
            comparatives_locked=comp.get("comparatives_locked", False),
        )

    # Surface any merge warnings that were recorded at staging time
    merge_warnings = staged.get("merge_warnings", [])
    for w in merge_warnings:
        messages.warning(request, w)

    # Clean up session
    request.session.pop("staged_tb_import", None)

    # Log and trigger risk engine
    _log_action(request, "import", f"Imported trial balance via wizard: {imported} lines", fy)
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "tb_import")

    messages.success(
        request,
        f"Imported {imported} lines. "
        f"{unmapped} unmapped accounts need attention."
    )
    if errors:
        for err in errors[:5]:
            messages.warning(request, err)

    return redirect("core:financial_year_detail", pk=pk)


@login_required
def trial_balance_view(request, pk):
    """Preview trial balance with comparative columns, grouped by section."""
    from collections import OrderedDict
    fy = get_financial_year_for_user(request, pk)
    tb_lines = fy.trial_balance_lines.select_related("mapped_line_item").order_by("account_code")
    _coa_lookup = _build_coa_section_lookup(fy.entity)
    SECTION_ORDER = [
        'Revenue', 'Income', 'Cost of Sales', 'Expenses',
        'Current Assets', 'Non-Current Assets',
        'Current Liabilities', 'Non-Current Liabilities',
        'Equity', 'Income Tax',
    ]
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }
    sections = OrderedDict()

    for line in tb_lines:
        # Display Dr/Cr: always use closing_balance (opening + movements)
        cb = line.closing_balance if line.closing_balance else Decimal('0')
        if cb > 0:
            line.display_dr = cb
            line.display_cr = Decimal('0')
        elif cb < 0:
            line.display_dr = Decimal('0')
            line.display_cr = abs(cb)
        else:
            line.display_dr = line.debit if line.debit else Decimal('0')
            line.display_cr = line.credit if line.credit else Decimal('0')
        # HARD RULE: HandiLedger numeric code range is authoritative for section.
        hl_section = _hl_section_for_code(line.account_code)
        if hl_section:
            display_section = hl_section
        elif line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = _coa_lookup.get(line.account_code, 'Unmapped')
        if display_section not in sections:
            sections[display_section] = []
        sections[display_section].append(line)

    # Sort sections by defined order
    ordered_sections = OrderedDict()
    seen = set()
    for s in SECTION_ORDER:
        ds = SECTION_DISPLAY.get(s, s)
        if ds not in seen and ds in sections:
            ordered_sections[ds] = sections[ds]
            seen.add(ds)
    for key in sections:
        if key not in ordered_sections:
            ordered_sections[key] = sections[key]

    # Year labels
    current_year_label = str(fy.year_label)
    year_digits = ''.join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        prior_year_label = f"FY{int(year_digits) - 1}" if fy.year_label.startswith('FY') else str(int(year_digits) - 1)
    elif fy.prior_year:
        prior_year_label = str(fy.prior_year.year_label)
    else:
        prior_year_label = 'Prior'

    # ---- Aggregate multiple sub-entries per account_code ----
    aggregated_sections = OrderedDict()
    for section_name, lines_list in ordered_sections.items():
        code_groups = OrderedDict()
        for line in lines_list:
            code_groups.setdefault(line.account_code, []).append(line)
        agg_lines = []
        for code, group in code_groups.items():
            if len(group) == 1:
                group[0].account_name = _resolve_account_name(fy.entity, group[0].account_code, group[0].account_name)
                group[0].sub_entries = []
                group[0].is_aggregated = False
                agg_lines.append(group[0])
            else:
                first = group[0]
                raw_dr = sum(l.display_dr or Decimal('0') for l in group)
                raw_cr = sum(l.display_cr or Decimal('0') for l in group)
                agg_prior_dr = sum(l.prior_debit or Decimal('0') for l in group)
                agg_prior_cr = sum(l.prior_credit or Decimal('0') for l in group)
                # Net the debits and credits so adjustments reduce the
                # original balance instead of showing a separate column
                net = raw_dr - raw_cr
                if net >= 0:
                    agg_dr = net
                    agg_cr = Decimal('0')
                else:
                    agg_dr = Decimal('0')
                    agg_cr = abs(net)
                class AggregatedLine:
                    pass
                agg = AggregatedLine()
                agg.account_code = code
                unique_names = list(dict.fromkeys(l.account_name for l in group if l.account_name))
                resolved_name = unique_names[0] if unique_names else code
                agg.account_name = _resolve_account_name(fy.entity, code, resolved_name)
                agg.display_dr = agg_dr
                agg.display_cr = agg_cr
                agg.prior_debit = agg_prior_dr
                agg.prior_credit = agg_prior_cr
                agg.mapped_line_item = first.mapped_line_item
                agg.is_adjustment = any(l.is_adjustment for l in group)
                agg.sub_entries = group
                agg.is_aggregated = True
                agg.sub_count = len(group)
                agg_lines.append(agg)
        aggregated_sections[section_name] = agg_lines

    # Recompute grand totals and net profit from AGGREGATED values so they
    # match the visible rows (adjustments netted per account code).
    pl_sections = {'Income', 'Cost of Sales', 'Expenses'}
    grand_total_dr = Decimal('0')
    grand_total_cr = Decimal('0')
    grand_total_prior_dr = Decimal('0')
    grand_total_prior_cr = Decimal('0')
    pl_dr = Decimal('0')
    pl_cr = Decimal('0')
    pl_prior_dr = Decimal('0')
    pl_prior_cr = Decimal('0')
    for section_name, agg_lines in aggregated_sections.items():
        is_pl = section_name in pl_sections
        for line in agg_lines:
            dr = line.display_dr or Decimal('0')
            cr = line.display_cr or Decimal('0')
            pdr = line.prior_debit or Decimal('0')
            pcr = line.prior_credit or Decimal('0')
            grand_total_dr += dr
            grand_total_cr += cr
            grand_total_prior_dr += pdr
            grand_total_prior_cr += pcr
            if is_pl:
                pl_dr += dr
                pl_cr += cr
                pl_prior_dr += pdr
                pl_prior_cr += pcr
    net_profit = pl_cr - pl_dr
    prior_net_profit = pl_prior_cr - pl_prior_dr

    return render(request, "core/trial_balance_view.html", {
        "fy": fy,
        "sections": aggregated_sections,
        "current_year_label": current_year_label,
        "prior_year_label": prior_year_label,
        "grand_total_dr": grand_total_dr,
        "grand_total_cr": grand_total_cr,
        "grand_total_prior_dr": grand_total_prior_dr,
        "grand_total_prior_cr": grand_total_prior_cr,
        "net_profit": net_profit,
        "net_profit_abs": abs(net_profit),
        "prior_net_profit": prior_net_profit,
        "prior_net_profit_abs": abs(prior_net_profit),
    })


# ---------------------------------------------------------------------------
# Account Code Breakdown (drill-down for multi-entry accounts)
# ---------------------------------------------------------------------------
@login_required
def account_code_breakdown(request, pk, account_code):
    """Show all individual TB lines for a given account code within a financial year."""
    fy = get_financial_year_for_user(request, pk)
    lines = fy.trial_balance_lines.filter(
        account_code=account_code
    ).select_related('mapped_line_item').order_by('account_name')

    if not lines.exists():
        messages.error(request, f"No trial balance entries found for account code {account_code}.")
        return redirect('core:financial_year_detail', pk=pk)

    # Compute display Dr/Cr for each line — always use closing_balance
    total_dr = Decimal('0')
    total_cr = Decimal('0')
    for line in lines:
        cb = line.closing_balance if line.closing_balance else Decimal('0')
        if cb > 0:
            line.display_dr = cb
            line.display_cr = Decimal('0')
        elif cb < 0:
            line.display_dr = Decimal('0')
            line.display_cr = abs(cb)
        else:
            line.display_dr = line.debit if line.debit else Decimal('0')
            line.display_cr = line.credit if line.credit else Decimal('0')
        total_dr += line.display_dr or Decimal('0')
        total_cr += line.display_cr or Decimal('0')

    # Year label
    current_year = str(fy.year_label)

    # Use the first line's mapped_line_item label as the account heading
    first_line = lines[0]
    mapped_label = first_line.mapped_line_item.line_item_label if first_line.mapped_line_item else 'Unmapped'

    # Get available accounts for reallocation dropdown
    # Use EntityChartOfAccount for the entity's own COA
    from .models import EntityChartOfAccount
    available_accounts = EntityChartOfAccount.objects.filter(
        entity=fy.entity,
        is_active=True,
    ).select_related('maps_to').order_by('account_code')

    # Fetch individual bank statement transactions coded to this account
    from review.models import PendingTransaction
    bank_txns = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        job__financial_year=fy,
        confirmed_code=account_code,
        is_confirmed=True,
    ).select_related('job').order_by('date', 'description')

    bank_txn_total = Decimal('0')
    for bt in bank_txns:
        bank_txn_total += bt.amount or Decimal('0')

    # For bank/balance-sheet accounts, transactions flow through via the
    # contra entry system — confirmed_code points to the expense/income
    # account, not the bank account.  Detect this via BankAccountMapping
    # and query all confirmed+posted transactions for matching jobs.
    is_bank_account = False
    bank_acct_txns = PendingTransaction.objects.none()
    bank_acct_txn_total = Decimal('0')
    bank_mappings = BankAccountMapping.objects.filter(
        entity=fy.entity, tb_account_code=account_code,
    )
    if bank_mappings.exists():
        is_bank_account = True
        from django.db.models import Q
        mapping_filter = Q()
        for bm in bank_mappings:
            q = Q()
            if bm.bsb and bm.account_number:
                q = Q(job__bsb=bm.bsb, job__account_number=bm.account_number)
            elif bm.bank_account_name:
                q = Q(job__bank_account_name=bm.bank_account_name)
            else:
                # catch-all mapping — no bsb/account_number/name, match all jobs for this entity
                q = Q(job__entity=fy.entity)
            mapping_filter |= q
        if mapping_filter:
            bank_acct_txns = PendingTransaction.objects.filter(
                mapping_filter,
                job__entity=fy.entity,
                job__financial_year=fy,
                is_confirmed=True,
                posted_to_tb=True,
            ).select_related('job').order_by('date', 'description')
            for bt in bank_acct_txns:
                bank_acct_txn_total += bt.amount or Decimal('0')

    # For the GST payable control account (3380), the individual bank
    # transactions are coded to expense/income accounts — not to 3380
    # itself.  The GST component is split out automatically.  So we
    # query all confirmed transactions that had a GST component to show
    # the individual GST movements.
    gst_txns = []
    gst_txn_total = Decimal('0')
    if account_code == '3380':
        from django.db.models import Q
        gst_txns = list(
            PendingTransaction.objects.filter(
                job__entity=fy.entity,
                is_confirmed=True,
            ).filter(
                Q(confirmed_gst_amount__gt=Decimal('0'))
            ).order_by('date', 'description')
        )
        for gt in gst_txns:
            gst_txn_total += gt.confirmed_gst_amount or Decimal('0')

    # Fetch journal entry movements for this account code
    journal_lines = JournalLine.objects.filter(
        journal__financial_year=fy,
        journal__status='posted',
        account_code=account_code,
    ).select_related('journal').order_by('journal__journal_date', 'journal__reference_number', 'line_number')

    journal_total_dr = Decimal('0')
    journal_total_cr = Decimal('0')
    for jl in journal_lines:
        journal_total_dr += jl.debit or Decimal('0')
        journal_total_cr += jl.credit or Decimal('0')

    # Determine if this is an asset account (show depreciation controls)
    is_asset_account = False
    entity_acct = EntityChartOfAccount.objects.filter(
        entity=fy.entity, account_code=account_code, is_active=True,
    ).first()
    if entity_acct and entity_acct.section == EntityChartOfAccount.StatementSection.ASSETS:
        is_asset_account = True

    # Build a set of transaction IDs already linked to depreciation assets
    txn_in_depreciation = set()
    if is_asset_account and bank_txns:
        bank_txn_ids = [bt.pk for bt in bank_txns]
        txn_in_depreciation = set(
            DepreciationAsset.objects.filter(
                source_transaction_id__in=bank_txn_ids,
            ).values_list('source_transaction_id', flat=True)
        )
    # Annotate each bank txn with a flag
    for bt in bank_txns:
        bt.in_depreciation = bt.pk in txn_in_depreciation

    return render(request, 'core/account_code_breakdown.html', {
        'fy': fy,
        'account_code': account_code,
        'account_name': first_line.account_name,
        'mapped_label': mapped_label,
        'lines': lines,
        'total_dr': total_dr,
        'total_cr': total_cr,
        'current_year': current_year,
        'entry_count': lines.count(),
        'available_accounts': available_accounts,
        'bank_txns': bank_txns,
        'bank_txn_count': bank_txns.count(),
        'bank_txn_total': bank_txn_total,
        'gst_txns': gst_txns,
        'gst_txn_count': len(gst_txns),
        'gst_txn_total': gst_txn_total,
        'journal_lines': journal_lines,
        'journal_line_count': journal_lines.count(),
        'journal_total_dr': journal_total_dr,
        'journal_total_cr': journal_total_cr,
        'is_asset_account': is_asset_account,
        'is_bank_account': is_bank_account,
        'bank_acct_txns': bank_acct_txns,
        'bank_acct_txn_count': bank_acct_txns.count() if is_bank_account else 0,
        'bank_acct_txn_total': bank_acct_txn_total,
    })


@login_required
def tb_line_reallocate(request, pk):
    """Reallocate a single trial balance line to a different account code/mapping."""
    if request.method != 'POST':
        return HttpResponseNotAllowed(['POST'])

    from .models import EntityChartOfAccount
    line = get_object_or_404(TrialBalanceLine, pk=pk)
    fy = line.financial_year

    # Security: ensure user has access
    get_financial_year_for_user(request, fy.pk)

    new_account_code = request.POST.get('new_account_code', '').strip()
    if not new_account_code:
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'message': 'Please select an account to reallocate to.'}, status=400)
        messages.error(request, "Please select an account to reallocate to.")
        return redirect('core:account_code_breakdown', pk=fy.pk, account_code=line.account_code)

    old_code = line.account_code
    old_name = line.account_name

    # Look up the target account from EntityChartOfAccount (entity-specific COA)
    coa_entry = EntityChartOfAccount.objects.filter(
        entity=fy.entity,
        account_code=new_account_code,
        is_active=True,
    ).select_related('maps_to').first()

    if not coa_entry:
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'message': f'Account code {new_account_code} not found in chart of accounts.'}, status=400)
        messages.error(request, f"Account code {new_account_code} not found in chart of accounts.")
        return redirect('core:account_code_breakdown', pk=fy.pk, account_code=old_code)

    # Update the TB line
    line.account_code = coa_entry.account_code
    line.account_name = coa_entry.account_name
    line.mapped_line_item = coa_entry.maps_to
    line.save(update_fields=['account_code', 'account_name', 'mapped_line_item'])

    # Also update the client account mapping
    ClientAccountMapping.objects.update_or_create(
        entity=fy.entity,
        client_account_code=coa_entry.account_code,
        defaults={
            'client_account_name': coa_entry.account_name,
            'mapped_line_item': coa_entry.maps_to,
        },
    )

    _log_action(
        request, 'adjustment',
        f"Reallocated TB line '{old_name}' from {old_code} to {coa_entry.account_code} ({coa_entry.account_name})",
        fy,
    )
    messages.success(
        request,
        f"Reallocated '{old_name}' from {old_code} to {coa_entry.account_code} — {coa_entry.account_name}."
    )

    # Handle AJAX requests (live reallocation from account breakdown page)
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        remaining = fy.trial_balance_lines.filter(account_code=old_code).count()
        return JsonResponse({
            'success': True,
            'message': f"Reallocated '{old_name}' from {old_code} to {coa_entry.account_code} — {coa_entry.account_name}.",
            'remaining_count': remaining,
            'redirect_url': reverse('core:financial_year_detail', kwargs={'pk': fy.pk}) if remaining == 0 else None,
        })

    # Redirect back to the original account code breakdown (which may now have fewer lines)
    remaining = fy.trial_balance_lines.filter(account_code=old_code).count()
    if remaining > 0:
        return redirect('core:account_code_breakdown', pk=fy.pk, account_code=old_code)
    else:
        return redirect('core:financial_year_detail', pk=fy.pk)


# ---------------------------------------------------------------------------
# Account Mapping
# ---------------------------------------------------------------------------
@login_required
def account_mapping_list(request):
    mappings = AccountMapping.objects.all()
    return render(request, "core/account_mapping_list.html", {"mappings": mappings})


@login_required
def account_mapping_create(request):
    if not request.user.is_admin:
        messages.error(request, "Only administrators can manage standard mappings.")
        return redirect("core:account_mapping_list")

    if request.method == "POST":
        form = AccountMappingForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, "Account mapping created.")
            return redirect("core:account_mapping_list")
    else:
        form = AccountMappingForm()
    return render(request, "core/account_mapping_form.html", {"form": form, "title": "Create Standard Mapping"})


@login_required
def map_client_accounts(request, pk):
    """Map unmapped client accounts to standard line items for a financial year."""
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    # Ensure every unmapped TB line has a ClientAccountMapping entry
    # (bulk journal uploads may create TB lines without a corresponding CAM)
    unmapped_tb_codes = (
        fy.trial_balance_lines
        .filter(mapped_line_item__isnull=True)
        .values_list('account_code', flat=True)
        .distinct()
    )
    existing_cam_codes = set(
        ClientAccountMapping.objects.filter(entity=entity)
        .values_list('client_account_code', flat=True)
    )
    for code in unmapped_tb_codes:
        if code not in existing_cam_codes:
            # Get the account name from the TB line
            tb_line = fy.trial_balance_lines.filter(
                account_code=code, mapped_line_item__isnull=True
            ).first()
            ClientAccountMapping.objects.create(
                entity=entity,
                client_account_code=code,
                client_account_name=tb_line.account_name if tb_line else code,
            )

    unmapped = ClientAccountMapping.objects.filter(
        entity=entity, mapped_line_item__isnull=True
    )
    standard_mappings = AccountMapping.objects.all()

    if request.method == "POST":
        mapped_count = 0
        for cam in unmapped:
            mapping_id = request.POST.get(f"mapping_{cam.pk}")
            if mapping_id:
                try:
                    cam.mapped_line_item = AccountMapping.objects.get(pk=mapping_id)
                    cam.save()
                    # Update trial balance lines with this mapping
                    TrialBalanceLine.objects.filter(
                        financial_year=fy,
                        account_code=cam.client_account_code,
                    ).update(mapped_line_item=cam.mapped_line_item)
                    mapped_count += 1
                except AccountMapping.DoesNotExist:
                    pass

        _log_action(request, "mapping_change", f"Mapped {mapped_count} accounts", fy)
        messages.success(request, f"Mapped {mapped_count} accounts.")
        return redirect("core:financial_year_detail", pk=pk)

    context = {
        "fy": fy,
        "unmapped": unmapped,
        "standard_mappings": standard_mappings,
    }
    return render(request, "core/map_client_accounts.html", context)


# ---------------------------------------------------------------------------
# Journal Entries (Enhanced)
# ---------------------------------------------------------------------------
@login_required
def adjustment_list(request, pk):
    fy = get_financial_year_for_user(request, pk)
    adjustments = fy.adjusting_journals.prefetch_related("lines").all()
    return render(request, "core/adjustment_list.html", {"fy": fy, "adjustments": adjustments})


@login_required
def adjustment_create(request, pk):
    fy = get_financial_year_for_user(request, pk)

    if fy.is_locked:
        messages.error(request, "Cannot add journals to a finalised year.")
        return redirect("core:financial_year_detail", pk=pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    # Get available accounts for the account picker.
    # Priority: entity-specific accounts from the trial balance or client
    # account mappings.  Fall back to the master Chart of Accounts only
    # when neither source has data (e.g. brand-new entity with no TB yet).
    entity_accounts = list(
        ClientAccountMapping.objects.filter(entity=fy.entity)
        .order_by("client_account_code")
        .values("client_account_code", "client_account_name")
    )

    if not entity_accounts:
        # Try unique accounts from the trial balance for this financial year
        from django.db.models import F
        tb_accounts = (
            TrialBalanceLine.objects.filter(financial_year=fy)
            .exclude(account_code="")
            .values("account_code", "account_name")
            .distinct()
            .order_by("account_code")
        )
        entity_accounts = [
            {"client_account_code": a["account_code"], "client_account_name": a["account_name"]}
            for a in tb_accounts
        ]
    # Always merge with the entity's Chart of Accounts so that accounts
    # which exist in the CoA but have no TB line yet are still selectable.
    entity_type = fy.entity.entity_type
    coa_accounts = (
        EntityChartOfAccount.objects.filter(entity=fy.entity, is_active=True)
        .order_by("account_code")
        .values("account_code", "account_name")
    )
    if not coa_accounts.exists():
        coa_accounts = (
            ChartOfAccount.objects.filter(entity_type=entity_type, is_active=True)
            .order_by("account_code")
            .values("account_code", "account_name")
        )
    # Build a merged, deduplicated list keyed by account_code.
    # TB lines take priority (they carry the entity-specific name).
    merged = {}
    for a in coa_accounts:
        merged[a["account_code"]] = {"client_account_code": a["account_code"], "client_account_name": a["account_name"]}
    for a in entity_accounts:
        merged[a["client_account_code"]] = a  # TB name wins
    accounts = sorted(merged.values(), key=lambda x: x["client_account_code"])

    if request.method == "POST":
        form = AdjustingJournalForm(request.POST)
        formset = JournalLineFormSet(request.POST)
        if form.is_valid() and formset.is_valid():
            journal = form.save(commit=False)
            journal.financial_year = fy
            journal.created_by = request.user
            journal.save()  # This auto-generates reference_number

            formset.instance = journal
            lines = formset.save()

            # Set line_number to preserve the order lines were entered
            for i, line in enumerate(lines, start=1):
                line.line_number = i
                line.save(update_fields=["line_number"])

            # Validate debits = credits
            total_dr = sum(l.debit for l in lines)
            total_cr = sum(l.credit for l in lines)
            if total_dr != total_cr:
                journal.delete()
                messages.error(request, f"Journal does not balance: Dr ${total_dr:,.2f} \u2260 Cr ${total_cr:,.2f}")
                return render(request, "core/adjustment_form.html", {
                    "form": form, "formset": formset, "fy": fy, "accounts": accounts
                })

            # Update cached totals
            journal.total_debit = total_dr
            journal.total_credit = total_cr
            journal.save(update_fields=["total_debit", "total_credit"])

            # Auto-post to Trial Balance and mark as POSTED in one atomic
            # block — if TB posting fails, the journal stays DRAFT and no
            # orphan TB lines are left behind.
            with db_transaction.atomic():
                _post_journal_to_tb(journal, fy)

                journal.status = AdjustingJournal.JournalStatus.POSTED
                journal.posted_by = request.user
                journal.posted_at = timezone.now()
                journal.save(update_fields=["status", "posted_by", "posted_at"])

            _log_action(request, "adjustment", f"Created and posted journal {journal.reference_number}: {journal.description}", journal)
            # Auto-trigger risk engine after journal post
            from core.signals import trigger_risk_recalc
            trigger_risk_recalc(fy, "journal_posted")
            messages.success(request, f"Journal {journal.reference_number} successfully posted and pushed to Trial Balance.")
            from django.urls import reverse
            return redirect(reverse("core:financial_year_detail", args=[pk]) + "?tab=journals")
    else:
        form = AdjustingJournalForm(initial={"journal_date": fy.end_date})
        formset = JournalLineFormSet()

    return render(request, "core/adjustment_form.html", {
        "form": form, "formset": formset, "fy": fy, "accounts": accounts
    })


@login_required
def journal_detail(request, pk):
    """View a single journal entry with all its lines and audit info."""
    journal = get_object_or_404(
        AdjustingJournal.objects.select_related(
            "financial_year", "financial_year__entity",
            "created_by", "posted_by"
        ).prefetch_related("lines"),
        pk=pk,
    )
    fy = journal.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check
    entity = fy.entity
    return render(request, "core/journal_detail.html", {
        "journal": journal, "fy": fy, "entity": entity,
    })


@login_required
def journal_post(request, pk):
    """Post a draft journal — creates adjustment TB lines."""
    journal = get_object_or_404(AdjustingJournal, pk=pk)
    fy = journal.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check

    if journal.status != AdjustingJournal.JournalStatus.DRAFT:
        messages.error(request, "Only draft journals can be posted.")
        return redirect("core:journal_detail", pk=pk)

    if not journal.is_balanced:
        messages.error(request, "Journal does not balance. Cannot post.")
        return redirect("core:journal_detail", pk=pk)

    if fy.is_locked:
        messages.error(request, "Cannot post to a finalised year.")
        return redirect("core:journal_detail", pk=pk)

    # Apply journal lines to Trial Balance and mark as POSTED atomically —
    # if TB posting fails, the journal stays DRAFT with no orphan TB lines.
    with db_transaction.atomic():
        _post_journal_to_tb(journal, fy)

        journal.status = AdjustingJournal.JournalStatus.POSTED
        journal.posted_by = request.user
        journal.posted_at = timezone.now()
        journal.save(update_fields=["status", "posted_by", "posted_at"])

    _log_action(request, "adjustment", f"Posted journal {journal.reference_number}", journal)
    # Auto-trigger risk engine after journal post
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "journal_posted")
    messages.success(request, f"Journal {journal.reference_number} has been posted.")
    return redirect("core:financial_year_detail", pk=fy.pk)


@login_required
def calculate_tax_journal(request, pk):
    """Calculate and post an income tax journal for company entities."""
    import math as _math

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    if entity.entity_type != "company":
        messages.error(request, "Tax journal calculation is only available for company entities.")
        return redirect("core:financial_year_detail", pk=fy.pk)

    if not fy.is_locked:
        messages.error(request, "Financial year must be finalised before calculating tax.")
        return redirect("core:financial_year_detail", pk=fy.pk)

    if entity.is_base_rate_entity is None:
        messages.error(request, "Please set the Base Rate Entity flag on the entity before calculating tax.")
        return redirect("core:financial_year_detail", pk=fy.pk)

    # Check for existing tax journal
    existing = AdjustingJournal.objects.filter(
        financial_year=fy, description__icontains="Income tax",
    ).exists()
    if existing:
        messages.warning(request, "An income tax journal already exists for this financial year.")
        return redirect("core:financial_year_detail", pk=fy.pk)

    # Calculate net profit from TB — aggregate all lines by P&L sections
    # statement_section values "Revenue" and "Income" both map to P&L
    pl_sections = {"Income", "Revenue", "Cost of Sales", "Expenses"}
    all_tb = TrialBalanceLine.objects.filter(
        financial_year=fy,
    ).select_related("mapped_line_item")
    pl_dr = Decimal("0")
    pl_cr = Decimal("0")
    for line in all_tb:
        mapping = line.mapped_line_item
        if mapping and mapping.statement_section in pl_sections:
            pl_dr += line.debit or Decimal("0")
            pl_cr += line.credit or Decimal("0")
    net_profit = pl_cr - pl_dr

    if net_profit <= 0:
        messages.info(request, "No tax payable — entity is in a loss position.")
        return redirect("core:financial_year_detail", pk=fy.pk)

    # Determine tax rate
    if entity.is_base_rate_entity:
        tax_rate = Decimal("0.25")
        rate_label = "25% (Base Rate Entity)"
    else:
        tax_rate = Decimal("0.30")
        rate_label = "30% (Standard Rate)"

    tax_amount = Decimal(_math.ceil(net_profit * tax_rate))

    # Ensure account codes 4110 and 3325 exist in entity CoA
    tax_accounts = [
        ("4110", "Income tax on profit", "expenses"),
        ("3325", "Taxation", "liabilities"),
    ]
    for code, name, section in tax_accounts:
        EntityChartOfAccount.objects.get_or_create(
            entity=entity,
            account_code=code,
            defaults={
                "account_name": name,
                "section": section,
                "is_custom": True,
            },
        )

    # Create and post the journal
    description = f"Income tax on profit for year ended {fy.end_date.strftime('%d %B %Y')}"
    with db_transaction.atomic():
        journal = AdjustingJournal.objects.create(
            financial_year=fy,
            journal_type="tax",
            journal_date=fy.end_date,
            description=description,
            created_by=request.user,
            status=AdjustingJournal.JournalStatus.POSTED,
            posted_by=request.user,
            posted_at=timezone.now(),
            total_debit=tax_amount,
            total_credit=tax_amount,
        )
        JournalLine.objects.create(
            journal=journal,
            account_code="4110",
            account_name="Income tax on profit",
            debit=tax_amount,
            credit=Decimal("0"),
            line_number=1,
        )
        JournalLine.objects.create(
            journal=journal,
            account_code="3325",
            account_name="Taxation",
            debit=Decimal("0"),
            credit=tax_amount,
            line_number=2,
        )
        _post_journal_to_tb(journal, fy)

    _log_action(request, "adjustment", f"Posted tax journal {journal.reference_number} — ${tax_amount:,.0f} at {rate_label}", journal)
    messages.success(request, f"Tax journal posted — ${tax_amount:,.0f} at {rate_label}")
    return redirect("core:financial_year_detail", pk=fy.pk)


# ---------------------------------------------------------------------------
# Auto Tax Provision — status check + post
# ---------------------------------------------------------------------------

def _resolve_tax_accounts(entity):
    """Resolve income tax expense and provision account codes for an entity.

    Uses the three-tier lookup:
      1. ClientAccountMapping → AccountMapping with standard_code IS-TAX-001 / BS-CL-003
      2. Fallback to well-known EntityChartOfAccount codes (4110, 3325)
      3. Return defaults if nothing found

    Returns (expense_code, expense_name, provision_code, provision_name).
    """
    expense_code, expense_name = "4110", "Income tax on profit"
    provision_code, provision_name = "3325", "Taxation"

    # Try IS-TAX-001 mapping
    cam_expense = (
        ClientAccountMapping.objects
        .filter(entity=entity, mapped_line_item__standard_code="IS-TAX-001")
        .select_related("mapped_line_item")
        .first()
    )
    if cam_expense:
        expense_code = cam_expense.client_account_code
        expense_name = cam_expense.client_account_name or expense_name

    # Try BS-CL-003 mapping
    cam_provision = (
        ClientAccountMapping.objects
        .filter(entity=entity, mapped_line_item__standard_code="BS-CL-003")
        .select_related("mapped_line_item")
        .first()
    )
    if cam_provision:
        provision_code = cam_provision.client_account_code
        provision_name = cam_provision.client_account_name or provision_name

    # If no ClientAccountMapping found, check EntityChartOfAccount for known codes
    if not cam_expense:
        ecoa = EntityChartOfAccount.objects.filter(entity=entity, account_code="4110").first()
        if ecoa:
            expense_name = ecoa.account_name or expense_name
    if not cam_provision:
        ecoa = EntityChartOfAccount.objects.filter(entity=entity, account_code="3325").first()
        if ecoa:
            provision_name = ecoa.account_name or provision_name

    return expense_code, expense_name, provision_code, provision_name


def _calculate_net_profit(fy):
    """Calculate net profit from the trial balance P&L sections.

    Mirrors the logic in calculate_tax_journal exactly.
    Returns Decimal.
    """
    pl_sections = {"Income", "Revenue", "Cost of Sales", "Expenses"}
    all_tb = TrialBalanceLine.objects.filter(
        financial_year=fy,
    ).select_related("mapped_line_item")
    pl_dr = Decimal("0")
    pl_cr = Decimal("0")
    for line in all_tb:
        mapping = line.mapped_line_item
        if mapping and mapping.statement_section in pl_sections:
            pl_dr += line.debit or Decimal("0")
            pl_cr += line.credit or Decimal("0")
    return pl_cr - pl_dr


def _calculate_existing_provision(fy, provision_code):
    """Calculate the existing tax provision balance from the trial balance.

    For a liability account, a credit balance is normal. Returns the net
    credit balance (positive means a credit-side provision exists).
    """
    agg = TrialBalanceLine.objects.filter(
        financial_year=fy,
        account_code=provision_code,
    ).aggregate(total_dr=Sum("debit"), total_cr=Sum("credit"))
    total_dr = agg["total_dr"] or Decimal("0")
    total_cr = agg["total_cr"] or Decimal("0")
    return total_cr - total_dr  # positive = credit balance (normal for liability)


@login_required
def tax_provision_status(request, pk):
    """Return JSON with tax provision eligibility and calculation details."""
    import math as _math

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    # Entity must be a company
    if entity.entity_type != "company":
        return JsonResponse({
            "eligible": False,
            "reason": "Tax provision is only available for company entities.",
        })

    # Financial year must not be locked
    if fy.is_locked:
        return JsonResponse({
            "eligible": False,
            "reason": "Financial year is locked.",
        })

    # Base rate entity flag must be set
    if entity.is_base_rate_entity is None:
        return JsonResponse({
            "eligible": False,
            "reason": "Please set the Base Rate Entity flag on the entity before calculating tax provision.",
        })

    # Check for existing tax provision journal
    existing_provision_journal = AdjustingJournal.objects.filter(
        financial_year=fy, journal_type="tax_provision",
    ).first()
    if existing_provision_journal:
        return JsonResponse({
            "eligible": False,
            "reason": "A tax provision journal already exists for this financial year.",
            "existing_journal_ref": existing_provision_journal.reference_number,
            "existing_journal_amount": str(existing_provision_journal.total_debit),
        })

    # Resolve accounts
    expense_code, expense_name, provision_code, provision_name = _resolve_tax_accounts(entity)

    # Calculate net profit
    net_profit = _calculate_net_profit(fy)

    if net_profit <= 0:
        return JsonResponse({
            "eligible": False,
            "reason": "No tax provision required — entity is in a loss position.",
            "net_profit": str(net_profit),
        })

    # Determine tax rate
    if entity.is_base_rate_entity:
        tax_rate = Decimal("0.25")
        rate_label = "25% (Base Rate Entity)"
    else:
        tax_rate = Decimal("0.30")
        rate_label = "30% (Standard Rate)"

    calculated_tax = Decimal(_math.ceil(net_profit * tax_rate))

    # Existing provision balance
    existing_provision = _calculate_existing_provision(fy, provision_code)

    # Adjustment required
    adjustment_required = calculated_tax - existing_provision

    return JsonResponse({
        "eligible": True,
        "net_profit": str(net_profit),
        "tax_rate": str(tax_rate),
        "rate_label": rate_label,
        "calculated_tax": str(calculated_tax),
        "existing_provision": str(existing_provision),
        "adjustment_required": str(adjustment_required),
        "expense_account": {"code": expense_code, "name": expense_name},
        "provision_account": {"code": provision_code, "name": provision_name},
        "year_end": fy.end_date.strftime("%d %B %Y"),
    })


@login_required
@require_POST
def auto_tax_provision(request, pk):
    """Create and post a tax provision journal."""
    import math as _math

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    if entity.entity_type != "company":
        return JsonResponse({"error": "Tax provision is only available for company entities."}, status=400)

    if fy.is_locked:
        return JsonResponse({"error": "Financial year is locked."}, status=400)

    if entity.is_base_rate_entity is None:
        return JsonResponse({"error": "Please set the Base Rate Entity flag on the entity."}, status=400)

    # Check for existing tax provision journal
    if AdjustingJournal.objects.filter(financial_year=fy, journal_type="tax_provision").exists():
        return JsonResponse({"error": "A tax provision journal already exists."}, status=400)

    # Resolve accounts
    expense_code, expense_name, provision_code, provision_name = _resolve_tax_accounts(entity)

    # Calculate net profit
    net_profit = _calculate_net_profit(fy)
    if net_profit <= 0:
        return JsonResponse({"error": "No tax provision required — entity is in a loss position."}, status=400)

    # Tax rate
    if entity.is_base_rate_entity:
        tax_rate = Decimal("0.25")
        rate_label = "25% (Base Rate Entity)"
    else:
        tax_rate = Decimal("0.30")
        rate_label = "30% (Standard Rate)"

    calculated_tax = Decimal(_math.ceil(net_profit * tax_rate))
    existing_provision = _calculate_existing_provision(fy, provision_code)
    adjustment_required = calculated_tax - existing_provision

    if adjustment_required <= 0:
        return JsonResponse({
            "error": f"No adjustment required — existing provision (${existing_provision:,.0f}) already covers calculated tax (${calculated_tax:,.0f}).",
        }, status=400)

    # Ensure account codes exist in entity CoA
    tax_accounts = [
        (expense_code, expense_name, "expenses"),
        (provision_code, provision_name, "liabilities"),
    ]
    for code, name, section in tax_accounts:
        EntityChartOfAccount.objects.get_or_create(
            entity=entity,
            account_code=code,
            defaults={
                "account_name": name,
                "section": section,
                "is_custom": True,
            },
        )

    # Create and post the journal
    description = f"Tax provision for year ended {fy.end_date.strftime('%d %B %Y')}"
    with db_transaction.atomic():
        journal = AdjustingJournal.objects.create(
            financial_year=fy,
            journal_type="tax_provision",
            journal_date=fy.end_date,
            description=description,
            created_by=request.user,
            status=AdjustingJournal.JournalStatus.POSTED,
            posted_by=request.user,
            posted_at=timezone.now(),
            total_debit=adjustment_required,
            total_credit=adjustment_required,
        )
        JournalLine.objects.create(
            journal=journal,
            account_code=expense_code,
            account_name=expense_name,
            debit=adjustment_required,
            credit=Decimal("0"),
            line_number=1,
            description=f"Income tax expense at {rate_label}",
        )
        JournalLine.objects.create(
            journal=journal,
            account_code=provision_code,
            account_name=provision_name,
            debit=Decimal("0"),
            credit=adjustment_required,
            line_number=2,
            description=f"Current tax liability provision",
        )
        _post_journal_to_tb(journal, fy)

    _log_action(
        request, "adjustment",
        f"Posted tax provision {journal.reference_number} — ${adjustment_required:,.0f} at {rate_label} "
        f"(calculated: ${calculated_tax:,.0f}, existing: ${existing_provision:,.0f})",
        journal,
    )

    return JsonResponse({
        "success": True,
        "journal_ref": journal.reference_number,
        "adjustment_amount": str(adjustment_required),
        "calculated_tax": str(calculated_tax),
        "existing_provision": str(existing_provision),
        "rate_label": rate_label,
        "message": f"Tax provision posted — ${adjustment_required:,.0f} at {rate_label}",
    })


@login_required
def journal_delete(request, pk):
    """Delete a journal entry. If posted, also removes its adjustment TB lines.

    The entire reversal + deletion is wrapped in a database transaction so
    that a partial failure (e.g. TB reversal succeeds but journal delete
    fails) cannot leave the database in an inconsistent state.
    """
    journal = get_object_or_404(
        AdjustingJournal.objects.prefetch_related("lines"),
        pk=pk,
    )
    fy = journal.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check

    if fy.is_locked:
        messages.error(request, "Cannot delete journals in a finalised year.")
        return redirect("core:journal_detail", pk=pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission to delete journals.")
        return redirect("core:journal_detail", pk=pk)

    ref = journal.reference_number
    status = journal.status

    with db_transaction.atomic():
        # If the journal was posted, reverse its effect on the Trial Balance
        # using the multi-tier strategy (FK → exact match → broad match).
        if status == AdjustingJournal.JournalStatus.POSTED:
            _reverse_journal_tb_lines(journal)

            # Clean up orphaned TB lines: adjustment lines matching this
            # journal's account codes and amounts but with a NULL
            # source_journal FK (e.g. from pre-FK posting or failed edits).
            _delete_orphaned_tb_lines_for_journal(journal)

        # Sever FK links before delete so SET_NULL doesn't leave orphaned
        # TB lines that we intentionally chose not to delete (shouldn't
        # happen after _reverse_journal_tb_lines, but belt-and-suspenders).
        TrialBalanceLine.objects.filter(source_journal=journal).update(
            source_journal=None
        )

        journal.delete()

        # Post-deletion balance verification
        if status == AdjustingJournal.JournalStatus.POSTED:
            _verify_tb_balance(fy)

    _log_action(request, "adjustment", f"Deleted {status} journal {ref}")
    # Auto-trigger risk engine after journal deletion
    if status == AdjustingJournal.JournalStatus.POSTED:
        from core.signals import trigger_risk_recalc
        trigger_risk_recalc(fy, "journal_deleted")
    messages.success(request, f"Journal {ref} has been deleted.")
    from django.urls import reverse
    return redirect(reverse("core:financial_year_detail", args=[fy.pk]) + "?tab=journals")


@login_required
def journal_edit(request, pk):
    """Edit a posted journal entry.

    Reverses the existing TB adjustment lines, applies the corrected lines,
    and writes a full before/after diff to the audit log.
    Only accountants can edit journals; locked years are protected.
    """
    journal = get_object_or_404(
        AdjustingJournal.objects.select_related(
            "financial_year", "financial_year__entity",
            "created_by", "posted_by",
        ).prefetch_related("lines"),
        pk=pk,
    )
    fy = journal.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check
    entity = fy.entity

    if fy.is_locked:
        messages.error(request, "Cannot edit journals in a finalised year.")
        return redirect("core:journal_detail", pk=pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission to edit journals.")
        return redirect("core:journal_detail", pk=pk)

    # Build account list for the picker: merge TB lines with the entity's
    # Chart of Accounts so that accounts which exist in the CoA but have
    # no TB line yet (e.g. 2895) are still selectable.
    tb_accounts = (
        TrialBalanceLine.objects.filter(financial_year=fy)
        .exclude(account_code="")
        .values("account_code", "account_name")
        .distinct()
        .order_by("account_code")
    )
    tb_list = [
        {"client_account_code": a["account_code"], "client_account_name": a["account_name"]}
        for a in tb_accounts
    ]
    coa_qs = EntityChartOfAccount.objects.filter(entity=entity, is_active=True).order_by("account_code").values("account_code", "account_name")
    if not coa_qs.exists():
        coa_qs = ChartOfAccount.objects.filter(entity_type=entity.entity_type, is_active=True).order_by("account_code").values("account_code", "account_name")
    merged = {}
    for a in coa_qs:
        merged[a["account_code"]] = {"client_account_code": a["account_code"], "client_account_name": a["account_name"]}
    for a in tb_list:
        merged[a["client_account_code"]] = a  # TB name wins
    accounts = sorted(merged.values(), key=lambda x: x["client_account_code"])

    # Snapshot the journal state before any changes (for audit log)
    before_lines = [
        {
            "account_code": line.account_code,
            "account_name": line.account_name,
            "description": line.description,
            "debit": str(line.debit),
            "credit": str(line.credit),
        }
        for line in journal.lines.all()
    ]
    before_header = {
        "journal_type": journal.journal_type,
        "journal_date": str(journal.journal_date),
        "description": journal.description,
        "narration": journal.narration,
    }

    # Build formset with extra=1 so there's always one blank row available
    EditJournalLineFormSet = forms.inlineformset_factory(
        AdjustingJournal,
        JournalLine,
        form=JournalLineForm,
        extra=1,
        can_delete=True,
    )

    if request.method == "POST":
        form = AdjustingJournalForm(request.POST, instance=journal)
        formset = EditJournalLineFormSet(request.POST, instance=journal)
        if form.is_valid() and formset.is_valid():
            # Wrap the entire reverse → save → re-apply cycle in an atomic
            # block so a failure at any step rolls back all changes.
            with db_transaction.atomic():
                # Reverse all existing TB adjustment lines for this journal
                # using the multi-tier strategy (FK → exact match → broad).
                _reverse_journal_tb_lines(journal)

                # Save the updated header
                updated_journal = form.save(commit=False)
                updated_journal.save()

                # Save the updated lines
                new_lines = formset.save()

                # Renumber lines
                for i, line in enumerate(journal.lines.order_by("line_number", "id"), start=1):
                    line.line_number = i
                    line.save(update_fields=["line_number"])

                # Validate balance
                all_lines = list(journal.lines.all())
                total_dr = sum(l.debit for l in all_lines)
                total_cr = sum(l.credit for l in all_lines)
                if total_dr != total_cr:
                    # Roll back the atomic block — this aborts the entire edit
                    # including the line saves, so the old TB state is preserved.
                    db_transaction.set_rollback(True)
                    messages.error(
                        request,
                        f"Journal does not balance: Dr ${total_dr:,.2f} \u2260 Cr ${total_cr:,.2f}. "
                        "No changes were saved."
                    )
                    return render(request, "core/journal_edit.html", {
                        "form": form, "formset": formset, "journal": journal,
                        "fy": fy, "entity": entity, "accounts": accounts,
                    })

                # Apply the new lines to the TB, aggregating by account code
                _post_journal_to_tb(journal, fy)

                # Update cached totals
                journal.total_debit = total_dr
                journal.total_credit = total_cr
                journal.save(update_fields=["total_debit", "total_credit"])

            # Build after snapshot for audit log (outside atomic — read-only)
            after_lines = [
                {
                    "account_code": line.account_code,
                    "account_name": line.account_name,
                    "description": line.description,
                    "debit": str(line.debit),
                    "credit": str(line.credit),
                }
                for line in journal.lines.all()
            ]
            after_header = {
                "journal_type": updated_journal.journal_type,
                "journal_date": str(updated_journal.journal_date),
                "description": updated_journal.description,
                "narration": updated_journal.narration,
            }

            # Build a human-readable diff for the audit log
            import json as _json
            diff_parts = []
            for key in ("journal_type", "journal_date", "description", "narration"):
                if before_header.get(key) != after_header.get(key):
                    diff_parts.append(
                        f"  {key}: '{before_header.get(key)}' -> '{after_header.get(key)}'"
                    )
            if before_lines != after_lines:
                diff_parts.append("  Lines changed:")
                diff_parts.append(f"    Before: {_json.dumps(before_lines)}")
                diff_parts.append(f"    After:  {_json.dumps(after_lines)}")

            audit_detail = "\n".join(diff_parts) if diff_parts else "No changes detected."
            _log_action(
                request, "adjustment",
                f"Edited posted journal {journal.reference_number}:\n{audit_detail}",
                journal,
            )

            # Trigger risk engine recalc
            from core.signals import trigger_risk_recalc
            trigger_risk_recalc(fy, "journal_edited")

            messages.success(
                request,
                f"Journal {journal.reference_number} has been updated and the Trial Balance recalculated."
            )
            return redirect("core:journal_detail", pk=journal.pk)

    else:
        form = AdjustingJournalForm(instance=journal)
        formset = EditJournalLineFormSet(instance=journal)

    return render(request, "core/journal_edit.html", {
        "form": form, "formset": formset, "journal": journal,
        "fy": fy, "entity": entity, "accounts": accounts,
    })


@login_required
def account_list_api(request, pk):
    """JSON API endpoint returning available accounts for a financial year's entity."""
    fy = get_financial_year_for_user(request, pk)
    accounts = list(
        ClientAccountMapping.objects.filter(entity=fy.entity)
        .order_by("client_account_code")
        .values("client_account_code", "client_account_name")
    )
    # Also include accounts from the trial balance that may not be mapped
    tb_accounts = list(
        fy.trial_balance_lines.filter(is_adjustment=False)
        .values("account_code", "account_name")
        .distinct()
        .order_by("account_code")
    )
    # Merge: use TB accounts as base, supplement with mapped accounts
    account_dict = {}
    for a in tb_accounts:
        account_dict[a["account_code"]] = a["account_name"]
    for a in accounts:
        if a["client_account_code"] not in account_dict:
            account_dict[a["client_account_code"]] = a["client_account_name"]

    result = [
        {"code": code, "name": name}
        for code, name in sorted(account_dict.items())
    ]
    return JsonResponse(result, safe=False)


# ---------------------------------------------------------------------------
# Financial Statements Preview
# ---------------------------------------------------------------------------
@login_required
def financial_statements_view(request, pk):
    """Render a preview of the financial statements on screen."""
    get_financial_year_for_user(request, pk)  # IDOR check
    fy = get_object_or_404(
        FinancialYear.objects.select_related("entity", "prior_year"),
        pk=pk,
    )

    # Aggregate trial balance by mapped line item
    from django.db.models import Sum as DSum
    current_data = (
        fy.trial_balance_lines
        .filter(mapped_line_item__isnull=False)
        .values(
            "mapped_line_item__standard_code",
            "mapped_line_item__line_item_label",
            "mapped_line_item__financial_statement",
            "mapped_line_item__statement_section",
            "mapped_line_item__display_order",
        )
        .annotate(total=DSum("closing_balance"))
        .order_by("mapped_line_item__display_order")
    )

    # Get prior year data if available
    prior_data = {}
    if fy.prior_year:
        prior_lines = (
            fy.prior_year.trial_balance_lines
            .filter(mapped_line_item__isnull=False)
            .values("mapped_line_item__standard_code")
            .annotate(total=DSum("closing_balance"))
        )
        prior_data = {item["mapped_line_item__standard_code"]: item["total"] for item in prior_lines}

    # Organise into statements
    income_statement = []
    balance_sheet = []
    for item in current_data:
        entry = {
            "code": item["mapped_line_item__standard_code"],
            "label": item["mapped_line_item__line_item_label"],
            "section": item["mapped_line_item__statement_section"],
            "current": item["total"],
            "prior": prior_data.get(item["mapped_line_item__standard_code"], Decimal("0")),
        }
        if item["mapped_line_item__financial_statement"] == "income_statement":
            income_statement.append(entry)
        elif item["mapped_line_item__financial_statement"] == "balance_sheet":
            balance_sheet.append(entry)

    has_prior = bool(
        fy.entity.include_comparative_figures
        and fy.prior_year
        and fy.prior_year.trial_balance_lines.exists()
    )
    context = {
        "fy": fy,
        "entity": fy.entity,
        "income_statement": income_statement,
        "balance_sheet": balance_sheet,
        "has_prior": has_prior,
    }
    return render(request, "core/financial_statements_view.html", context)


# ---------------------------------------------------------------------------
# Line Item Breakdown (drill-down from Financial Statements preview)
# ---------------------------------------------------------------------------
@login_required
def line_item_breakdown(request, pk, standard_code):
    """Show all account codes that contribute to a given financial statement
    line item (identified by its AccountMapping standard_code).

    This enables drill-down from the Financial Statements preview page,
    where figures are aggregated by mapped line item.
    """
    fy = get_financial_year_for_user(request, pk)
    mapping = get_object_or_404(AccountMapping, standard_code=standard_code)

    # Get all TB lines mapped to this line item
    lines = (
        fy.trial_balance_lines
        .filter(mapped_line_item=mapping)
        .select_related('mapped_line_item')
        .order_by('account_code')
    )

    # Compute display Dr/Cr for each line
    total_dr = Decimal('0')
    total_cr = Decimal('0')
    for line in lines:
        cb = line.closing_balance if line.closing_balance else Decimal('0')
        if cb > 0:
            line.display_dr = cb
            line.display_cr = Decimal('0')
        elif cb < 0:
            line.display_dr = Decimal('0')
            line.display_cr = abs(cb)
        else:
            line.display_dr = line.debit if line.debit else Decimal('0')
            line.display_cr = line.credit if line.credit else Decimal('0')
        total_dr += line.display_dr or Decimal('0')
        total_cr += line.display_cr or Decimal('0')

    # Net total for the line item
    net_total = total_dr - total_cr

    # Prior year data for this line item
    prior_total = Decimal('0')
    if fy.prior_year:
        from django.db.models import Sum as DSum
        prior_agg = (
            fy.prior_year.trial_balance_lines
            .filter(mapped_line_item=mapping)
            .aggregate(total=DSum('closing_balance'))
        )
        prior_total = prior_agg['total'] or Decimal('0')

    # Year label
    current_year = str(fy.year_label)

    return render(request, 'core/line_item_breakdown.html', {
        'fy': fy,
        'mapping': mapping,
        'lines': lines,
        'total_dr': total_dr,
        'total_cr': total_cr,
        'net_total': net_total,
        'prior_total': prior_total,
        'current_year': current_year,
        'entry_count': lines.count(),
    })


# ---------------------------------------------------------------------------
# Document Generation (placeholder for Phase 2)
# ---------------------------------------------------------------------------
@login_required
def generate_document(request, pk):
    fy = get_financial_year_for_user(request, pk)

    # Check that there are trial balance lines
    if not fy.trial_balance_lines.exists():
        messages.error(request, "Cannot generate statements: no trial balance data loaded.")
        return redirect("core:financial_year_detail", pk=pk)

    # Determine requested format (default: docx)
    fmt = request.GET.get("format", "docx").lower()
    if fmt not in ("docx", "pdf"):
        fmt = "docx"

    # include_watermark: suppress DRAFT watermark once Eva has cleared the year
    include_watermark = not fy.can_assemble_package

    # Build filename
    entity_name = fy.entity.entity_name.replace(" ", "_")
    base_filename = f"{entity_name}_Financial_Statements_{fy.year_label}"

    if fmt == "pdf":
        from .fs_template_service import generate_combined_pdf
        try:
            pdf_buffer = generate_combined_pdf(fy.pk, include_watermark=include_watermark)
            pdf_bytes = pdf_buffer.getvalue()

            filename = f"{base_filename}.pdf"
            response = HttpResponse(pdf_bytes, content_type="application/pdf")
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            file_content = pdf_bytes
            file_format = "pdf"
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"PDF generation failed: {e}")
            messages.error(request, f"PDF generation failed: {e}")
            return redirect("core:financial_year_detail", pk=pk)
    else:
        from .fs_template_service import generate_financial_statements, DOCUMENT_TYPE_ORDER
        try:
            docs = generate_financial_statements(fy.pk, include_watermark=include_watermark)
            if not docs:
                raise RuntimeError("No templates rendered")
            # Serve the first rendered template as individual docx
            first_key = next(dt for dt in DOCUMENT_TYPE_ORDER if dt in docs)
            docx_buffer = docs[first_key]
            filename = f"{base_filename}.docx"
            response = HttpResponse(
                docx_buffer.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            file_content = docx_buffer.getvalue()
            file_format = "docx"
        except Exception as e:
            messages.error(request, f"Document generation failed: {e}")
            return redirect("core:financial_year_detail", pk=pk)

    # Log the generation
    _log_action(request, "generate", f"Generated financial statements ({file_format.upper()}) for {fy}", fy)

    # Save record
    from django.core.files.base import ContentFile
    doc = GeneratedDocument(
        financial_year=fy,
        file_format=file_format,
        generated_by=request.user,
    )
    doc.file.save(filename, ContentFile(file_content), save=True)
    # Also create/update a LegalDocument record so the package assembly
    # checklist can detect that distribution minutes have been generated.
    entity = fy.entity
    if entity.entity_type == "trust":
        from core.models import LegalDocument
        from core.views_compliance_docs import _sanitise_context_for_storage
        try:
            from core.document_context_builder import DocumentContextBuilder
            _dcb = DocumentContextBuilder(entity, financial_year=fy)
            _dist_ctx = _dcb.build("distribution_minutes")
        except Exception:
            _dist_ctx = {
                "entity_name": entity.entity_name,
                "financial_year": str(fy.end_date.year),
                "financial_year_end": str(fy.end_date),
            }
        LegalDocument.objects.update_or_create(
            financial_year=fy,
            entity=entity,
            document_type="distribution_minutes",
            defaults={
                "title": f"Trust Distribution Minutes — {entity.entity_name} — {fy.end_date.year}",
                "context_data": _sanitise_context_for_storage(_dist_ctx),
                "generated_by": request.user,
                "status": "generated",
            },
        )

    return response


# ---------------------------------------------------------------------------
# Management Accounts Generation (Period-Scoped Draft)
# ---------------------------------------------------------------------------
@login_required
def generate_management_accounts_view(request, pk):
    """Generate period-scoped, watermarked management accounts (Cover + B/S + P&L)."""
    from datetime import date as _date
    from .mgmt_accounts import generate_management_accounts

    fy = get_financial_year_for_user(request, pk)

    # Parse query parameters
    period_start_str = request.GET.get('period_start', '')
    period_end_str = request.GET.get('period_end', '')
    output_type = request.GET.get('output_type', 'bs_pnl')
    fmt = request.GET.get('format', 'docx').lower()

    if fmt not in ('docx', 'pdf'):
        fmt = 'docx'
    if output_type not in ('bs_pnl', 'pnl_only'):
        output_type = 'bs_pnl'

    # Parse dates
    try:
        period_start = _date.fromisoformat(period_start_str) if period_start_str else fy.start_date
        period_end = _date.fromisoformat(period_end_str) if period_end_str else fy.end_date
    except ValueError:
        messages.error(request, "Invalid date format. Use YYYY-MM-DD.")
        return redirect('core:financial_year_detail', pk=pk)

    # Generate
    try:
        buffer, tb_source = generate_management_accounts(
            financial_year_id=fy.pk,
            period_start=period_start,
            period_end=period_end,
            user=request.user,
            output_type=output_type,
        )
    except ValueError as e:
        messages.error(request, str(e))
        return redirect('core:financial_year_detail', pk=pk)
    except Exception as e:
        import logging
        logging.getLogger(__name__).error(f"Management accounts generation failed: {e}")
        messages.error(request, f"Management accounts generation failed: {e}")
        return redirect('core:financial_year_detail', pk=pk)

    # Build filename
    entity_name = fy.entity.entity_name.replace(' ', '_')
    period_label = f"{period_start.strftime('%b%Y')}_to_{period_end.strftime('%b%Y')}"
    base_filename = f"{entity_name}_Management_Accounts_{period_label}"

    if fmt == 'pdf':
        import subprocess, tempfile, os
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, f"{base_filename}.docx")
                with open(docx_path, 'wb') as f:
                    f.write(buffer.getvalue())

                from core.libreoffice_utils import convert_docx_to_pdf
                convert_docx_to_pdf(docx_path, tmpdir, timeout=120)
                pdf_path = os.path.join(tmpdir, f"{base_filename}.pdf")
                if not os.path.exists(pdf_path):
                    raise RuntimeError('PDF conversion failed')
                with open(pdf_path, 'rb') as f:
                    pdf_bytes = f.read()

            filename = f"{base_filename}.pdf"
            response = HttpResponse(pdf_bytes, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            file_content = pdf_bytes
            file_format = 'pdf'
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f'PDF conversion failed for mgmt accounts: {e}')
            messages.warning(request, f'PDF conversion failed: {e}. Falling back to DOCX.')
            filename = f"{base_filename}.docx"
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            file_content = buffer.getvalue()
            file_format = 'docx'
    else:
        filename = f"{base_filename}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        file_content = buffer.getvalue()
        file_format = 'docx'

    # Audit log
    _log_action(
        request, 'generate',
        f'Generated management accounts ({file_format.upper()}, {tb_source}) '
        f'for {period_start} to {period_end} — {fy}',
        fy,
    )

    # Activity log
    ActivityLog.objects.create(
        user=request.user,
        event_type=ActivityLog.EventType.MGMT_ACCOUNTS_GENERATED,
        title=f'Management Accounts generated ({file_format.upper()})',
        description=(
            f'Period: {period_start.strftime("%d %b %Y")} to {period_end.strftime("%d %b %Y")}. '
            f'Source: {tb_source}. Output: {output_type}.'
        ),
        entity=fy.entity,
        financial_year=fy,
        url=request.build_absolute_uri(
            reverse('core:financial_year_detail', args=[fy.pk])
        ),
    )

    # Save to GeneratedDocument (as management_accounts type)
    from django.core.files.base import ContentFile
    doc = GeneratedDocument(
        financial_year=fy,
        file_format=file_format,
        document_type=GeneratedDocument.DocumentType.MANAGEMENT_ACCOUNTS,
        status=GeneratedDocument.DocumentStatus.DRAFT,
        generated_by=request.user,
        change_summary=f'Period: {period_start} to {period_end}. Source: {tb_source}.',
    )
    doc.file.save(filename, ContentFile(file_content), save=True)

    return response


# ---------------------------------------------------------------------------
# Distribution Minutes Generation
# ---------------------------------------------------------------------------
@login_required
def generate_distribution_minutes(request, pk):
    """Generate distribution minutes document for a trust entity."""
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    # Only trusts have distribution minutes
    _TRUST_TYPES = ('trust', 'trust_unit', 'trust_discretionary', 'trust_hybrid')
    if entity.entity_type not in _TRUST_TYPES:
        messages.error(request, "Distribution minutes are only applicable to trust entities.")
        return redirect("core:financial_year_detail", pk=pk)

    fmt = request.GET.get("format", "docx").lower()
    if fmt not in ("docx", "pdf"):
        fmt = "docx"

    from .template_docgen import generate_from_template

    try:
        buffer = generate_from_template("distribution_minutes", "trust", fy.pk)
    except (ValueError, FileNotFoundError) as e:
        messages.error(request, f"Distribution minutes generation failed: {e}")
        return redirect("core:financial_year_detail", pk=pk)
    except Exception as e:
        messages.error(request, f"Unexpected error generating distribution minutes: {e}")
        return redirect("core:financial_year_detail", pk=pk)

    entity_name = entity.entity_name.replace(" ", "_")
    base_filename = f"{entity_name}_Distribution_Minutes_{fy.year_label}"

    if fmt == "pdf":
        import subprocess, tempfile, os
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, f"{base_filename}.docx")
                with open(docx_path, "wb") as f:
                    f.write(buffer.getvalue())

                from core.libreoffice_utils import convert_docx_to_pdf
                convert_docx_to_pdf(docx_path, tmpdir, timeout=120)
                pdf_path = os.path.join(tmpdir, f"{base_filename}.pdf")
                if not os.path.exists(pdf_path):
                    raise RuntimeError("PDF conversion failed.")
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()

            filename = f"{base_filename}.pdf"
            response = HttpResponse(pdf_bytes, content_type="application/pdf")
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            file_content = pdf_bytes
            file_format = "pdf"
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"PDF conversion failed: {e}")
            messages.error(request, f"PDF conversion failed: {e}. Falling back to DOCX.")
            filename = f"{base_filename}.docx"
            response = HttpResponse(
                buffer.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            response["Content-Disposition"] = f'attachment; filename="{filename}"'
            file_content = buffer.getvalue()
            file_format = "docx"
    else:
        filename = f"{base_filename}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        file_content = buffer.getvalue()
        file_format = "docx"

    _log_action(request, "generate", f"Generated distribution minutes ({file_format.upper()}) for {fy}", fy)

    # Save record
    from django.core.files.base import ContentFile
    doc = GeneratedDocument(
        financial_year=fy,
        file_format=file_format,
        generated_by=request.user,
    )
    doc.file.save(filename, ContentFile(file_content), save=True)

    # Also create/update a LegalDocument record so the package assembly
    # checklist detects distribution_minutes as present.
    from core.models import LegalDocument
    legal_doc, _created = LegalDocument.objects.update_or_create(
        financial_year=fy,
        document_type="distribution_minutes",
        defaults={
            "entity": entity,
            "title": f"Distribution Minutes — {entity.entity_name} — {fy.year_label}",
            "status": LegalDocument.Status.GENERATED,
            "generated_by": request.user,
        },
    )
    if file_format == "pdf":
        from django.core.files.base import ContentFile as CF
        legal_doc.pdf_file.save(filename, CF(file_content), save=True)
    else:
        from django.core.files.base import ContentFile as CF
        legal_doc.generated_file.save(filename, CF(file_content), save=True)

    return response


# ---------------------------------------------------------------------------
# HTMX Partials
# ---------------------------------------------------------------------------
@login_required
def htmx_client_search(request):
    """HTMX search endpoint for entities (replaces client search)."""
    query = request.GET.get("q", "").strip()
    entity_type = request.GET.get("entity_type", "")

    if request.user.can_view_all_entities:
        entities = Entity.objects.filter(is_archived=False)
    else:
        entities = Entity.objects.filter(
            assigned_accountant=request.user, is_archived=False
        )

    if entity_type:
        entities = entities.filter(entity_type=entity_type)

    if len(query) >= 2:
        entities = entities.filter(
            Q(entity_name__icontains=query)
            | Q(abn__icontains=query)
            | Q(trading_as__icontains=query)
        )

    entities = entities.distinct().annotate(num_fys=Count("financial_years")).order_by("entity_name")[:50]

    return render(request, "partials/entity_list_rows.html", {"entities": entities})


@login_required
def htmx_map_tb_line(request, pk):
    """HTMX endpoint to map a single trial balance line."""
    line = get_object_or_404(TrialBalanceLine, pk=pk)
    get_financial_year_for_user(request, line.financial_year.pk)  # IDOR check
    if not request.user.can_do_accounting:
        return HttpResponse("Permission denied", status=403)
    mapping_id = request.POST.get("mapped_line_item")

    if mapping_id:
        try:
            mapping = AccountMapping.objects.get(pk=mapping_id)
            line.mapped_line_item = mapping
            line.save()

            # Also save to client account mapping for reuse
            ClientAccountMapping.objects.update_or_create(
                entity=line.financial_year.entity,
                client_account_code=line.account_code,
                defaults={
                    "client_account_name": line.account_name,
                    "mapped_line_item": mapping,
                },
            )
        except AccountMapping.DoesNotExist:
            pass

    mappings = AccountMapping.objects.all()
    return render(request, "partials/tb_line_row.html", {
        "line": line, "mappings": mappings
    })


# ---------------------------------------------------------------------------
# Entity Officers / Signatories
# ---------------------------------------------------------------------------
@login_required
def entity_officers(request, pk):
    """List all officers/signatories for an entity."""
    entity = get_entity_for_user(request, pk)
    officers = entity.officers.all()

    officer_label_map = {
        "company": "Director / Officer",
        "trust": "Trustee / Beneficiary",
        "partnership": "Partner",
        "sole_trader": "Proprietor",
        "smsf": "Trustee / Director",
    }
    officer_label = officer_label_map.get(entity.entity_type, "Officer")

    # Calculate distribution totals for trust unit holders/beneficiaries
    dist_active = 0
    dist_ceased = 0
    dist_total = 0
    if entity.entity_type == "trust":
        from decimal import Decimal
        from django.db.models import Q, Sum
        from django.utils import timezone
        today = timezone.now().date()
        distribution_roles = [
            EntityOfficer.OfficerRole.UNIT_HOLDER,
            EntityOfficer.OfficerRole.BENEFICIARY,
        ]
        base_qs = officers.filter(
            role__in=distribution_roles,
            distribution_percentage__isnull=False,
        )
        dist_active = base_qs.filter(
            Q(date_ceased__isnull=True) | Q(date_ceased__gt=today)
        ).aggregate(total=Sum("distribution_percentage"))["total"] or Decimal("0")
        dist_ceased = base_qs.filter(
            date_ceased__isnull=False, date_ceased__lte=today,
        ).aggregate(total=Sum("distribution_percentage"))["total"] or Decimal("0")
        dist_total = dist_active + dist_ceased

    return render(request, "core/entity_officers.html", {
        "entity": entity,
        "officers": officers,
        "officer_label": officer_label,
        "dist_active": dist_active,
        "dist_ceased": dist_ceased,
        "dist_total": dist_total,
    })


def _handle_ceased_redistribution(request, officer):
    """Auto-redistribute distribution % when a unit holder/beneficiary is ceased."""
    from django.db import transaction
    from django.db.models import Q, Sum
    from django.utils import timezone
    from decimal import Decimal

    if officer.role not in EntityOfficer.DISTRIBUTION_ROLES:
        return
    today = timezone.now().date()
    if not officer.date_ceased or officer.date_ceased > today:
        return
    if not officer.distribution_percentage or officer.distribution_percentage <= 0:
        return

    ceased_pct = officer.distribution_percentage
    entity = officer.entity
    remaining = EntityOfficer.objects.filter(
        entity=entity,
        role__in=list(EntityOfficer.DISTRIBUTION_ROLES),
    ).filter(
        Q(date_ceased__isnull=True) | Q(date_ceased__gt=today)
    ).exclude(pk=officer.pk)

    count = remaining.count()
    if count == 0:
        return
    if count == 1:
        sole = remaining.first()
        with transaction.atomic():
            sole.distribution_percentage = Decimal("100.00")
            sole._updated_by = getattr(request, "user", None)
            sole.save()
        messages.info(
            request,
            f"{sole.full_name} is now the sole active unit holder and has been "
            f"set to 100.00% distribution."
        )
    else:
        active_total = remaining.filter(
            distribution_percentage__isnull=False,
        ).aggregate(total=Sum("distribution_percentage"))["total"] or Decimal("0")
        messages.warning(
            request,
            f"Warning: {officer.full_name} has been ceased. Their {ceased_pct}% "
            f"distribution has not been reallocated. Active unit holders currently "
            f"total {active_total}%. Please update distribution percentages manually."
        )


@login_required
def entity_officer_create(request, entity_pk):
    """Add a new officer/signatory to an entity."""
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_officers", pk=entity.pk)

    if request.method == "POST":
        form = EntityOfficerForm(request.POST, entity_type=entity.entity_type)
        # Set entity on instance before validation so clean() can query siblings
        form.instance.entity = entity
        if form.is_valid():
            officer = form.save(commit=False)
            officer.entity = entity
            officer.roles = form.cleaned_data["roles_multi"]
            officer._updated_by = request.user
            officer.save()
            _log_action(request, "user_change",
                        f"Added officer {officer.full_name} to {entity.entity_name}",
                        officer)
            messages.success(request, f"Added {officer.full_name} as {officer.get_role_display()}.")
            return redirect("core:entity_officers", pk=entity.pk)
    else:
        next_order = entity.officers.count() + 1
        form = EntityOfficerForm(
            entity_type=entity.entity_type,
            initial={"display_order": next_order},
        )

    return render(request, "core/entity_officer_form.html", {
        "form": form,
        "entity": entity,
    })


@login_required
def entity_officer_edit(request, pk):
    """Edit an existing officer/signatory."""
    officer = get_object_or_404(EntityOfficer, pk=pk)
    entity = officer.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_officers", pk=entity.pk)

    if request.method == "POST":
        form = EntityOfficerForm(request.POST, instance=officer, entity_type=entity.entity_type)
        if form.is_valid():
            obj = form.save(commit=False)
            obj.roles = form.cleaned_data["roles_multi"]
            obj._updated_by = request.user
            obj.save()
            _log_action(request, "user_change",
                        f"Updated officer {officer.full_name} for {entity.entity_name}",
                        officer)
            messages.success(request, f"Updated {officer.full_name}.")
            # Auto-redistribute distribution % when a unit holder/beneficiary is ceased
            _handle_ceased_redistribution(request, obj)
            return redirect("core:entity_officers", pk=entity.pk)
    else:
        form = EntityOfficerForm(instance=officer, entity_type=entity.entity_type)

    return render(request, "core/entity_officer_form.html", {
        "form": form,
        "entity": entity,
    })


@login_required
@require_POST
def entity_officer_delete(request, pk):
    """Delete an officer/signatory."""
    officer = get_object_or_404(EntityOfficer, pk=pk)
    entity = officer.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_officers", pk=entity.pk)
    name = officer.full_name
    _log_action(request, "user_change",
                f"Removed officer {name} from {entity.entity_name}",
                officer)
    from django.db import transaction as txn
    from core.models import OfficerDistributionHistory, EntityChartOfAccount
    with txn.atomic():
        OfficerDistributionHistory.objects.filter(officer=officer).delete()
        EntityChartOfAccount.objects.filter(
            beneficiary_officer=officer, auto_provisioned=True,
        ).delete()
        officer.delete()
    messages.success(request, f"Removed {name}.")
    return redirect("core:entity_officers", pk=entity.pk)


# ---------------------------------------------------------------------------
# Access Ledger Import
# ---------------------------------------------------------------------------
@login_required
def access_ledger_import(request):
    """Import an Access Ledger ZIP export."""
    from .access_ledger_import import import_access_ledger_zip

    result = None

    if request.method == "POST":
        zip_file = request.FILES.get("zip_file")
        if not zip_file:
            messages.error(request, "Please select a ZIP file.")
        elif not zip_file.name.lower().endswith(".zip"):
            messages.error(request, "File must be a .zip file.")
        else:
            replace = request.POST.get("replace_existing") == "1"
            import_finalised = request.POST.get("import_as_finalised") == "1"
            try:
                result = import_access_ledger_zip(
                    zip_file,
                    replace_existing=replace,
                    import_as_finalised=import_finalised,
                )
                if result["errors"]:
                    messages.warning(
                        request,
                        f"Import completed with {len(result['errors'])} error(s)."
                    )
                else:
                    status_note = " (all years finalised — ready to roll over)" if import_finalised else ""
                    messages.success(
                        request,
                        f"Successfully imported {result['entity'].entity_name}: "
                        f"{result['years_imported']} years, "
                        f"{result['total_tb_lines']} TB lines, "
                        f"{result['total_dep_assets']} depreciation assets."
                        f"{status_note}"
                    )
                _log_action(
                    request, "import",
                    f"Imported Access Ledger ZIP: {zip_file.name} "
                    f"({result['years_imported']} years)"
                    f"{' (import as finalised)' if import_finalised else ''}",
                    result.get("entity"),
                )
                # Auto-trigger risk engine for all imported years
                from core.signals import trigger_risk_recalc
                for _imp_fy in result.get('financial_years', []):
                    trigger_risk_recalc(_imp_fy, "access_ledger_import")
            except Exception as e:
                messages.error(request, "Import failed. Please check the file format and try again.")

    return render(request, "core/access_ledger_import.html", {
        "result": result,
    })


# ============================================================
# PDF Downloads: Trial Balance & Journals List
# ============================================================

@login_required
def trial_balance_pdf(request, pk):
    """Generate a comparative trial balance PDF matching professional accounting format."""
    from io import BytesIO
    from collections import OrderedDict
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        BaseDocTemplate, Frame, PageTemplate, Table, TableStyle,
        Paragraph, Spacer, Image, NextPageTemplate, PageBreak,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    import os
    from django.conf import settings

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    show_comparative = bool(
        entity.include_comparative_figures
        and fy.prior_year
        and fy.prior_year.trial_balance_lines.exists()
    )
    tb_lines = TrialBalanceLine.objects.filter(financial_year=fy).select_related('mapped_line_item').order_by('account_code')
    _coa_lookup = _build_coa_section_lookup(entity)
    # Determine section ordering and groupingg
    SECTION_ORDER = [
        'Revenue', 'Income', 'Cost of Sales', 'Expenses',
        'Current Assets', 'Non-Current Assets',
        'Current Liabilities', 'Non-Current Liabilities',
        'Equity', 'Income Tax',
    ]
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }

    # Group lines by section
    sections = OrderedDict()
    unmapped_lines = []
    for line in tb_lines:
        # HARD RULE: HandiLedger numeric code range is authoritative for section.
        hl_section = _hl_section_for_code(line.account_code)
        if hl_section:
            display_section = hl_section
        elif line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = _coa_lookup.get(line.account_code, 'Unmapped')
        if display_section not in sections:
            sections[display_section] = []
        sections[display_section].append(line)
    # Sort sections by defined orderr
    ordered_sections = OrderedDict()
    section_keys_ordered = []
    for s in SECTION_ORDER:
        ds = SECTION_DISPLAY.get(s, s)
        if ds not in section_keys_ordered:
            section_keys_ordered.append(ds)
    for key in section_keys_ordered:
        if key in sections:
            ordered_sections[key] = sections[key]
    # Add any remaining sections
    for key in sections:
        if key not in ordered_sections:
            ordered_sections[key] = sections[key]

    # Year labels
    current_year = str(fy.year_label)
    # Extract year number from label like "FY2026" or "2026"
    year_digits = ''.join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        prior_year = f"FY{int(year_digits) - 1}" if fy.year_label.startswith('FY') else str(int(year_digits) - 1)
    elif fy.prior_year:
        prior_year = str(fy.prior_year.year_label)
    else:
        prior_year = 'Prior'

    # ABN
    abn = entity.abn if hasattr(entity, 'abn') and entity.abn else ''
    abn_display = ''
    if abn:
        abn_clean = abn.replace(' ', '')
        if len(abn_clean) == 11:
            abn_display = f'ABN {abn_clean[:2]} {abn_clean[2:5]} {abn_clean[5:8]} {abn_clean[8:11]}'
        else:
            abn_display = f'ABN {abn}'

    buffer = BytesIO()

    # Custom page template with header and footer on every page
    class TBDocTemplate(BaseDocTemplate):
        def __init__(self, *args, **kwargs):
            self.entity_name = kwargs.pop('entity_name', '')
            self.abn_display = kwargs.pop('abn_display', '')
            self.end_date_str = kwargs.pop('end_date_str', '')
            self.current_year = kwargs.pop('current_year', '')
            self.prior_year = kwargs.pop('prior_year', '')
            super().__init__(*args, **kwargs)

        def afterPage(self):
            """Draw footer on every page."""
            canvas = self.canv
            canvas.saveState()
            # Footer line
            canvas.setStrokeColor(colors.HexColor('#333333'))
            canvas.setLineWidth(0.5)
            canvas.line(15*mm, 18*mm, A4[0] - 15*mm, 18*mm)
            # Footer text
            canvas.setFont('Helvetica-Bold', 7)
            canvas.setFillColor(colors.HexColor('#333333'))
            footer_text = (
                "These financial statements are unaudited. They must be read in conjunction with the attached "
                "Accountant's Compilation Report and Notes which form part of these financial statements."
            )
            # Center the text
            text_width = canvas.stringWidth(footer_text, 'Helvetica-Bold', 7)
            page_width = A4[0]
            if text_width > page_width - 30*mm:
                # Two-line footer
                line1 = "These financial statements are unaudited. They must be read in conjunction with the attached Accountant's"
                line2 = "Compilation Report and Notes which form part of these financial statements."
                w1 = canvas.stringWidth(line1, 'Helvetica-Bold', 7)
                w2 = canvas.stringWidth(line2, 'Helvetica-Bold', 7)
                canvas.drawString((page_width - w1) / 2, 13*mm, line1)
                canvas.drawString((page_width - w2) / 2, 10*mm, line2)
            else:
                canvas.drawString((page_width - text_width) / 2, 12*mm, footer_text)
            canvas.restoreState()

    doc = TBDocTemplate(
        buffer, pagesize=A4, topMargin=15*mm, bottomMargin=25*mm,
        leftMargin=15*mm, rightMargin=15*mm,
        entity_name=entity.entity_name.upper(),
        abn_display=abn_display,
        end_date_str=fy.end_date.strftime('%d %B %Y'),
        current_year=current_year,
        prior_year=prior_year,
    )

    frame = Frame(
        doc.leftMargin, doc.bottomMargin,
        doc.width, doc.height,
        id='main'
    )
    doc.addPageTemplates([PageTemplate(id='main', frames=[frame])])

    styles = getSampleStyleSheet()
    elements = []

    # Styles
    s_entity = ParagraphStyle('Entity', fontName='Helvetica-Bold', fontSize=14,
                               alignment=TA_CENTER, spaceAfter=2*mm)
    s_abn = ParagraphStyle('ABN', fontName='Helvetica', fontSize=10,
                            alignment=TA_CENTER, spaceAfter=2*mm)
    s_title = ParagraphStyle('TBTitle', fontName='Helvetica-Bold', fontSize=11,
                              alignment=TA_CENTER, spaceAfter=6*mm)
    s_section = ParagraphStyle('Section', fontName='Helvetica-Bold', fontSize=11,
                                spaceBefore=5*mm, spaceAfter=2*mm)
    s_cell = ParagraphStyle('Cell', fontName='Helvetica', fontSize=9)
    s_cell_bold = ParagraphStyle('CellBold', fontName='Helvetica-Bold', fontSize=9)
    s_num = ParagraphStyle('Num', fontName='Helvetica', fontSize=9, alignment=TA_RIGHT)
    s_num_bold = ParagraphStyle('NumBold', fontName='Helvetica-Bold', fontSize=9, alignment=TA_RIGHT)

    # Header
    elements.append(Paragraph(entity.entity_name.upper(), s_entity))
    if abn_display:
        elements.append(Paragraph(abn_display, s_abn))
    tb_title = (
        f"Comparative Trial Balance as at {fy.end_date.strftime('%d %B %Y')}"
        if show_comparative
        else f"Trial Balance as at {fy.end_date.strftime('%d %B %Y')}"
    )
    elements.append(Paragraph(tb_title, s_title))

    # Column header table — 6 columns when comparative, 4 when not
    if show_comparative:
        col_widths = [50, 165, 75, 75, 75, 75]
        header_data = [
            ['', '', current_year, current_year, prior_year, prior_year],
            ['', '', '$ Dr', '$ Cr', '$ Dr', '$ Cr'],
        ]
    else:
        col_widths = [50, 265, 75, 75]
        header_data = [
            ['', '', current_year, current_year],
            ['', '', '$ Dr', '$ Cr'],
        ]
    header_table = Table(header_data, colWidths=col_widths)
    header_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
        ('LINEBELOW', (2, 1), (-1, 1), 0.8, colors.HexColor('#333333')),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    elements.append(header_table)

    def fmt(val):
        """Format a decimal value with commas, or return empty string if zero."""
        if val and val != Decimal('0'):
            return f"{val:,.2f}"
        return ''

    # Totals accumulators
    grand_total_dr = Decimal('0')
    grand_total_cr = Decimal('0')
    grand_total_prior_dr = Decimal('0')
    grand_total_prior_cr = Decimal('0')

    # P&L accumulators for net profit calculation
    PL_SECTIONS = {'Income', 'Cost of Sales', 'Expenses'}
    pl_dr = Decimal('0')
    pl_cr = Decimal('0')
    pl_prior_dr = Decimal('0')
    pl_prior_cr = Decimal('0')

    # Aggregate lines by account_code to net adjustments (journal entries)
    aggregated_sections = _aggregate_tb_lines(ordered_sections, entity=entity)
    # Build section tables
    for section_name, lines in aggregated_sections.items():
        is_pl_section = section_name in PL_SECTIONS
        elements.append(Paragraph(f"<b>{section_name}</b>", s_section))

        data = []
        for line in lines:
            dr = line._agg_dr
            cr = line._agg_cr
            prior_dr = line._agg_prior_dr
            prior_cr = line._agg_prior_cr

            grand_total_dr += dr
            grand_total_cr += cr
            grand_total_prior_dr += prior_dr
            grand_total_prior_cr += prior_cr

            if is_pl_section:
                pl_dr += dr
                pl_cr += cr
                pl_prior_dr += prior_dr
                pl_prior_cr += prior_cr

            if show_comparative:
                row = [
                    Paragraph(f"<b>{line.account_code}</b>", ParagraphStyle('Code', fontName='Helvetica-Bold', fontSize=9)),
                    Paragraph(line.account_name, s_cell),
                    fmt(dr),
                    fmt(cr),
                    fmt(prior_dr),
                    fmt(prior_cr),
                ]
            else:
                row = [
                    Paragraph(f"<b>{line.account_code}</b>", ParagraphStyle('Code', fontName='Helvetica-Bold', fontSize=9)),
                    Paragraph(line.account_name, s_cell),
                    fmt(dr),
                    fmt(cr),
                ]
            data.append(row)

        if data:
            section_table = Table(data, colWidths=col_widths)
            style_cmds = [
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                ('LINEBELOW', (0, 0), (-1, -1), 0.3, colors.HexColor('#cccccc')),
            ]
            section_table.setStyle(TableStyle(style_cmds))
            elements.append(section_table)

    # Grand totals
    elements.append(Spacer(1, 3*mm))
    if show_comparative:
        totals_data = [[
            '', '',
            f"{grand_total_dr:,.2f}",
            f"{grand_total_cr:,.2f}",
            f"{grand_total_prior_dr:,.2f}",
            f"{grand_total_prior_cr:,.2f}",
        ]]
    else:
        totals_data = [[
            '', '',
            f"{grand_total_dr:,.2f}",
            f"{grand_total_cr:,.2f}",
        ]]
    totals_table = Table(totals_data, colWidths=col_widths)
    totals_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
        ('LINEABOVE', (2, 0), (-1, 0), 1.0, colors.HexColor('#333333')),
        ('LINEBELOW', (2, 0), (-1, 0), 0.5, colors.HexColor('#333333')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(totals_table)

    # Net Profit (P&L sections only: Income Cr - Expenses Dr)
    net_profit_current = pl_cr - pl_dr
    net_profit_prior = pl_prior_cr - pl_prior_dr
    elements.append(Spacer(1, 3*mm))

    # Format with profit/loss indicator
    if net_profit_current >= 0:
        np_current_str = f"${net_profit_current:,.2f}"
    else:
        np_current_str = f"(${abs(net_profit_current):,.2f})"
    if net_profit_prior >= 0:
        np_prior_str = f"${net_profit_prior:,.2f}"
    else:
        np_prior_str = f"(${abs(net_profit_prior):,.2f})"

    if show_comparative:
        profit_data = [[
            '', Paragraph('<b>Net Profit / (Loss)</b>', s_cell_bold),
            '', np_current_str,
            '', np_prior_str,
        ]]
    else:
        profit_data = [[
            '', Paragraph('<b>Net Profit / (Loss)</b>', s_cell_bold),
            '', np_current_str,
        ]]
    profit_table = Table(profit_data, colWidths=col_widths)
    profit_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (2, 0), (-1, -1), 'RIGHT'),
        ('LINEBELOW', (2, 0), (-1, 0), 0.8, colors.HexColor('#333333')),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    elements.append(profit_table)

    doc.build(elements)
    buffer.seek(0)

    tb_prefix = "Comparative_TB" if show_comparative else "Trial_Balance"
    filename = f"{tb_prefix}_{entity.entity_name.replace(' ', '_')}_{fy.year_label}.pdf"
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def journals_pdf(request, pk):
    """Generate a PDF listing all journal entries for a financial year."""
    from io import BytesIO
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, KeepTogether, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    import os

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    journals = AdjustingJournal.objects.filter(financial_year=fy).prefetch_related('lines').order_by('created_at')

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=20*mm, bottomMargin=20*mm,
                            leftMargin=15*mm, rightMargin=15*mm)
    styles = getSampleStyleSheet()
    elements = []

    # Logo — prefer FirmSettings upload, fall back to static file
    from django.conf import settings
    _logo_path = None
    try:
        from core.models import FirmSettings
        _logo_path = FirmSettings.get().logo_path
    except Exception:
        pass
    if not _logo_path:
        _logo_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'mcs_logo.png')
    if _logo_path and os.path.exists(_logo_path):
        logo = Image(_logo_path, width=40*mm, height=40*mm)
        logo.hAlign = 'LEFT'
        elements.append(logo)
        elements.append(Spacer(1, 3*mm))

    # Title
    title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=16, spaceAfter=4*mm)
    elements.append(Paragraph(f"{entity.entity_name}", title_style))

    subtitle_style = ParagraphStyle('Subtitle', parent=styles['Normal'], fontSize=11, textColor=colors.grey, spaceAfter=2*mm)
    elements.append(Paragraph(f"Journal Entries — {fy.year_label}", subtitle_style))
    elements.append(Paragraph(
        f"{fy.start_date.strftime('%d %b %Y')} to {fy.end_date.strftime('%d %b %Y')} · Status: {fy.get_status_display()}",
        ParagraphStyle('DateLine', parent=styles['Normal'], fontSize=9, textColor=colors.grey, spaceAfter=6*mm)
    ))

    if not journals.exists():
        elements.append(Paragraph("No journal entries recorded for this financial year.", styles['Normal']))
    else:
        # Summary table
        summary_data = [
            ['Total Journals', str(journals.count())],
            ['Posted', str(journals.filter(status='posted').count())],
            ['Draft', str(journals.filter(status='draft').count())],
        ]
        summary_table = Table(summary_data, colWidths=[100, 80])
        summary_table.setStyle(TableStyle([
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f8f9fa')),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 6*mm))

        # Each journal
        for journal in journals:
            journal_elements = []

            # Journal header
            ref = journal.reference_number or 'DRAFT'
            status_text = journal.get_status_display()
            jtype = journal.get_journal_type_display()
            header_text = (
                f"<b>{ref}</b> · {jtype} · {status_text} · "
                f"{journal.journal_date.strftime('%d %b %Y')}"
            )
            journal_elements.append(Paragraph(header_text, ParagraphStyle(
                'JournalHeader', fontSize=10, spaceAfter=1*mm,
                textColor=colors.HexColor('#212529')
            )))

            if journal.description:
                journal_elements.append(Paragraph(
                    journal.description,
                    ParagraphStyle('JDesc', fontSize=8, textColor=colors.grey, spaceAfter=2*mm)
                ))

            # Lines table
            lines = journal.lines.all()
            line_header = ['Account', 'Name', 'Description', 'Debit', 'Credit']
            line_data = [line_header]
            total_dr = Decimal('0')
            total_cr = Decimal('0')

            for line in lines:
                line_data.append([
                    line.account_code,
                    Paragraph(line.account_name, ParagraphStyle('LCell', fontSize=7)),
                    Paragraph(line.description or '', ParagraphStyle('LCell2', fontSize=7, textColor=colors.grey)),
                    f"{line.debit:,.2f}" if line.debit else '',
                    f"{line.credit:,.2f}" if line.credit else '',
                ])
                total_dr += line.debit or Decimal('0')
                total_cr += line.credit or Decimal('0')

            line_data.append(['', '', Paragraph('<b>Total</b>', ParagraphStyle('Bold', fontSize=7)),
                              f"{total_dr:,.2f}", f"{total_cr:,.2f}"])

            line_table = Table(line_data, colWidths=[50, 120, 130, 65, 65], repeatRows=1)
            line_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#495057')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 7),
                ('FONTSIZE', (0, 1), (-1, -1), 7),
                ('ALIGN', (3, 0), (4, -1), 'RIGHT'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -2), [colors.white, colors.HexColor('#f8f9fa')]),
                ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#e9ecef')),
                ('TOPPADDING', (0, 0), (-1, -1), 2),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            journal_elements.append(line_table)

            # Created by info
            created_by = journal.created_by.get_full_name() if journal.created_by else 'System'
            journal_elements.append(Paragraph(
                f"Created by {created_by} on {journal.created_at.strftime('%d/%m/%Y %H:%M')}",
                ParagraphStyle('CreatedBy', fontSize=6, textColor=colors.grey, spaceBefore=1*mm)
            ))
            journal_elements.append(Spacer(1, 5*mm))

            elements.append(KeepTogether(journal_elements))

    # Footer
    elements.append(Spacer(1, 4*mm))
    footer_style = ParagraphStyle('Footer', fontSize=7, textColor=colors.grey)
    try:
        from core.models import FirmSettings
        _firm_name_footer = FirmSettings.get().firm_name or "MC & S Pty Ltd"
    except Exception:
        _firm_name_footer = "MC & S Pty Ltd"
    elements.append(Paragraph(
        f"Generated by StatementHub on {timezone.now().strftime('%d %b %Y at %H:%M')} \u00b7 {_firm_name_footer}",
        footer_style
    ))

    doc.build(elements)
    buffer.seek(0)

    filename = f"Journals_{entity.entity_name.replace(' ', '_')}_{fy.year_label}.pdf"
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Client Associates (Family & Business)
# ---------------------------------------------------------------------------
@login_required
def associate_create(request, entity_pk):
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to add associates.")
        return redirect("core:entity_detail", pk=entity_pk)

    if request.method == "POST":
        form = ClientAssociateForm(request.POST)
        if form.is_valid():
            assoc = form.save(commit=False)
            assoc.entity = entity
            assoc.save()
            _log_action(request, "import", f"Added associate: {assoc.name} ({assoc.get_relationship_type_display()})", assoc)
            messages.success(request, f"Associate '{assoc.name}' added.")
            return redirect("core:entity_detail", pk=entity_pk)
    else:
        form = ClientAssociateForm()
    return render(request, "core/associate_form.html", {
        "form": form, "entity": entity, "title": "Add Associate"
    })


@login_required
def associate_edit(request, pk):
    assoc = get_object_or_404(ClientAssociate, pk=pk)
    entity = assoc.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to edit associates.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        form = ClientAssociateForm(request.POST, instance=assoc)
        if form.is_valid():
            form.save()
            messages.success(request, f"Associate '{assoc.name}' updated.")
            return redirect("core:entity_detail", pk=entity.pk)
    else:
        form = ClientAssociateForm(instance=assoc)
    return render(request, "core/associate_form.html", {
        "form": form, "entity": entity, "title": f"Edit: {assoc.name}"
    })


@login_required
def associate_delete(request, pk):
    assoc = get_object_or_404(ClientAssociate, pk=pk)
    entity = assoc.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to delete associates.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        name = assoc.name
        assoc.delete()
        messages.success(request, f"Associate '{name}' removed.")
    return redirect("core:entity_detail", pk=entity.pk)


# ---------------------------------------------------------------------------
# Entity-to-Entity Relationships
# ---------------------------------------------------------------------------
@login_required
def entity_link_search(request):
    """HTMX/JSON endpoint to search entities for the link modal."""
    query = request.GET.get("q", "")
    exclude_pk = request.GET.get("exclude", "")
    if len(query) < 2:
        return JsonResponse([], safe=False)
    entities = Entity.objects.filter(is_archived=False).filter(
        Q(entity_name__icontains=query)
        | Q(abn__icontains=query)
        | Q(trading_as__icontains=query)
    )
    if exclude_pk:
        entities = entities.exclude(pk=exclude_pk)
    entities = entities[:15]
    results = [
        {
            "id": str(e.pk),
            "name": e.entity_name,
            "type": e.get_entity_type_display(),
            "abn": e.abn or "",
        }
        for e in entities
    ]
    return JsonResponse(results, safe=False)


@login_required
def entity_link_create(request, entity_pk):
    """Create an entity-to-entity relationship."""
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to link entities.")
        return redirect("core:entity_detail", pk=entity_pk)
    if request.method == "POST":
        to_entity_id = request.POST.get("to_entity")
        relationship_type = request.POST.get("relationship_type", "associated_entity")
        notes = request.POST.get("notes", "")
        if not to_entity_id:
            messages.error(request, "Please select an entity to link.")
            return redirect("core:entity_detail", pk=entity_pk)
        try:
            to_entity = Entity.objects.get(pk=to_entity_id)
        except Entity.DoesNotExist:
            messages.error(request, "Selected entity not found.")
            return redirect("core:entity_detail", pk=entity_pk)
        if to_entity.pk == entity.pk:
            messages.error(request, "Cannot link an entity to itself.")
            return redirect("core:entity_detail", pk=entity_pk)
        from core.models import EntityRelationship
        _, created = EntityRelationship.objects.get_or_create(
            from_entity=entity,
            to_entity=to_entity,
            relationship_type=relationship_type,
            defaults={"notes": notes, "created_by": request.user},
        )
        if created:
            _log_action(request, "import", f"Linked entity: {entity.entity_name} → {relationship_type} → {to_entity.entity_name}", entity)
            messages.success(request, f"Linked to '{to_entity.entity_name}' as {dict(EntityRelationship.RelationshipType.choices).get(relationship_type, relationship_type)}.")
        else:
            messages.info(request, "This relationship already exists.")
    return redirect("core:entity_detail", pk=entity_pk)


@login_required
def entity_link_delete(request, pk):
    """Delete an entity-to-entity relationship."""
    from core.models import EntityRelationship
    rel = get_object_or_404(EntityRelationship, pk=pk)
    # Check user has access to at least one side
    try:
        get_entity_for_user(request, rel.from_entity.pk)
    except Exception:
        get_entity_for_user(request, rel.to_entity.pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to unlink entities.")
        return redirect("core:entity_detail", pk=rel.from_entity.pk)
    if request.method == "POST":
        entity_pk = request.POST.get("return_to", rel.from_entity.pk)
        other_name = rel.to_entity.entity_name if str(rel.from_entity.pk) == str(entity_pk) else rel.from_entity.entity_name
        rel.delete()
        messages.success(request, f"Unlinked from '{other_name}'.")
        return redirect("core:entity_detail", pk=entity_pk)
    return redirect("core:entity_detail", pk=rel.from_entity.pk)


# ---------------------------------------------------------------------------
# Accounting Software
# ---------------------------------------------------------------------------
@login_required
def software_create(request, entity_pk):
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to add software configs.")
        return redirect("core:entity_detail", pk=entity_pk)

    if request.method == "POST":
        form = AccountingSoftwareForm(request.POST)
        if form.is_valid():
            sw = form.save(commit=False)
            sw.entity = entity
            sw.save()
            _log_action(request, "import", f"Added software: {sw.get_software_type_display()}", sw)
            messages.success(request, f"Software '{sw.get_software_type_display()}' added.")
            return redirect("core:entity_detail", pk=entity_pk)
    else:
        form = AccountingSoftwareForm()
    return render(request, "core/software_form.html", {
        "form": form, "entity": entity, "title": "Add Accounting Software", "is_new": True
    })


@login_required
def software_edit(request, pk):
    sw = get_object_or_404(AccountingSoftware, pk=pk)
    entity = sw.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to edit software configs.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        form = AccountingSoftwareForm(request.POST, instance=sw)
        if form.is_valid():
            form.save()
            messages.success(request, f"Software '{sw.get_software_type_display()}' updated.")
            return redirect("core:entity_detail", pk=entity.pk)
    else:
        form = AccountingSoftwareForm(instance=sw)
    return render(request, "core/software_form.html", {
        "form": form, "entity": entity, "title": f"Edit: {sw.get_software_type_display()}", "is_new": False
    })


@login_required
def software_delete(request, pk):
    sw = get_object_or_404(AccountingSoftware, pk=pk)
    entity = sw.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission to delete software configs.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        label = sw.get_software_type_display()
        sw.delete()
        messages.success(request, f"Software '{label}' removed.")
    return redirect("core:entity_detail", pk=entity.pk)


# ---------------------------------------------------------------------------
# Meeting Notes
# ---------------------------------------------------------------------------
@login_required
def meeting_note_create(request, entity_pk):
    entity = get_entity_for_user(request, entity_pk)
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=entity_pk)

    if request.method == "POST":
        form = MeetingNoteForm(request.POST)
        if form.is_valid():
            note = form.save(commit=False)
            note.entity = entity
            note.created_by = request.user
            note.save()
            _log_action(request, "import", f"Added meeting note: {note.title}", note)
            messages.success(request, f"Meeting note '{note.title}' created.")
            return redirect("core:entity_detail", pk=entity_pk)
    else:
        form = MeetingNoteForm(initial={"meeting_date": timezone.now().date()})
    return render(request, "core/meeting_note_form.html", {
        "form": form, "entity": entity, "title": "New Meeting Note"
    })


@login_required
def meeting_note_edit(request, pk):
    note = get_object_or_404(MeetingNote, pk=pk)
    entity = note.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        form = MeetingNoteForm(request.POST, instance=note)
        if form.is_valid():
            form.save()
            messages.success(request, f"Meeting note '{note.title}' updated.")
            return redirect("core:entity_detail", pk=entity.pk)
    else:
        form = MeetingNoteForm(instance=note)
    return render(request, "core/meeting_note_form.html", {
        "form": form, "entity": entity, "title": f"Edit: {note.title}"
    })


@login_required
def meeting_note_detail(request, pk):
    note = get_object_or_404(MeetingNote, pk=pk)
    entity = note.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    return render(request, "core/meeting_note_detail.html", {
        "note": note, "entity": entity,
    })


@login_required
def meeting_note_delete(request, pk):
    note = get_object_or_404(MeetingNote, pk=pk)
    entity = note.entity
    get_entity_for_user(request, entity.pk)  # IDOR check
    if not request.user.can_edit:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=entity.pk)

    if request.method == "POST":
        title = note.title
        note.delete()
        messages.success(request, f"Meeting note '{title}' deleted.")
    return redirect("core:entity_detail", pk=entity.pk)


@login_required
def meeting_note_toggle_followup(request, pk):
    """HTMX endpoint to toggle follow-up completion."""
    note = get_object_or_404(MeetingNote, pk=pk)
    get_entity_for_user(request, note.entity.pk)  # IDOR check
    if not request.user.can_edit:
        return JsonResponse({"error": "Permission denied"}, status=403)
    note.follow_up_completed = not note.follow_up_completed
    note.save(update_fields=["follow_up_completed"])
    return JsonResponse({"completed": note.follow_up_completed})



# ---------------------------------------------------------------------------
# GST Activity Statement
# ---------------------------------------------------------------------------
@login_required
def gst_activity_statement(request, pk):
    """
    Generate a GST Activity Statement (BAS summary) for a financial year.
    Maps trial balance lines to BAS labels using ChartOfAccount tax codes.
    Supports both Simpler BAS (G1, 1A, 1B) and Full BAS (G1-G20).
    """
    fy = get_financial_year_for_user(request, pk)
    # Re-fetch with select_related for efficiency
    fy = get_object_or_404(
        FinancialYear.objects.select_related("entity", "entity__client"),
        pk=pk,
    )
    entity = fy.entity
    entity_type = entity.entity_type  # company, trust, partnership, sole_trader

    # Build a lookup: account_code -> ChartOfAccount for this entity type
    # First load template COA, then overlay with entity-specific COA
    coa_lookup = {}
    for coa in ChartOfAccount.objects.filter(entity_type=entity_type, is_active=True):
        coa_lookup[coa.account_code] = coa

    # Also load entity-specific chart of accounts (these take priority)
    entity_coa_lookup = {}
    for ecoa in EntityChartOfAccount.objects.filter(entity=entity):
        entity_coa_lookup[ecoa.account_code] = ecoa

    # Get all trial balance lines for this financial year
    tb_lines = fy.trial_balance_lines.all()

    # Initialize BAS labels
    g1 = Decimal("0")   # Total sales (including GST)
    g2 = Decimal("0")   # Export sales
    g3 = Decimal("0")   # Other GST-free sales
    g4 = Decimal("0")   # Input taxed sales
    g10 = Decimal("0")  # Capital purchases (including GST)
    g11 = Decimal("0")  # Non-capital purchases (including GST)
    g13 = Decimal("0")  # Purchases for making input taxed sales
    g14 = Decimal("0")  # Purchases without GST (GST-free purchases)
    g15 = Decimal("0")  # Private use / not income tax deductible
    g7 = Decimal("0")   # Sales adjustments
    g18 = Decimal("0")  # Purchase adjustments

    # Detailed line items for the breakdown
    sales_lines = []
    purchase_lines = []
    capital_lines = []
    excluded_lines = []

    for line in tb_lines:
        coa = coa_lookup.get(line.account_code)
        ecoa = entity_coa_lookup.get(line.account_code)

        # --- Determine section and tax code ---
        # For lines in COA, use COA section/tax. For unmapped lines with
        # tax_type from bank statement review, infer section from tax_type.
        line_tax = (getattr(line, 'tax_type', '') or '').strip()

        if not coa and not ecoa:
            # Not in any COA — check if we can infer from tax_type
            # GST control/clearing accounts are excluded from BAS
            # (GST is calculated from the revenue/expense gross amounts)
            if line.account_code in ('3380', '9100', '9110'):
                excluded_lines.append({
                    "code": line.account_code,
                    "name": line.account_name,
                    "amount": abs(line.closing_balance),
                    "reason": "GST clearing account",
                })
                continue

            # Infer section and tax_code from the line-level tax_type
            tax_type_map = {
                'GST on Income': ('revenue', 'GST'),
                'GST on Expenses': ('expenses', 'INP'),
                'GST Free Income': ('revenue', 'FRE'),
                'GST Free Expenses': ('expenses', 'FRE'),
                'BAS Excluded': (None, 'N-T'),
                'N-T': (None, 'N-T'),
                'Input Taxed': (None, 'ITS'),
            }
            mapped = tax_type_map.get(line_tax)
            if mapped and mapped[0]:
                section = mapped[0]
                tax_code = mapped[1]
            else:
                # Truly unmapped — exclude
                excluded_lines.append({
                    "code": line.account_code,
                    "name": line.account_name,
                    "amount": abs(line.closing_balance),
                    "reason": f"Not in chart of accounts" + (f" (tax: {line_tax})" if line_tax else ""),
                })
                continue
        else:
            # Determine section: prefer entity COA, fall back to template COA
            if ecoa:
                section = ecoa.section
            else:
                section = coa.section

            # Determine tax code: prefer entity COA tax_code, then template, then line-level
            ecoa_tax = (ecoa.tax_code or "").upper().strip() if ecoa else ""
            coa_tax = (coa.tax_code or "").upper().strip() if coa else ""

            # Priority: entity COA tax > template COA tax > line-level tax
            base_tax = ecoa_tax or coa_tax
            if not base_tax and line_tax:
                tax_map = {
                    'GST on Income': 'GST',
                    'GST on Expenses': 'INP',
                    'GST Free Income': 'FRE',
                    'GST Free Expenses': 'FRE',
                    'BAS Excluded': 'N-T',
                    'N-T': 'N-T',
                }
                tax_code = tax_map.get(line_tax, base_tax)
            else:
                tax_code = base_tax

        # Use closing balance; fall back to debit/credit when closing_balance is 0
        # (bank statement TB lines only set debit/credit, not closing_balance)
        if line.closing_balance != 0:
            amount = abs(line.closing_balance)
        else:
            amount = max(line.debit, line.credit)

        # BAS requires GROSS amounts (including GST).
        # Bank statement TB lines store NET amounts (ex-GST) with GST in 3380 (GST payable control account).
        # For GST-coded lines, gross up: net * 11/10
        if tax_code in ('INP', 'GST') and line.source == 'bank_statement':
            amount = (amount * Decimal('11') / Decimal('10')).quantize(Decimal('0.01'))

        # Section values are lowercase DB values from StatementSection TextChoices
        if section in ("revenue", "Revenue"):
            # All revenue goes to G1
            g1 += amount
            bas_label = "G1"

            if tax_code == "GST":
                bas_label = "G1 (Taxable)"
            elif tax_code == "ITS":
                g4 += amount
                bas_label = "G4 (Input Taxed)"
            elif tax_code == "ADS":
                g7 += amount
                bas_label = "G7 (Adjustment)"
            elif tax_code in ("", "FRE", "N-T"):
                g3 += amount
                bas_label = "G3 (GST-Free)"

            sales_lines.append({
                "code": line.account_code,
                "name": line.account_name,
                "tax_code": tax_code or "N-T",
                "amount": amount,
                "bas_label": bas_label,
            })

        elif section in ("expenses", "Expenses", "cost_of_sales", "Cost of Sales"):
            # Non-capital purchases go to G11
            if tax_code in ("INP", "GST"):
                g11 += amount
                bas_label = "G11 (Non-Capital)"
            elif tax_code in ("IOA",):
                g11 += amount
                g13 += amount
                bas_label = "G11/G13 (Input Taxed)"
            elif tax_code in ("FOA", "FRE"):
                g11 += amount
                g14 += amount
                bas_label = "G11/G14 (GST-Free)"
            elif tax_code == "ADS":
                g11 += amount
                g18 += amount
                bas_label = "G11/G18 (Adjustment)"
            else:
                # No tax code - still include in G11 but also G14
                g11 += amount
                g14 += amount
                bas_label = "G11/G14 (No GST)"

            purchase_lines.append({
                "code": line.account_code,
                "name": line.account_name,
                "tax_code": tax_code or "N-T",
                "amount": amount,
                "bas_label": bas_label,
            })

        elif section in ("assets", "Assets"):
            if tax_code == "CAP":
                g10 += amount
                bas_label = "G10 (Capital)"
                capital_lines.append({
                    "code": line.account_code,
                    "name": line.account_name,
                    "tax_code": tax_code,
                    "amount": amount,
                    "bas_label": bas_label,
                })
            elif tax_code == "FCA":
                g10 += amount
                g14 += amount
                bas_label = "G10/G14 (GST-Free Capital)"
                capital_lines.append({
                    "code": line.account_code,
                    "name": line.account_name,
                    "tax_code": tax_code,
                    "amount": amount,
                    "bas_label": bas_label,
                })
            # Other assets (no tax code) are balance sheet items, not on BAS

        # Liabilities, Equity, Capital Accounts - not on BAS

    # Calculated fields
    g5 = g2 + g3 + g4                    # Total non-taxable sales
    g6 = g1 - g5                          # Sales subject to GST
    g8 = g6 + g7                          # Sales subject to GST after adjustments
    g9 = (g8 / Decimal("11")).quantize(Decimal("0.01")) if g8 else Decimal("0")  # GST on sales

    g12 = g10 + g11                       # Total purchases
    g16 = g13 + g14 + g15                 # Non-creditable purchases
    g17 = g12 - g16                       # Purchases subject to GST
    g19 = g17 + g18                       # Purchases subject to GST after adjustments
    g20 = (g19 / Decimal("11")).quantize(Decimal("0.01")) if g19 else Decimal("0")  # GST on purchases

    label_1a = g9                         # GST on sales (payable)
    label_1b = g20                        # GST on purchases (credit)
    gst_payable = label_1a - label_1b     # Net GST payable (or refund if negative)

    # Build the context
    bas_data = {
        # Sales section
        "G1": g1, "G2": g2, "G3": g3, "G4": g4,
        "G5": g5, "G6": g6, "G7": g7, "G8": g8, "G9": g9,
        # Purchases section
        "G10": g10, "G11": g11, "G12": g12, "G13": g13,
        "G14": g14, "G15": g15, "G16": g16, "G17": g17,
        "G18": g18, "G19": g19, "G20": g20,
        # Summary
        "1A": label_1a, "1B": label_1b,
        "gst_payable": gst_payable,
    }

    context = {
        "fy": fy,
        "entity": entity,
        "bas_data": bas_data,
        "sales_lines": sales_lines,
        "purchase_lines": purchase_lines,
        "capital_lines": capital_lines,
        "excluded_lines": excluded_lines,
        "is_gst_registered": entity.is_gst_registered,
    }
    return render(request, "core/gst_activity_statement.html", context)


@login_required
def gst_activity_statement_download(request, pk):
    """Download GST Activity Statement as Excel or PDF."""
    import io

    fmt = request.GET.get("format", "excel").lower()

    fy = get_object_or_404(
        FinancialYear.objects.select_related("entity", "entity__client"),
        pk=pk,
    )
    entity = fy.entity
    entity_type = entity.entity_type

    # Re-run the calculation (same logic as the view)
    coa_lookup = {}
    for coa in ChartOfAccount.objects.filter(entity_type=entity_type, is_active=True):
        coa_lookup[coa.account_code] = coa

    # Also load entity-specific chart of accounts (these take priority)
    entity_coa_lookup = {}
    for ecoa in EntityChartOfAccount.objects.filter(entity=entity):
        entity_coa_lookup[ecoa.account_code] = ecoa

    tb_lines = fy.trial_balance_lines.all()

    g1 = g2 = g3 = g4 = g7 = Decimal("0")
    g10 = g11 = g13 = g14 = g15 = g18 = Decimal("0")
    detail_rows = []

    for line in tb_lines:
        coa = coa_lookup.get(line.account_code)
        ecoa = entity_coa_lookup.get(line.account_code)
        line_tax = (getattr(line, 'tax_type', '') or '').strip()

        if not coa and not ecoa:
            # Skip GST control/clearing accounts
            if line.account_code in ('3380', '9100', '9110'):
                continue
            # Infer from tax_type
            tax_type_map = {
                'GST on Income': ('revenue', 'GST'),
                'GST on Expenses': ('expenses', 'INP'),
                'GST Free Income': ('revenue', 'FRE'),
                'GST Free Expenses': ('expenses', 'FRE'),
            }
            mapped = tax_type_map.get(line_tax)
            if mapped and mapped[0]:
                section = mapped[0]
                tax_code = mapped[1]
            else:
                continue
        else:
            if ecoa:
                section = ecoa.section
            else:
                section = coa.section
            ecoa_tax = (ecoa.tax_code or "").upper().strip() if ecoa else ""
            coa_tax = (coa.tax_code or "").upper().strip() if coa else ""
            base_tax = ecoa_tax or coa_tax
            if not base_tax and line_tax:
                tax_map = {
                    'GST on Income': 'GST',
                    'GST on Expenses': 'INP',
                    'GST Free Income': 'FRE',
                    'GST Free Expenses': 'FRE',
                    'BAS Excluded': 'N-T',
                    'N-T': 'N-T',
                }
                tax_code = tax_map.get(line_tax, base_tax)
            else:
                tax_code = base_tax

        # Use closing balance; fall back to debit/credit when closing_balance is 0
        if line.closing_balance != 0:
            amount = abs(line.closing_balance)
        else:
            amount = max(line.debit, line.credit)

        # BAS requires GROSS amounts (including GST).
        # Bank statement TB lines store NET amounts (ex-GST) with GST in 3380 (GST payable control account).
        if tax_code in ('INP', 'GST') and line.source == 'bank_statement':
            amount = (amount * Decimal('11') / Decimal('10')).quantize(Decimal('0.01'))

        bas_label = ""
        if section in ("revenue", "Revenue"):
            g1 += amount
            if tax_code == "GST":
                bas_label = "G1"
            elif tax_code == "ITS":
                g4 += amount
                bas_label = "G4"
            elif tax_code == "ADS":
                g7 += amount
                bas_label = "G7"
            else:
                g3 += amount
                bas_label = "G3"
        elif section in ("expenses", "Expenses", "cost_of_sales", "Cost of Sales"):
            if tax_code in ("INP", "GST"):
                g11 += amount
                bas_label = "G11"
            elif tax_code == "IOA":
                g11 += amount
                g13 += amount
                bas_label = "G11/G13"
            elif tax_code in ("FOA", "FRE"):
                g11 += amount
                g14 += amount
                bas_label = "G11/G14"
            elif tax_code == "ADS":
                g11 += amount
                g18 += amount
                bas_label = "G11/G18"
            else:
                g11 += amount
                g14 += amount
                bas_label = "G11/G14"
        elif section in ("assets", "Assets"):
            if tax_code == "CAP":
                g10 += amount
                bas_label = "G10"
            elif tax_code == "FCA":
                g10 += amount
                g14 += amount
                bas_label = "G10/G14"

        if bas_label:
            detail_rows.append({
                "code": line.account_code,
                "name": line.account_name,
                "tax_code": tax_code or "N-T",
                "amount": float(amount),
                "bas_label": bas_label,
            })

    g5 = g2 + g3 + g4
    g6 = g1 - g5
    g8 = g6 + g7
    g9 = (g8 / Decimal("11")).quantize(Decimal("0.01")) if g8 else Decimal("0")
    g12 = g10 + g11
    g16 = g13 + g14 + g15
    g17 = g12 - g16
    g19 = g17 + g18
    g20 = (g19 / Decimal("11")).quantize(Decimal("0.01")) if g19 else Decimal("0")
    label_1a = g9
    label_1b = g20

    # PDF format
    if fmt == "pdf":
        return _gst_download_pdf(
            fy, entity, detail_rows,
            g1, g2, g3, g4, g5, g6, g7, g8, g9,
            g10, g11, g12, g13, g14, g15, g16, g17, g18, g19, g20,
            label_1a, label_1b,
        )

    # Create Excel workbook
    wb = openpyxl.Workbook()

    # Sheet 1: BAS Summary
    ws = wb.active
    ws.title = "GST Activity Statement"

    # Header
    ws.append([f"GST Activity Statement — {entity.entity_name}"])
    ws.append([f"Period: {fy.start_date.strftime('%d/%m/%Y')} to {fy.end_date.strftime('%d/%m/%Y')}"])
    ws.append([f"ABN: {entity.abn or 'N/A'}"])
    ws.append([])

    # GST on Sales
    ws.append(["GST ON SALES"])
    ws.append(["Label", "Description", "Amount"])
    ws.append(["G1", "Total sales (including any GST)", float(g1)])
    ws.append(["G2", "Export sales", float(g2)])
    ws.append(["G3", "Other GST-free sales", float(g3)])
    ws.append(["G4", "Input taxed sales", float(g4)])
    ws.append(["G5", "G2 + G3 + G4", float(g5)])
    ws.append(["G6", "Total sales subject to GST (G1 - G5)", float(g6)])
    ws.append(["G7", "Adjustments", float(g7)])
    ws.append(["G8", "Total sales subject to GST after adjustments (G6 + G7)", float(g8)])
    ws.append(["G9", "GST on sales (G8 ÷ 11)", float(g9)])
    ws.append([])

    # GST on Purchases
    ws.append(["GST ON PURCHASES"])
    ws.append(["Label", "Description", "Amount"])
    ws.append(["G10", "Capital purchases (including any GST)", float(g10)])
    ws.append(["G11", "Non-capital purchases (including any GST)", float(g11)])
    ws.append(["G12", "G10 + G11", float(g12)])
    ws.append(["G13", "Purchases for making input taxed sales", float(g13)])
    ws.append(["G14", "Purchases without GST in the price", float(g14)])
    ws.append(["G15", "Estimated purchases for private use", float(g15)])
    ws.append(["G16", "G13 + G14 + G15", float(g16)])
    ws.append(["G17", "Total purchases subject to GST (G12 - G16)", float(g17)])
    ws.append(["G18", "Adjustments", float(g18)])
    ws.append(["G19", "Total purchases subject to GST after adjustments (G17 + G18)", float(g19)])
    ws.append(["G20", "GST on purchases (G19 ÷ 11)", float(g20)])
    ws.append([])

    # Summary
    ws.append(["BAS SUMMARY"])
    ws.append(["Label", "Description", "Amount"])
    ws.append(["1A", "GST on sales", float(label_1a)])
    ws.append(["1B", "GST on purchases", float(label_1b)])
    ws.append(["", "Net GST payable / (refundable)", float(label_1a - label_1b)])

    # Format columns
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 55
    ws.column_dimensions["C"].width = 18

    # Bold headers
    from openpyxl.styles import Font, numbers
    bold = Font(bold=True)
    for row_idx in [1, 2, 3, 5, 6, 17, 18, 31, 32]:
        for cell in ws[row_idx]:
            cell.font = bold

    # Number format for amount column
    for row in ws.iter_rows(min_row=7, max_col=3, max_row=ws.max_row):
        cell = row[2]
        if isinstance(cell.value, (int, float)):
            cell.number_format = '#,##0'

    # Sheet 2: Detail Breakdown
    ws2 = wb.create_sheet("Detail Breakdown")
    ws2.append(["Account Code", "Account Name", "Tax Code", "Amount", "BAS Label"])
    for row in detail_rows:
        ws2.append([row["code"], row["name"], row["tax_code"], row["amount"], row["bas_label"]])

    ws2.column_dimensions["A"].width = 14
    ws2.column_dimensions["B"].width = 40
    ws2.column_dimensions["C"].width = 12
    ws2.column_dimensions["D"].width = 18
    ws2.column_dimensions["E"].width = 14

    for cell in ws2[1]:
        cell.font = bold
    for row in ws2.iter_rows(min_row=2, min_col=4, max_col=4, max_row=ws2.max_row):
        row[0].number_format = '#,##0'

    # Write to response
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    entity_name = entity.entity_name.replace(" ", "_")
    filename = f"GST_Activity_Statement_{entity_name}_{fy.start_date.strftime('%Y%m%d')}_{fy.end_date.strftime('%Y%m%d')}.xlsx"

    response = HttpResponse(
        buffer.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response



def _gst_download_pdf(fy, entity, detail_rows,
                      g1, g2, g3, g4, g5, g6, g7, g8, g9,
                      g10, g11, g12, g13, g14, g15, g16, g17, g18, g19, g20,
                      label_1a, label_1b):
    """Generate GST Activity Statement as PDF using WeasyPrint."""
    import io
    import weasyprint
    from django.utils.html import escape as html_escape

    abn = html_escape(entity.abn or 'N/A')
    period = f"{fy.start_date.strftime('%d/%m/%Y')} to {fy.end_date.strftime('%d/%m/%Y')}"
    net_gst = label_1a - label_1b
    net_label = "Net GST Payable to ATO" if net_gst > 0 else "Net GST Refundable from ATO"

    def fmt(v):
        return f"${v:,.0f}"

    # Build detail rows HTML (escaped to prevent HTML injection)
    detail_html = ""
    for row in detail_rows:
        detail_html += f"""<tr>
            <td>{html_escape(str(row['code']))}</td><td>{html_escape(str(row['name']))}</td>
            <td>{html_escape(str(row['tax_code']))}</td><td class="r">${row['amount']:,.0f}</td>
            <td>{html_escape(str(row['bas_label']))}</td>
        </tr>"""

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
    @page {{ size: A4; margin: 15mm; }}
    body {{ font-family: 'Times New Roman', serif; font-size: 10pt; color: #333; }}
    h1 {{ font-size: 14pt; text-align: center; margin-bottom: 2mm; }}
    h2 {{ font-size: 11pt; margin-top: 6mm; margin-bottom: 3mm; border-bottom: 1px solid #333; padding-bottom: 2mm; }}
    .sub {{ text-align: center; font-size: 10pt; color: #666; margin-bottom: 4mm; }}
    table {{ width: 100%; border-collapse: collapse; margin-bottom: 4mm; }}
    th, td {{ padding: 3px 6px; font-size: 9pt; border-bottom: 1px solid #ddd; }}
    th {{ background: #f5f5f5; text-align: left; font-weight: bold; }}
    .r {{ text-align: right; }}
    .bold {{ font-weight: bold; }}
    .highlight {{ background: #e8f4fd; }}
    .summary {{ background: #f8f9fa; }}
    .total-row {{ border-top: 2px solid #333; font-weight: bold; }}
</style></head><body>

<h1>GST Activity Statement</h1>
<p class="sub">{html_escape(entity.entity_name)} &mdash; ABN: {abn}</p>
<p class="sub">Period: {period}</p>

<h2>GST on Sales</h2>
<table>
    <tr><th>Label</th><th>Description</th><th class="r">Amount</th></tr>
    <tr><td class="bold">G1</td><td>Total sales (including any GST)</td><td class="r">{fmt(g1)}</td></tr>
    <tr><td>G2</td><td>Export sales</td><td class="r">{fmt(g2)}</td></tr>
    <tr><td>G3</td><td>Other GST-free sales</td><td class="r">{fmt(g3)}</td></tr>
    <tr><td>G4</td><td>Input taxed sales</td><td class="r">{fmt(g4)}</td></tr>
    <tr class="summary"><td class="bold">G5</td><td>G2 + G3 + G4</td><td class="r bold">{fmt(g5)}</td></tr>
    <tr><td class="bold">G6</td><td>Total sales subject to GST (G1 &minus; G5)</td><td class="r bold">{fmt(g6)}</td></tr>
    <tr><td>G7</td><td>Adjustments</td><td class="r">{fmt(g7)}</td></tr>
    <tr class="summary"><td class="bold">G8</td><td>Total sales subject to GST after adj. (G6 + G7)</td><td class="r bold">{fmt(g8)}</td></tr>
    <tr class="highlight"><td class="bold">G9</td><td>GST on sales (G8 &divide; 11)</td><td class="r bold">{fmt(g9)}</td></tr>
</table>

<h2>GST on Purchases</h2>
<table>
    <tr><th>Label</th><th>Description</th><th class="r">Amount</th></tr>
    <tr><td class="bold">G10</td><td>Capital purchases (including any GST)</td><td class="r">{fmt(g10)}</td></tr>
    <tr><td class="bold">G11</td><td>Non-capital purchases (including any GST)</td><td class="r">{fmt(g11)}</td></tr>
    <tr class="summary"><td class="bold">G12</td><td>G10 + G11</td><td class="r bold">{fmt(g12)}</td></tr>
    <tr><td>G13</td><td>Purchases for making input taxed sales</td><td class="r">{fmt(g13)}</td></tr>
    <tr><td>G14</td><td>Purchases without GST in the price</td><td class="r">{fmt(g14)}</td></tr>
    <tr><td>G15</td><td>Estimated purchases for private use</td><td class="r">{fmt(g15)}</td></tr>
    <tr class="summary"><td class="bold">G16</td><td>G13 + G14 + G15</td><td class="r bold">{fmt(g16)}</td></tr>
    <tr><td class="bold">G17</td><td>Total purchases subject to GST (G12 &minus; G16)</td><td class="r bold">{fmt(g17)}</td></tr>
    <tr><td>G18</td><td>Adjustments</td><td class="r">{fmt(g18)}</td></tr>
    <tr class="summary"><td class="bold">G19</td><td>Total purchases subject to GST after adj. (G17 + G18)</td><td class="r bold">{fmt(g19)}</td></tr>
    <tr class="highlight"><td class="bold">G20</td><td>GST on purchases (G19 &divide; 11)</td><td class="r bold">{fmt(g20)}</td></tr>
</table>

<h2>Activity Statement Summary</h2>
<table style="max-width: 400px;">
    <tr><td class="bold">1A</td><td>GST on sales</td><td class="r bold">{fmt(label_1a)}</td></tr>
    <tr><td class="bold">1B</td><td>GST on purchases (credit)</td><td class="r bold">{fmt(label_1b)}</td></tr>
    <tr class="total-row"><td colspan="2">{net_label}</td><td class="r">{fmt(abs(net_gst))}</td></tr>
</table>

<h2>Detail Breakdown</h2>
<table>
    <tr><th>Code</th><th>Account Name</th><th>Tax Code</th><th class="r">Amount</th><th>BAS Label</th></tr>
    {detail_html}
</table>

</body></html>"""

    pdf_bytes = weasyprint.HTML(string=html).write_pdf()

    entity_name = entity.entity_name.replace(' ', '_')
    filename = f"GST_Activity_Statement_{entity_name}_{fy.start_date.strftime('%Y%m%d')}_{fy.end_date.strftime('%Y%m%d')}.pdf"

    response = HttpResponse(pdf_bytes, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Depreciation Asset AJAX endpoints
# ---------------------------------------------------------------------------
@login_required
def depreciation_add(request, pk):
    """Add a new depreciation asset to a financial year."""
    fy = get_financial_year_for_user(request, pk)
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        asset = DepreciationAsset.objects.create(
            financial_year=fy,
            category=request.POST.get("category", "Other"),
            asset_name=request.POST.get("asset_name", ""),
            purchase_date=request.POST.get("purchase_date") or None,
            total_cost=Decimal(request.POST.get("total_cost", "0") or "0"),
            private_use_pct=Decimal(request.POST.get("private_use_pct", "0") or "0"),
            opening_wdv=Decimal(request.POST.get("opening_wdv", "0") or "0"),
            method=request.POST.get("method", "D"),
            rate=Decimal(request.POST.get("rate", "0") or "0"),
            addition_cost=Decimal(request.POST.get("addition_cost", "0") or "0"),
            addition_date=request.POST.get("addition_date") or None,
            disposal_date=request.POST.get("disposal_date") or None,
            disposal_consideration=Decimal(request.POST.get("disposal_consideration", "0") or "0"),
            asset_account_code=request.POST.get("asset_account_code", "").strip(),
            asset_account_name=request.POST.get("asset_account_name", "").strip(),
            accum_dep_code=request.POST.get("accum_dep_code", "").strip(),
            accum_dep_name=request.POST.get("accum_dep_name", "").strip(),
            dep_expense_code=request.POST.get("dep_expense_code", "").strip(),
            dep_expense_name=request.POST.get("dep_expense_name", "").strip(),
        )
    except (InvalidOperation, ValueError):
        messages.error(request, "Invalid numeric value provided.")
        return redirect("core:financial_year_detail", pk=pk)
    # Calculate depreciation
    _calc_depreciation(asset)
    asset.save()

    _log_action(request, "create", f"Added depreciation asset: {asset.asset_name}", asset)
    messages.success(request, f"Asset '{asset.asset_name}' added.")
    return redirect("core:financial_year_detail", pk=pk)



@login_required
def depreciation_suggest_account_code(request, pk):
    """
    AJAX endpoint: suggest the next available account code for a new
    depreciation asset account or accumulated depreciation account.
    Delegates to the existing entity_coa_suggest_code logic so the
    alphabetical-placement algorithm is reused exactly.
    Expects GET params: account_name, section (assets|liabilities|expenses)
    """
    from django.http import QueryDict
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    section = request.GET.get('section', 'assets').strip()
    account_name = (request.GET.get('account_name') or request.GET.get('name') or '').strip()
    if not account_name:
        return JsonResponse({'suggested_code': '', 'position_info': 'Enter an account name.'})
    # Reuse the existing suggest-code logic by forwarding to the same function
    # but constructing a fake GET with the right params.
    fake_get = QueryDict(mutable=True)
    fake_get['section'] = section
    fake_get['account_name'] = account_name
    request.GET = fake_get
    return entity_coa_suggest_code(request, pk)


@login_required
def depreciation_create_account(request, pk):
    """
    AJAX POST endpoint: create a new EntityChartOfAccount directly from
    journal entry, depreciation, or any account picker modal.

    Accepts both form-encoded POST and JSON body.  Uses the same fields,
    validation, and model save as the full-page entity_coa_add view so
    there is no parallel code path.

    Required: account_name, account_code, section
    Optional: tax_code, classification, maps_to (AccountMapping PK)
    """
    import json as _json

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    if not request.user.can_do_accounting:
        return JsonResponse({'error': 'Permission denied'}, status=403)
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)

    # Accept both form-encoded and JSON bodies
    if request.content_type and 'application/json' in request.content_type:
        try:
            data = _json.loads(request.body)
        except (ValueError, TypeError):
            return JsonResponse({'error': 'Invalid JSON'}, status=400)
    else:
        data = request.POST

    account_name = (data.get('account_name') or data.get('name') or '').strip()
    account_code = (data.get('account_code') or data.get('code') or '').strip()
    section = (data.get('section') or '').strip()
    tax_code = (data.get('tax_code') or '').strip()
    classification = (data.get('classification') or '').strip()
    maps_to_id = (data.get('maps_to') or '').strip()

    # --- Same validation as entity_coa_add ---
    if not account_code or not account_name or not section:
        return JsonResponse(
            {'error': 'Account code, name, and section are required.'},
            status=400,
        )

    valid_sections = [s[0] for s in EntityChartOfAccount.StatementSection.choices]
    if section not in valid_sections:
        return JsonResponse({'error': f'Invalid section: {section}'}, status=400)

    # Check for duplicate code
    if EntityChartOfAccount.objects.filter(entity=entity, account_code=account_code).exists():
        existing = EntityChartOfAccount.objects.get(entity=entity, account_code=account_code)
        return JsonResponse({
            'success': True,
            'code': existing.account_code,
            'name': existing.account_name,
            # Backwards-compat keys for depreciation modal in financial_year_detail
            'account_code': existing.account_code,
            'account_name': existing.account_name,
            'already_existed': True,
        })

    # Resolve maps_to (same as entity_coa_add)
    maps_to = None
    if maps_to_id:
        try:
            maps_to = AccountMapping.objects.get(pk=maps_to_id)
        except AccountMapping.DoesNotExist:
            pass

    # --- Same model save as entity_coa_add ---
    new_account = EntityChartOfAccount.objects.create(
        entity=entity,
        account_code=account_code,
        account_name=account_name,
        section=section,
        classification=classification,
        tax_code=tax_code,
        maps_to=maps_to,
        is_active=True,
        is_custom=True,
    )
    _log_action(
        request, 'create',
        f'Created entity account from modal: {account_code} — {account_name} ({section})',
        fy,
    )
    return JsonResponse({
        'success': True,
        'code': new_account.account_code,
        'name': new_account.account_name,
        # Backwards-compat keys for depreciation modal in financial_year_detail
        'account_code': new_account.account_code,
        'account_name': new_account.account_name,
        'already_existed': False,
    })

@login_required
def depreciation_edit(request, pk):
    """Edit a depreciation asset."""
    asset = get_object_or_404(DepreciationAsset, pk=pk)
    fy_pk = asset.financial_year.pk
    get_financial_year_for_user(request, fy_pk)  # IDOR check
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        asset.category = request.POST.get("category", asset.category)
        asset.asset_name = request.POST.get("asset_name", asset.asset_name)
        asset.purchase_date = request.POST.get("purchase_date") or asset.purchase_date
        asset.total_cost = Decimal(request.POST.get("total_cost", "0") or "0")
        asset.private_use_pct = Decimal(request.POST.get("private_use_pct", "0") or "0")
        asset.opening_wdv = Decimal(request.POST.get("opening_wdv", "0") or "0")
        asset.method = request.POST.get("method", asset.method)
        asset.rate = Decimal(request.POST.get("rate", "0") or "0")
        asset.addition_cost = Decimal(request.POST.get("addition_cost", "0") or "0")
        asset.addition_date = request.POST.get("addition_date") or None
        asset.disposal_date = request.POST.get("disposal_date") or None
        asset.disposal_consideration = Decimal(request.POST.get("disposal_consideration", "0") or "0")
        # Account mapping fields
        if request.POST.get("asset_account_code"):
            asset.asset_account_code = request.POST.get("asset_account_code", "").strip()
            asset.asset_account_name = request.POST.get("asset_account_name", "").strip()
        if request.POST.get("accum_dep_code"):
            asset.accum_dep_code = request.POST.get("accum_dep_code", "").strip()
            asset.accum_dep_name = request.POST.get("accum_dep_name", "").strip()
        if request.POST.get("dep_expense_code"):
            asset.dep_expense_code = request.POST.get("dep_expense_code", "").strip()
            asset.dep_expense_name = request.POST.get("dep_expense_name", "").strip()
    except (InvalidOperation, ValueError):
        messages.error(request, "Invalid numeric value provided.")
        return redirect("core:financial_year_detail", pk=fy_pk)
    _calc_depreciation(asset)
    asset.save()

    _log_action(request, "update", f"Updated depreciation asset: {asset.asset_name}", asset)
    messages.success(request, f"Asset '{asset.asset_name}' updated.")
    return redirect("core:financial_year_detail", pk=fy_pk)


@login_required
@require_POST
def depreciation_delete(request, pk):
    """Delete a depreciation asset."""
    asset = get_object_or_404(DepreciationAsset, pk=pk)
    fy_pk = asset.financial_year.pk
    get_financial_year_for_user(request, fy_pk)  # IDOR check
    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=fy_pk)
    name = asset.asset_name
    asset.delete()
    _log_action(request, "delete", f"Deleted depreciation asset: {name}")
    messages.success(request, f"Asset '{name}' deleted.")
    return redirect("core:financial_year_detail", pk=fy_pk)


@login_required
@require_POST
def depreciation_roll_forward(request, pk):
    """Roll forward depreciation schedule from prior year."""
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)
    if not fy.prior_year:
        messages.error(request, "No prior year linked. Cannot roll forward depreciation.")
        return redirect("core:financial_year_detail", pk=pk)

    prior_assets = DepreciationAsset.objects.filter(financial_year=fy.prior_year)
    if not prior_assets.exists():
        messages.warning(request, "No depreciation assets in the prior year to roll forward.")
        return redirect("core:financial_year_detail", pk=pk)

    # Clear existing assets if any
    DepreciationAsset.objects.filter(financial_year=fy).delete()

    count = 0
    for pa in prior_assets:
        if pa.closing_wdv <= 0 and not pa.disposal_date:
            continue  # Skip fully depreciated with no disposal
        new_asset = DepreciationAsset(
            financial_year=fy,
            category=pa.category,
            asset_name=pa.asset_name,
            purchase_date=pa.purchase_date,
            total_cost=pa.total_cost,
            private_use_pct=pa.private_use_pct,
            opening_wdv=pa.closing_wdv,  # Prior closing = new opening
            method=pa.method,
            rate=pa.rate,
            display_order=pa.display_order,
            # Carry forward account mapping so journal posting still works
            asset_account_code=pa.asset_account_code,
            asset_account_name=pa.asset_account_name,
            accum_dep_code=pa.accum_dep_code,
            accum_dep_name=pa.accum_dep_name,
            dep_expense_code=pa.dep_expense_code,
            dep_expense_name=pa.dep_expense_name,
            notes=f"Rolled forward from FY{fy.prior_year.year_label}",
        )
        # Pre-calculate depreciation so the schedule is immediately populated
        _calc_depreciation(new_asset)
        new_asset.save()
        count += 1

    _log_action(request, "roll_forward", f"Rolled forward {count} depreciation assets to {fy}", fy)
    messages.success(request, f"Rolled forward {count} depreciation assets from prior year.")
    return redirect("core:financial_year_detail", pk=pk)


def _calc_depreciation(asset):
    """Calculate depreciation amount and closing WDV for an asset."""
    depreciable = asset.opening_wdv + asset.addition_cost
    if asset.disposal_date:
        # On disposal, depreciation is up to disposal date
        # Profit/loss = disposal consideration - WDV at disposal
        asset.depreciation_amount = Decimal("0")
        asset.closing_wdv = Decimal("0")
        wdv_at_disposal = depreciable
        asset.profit_on_disposal = max(Decimal("0"), asset.disposal_consideration - wdv_at_disposal)
        asset.loss_on_disposal = max(Decimal("0"), wdv_at_disposal - asset.disposal_consideration)
    elif asset.method == "W":
        # Written off entirely
        asset.depreciation_amount = depreciable
        asset.closing_wdv = Decimal("0")
    elif asset.method == "D":
        # Diminishing value
        dep = (depreciable * asset.rate / Decimal("100")).quantize(Decimal("0.01"))
        asset.depreciation_amount = dep
        asset.closing_wdv = (depreciable - dep).quantize(Decimal("0.01"))
    elif asset.method == "P":
        # Prime cost (straight line)
        dep = (asset.total_cost * asset.rate / Decimal("100")).quantize(Decimal("0.01"))
        asset.depreciation_amount = min(dep, depreciable)
        asset.closing_wdv = (depreciable - asset.depreciation_amount).quantize(Decimal("0.01"))
    else:
        asset.depreciation_amount = Decimal("0")
        asset.closing_wdv = depreciable

    # Private use adjustment
    if asset.private_use_pct > 0:
        asset.private_depreciation = (
            asset.depreciation_amount * asset.private_use_pct / Decimal("100")
        ).quantize(Decimal("0.01"))
    else:
        asset.private_depreciation = Decimal("0")

    asset.depreciable_value = depreciable


# ---------------------------------------------------------------------------
# Add Depreciation Asset from Bank Statement Transaction
# ---------------------------------------------------------------------------
@login_required
def depreciation_add_from_transaction(request, pk):
    """
    Pre-populate the depreciation add form from a bank statement transaction.
    GET:  Renders the depreciation schedule tab with the add-asset modal pre-filled.
    POST: Creates the asset linked to the source transaction, then redirects back.
    """
    from review.models import PendingTransaction
    txn = get_object_or_404(PendingTransaction, pk=pk)
    fy_entity = txn.job.entity

    # Find the financial year that covers this transaction
    # Use the FY linked to the job's entity — pick the one whose year range covers the txn date
    fy = None
    if txn.date:
        import datetime
        try:
            txn_date = datetime.datetime.strptime(txn.date, "%Y-%m-%d").date()
        except (ValueError, TypeError):
            txn_date = None
        if txn_date:
            fy = FinancialYear.objects.filter(
                entity=fy_entity,
                start_date__lte=txn_date,
                end_date__gte=txn_date,
            ).first()
    # Fallback: most recent FY for the entity
    if not fy:
        fy = FinancialYear.objects.filter(entity=fy_entity).order_by('-end_date').first()
    if not fy:
        messages.error(request, "No financial year found for this entity.")
        return redirect('core:entity_detail', pk=fy_entity.pk)

    # IDOR check
    get_financial_year_for_user(request, fy.pk)

    if request.method == "POST":
        if not request.user.can_do_accounting:
            return JsonResponse({"error": "Permission denied"}, status=403)
        try:
            asset = DepreciationAsset.objects.create(
                financial_year=fy,
                category=request.POST.get("category", "Other"),
                asset_name=request.POST.get("asset_name", ""),
                purchase_date=request.POST.get("purchase_date") or None,
                total_cost=Decimal(request.POST.get("total_cost", "0") or "0"),
                private_use_pct=Decimal(request.POST.get("private_use_pct", "0") or "0"),
                opening_wdv=Decimal(request.POST.get("opening_wdv", "0") or "0"),
                method=request.POST.get("method", "D"),
                rate=Decimal(request.POST.get("rate", "0") or "0"),
                addition_cost=Decimal(request.POST.get("addition_cost", "0") or "0"),
                addition_date=request.POST.get("addition_date") or None,
                disposal_date=request.POST.get("disposal_date") or None,
                disposal_consideration=Decimal(request.POST.get("disposal_consideration", "0") or "0"),
                source_transaction=txn,
                asset_account_code=request.POST.get("asset_account_code", "").strip(),
                asset_account_name=request.POST.get("asset_account_name", "").strip(),
                accum_dep_code=request.POST.get("accum_dep_code", "").strip(),
                accum_dep_name=request.POST.get("accum_dep_name", "").strip(),
                dep_expense_code=request.POST.get("dep_expense_code", "").strip(),
                dep_expense_name=request.POST.get("dep_expense_name", "").strip(),
            )
        except (InvalidOperation, ValueError):
            messages.error(request, "Invalid numeric value provided.")
            return redirect("core:financial_year_detail", pk=fy.pk)
        _calc_depreciation(asset)
        asset.save()
        _log_action(request, "create", f"Added depreciation asset from bank txn: {asset.asset_name}", asset)
        messages.success(request, f"Asset '{asset.asset_name}' added to depreciation schedule from bank transaction.")
        return redirect(reverse("core:financial_year_detail", args=[fy.pk]) + "?tab=depreciation")

    # GET — render a standalone page with the pre-filled form
    # Calculate net amount (use absolute value since expenses are negative)
    net_amount = abs(txn.net_amount) if txn.net_amount else abs(txn.amount) if txn.amount else Decimal("0")

    # Try to guess a category from the account name
    account_name = txn.confirmed_name or ""
    category_guess = "Other"
    name_lower = account_name.lower()
    if "motor" in name_lower or "vehicle" in name_lower:
        category_guess = "Motor Vehicles"
    elif "computer" in name_lower or "laptop" in name_lower or "apple" in name_lower:
        category_guess = "Computer Equipment"
    elif "furniture" in name_lower or "fixture" in name_lower:
        category_guess = "Furniture and Fixtures"
    elif "office" in name_lower:
        category_guess = "Office Equipment"
    elif "plant" in name_lower or "equipment" in name_lower:
        category_guess = "Plant and Equipment"
    elif "leasehold" in name_lower or "improvement" in name_lower:
        category_guess = "Leasehold Improvements"

    # ── Auto-detect account codes from entity COA ──
    asset_acct_code = txn.confirmed_code or ""
    asset_acct_name = txn.confirmed_name or ""
    accum_dep_code = ""
    accum_dep_name = ""
    dep_expense_code = ""
    dep_expense_name = ""

    # Find the paired accumulated depreciation account:
    # Look for the nearest "accumulated depreciation/amortisation" account
    # in the entity COA whose code is > the asset account code.
    if asset_acct_code:
        entity_asset_accounts = list(
            EntityChartOfAccount.objects.filter(
                entity=fy.entity, is_active=True, section="assets",
            ).order_by("account_code")
        )
        # Find the accumulated dep account that sits after this asset account
        found_asset = False
        for acct in entity_asset_accounts:
            if acct.account_code == asset_acct_code:
                found_asset = True
                continue
            if found_asset:
                acct_lower = acct.account_name.lower()
                if "accum" in acct_lower or "amortis" in acct_lower:
                    accum_dep_code = acct.account_code
                    accum_dep_name = acct.account_name
                    break
                # If we hit another non-accumulated asset account, stop looking
                if "less:" not in acct_lower and "accumulated" not in acct_lower:
                    break

        # If not found after the asset, try any accumulated dep account in assets section
        if not accum_dep_code:
            for acct in entity_asset_accounts:
                acct_lower = acct.account_name.lower()
                if ("accum" in acct_lower and "deprec" in acct_lower) or \
                   ("accum" in acct_lower and "amortis" in acct_lower):
                    accum_dep_code = acct.account_code
                    accum_dep_name = acct.account_name
                    break

    # Auto-detect depreciation expense account
    dep_expense_acct = EntityChartOfAccount.objects.filter(
        entity=fy.entity, is_active=True, section="expenses",
        account_name__icontains="depreciation",
    ).exclude(
        account_name__icontains="accum"
    ).first()
    if dep_expense_acct:
        dep_expense_code = dep_expense_acct.account_code
        dep_expense_name = dep_expense_acct.account_name

    # Build list of all asset-section accounts for the override dropdowns
    all_asset_accounts = list(
        EntityChartOfAccount.objects.filter(
            entity=fy.entity, is_active=True, section="assets",
        ).order_by("account_code").values_list("account_code", "account_name")
    )
    all_expense_accounts = list(
        EntityChartOfAccount.objects.filter(
            entity=fy.entity, is_active=True, section="expenses",
        ).order_by("account_code").values_list("account_code", "account_name")
    )

    context = {
        "fy": fy,
        "txn": txn,
        "prefill": {
            "asset_name": txn.description or "",
            "category": category_guess,
            "purchase_date": txn.date or "",
            "total_cost": net_amount,
            "addition_cost": net_amount,
            "addition_date": txn.date or "",
            "asset_account_code": asset_acct_code,
            "asset_account_name": asset_acct_name,
            "accum_dep_code": accum_dep_code,
            "accum_dep_name": accum_dep_name,
            "dep_expense_code": dep_expense_code,
            "dep_expense_name": dep_expense_name,
        },
        "all_asset_accounts": all_asset_accounts,
        "all_expense_accounts": all_expense_accounts,
    }
    return render(request, "core/depreciation_add_from_txn.html", context)


# ---------------------------------------------------------------------------
# Post Depreciation to Trial Balance
# ---------------------------------------------------------------------------
@login_required
@require_POST
def depreciation_post_to_tb(request, pk):
    """
    Post the depreciation schedule totals to the trial balance as a journal entry.
    Groups assets by their account mapping and creates per-account journal lines:
      - Dr  Depreciation Expense account(s)
      - Cr  Accumulated Depreciation account(s) (paired with each asset account)
    Assets with explicit account mappings use those; assets without mappings
    fall back to auto-detected global accounts.
    The journal is auto-posted immediately.
    """
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    if fy.is_locked:
        messages.error(request, "Cannot post to a finalised year.")
        return redirect("core:financial_year_detail", pk=pk)

    assets = list(DepreciationAsset.objects.filter(financial_year=fy))
    if not assets:
        messages.warning(request, "No depreciation assets to post.")
        return redirect("core:financial_year_detail", pk=pk)

    # Calculate total business depreciation
    total_depreciation = Decimal("0")
    for asset in assets:
        business_dep = asset.depreciation_amount - asset.private_depreciation
        total_depreciation += business_dep

    if total_depreciation <= 0:
        messages.warning(request, "Total business depreciation is zero. Nothing to post.")
        return redirect("core:financial_year_detail", pk=pk)

    # Check for existing depreciation journal to avoid double-posting
    existing = AdjustingJournal.objects.filter(
        financial_year=fy,
        journal_type=AdjustingJournal.JournalType.DEPRECIATION,
        status=AdjustingJournal.JournalStatus.POSTED,
    ).first()
    if existing:
        messages.warning(
            request,
            f"A depreciation journal ({existing.reference_number}) has already been posted. "
            f"Delete it first if you want to re-post."
        )
        return redirect("core:financial_year_detail", pk=pk)

    # ── Auto-detect fallback accounts from entity COA ──
    # These are used for assets that don't have explicit account mappings.
    fallback_dep_expense_code = ""
    fallback_dep_expense_name = ""
    fallback_accum_dep_code = ""
    fallback_accum_dep_name = ""

    dep_coa = EntityChartOfAccount.objects.filter(
        entity=fy.entity, is_active=True,
        account_name__icontains="depreciation"
    ).exclude(account_name__icontains="accumulated").exclude(
        account_name__icontains="accum"
    ).first()
    if dep_coa:
        fallback_dep_expense_code = dep_coa.account_code
        fallback_dep_expense_name = dep_coa.account_name
    else:
        dep_mapping = ClientAccountMapping.objects.filter(
            entity=fy.entity,
            client_account_name__icontains="depreciation"
        ).exclude(client_account_name__icontains="accumulated").exclude(
            client_account_name__icontains="accum"
        ).first()
        if dep_mapping:
            fallback_dep_expense_code = dep_mapping.client_account_code
            fallback_dep_expense_name = dep_mapping.client_account_name
        else:
            dep_tb = TrialBalanceLine.objects.filter(
                financial_year=fy,
                account_name__icontains="depreciation"
            ).exclude(account_name__icontains="accumulated").exclude(
                account_name__icontains="accum"
            ).first()
            if dep_tb:
                fallback_dep_expense_code = dep_tb.account_code
                fallback_dep_expense_name = dep_tb.account_name

    accum_coa = EntityChartOfAccount.objects.filter(
        entity=fy.entity, is_active=True,
        account_name__icontains="accum"
    ).filter(
        Q(account_name__icontains="depreciation") |
        Q(account_name__icontains="amortis")
    ).first()
    if accum_coa:
        fallback_accum_dep_code = accum_coa.account_code
        fallback_accum_dep_name = accum_coa.account_name
    else:
        accum_mapping = ClientAccountMapping.objects.filter(
            entity=fy.entity,
            client_account_name__icontains="accum"
        ).filter(client_account_name__icontains="depreciation").first()
        if accum_mapping:
            fallback_accum_dep_code = accum_mapping.client_account_code
            fallback_accum_dep_name = accum_mapping.client_account_name
        else:
            accum_tb = TrialBalanceLine.objects.filter(
                financial_year=fy,
                account_name__icontains="accum"
            ).filter(account_name__icontains="depreciation").first()
            if accum_tb:
                fallback_accum_dep_code = accum_tb.account_code
                fallback_accum_dep_name = accum_tb.account_name

    if not fallback_dep_expense_code or not fallback_accum_dep_code:
        missing = []
        if not fallback_dep_expense_code:
            missing.append("Depreciation Expense")
        if not fallback_accum_dep_code:
            missing.append("Accumulated Depreciation")
        messages.error(
            request,
            f"Could not auto-detect fallback account codes for: {', '.join(missing)}. "
            f"Please ensure these accounts exist in the Chart of Accounts, or set "
            f"account mappings on each asset in the depreciation schedule."
        )
        return redirect("core:financial_year_detail", pk=pk)

    # ── Build a lookup of entity asset accounts for auto-pairing ──
    entity_asset_accounts = list(
        EntityChartOfAccount.objects.filter(
            entity=fy.entity, is_active=True, section="assets",
        ).order_by("account_code")
    )

    def _find_paired_accum_dep(asset_code):
        """
        Given an asset account code (e.g. '2870'), find the nearest
        accumulated depreciation account that sits after it in the COA.
        Returns (code, name) or (None, None).
        """
        if not asset_code:
            return None, None
        found_asset = False
        for acct in entity_asset_accounts:
            if acct.account_code == asset_code:
                found_asset = True
                continue
            if found_asset:
                acct_lower = acct.account_name.lower()
                if "accum" in acct_lower or "amortis" in acct_lower:
                    return acct.account_code, acct.account_name
                # If we hit a non-accumulated account, stop looking
                if "less:" not in acct_lower:
                    break
        return None, None

    # ── Group assets by their (dep_expense, accum_dep) account pair ──
    # Each group will become a pair of journal lines (Dr expense, Cr accum dep).
    from collections import defaultdict
    # Key: (dep_expense_code, dep_expense_name, accum_dep_code, accum_dep_name)
    account_groups = defaultdict(Decimal)

    for asset in assets:
        business_dep = asset.depreciation_amount - asset.private_depreciation
        if business_dep <= 0:
            continue

        # Determine the accounts for this asset, with intelligent fallback:
        # 1. Use explicit mapping on the asset (set via form)
        # 2. Try to detect from the source_transaction's confirmed_code
        # 3. Fall back to global auto-detected account

        a_dep_code = asset.dep_expense_code
        a_dep_name = asset.dep_expense_name
        a_accum_code = asset.accum_dep_code
        a_accum_name = asset.accum_dep_name

        # If no explicit accum_dep mapping, try to detect from source transaction
        if not a_accum_code and asset.source_transaction_id:
            txn_code = getattr(asset.source_transaction, 'confirmed_code', '') if hasattr(asset, '_source_txn_cache') else ''
            if not txn_code:
                # Fetch the transaction's confirmed_code
                from review.models import PendingTransaction
                try:
                    txn = PendingTransaction.objects.only('confirmed_code', 'confirmed_name').get(pk=asset.source_transaction_id)
                    txn_code = txn.confirmed_code or ''
                except PendingTransaction.DoesNotExist:
                    txn_code = ''
            if txn_code:
                paired_code, paired_name = _find_paired_accum_dep(txn_code)
                if paired_code:
                    a_accum_code = paired_code
                    a_accum_name = paired_name
                    # Also back-fill the asset record so next time it's explicit
                    asset.accum_dep_code = paired_code
                    asset.accum_dep_name = paired_name
                    if not asset.asset_account_code:
                        asset.asset_account_code = txn_code
                        asset.asset_account_name = getattr(txn, 'confirmed_name', '') if txn_code else ''
                    asset.save(update_fields=[
                        'accum_dep_code', 'accum_dep_name',
                        'asset_account_code', 'asset_account_name',
                    ])

        # If still no explicit mapping, try to detect from asset_account_code
        if not a_accum_code and asset.asset_account_code:
            paired_code, paired_name = _find_paired_accum_dep(asset.asset_account_code)
            if paired_code:
                a_accum_code = paired_code
                a_accum_name = paired_name

        # Final fallback to global accounts
        if not a_dep_code:
            a_dep_code = fallback_dep_expense_code
            a_dep_name = fallback_dep_expense_name or "Depreciation"
        if not a_accum_code:
            a_accum_code = fallback_accum_dep_code
            a_accum_name = fallback_accum_dep_name or "Less: Accumulated depreciation"

        key = (a_dep_code, a_dep_name, a_accum_code, a_accum_name)
        account_groups[key] += business_dep

    # ── Create the journal ──
    journal = AdjustingJournal(
        financial_year=fy,
        journal_type=AdjustingJournal.JournalType.DEPRECIATION,
        status=AdjustingJournal.JournalStatus.DRAFT,
        journal_date=fy.end_date,
        description=f"Depreciation for year ended {fy.end_date.strftime('%d/%m/%Y')}",
        narration=(
            f"Auto-generated from depreciation schedule. "
            f"Total depreciation: ${total_depreciation:,.2f} "
            f"(business portion only, private use excluded). "
            f"Grouped into {len(account_groups)} account pair(s)."
        ),
        total_debit=total_depreciation,
        total_credit=total_depreciation,
        created_by=request.user,
    )
    journal.save()  # Auto-generates reference_number

    # ── Create journal lines per account group ──
    line_number = 0
    line_descriptions = []
    for (dep_code, dep_name, accum_code, accum_name), amount in account_groups.items():
        line_number += 1
        JournalLine.objects.create(
            journal=journal,
            line_number=line_number,
            account_code=dep_code,
            account_name=dep_name,
            description=f"Depreciation charge — {dep_name}",
            debit=amount,
            credit=Decimal("0"),
        )
        line_number += 1
        JournalLine.objects.create(
            journal=journal,
            line_number=line_number,
            account_code=accum_code,
            account_name=accum_name,
            description=f"Accumulated depreciation — {accum_name}",
            debit=Decimal("0"),
            credit=amount,
        )
        line_descriptions.append(
            f"Dr {dep_code} {dep_name} ${amount:,.2f} / Cr {accum_code} {accum_name} ${amount:,.2f}"
        )

    # Auto-post and mark POSTED atomically — no orphan TB lines on failure
    with db_transaction.atomic():
        _post_journal_to_tb(journal, fy)

        journal.status = AdjustingJournal.JournalStatus.POSTED
        journal.posted_by = request.user
        journal.posted_at = timezone.now()
        journal.save(update_fields=["status", "posted_by", "posted_at"])

    _log_action(
        request, "adjustment",
        f"Posted depreciation journal {journal.reference_number} with "
        f"{len(account_groups)} account group(s): " + "; ".join(line_descriptions),
        journal,
    )
    # Auto-trigger risk engine after depreciation post
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "depreciation_post")

    if len(account_groups) == 1:
        # Simple message for single-group posting
        (dep_code, dep_name, accum_code, accum_name) = list(account_groups.keys())[0]
        messages.success(
            request,
            f"Depreciation journal {journal.reference_number} posted: "
            f"Dr {dep_name} ${total_depreciation:,.2f} / "
            f"Cr {accum_name} ${total_depreciation:,.2f}"
        )
    else:
        # Detailed message for multi-group posting
        messages.success(
            request,
            f"Depreciation journal {journal.reference_number} posted with "
            f"{len(account_groups)} account groups, "
            f"total ${total_depreciation:,.2f}. "
            f"See journal details for per-account breakdown."
        )
    return redirect("core:financial_year_detail", pk=pk)


# ---------------------------------------------------------------------------
# Depreciation Schedule PDF Export
# ---------------------------------------------------------------------------
@login_required
def depreciation_pdf(request, pk):
    """
    Generate and download a PDF of the depreciation schedule for a financial year.
    Uses weasyprint to render an HTML template to PDF.
    """
    import weasyprint
    from html import escape as html_escape
    from collections import OrderedDict

    fy = get_financial_year_for_user(request, pk)

    # Gather assets grouped by category (same logic as the main view)
    dep_assets = DepreciationAsset.objects.filter(
        financial_year=fy
    ).order_by('category', 'display_order', 'asset_name')

    dep_categories = OrderedDict()
    dep_total_opening = Decimal('0')
    dep_total_depreciation = Decimal('0')
    dep_total_closing = Decimal('0')
    for asset in dep_assets:
        if asset.category not in dep_categories:
            dep_categories[asset.category] = []
        dep_categories[asset.category].append(asset)
        dep_total_opening += asset.opening_wdv
        dep_total_depreciation += asset.depreciation_amount
        dep_total_closing += asset.closing_wdv

    # Build table rows HTML
    rows_html = ""
    for category, assets in dep_categories.items():
        rows_html += f'<tr class="category-row"><td colspan="7"><strong>{html_escape(category)}</strong></td></tr>\n'
        for asset in assets:
            badges = ""
            if asset.disposal_date:
                badges += ' <span class="badge disposed">Disposed</span>'
            if asset.addition_cost > 0:
                badges += ' <span class="badge addition">Addition</span>'
            addition_cell = f"${asset.addition_cost:,.2f}" if asset.addition_cost > 0 else ""
            rows_html += f"""<tr>
                <td>{html_escape(asset.asset_name)}{badges}</td>
                <td>{html_escape(asset.get_method_display())}</td>
                <td class="r">{asset.rate:.2f}%</td>
                <td class="r">${asset.opening_wdv:,.2f}</td>
                <td class="r">{addition_cell}</td>
                <td class="r">${asset.depreciation_amount:,.2f}</td>
                <td class="r">${asset.closing_wdv:,.2f}</td>
            </tr>\n"""

    now_str = timezone.now().strftime("%d/%m/%Y %H:%M")
    entity_name = html_escape(fy.entity.entity_name)
    fy_label = html_escape(fy.year_label)
    period_label = html_escape(
        f"{fy.start_date.strftime('%d/%m/%Y')} — {fy.end_date.strftime('%d/%m/%Y')}"
    )

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
    @page {{
        size: A4 landscape;
        margin: 15mm 12mm;
        @bottom-right {{ content: "Page " counter(page) " of " counter(pages); font-size: 7pt; color: #999; }}
        @bottom-left {{ content: "Generated {now_str}"; font-size: 7pt; color: #999; }}
    }}
    body {{ font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; font-size: 9pt; color: #333; }}
    .header {{ display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 5mm; border-bottom: 2px solid #1a6b3c; padding-bottom: 3mm; }}
    .header-left h1 {{ font-size: 15pt; margin: 0 0 2px 0; color: #1a1a2e; }}
    .header-left .subtitle {{ font-size: 10pt; color: #555; }}
    .header-right {{ text-align: right; font-size: 8pt; color: #666; }}
    table {{ width: 100%; border-collapse: collapse; margin-bottom: 5mm; }}
    th {{ background: #f0f4f0; font-weight: 700; font-size: 8pt; text-transform: uppercase;
          letter-spacing: 0.3px; color: #444; padding: 4px 6px; border-bottom: 2px solid #ccc;
          text-align: left; white-space: nowrap; }}
    th.r {{ text-align: right; }}
    td {{ padding: 3px 6px; font-size: 8.5pt; border-bottom: 1px solid #eee; vertical-align: middle; }}
    .r {{ text-align: right; }}
    .category-row td {{ background: #f7f7f7; font-size: 9pt; padding: 5px 6px;
                        border-top: 1px solid #ccc; border-bottom: 1px solid #ccc; }}
    .total-row {{ border-top: 2px solid #333; }}
    .total-row td {{ font-weight: 700; padding: 5px 6px; font-size: 9pt; }}
    .badge {{ display: inline-block; font-size: 6.5pt; padding: 1px 4px; border-radius: 3px;
              font-weight: 600; margin-left: 4px; vertical-align: middle; }}
    .disposed {{ background: #fde8e8; color: #c62828; }}
    .addition {{ background: #e3f2fd; color: #1565c0; }}
    .firm {{ font-size: 7.5pt; color: #888; margin-top: 3mm; }}
</style></head><body>
<div class="header">
    <div class="header-left">
        <h1>{entity_name}</h1>
        <div class="subtitle">Depreciation Schedule &mdash; {fy_label}</div>
    </div>
    <div class="header-right">
        Period: {period_label}<br>
        Prepared by: MC &amp; S Pty Ltd
    </div>
</div>
<table>
    <thead>
        <tr>
            <th>Asset</th>
            <th>Method</th>
            <th class="r">Rate %</th>
            <th class="r">Opening WDV</th>
            <th class="r">Additions</th>
            <th class="r">Depreciation</th>
            <th class="r">Closing WDV</th>
        </tr>
    </thead>
    <tbody>
        {rows_html}
        <tr class="total-row">
            <td colspan="3" class="r">Totals</td>
            <td class="r">${dep_total_opening:,.2f}</td>
            <td></td>
            <td class="r">${dep_total_depreciation:,.2f}</td>
            <td class="r">${dep_total_closing:,.2f}</td>
        </tr>
    </tbody>
</table>
<div class="firm">MC &amp; S Pty Ltd &mdash; Confidential &mdash; For client use only</div>
</body></html>"""

    pdf_bytes = weasyprint.HTML(string=html).write_pdf()
    entity_slug = fy.entity.entity_name.replace(" ", "_").replace("/", "-")[:40]
    filename = f"{entity_slug}_Depreciation_Schedule_{fy.year_label}.pdf"
    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Stock Item AJAX endpoints
# ---------------------------------------------------------------------------
@login_required
def stock_add(request, pk):
    """Add a stock item to a financial year."""
    fy = get_financial_year_for_user(request, pk)
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        StockItem.objects.create(
            financial_year=fy,
            item_name=request.POST.get("item_name", ""),
            opening_quantity=Decimal(request.POST.get("opening_quantity", "0") or "0"),
            opening_value=Decimal(request.POST.get("opening_value", "0") or "0"),
            closing_quantity=Decimal(request.POST.get("closing_quantity", "0") or "0"),
            closing_value=Decimal(request.POST.get("closing_value", "0") or "0"),
            notes=request.POST.get("notes", ""),
        )
    except (InvalidOperation, ValueError):
        messages.error(request, "Invalid numeric value provided.")
        return redirect("core:financial_year_detail", pk=pk)
    messages.success(request, "Stock item added.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
def stock_edit(request, pk):
    """Edit a stock item."""
    item = get_object_or_404(StockItem, pk=pk)
    fy_pk = item.financial_year.pk
    get_financial_year_for_user(request, fy_pk)  # IDOR check
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        item.item_name = request.POST.get("item_name", item.item_name)
        item.opening_quantity = Decimal(request.POST.get("opening_quantity", "0") or "0")
        item.opening_value = Decimal(request.POST.get("opening_value", "0") or "0")
        item.closing_quantity = Decimal(request.POST.get("closing_quantity", "0") or "0")
        item.closing_value = Decimal(request.POST.get("closing_value", "0") or "0")
        item.notes = request.POST.get("notes", "")
    except (InvalidOperation, ValueError):
        messages.error(request, "Invalid numeric value provided.")
        return redirect("core:financial_year_detail", pk=fy_pk)
    item.save()
    messages.success(request, f"Stock item '{item.item_name}' updated.")
    return redirect("core:financial_year_detail", pk=fy_pk)


@login_required
@require_POST
def stock_delete(request, pk):
    """Delete a stock item."""
    item = get_object_or_404(StockItem, pk=pk)
    fy_pk = item.financial_year.pk
    get_financial_year_for_user(request, fy_pk)  # IDOR check
    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=fy_pk)
    name = item.item_name
    item.delete()
    messages.success(request, f"Stock item '{name}' deleted.")
    return redirect("core:financial_year_detail", pk=fy_pk)


@login_required
def stock_push_to_tb(request, pk):
    """Push stock values to the trial balance."""
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)
    stock_items = StockItem.objects.filter(financial_year=fy)

    if not stock_items.exists():
        messages.warning(request, "No stock items to push.")
        return redirect("core:financial_year_detail", pk=pk)

    total_opening = sum(s.opening_value for s in stock_items)
    total_closing = sum(s.closing_value for s in stock_items)

    # Remove any existing stock TB lines
    TrialBalanceLine.objects.filter(
        financial_year=fy,
        account_name__in=["Opening Stock", "Closing Stock"],
    ).delete()

    # Create Opening Stock (debit = cost of goods)
    if total_opening > 0:
        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code="5100",
            account_name="Opening Stock",
            debit=total_opening,
            credit=Decimal("0"),
            source='manual_journal',
        )

    # Create Closing Stock (credit to P&L, debit to Balance Sheet)
    if total_closing > 0:
        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code="5200",
            account_name="Closing Stock",
            debit=Decimal("0"),
            credit=total_closing,
            source='manual_journal',
        )
        # Balance sheet current asset
        TrialBalanceLine.objects.create(
            financial_year=fy,
            account_code="1300",
            account_name="Stock on Hand",
            debit=total_closing,
            credit=Decimal("0"),
            source='manual_journal',
        )

    # Mark as pushed
    stock_items.update(pushed_to_tb=True)

    # Auto-trigger risk engine after stock push
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "stock_push")
    messages.success(request, f"Stock pushed to trial balance: Opening ${total_opening}, Closing ${total_closing}.")
    return redirect("core:financial_year_detail", pk=pk)


# ---------------------------------------------------------------------------
# Review → Trial Balance Push
# ---------------------------------------------------------------------------
@login_required
def review_push_to_tb(request, pk):
    """Push confirmed review transactions to the trial balance as journal entries."""
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)
    from review.models import PendingTransaction

    confirmed = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=True,
    )

    if not confirmed.exists():
        messages.warning(request, "No confirmed transactions to push.")
        return redirect("core:financial_year_detail", pk=pk)

    # Push each confirmed transaction individually using the centralised
    # helper so that expense/income + GST + bank contra are always posted.
    tb_count = 0
    for txn in confirmed:
        if not txn.confirmed_code or txn.amount == 0:
            continue
        has_gst = txn.confirmed_gst_amount and txn.confirmed_gst_amount > 0
        if _post_txn_to_tb(txn, fy, has_gst):
            tb_count += 1

    # Auto-trigger risk engine after review push to TB
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "review_push")
    messages.success(request, f"Pushed {tb_count} transactions to trial balance from {confirmed.count()} confirmed.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
def review_approve_transaction(request, pk):
    """Approve a single pending transaction (AJAX).
    Also auto-pushes the transaction to the trial balance so TB and GST/BAS
    update live without needing a separate 'Push to TB' step.
    """
    from review.models import PendingTransaction
    txn = get_object_or_404(PendingTransaction, pk=pk)

    # IDOR check: verify user has access to the linked entity
    if txn.job and txn.job.entity:
        get_entity_for_user(request, txn.job.entity.pk)

    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=405)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    txn.confirmed_code = request.POST.get("confirmed_code", txn.ai_suggested_code)
    txn.confirmed_name = request.POST.get("confirmed_name", txn.ai_suggested_name)
    txn.confirmed_tax_type = request.POST.get("confirmed_tax_type", txn.ai_suggested_tax_type)

    # Handle GST toggle from the review form
    has_gst = request.POST.get("has_gst", "0") == "1"
    if has_gst:
        abs_amount = abs(txn.amount)
        gst_amount = (abs_amount / Decimal("11")).quantize(Decimal("0.01"))
        net_amount = abs_amount - gst_amount
        txn.gst_amount = gst_amount
        txn.net_amount = net_amount
        txn.confirmed_gst_amount = gst_amount
        # Ensure tax type is set correctly for GST
        if not txn.confirmed_tax_type or 'Free' in (txn.confirmed_tax_type or '') or txn.confirmed_tax_type in ('BAS Excluded', 'N-T', ''):
            txn.confirmed_tax_type = 'GST on Expenses' if txn.amount < 0 else 'GST on Income'
    else:
        txn.gst_amount = Decimal("0.00")
        txn.net_amount = abs(txn.amount)
        txn.confirmed_gst_amount = Decimal("0.00")
        # Set GST-free tax type if currently GST
        if txn.confirmed_tax_type in ('GST on Income', 'GST on Expenses'):
            txn.confirmed_tax_type = 'GST Free Expenses' if txn.amount < 0 else 'GST Free Income'

    txn.is_confirmed = True
    txn.save()

    # Auto-push to trial balance immediately
    tb_updated = False
    if txn.job and txn.job.entity:
        # Find the financial year for this entity that covers this transaction
        from datetime import datetime as dt
        entity = txn.job.entity
        fys = FinancialYear.objects.filter(entity=entity, status__in=['draft', 'in_review', 'reopened'])
        target_fy = None
        for fy_candidate in fys:
            # Try to parse the transaction date
            txn_date = None
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d %b %Y"):
                try:
                    txn_date = dt.strptime(txn.date.strip(), fmt).date()
                    break
                except (ValueError, AttributeError):
                    continue
            if txn_date and fy_candidate.start_date <= txn_date <= fy_candidate.end_date:
                target_fy = fy_candidate
                break
        if not target_fy and fys.exists():
            target_fy = fys.order_by('-end_date').first()

        if target_fy and txn.confirmed_code:
            tb_updated = _post_txn_to_tb(txn, target_fy, has_gst)

    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        # Return rich response for live UI updates
        remaining_pending = 0
        remaining_confirmed = 0
        if txn.job and txn.job.entity:
            from review.models import PendingTransaction as PT
            remaining_pending = PT.objects.filter(
                job__entity=txn.job.entity, is_confirmed=False
            ).count()
            remaining_confirmed = PT.objects.filter(
                job__entity=txn.job.entity, is_confirmed=True
            ).count()
        return JsonResponse({
            "status": "success",
            "id": str(txn.pk),
            "tb_updated": tb_updated,
            "has_gst": has_gst,
            "gst_amount": str(txn.confirmed_gst_amount or Decimal("0.00")),
            "net_amount": str(txn.net_amount or abs(txn.amount)),
            "pending_count": remaining_pending,
            "confirmed_count": remaining_confirmed,
            "confirmed_code": txn.confirmed_code,
            "confirmed_name": txn.confirmed_name,
            "confirmed_tax_type": txn.confirmed_tax_type,
            "amount": str(txn.amount),
            "date": txn.date,
            "description": txn.description,
        })

    messages.success(request, f"Transaction approved: {txn.description[:50]}")
    return redirect(request.META.get("HTTP_REFERER", "/"))


@login_required
@require_POST
def review_unconfirm_transaction(request, pk):
    """Unconfirm a previously approved transaction (AJAX).
    Reverses the TB line amounts and resets the transaction to unclassified.
    """
    import traceback
    from django.urls import reverse
    from review.models import PendingTransaction

    try:
        txn = get_object_or_404(PendingTransaction, pk=pk)

        # IDOR check
        if txn.job and txn.job.entity:
            get_entity_for_user(request, txn.job.entity.pk)

        if not request.user.can_do_accounting:
            return JsonResponse({"error": "Permission denied"}, status=403)

        if not txn.is_confirmed:
            return JsonResponse({"error": "Transaction is not confirmed"}, status=400)

        # Reverse the TB line amounts
        if txn.job and txn.job.entity and txn.confirmed_code:
            from datetime import datetime as dt
            entity = txn.job.entity
            fys = FinancialYear.objects.filter(entity=entity, status__in=['draft', 'in_review', 'reopened'])
            target_fy = None
            for fy_candidate in fys:
                txn_date = None
                for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d %b %Y"):
                    try:
                        txn_date = dt.strptime(txn.date.strip(), fmt).date()
                        break
                    except (ValueError, AttributeError):
                        continue
                if txn_date and fy_candidate.start_date <= txn_date <= fy_candidate.end_date:
                    target_fy = fy_candidate
                    break
            if not target_fy and fys.exists():
                target_fy = fys.order_by('-end_date').first()

            if target_fy:
                _reverse_tb_for_transaction(txn, target_fy)

        # Reset the transaction
        txn.is_confirmed = False
        txn.confirmed_code = ''
        txn.confirmed_name = ''
        txn.confirmed_tax_type = ''
        txn.confirmed_gst_amount = Decimal("0.00")
        txn.posted_to_tb = False
        txn.save()

        # Return remaining counts
        remaining_pending = 0
        remaining_confirmed = 0
        if txn.job and txn.job.entity:
            from review.models import PendingTransaction as PT
            remaining_pending = PT.objects.filter(
                job__entity=txn.job.entity, is_confirmed=False
            ).count()
            remaining_confirmed = PT.objects.filter(
                job__entity=txn.job.entity, is_confirmed=True
            ).count()

        # Build URLs for the pending row re-creation
        approve_url = reverse("core:review_approve_transaction", kwargs={"pk": txn.pk})
        unconfirm_url = reverse("core:review_unconfirm_transaction", kwargs={"pk": txn.pk})

        return JsonResponse({
            "status": "success",
            "id": str(txn.pk),
            "pending_count": remaining_pending,
            "confirmed_count": remaining_confirmed,
            "message": f"Transaction unconfirmed: {txn.description[:50]}",
            "date": txn.date,
            "description": txn.description,
            "amount": str(txn.amount),
            "ai_suggested_code": txn.ai_suggested_code or '',
            "ai_suggested_name": txn.ai_suggested_name or '',
            "ai_suggested_tax_type": txn.ai_suggested_tax_type or '',
            "ai_code": txn.ai_suggested_code or '',
            "ai_name": txn.ai_suggested_name or '',
            "ai_tax": txn.ai_suggested_tax_type or '',
            "gst_amount": str(txn.gst_amount or Decimal('0.00')),
            "net_amount": str(txn.net_amount or abs(txn.amount)),
            "ai_confidence": txn.ai_confidence or 0,
            "approve_url": approve_url,
            "unconfirm_url": unconfirm_url,
        })

    except Exception as exc:
        import logging
        logger = logging.getLogger(__name__)
        logger.error("Unconfirm transaction %s failed: %s\n%s", pk, exc, traceback.format_exc())
        return JsonResponse(
            {"status": "error", "error": str(exc)},
            status=500,
        )
@login_required
def review_approve_all(request, pk):
    """Approve all pending transactions for a financial year's entity.
    Also auto-pushes all approved transactions to the trial balance.
    """
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)
    from review.models import PendingTransaction

    pending = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=False,
    )
    count = 0
    tb_count = 0
    for txn in pending:
        if txn.ai_suggested_code:
            txn.confirmed_code = txn.ai_suggested_code
            txn.confirmed_name = txn.ai_suggested_name
            txn.confirmed_tax_type = txn.ai_suggested_tax_type
            txn.is_confirmed = True

            # Preserve GST amounts
            if txn.ai_suggested_tax_type in ('GST on Income', 'GST on Expenses'):
                txn.confirmed_gst_amount = txn.gst_amount
            else:
                txn.confirmed_gst_amount = Decimal('0.00')

            txn.save()
            count += 1

            # Auto-push to trial balance (centralised helper handles
            # expense/income + GST + bank contra in one call)
            has_gst = txn.confirmed_gst_amount and txn.confirmed_gst_amount > 0
            if _post_txn_to_tb(txn, fy, has_gst):
                tb_count += 1

    # Auto-trigger risk engine after bulk approve
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "review_approve_all")

    if request.headers.get("X-Requested-With") == "XMLHttpRequest":
        from review.models import PendingTransaction as PT
        remaining_pending = PT.objects.filter(
            job__entity=fy.entity, is_confirmed=False
        ).count()
        remaining_confirmed = PT.objects.filter(
            job__entity=fy.entity, is_confirmed=True
        ).count()
        return JsonResponse({
            "status": "ok",
            "approved_count": count,
            "tb_count": tb_count,
            "remaining_pending": remaining_pending,
            "remaining_confirmed": remaining_confirmed,
            "message": f"Approved {count} transactions. {tb_count} lines pushed to trial balance.",
        })
    messages.success(request, f"Approved {count} transactions with AI suggestions. {tb_count} lines pushed to trial balance.")
    return redirect("core:financial_year_detail", pk=pk)


@login_required
@require_POST
def review_approve_selected(request, pk):
    """Approve selected pending transactions by ID list (AJAX).
    Each transaction is approved using its current form data (GST settings,
    account code, etc.) and auto-pushed to the trial balance.
    """
    import json
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"status": "error", "message": "Permission denied."}, status=403)

    from review.models import PendingTransaction

    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"status": "error", "message": "Invalid JSON."}, status=400)

    transactions = body.get("transactions", [])
    if not transactions:
        return JsonResponse({"status": "error", "message": "No transactions provided."}, status=400)

    # Build a lookup of per-transaction form data
    txn_data_map = {t["id"]: t for t in transactions if isinstance(t, dict) and "id" in t}
    txn_ids = list(txn_data_map.keys())

    pending = PendingTransaction.objects.filter(
        pk__in=txn_ids,
        job__entity=fy.entity,
        is_confirmed=False,
    )

    count = 0
    tb_count = 0
    approved_ids = []
    from datetime import datetime as dt

    for txn in pending:
        td = txn_data_map.get(str(txn.pk), {})

        txn.confirmed_code = td.get("confirmed_code", txn.ai_suggested_code)
        txn.confirmed_name = td.get("confirmed_name", txn.ai_suggested_name)
        txn.confirmed_tax_type = td.get("confirmed_tax_type", txn.ai_suggested_tax_type)

        has_gst = td.get("has_gst", "0") == "1"
        if has_gst:
            abs_amount = abs(txn.amount)
            cred_pct = Decimal(str(td.get("creditable_percentage", "100")))
            full_gst = (abs_amount / Decimal("11")).quantize(Decimal("0.01"))
            gst_amount = (full_gst * cred_pct / Decimal("100")).quantize(Decimal("0.01"))
            net_amount = abs_amount - gst_amount
            txn.gst_amount = gst_amount
            txn.net_amount = net_amount
            txn.confirmed_gst_amount = gst_amount
            txn.creditable_percentage = cred_pct
            if not txn.confirmed_tax_type or 'Free' in (txn.confirmed_tax_type or '') or txn.confirmed_tax_type in ('BAS Excluded', 'N-T', ''):
                txn.confirmed_tax_type = 'GST on Expenses' if txn.amount < 0 else 'GST on Income'
        else:
            txn.gst_amount = Decimal("0.00")
            txn.net_amount = abs(txn.amount)
            txn.confirmed_gst_amount = Decimal("0.00")
            if txn.confirmed_tax_type in ('GST on Income', 'GST on Expenses'):
                txn.confirmed_tax_type = 'GST Free Expenses' if txn.amount < 0 else 'GST Free Income'

        txn.is_confirmed = True
        txn.save()
        count += 1
        approved_ids.append(str(txn.pk))

        # Auto-push to trial balance (centralised helper handles
        # expense/income + GST + bank contra in one call)
        if _post_txn_to_tb(txn, fy, has_gst):
            tb_count += 1

    # Auto-trigger risk engine after bulk approve
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "review_approve_selected")

    remaining_pending = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=False
    ).count()
    remaining_confirmed = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=True
    ).count()

    return JsonResponse({
        "status": "ok",
        "approved_count": count,
        "tb_count": tb_count,
        "approved_ids": approved_ids,
        "remaining_pending": remaining_pending,
        "remaining_confirmed": remaining_confirmed,
        "message": f"Approved {count} transactions. {tb_count} lines pushed to trial balance.",
    })


# ---------------------------------------------------------------------------
# Bank Account Mapping
# ---------------------------------------------------------------------------
@login_required
@require_POST
def review_bank_account_mapping(request, pk):
    """Create or update a bank account mapping for the entity (AJAX).
    Links a physical bank account (BSB + account number) to a TB account code
    so that double-entry contra-entries can be generated.
    """
    import json
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"status": "error", "message": "Permission denied."}, status=403)

    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        body = request.POST

    bsb = str(body.get('bsb', '')).strip()
    account_number = str(body.get('account_number', '')).strip()
    bank_account_name = str(body.get('bank_account_name', '')).strip()
    tb_account_code = str(body.get('tb_account_code', '')).strip()
    tb_account_name = str(body.get('tb_account_name', '')).strip()
    is_default = body.get('is_default', False)

    if not tb_account_code:
        return JsonResponse({"status": "error", "message": "TB account code is required."}, status=400)

    # Resolve mapped_line_item from the entity CoA
    mapped_item = None
    ecoa = EntityChartOfAccount.objects.filter(
        entity=fy.entity, account_code=tb_account_code, is_active=True,
    ).select_related('maps_to').first()
    if ecoa and ecoa.maps_to:
        mapped_item = ecoa.maps_to
    else:
        # Try master CoA
        coa = ChartOfAccount.objects.filter(
            entity_type=fy.entity.entity_type, account_code=tb_account_code, is_active=True,
        ).select_related('maps_to').first()
        if coa and coa.maps_to:
            mapped_item = coa.maps_to

    # If setting as default, clear other defaults
    if is_default:
        BankAccountMapping.objects.filter(
            entity=fy.entity, is_default=True,
        ).update(is_default=False)

    mapping, created = BankAccountMapping.objects.update_or_create(
        entity=fy.entity,
        bsb=bsb,
        account_number=account_number,
        defaults={
            'bank_account_name': bank_account_name,
            'tb_account_code': tb_account_code,
            'tb_account_name': tb_account_name,
            'mapped_line_item': mapped_item,
            'is_default': is_default,
        },
    )

    return JsonResponse({
        "status": "ok",
        "created": created,
        "mapping_id": str(mapping.pk),
        "tb_account_code": mapping.tb_account_code,
        "tb_account_name": mapping.tb_account_name,
        "message": f"Bank account mapped to {mapping.tb_account_code} — {mapping.tb_account_name}",
    })


# ---------------------------------------------------------------------------
# Recalculate Bank Contra Entries
# ---------------------------------------------------------------------------
def _recalc_bank_contra(fy):
    """
    Internal helper: recalculate bank contra TB line from scratch for all
    confirmed+posted transactions in this financial year.

    Fully idempotent — calculates correct totals from scratch and uses
    update_or_create to atomically SET values (never increment).
    Calling it 1 or 1000 times produces identical results.
    """
    import logging
    logger = logging.getLogger('core.views')

    from review.models import PendingTransaction

    confirmed_txns = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        job__financial_year=fy,
        is_confirmed=True,
        posted_to_tb=True,
    )

    if not confirmed_txns.exists():
        return {"status": "ok", "posted": 0}

    sample_txn = confirmed_txns.select_related('job').first()
    bank_mapping = _get_bank_mapping_for_txn(sample_txn)
    if not bank_mapping:
        logger.warning(
            f"_recalc_bank_contra: No bank mapping for entity {fy.entity.pk} FY {fy.pk}"
        )
        return {"status": "no_mapping", "posted": 0}

    bank_code = bank_mapping.tb_account_code
    bank_name = bank_mapping.tb_account_name

    # Calculate correct totals from scratch
    # Receipts (amount > 0) → debit the bank account
    # Payments (amount < 0) → credit the bank account
    total_debit = Decimal('0')
    total_credit = Decimal('0')
    for txn in confirmed_txns:
        gross = abs(txn.amount)
        if txn.amount > 0:
            total_debit += gross
        elif txn.amount < 0:
            total_credit += gross

    # Find the bank_statement TB line (or create it) and SET values from scratch.
    # Multiple bank_statement lines may exist for the same account code, so we
    # cannot rely on update_or_create alone.  Filter explicitly and consolidate.
    bs_lines = TrialBalanceLine.objects.filter(
        financial_year=fy,
        account_code=bank_code,
        source='bank_statement',
    )
    if bs_lines.count() > 1:
        # Consolidate: keep the first, delete the rest
        keep = bs_lines.first()
        bs_lines.exclude(pk=keep.pk).delete()
        tb_line = keep
        created = False
    elif bs_lines.count() == 1:
        tb_line = bs_lines.first()
        created = False
    else:
        tb_line = TrialBalanceLine(
            financial_year=fy,
            account_code=bank_code,
            source='bank_statement',
        )
        created = True

    tb_line.account_name = bank_name
    tb_line.debit = total_debit
    tb_line.credit = total_credit
    tb_line.closing_balance = total_debit - total_credit
    tb_line.tax_type = ""
    tb_line.save()

    logger.info(
        f"Recalculated bank contra for entity {fy.entity.pk} FY {fy.pk}: "
        f"SET Dr={total_debit}, Cr={total_credit} on {bank_code} ({bank_name})"
    )

    return {
        "status": "ok",
        "posted_debit": str(total_debit),
        "posted_credit": str(total_credit),
        "bank_code": bank_code,
        "bank_name": bank_name,
    }


@login_required
@require_POST
def recalculate_bank_contra_entries(request, pk):
    """HTTP endpoint for Recalc Contra button — delegates to _recalc_bank_contra."""
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"status": "error", "message": "Permission denied."}, status=403)

    result = _recalc_bank_contra(fy)

    if result.get("status") == "no_mapping":
        return JsonResponse({
            "status": "error",
            "message": (
                "No bank account mapping found for this entity. "
                "Please configure a bank account mapping first."
            ),
        }, status=400)

    if result.get("posted", 0) == 0 and "posted_debit" not in result:
        if not result.get("bank_code"):
            return JsonResponse({
                "status": "ok",
                "message": "No confirmed transactions found.",
                "posted": 0,
            })
        return JsonResponse({
            "status": "ok",
            "message": "Bank contra entries are already balanced. No adjustment needed.",
            "posted": 0,
            "bank_code": result["bank_code"],
            "bank_name": result["bank_name"],
        })

    return JsonResponse({
        "status": "ok",
        "message": (
            f"Bank contra entries recalculated. "
            f"Posted ${Decimal(result['posted_debit']):,.2f} Dr / "
            f"${Decimal(result['posted_credit']):,.2f} Cr "
            f"to {result['bank_code']} — {result['bank_name']}."
        ),
        "posted_debit": result["posted_debit"],
        "posted_credit": result["posted_credit"],
        "bank_code": result["bank_code"],
        "bank_name": result["bank_name"],
    })


@login_required
def bank_statement_template_download(request):
    """Serve the bank statement import template .xlsx file."""
    import os
    from django.conf import settings as _settings
    template_path = os.path.join(_settings.BASE_DIR, "static", "bank_statement_template.xlsx")
    if not os.path.exists(template_path):
        from django.http import Http404
        raise Http404("Template file not found.")
    with open(template_path, "rb") as f:
        response = HttpResponse(
            f.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = 'attachment; filename="StatementHub_Bank_Statement_Template.xlsx"'
        return response


@login_required
@require_POST
def review_validate_opening_balance(request, pk):
    """Validate that the opening balance from a bank statement import matches
    the current trial balance closing balance for the mapped bank account.

    Returns JSON with:
    - tb_balance: the current netted balance in the TB for this account
    - import_opening: the opening balance from the import
    - matches: whether they match (within $0.01 tolerance)
    - account_exists: whether the account exists in the TB at all
    """
    import json as _json
    fy = get_financial_year_for_user(request, pk)

    try:
        body = _json.loads(request.body)
    except (ValueError, _json.JSONDecodeError):
        return JsonResponse({"status": "error", "message": "Invalid JSON."}, status=400)

    tb_account_code = str(body.get('tb_account_code', '')).strip()
    import_opening_balance = body.get('opening_balance', 0)

    if not tb_account_code:
        return JsonResponse({"status": "error", "message": "No account code provided."}, status=400)

    try:
        import_opening = Decimal(str(import_opening_balance)).quantize(Decimal('0.01'))
    except Exception:
        import_opening = Decimal('0')

    # Find all TB lines for this account code in this financial year
    tb_lines = TrialBalanceLine.objects.filter(
        financial_year=fy,
        account_code=tb_account_code,
    )

    if not tb_lines.exists():
        # Account doesn't exist in TB yet
        return JsonResponse({
            "status": "ok",
            "account_exists": False,
            "tb_balance": "0.00",
            "import_opening": str(import_opening),
            "matches": True,
            "message": f"Account {tb_account_code} does not exist in the trial balance yet. "
                       f"It will be created when transactions are posted.",
        })

    # Calculate the netted balance: sum of (opening_balance + debit - credit) across all lines
    # For bank accounts (assets), the balance = sum of debits - sum of credits + opening
    from django.db.models import Sum
    agg = tb_lines.aggregate(
        total_opening=Sum('opening_balance'),
        total_debit=Sum('debit'),
        total_credit=Sum('credit'),
    )
    total_opening = agg['total_opening'] or Decimal('0')
    total_debit = agg['total_debit'] or Decimal('0')
    total_credit = agg['total_credit'] or Decimal('0')

    # The current TB closing balance for this account
    tb_closing = (total_opening + total_debit - total_credit).quantize(Decimal('0.01'))

    # Check if the import's opening balance matches the TB closing balance
    # The import opening balance should match the TB's current closing balance
    # because the import adds new movements on top of the existing balance
    difference = abs(tb_closing - import_opening)
    matches = difference <= Decimal('0.01')

    response_data = {
        "status": "ok",
        "account_exists": True,
        "tb_balance": str(tb_closing),
        "import_opening": str(import_opening),
        "difference": str(difference),
        "matches": matches,
    }

    if not matches:
        response_data["message"] = (
            f"Opening balance mismatch: The trial balance shows a closing balance of "
            f"${tb_closing:,.2f} for account {tb_account_code}, but the bank statement "
            f"opening balance is ${import_opening:,.2f}. Difference: ${difference:,.2f}. "
            f"Please check for missing bank statements or adjust the trial balance."
        )
    else:
        response_data["message"] = (
            f"Opening balance matches. TB balance: ${tb_closing:,.2f}, "
            f"Import opening: ${import_opening:,.2f}."
        )

    return JsonResponse(response_data)


@login_required
@require_POST
def review_post_opening_balance(request, pk):
    """Create and auto-post an opening balance journal entry.
    Accepts a list of balance lines (account_code, account_name, debit, credit).
    Creates an AdjustingJournal with type 'general', auto-posts it, and
    applies the lines to the Trial Balance.
    """
    import json as _json
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"status": "error", "message": "Permission denied."}, status=403)
    if fy.is_locked:
        return JsonResponse({"status": "error", "message": "Financial year is locked."}, status=400)

    try:
        body = _json.loads(request.body)
    except (ValueError, _json.JSONDecodeError):
        return JsonResponse({"status": "error", "message": "Invalid JSON."}, status=400)

    lines_data = body.get('lines', [])
    if not lines_data:
        return JsonResponse({"status": "error", "message": "No lines provided."}, status=400)

    # Validate lines and calculate totals
    total_debit = Decimal('0')
    total_credit = Decimal('0')
    validated_lines = []
    for i, line in enumerate(lines_data):
        code = str(line.get('account_code', '')).strip()
        name = str(line.get('account_name', '')).strip()
        try:
            dr = Decimal(str(line.get('debit', 0) or 0)).quantize(Decimal('0.01'))
            cr = Decimal(str(line.get('credit', 0) or 0)).quantize(Decimal('0.01'))
        except Exception:
            return JsonResponse({"status": "error", "message": f"Invalid amount on line {i+1}."}, status=400)
        if not code:
            return JsonResponse({"status": "error", "message": f"Account code missing on line {i+1}."}, status=400)
        if dr == 0 and cr == 0:
            continue  # skip zero lines
        total_debit += dr
        total_credit += cr
        validated_lines.append({
            'account_code': code,
            'account_name': name,
            'debit': dr,
            'credit': cr,
        })

    if not validated_lines:
        return JsonResponse({"status": "error", "message": "All lines are zero."}, status=400)

    if total_debit != total_credit:
        diff = abs(total_debit - total_credit)
        return JsonResponse({
            "status": "error",
            "message": f"Journal does not balance. Debit: {total_debit}, Credit: {total_credit}, Difference: {diff}",
        }, status=400)

    # Create the journal
    journal = AdjustingJournal(
        financial_year=fy,
        journal_type=AdjustingJournal.JournalType.GENERAL,
        journal_date=fy.start_date,
        description='Opening balance — brought forward from prior records',
        narration='Auto-generated opening balance journal from bank statement review.',
        total_debit=total_debit,
        total_credit=total_credit,
        created_by=request.user,
    )
    journal.save()

    # Create journal lines
    journal_lines = []
    for i, line in enumerate(validated_lines):
        journal_lines.append(JournalLine(
            journal=journal,
            line_number=i + 1,
            account_code=line['account_code'],
            account_name=line['account_name'],
            debit=line['debit'],
            credit=line['credit'],
        ))
    JournalLine.objects.bulk_create(journal_lines)

    # Auto-post and mark POSTED atomically — no orphan TB lines on failure
    with db_transaction.atomic():
        _post_journal_to_tb(journal, fy)

        journal.status = AdjustingJournal.JournalStatus.POSTED
        journal.posted_by = request.user
        journal.posted_at = timezone.now()
        journal.save(update_fields=['status', 'posted_by', 'posted_at'])

    _log_action(request, 'adjustment', f'Posted opening balance journal {journal.reference_number}', journal)

    # Trigger risk recalc
    try:
        from core.signals import trigger_risk_recalc
        trigger_risk_recalc(fy, 'journal_posted')
    except Exception:
        pass

    return JsonResponse({
        "status": "ok",
        "journal_id": str(journal.pk),
        "reference": journal.reference_number,
        "total_debit": str(total_debit),
        "total_credit": str(total_credit),
        "message": f"Opening balance journal {journal.reference_number} posted successfully.",
    })


# ---------------------------------------------------------------------------
# Delete Transactions
# ----------------------------------------------------------------------------
@login_required
@require_POST
def review_delete_transaction(request, pk, txn_pk):
    """Delete a single imported bank transaction (AJAX).
    If the transaction was confirmed and pushed to TB, reverses the TB impact first.
    """
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    from review.models import PendingTransaction
    txn = get_object_or_404(PendingTransaction, pk=txn_pk, job__entity=fy.entity)

    # If confirmed, reverse TB impact before deleting
    if txn.is_confirmed and txn.confirmed_code:
        _reverse_tb_for_transaction(txn, fy)

    desc_preview = txn.description[:50]
    # Also delete any split children
    split_children = PendingTransaction.objects.filter(split_parent=txn)
    for child in split_children:
        if child.is_confirmed and child.confirmed_code:
            _reverse_tb_for_transaction(child, fy)
        child.delete()

    txn.delete()

    # Return updated counts
    remaining_pending = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=False
    ).count()
    remaining_confirmed = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=True
    ).count()

    return JsonResponse({
        "status": "ok",
        "message": f"Deleted: {desc_preview}",
        "remaining_pending": remaining_pending,
        "remaining_confirmed": remaining_confirmed,
    })


@login_required
@require_POST
def review_delete_all_transactions(request, pk):
    """Delete all imported bank transactions for a financial year's entity (AJAX).
    Supports filtering by status: 'pending', 'confirmed', or 'all'.
    Reverses TB impact for any confirmed transactions before deleting.
    """
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    import json
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        body = {}
    scope = body.get('scope', request.POST.get('scope', 'all'))

    from review.models import PendingTransaction
    qs = PendingTransaction.objects.filter(job__entity=fy.entity)

    if scope == 'pending':
        qs = qs.filter(is_confirmed=False)
    elif scope == 'confirmed':
        qs = qs.filter(is_confirmed=True)
    # else 'all' — no filter

    # Reverse TB for confirmed transactions
    confirmed_txns = qs.filter(is_confirmed=True, confirmed_code__gt='')
    for txn in confirmed_txns:
        _reverse_tb_for_transaction(txn, fy)

    count = qs.count()
    qs.delete()

    # Return updated counts
    remaining_pending = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=False
    ).count()
    remaining_confirmed = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=True
    ).count()

    return JsonResponse({
        "status": "ok",
        "message": f"Deleted {count} {scope} transactions.",
        "deleted_count": count,
        "remaining_pending": remaining_pending,
        "remaining_confirmed": remaining_confirmed,
    })


@login_required
@require_POST
def review_delete_selected_transactions(request, pk):
    """Delete selected imported bank transactions by ID list (AJAX).
    Reverses TB impact for any confirmed transactions before deleting.
    """
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    import json
    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        body = {}
    txn_ids = body.get('transaction_ids', [])
    if not txn_ids:
        return JsonResponse({"error": "No transactions selected"}, status=400)

    from review.models import PendingTransaction
    qs = PendingTransaction.objects.filter(pk__in=txn_ids, job__entity=fy.entity)

    # Reverse TB for confirmed transactions
    confirmed_txns = qs.filter(is_confirmed=True, confirmed_code__gt='')
    for txn in confirmed_txns:
        _reverse_tb_for_transaction(txn, fy)

    count = qs.count()
    qs.delete()

    # Return updated counts
    remaining_pending = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=False
    ).count()
    remaining_confirmed = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=True
    ).count()

    return JsonResponse({
        "status": "ok",
        "message": f"Deleted {count} selected transactions.",
        "deleted_count": count,
        "remaining_pending": remaining_pending,
        "remaining_confirmed": remaining_confirmed,
    })


def _reverse_tb_for_transaction(txn, fy):
    """Reverse the trial balance impact of a single confirmed transaction
    by decrementing the accumulated values on the original TB lines.

    _post_txn_to_tb() accumulates debit/credit/closing_balance directly
    into existing TB lines.  The reversal MUST undo that accumulation
    (not create separate reversal lines) to prevent closing_balance drift
    after unconfirm→re-approve cycles.
    """
    from django.db import transaction as db_transaction

    if not txn.confirmed_code:
        return

    with db_transaction.atomic():
        # --- 1. Reverse the main account ---
        has_gst = txn.confirmed_gst_amount and txn.confirmed_gst_amount > 0
        net_for_tb = txn.net_amount if has_gst else abs(txn.amount)

        tb_line = TrialBalanceLine.objects.filter(
            financial_year=fy,
            account_code=txn.confirmed_code,
            is_adjustment=False,
        ).first()
        if not tb_line:
            tb_line = TrialBalanceLine.objects.filter(
                financial_year=fy,
                account_code=txn.confirmed_code,
            ).first()
        if tb_line:
            if txn.amount < 0:
                # Original posting debited expense — reverse
                tb_line.debit = max(Decimal("0"), tb_line.debit - net_for_tb)
                tb_line.closing_balance -= net_for_tb
            else:
                # Original posting credited income — reverse
                tb_line.credit = max(Decimal("0"), tb_line.credit - net_for_tb)
                tb_line.closing_balance += net_for_tb
            tb_line.save(update_fields=["debit", "credit", "closing_balance"])

        # --- 2. Reverse the GST clearing account ---
        if has_gst:
            gst_amt = txn.confirmed_gst_amount
            gst_line = TrialBalanceLine.objects.filter(
                financial_year=fy,
                account_code="3380",
                is_adjustment=False,
            ).first()
            if not gst_line:
                gst_line = TrialBalanceLine.objects.filter(
                    financial_year=fy,
                    account_code="3380",
                ).first()
            if gst_line:
                if txn.amount > 0:
                    # Original: receipt credited 3380 — reverse
                    gst_line.credit = max(Decimal("0"), gst_line.credit - gst_amt)
                    gst_line.closing_balance += gst_amt
                else:
                    # Original: payment debited 3380 — reverse
                    gst_line.debit = max(Decimal("0"), gst_line.debit - gst_amt)
                    gst_line.closing_balance -= gst_amt
                gst_line.save(update_fields=["debit", "credit", "closing_balance"])

        # --- 3. Reverse the bank contra-entry ---
        _reverse_bank_contra_entry(txn, fy)

        # Reset the posted_to_tb flag so the transaction can be re-posted
        if hasattr(txn, 'posted_to_tb'):
            txn.posted_to_tb = False
            txn.save(update_fields=['posted_to_tb'])


# ---------------------------------------------------------------------------
# Bulk Edit Transactions (Assign Account)
# ---------------------------------------------------------------------------
@login_required
@require_POST
def review_bulk_edit_transactions(request, pk):
    """Bulk assign an account code to selected transactions (AJAX).
    Optionally approves them and pushes to TB in one step.
    """
    import json
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"error": "Permission denied"}, status=403)

    try:
        body = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    txn_ids = body.get('transaction_ids', [])
    account_code = body.get('account_code', '').strip()
    account_name = body.get('account_name', '').strip()
    tax_code = body.get('tax_code', '').strip()
    approve = body.get('approve', False)

    if not txn_ids:
        return JsonResponse({"error": "No transactions selected"}, status=400)
    if not account_code:
        return JsonResponse({"error": "No account code specified"}, status=400)

    from review.models import PendingTransaction
    txns = PendingTransaction.objects.filter(pk__in=txn_ids, job__entity=fy.entity, is_confirmed=False)

    # Map tax_code to tax_type for TB posting
    TAX_CODE_MAP = {
        'GST': {'expense': 'GST on Expenses', 'income': 'GST on Income'},
        'INP': {'expense': 'GST on Expenses', 'income': 'GST on Income'},
        'FRE': {'expense': 'GST Free Expenses', 'income': 'GST Free Income'},
        'ITS': {'expense': 'Input Taxed', 'income': 'Input Taxed'},
        'N-T': {'expense': 'N-T', 'income': 'N-T'},
        'CAP': {'expense': 'GST on Expenses', 'income': 'GST on Income'},
        'GNR': {'expense': 'GST on Expenses', 'income': 'GST on Income'},
        'ADS': {'expense': 'GST on Expenses', 'income': 'GST on Income'},
    }

    count = 0
    tb_count = 0
    for txn in txns:
        direction = 'expense' if txn.amount < 0 else 'income'
        tax_type = ''
        if tax_code and tax_code in TAX_CODE_MAP:
            tax_type = TAX_CODE_MAP[tax_code][direction]
        elif txn.ai_suggested_tax_type:
            tax_type = txn.ai_suggested_tax_type

        # Determine GST
        has_gst = tax_code in ('GST', 'INP', 'CAP', 'GNR', 'ADS')
        abs_amount = abs(txn.amount)
        if has_gst:
            gst_amount = (abs_amount / Decimal('11')).quantize(Decimal('0.01'))
            net_amount = abs_amount - gst_amount
        else:
            gst_amount = Decimal('0.00')
            net_amount = abs_amount

        # Auto-set gst_treatment based on the account's tax_code
        TAX_CODE_TO_GST_TREATMENT = {
            'GST': 'gst', 'INP': 'gst', 'CAP': 'gst', 'GNR': 'gst', 'ADS': 'gst',
            'FRE': 'gst_free', 'FOA': 'gst_free',
            'ITS': 'input_taxed', 'IOA': 'input_taxed',
            'N-T': 'out_of_scope',
        }
        gst_treatment = TAX_CODE_TO_GST_TREATMENT.get(tax_code.upper(), '') if tax_code else ''
        if gst_treatment:
            txn.gst_treatment = gst_treatment

        # Update the AI suggestion fields so the row displays correctly
        txn.ai_suggested_code = account_code
        txn.ai_suggested_name = account_name
        txn.ai_suggested_tax_type = tax_type
        txn.gst_amount = gst_amount
        txn.net_amount = net_amount

        if approve:
            txn.confirmed_code = account_code
            txn.confirmed_name = account_name
            txn.confirmed_tax_type = tax_type
            txn.confirmed_gst_amount = gst_amount
            txn.is_confirmed = True
            txn.save()

            # Auto-push to trial balance (centralised helper handles
            # expense/income + GST + bank contra in one call)
            if _post_txn_to_tb(txn, fy, has_gst):
                tb_count += 1
        else:
            txn.save()

        count += 1

    # Auto-trigger risk engine after bulk edit
    if approve:
        from core.signals import trigger_risk_recalc
        trigger_risk_recalc(fy, "review_bulk_edit")

    remaining_pending = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=False
    ).count()
    remaining_confirmed = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=True
    ).count()

    action = 'assigned and approved' if approve else 'assigned'
    return JsonResponse({
        "status": "ok",
        "message": f"Account {account_code} {action} to {count} transactions." + (f" {tb_count} pushed to TB." if approve else ""),
        "count": count,
        "tb_count": tb_count,
        "remaining_pending": remaining_pending,
        "remaining_confirmed": remaining_confirmed,
    })


# ---------------------------------------------------------------------------
# Activity Log / Notifications
# ---------------------------------------------------------------------------
@login_required
@require_POST
def mark_notification_read(request, pk):
    """Mark a single activity log entry as read."""
    activity = get_object_or_404(ActivityLog, pk=pk, user=request.user)
    activity.is_read = True
    activity.save()
    return JsonResponse({"status": "ok"})


@login_required
@require_POST
def mark_all_notifications_read(request):
    """Mark all unread notifications as read for the current user."""
    ActivityLog.objects.filter(is_read=False, user=request.user).update(is_read=True)
    return JsonResponse({"status": "ok"})


@login_required
def risk_badge_api(request, pk):
    """JSON API for polling risk flag badge count and engine status.
    Returns open flag count, severity breakdown, and whether the engine
    is currently running (debounce pending).
    """
    fy = get_financial_year_for_user(request, pk)
    from django.core.cache import cache
    flags = fy.risk_flags.filter(status='open')
    total = flags.count()
    critical = flags.filter(severity='CRITICAL').count()
    high = flags.filter(severity='HIGH').count()
    medium = flags.filter(severity='MEDIUM').count()
    low = flags.filter(severity='LOW').count()

    # Count distinct flagged accounts (accounts with open flags)
    flagged_codes = set()
    for f in fy.risk_flags.filter(status__in=['open', 'reviewed']):
        for acc in (f.affected_accounts or []):
            code = acc.get('account_code', '') if isinstance(acc, dict) else str(acc)
            if code:
                flagged_codes.add(code)
    flagged_account_count = len(flagged_codes)

    # Check if the engine is currently in debounce (pending run)
    engine_pending = cache.get(f'risk_engine_pending_{fy.pk}', False)
    # Check last run timestamp
    last_run = cache.get(f'risk_engine_last_run_{fy.pk}')

    return JsonResponse({
        'open_count': total,
        'flagged_account_count': flagged_account_count,
        'critical': critical,
        'high': high,
        'medium': medium,
        'low': low,
        'engine_pending': engine_pending,
        'last_run': last_run.isoformat() if last_run else None,
        'status': fy.status,
    })


@login_required
def notifications_api(request):
    """Return recent unread notifications as JSON for polling (scoped to current user)."""
    activities = (
        ActivityLog.objects.filter(is_read=False, user=request.user)
        .order_by("-created_at")[:10]
    )
    data = []
    for a in activities:
        data.append({
            "id": str(a.pk),
            "event_type": a.event_type,
            "title": a.title,
            "description": a.description,
            "url": a.url,
            "created_at": a.created_at.isoformat(),
        })
    return JsonResponse({"unread_count": ActivityLog.objects.filter(is_read=False, user=request.user).count(), "items": data})


# ============================================================
# Bulk Client Actions (Delete / Archive)
# ============================================================

@login_required
def client_bulk_action(request):
    """Bulk delete/archive actions on selected entities (top-level)."""
    if request.method != "POST":
        return redirect("core:entity_list")

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_list")

    action = request.POST.get("bulk_action")
    entity_ids = request.POST.getlist("entity_ids")

    if not entity_ids:
        messages.warning(request, "No entities selected.")
        return redirect("core:entity_list")

    entities = Entity.objects.filter(pk__in=entity_ids)

    # Ownership check: non-senior users can only act on their own entities
    if not request.user.is_senior:
        entities = entities.filter(
            Q(assigned_accountant=request.user) |
            Q(client__assigned_accountant=request.user)
        )

    if action == "archive":
        count = entities.update(is_archived=True)
        messages.success(request, f"Archived {count} entity/entities.")
        _log_action(request, "update", f"Archived {count} entity/entities")
    elif action == "unarchive":
        count = entities.update(is_archived=False)
        messages.success(request, f"Unarchived {count} entity/entities.")
        _log_action(request, "update", f"Unarchived {count} entity/entities")
    elif action == "delete":
        count = entities.count()
        for e in entities:
            _log_action(request, "delete", f"Deleted entity: {e.entity_name}")
        entities.delete()
        messages.success(request, f"Deleted {count} entity/entities and all associated data.")
    else:
        messages.error(request, "Invalid action.")

    return redirect("core:entity_list")


# Alias for backward compatibility
entity_bulk_action = client_bulk_action


# ============================================================
# Entity-level HandiLedger Import
# ============================================================

@login_required
def entity_import_handiledger(request, pk):
    """Import HandiLedger ZIP for a specific entity."""
    from .access_ledger_import import import_access_ledger_zip

    entity = get_entity_for_user(request, pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=pk)

    result = None

    if request.method == "POST":
        zip_file = request.FILES.get("zip_file")
        if not zip_file:
            messages.error(request, "Please select a ZIP file.")
        elif not zip_file.name.lower().endswith(".zip"):
            messages.error(request, "File must be a .zip file.")
        else:
            replace = request.POST.get("replace_existing") == "1"
            import_finalised = request.POST.get("import_as_finalised") == "1"
            try:
                result = import_access_ledger_zip(
                    zip_file,
                    client=entity.client if entity.client else None,
                    entity=entity,
                    replace_existing=replace,
                    import_as_finalised=import_finalised,
                )
                if result["errors"]:
                    messages.warning(
                        request,
                        f"Import completed with {len(result['errors'])} error(s)."
                    )
                else:
                    status_note = " (all years finalised — ready to roll over)" if import_finalised else ""
                    messages.success(
                        request,
                        f"Successfully imported: "
                        f"{result['years_imported']} years, "
                        f"{result['total_tb_lines']} TB lines, "
                        f"{result['total_dep_assets']} depreciation assets."
                        f"{status_note}"
                    )
                if result.get("warnings"):
                    for w in result["warnings"]:
                        messages.info(request, w)
                _log_action(
                    request, "import",
                    f"Imported HandiLedger ZIP for {entity.entity_name}: "
                    f"{result['years_imported']} years"
                    f"{' (import as finalised)' if import_finalised else ''}",
                    entity,
                )
            except Exception as e:
                messages.error(request, "Import failed. Please check the file format and try again.")

    return render(request, "core/entity_import_handiledger.html", {
        "entity": entity,
        "result": result,
    })


# ============================================================
# Delete Unfinalised FY Data
# ============================================================

@login_required
def delete_unfinalised_fy(request, pk):
    """Delete all unfinalised financial years and their data for an entity."""
    entity = get_entity_for_user(request, pk)

    if request.method != "POST":
        return redirect("core:entity_detail", pk=pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:entity_detail", pk=pk)

    unfinalised = entity.financial_years.exclude(status="finalised")
    count = unfinalised.count()
    if count == 0:
        messages.info(request, "No unfinalised financial years to delete.")
        return redirect("core:entity_detail", pk=pk)

    for fy in unfinalised:
        _log_action(request, "delete", f"Deleted unfinalised FY: {fy.year_label} for {entity.entity_name}", fy)

    # ReviewJob.financial_year FK has on_delete=CASCADE, so associated
    # ReviewJobs (and their PendingTransactions) are deleted automatically.
    unfinalised.delete()
    messages.success(request, f"Deleted {count} unfinalised financial year(s) and all associated data.")
    return redirect("core:entity_detail", pk=pk)


# ============================================================
# HTMX: Update TB Line Mapping (clickable AI suggestion)
# ============================================================

@login_required
def htmx_update_tb_mapping(request, pk):
    """HTMX endpoint to update the mapping of a trial balance line via dropdown."""
    line = get_object_or_404(TrialBalanceLine, pk=pk)
    get_financial_year_for_user(request, line.financial_year.pk)  # IDOR check
    if not request.user.can_do_accounting:
        return HttpResponse("Permission denied", status=403)

    if request.method == "POST":
        mapping_id = request.POST.get("mapped_line_item")
        if mapping_id:
            try:
                mapping = AccountMapping.objects.get(pk=mapping_id)
                line.mapped_line_item = mapping
                line.save()

                # Also save to client account mapping for reuse
                ClientAccountMapping.objects.update_or_create(
                    entity=line.financial_year.entity,
                    client_account_code=line.account_code,
                    defaults={
                        "client_account_name": line.account_name,
                        "mapped_line_item": mapping,
                    },
                )
            except AccountMapping.DoesNotExist:
                pass

    # Return the updated row
    entity_type = line.financial_year.entity.entity_type
    coa_items = ChartOfAccount.objects.filter(
        entity_type=entity_type, is_active=True
    ).select_related("maps_to").order_by("section", "account_code")

    return render(request, "partials/tb_line_row_review.html", {
        "line": line,
        "coa_items": coa_items,
    })


# ============================================================
# Chart of Accounts API for searchable dropdown
# ============================================================

@login_required
def coa_search_api(request):
    """JSON API for searching chart of accounts — used by the review tab dropdown."""
    entity_type = request.GET.get("entity_type", "company")
    q = request.GET.get("q", "")

    qs = ChartOfAccount.objects.filter(
        entity_type=entity_type, is_active=True
    ).select_related("maps_to").order_by("section", "account_code")

    if q:
        qs = qs.filter(
            Q(account_code__icontains=q) | Q(account_name__icontains=q)
        )

    items = []
    for a in qs[:200]:
        items.append({
            "id": str(a.maps_to.pk) if a.maps_to else "",
            "code": a.account_code,
            "name": a.account_name,
            "section": a.get_section_display(),
            "section_value": a.section,
            "classification": a.classification or "",
            "tax_code": a.tax_code or "",
            "maps_to_id": str(a.maps_to.pk) if a.maps_to else "",
            "mapping_label": a.maps_to.line_item_label if a.maps_to else "Unmapped",
        })
    return JsonResponse({"items": items})


@login_required
def entity_coa_search_api(request, pk):
    """JSON API for searching an entity's own chart of accounts.
    Used by the bank statement review Change Account dropdown.
    Searches EntityChartOfAccount by code or name.
    Includes the netted trial balance amount for each account so
    accountants can see current balances when choosing between
    similar accounts.
    """
    fy = get_financial_year_for_user(request, pk)
    q = request.GET.get("q", "")

    qs = EntityChartOfAccount.objects.filter(
        entity=fy.entity, is_active=True
    ).select_related("maps_to").order_by("section", "account_code")

    if q:
        qs = qs.filter(
            Q(account_code__icontains=q) | Q(account_name__icontains=q)
        )

    # Build a lookup of netted TB balances per account_code.
    # Multiple TrialBalanceLine rows may exist for the same code
    # (original + adjustments), so we aggregate them.
    tb_balances = {}
    tb_lines = fy.trial_balance_lines.all()
    for line in tb_lines:
        code = line.account_code
        cb = line.closing_balance if line.closing_balance else Decimal('0')
        if cb != 0:
            tb_balances[code] = tb_balances.get(code, Decimal('0')) + cb
        else:
            dr = line.debit if line.debit else Decimal('0')
            cr = line.credit if line.credit else Decimal('0')
            tb_balances[code] = tb_balances.get(code, Decimal('0')) + (dr - cr)

    # Map sections to P&L or B/S for the frontend
    _PL_SECTIONS = {'revenue', 'cost_of_sales', 'expenses'}
    _BS_SECTIONS = {'assets', 'liabilities', 'equity', 'capital_accounts'}

    items = []
    for a in qs[:200]:
        balance = tb_balances.get(a.account_code)
        if a.section in _PL_SECTIONS:
            stmt_type = 'P&L'
        elif a.section in _BS_SECTIONS:
            stmt_type = 'B/S'
        else:
            stmt_type = ''
        items.append({
            "id": str(a.pk),
            "code": a.account_code,
            "name": a.account_name,
            "section": a.get_section_display(),
            "section_value": a.section,
            "stmt_type": stmt_type,
            "classification": a.classification or "",
            "tax_code": a.tax_code or "",
            "maps_to_id": str(a.maps_to.pk) if a.maps_to else "",
            "mapping_label": a.maps_to.line_item_label if a.maps_to else "Unmapped",
            "tb_balance": str(balance.quantize(Decimal('0.01'))) if balance is not None else None,
        })
    return JsonResponse({"items": items})


# ============================================================
# XRM Pull (Xero Practice Manager)
# ============================================================

@login_required
def xrm_search(request, pk):
    """
    Search XPM for clients matching this entity's name.
    Returns a list of potential matches for the user to confirm.
    """
    entity = get_entity_for_user(request, pk)

    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)

    if not request.user.can_do_accounting:
        return JsonResponse({"status": "error", "message": "Permission denied"}, status=403)

    from integrations.models import XPMConnection
    from integrations.xpm_sync import _ensure_valid_token, _xpm_get, _xml_text

    connection = XPMConnection.objects.filter(status="active").first()
    if not connection:
        return JsonResponse({
            "status": "error",
            "message": "No active XPM connection. Go to Integrations > XPM to connect first."
        })

    if not _ensure_valid_token(connection):
        return JsonResponse({
            "status": "error",
            "message": "XPM token has expired. Please reconnect via Integrations > XPM."
        })

    try:
        # Search XPM by entity name
        search_name = entity.entity_name or entity.trading_name or ""
        root = _xpm_get(connection, "client.api/search", params={"query": search_name})

        matches = []
        # XPM returns <Clients><Client>...</Client></Clients> or <Response><Clients>...
        clients_el = root.find(".//Clients")
        if clients_el is None:
            clients_el = root

        for client_el in clients_el.findall("Client"):
            c_uuid = _xml_text(client_el, "UUID")
            c_name = _xml_text(client_el, "Name")
            c_email = _xml_text(client_el, "Email")
            c_phone = _xml_text(client_el, "Phone")
            c_structure = _xml_text(client_el, "BusinessStructure")
            c_abn = _xml_text(client_el, "BusinessNumber")
            if c_uuid and c_name:
                matches.append({
                    "uuid": c_uuid,
                    "name": c_name,
                    "email": c_email,
                    "phone": c_phone,
                    "structure": c_structure,
                    "abn": c_abn,
                })

        if not matches:
            return JsonResponse({
                "status": "no_results",
                "message": f"No clients found in XPM matching '{search_name}'.",
                "matches": [],
            })

        return JsonResponse({
            "status": "success",
            "message": f"Found {len(matches)} match(es) in XPM.",
            "matches": matches,
            "search_term": search_name,
        })

    except Exception as e:
        import traceback
        logger.error(f"XRM search failed: {e}\n{traceback.format_exc()}")
        return JsonResponse({
            "status": "error",
            "message": "XPM search failed. Please try again later."
        })


@login_required
def xrm_pull(request, pk):
    """Pull entity data from Xero Practice Manager for a single entity."""
    entity = get_entity_for_user(request, pk)

    if request.method != "POST":
        return JsonResponse({"status": "error", "message": "POST required"}, status=405)

    if not request.user.can_do_accounting:
        return JsonResponse({"status": "error", "message": "Permission denied"}, status=403)

    # Accept xpm_client_id from POST body (set during confirmation)
    import json as json_mod
    try:
        body = json_mod.loads(request.body)
    except (json_mod.JSONDecodeError, ValueError):
        body = {}

    xpm_id_from_body = body.get("xpm_client_id", "")
    if xpm_id_from_body:
        # User confirmed a match — save the XPM Client ID to the entity
        entity.xpm_client_id = xpm_id_from_body
        entity.save(update_fields=["xpm_client_id"])

    if not entity.xpm_client_id:
        return JsonResponse({
            "status": "error",
            "message": "This entity does not have an XPM Client ID. "
                       "Use the XRM button to search and link a client first."
        })

    # Get active XPM connection
    from integrations.models import XPMConnection
    from integrations.xpm_sync import (
        _ensure_valid_token, _xpm_get, _xml_text, _xml_name,
        STRUCTURE_MAP, RELATIONSHIP_MAP,
    )
    import xml.etree.ElementTree as ET

    connection = XPMConnection.objects.filter(status="active").first()
    if not connection:
        return JsonResponse({
            "status": "error",
            "message": "No active XPM connection. Go to Integrations > XPM to connect first."
        })

    # Ensure token is valid (refresh if needed)
    if not _ensure_valid_token(connection):
        return JsonResponse({
            "status": "error",
            "message": "XPM token has expired and could not be refreshed. "
                       "Please reconnect via Integrations > XPM."
        })

    def _normalize_name(name):
        """Normalize a name for comparison: handle 'SURNAME, First' vs 'First Surname',
        strip extra whitespace, and lowercase."""
        if not name:
            return ""
        name = name.strip()
        # Handle 'SURNAME, FirstName' format
        if "," in name:
            parts = [p.strip() for p in name.split(",", 1)]
            if len(parts) == 2:
                name = f"{parts[1]} {parts[0]}"
        # Collapse multiple spaces, lowercase
        return " ".join(name.split()).lower()

    def _officer_exists(entity, name):
        """Check if an officer with a matching name already exists."""
        norm = _normalize_name(name)
        if not norm:
            return True  # Skip empty names
        for officer in EntityOfficer.objects.filter(entity=entity):
            if _normalize_name(officer.full_name) == norm:
                return True
        return False

    def _associate_exists(entity, name):
        """Check if a relationship/associate with a matching name already exists."""
        norm = _normalize_name(name)
        if not norm:
            return None
        for assoc in ClientAssociate.objects.filter(entity=entity):
            if _normalize_name(assoc.name) == norm:
                return assoc
        return None

    try:
        # Fetch detailed client data from XPM
        root = _xpm_get(connection, f"client.api/get/{entity.xpm_client_id}")
        client_el = root.find(".//Client")
        if client_el is None:
            client_el = root

        updated_fields = []

        # --- Entity Info ---
        abn = _xml_text(client_el, "BusinessNumber")
        acn = _xml_text(client_el, "CompanyNumber")
        tax_number = _xml_text(client_el, "TaxNumber")
        email = _xml_text(client_el, "Email")
        phone = _xml_text(client_el, "Phone")
        structure = _xml_text(client_el, "BusinessStructure")
        gst_registered = _xml_text(client_el, "GSTRegistered", "").lower() == "yes"

        # Address — XPM uses Address element with sub-elements
        address_el = client_el.find("Address")
        addr_line1 = addr_line2 = city = region = postcode = country = ""
        if address_el is not None:
            addr_line1 = _xml_text(address_el, "Address") or _xml_text(address_el, "Line1")
            addr_line2 = _xml_text(address_el, "Line2")
            city = _xml_text(address_el, "City")
            region = _xml_text(address_el, "Region")
            postcode = _xml_text(address_el, "PostCode")
            country = _xml_text(address_el, "Country")

        # Update entity fields (only overwrite if XPM has data)
        if abn and abn != entity.abn:
            entity.abn = abn; updated_fields.append("ABN")
        if acn and acn != entity.acn:
            entity.acn = acn; updated_fields.append("ACN")
        if tax_number and tax_number != entity.tfn:
            entity.tfn = tax_number; updated_fields.append("TFN")
        if email and email != entity.contact_email:
            entity.contact_email = email; updated_fields.append("Email")
        if phone and phone != entity.contact_phone:
            entity.contact_phone = phone; updated_fields.append("Phone")
        if gst_registered != entity.is_gst_registered:
            entity.is_gst_registered = gst_registered; updated_fields.append("GST Status")
        if structure:
            entity_type = STRUCTURE_MAP.get(structure, "")
            if entity_type and entity_type != entity.entity_type:
                entity.entity_type = entity_type; updated_fields.append("Entity Type")
        if addr_line1 and addr_line1 != entity.address_line_1:
            entity.address_line_1 = addr_line1; updated_fields.append("Address Line 1")
        if addr_line2 and addr_line2 != entity.address_line_2:
            entity.address_line_2 = addr_line2; updated_fields.append("Address Line 2")
        if city and city != entity.suburb:
            entity.suburb = city; updated_fields.append("Suburb")
        if region and region != entity.state:
            entity.state = region; updated_fields.append("State")
        if postcode and postcode != entity.postcode:
            entity.postcode = postcode; updated_fields.append("Postcode")
        if country and country != entity.country:
            entity.country = country; updated_fields.append("Country")

        entity.save()

        # --- Sync Contacts as Relationships (ClientAssociate) ---
        contacts_synced = 0
        contacts_el = client_el.find("Contacts")
        if contacts_el is not None:
            for contact_el in contacts_el.findall("Contact"):
                c_uuid = _xml_text(contact_el, "UUID")
                c_name = _xml_text(contact_el, "Name")
                c_email = _xml_text(contact_el, "Email")
                c_phone = _xml_text(contact_el, "Phone")
                c_mobile = _xml_text(contact_el, "Mobile")
                c_position = _xml_text(contact_el, "Position")
                if not c_name:
                    continue

                # Determine relationship type from position
                pos_lower = (c_position or "").lower()
                if "spouse" in pos_lower or "wife" in pos_lower or "husband" in pos_lower:
                    rel_type = "spouse"
                elif "child" in pos_lower or "son" in pos_lower or "daughter" in pos_lower:
                    rel_type = "child"
                elif "director" in pos_lower:
                    rel_type = "director"
                elif "shareholder" in pos_lower:
                    rel_type = "shareholder"
                elif "trustee" in pos_lower:
                    rel_type = "trustee"
                elif "beneficiary" in pos_lower:
                    rel_type = "beneficiary"
                elif "partner" in pos_lower:
                    rel_type = "partner_biz"
                elif "secretary" in pos_lower:
                    rel_type = "trustee"
                else:
                    rel_type = "other"

                # Find or create associate (deduplicate by UUID, then normalized name)
                assoc = None
                if c_uuid:
                    assoc = ClientAssociate.objects.filter(
                        entity=entity, xpm_contact_uuid=c_uuid
                    ).first()
                if not assoc:
                    assoc = _associate_exists(entity, c_name)
                if assoc:
                    changed = False
                    if c_uuid and not assoc.xpm_contact_uuid:
                        assoc.xpm_contact_uuid = c_uuid; changed = True
                    if c_email and not assoc.email:
                        assoc.email = c_email; changed = True
                    if (c_phone or c_mobile) and not assoc.phone:
                        assoc.phone = c_phone or c_mobile; changed = True
                    if c_position and not assoc.occupation:
                        assoc.occupation = c_position; changed = True
                    if changed:
                        assoc.save()
                else:
                    ClientAssociate.objects.create(
                        entity=entity,
                        name=c_name,
                        relationship_type=rel_type,
                        email=c_email,
                        phone=c_phone or c_mobile,
                        occupation=c_position,
                        xpm_contact_uuid=c_uuid,
                    )
                contacts_synced += 1

        # --- Sync Relationships ---
        relationships_synced = 0
        relationships_el = client_el.find("Relationships")
        if relationships_el is not None:
            for rel_el in relationships_el.findall("Relationship"):
                rel_type_xpm = _xml_text(rel_el, "Type")
                related_name_el = rel_el.find("RelatedClient")
                related_name = ""
                if related_name_el is not None:
                    name_el = related_name_el.find("Name")
                    if name_el is not None and name_el.text:
                        related_name = name_el.text.strip()
                if not related_name:
                    continue

                rel_type = RELATIONSHIP_MAP.get(rel_type_xpm, "related_entity")
                assoc = _associate_exists(entity, related_name)
                if assoc:
                    if rel_type and assoc.relationship_type == "other":
                        assoc.relationship_type = rel_type
                        assoc.save()
                else:
                    ClientAssociate.objects.create(
                        entity=entity,
                        name=related_name,
                        relationship_type=rel_type,
                    )
                relationships_synced += 1

        # --- Sync Officers from Contacts with officer-like positions ---
        officers_synced = 0
        if contacts_el is not None:
            for contact_el in contacts_el.findall("Contact"):
                c_name = _xml_text(contact_el, "Name")
                c_position = _xml_text(contact_el, "Position")
                if not c_name or not c_position:
                    continue
                pos_lower = c_position.lower()
                officer_role = None
                if "director" in pos_lower:
                    officer_role = "director"
                elif "partner" in pos_lower:
                    officer_role = "partner"
                elif "trustee" in pos_lower:
                    officer_role = "trustee"
                elif "secretary" in pos_lower:
                    officer_role = "secretary"
                elif "public officer" in pos_lower:
                    officer_role = "public_officer"
                if officer_role:
                    if not _officer_exists(entity, c_name):
                        EntityOfficer.objects.create(
                            entity=entity,
                            full_name=c_name,
                            role=officer_role,
                            title=c_position,
                        )
                        officers_synced += 1

        # Build summary
        summary_parts = []
        if updated_fields:
            summary_parts.append(f"Updated: {', '.join(updated_fields)}")
        if contacts_synced:
            summary_parts.append(f"{contacts_synced} contacts synced")
        if relationships_synced:
            summary_parts.append(f"{relationships_synced} relationships synced")
        if officers_synced:
            summary_parts.append(f"{officers_synced} officers added")
        if not summary_parts:
            summary_parts.append("Entity is already up to date")

        _log_action(
            request, "xrm_pull",
            f"XRM pull completed for {entity.entity_name}: "
            f"{len(updated_fields)} fields, {contacts_synced} contacts, "
            f"{relationships_synced} relationships, {officers_synced} officers",
            entity,
        )

        return JsonResponse({
            "status": "success",
            "message": f"XRM pull completed for {entity.entity_name}. "
                       + ". ".join(summary_parts) + "."
        })

    except Exception as e:
        import traceback
        logger.error(f"XRM pull failed for {entity.entity_name}: {e}\n{traceback.format_exc()}")
        return JsonResponse({
            "status": "error",
            "message": "XRM pull failed. Please try again later."
        })


@login_required
def trial_balance_download(request, pk):
    """
    Download comparative trial balance as Word (.docx) or PDF.
    Uses ?format=docx or ?format=pdf query parameter.
    """
    from io import BytesIO
    from collections import OrderedDict
    from decimal import Decimal

    fmt = request.GET.get("format", "pdf").lower()
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    show_comparative = bool(
        entity.include_comparative_figures
        and fy.prior_year
        and fy.prior_year.trial_balance_lines.exists()
    )
    tb_lines = list(TrialBalanceLine.objects.filter(
        financial_year=fy
    ).select_related('mapped_line_item').order_by('account_code', 'source'))

    for line in tb_lines:
        if line.source == 'rollover':
            line._cy = Decimal('0')
            line._py = (line.prior_debit or Decimal('0')) - (line.prior_credit or Decimal('0'))
        else:
            line._cy = line.closing_balance or Decimal('0')
            line._py = Decimal('0')

    _coa_lookup = _build_coa_section_lookup(entity)
    # Section ordering
    SECTION_ORDER = [
        'Revenue', 'Income', 'Cost of Sales', 'Expenses',
        'Current Assets', 'Non-Current Assets',
        'Current Liabilities', 'Non-Current Liabilities',
        'Equity', 'Income Tax',
    ]
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }

    sections = OrderedDict()
    grand_dr = Decimal('0')
    grand_cr = Decimal('0')
    grand_prior_dr = Decimal('0')
    grand_prior_cr = Decimal('0')
    for line in tb_lines:
        # HARD RULE: HandiLedger numeric code range is authoritative for section.
        hl_section = _hl_section_for_code(line.account_code)
        if hl_section:
            display_section = hl_section
        elif line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = _coa_lookup.get(line.account_code, 'Unmapped')
        sections.setdefault(display_section, []).append(line)
        cy = line._cy
        if cy > 0:
            grand_dr += cy
        elif cy < 0:
            grand_cr += abs(cy)
        py = line._py
        if py > 0:
            grand_prior_dr += py
        elif py < 0:
            grand_prior_cr += abs(py)

    ordered = OrderedDict()
    seen = set()
    for s in SECTION_ORDER:
        ds = SECTION_DISPLAY.get(s, s)
        if ds not in seen and ds in sections:
            ordered[ds] = sections[ds]
            seen.add(ds)
    for key in sections:
        if key not in ordered:
            ordered[key] = sections[key]

    # Aggregate lines by account_code to net adjustments (journal entries)
    aggregated = _aggregate_tb_lines(ordered, entity=entity)
    # Recompute grand totals from aggregated lines
    grand_dr = Decimal('0')
    grand_cr = Decimal('0')
    grand_prior_dr = Decimal('0')
    grand_prior_cr = Decimal('0')
    for _sec_name, _sec_lines in aggregated.items():
        for _line in _sec_lines:
            grand_dr += _line._agg_dr
            grand_cr += _line._agg_cr
            grand_prior_dr += _line._agg_prior_dr
            grand_prior_cr += _line._agg_prior_cr

    # Year labels
    current_year = str(fy.year_label)
    year_digits = ''.join(c for c in fy.year_label if c.isdigit())
    if year_digits:
        prior_year = f"FY{int(year_digits) - 1}" if fy.year_label.startswith('FY') else str(int(year_digits) - 1)
    elif fy.prior_year:
        prior_year = str(fy.prior_year.year_label)
    else:
        prior_year = 'Prior'

    # ABN formatting
    abn = entity.abn or ''
    abn_display = ''
    if abn:
        abn_clean = abn.replace(' ', '')
        if len(abn_clean) == 11:
            abn_display = f'ABN {abn_clean[:2]} {abn_clean[2:5]} {abn_clean[5:8]} {abn_clean[8:11]}'
        else:
            abn_display = f'ABN {abn}'

    safe_name = entity.entity_name.replace(' ', '_')

    if fmt == "docx":
        return _tb_download_word(
            fy, entity, aggregated, current_year, prior_year,
            abn_display, grand_dr, grand_cr, grand_prior_dr, grand_prior_cr, safe_name,
            show_comparative=show_comparative,
        )
    elif fmt == "xlsx":
        return _tb_download_excel(
            fy, entity, aggregated, current_year, prior_year,
            abn_display, grand_dr, grand_cr, grand_prior_dr, grand_prior_cr, safe_name,
            show_comparative=show_comparative,
        )
    else:
        # Reuse existing trial_balance_pdf
        return trial_balance_pdf(request, pk)


def _tb_download_word(fy, entity, sections, current_year, prior_year,
                      abn_display, grand_dr, grand_cr, grand_prior_dr, grand_prior_cr, safe_name,
                      show_comparative=True):
    """Generate a Word document for the trial balance (comparative or single-year)."""
    from io import BytesIO
    from decimal import Decimal
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    import os
    from django.conf import settings

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Entity name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(entity.entity_name.upper())
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'

    # ABN
    if abn_display:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(abn_display)
        run.font.size = Pt(10)
        run.font.name = 'Arial'

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tb_title = (
        f"Comparative Trial Balance as at {fy.end_date.strftime('%d %B %Y')}"
        if show_comparative
        else f"Trial Balance as at {fy.end_date.strftime('%d %B %Y')}"
    )
    run = p.add_run(tb_title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.space_after = Pt(12)

    def fmt_val(val):
        if val and val != Decimal('0'):
            return f"{val:,.2f}"
        return ''

    # Column header table — 6 columns when comparative, 4 when not
    num_cols = 6 if show_comparative else 4
    header_table = doc.add_table(rows=2, cols=num_cols)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = True

    # Header row 1
    cells = header_table.rows[0].cells
    cells[0].text = ''
    cells[1].text = ''
    cells[2].text = current_year
    cells[3].text = current_year
    if show_comparative:
        cells[4].text = prior_year
        cells[5].text = prior_year

    # Header row 2
    cells = header_table.rows[1].cells
    cells[0].text = ''
    cells[1].text = ''
    cells[2].text = '$ Dr'
    cells[3].text = '$ Cr'
    if show_comparative:
        cells[4].text = '$ Dr'
        cells[5].text = '$ Cr'

    # Style header rows
    for row_idx in range(2):
        for col_idx in range(num_cols):
            cell = header_table.rows[row_idx].cells[col_idx]
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
                if col_idx >= 2:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Section tables
    for section_name, lines in sections.items():
        # Section heading
        p = doc.add_paragraph()
        run = p.add_run(section_name)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        p.space_before = Pt(8)
        p.space_after = Pt(4)

        table = doc.add_table(rows=len(lines), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for i, line in enumerate(lines):
            # Use pre-aggregated values from _aggregate_tb_lines
            dr = line._agg_dr
            cr = line._agg_cr
            prior_dr = line._agg_prior_dr
            prior_cr = line._agg_prior_cr

            cells = table.rows[i].cells
            cells[0].text = line.account_code
            cells[1].text = line.account_name
            cells[2].text = fmt_val(dr)
            cells[3].text = fmt_val(cr)
            if show_comparative:
                cells[4].text = fmt_val(prior_dr)
                cells[5].text = fmt_val(prior_cr)

            for col_idx in range(num_cols):
                for paragraph in cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'
                        if col_idx == 0:
                            run.bold = True
                    if col_idx >= 2:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Grand totals
    p = doc.add_paragraph()
    p.space_before = Pt(8)
    totals_table = doc.add_table(rows=1, cols=num_cols)
    totals_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = totals_table.rows[0].cells
    cells[0].text = ''
    cells[1].text = 'TOTALS'
    cells[2].text = f"{grand_dr:,.2f}"
    cells[3].text = f"{grand_cr:,.2f}"
    if show_comparative:
        cells[4].text = f"{grand_prior_dr:,.2f}"
        cells[5].text = f"{grand_prior_cr:,.2f}"
    for col_idx in range(num_cols):
        for paragraph in cells[col_idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Arial'
            if col_idx >= 2:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Net profit
    net_profit_current = grand_cr - grand_dr
    net_profit_prior = grand_prior_cr - grand_prior_dr
    profit_table = doc.add_table(rows=1, cols=num_cols)
    profit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = profit_table.rows[0].cells
    cells[0].text = ''
    cells[1].text = 'Net Profit'
    cells[2].text = ''
    cells[3].text = f"{abs(net_profit_current):,.2f}"
    if show_comparative:
        cells[4].text = ''
        cells[5].text = f"{abs(net_profit_prior):,.2f}"
    for col_idx in range(num_cols):
        for paragraph in cells[col_idx].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = 'Arial'
            if col_idx >= 2:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    tb_prefix = "Comparative_TB" if show_comparative else "Trial_Balance"
    filename = f"{tb_prefix}_{safe_name}_{fy.year_label}.docx"
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _tb_download_excel(fy, entity, sections, current_year, prior_year,
                       abn_display, grand_dr, grand_cr, grand_prior_dr, grand_prior_cr, safe_name,
                       show_comparative=True):
    """Generate an Excel workbook for the trial balance (comparative or single-year)."""
    from io import BytesIO
    from decimal import Decimal
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers

    wb = Workbook()
    ws = wb.active
    ws.title = "Comparative Trial Balance"

    # Column widths
    ws.column_dimensions['A'].width = 14
    ws.column_dimensions['B'].width = 45
    ws.column_dimensions['C'].width = 16
    ws.column_dimensions['D'].width = 16
    ws.column_dimensions['E'].width = 16
    ws.column_dimensions['F'].width = 16
    ws.column_dimensions['G'].width = 14
    ws.column_dimensions['H'].width = 10

    # Styles
    title_font = Font(name='Calibri', size=14, bold=True)
    subtitle_font = Font(name='Calibri', size=11, color='555555')
    header_font = Font(name='Calibri', size=10, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='2C3E50', end_color='2C3E50', fill_type='solid')
    section_font = Font(name='Calibri', size=10, bold=True, color='1A5276')
    section_fill = PatternFill(start_color='EBF5FB', end_color='EBF5FB', fill_type='solid')
    data_font = Font(name='Calibri', size=10)
    code_font = Font(name='Calibri', size=10, bold=True)
    total_font = Font(name='Calibri', size=10, bold=True)
    total_fill = PatternFill(start_color='D5F5E3', end_color='D5F5E3', fill_type='solid')
    thin_border = Border(
        bottom=Side(style='thin', color='CCCCCC')
    )
    total_border = Border(
        top=Side(style='double', color='333333'),
        bottom=Side(style='double', color='333333')
    )
    num_fmt = '#,##0.00'

    # Title rows
    row = 1
    ws.merge_cells('A1:F1')
    ws['A1'] = entity.entity_name.upper()
    ws['A1'].font = title_font
    ws['A1'].alignment = Alignment(horizontal='center')

    row = 2
    if abn_display:
        ws.merge_cells('A2:F2')
        ws['A2'] = abn_display
        ws['A2'].font = subtitle_font
        ws['A2'].alignment = Alignment(horizontal='center')
        row = 3

    ws.merge_cells(f'A{row}:F{row}')
    tb_title = (
        f"Comparative Trial Balance as at {fy.end_date.strftime('%d %B %Y')}"
        if show_comparative
        else f"Trial Balance as at {fy.end_date.strftime('%d %B %Y')}"
    )
    ws[f'A{row}'] = tb_title
    ws[f'A{row}'].font = Font(name='Calibri', size=11, bold=True)
    ws[f'A{row}'].alignment = Alignment(horizontal='center')
    row += 2

    # Header rows — include prior year columns only when show_comparative is True
    if show_comparative:
        headers_row1 = ['', '', current_year, current_year, prior_year, prior_year, '', '']
        headers_row2 = ['Code', 'Account Name', '$ Dr', '$ Cr', '$ Dr', '$ Cr', 'Variance $', 'Var %']
    else:
        headers_row1 = ['', '', current_year, current_year]
        headers_row2 = ['Code', 'Account Name', '$ Dr', '$ Cr']

    for col_idx, val in enumerate(headers_row1, 1):
        cell = ws.cell(row=row, column=col_idx, value=val)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center' if col_idx >= 3 else 'left')
    row += 1
    for col_idx, val in enumerate(headers_row2, 1):
        cell = ws.cell(row=row, column=col_idx, value=val)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='right' if col_idx >= 3 else 'left')
    row += 1

    # Data rows by section
    for section_name, lines in sections.items():
        # Section header row
        ws.merge_cells(f'A{row}:F{row}')
        cell = ws.cell(row=row, column=1, value=section_name)
        cell.font = section_font
        cell.fill = section_fill
        for col_idx in range(1, 9):
            ws.cell(row=row, column=col_idx).fill = section_fill
        row += 1

        for line in lines:
            # Use pre-aggregated values from _aggregate_tb_lines
            dr_val = line._agg_dr
            cr_val = line._agg_cr
            dr = float(dr_val) if dr_val != 0 else None
            cr = float(cr_val) if cr_val != 0 else None

            prior_dr_val = line._agg_prior_dr
            prior_cr_val = line._agg_prior_cr
            prior_dr = float(prior_dr_val) if prior_dr_val != 0 else None
            prior_cr = float(prior_cr_val) if prior_cr_val != 0 else None

            # Write row
            code_cell = ws.cell(row=row, column=1, value=line.account_code)
            code_cell.font = code_font
            code_cell.border = thin_border

            name_cell = ws.cell(row=row, column=2, value=line.account_name)
            name_cell.font = data_font
            name_cell.border = thin_border

            if show_comparative:
                # Variance calculation
                current_net = float(dr_val - cr_val)
                prior_net = float(prior_dr_val - prior_cr_val)
                variance = current_net - prior_net
                var_pct = (variance / abs(prior_net)) * 100 if prior_net != 0 else None

                for col_idx, val in [(3, dr), (4, cr), (5, prior_dr), (6, prior_cr)]:
                    cell = ws.cell(row=row, column=col_idx, value=val)
                    cell.font = data_font
                    cell.number_format = num_fmt
                    cell.alignment = Alignment(horizontal='right')
                    cell.border = thin_border

                var_cell = ws.cell(row=row, column=7, value=round(variance, 2) if variance != 0 else None)
                var_cell.font = data_font
                var_cell.number_format = num_fmt
                var_cell.alignment = Alignment(horizontal='right')
                var_cell.border = thin_border

                pct_cell = ws.cell(row=row, column=8, value=round(var_pct, 1) if var_pct is not None else None)
                pct_cell.font = data_font
                pct_cell.number_format = '0.0%' if var_pct is not None else 'General'
                if var_pct is not None:
                    pct_cell.value = round(var_pct / 100, 3)  # Store as decimal for % format
                pct_cell.alignment = Alignment(horizontal='right')
                pct_cell.border = thin_border
            else:
                for col_idx, val in [(3, dr), (4, cr)]:
                    cell = ws.cell(row=row, column=col_idx, value=val)
                    cell.font = data_font
                    cell.number_format = num_fmt
                    cell.alignment = Alignment(horizontal='right')
                    cell.border = thin_border

            row += 1

    # Grand totals row
    row += 1
    ws.cell(row=row, column=1, value='').border = total_border
    total_label = ws.cell(row=row, column=2, value='TOTALS')
    total_label.font = total_font
    total_label.fill = total_fill
    total_label.border = total_border

    if show_comparative:
        for col_idx, val in [(3, float(grand_dr)), (4, float(grand_cr)),
                              (5, float(grand_prior_dr)), (6, float(grand_prior_cr))]:
            cell = ws.cell(row=row, column=col_idx, value=val)
            cell.font = total_font
            cell.fill = total_fill
            cell.number_format = num_fmt
            cell.alignment = Alignment(horizontal='right')
            cell.border = total_border
        for col_idx in [1, 7, 8]:
            ws.cell(row=row, column=col_idx).fill = total_fill
            ws.cell(row=row, column=col_idx).border = total_border
    else:
        for col_idx, val in [(3, float(grand_dr)), (4, float(grand_cr))]:
            cell = ws.cell(row=row, column=col_idx, value=val)
            cell.font = total_font
            cell.fill = total_fill
            cell.number_format = num_fmt
            cell.alignment = Alignment(horizontal='right')
            cell.border = total_border
        ws.cell(row=row, column=1).fill = total_fill
        ws.cell(row=row, column=1).border = total_border

    # Net profit row - calculate from P&L sections only
    row += 1
    pl_section_names = {'Income', 'Cost of Sales', 'Expenses'}
    pl_dr_total = Decimal('0')
    pl_cr_total = Decimal('0')
    pl_prior_dr_total = Decimal('0')
    pl_prior_cr_total = Decimal('0')
    for section_name, lines in sections.items():
        if section_name in pl_section_names:
            for line in lines:
                pl_dr_total += line._agg_dr
                pl_cr_total += line._agg_cr
                pl_prior_dr_total += line._agg_prior_dr
                pl_prior_cr_total += line._agg_prior_cr
    net_current = float(pl_cr_total - pl_dr_total)
    net_prior = float(pl_prior_cr_total - pl_prior_dr_total)
    profit_font = Font(name='Calibri', size=10, bold=True, color='198754')
    loss_font = Font(name='Calibri', size=10, bold=True, color='DC3545')
    ws.cell(row=row, column=2, value='Net Profit / (Loss)').font = profit_font
    if net_current >= 0:
        net_cell = ws.cell(row=row, column=4, value=net_current)
        net_cell.font = profit_font
    else:
        net_cell = ws.cell(row=row, column=4, value=net_current)
        net_cell.font = loss_font
    net_cell.number_format = num_fmt
    net_cell.alignment = Alignment(horizontal='right')
    if show_comparative:
        if net_prior >= 0:
            prior_net_cell = ws.cell(row=row, column=6, value=net_prior)
            prior_net_cell.font = profit_font
        else:
            prior_net_cell = ws.cell(row=row, column=6, value=net_prior)
            prior_net_cell.font = loss_font
        prior_net_cell.number_format = num_fmt
        prior_net_cell.alignment = Alignment(horizontal='right')

    # Freeze panes (freeze below header)
    ws.freeze_panes = 'A8' if abn_display else 'A7'

    # Print setup
    ws.sheet_properties.pageSetUpPr = None
    ws.page_setup.orientation = 'landscape'
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    tb_prefix = "Comparative_TB" if show_comparative else "Trial_Balance"
    filename = f"{tb_prefix}_{safe_name}_{fy.year_label}.xlsx"
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
@require_POST
def delete_document(request, pk):
    """Delete a generated document."""
    from .models import GeneratedDocument
    doc = get_object_or_404(GeneratedDocument, pk=pk)
    fy_pk = doc.financial_year.pk
    # Delete the file from storage
    if doc.file:
        doc.file.delete(save=False)
    doc.delete()
    messages.success(request, "Document deleted.")
    return redirect("core:financial_year_detail", pk=fy_pk)



@login_required
def trial_balance_template_download(request):
    """Serve the trial balance import template .xlsx file."""
    import os
    from django.conf import settings as _settings
    template_path = os.path.join(_settings.BASE_DIR, "static", "trial_balance_template.xlsx")
    if not os.path.exists(template_path):
        from django.http import Http404
        raise Http404("Template file not found.")
    with open(template_path, "rb") as f:
        response = HttpResponse(
            f.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = 'attachment; filename="StatementHub_Trial_Balance_Template.xlsx"'
        return response


# ============================================================
# AI Classification for Bank Statement Review (Financial Year Tab)
# ============================================================

# In-memory store for background AI classification task status
# Key: str(financial_year.pk), Value: dict with status info
_ai_classify_tasks = {}


def _run_ai_classification_background(fy_pk, entity_pk, entity_type, user_pk):
    """
    Background worker that classifies ALL unclassified transactions for an entity.
    Runs in a separate thread so the user can continue working.
    Updates _ai_classify_tasks with progress and creates an ActivityLog on completion.
    """
    import django
    from django.db import connection
    task_key = str(fy_pk)

    try:
        from review.models import PendingTransaction, ReviewJob
        from review.email_ingestion import classify_transactions
        from .models import FinancialYear, Entity, ActivityLog
        from django.contrib.auth import get_user_model
        User = get_user_model()

        entity = Entity.objects.get(pk=entity_pk)
        fy = FinancialYear.objects.get(pk=fy_pk)
        user = User.objects.get(pk=user_pk)

        # Determine GST registration
        job = ReviewJob.objects.filter(entity=entity).first()
        is_gst = job.is_gst_registered if job else True

        total_classified = 0
        batch_size = 20

        while True:
            # Fetch next batch of unclassified transactions
            unclassified = list(
                PendingTransaction.objects.filter(
                    job__entity=entity,
                    is_confirmed=False,
                ).filter(
                    Q(ai_suggested_code__isnull=True) | Q(ai_suggested_code="")
                ).order_by("date")[:batch_size]
            )

            if not unclassified:
                break

            # Build transaction dicts
            txn_dicts = []
            txn_map = {}
            for txn in unclassified:
                txn_dicts.append({
                    "date": str(txn.date),
                    "description": txn.description,
                    "amount": float(txn.amount),
                })
                txn_map[len(txn_dicts) - 1] = txn

            try:
                classifications = classify_transactions(
                    txn_dicts, entity=entity, is_gst_registered=is_gst
                )
            except Exception as exc:
                logger.error(f"AI classification batch failed: {exc}")
                _ai_classify_tasks[task_key]["status"] = "error"
                _ai_classify_tasks[task_key]["message"] = f"Classification error: {str(exc)}"
                return

            # Apply classifications
            batch_classified = 0
            for i, result in enumerate(classifications):
                txn = txn_map.get(i)
                if not txn or result is None:
                    continue
                code = result.get("account_code", "") or result.get("code", "")
                name = result.get("account_name", "") or result.get("name", "")
                confidence = result.get("confidence", 0)
                tax_type = result.get("tax_type", "")
                from_learning = result.get("from_learning", False)

                # Enforce: interest is ALWAYS GST-Free
                from review.email_ingestion import _is_interest_transaction
                desc_upper = (txn.description or "").upper()
                if _is_interest_transaction(desc_upper):
                    is_income = txn.amount >= 0
                    if is_gst:
                        tax_type = "GST Free Income" if is_income else "GST Free Expenses"
                    else:
                        tax_type = "BAS Excluded"

                if code:
                    txn.ai_suggested_code = code
                    txn.ai_suggested_name = name
                    txn.ai_confidence = 5 if from_learning else confidence
                    txn.ai_suggested_tax_type = tax_type

                    # Set gst_treatment from classification (enforced by account's tax_code)
                    gst_treatment = result.get("gst_treatment", "")
                    if gst_treatment:
                        txn.gst_treatment = gst_treatment
                    elif tax_type:
                        # Derive gst_treatment from tax_type
                        _TAX_TYPE_TO_TREATMENT = {
                            "GST on Income": "taxable",
                            "GST on Expenses": "taxable",
                            "GST Free Income": "gst_free",
                            "GST Free Expenses": "gst_free",
                            "Input Taxed": "input_taxed",
                            "BAS Excluded": "out_of_scope",
                            "N-T": "out_of_scope",
                        }
                        txn.gst_treatment = _TAX_TYPE_TO_TREATMENT.get(tax_type, "")

                    # Calculate GST amounts
                    abs_amount = abs(txn.amount)
                    if tax_type in ("GST on Income", "GST on Expenses"):
                        txn.gst_amount = (abs_amount / Decimal("11")).quantize(Decimal("0.01"))
                        txn.net_amount = (abs_amount - txn.gst_amount).quantize(Decimal("0.01"))
                    else:
                        txn.gst_amount = Decimal("0.00")
                        txn.net_amount = abs_amount

                    txn.save()
                    batch_classified += 1

            total_classified += batch_classified

            # Update progress
            total_pending = PendingTransaction.objects.filter(
                job__entity=entity, is_confirmed=False
            ).count()
            done = PendingTransaction.objects.filter(
                job__entity=entity, is_confirmed=False
            ).exclude(ai_suggested_code__isnull=True).exclude(ai_suggested_code="").count()
            _ai_classify_tasks[task_key].update({
                "total_classified": done,
                "total_pending": total_pending,
                "remaining": total_pending - done,
                "message": f"Classified {done} of {total_pending} transactions...",
            })

        # Final counts
        total_pending = PendingTransaction.objects.filter(
            job__entity=entity, is_confirmed=False
        ).count()
        done = PendingTransaction.objects.filter(
            job__entity=entity, is_confirmed=False
        ).exclude(ai_suggested_code__isnull=True).exclude(ai_suggested_code="").count()

        _ai_classify_tasks[task_key].update({
            "status": "complete",
            "total_classified": done,
            "total_pending": total_pending,
            "remaining": 0,
            "message": f"All {done} transactions classified successfully.",
        })

        # Create notification via ActivityLog
        ActivityLog.objects.create(
            user=user,
            event_type=ActivityLog.EventType.CLASSIFY_COMPLETE,
            title=f"AI Classification Complete — {entity.entity_name}",
            description=f"Classified {total_classified} transactions for {entity.entity_name} ({fy.year_label}).",
            entity=entity,
            financial_year=fy,
            url=f"/entities/years/{fy.pk}/",
        )

    except Exception as exc:
        logger.exception(f"Background AI classification failed: {exc}")
        _ai_classify_tasks[task_key] = {
            "status": "error",
            "message": f"Background classification failed: {str(exc)}",
            "total_classified": 0,
            "total_pending": 0,
            "remaining": 0,
        }
    finally:
        connection.close()


@login_required
@require_POST
def review_classify_ai(request, pk):
    """
    AJAX endpoint to kick off background AI classification of unclassified
    bank statement transactions. Returns immediately so the user can
    continue working. Progress is polled via review_classify_status.
    """
    import threading
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    task_key = str(fy.pk)

    # If already running, return current status
    if task_key in _ai_classify_tasks and _ai_classify_tasks[task_key].get("status") == "running":
        return JsonResponse({
            "status": "running",
            "message": "Classification is already in progress.",
            **{k: _ai_classify_tasks[task_key].get(k, 0) for k in ("total_classified", "total_pending", "remaining")},
        })

    from review.models import PendingTransaction

    # Check if there's anything to classify
    unclassified_count = PendingTransaction.objects.filter(
        job__entity=entity,
        is_confirmed=False,
    ).filter(
        Q(ai_suggested_code__isnull=True) | Q(ai_suggested_code="")
    ).count()

    if unclassified_count == 0:
        total_pending = PendingTransaction.objects.filter(
            job__entity=entity, is_confirmed=False
        ).count()
        return JsonResponse({
            "status": "complete",
            "total_classified": total_pending,
            "total_pending": total_pending,
            "remaining": 0,
            "message": "All transactions have already been classified.",
        })

    total_pending = PendingTransaction.objects.filter(
        job__entity=entity, is_confirmed=False
    ).count()

    # Initialise task status
    _ai_classify_tasks[task_key] = {
        "status": "running",
        "total_classified": total_pending - unclassified_count,
        "total_pending": total_pending,
        "remaining": unclassified_count,
        "message": f"Starting classification of {unclassified_count} transactions...",
    }

    # Log the start
    ActivityLog.objects.create(
        user=request.user,
        event_type=ActivityLog.EventType.CLASSIFY_STARTED,
        title=f"AI Classification Started — {entity.entity_name}",
        description=f"Classifying {unclassified_count} transactions for {entity.entity_name} ({fy.year_label}).",
        entity=entity,
        financial_year=fy,
        url=f"/entities/years/{fy.pk}/",
    )

    # Launch background thread
    thread = threading.Thread(
        target=_run_ai_classification_background,
        args=(fy.pk, entity.pk, entity.entity_type, request.user.pk),
        daemon=True,
    )
    thread.start()

    return JsonResponse({
        "status": "running",
        "total_classified": total_pending - unclassified_count,
        "total_pending": total_pending,
        "remaining": unclassified_count,
        "message": f"Classification started for {unclassified_count} transactions. You can continue working.",
    })


@login_required
def review_classify_status(request, pk):
    """
    AJAX endpoint to check the status of a background AI classification task.
    Polled by the frontend to update progress and detect completion.
    """
    fy = get_financial_year_for_user(request, pk)
    task_key = str(fy.pk)

    if task_key in _ai_classify_tasks:
        task = _ai_classify_tasks[task_key]
        return JsonResponse(task)

    # No task found — check actual DB state
    from review.models import PendingTransaction
    entity = fy.entity
    total_pending = PendingTransaction.objects.filter(
        job__entity=entity, is_confirmed=False
    ).count()
    total_classified = PendingTransaction.objects.filter(
        job__entity=entity, is_confirmed=False
    ).exclude(ai_suggested_code__isnull=True).exclude(ai_suggested_code="").count()

    return JsonResponse({
        "status": "idle",
        "total_classified": total_classified,
        "total_pending": total_pending,
        "remaining": total_pending - total_classified,
        "message": "",
    })


@login_required
@require_POST
def review_bulk_approve_group(request, pk):
    """
    AJAX endpoint to bulk-approve all pending transactions that share
    the same AI-suggested account code for a financial year's entity.
    Also auto-pushes approved transactions to the trial balance.
    """
    import json
    fy = get_financial_year_for_user(request, pk)
    if not request.user.can_do_accounting:
        return JsonResponse({"status": "error", "message": "Permission denied."}, status=403)

    from review.models import PendingTransaction

    body = json.loads(request.body) if request.content_type == "application/json" else request.POST
    account_code = body.get("account_code", "")
    # Optional override: accountant can change the AI suggestion before approving
    override_code = body.get("override_code", "").strip()
    override_name = body.get("override_name", "").strip()

    if not account_code:
        return JsonResponse({"status": "error", "message": "No account code specified."}, status=400)

    pending = PendingTransaction.objects.filter(
        job__entity=fy.entity,
        is_confirmed=False,
        ai_suggested_code=account_code,
    )

    # Determine the final code/name to use
    final_code = override_code or account_code
    final_name = override_name

    count = 0
    tb_count = 0
    for txn in pending:
        txn.confirmed_code = final_code
        txn.confirmed_name = final_name or txn.ai_suggested_name
        txn.confirmed_tax_type = txn.ai_suggested_tax_type
        txn.is_confirmed = True

        # Preserve GST amounts
        if txn.ai_suggested_tax_type in ("GST on Income", "GST on Expenses"):
            txn.confirmed_gst_amount = txn.gst_amount
        else:
            txn.confirmed_gst_amount = Decimal("0.00")

        txn.save()
        count += 1

        # Auto-push to trial balance (centralised helper handles
        # expense/income + GST + bank contra in one call)
        has_gst = txn.confirmed_gst_amount and txn.confirmed_gst_amount > 0
        if _post_txn_to_tb(txn, fy, has_gst):
            tb_count += 1

    # Auto-trigger risk engine after bulk group approve
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "review_bulk_approve_group")

    # Get remaining counts
    remaining_pending = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=False
    ).count()
    remaining_confirmed = PendingTransaction.objects.filter(
        job__entity=fy.entity, is_confirmed=True
    ).count()

    return JsonResponse({
        "status": "ok",
        "approved_count": count,
        "tb_count": tb_count,
        "remaining_pending": remaining_pending,
        "remaining_confirmed": remaining_confirmed,
        "message": f"Approved {count} transactions for {final_code}. {tb_count} pushed to TB.",
    })


# ---------------------------------------------------------------------------
# Entity Chart of Accounts (per-entity customisation)
# ---------------------------------------------------------------------------
@login_required
def entity_coa_add(request, pk):
    """Add a new account to the entity's chart of accounts (full-page form)."""
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    section_choices = EntityChartOfAccount.StatementSection.choices
    tax_code_choices = ['GST', 'ADS', 'ITS', 'FRE', 'CAP', 'INP', 'GNR', 'N-T']
    mapping_options = AccountMapping.objects.filter(
        applicable_entities__contains=entity.entity_type
    ).order_by('financial_statement', 'line_item_label')

    if request.method == 'POST':
        account_code = request.POST.get("account_code", "").strip()
        account_name = request.POST.get("account_name", "").strip()
        section = request.POST.get("section", "")
        classification = request.POST.get("classification", "").strip()
        tax_code = request.POST.get("tax_code", "").strip()
        maps_to_id = request.POST.get("maps_to", "").strip()

        if not account_code or not account_name or not section:
            messages.error(request, "Account code, name, and section are required.")
            return redirect("core:entity_coa_add", pk=pk)

        if EntityChartOfAccount.objects.filter(entity=entity, account_code=account_code).exists():
            messages.error(request, f"Account code {account_code} already exists for this entity.")
            return redirect("core:entity_coa_add", pk=pk)

        maps_to = None
        if maps_to_id:
            try:
                maps_to = AccountMapping.objects.get(pk=maps_to_id)
            except AccountMapping.DoesNotExist:
                pass

        EntityChartOfAccount.objects.create(
            entity=entity,
            account_code=account_code,
            account_name=account_name,
            section=section,
            classification=classification,
            tax_code=tax_code,
            maps_to=maps_to,
            is_active=True,
            is_custom=True,
            is_non_deductible=request.POST.get("is_non_deductible") == "on",
            is_non_assessable=request.POST.get("is_non_assessable") == "on",
            is_cgt=request.POST.get("is_cgt") == "on",
            is_franked_dividend=request.POST.get("is_franked_dividend") == "on",
            is_franking_credit=request.POST.get("is_franking_credit") == "on",
        )
        _log_action(request, "create", f"Added entity account: {account_code} — {account_name}", fy)

        # Handle sub-accounts if this is a control account
        is_control = request.POST.get("is_control_account") == "on"
        sub_count = int(request.POST.get("sub_account_count", 0) or 0)
        sub_created = 0
        if is_control and sub_count > 0:
            for i in range(1, sub_count + 1):
                sub_code = request.POST.get(f"sub_code_{i}", "").strip()
                sub_name = request.POST.get(f"sub_name_{i}", "").strip()
                if sub_code and sub_name:
                    if not EntityChartOfAccount.objects.filter(entity=entity, account_code=sub_code).exists():
                        EntityChartOfAccount.objects.create(
                            entity=entity,
                            account_code=sub_code,
                            account_name=sub_name,
                            section=section,
                            classification=classification,
                            tax_code=tax_code,
                            maps_to=maps_to,
                            is_active=True,
                            is_custom=True,
                        )
                        sub_created += 1
                        _log_action(request, "create", f"Added sub-account: {sub_code} — {sub_name} (parent: {account_code})", fy)

        if sub_created > 0:
            messages.success(request, f"Control account {account_code} — {account_name} added with {sub_created} sub-account(s).")
        else:
            messages.success(request, f"Account {account_code} — {account_name} added.")
        # Redirect to the COA tab with the new account highlighted
        return redirect(f"{reverse('core:financial_year_detail', kwargs={'pk': pk})}?tab=coa&highlight={account_code}")

    # GET — show the form
    return render(request, "core/entity_coa_form.html", {
        "fy": fy,
        "entity": entity,
        "is_edit": False,
        "section_choices": section_choices,
        "tax_code_choices": tax_code_choices,
        "mapping_options": mapping_options,
    })


@login_required
def entity_coa_edit(request, pk):
    """Edit an existing entity chart of accounts entry (full-page form)."""
    acct = get_object_or_404(EntityChartOfAccount, pk=pk)
    entity = acct.entity

    # Security: ensure user has access to this entity
    fy = entity.financial_years.order_by("-end_date").first()
    if fy:
        get_financial_year_for_user(request, fy.pk)

    section_choices = EntityChartOfAccount.StatementSection.choices
    tax_code_choices = ['GST', 'ADS', 'ITS', 'FRE', 'CAP', 'INP', 'GNR', 'N-T']
    mapping_options = AccountMapping.objects.filter(
        applicable_entities__contains=entity.entity_type
    ).order_by('financial_statement', 'line_item_label')

    if request.method == 'POST':
        old_code = acct.account_code
        acct.account_code = request.POST.get("account_code", acct.account_code).strip()
        acct.account_name = request.POST.get("account_name", acct.account_name).strip()
        acct.section = request.POST.get("section", acct.section)
        acct.classification = request.POST.get("classification", "").strip()
        acct.tax_code = request.POST.get("tax_code", "").strip()

        maps_to_id = request.POST.get("maps_to", "").strip()
        if maps_to_id:
            try:
                acct.maps_to = AccountMapping.objects.get(pk=maps_to_id)
            except AccountMapping.DoesNotExist:
                acct.maps_to = None
        else:
            acct.maps_to = None

        # Trust tax planning tags
        acct.is_non_deductible = request.POST.get("is_non_deductible") == "on"
        acct.is_non_assessable = request.POST.get("is_non_assessable") == "on"
        acct.is_cgt = request.POST.get("is_cgt") == "on"
        acct.is_franked_dividend = request.POST.get("is_franked_dividend") == "on"
        acct.is_franking_credit = request.POST.get("is_franking_credit") == "on"

        acct.save()

        # ── Sync TB lines ────────────────────────────────────────────────────
        # Propagate name/code changes to every TrialBalanceLine for this entity
        # so the TB always reflects the current CoA.
        tb_qs = TrialBalanceLine.objects.filter(
            financial_year__entity=entity,
            account_code=old_code,
        )
        tb_updated = 0
        if old_code != acct.account_code:
            # Code changed: update both code and name
            tb_updated = tb_qs.update(
                account_code=acct.account_code,
                account_name=acct.account_name,
            )
            # Also update ClientAccountMapping
            ClientAccountMapping.objects.filter(
                entity=entity, client_account_code=old_code
            ).update(
                client_account_code=acct.account_code,
                client_account_name=acct.account_name,
            )
        else:
            # Name-only change
            tb_updated = tb_qs.update(account_name=acct.account_name)
            ClientAccountMapping.objects.filter(
                entity=entity, client_account_code=acct.account_code
            ).update(client_account_name=acct.account_name)
        # ────────────────────────────────────────────────────────────────────

        if fy:
            _log_action(
                request, "update",
                f"Edited entity account: {old_code} → {acct.account_code} — {acct.account_name}"
                + (f" (synced {tb_updated} TB line(s))" if tb_updated else ""),
                fy,
            )
        messages.success(
            request,
            f"Account {acct.account_code} updated."
            + (f" {tb_updated} trial balance line(s) synced." if tb_updated else ""),
        )
        if fy:
            return redirect("core:financial_year_detail", pk=fy.pk)
        return redirect("core:entity_detail", pk=entity.pk)

    # GET — show the edit form
    return render(request, "core/entity_coa_form.html", {
        "fy": fy,
        "entity": entity,
        "account": acct,
        "is_edit": True,
        "section_choices": section_choices,
        "tax_code_choices": tax_code_choices,
        "mapping_options": mapping_options,
    })


@login_required
@require_POST
def entity_coa_delete(request, pk):
    """Delete an entity chart of accounts entry."""
    acct = get_object_or_404(EntityChartOfAccount, pk=pk)
    entity = acct.entity

    # Security: ensure user has access
    fy = entity.financial_years.order_by("-end_date").first()
    if fy:
        get_financial_year_for_user(request, fy.pk)

    code = acct.account_code
    name = acct.account_name
    acct.delete()

    if fy:
        _log_action(request, "delete", f"Deleted entity account: {code} — {name}", fy)
    messages.success(request, f"Account {code} — {name} deleted.")

    if fy:
        return redirect("core:financial_year_detail", pk=fy.pk)
    return redirect("core:entity_detail", pk=entity.pk)


# ---------------------------------------------------------------------------
# Entity COA: Auto-code suggestion and code availability check
# ---------------------------------------------------------------------------
@login_required
def entity_coa_suggest_code(request, pk):
    """
    AJAX endpoint to suggest the next available account code for an entity.

    Logic:
      1. DUPLICATE CHECK: If an account with a matching name already exists
         in the entity's COA, return its code with an 'existing_match' flag
         so the UI can offer to reuse it instead of creating a duplicate.
      2. ALPHABETICAL PLACEMENT: Sort ALL known accounts in the section by
         account code numerically, then find where the new name would sit
         alphabetically among the neighbours at each code position. Pick a
         code that keeps the list as close to alphabetical order as possible.
      3. FALLBACK: If no ideal slot exists, find the nearest available code
         to the alphabetical neighbours.
    """
    from difflib import SequenceMatcher

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    section = request.GET.get('section', '').strip()
    account_name = request.GET.get('account_name', '').strip()
    exclude_pk = request.GET.get('exclude', '').strip()

    if not section or not account_name:
        return JsonResponse({'suggested_code': '', 'position_info': 'Enter a section and account name.'})

    # Code ranges per section
    sub_ranges = {
        'revenue': (0, 999), 'cost_of_sales': (0, 999),
        'expenses': (1000, 1999),
        'current_assets': (2000, 2499),
        'non_current_assets': (2500, 2999),
        'assets': (2000, 2999),
        'current_liabilities': (3000, 3499),
        'non_current_liabilities': (3500, 3999),
        'liabilities': (3000, 3999),
        'equity': (4000, 4999), 'capital_accounts': (4000, 4999), 'pl_appropriation': (4000, 4999),
        'suspense': (9000, 9999),
    }
    lo, hi = sub_ranges.get(section, (0, 9999))

    # Shared sections that use the same code range
    shared_sections = {
        'revenue': ['revenue', 'cost_of_sales'],
        'cost_of_sales': ['revenue', 'cost_of_sales'],
        'current_assets': ['current_assets', 'assets'],
        'non_current_assets': ['non_current_assets', 'assets'],
        'assets': ['current_assets', 'non_current_assets', 'assets'],
        'current_liabilities': ['current_liabilities', 'liabilities'],
        'non_current_liabilities': ['non_current_liabilities', 'liabilities'],
        'liabilities': ['current_liabilities', 'non_current_liabilities', 'liabilities'],
        'equity': ['equity', 'capital_accounts', 'pl_appropriation'],
        'capital_accounts': ['equity', 'capital_accounts', 'pl_appropriation'],
        'pl_appropriation': ['equity', 'capital_accounts', 'pl_appropriation'],
    }
    related = shared_sections.get(section, [section])

    # ── Build combined list of all known accounts in the code range ──
    combined = {}  # code_str -> (code_str, account_name)

    # 1. Entity COA accounts in this section
    ecoa_qs = EntityChartOfAccount.objects.filter(
        entity=entity, section__in=related, is_active=True
    )
    if exclude_pk:
        ecoa_qs = ecoa_qs.exclude(pk=exclude_pk)
    for a in ecoa_qs:
        combined[a.account_code] = (a.account_code, a.account_name)

    # 2. Client account mappings in the code range
    for m in ClientAccountMapping.objects.filter(entity=entity):
        code_part = m.client_account_code.split('.')[0]
        if code_part.isdigit() and lo <= int(code_part) <= hi:
            if m.client_account_code not in combined:
                combined[m.client_account_code] = (m.client_account_code, m.client_account_name)

    # 3. Trial balance lines in the code range (latest FY for this entity)
    for tb in TrialBalanceLine.objects.filter(financial_year=fy, is_adjustment=False):
        code_part = tb.account_code.split('.')[0]
        if code_part.isdigit() and lo <= int(code_part) <= hi:
            if tb.account_code not in combined:
                combined[tb.account_code] = (tb.account_code, tb.account_name)

    # ── STEP 1: Duplicate / near-duplicate detection ──
    name_lower = account_name.lower().strip()
    best_match = None
    best_ratio = 0.0
    for code_str, acc_name in combined.values():
        # Exact match (case-insensitive)
        if acc_name.lower().strip() == name_lower:
            fmt_code = code_str
            if code_str.split('.')[0].isdigit() and int(code_str.split('.')[0]) >= 1000:
                fmt_code = code_str.split('.')[0].zfill(4)
                if '.' in code_str:
                    fmt_code += '.' + code_str.split('.', 1)[1]
            return JsonResponse({
                'suggested_code': fmt_code,
                'position_info': f'Existing account found: {code_str} ({acc_name})',
                'existing_match': True,
                'existing_code': fmt_code,
                'existing_name': acc_name,
            })
        # Fuzzy match
        ratio = SequenceMatcher(None, name_lower, acc_name.lower().strip()).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_match = (code_str, acc_name)

    # ── STEP 2: Build sorted-by-code list for alphabetical placement ──
    # Sort accounts by their integer code
    accts_by_code = []
    for code_str, acc_name in combined.values():
        code_part = code_str.split('.')[0]
        if code_part.isdigit():
            accts_by_code.append((int(code_part), code_str, acc_name))
    accts_by_code.sort(key=lambda x: x[0])

    # Build set of ALL used integer codes
    used_codes = set()
    for code_int, _, _ in accts_by_code:
        used_codes.add(code_int)

    # Find the alphabetical insertion point among existing accounts
    # We want to find the two neighbours whose names bracket the new name
    # alphabetically, then find a code near them.
    before = None  # (code_int, code_str, name) — closest alphabetically before
    after = None   # (code_int, code_str, name) — closest alphabetically after

    # Sort by name for alphabetical neighbour finding
    accts_by_name = sorted(accts_by_code, key=lambda x: x[2].lower())
    for code_int, code_str, acc_name in accts_by_name:
        if acc_name.lower() < name_lower:
            before = (code_int, code_str, acc_name)
        elif acc_name.lower() > name_lower:
            after = (code_int, code_str, acc_name)
            break

    # ── STEP 3: Find the best available code ──
    suggested_code = ''
    position_info = ''

    if before and after:
        code_before = before[0]
        code_after = after[0]
        position_info = f'Between {before[1]} ({before[2]}) and {after[1]} ({after[2]})'

        # Ensure lo_target <= hi_target regardless of code order
        target_lo = min(code_before, code_after) + 1
        target_hi = max(code_before, code_after) - 1

        # Try midpoint first
        if target_lo <= target_hi:
            mid = (target_lo + target_hi) // 2
            if mid not in used_codes and lo <= mid <= hi:
                suggested_code = str(mid)
            else:
                for c in range(target_lo, target_hi + 1):
                    if c not in used_codes and lo <= c <= hi:
                        suggested_code = str(c)
                        break

        # If no space between neighbours, search outward from the nearest
        if not suggested_code:
            anchor = (code_before + code_after) // 2
            for offset in range(1, hi - lo + 2):
                for candidate in [anchor + offset, anchor - offset]:
                    if lo <= candidate <= hi and candidate not in used_codes:
                        suggested_code = str(candidate)
                        position_info += ' (nearest available)'
                        break
                if suggested_code:
                    break

    elif before and not after:
        code_before = before[0]
        position_info = f'After {before[1]} ({before[2]})'
        # Try code_before + 1, then expand
        for c in range(code_before + 1, hi + 1):
            if c not in used_codes:
                suggested_code = str(c)
                break
        if not suggested_code:
            for c in range(code_before - 1, lo - 1, -1):
                if c not in used_codes:
                    suggested_code = str(c)
                    position_info += ' (nearest available)'
                    break

    elif after and not before:
        code_after = after[0]
        position_info = f'Before {after[1]} ({after[2]})'
        # Try code_after - 1, then expand
        for c in range(code_after - 1, lo - 1, -1):
            if c not in used_codes:
                suggested_code = str(c)
                break
        if not suggested_code:
            for c in range(code_after + 1, hi + 1):
                if c not in used_codes:
                    suggested_code = str(c)
                    position_info += ' (nearest available)'
                    break

    else:
        # No existing accounts — pick midpoint of range
        position_info = 'First account in this section'
        mid = (lo + hi) // 2
        if mid not in used_codes:
            suggested_code = str(mid)
        else:
            for c in range(lo, hi + 1):
                if c not in used_codes:
                    suggested_code = str(c)
                    break

    # Final fallback
    if not suggested_code:
        for c in range(lo, hi + 1):
            if c not in used_codes:
                suggested_code = str(c)
                position_info = 'Next available code in section'
                break

    if not suggested_code:
        return JsonResponse({
            'suggested_code': '',
            'position_info': 'No available codes in this section range.',
            'error': True,
        })

    # Zero-pad 4-digit codes
    if int(suggested_code) >= 1000:
        suggested_code = suggested_code.zfill(4)

    # Include near-match warning if similarity > 80%
    result = {
        'suggested_code': suggested_code,
        'position_info': position_info,
    }
    if best_match and best_ratio >= 0.80:
        result['similar_match'] = True
        result['similar_code'] = best_match[0]
        result['similar_name'] = best_match[1]
        result['similar_ratio'] = round(best_ratio * 100)

    return JsonResponse(result)


@login_required
def entity_coa_check_code(request, pk):
    """
    AJAX endpoint to check if an account code is available for an entity.
    """
    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity

    code = request.GET.get('code', '').strip()
    exclude_pk = request.GET.get('exclude', '').strip()

    result = {'available': False, 'error': ''}

    if not code:
        result['error'] = 'Enter an account code.'
        return JsonResponse(result)

    qs = EntityChartOfAccount.objects.filter(entity=entity, account_code=code)
    if exclude_pk:
        qs = qs.exclude(pk=exclude_pk)
    if qs.exists():
        existing = qs.first()
        result['error'] = f'Code {code} is already used by "{existing.account_name}".'
        return JsonResponse(result)

    result['available'] = True
    return JsonResponse(result)


# ============================================================
# Journal Upload (Bulk DR/CR from Excel)
# ============================================================

@login_required
def journal_upload(request, pk):
    """Upload journals from an Excel file. Parses the file and stages
    the data in the session, then redirects to the review wizard."""
    fy = get_financial_year_for_user(request, pk)

    if fy.is_locked:
        messages.error(request, "Cannot upload journals to a finalised financial year.")
        return redirect("core:financial_year_detail", pk=pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission.")
        return redirect("core:financial_year_detail", pk=pk)

    if request.method == "POST":
        uploaded_file = request.FILES.get("file")
        if not uploaded_file:
            messages.error(request, "Please select a file to upload.")
            return redirect("core:journal_upload", pk=pk)

        if uploaded_file.size > 20 * 1024 * 1024:
            messages.error(request, "File too large. Maximum size is 20MB.")
            return redirect("core:journal_upload", pk=pk)

        import os as _os
        file_ext = _os.path.splitext(uploaded_file.name)[1].lower()
        if file_ext not in {".xlsx", ".xls"}:
            messages.error(request, f"Unsupported file type: {file_ext}. Only Excel files (.xlsx) are supported.")
            return redirect("core:journal_upload", pk=pk)

        try:
            raw_lines = _parse_journal_excel(fy, uploaded_file)
            if not raw_lines:
                messages.warning(request, "No journal lines found in the uploaded file. Check the format and try again.")
                return redirect("core:journal_upload", pk=pk)

            # Apply learned mappings (same pattern as TB import)
            staged_lines = _apply_journal_learned_mappings(fy, raw_lines)

            # Store in session for the review wizard
            request.session["staged_journal_upload"] = {
                "fy_pk": str(pk),
                "filename": uploaded_file.name,
                "lines": staged_lines,
            }
            # Force session save to DB before redirect so the next
            # request (possibly handled by a different Gunicorn worker)
            # can read the staged data from the database backend.
            request.session.modified = True
            request.session.save()

            return redirect("core:review_journal_upload", pk=pk)
        except Exception as e:
            messages.error(request, f"Journal upload failed: {e}")
            return redirect("core:journal_upload", pk=pk)

    return render(request, "core/journal_upload.html", {"fy": fy})


def _parse_journal_excel(fy, file):
    """Parse a journal upload Excel file into a list of raw line dicts.

    Expected columns: JOURNAL DATE | DESCRIPTION | ACCOUNT CODE | DEBIT | CREDIT
    Returns a list of dicts with keys: account_code, account_name, description,
    journal_date, debit, credit.
    """
    from datetime import datetime, date

    wb = openpyxl.load_workbook(file, data_only=True)
    ws = wb.active

    rows = list(ws.iter_rows(min_row=1, values_only=True))
    if not rows:
        return []

    # Find the header row (look for "account code" or "debit" in the first 5 rows)
    header_row_idx = 0
    for i, row in enumerate(rows[:5]):
        row_lower = [str(c).lower().strip() if c else "" for c in row]
        if any(kw in col for col in row_lower for kw in ["account code", "account", "debit", "journal date"]):
            header_row_idx = i
            break

    data_rows = rows[header_row_idx + 1:]

    def _parse_amount(val):
        if val is None:
            return Decimal("0")
        if isinstance(val, (int, float)):
            return Decimal(str(round(val, 2)))
        val_str = str(val).strip().replace(",", "").replace("$", "")
        if not val_str or val_str == "-":
            return Decimal("0")
        try:
            return Decimal(val_str).quantize(Decimal("0.01"))
        except (InvalidOperation, ValueError):
            return Decimal("0")

    raw_lines = []

    for row in data_rows:
        if not row or len(row) < 5:
            continue

        # Parse date (column 0) — optional, skip row if unparseable
        raw_date = row[0]
        if raw_date is None:
            continue
        journal_date_str = ""
        if isinstance(raw_date, datetime):
            journal_date_str = raw_date.strftime("%d/%m/%Y")
        elif isinstance(raw_date, date):
            journal_date_str = raw_date.strftime("%d/%m/%Y")
        elif isinstance(raw_date, str):
            raw_date = raw_date.strip()
            if not raw_date:
                continue
            parsed = False
            for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%d/%m/%y", "%d-%m-%y"):
                try:
                    datetime.strptime(raw_date, fmt)
                    journal_date_str = raw_date
                    parsed = True
                    break
                except ValueError:
                    continue
            if not parsed:
                continue
        else:
            continue

        # Parse description (column 1)
        description = str(row[1]).strip() if row[1] else "Journal Entry"

        # Parse account code (column 2)
        raw_code = row[2]
        if raw_code is None:
            continue
        account_code = str(raw_code).strip()
        if account_code.endswith(".0"):
            account_code = account_code[:-2]
        if not account_code:
            continue

        # Parse debit (column 3) and credit (column 4)
        debit = _parse_amount(row[3] if len(row) > 3 else None)
        credit = _parse_amount(row[4] if len(row) > 4 else None)

        # Skip rows where both debit and credit are zero
        if debit == 0 and credit == 0:
            continue

        # Resolve account name — use the description column as a fallback
        # since it often contains the account name from the source system
        account_name = _resolve_account_name(fy.entity, account_code, description)

        raw_lines.append({
            "account_code": account_code,
            "account_name": account_name,
            "description": description,
            "journal_date": journal_date_str,
            "debit": str(debit),
            "credit": str(credit),
        })

    return raw_lines


def _apply_journal_learned_mappings(fy, raw_lines):
    """Apply learned mappings and entity COA matching to parsed journal lines.
    Returns a list of staged line dicts ready for the review wizard.

    Account name resolution priority:
      1. Existing TrialBalanceLine for this financial year (source of truth —
         matches what is already displayed on the TB).
      2. Entity COA (EntityChartOfAccount) — used for accounts not yet in the TB.
      3. Raw Excel name — last resort for genuinely new accounts.
    """
    entity = fy.entity

    existing_mappings = {
        cam.client_account_code: cam
        for cam in ClientAccountMapping.objects.filter(entity=entity)
        .select_related("mapped_line_item")
    }
    entity_coa = {
        ea.account_code.lower(): ea
        for ea in EntityChartOfAccount.objects.filter(entity=entity)
        .select_related("maps_to")
    }
    # Build a lookup of canonical account names from the existing (non-adjustment)
    # TB lines for this financial year.  These are the names already displayed on
    # the TB and are the source of truth.
    tb_names = {
        tbl["account_code"].lower(): tbl["account_name"]
        for tbl in TrialBalanceLine.objects.filter(
            financial_year=fy, is_adjustment=False
        ).values("account_code", "account_name").distinct()
    }

    staged = []
    for line in raw_lines:
        code = line["account_code"]
        cam = existing_mappings.get(code)

        # Resolve the canonical account name:
        #   1. Use the existing TB name if the account is already in the TB.
        #   2. Fall back to the entity COA name for accounts not yet in the TB.
        #   3. Fall back to the raw Excel name for genuinely new accounts.
        tb_name = tb_names.get(code.lower())
        if tb_name:
            resolved_name = tb_name
        else:
            ea_lookup = entity_coa.get(code.lower())
            resolved_name = (ea_lookup.account_name if ea_lookup else None) or line["account_name"]

        staged_line = {
            "account_code": code,
            "account_name": resolved_name,
            "description": line.get("description", ""),
            "journal_date": line.get("journal_date", ""),
            "debit": line["debit"],
            "credit": line["credit"],
            "mapped_id": "",
            "mapped_label": "",
            "confidence": "new",
            "entity_acct_code": "",
            "entity_acct_name": "",
        }
        # Check learned mappings from ClientAccountMapping
        if cam and cam.mapped_line_item:
            staged_line["mapped_id"] = str(cam.mapped_line_item.pk)
            staged_line["mapped_label"] = (
                f"{cam.mapped_line_item.standard_code} - "
                f"{cam.mapped_line_item.line_item_label}"
            )
            staged_line["confidence"] = "learned"
        # Try to match entity COA by code
        ea = entity_coa.get(code.lower())
        if ea:
            staged_line["entity_acct_code"] = ea.account_code
            # Use the TB name for the entity account display if available,
            # since the TB is the source of truth and the COA may be stale.
            staged_line["entity_acct_name"] = tb_names.get(code.lower()) or ea.account_name
            if staged_line["confidence"] == "new":
                staged_line["confidence"] = "matched"
            if ea.maps_to and not staged_line["mapped_id"]:
                staged_line["mapped_id"] = str(ea.maps_to.pk)
                staged_line["mapped_label"] = (
                    f"{ea.maps_to.standard_code} - {ea.maps_to.line_item_label}"
                )
        staged.append(staged_line)
    return staged


@login_required
def review_journal_upload(request, pk):
    """Review wizard for journal uploads — same UX as TB import wizard."""
    import json as _json
    fy = get_financial_year_for_user(request, pk)
    staged = request.session.get("staged_journal_upload")

    if not staged or staged.get("fy_pk") != str(pk):
        messages.error(request, "No staged journal upload data found. Please upload again.")
        return redirect("core:journal_upload", pk=pk)

    lines = staged["lines"]
    entity = fy.entity

    # Standard accounts for statement line mapping dropdown
    standard_accounts = list(
        AccountMapping.objects.values("id", "standard_code", "line_item_label", "statement_section")
        .order_by("financial_statement", "display_order")
    )
    for sa in standard_accounts:
        sa["id"] = str(sa["id"])

    # Entity accounts for the searchable COA dropdown
    entity_accts = []
    for ea in EntityChartOfAccount.objects.filter(entity=entity).select_related("maps_to").order_by("account_code"):
        entity_accts.append({
            "code": ea.account_code,
            "name": ea.account_name,
            "section": ea.get_section_display(),
            "section_key": ea.section,
            "maps_to_id": str(ea.maps_to.pk) if ea.maps_to else "",
        })

    total = len(lines)
    auto_mapped = sum(1 for l in lines if l.get("mapped_id") or l.get("entity_acct_code"))
    unmapped = total - auto_mapped

    context = {
        "fy": fy,
        "lines": lines,
        "standard_accounts_json": _json.dumps(standard_accounts),
        "entity_accounts_json": _json.dumps(entity_accts),
        "total": total,
        "auto_mapped": auto_mapped,
        "unmapped": unmapped,
        "source_name": staged.get("filename", "Journal Upload"),
    }
    return render(request, "core/review_journal_upload.html", context)


@login_required
@require_POST
def commit_journal_upload(request, pk):
    """Commit the reviewed journal upload. Creates a BulkJournalUpload record,
    creates adjustment TrialBalanceLine records linked to it, updates
    ClientAccountMapping (learning system), and triggers risk engine."""
    fy = get_financial_year_for_user(request, pk)
    staged = request.session.get("staged_journal_upload")

    if not staged or staged.get("fy_pk") != str(pk):
        messages.error(request, "No staged journal upload data found.")
        return redirect("core:financial_year_detail", pk=pk)

    if fy.is_locked:
        messages.error(request, "Cannot upload journals to a finalised financial year.")
        return redirect("core:financial_year_detail", pk=pk)

    entity = fy.entity
    staged_lines = staged["lines"]
    lines_created = 0
    total_debit = Decimal('0')
    total_credit = Decimal('0')
    errors = []

    # Create the BulkJournalUpload tracking record
    bulk_upload = BulkJournalUpload(
        financial_year=fy,
        filename=staged.get('filename', 'Journal Upload'),
        description=f"Bulk Journal Upload — {staged.get('filename', 'Excel')}",
        uploaded_by=request.user,
    )
    bulk_upload.save()  # Triggers auto-reference generation (Bulk JNLS-001, etc.)

    for i, line in enumerate(staged_lines):
        mapping_id = request.POST.get(f"mapping_{i}", "").strip()
        entity_acct_code = request.POST.get(f"entity_acct_{i}", "").strip()
        mapped_item = None

        if mapping_id:
            try:
                mapped_item = AccountMapping.objects.get(pk=mapping_id)
            except AccountMapping.DoesNotExist:
                pass

        # If an entity account was assigned, look it up for its maps_to
        if entity_acct_code and not mapped_item:
            try:
                ea = EntityChartOfAccount.objects.select_related("maps_to").get(
                    entity=entity, account_code=entity_acct_code
                )
                if ea.maps_to:
                    mapped_item = ea.maps_to
            except EntityChartOfAccount.DoesNotExist:
                pass

        try:
            debit = Decimal(str(line.get("debit", "0")))
            credit = Decimal(str(line.get("credit", "0")))
            account_code = line["account_code"]
            account_name = line["account_name"]
            account_name = _resolve_account_name(entity, account_code, account_name)

            # If the user assigned an entity account, resolve the name from it
            if entity_acct_code:
                try:
                    ea = EntityChartOfAccount.objects.get(
                        entity=entity, account_code=entity_acct_code
                    )
                    account_name = ea.account_name or account_name
                except EntityChartOfAccount.DoesNotExist:
                    pass

            # Apply to Trial Balance as adjustment line, linked to bulk upload
            line_description = line.get('description', '')
            _apply_journal_line_to_tb(
                fy, account_code, account_name,
                debit, credit, source='journal_upload',
                bulk_upload=bulk_upload,
                description=line_description,
            )

            # Update the learning system
            if mapped_item:
                ClientAccountMapping.objects.update_or_create(
                    entity=entity,
                    client_account_code=account_code,
                    defaults={
                        "client_account_name": account_name,
                        "mapped_line_item": mapped_item,
                    },
                )

            total_debit += debit
            total_credit += credit
            lines_created += 1
        except Exception as e:
            errors.append(f"Line {i + 1} ({line.get('account_code', '?')}): {str(e)}")

    # Update the bulk upload record with totals
    bulk_upload.lines_count = lines_created
    bulk_upload.total_debit = total_debit
    bulk_upload.total_credit = total_credit
    bulk_upload.save(update_fields=['lines_count', 'total_debit', 'total_credit'])

    # Clean up session
    request.session.pop("staged_journal_upload", None)

    # Log and trigger risk engine
    _log_action(request, "journal_upload", f"Uploaded {lines_created} journal lines via wizard from {staged.get('filename', 'Excel')} as {bulk_upload.reference_number}", fy)
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "journal_uploaded")

    messages.success(
        request,
        f"Successfully posted {lines_created} journal line{'s' if lines_created != 1 else ''} "
        f"to the Trial Balance as {bulk_upload.reference_number}."
    )
    if errors:
        for err in errors[:5]:
            messages.warning(request, err)

    return redirect("core:financial_year_detail", pk=pk)


@login_required
def journal_template_download(request):
    """Serve the journal upload template .xlsx file."""
    import os
    from django.conf import settings as _settings
    template_path = os.path.join(_settings.BASE_DIR, "static", "journal_upload_template.xlsx")
    if not os.path.exists(template_path):
        from django.http import Http404
        raise Http404("Journal template file not found.")
    with open(template_path, "rb") as f:
        response = HttpResponse(
            f.read(),
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        response["Content-Disposition"] = 'attachment; filename="StatementHub_Journal_Upload_Template.xlsx"'
        return response


@login_required
def net_profit_api(request, pk):
    """Return the current net profit for a financial year (JSON).

    Used by the live net profit banner in the Trial Balance tab to
    refresh after bank statement approvals without a full page reload.
    """
    fy = get_financial_year_for_user(request, pk)
    tb_lines = fy.trial_balance_lines.select_related("mapped_line_item").all()
    _coa_lookup = _build_coa_section_lookup(fy.entity)
    SECTION_DISPLAY = {
        'Revenue': 'Income', 'Income': 'Income',
        'Cost of Sales': 'Cost of Sales', 'Expenses': 'Expenses',
        'Current Assets': 'Current Assets', 'Non-Current Assets': 'Non Current Assets',
        'Current Liabilities': 'Current Liabilities', 'Non-Current Liabilities': 'Non Current Liabilities',
        'Equity': 'Equity', 'Income Tax': 'Equity',
    }
    pl_sections = {'Income', 'Cost of Sales', 'Expenses'}
    pl_dr = Decimal('0')
    pl_cr = Decimal('0')
    for line in tb_lines:
        # HARD RULE: HandiLedger numeric code range is authoritative for section.
        hl_section = _hl_section_for_code(line.account_code)
        if hl_section:
            display_section = hl_section
        elif line.mapped_line_item:
            raw_section = line.mapped_line_item.statement_section
            display_section = SECTION_DISPLAY.get(raw_section, raw_section)
        else:
            display_section = _coa_lookup.get(line.account_code, 'Unmapped')
        if display_section in pl_sections:
            # Compute display_dr / display_cr — always use closing_balance
            cb = line.closing_balance if line.closing_balance else Decimal('0')
            if cb > 0:
                dr = cb
                cr = Decimal('0')
            elif cb < 0:
                dr = Decimal('0')
                cr = abs(cb)
            else:
                dr = line.debit if line.debit else Decimal('0')
                cr = line.credit if line.credit else Decimal('0')
            pl_dr += dr or Decimal('0')
            pl_cr += cr or Decimal('0')

    net_profit = pl_cr - pl_dr

    # Also compute totals for the balance check
    total_dr = Decimal('0')
    total_cr = Decimal('0')
    for line in tb_lines:
        # Always use closing_balance for display
        cb = line.closing_balance if line.closing_balance else Decimal('0')
        if cb > 0:
            total_dr += cb
        elif cb < 0:
            total_cr += abs(cb)
        else:
            total_dr += line.debit or Decimal('0')
            total_cr += line.credit or Decimal('0')

    return JsonResponse({
        "net_profit": str(net_profit),
        "total_debit": str(total_dr),
        "total_credit": str(total_cr),
        "balanced": total_dr == total_cr,
    })


# ---------------------------------------------------------------------------
# Bulk Journal Upload Views
# ---------------------------------------------------------------------------

@login_required
def bulk_journal_detail(request, pk):
    """View a bulk journal upload with all its linked trial balance lines."""
    bulk = get_object_or_404(
        BulkJournalUpload.objects.select_related(
            "financial_year", "financial_year__entity", "uploaded_by"
        ),
        pk=pk,
    )
    fy = bulk.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check
    entity = fy.entity

    lines = TrialBalanceLine.objects.filter(
        bulk_journal_upload=bulk
    ).select_related('mapped_line_item').order_by('account_code')

    return render(request, "core/bulk_journal_detail.html", {
        "bulk": bulk, "fy": fy, "entity": entity, "lines": lines,
    })


@login_required
@require_POST
def bulk_journal_delete(request, pk):
    """Delete an entire bulk journal upload and all its linked TB adjustment lines."""
    bulk = get_object_or_404(
        BulkJournalUpload.objects.select_related("financial_year"),
        pk=pk,
    )
    fy = bulk.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check

    if fy.is_locked:
        messages.error(request, "Cannot delete journals in a finalised year.")
        return redirect("core:bulk_journal_detail", pk=pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission to delete journals.")
        return redirect("core:bulk_journal_detail", pk=pk)

    ref = bulk.reference_number
    lines_count = bulk.trial_balance_lines.count()

    # Delete all linked TB adjustment lines
    bulk.trial_balance_lines.all().delete()

    # Delete the bulk upload record
    bulk.delete()

    _log_action(request, "journal_upload", f"Deleted bulk journal upload {ref} ({lines_count} lines removed)")

    # Trigger risk engine recalculation
    from core.signals import trigger_risk_recalc
    trigger_risk_recalc(fy, "bulk_journal_deleted")

    messages.success(request, f"Bulk journal upload {ref} and all {lines_count} lines have been deleted.")
    return redirect(reverse("core:financial_year_detail", args=[fy.pk]) + "?tab=journals")


@login_required
@require_POST
def bulk_journal_line_delete(request, pk):
    """Delete a single line from a bulk journal upload."""
    line = get_object_or_404(
        TrialBalanceLine.objects.select_related(
            "financial_year", "bulk_journal_upload"
        ),
        pk=pk,
    )
    fy = line.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check

    if fy.is_locked:
        messages.error(request, "Cannot delete journal lines in a finalised year.")
        if line.bulk_journal_upload:
            return redirect("core:bulk_journal_detail", pk=line.bulk_journal_upload.pk)
        return redirect("core:financial_year_detail", pk=fy.pk)

    if not request.user.can_do_accounting:
        messages.error(request, "You do not have permission to delete journal lines.")
        if line.bulk_journal_upload:
            return redirect("core:bulk_journal_detail", pk=line.bulk_journal_upload.pk)
        return redirect("core:financial_year_detail", pk=fy.pk)

    bulk = line.bulk_journal_upload
    account_code = line.account_code
    debit = line.debit
    credit = line.credit

    line.delete()

    # Update the bulk upload totals
    if bulk:
        remaining = TrialBalanceLine.objects.filter(bulk_journal_upload=bulk)
        bulk.lines_count = remaining.count()
        bulk.total_debit = sum(l.debit for l in remaining) or Decimal('0')
        bulk.total_credit = sum(l.credit for l in remaining) or Decimal('0')
        bulk.save(update_fields=['lines_count', 'total_debit', 'total_credit'])

        _log_action(
            request, "journal_upload",
            f"Deleted line {account_code} (Dr ${debit}, Cr ${credit}) from {bulk.reference_number}"
        )

        # If no lines remain, delete the bulk upload record too
        if bulk.lines_count == 0:
            ref = bulk.reference_number
            bulk.delete()
            messages.success(request, f"Last line removed. {ref} has been automatically deleted.")
            return redirect(reverse("core:financial_year_detail", args=[fy.pk]) + "?tab=journals")

        # Trigger risk engine recalculation
        from core.signals import trigger_risk_recalc
        trigger_risk_recalc(fy, "bulk_journal_line_deleted")

        messages.success(request, f"Line {account_code} deleted from {bulk.reference_number}.")
        return redirect("core:bulk_journal_detail", pk=bulk.pk)

    messages.success(request, f"Journal line {account_code} deleted.")
    return redirect(reverse("core:financial_year_detail", args=[fy.pk]) + "?tab=journals")


@login_required
@require_POST
def bulk_journal_reallocate(request, pk):
    """Re-allocate (re-map) lines within a bulk journal upload to different account mappings."""
    bulk = get_object_or_404(
        BulkJournalUpload.objects.select_related("financial_year"),
        pk=pk,
    )
    fy = bulk.financial_year
    get_financial_year_for_user(request, fy.pk)  # IDOR check

    if fy.is_locked:
        messages.error(request, "Cannot modify journals in a finalised year.")
        return redirect("core:bulk_journal_detail", pk=pk)

    lines = TrialBalanceLine.objects.filter(bulk_journal_upload=bulk)
    updated = 0

    for line in lines:
        mapping_id = request.POST.get(f"mapping_{line.pk}", "").strip()
        if mapping_id:
            try:
                new_mapping = AccountMapping.objects.get(pk=mapping_id)
                if line.mapped_line_item != new_mapping:
                    line.mapped_line_item = new_mapping
                    line.save(update_fields=["mapped_line_item"])
                    updated += 1
            except AccountMapping.DoesNotExist:
                pass

    if updated:
        _log_action(
            request, "journal_upload",
            f"Re-allocated {updated} lines in {bulk.reference_number}"
        )
        from core.signals import trigger_risk_recalc
        trigger_risk_recalc(fy, "bulk_journal_reallocated")
        messages.success(request, f"Re-allocated {updated} line{'s' if updated != 1 else ''} in {bulk.reference_number}.")
    else:
        messages.info(request, "No changes were made.")

    return redirect("core:bulk_journal_detail", pk=pk)


# ---------------------------------------------------------------------------
# PDF Export — Bank Statement Review (Pending / Confirmed) from FY detail
# ---------------------------------------------------------------------------

@login_required
def review_export_pdf(request, pk):
    """
    Generate a PDF of bank statement review transactions for a financial year.
    Query params:
        filter=pending|confirmed|all  (default: all)
    """
    import weasyprint
    from django.utils.html import escape as html_escape
    from review.models import PendingTransaction

    fy = get_financial_year_for_user(request, pk)
    entity = fy.entity
    filter_type = request.GET.get("filter", "all")

    # Get all transactions for this entity
    base_qs = PendingTransaction.objects.filter(
        job__entity=entity,
    ).order_by("date", "description")

    if filter_type == "pending":
        txns = base_qs.filter(is_confirmed=False)
        section_title = "Pending Review"
    elif filter_type == "confirmed":
        txns = base_qs.filter(is_confirmed=True)
        section_title = "Confirmed Transactions"
    else:
        txns = base_qs
        section_title = "All Transactions"

    total_count = base_qs.count()
    confirmed_count = base_qs.filter(is_confirmed=True).count()
    pending_count = total_count - confirmed_count

    # Build transaction rows
    rows_html = ""
    total_gross = Decimal("0.00")
    total_gst = Decimal("0.00")
    total_net = Decimal("0.00")

    for txn in txns:
        total_gross += txn.amount
        total_gst += txn.gst_amount or Decimal("0.00")
        total_net += txn.net_amount or Decimal("0.00")

        amt_class = "debit" if txn.amount < 0 else "credit"

        if txn.is_confirmed:
            acct_display = html_escape(
                f"{txn.confirmed_code} — {txn.confirmed_name}"
                if txn.confirmed_code else "—"
            )
            tax_display = html_escape(txn.confirmed_tax_type or "—")
        else:
            acct_display = html_escape(
                f"{txn.ai_suggested_code} — {txn.ai_suggested_name}"
                if txn.ai_suggested_code else "Unclassified"
            )
            tax_display = html_escape(txn.ai_suggested_tax_type or "—")

        status = "Confirmed" if txn.is_confirmed else "Pending"

        gst_amt = f"${txn.gst_amount:,.2f}" if txn.gst_amount else "—"
        net_amt = f"${txn.net_amount:,.2f}" if txn.net_amount else "—"
        cred_pct = f"{txn.creditable_percentage:.0f}%" if txn.creditable_percentage else "—"

        rows_html += f"""<tr>
            <td>{html_escape(str(txn.date))}</td>
            <td>{html_escape(txn.description)}</td>
            <td class="r {amt_class}">${txn.amount:,.2f}</td>
            <td class="r">{gst_amt}</td>
            <td class="r">{net_amt}</td>
            <td class="r">{cred_pct}</td>
            <td>{acct_display}</td>
            <td>{tax_display}</td>
            <td class="status-{status.lower()}">{status}</td>
        </tr>"""

    now_str = timezone.now().strftime("%d/%m/%Y %H:%M")
    entity_name = html_escape(entity.entity_name)
    fy_label = html_escape(f"{fy.start_date.strftime('%d/%m/%Y')} — {fy.end_date.strftime('%d/%m/%Y')}")

    html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
    @page {{
        size: A4 landscape;
        margin: 12mm 10mm;
        @bottom-right {{ content: "Page " counter(page) " of " counter(pages); font-size: 7pt; color: #999; }}
        @bottom-left {{ content: "Generated {now_str}"; font-size: 7pt; color: #999; }}
    }}
    body {{ font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; font-size: 8pt; color: #333; }}
    .header {{ display: flex; justify-content: space-between; align-items: flex-start; margin-bottom: 4mm; border-bottom: 2px solid #1565c0; padding-bottom: 3mm; }}
    .header-left {{ }}
    .header-right {{ text-align: right; font-size: 7pt; color: #666; }}
    h1 {{ font-size: 14pt; margin: 0; color: #1a1a2e; }}
    .subtitle {{ font-size: 9pt; color: #666; margin-top: 2px; }}
    .count-summary {{ font-size: 8pt; color: #555; margin-bottom: 3mm; }}
    table {{ width: 100%; border-collapse: collapse; margin-bottom: 3mm; }}
    th {{ background: #f5f5f5; font-weight: 700; font-size: 7pt; text-transform: uppercase; letter-spacing: 0.3px; color: #555; padding: 3px 4px; border-bottom: 2px solid #ddd; text-align: left; white-space: nowrap; }}
    td {{ padding: 2.5px 4px; font-size: 7.5pt; border-bottom: 1px solid #eee; vertical-align: top; }}
    .r {{ text-align: right; }}
    .bold {{ font-weight: 700; }}
    .debit {{ color: #c62828; }}
    .credit {{ color: #2e7d32; }}
    .total-row {{ border-top: 2px solid #333; background: #f8f9fa; }}
    .total-row td {{ font-weight: 700; padding: 4px; }}
    .status-confirmed {{ color: #2e7d32; font-weight: 600; }}
    .status-pending {{ color: #e65100; font-weight: 600; }}
</style></head><body>

<div class="header">
    <div class="header-left">
        <h1>{entity_name}</h1>
        <div class="subtitle">{section_title}</div>
    </div>
    <div class="header-right">
        Financial Year: {fy_label}
    </div>
</div>

<div class="count-summary">
    Showing <strong>{txns.count()}</strong> transactions
    &nbsp;|&nbsp; Total: {total_count}
    &nbsp;|&nbsp; Confirmed: {confirmed_count}
    &nbsp;|&nbsp; Pending: {pending_count}
</div>

<table>
    <thead>
        <tr>
            <th>Date</th>
            <th>Description</th>
            <th class="r">Amount</th>
            <th class="r">GST</th>
            <th class="r">Net</th>
            <th class="r">Cred.%</th>
            <th>Account</th>
            <th>Tax Type</th>
            <th>Status</th>
        </tr>
    </thead>
    <tbody>
        {rows_html}
        <tr class="total-row">
            <td colspan="2" class="bold">TOTALS</td>
            <td class="r bold {'debit' if total_gross < 0 else 'credit'}">${total_gross:,.2f}</td>
            <td class="r bold">${total_gst:,.2f}</td>
            <td class="r bold">${total_net:,.2f}</td>
            <td colspan="4"></td>
        </tr>
    </tbody>
</table>

</body></html>"""

    pdf_bytes = weasyprint.HTML(string=html).write_pdf()

    entity_slug = entity.entity_name.replace(" ", "_").replace("/", "-")[:40]
    filename = f"{entity_slug}_{section_title.replace(' ', '_')}_{timezone.now().strftime('%Y%m%d')}.pdf"

    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response
